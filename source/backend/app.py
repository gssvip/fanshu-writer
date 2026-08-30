import os, sys
if __name__ == '__main__': sys.modules.setdefault('app', sys.modules[__name__])  # 防双加载
import re
import json
import uuid
import shutil
import zipfile
import tempfile
import hashlib
import functools
import requests
from datetime import datetime, timezone, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional  # 必须在顶层导入，CI高版本解释器检查类型注解
from flask import Flask, request, jsonify, send_file, send_from_directory, stream_with_context
from flask_cors import CORS; from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

# P0-1：确定性后写校验器（零 LLM 成本，章节生成后检测 AI 痕迹）
try:
    from post_write_validator import validate_chapter, validate_chapter_with_bible, get_repair_hints, validate_chapter_with_drift, compute_style_fingerprint
except ImportError:
    # 静默降级：模块缺失时不阻断主流程
    validate_chapter = None
    validate_chapter_with_bible = None
    get_repair_hints = None
    validate_chapter_with_drift = None
    compute_style_fingerprint = None

# P0-2：伏笔 DAG 管理器
try:
    from foreshadowing_manager import parse_text_to_dag, ForeshadowingGraph, get_hooks_for_chapter
except ImportError:
    parse_text_to_dag = None
    ForeshadowingGraph = None
    get_hooks_for_chapter = None

# P1-4：四级大纲层级构建器
try:
    from outline_hierarchy_builder import build_outline_hierarchy, build_dramatic_position_prompt, get_dramatic_context
except ImportError:
    build_outline_hierarchy = None
    build_dramatic_position_prompt = None
    get_dramatic_context = None

# P1-5：节拍模板加载器
try:
    from beat_template_loader import build_beat_prompt
except ImportError:
    build_beat_prompt = None

# P1-6 + P1-7：章级变更处理器（12类CHANGES解析 + delta回写）
try:
    from chapter_changes_processor import extract_changes, apply_chapter_changes, remove_chapter_changes, build_changes_prompt_template
except ImportError:
    extract_changes = None
    apply_chapter_changes = None
    remove_chapter_changes = None
    build_changes_prompt_template = None

# P2-9：Spot-Fix 修订路由器
try:
    from revise import route_revision, build_spot_fix_prompt, apply_spot_fix_patches, estimate_token_saving
except ImportError:
    route_revision = None
    build_spot_fix_prompt = None
    apply_spot_fix_patches = None
    estimate_token_saving = None

# P2-8：审计-修订闭环
try:
    from chapter_review_cycle import run_review_cycle, run_review_cycle_with_bible, PASS_SCORE
except ImportError:
    run_review_cycle = None
    run_review_cycle_with_bible = None
    PASS_SCORE = 85

# P2-10 + P2-11：落地门禁 + PRE_WRITE_CHECK
try:
    from generation_gate import run_all_gates, build_pre_write_check_prompt
except ImportError:
    run_all_gates = None
    build_pre_write_check_prompt = None

# P3：LLM Gateway 统一入口（错误分类 + 智能重试 + 空内容检测）
try:
    from llm_gateway import LLMGateway, ModelResult, FailureClass, LLMError, get_llm_config, create_gateway, build_auth_headers, get_output_limit
except ImportError:
    LLMGateway = None
    ModelResult = None
    FailureClass = None
    LLMError = None
    get_llm_config = None
    create_gateway = None
    def get_output_limit(base_url: str, model: str) -> int:
        """导入失败时的回退实现：未知上限，不钳制。"""
        return 0
    def build_auth_headers(api_key: str, content_type: bool = True) -> dict:
        """导入失败时的回退实现：仅下发 Authorization: Bearer（标准 OpenAI 兼容）。"""
        headers = {'Authorization': f'Bearer {api_key}'}
        if content_type:
            headers['Content-Type'] = 'application/json'
        return headers

# P3：Context Manifest（章节生成前记录上下文来源 + hash + token 预算）
try:
    from context_manifest import ContextManifest, ContextOrchestrator
except ImportError:
    ContextManifest = None
    ContextOrchestrator = None

# P3：PromptSpec 编译器（prompt 可测试 + golden 断言）
try:
    from prompt_spec import PromptSpec, PromptCompiler, load_prompt_spec
except ImportError:
    PromptSpec = None
    PromptCompiler = None
    load_prompt_spec = None

# P4：运行恢复（幂等键 + 单步 retry/resume）
try:
    from run_recovery import IdempotencyKey, RunRecoveryService
except ImportError:
    IdempotencyKey = None
    RunRecoveryService = None

app = Flask(__name__, static_folder=None)
app.config['SECRET_KEY'] = 'fanshu-writer-secret-key'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

# 前端构建产物目录：FANSHU_FRONTEND_DIST 环境变量 > backend/static（git 预构建）> frontend/dist（本地开发）
_static_dist = Path(__file__).parent / 'static'
_frontend_dist_env = os.environ.get('FANSHU_FRONTEND_DIST')
if _frontend_dist_env:
    FRONTEND_DIST = Path(_frontend_dist_env)
elif (_static_dist / 'index.html').exists() and (_static_dist / 'assets').exists():
    # backend/static 完整可用（生产环境 git 跟踪的产物）
    FRONTEND_DIST = _static_dist
else:
    # 本地开发：使用 frontend/dist（vite dev/build 产物）
    FRONTEND_DIST = Path(__file__).parent.parent / 'frontend' / 'dist'

# 数据持久化目录：HF Spaces 用 /data（设 FANSHU_DATA_DIR）、Render 按环境变量、本地 ~/.fanshu-writer
DATA_DIR = Path(os.environ.get('FANSHU_DATA_DIR', Path.home() / '.fanshu-writer'))
DATA_DIR.mkdir(parents=True, exist_ok=True)

# ============================================================================
# 铁律：用户数据（账号及所有作品数据）必须持久化，绝不能因部署/重启而丢失。
# 实现方式：生产环境强制使用外部 PostgreSQL（DATABASE_URL），禁止回退到 SQLite。
# SQLite 仅用于本地开发。云平台（Render/HF Spaces/Railway 等）无 DATABASE_URL
# 时直接拒绝启动，避免出现"注册成功但下次部署数据消失"的假象。
# ============================================================================

DATABASE_URL = os.environ.get('DATABASE_URL', '').strip()
# 检测生产环境：Render 自动设置 RENDER=true；其他云平台通常设置 PORT 或 CI=true
_IS_PROD_ENV = (
    os.environ.get('RENDER', '').lower() in ('true', '1', 'yes')
    or os.environ.get('PORT', '').strip() != ''
    or os.environ.get('HF_SPACE_ID', '').strip() != ''
    or os.environ.get('RAILWAY_PROJECT_ID', '').strip() != ''
)

if DATABASE_URL:
    # 兼容 Render/Heroku 的 postgres:// 前缀（SQLAlchemy 2.0+ 需要 postgresql://）
    if DATABASE_URL.startswith('postgres://'):
        DATABASE_URL = DATABASE_URL.replace('postgres://', 'postgresql://', 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_URL
    print(f'[DB] ✅ 使用 PostgreSQL 外部数据库（铁律：数据持久化，部署不丢失）', flush=True)
else:
    if _IS_PROD_ENV:
        # 生产环境铁律：必须有 DATABASE_URL，否则拒绝启动
        print('=' * 70, flush=True)
        print('[DB][铁律违规] 生产环境未配置 DATABASE_URL，拒绝启动！', flush=True)
        print('[DB][铁律违规] 用户数据必须存到外部 PostgreSQL，绝不能用 SQLite。', flush=True)
        print('[DB][铁律违规] 请在 Render Dashboard → Environment 添加：', flush=True)
        print('[DB][铁律违规]   key=DATABASE_URL  value=postgresql://user:pass@host/dbname?sslmode=require', flush=True)
        print('[DB][铁律违规] 推荐用 Neon 免费版（永久免费 0.5GB）：https://neon.tech', flush=True)
        print('=' * 70, flush=True)
        raise SystemExit('[铁律] 生产环境必须配置 DATABASE_URL，进程退出以保护用户数据')
    # 本地开发：允许 SQLite
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{DATA_DIR}/fanshu.db'
    print(f'[DB] ⚠️ 本地开发模式使用 SQLite：{DATA_DIR}/fanshu.db（生产环境会拒绝启动）', flush=True)

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
# 连接池：防 Neon/Render PG 空闲断连与 SSL 大包断连（pre_ping 剔死连接 / 15min 回收 /
# keepalives 防代理掐断），pool_size 适配免费层
if DATABASE_URL:
    _engine_opts = {
        'pool_pre_ping': True,
        'pool_recycle': 900,  # 15 分钟回收，小于 Neon/Render 空闲超时
        'pool_size': 5,
        'max_overflow': 5,
    }
    # psycopg2 专属：keepalives 防止中间代理（Render/Cloudflare/PGBouncer）在大事务时静默断连
    # 典型触发：UPDATE ai_sessions SET messages_json 写 100KB+ 大包，SSL 管道被掐
    if 'psycopg2' in DATABASE_URL or DATABASE_URL.startswith('postgresql://') or DATABASE_URL.startswith('postgres://'):
        _engine_opts['connect_args'] = {
            'connect_timeout': 15,          # 建连超时 15s
            'keepalives': 1,                # 启用 TCP keepalive
            'keepalives_idle': 30,          # 空闲 30s 开始探测
            'keepalives_interval': 10,      # 每 10s 一个探测包
            'keepalives_count': 5,          # 5 次探测失败判死（总 80s）
            'tcp_user_timeout': 120000,     # 大包未 ACK 2 分钟判死（毫秒，Linux ≥ 2.6.37）
        }
    app.config['SQLALCHEMY_ENGINE_OPTIONS'] = _engine_opts

CORS(app, resources={r"/api/*": {"origins": "*"}})
db = SQLAlchemy(app)

try:
    from blueprints.health_bp import health_bp; app.register_blueprint(health_bp)
    from blueprints.ai_config_bp import ai_config_bp; app.register_blueprint(ai_config_bp)
    from blueprints.chat_collab_bp import chat_collab_bp; app.register_blueprint(chat_collab_bp)
except ImportError:
    pass

# gunicorn prefork 使 _app_engines(WeakKeyDictionary) weakref 失效，请求前检查重注册
@app.before_request
def _ensure_db():
    if app not in db._app_engines: app.extensions.pop('sqlalchemy', None); db.init_app(app)

EXPORTS_DIR = DATA_DIR / 'exports'
EXPORTS_DIR.mkdir(exist_ok=True)
COVERS_DIR = DATA_DIR / 'covers'
COVERS_DIR.mkdir(exist_ok=True)
TEMPLATES_DIR = DATA_DIR / 'templates'
TEMPLATES_DIR.mkdir(exist_ok=True)

class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    username = db.Column(db.String(50), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    email = db.Column(db.String(100), default='')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

    def to_dict(self):
        return {'id': self.id, 'username': self.username, 'email': self.email,
                'created_at': self.created_at.isoformat() if self.created_at else None}

class AuthToken(db.Model):
    __tablename__ = 'auth_tokens'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = db.Column(db.String(36), db.ForeignKey('users.id'), nullable=False)
    token = db.Column(db.String(100), unique=True, nullable=False)
    expires_at = db.Column(db.DateTime, nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

class PasswordResetToken(db.Model):
    __tablename__ = 'password_reset_tokens'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = db.Column(db.String(36), db.ForeignKey('users.id'), nullable=False)
    token = db.Column(db.String(100), unique=True, nullable=False)
    used = db.Column(db.Boolean, default=False)
    expires_at = db.Column(db.DateTime, nullable=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

def generate_token():
    return hashlib.sha256(os.urandom(32)).hexdigest()

def login_required(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            token = request.args.get('token', '')
        if not token:
            return jsonify({'error': '请先登录'}), 401
        at = AuthToken.query.filter_by(token=token).first()
        now = datetime.now(timezone.utc)
        if not at:
            return jsonify({'error': '登录已过期，请重新登录'}), 401
        exp = at.expires_at
        if exp.tzinfo is None:
            exp = exp.replace(tzinfo=timezone.utc)
        if exp < now:
            return jsonify({'error': '登录已过期，请重新登录'}), 401
        request.current_user_id = at.user_id
        return f(*args, **kwargs)
    return decorated

def optional_login(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        at = AuthToken.query.filter_by(token=token).first() if token else None
        now = datetime.now(timezone.utc)
        if at:
            exp = at.expires_at
            if exp.tzinfo is None:
                exp = exp.replace(tzinfo=timezone.utc)
            request.current_user_id = at.user_id if exp > now else None
        else:
            request.current_user_id = None
        return f(*args, **kwargs)
    return decorated

class Book(db.Model):
    __tablename__ = 'books'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    user_id = db.Column(db.String(36), db.ForeignKey('users.id'), default='')
    title = db.Column(db.String(200), nullable=False)
    author = db.Column(db.String(100), default='')
    genre = db.Column(db.String(50), default='other')
    book_type = db.Column(db.String(20), default='novel')  # novel, short_story, script
    synopsis = db.Column(db.Text, default='')
    cover_path = db.Column(db.String(500), default='')
    template_id = db.Column(db.String(36), default='')
    word_count = db.Column(db.Integer, default=0)
    chapter_count = db.Column(db.Integer, default=0)
    status = db.Column(db.String(20), default='draft')  # draft, writing, completed
    target_words = db.Column(db.Integer, default=0)
    # 总卷数（用户自定义，不设上限）：作为五幕总纲/剧情大纲生成的核心依据
    total_volumes = db.Column(db.Integer, default=10)
    # 小说风格流派（JSON 数组字符串，最多3种叠加）：爽文流/虐文流/系统流等
    novel_styles = db.Column(db.Text, default='[]')
    # 技能包三类划分：构思类/文风类/审查类 各自独立的 ID 列表（JSON 数组字符串）
    # 三类在各创作阶段无污染隔离：构思类→大纲/规划，文风类→正文生成，审查类→去AI味/一致性
    master_skill_ids = db.Column(db.Text, default='[]')  # 构思类技能包
    style_skill_ids = db.Column(db.Text, default='[]')   # 文风类技能包（通常选1个）
    review_skill_ids = db.Column(db.Text, default='[]')  # 审查类技能包
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))
    metadata_json = db.Column(db.Text, default='{}')

    chapters = db.relationship('Chapter', backref='book', lazy=True, cascade='all, delete-orphan', order_by='Chapter.order_index')
    characters = db.relationship('Character', backref='book', lazy=True, cascade='all, delete-orphan')
    outlines = db.relationship('Outline', backref='book', lazy=True, cascade='all, delete-orphan', order_by='Outline.order_index')
    daily_stats = db.relationship('DailyStats', backref='book', lazy=True, cascade='all, delete-orphan')

    def to_dict(self):
        return {
            'id': self.id, 'user_id': self.user_id, 'title': self.title, 'author': self.author,
            'genre': self.genre, 'book_type': self.book_type, 'synopsis': self.synopsis,
            'cover_path': self.cover_path, 'template_id': self.template_id,
            'word_count': self.word_count, 'chapter_count': self.chapter_count,
            'status': self.status, 'target_words': self.target_words,
            'total_volumes': self.total_volumes if self.total_volumes else 10,
            'novel_styles': json.loads(self.novel_styles or '[]'),
            'master_skill_ids': json.loads(self.master_skill_ids or '[]'),
            'style_skill_ids': json.loads(self.style_skill_ids or '[]'),
            'review_skill_ids': json.loads(self.review_skill_ids or '[]'),
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None,
            'metadata': json.loads(self.metadata_json or '{}')
        }

class Chapter(db.Model):
    __tablename__ = 'chapters'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    book_id = db.Column(db.String(36), db.ForeignKey('books.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False, default='')
    content = db.Column(db.Text, default='')
    order_index = db.Column(db.Integer, default=0)
    word_count = db.Column(db.Integer, default=0)
    status = db.Column(db.String(20), default='draft')
    is_volume = db.Column(db.Boolean, default=False)
    parent_id = db.Column(db.String(36), default='')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))
    notes = db.Column(db.Text, default='')
    review_snapshots = db.Column(db.Text, default='')  # P1-5：审计-修订闭环 best snapshot 历史（JSON）
    summary = db.Column(db.String(500), default='')  # 章节摘要（用于上下文构建，避免塞入完整正文）
    # M1a: 本章埋/收的伏笔索引 + 本章抽取的事件ID列表（支持跨章溯源与任务清单）
    hooks_set_json = db.Column(db.Text, default='')
    events_extracted_json = db.Column(db.Text, default='')

    versions = db.relationship('ChapterVersion', backref='chapter', lazy=True, cascade='all, delete-orphan', order_by='ChapterVersion.version_num.desc()')

    def to_dict(self, include_content=True):
        d = {
            'id': self.id, 'book_id': self.book_id, 'title': self.title,
            'order_index': self.order_index, 'word_count': self.word_count,
            'status': self.status, 'is_volume': self.is_volume, 'parent_id': self.parent_id,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None,
            'notes': self.notes,
            'summary': self.summary,
            'hooks_set': json.loads(self.hooks_set_json or '[]'),
            'events_extracted': json.loads(self.events_extracted_json or '[]'),
        }
        if include_content:
            d['content'] = self.content
        return d

class ChapterVersion(db.Model):
    __tablename__ = 'chapter_versions'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    chapter_id = db.Column(db.String(36), db.ForeignKey('chapters.id'), nullable=False)
    content = db.Column(db.Text, default='')
    version_num = db.Column(db.Integer, default=1)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    summary = db.Column(db.String(200), default='')

    def to_dict(self):
        return {
            'id': self.id, 'chapter_id': self.chapter_id, 'version_num': self.version_num,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'summary': self.summary, 'content': self.content
        }

class Character(db.Model):
    __tablename__ = 'characters'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    book_id = db.Column(db.String(36), db.ForeignKey('books.id'), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    role = db.Column(db.String(50), default='supporting')  # protagonist, antagonist, supporting
    description = db.Column(db.Text, default='')
    appearance = db.Column(db.Text, default='')
    personality = db.Column(db.Text, default='')
    background = db.Column(db.Text, default='')
    relationships_json = db.Column(db.Text, default='[]')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    def to_dict(self):
        return {
            'id': self.id, 'book_id': self.book_id, 'name': self.name, 'role': self.role,
            'description': self.description, 'appearance': self.appearance,
            'personality': self.personality, 'background': self.background,
            'relationships': json.loads(self.relationships_json or '[]'),
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class Outline(db.Model):
    __tablename__ = 'outlines'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    book_id = db.Column(db.String(36), db.ForeignKey('books.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False, default='')
    content = db.Column(db.Text, default='')
    order_index = db.Column(db.Integer, default=0)
    level = db.Column(db.Integer, default=0)  # 0=act, 1=chapter, 2=scene
    parent_id = db.Column(db.String(36), default='')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    def to_dict(self):
        return {
            'id': self.id, 'book_id': self.book_id, 'title': self.title,
            'content': self.content, 'order_index': self.order_index, 'level': self.level,
            'parent_id': self.parent_id,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class DailyStats(db.Model):
    __tablename__ = 'daily_stats'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    book_id = db.Column(db.String(36), db.ForeignKey('books.id'), nullable=False)
    date = db.Column(db.Date, nullable=False)
    words_written = db.Column(db.Integer, default=0)
    time_spent_minutes = db.Column(db.Integer, default=0)
    chapters_completed = db.Column(db.Integer, default=0)

    def to_dict(self):
        return {
            'id': self.id, 'book_id': self.book_id, 'date': self.date.isoformat(),
            'words_written': self.words_written, 'time_spent_minutes': self.time_spent_minutes,
            'chapters_completed': self.chapters_completed
        }

class Template(db.Model):
    __tablename__ = 'templates'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, default='')
    genre = db.Column(db.String(50), default='other')
    book_type = db.Column(db.String(20), default='novel')
    structure_json = db.Column(db.Text, default='[]')
    prompts_json = db.Column(db.Text, default='{}')
    is_builtin = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

    def to_dict(self):
        return {
            'id': self.id, 'name': self.name, 'description': self.description,
            'genre': self.genre, 'book_type': self.book_type,
            'structure': json.loads(self.structure_json or '[]'),
            'prompts': json.loads(self.prompts_json or '{}'),
            'is_builtin': self.is_builtin,
            'created_at': self.created_at.isoformat() if self.created_at else None
        }

class AISession(db.Model):
    __tablename__ = 'ai_sessions'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    book_id = db.Column(db.String(36), db.ForeignKey('books.id'), nullable=True)
    scope = db.Column(db.String(50), default='general')  # general, character, plot, chapter
    scope_id = db.Column(db.String(36), default='')
    title = db.Column(db.String(200), default='')
    messages_json = db.Column(db.Text, default='[]')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    def to_dict(self):
        return {
            'id': self.id, 'book_id': self.book_id, 'scope': self.scope,
            'scope_id': self.scope_id, 'title': self.title,
            'messages': json.loads(self.messages_json or '[]'),
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class AIConfig(db.Model):
    __tablename__ = 'ai_config'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    name = db.Column(db.String(50), default='默认配置')  # 多配置标识
    is_active = db.Column(db.Boolean, default=True, index=True)  # 是否激活（同时仅一个）
    provider = db.Column(db.String(50), default='deepseek')
    model = db.Column(db.String(100), default='deepseek-chat')
    recognition_model = db.Column(db.String(100), default='')  # AI识别专用模型，为空时使用model
    api_key = db.Column(db.String(200), default='')
    base_url = db.Column(db.String(300), default='https://api.deepseek.com')
    temperature = db.Column(db.Float, default=0.7)
    max_tokens = db.Column(db.Integer, default=4096)

    def to_dict(self):
        return {
            'id': self.id, 'name': self.name or '默认配置', 'is_active': self.is_active,
            'provider': self.provider, 'model': self.model,
            'recognition_model': self.recognition_model or '',
            'api_key': '***' if self.api_key else '', 'base_url': self.base_url,
            'temperature': self.temperature, 'max_tokens': self.max_tokens,
            'has_key': bool(self.api_key)
        }

    def get_model_for_task(self, task_type='creation'):
        """根据任务类型返回对应模型：recognition或creation"""
        if task_type == 'recognition' and self.recognition_model:
            return self.recognition_model
        return self.model

    @classmethod
    def get_active(cls):
        cfg = cls.query.filter_by(is_active=True).first()
        if cfg:
            return cfg
        cfg = cls.query.order_by(cls.id.asc()).first()
        if cfg:
            cfg.is_active = True
            db.session.commit()
            return cfg
        cfg = cls(name='默认配置', is_active=True)
        db.session.add(cfg)
        db.session.commit()
        return cfg


    @classmethod
    def get_by_id(cls, cfg_id):
        """P1-1 会话级切模型：按ID取指定配置（找不到返回None），绝不修改全局激活。"""
        if not cfg_id: return None
        try:
            return cls.query.filter_by(id=str(cfg_id)).first()
        except Exception:
            return None


class AppPreference(db.Model):
    __tablename__ = 'app_preferences'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    key = db.Column(db.String(100), unique=True, nullable=False)
    value = db.Column(db.Text, default='')

    @staticmethod
    def get(key, default=None):
        pref = AppPreference.query.filter_by(key=key).first()
        return pref.value if pref else default

    @staticmethod
    def set(key, value):
        pref = AppPreference.query.filter_by(key=key).first()
        if pref:
            pref.value = value
        else:
            db.session.add(AppPreference(key=key, value=value))
        db.session.commit()

class StageContent(db.Model):
    __tablename__ = 'stage_contents'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    book_id = db.Column(db.String(36), db.ForeignKey('books.id'), nullable=False)
    stage_key = db.Column(db.String(100), nullable=False)
    content = db.Column(db.Text, default='')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    __table_args__ = (db.UniqueConstraint('book_id', 'stage_key'),)

    def to_dict(self):
        return {
            'id': self.id, 'book_id': self.book_id, 'stage_key': self.stage_key,
            'content': self.content,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class BookBible(db.Model):
    __tablename__ = 'book_bible'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    book_id = db.Column(db.String(36), db.ForeignKey('books.id'), nullable=False, unique=True)
    worldbuilding = db.Column(db.Text, default='')
    character_profiles = db.Column(db.Text, default='')
    timeline = db.Column(db.Text, default='')
    foreshadowing = db.Column(db.Text, default='')
    style_guide = db.Column(db.Text, default='')
    key_rules = db.Column(db.Text, default='')
    locations = db.Column(db.Text, default='')
    concept = db.Column(db.Text, default='')
    plot_design = db.Column(db.Text, default='')
    generated_summary = db.Column(db.Text, default='')
    # 关系图谱专用字段（与 character_profiles 解耦，避免互相覆盖导致角色丢失）
    relation_graph = db.Column(db.Text, default='')
    # 物资库：JSON 数组，按卷存储势力/角色的物品、功法、法宝、境界等
    inventory = db.Column(db.Text, default='')
    # 人物按卷：JSON 数组，每卷的人物档案
    character_volumes = db.Column(db.Text, default='')
    # 动态文件按卷：JSON 数组，每卷的动态分类摘要
    dynamic_volumes = db.Column(db.Text, default='')
    # 伏笔按卷：JSON 数组，每卷的伏笔识别数据
    foreshadowing_volumes = db.Column(db.Text, default='')
    # 地图按卷：JSON 数组，每卷的地点识别数据
    locations_volumes = db.Column(db.Text, default='')
    # 防遗忘检查报告历史：JSON 数组，每份报告含 id/title/checked_at/scope/volume_ids/report_json
    anti_forget_reports = db.Column(db.Text, default='')
    # P0-2：伏笔 DAG（JSON，结构化伏笔图，与 foreshadowing 文本字段并存）
    # 文本字段供前端展示，DAG 供后端状态追踪/prompt注入/逾期检测用
    foreshadowing_graph = db.Column(db.Text, default='')
    # P1-4：四级大纲层级（JSON，master→arc→section→chapter，由 timeline 自动构建）
    outline_hierarchy = db.Column(db.Text, default='')
    # P1-6：章级变更日志（JSON 数组，每章的 12 类 CHANGES delta，支持重写回滚）
    chapter_changes_log = db.Column(db.Text, default='')
    # 借鉴 PlotPilot 检查点快照：每5章自动备份 BookBible+DynamicMemory 关键字段，支持回滚
    state_snapshots = db.Column(db.Text, default='')
    # M1a: 全书事件日志（时间序列），每章写作后自动抽取事件追加
    event_log_json = db.Column(db.Text, default='')
    # M1b: 实体注册表（JSON），跨维度统一追踪人/地/物/势力/技能及其别名
    entity_registry_json = db.Column(db.Text, default='')
    # M4: 失败记录库（JSON），供 Meta-LLM 分析并优化 prompt
    failure_log_json = db.Column(db.Text, default='')
    # M4b: 用户已采纳的 prompt 补丁列表（JSON 数组），每条含 id/category/patch_text/handled_bucket_key/applied_at
    # 生成任何维度/章节时自动追加到 system prompt 末尾（tail-rule 之后）
    prompt_patches_json = db.Column(db.Text, default='')
    # M4c: 忽略的失败模式 bucket（JSON 数组），被忽略的 bucket 不出现在 optimization-report 里
    ignored_failure_buckets_json = db.Column(db.Text, default='')
    # 总卷数 + 风格流派（与 Book 表同步，创作时从 bible 直接读取注入各维度）
    total_volumes = db.Column(db.Integer, default=10)
    novel_styles = db.Column(db.Text, default='[]')
    last_synced_at = db.Column(db.DateTime)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    def to_dict(self):
        return {
            'id': self.id, 'book_id': self.book_id, 'worldbuilding': self.worldbuilding,
            'character_profiles': self.character_profiles, 'timeline': self.timeline,
            'foreshadowing': self.foreshadowing, 'style_guide': self.style_guide,
            'key_rules': self.key_rules, 'locations': self.locations,
            'concept': self.concept, 'plot_design': self.plot_design,
            'generated_summary': self.generated_summary,
            'relation_graph': self.relation_graph,
            'inventory': self.inventory or '',
            'character_volumes': self.character_volumes or '',
            'dynamic_volumes': self.dynamic_volumes or '',
            'foreshadowing_volumes': self.foreshadowing_volumes or '',
            'locations_volumes': self.locations_volumes or '',
            'anti_forget_reports': self.anti_forget_reports or '',
            # P0-2/P1-4/P1-6 新增字段
            'foreshadowing_graph': self.foreshadowing_graph or '',
            'outline_hierarchy': self.outline_hierarchy or '',
            'chapter_changes_log': self.chapter_changes_log or '',
            'state_snapshots': self.state_snapshots or '',
            'event_log': json.loads(self.event_log_json or '[]'),
            'entity_registry': json.loads(self.entity_registry_json or '{}'),
            'failure_log': json.loads(self.failure_log_json or '[]'),
            'prompt_patches': json.loads(self.prompt_patches_json or '[]'),
            'ignored_failure_buckets': json.loads(self.ignored_failure_buckets_json or '[]'),
            'last_synced_at': self.last_synced_at.isoformat() if self.last_synced_at else None,
            # P1-2修复：补 total_volumes / novel_styles，让前端可见 bible 权威值
            'total_volumes': self.total_volumes if self.total_volumes else 10,
            'novel_styles': self.novel_styles or '[]'
        }

class SkillPack(db.Model):
    __tablename__ = 'skill_packs'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    name = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, default='')
    genre = db.Column(db.String(50), default='other')
    book_type = db.Column(db.String(20), default='short_story')
    stage_keys_json = db.Column(db.Text, default='[]')
    workflow_json = db.Column(db.Text, default='[]')
    prompts_json = db.Column(db.Text, default='{}')
    is_builtin = db.Column(db.Boolean, default=False)
    icon = db.Column(db.String(10), default='📦')
    github_source = db.Column(db.String(500), default='')  # GitHub 仓库地址，用于拉取更新
    github_synced_at = db.Column(db.DateTime, nullable=True)  # 上次同步时间
    # 技能包三类划分：master=构思类 / style=文风类 / review=审查类
    category = db.Column(db.String(20), default='master')
    # 文风类专属：题材目标（fantasy/urban/mystery/history...），构思/审查类为空
    genre_target = db.Column(db.String(50), default='')
    # 同类多包时的注入优先级（数字小的先注入），默认100
    priority = db.Column(db.Integer, default=100)
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

    def to_dict(self):
        return {
            'id': self.id, 'name': self.name, 'description': self.description,
            'genre': self.genre, 'book_type': self.book_type,
            'stage_keys': json.loads(self.stage_keys_json or '[]'),
            'workflow': json.loads(self.workflow_json or '[]'),
            'prompts': json.loads(self.prompts_json or '{}'),
            'is_builtin': self.is_builtin, 'icon': self.icon,
            'github_source': self.github_source or '',
            'category': self.category or 'master',
            'genre_target': self.genre_target or '',
            'priority': self.priority if self.priority is not None else 100,
            'github_synced_at': self.github_synced_at.isoformat() if self.github_synced_at else None,
            'created_at': self.created_at.isoformat() if self.created_at else None
        }

class DynamicMemory(db.Model):
    """动态文件库 - 长篇小说防遗忘系统（5文件版）"""
    __tablename__ = 'dynamic_memory'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    book_id = db.Column(db.String(36), db.ForeignKey('books.id'), nullable=False, unique=True)
    # 5个动态文件（JSON文本）
    narrative_engine = db.Column(db.Text, default='')       # 叙事引擎
    foreshadowing_tracker = db.Column(db.Text, default='')  # 伏笔追踪器
    character_ecosystem = db.Column(db.Text, default='')    # 角色生态系统
    ability_world = db.Column(db.Text, default='')          # 能力与世界观
    health_dashboard = db.Column(db.Text, default='')       # 健康度仪表盘
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    FILE_KEYS = ['narrative_engine', 'foreshadowing_tracker', 'character_ecosystem', 'ability_world', 'health_dashboard']

    def to_dict(self):
        return {
            'id': self.id, 'book_id': self.book_id,
            'narrative_engine': self.narrative_engine or '',
            'foreshadowing_tracker': self.foreshadowing_tracker or '',
            'character_ecosystem': self.character_ecosystem or '',
            'ability_world': self.ability_world or '',
            'health_dashboard': self.health_dashboard or '',
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

    @staticmethod
    def get_empty_template(file_key):
        """返回各文件的空模板"""
        templates = {
            'narrative_engine': json.dumps({
                'state': {
                    'book_name': '', 'current_chapter': 0, 'current_volume': 1,
                    'current_node': '', 'mc_status': '', 'recent_events': [],
                    'active_hooks': [], 'pending_foreshadowing': [], 'cost_points': 0, 'cost_threshold': 10
                },
                'timeline': [],
                'chapters': []
            }, ensure_ascii=False, indent=2),
            'foreshadowing_tracker': json.dumps({
                'foreshadowing': [],
                'scan_rules': {
                    'short_cycle': 20, 'mid_cycle': 75, 'long_cycle': 150,
                    'alert_levels': {'warning': '待回收', 'danger': '高危遗忘', 'critical': '核心悬念'}
                }
            }, ensure_ascii=False, indent=2),
            'character_ecosystem': json.dumps({
                'characters': [],
                'relationships': []
            }, ensure_ascii=False, indent=2),
            'ability_world': json.dumps({
                'ability_log': [],
                'world_facts': []
            }, ensure_ascii=False, indent=2),
            'health_dashboard': json.dumps({
                'thread_health': {
                    'main_line': {'last_chapter': 0, 'status': 'ok'},
                    'sub_line_a': {'last_chapter': 0, 'status': 'ok'},
                    'sub_line_b': {'last_chapter': 0, 'status': 'ok'},
                    'dark_line': {'last_chapter': 0, 'status': 'ok'}
                },
                'chapter_type_distribution': {},
                'foreshadowing_aging': {},
                'ai_flavor_trend': [],
                'dialogue_ratio_trend': [],
                'character_appearances': {},
                'alerts': []
            }, ensure_ascii=False, indent=2),
        }
        return templates.get(file_key, '{}')

class DynamicReport(db.Model):
    """动态文件报告 - 长篇小说防遗忘摘要（每5-10章自动生成）"""
    __tablename__ = 'dynamic_reports'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    book_id = db.Column(db.String(36), db.ForeignKey('books.id'), nullable=False)
    title = db.Column(db.String(200), default='')        # 如 "动态-(1-5章)"
    content = db.Column(db.Text, default='')             # 汇总报告，≤500字
    chapter_start = db.Column(db.Integer, default=0)     # 起始章号
    chapter_end = db.Column(db.Integer, default=0)       # 结束章号
    auto_generated = db.Column(db.Boolean, default=False) # 是否自动生成
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))
    updated_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc), onupdate=lambda: datetime.now(timezone.utc))

    def to_dict(self):
        return {
            'id': self.id, 'book_id': self.book_id,
            'title': self.title, 'content': self.content,
            'chapter_start': self.chapter_start, 'chapter_end': self.chapter_end,
            'auto_generated': self.auto_generated,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None,
        }

class PromptTemplate(db.Model):
    __tablename__ = 'prompt_templates'
    id = db.Column(db.String(36), primary_key=True, default=lambda: str(uuid.uuid4()))
    name = db.Column(db.String(200), nullable=False)
    agent_id = db.Column(db.String(100), nullable=False)
    book_type = db.Column(db.String(20), default='short_story')
    genre = db.Column(db.String(50), default='other')
    content = db.Column(db.Text, default='')
    is_builtin = db.Column(db.Boolean, default=False)
    description = db.Column(db.Text, default='')
    created_at = db.Column(db.DateTime, default=lambda: datetime.now(timezone.utc))

    def to_dict(self):
        return {
            'id': self.id, 'name': self.name, 'agent_id': self.agent_id,
            'book_type': self.book_type, 'genre': self.genre,
            'content': self.content, 'is_builtin': self.is_builtin,
            'description': self.description,
            'created_at': self.created_at.isoformat() if self.created_at else None
        }

def count_words(text):
    """统计字数：中文按字符计数，英文按单词计数"""
    if not text:
        return 0
    import re
    # 移除空白字符
    cleaned = re.sub(r'\s+', '', text)
    # 统计中文字符数（含中文标点）
    cn_chars = len(re.findall(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]', cleaned))
    # 统计英文单词数
    en_words = len(re.findall(r'[a-zA-Z]+', text))
    # 统计数字串
    numbers = len(re.findall(r'\d+', text))
    return cn_chars + en_words + numbers

def _count_cn_chars(s):
    """统计中文字符数（含标点），去除空白。四模式字数校验统一用此函数。"""
    if not s:
        return 0
    return len(re.sub(r'\s', '', s))

def _llm_chat(messages, api_key=None, base_url=None, model=None,
              temperature=0.7, max_tokens=4096, timeout=180):
    """【P3 统一 LLM 入口】所有 LLM 调用应经此函数。

    封装 LLMGateway，提供：
      - 错误分类（auth/quota/timeout/empty/...）
      - 智能重试（timeout/unavailable 自动重试 2 次，auth 不重试）
      - 空内容检测（自动标记 empty_response）

    返回 (content: str, error: str)。
      - 成功：content 非空, error 为空串
      - 失败：content 为空, error 含错误描述
    """
    if LLMGateway is None:
        # gateway 不可用时回退到直接 requests.post（兼容旧环境）
        try:
            resp = requests.post(f'{base_url}/chat/completions',
                headers=build_auth_headers(api_key),
                json={'model': model, 'messages': messages,
                      'temperature': temperature, 'max_tokens': max_tokens},
                timeout=timeout)
            result = resp.json()
            content = result['choices'][0]['message']['content'] or ''
            if result.get('error'):
                return '', result['error'].get('message', str(result['error']))[:200]
            return content, ''
        except Exception as e:
            return '', f'LLM 调用失败：{str(e)[:200]}'

    # 确保 base_url 以 /v1 结尾
    if base_url and not base_url.rstrip('/').endswith('/v1'):
        base_url = base_url.rstrip('/') + '/v1'
    gw = LLMGateway(base_url, api_key, model, timeout=timeout)
    result = gw.chat(messages, temperature=temperature, max_tokens=max_tokens)
    if result.ok:
        return result.content, ''
    return '', result.error or 'LLM 返回空内容'

def _ensure_word_count(content, api_key, base_url, model, max_tokens=12000, chapter_num=0, count_fn=None):
    """【字数铁律】公共字数修正函数：初稿字数不在 2300-2500 区间时调 AI 重写。
    四种创作模式（多Agent/流式/连续/连续流式）统一调用此函数。
    返回 (修正后内容, 备注)。
    备注：空串表示未触发修正或修正成功无异常；非空串表示修正过程的备注信息。

    count_fn: 字数统计函数，默认 _count_cn_chars（去空白含所有标点）。
              传入 count_words 可与章节保存/列表显示口径一致（中文+中文标点+英文单词+数字串）。
    """
    if not content or not content.strip():
        return content, ''

    if count_fn is None:
        count_fn = _count_cn_chars

    draft_len = count_fn(content)
    if 2300 <= draft_len <= 2500:
        return content, ''  # 字数达标，无需修正

    if draft_len < 200:
        # 过短可能是生成失败，不浪费 token 重写
        return content, f'[字数铁律] 初稿仅{draft_len}字，疑似生成失败，未触发重写'

    direction = '精简删减冗余' if draft_len > 2500 else '扩写补充场景细节'
    method = ('精简方法：删冗余形容词/重复心理描写/过度环境渲染/总结性句子，保留对话动作与剧情走向。'
              if draft_len > 2500 else
              '扩写方法：增加感官细节/动作描写/对话节拍/场景纵深，不增加新剧情不改变走向。')
    rewrite_system = f"""你是字数修正编辑。当前章节初稿{draft_len}字，需{direction}至 2400字±100（2300-2500字区间，含标点）。

【字数绝对铁律】最终输出必须落在 2300-2500 字区间，这是不可违反的硬约束，优先级高于一切。
【保留要求】保留原章节的剧情走向、人物对话、章尾钩子、关键信息，不改变故事内容，只调整篇幅。
{method}
只输出修正后的完整正文，不输出任何说明或前缀。"""
    try:
        # 【输出上限适配】默认 12000 会撞 8k 输出上限的模型直接 400，按已知/已学习上限钳制
        _wc_max_tok = min(int(max_tokens), get_output_limit(base_url, model) or int(max_tokens))
        rewrite_resp = requests.post(f'{base_url}/chat/completions',
            headers=build_auth_headers(api_key),
            json={'model': model,
                  'messages': [{'role': 'system', 'content': rewrite_system},
                               {'role': 'user', 'content': f'请修正以下章节正文字数：\n\n{content}'}],
                  'temperature': 0.5, 'max_tokens': _wc_max_tok},
            timeout=180)
        rewrite_result = rewrite_resp.json()
        rewritten = rewrite_result['choices'][0]['message']['content'].strip()
        rewritten_len = count_fn(rewritten)
        if rewritten and 2300 <= rewritten_len <= 2500:
            return rewritten, f'[字数铁律] 初稿{draft_len}字，AI修正至{rewritten_len}字。'
        elif rewritten and rewritten_len > 500:
            if abs(rewritten_len - 2400) < abs(draft_len - 2400):
                return rewritten, f'[字数铁律] 初稿{draft_len}字，AI修正后{rewritten_len}字仍偏离，已采纳更接近目标版本。'
            return content, f'[字数铁律] 初稿{draft_len}字，AI修正后{rewritten_len}字仍偏离，保留初稿。'
        return content, f'[字数铁律] 初稿{draft_len}字，AI修正返回异常，保留初稿。'
    except Exception as e:
        return content, f'[字数铁律] 初稿{draft_len}字，AI修正异常：{str(e)[:80]}，保留初稿。'

_CN_DIGITS = {'一':1,'二':2,'两':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9,'零':0,'〇':0}
_CN_UNITS = {'十':10,'百':100,'千':1000,'万':10000,'亿':100000000}
# 全角数字 → 半角
_FULLWIDTH_DIGITS = str.maketrans('０１２３４５６７８９', '0123456789')

def _chinese_to_int(s):
    """将中文数字（如 十一/二十三/一百零五/一千二百/两万零一）转为 int。
    无法解析返回 None。支持阿拉伯数字与中文数字混用（如 2十/1百2/12）。"""
    if not s:
        return None
    total = 0
    cur = 0
    wan_part = 0  # 万位以上的累积
    yi_part = 0   # 亿位以上的累积
    has_digit = False
    has_value = False  # 是否出现过非零值
    i = 0
    n = len(s)
    while i < n:
        ch = s[i]
        if ch.isdigit():
            # 连续阿拉伯数字作为一个整体整数
            j = i
            while j < n and s[j].isdigit():
                j += 1
            cur = int(s[i:j])
            has_digit = True
            if cur != 0:
                has_value = True
            i = j
            continue
        elif ch in _CN_DIGITS:
            cur = _CN_DIGITS[ch]
            has_digit = True
            if cur != 0:
                has_value = True
            i += 1
        elif ch == '亿':
            if cur == 0 and total == 0:
                return None
            if cur == 0:
                cur = 1
            yi_part = (yi_part + wan_part + total + cur) * 100000000
            wan_part = 0
            total = 0
            cur = 0
            has_value = True
            i += 1
        elif ch == '万':
            if cur == 0 and total == 0:
                return None
            if cur == 0:
                cur = 1
            wan_part = (wan_part + total + cur) * 10000
            total = 0
            cur = 0
            has_value = True
            i += 1
        elif ch in _CN_UNITS:
            if cur == 0:
                cur = 1  # "十" 单独出现视为 10
            total += cur * _CN_UNITS[ch]
            cur = 0
            has_digit = True
            has_value = True
            i += 1
        else:
            return None  # 含非数字字符，无法解析
    if not has_digit:
        return None
    result = yi_part + wan_part + total + cur
    # 允许 0（如"第零章"），仅当原串确实含数字时
    if result == 0:
        return 0 if has_digit else None
    return result

def parse_chapter_number(title):
    """从章节标题解析章节号，返回 int 或 None。
    支持的格式（阿拉伯/中文数字可混用）：
      第N章/第N章/第N回/第N节/第N话/第N集/第N幕/第N折/第N卷/第N部/第N篇...
      Chapter N / Ch.N / CHAPTER N / Episode N / Ep.N
      N.标题 / N、标题 / N:标题 / N-标题 / N 标题 / N标题
      十一 标题（行首纯中文数字）
    支持括号包裹：【第11章】、（第十一章）、[Chapter 11]
    多个章节号时取最后一个（如 第3卷第5章 → 5）。"""
    if not title:
        return None
    import re
    # 全角数字转半角
    t = title.strip().translate(_FULLWIDTH_DIGITS)

    # 收集所有“第...后缀”模式的章节号，取最后一个（最细粒度）
    # 后缀字符集：章/节/回/卷/部/篇/话/集/幕/折/更/段/讲/课/夜/日/年/季/场 等
    suffix_class = '章节回卷部篇话集幕折更段讲课夜日年季场'
    matches = re.findall(r'第\s*([0-9零一二三四五六七八九十百千万亿两〇]+)\s*[' + suffix_class + r']', t)
    if matches:
        # 取最后一个匹配（如 第3卷第5章 → 取5）
        n = _chinese_to_int(matches[-1])
        if n is not None:
            return n

    # Chapter N / Ch.N / CHAPTER N / Episode N / Ep.N（大小写不敏感）
    m = re.search(r'(?:chapter|ch|episode|ep)\.?\s*(\d+)', t, re.IGNORECASE)
    if m:
        return int(m.group(1))

    # 行首：数字（阿拉伯）+ 可选分隔符 + 后续标题
    # 分隔符：. 、 : ： - ) ] 】 空格 ，或直接跟中文/字母
    # 排除日期格式：数字紧接 年/月/日/时/分/秒/号 或 N-N-N 连续日期格式 视为非章节号
    m = re.match(r'\s*(\d+)(?:\s*[\.、:：\-\)\]】，;；]|\s+|\s*(?=[\u4e00-\u9fffA-Za-z]))', t)
    if m:
        num_end = m.end(1)
        rest = t[num_end:]
        is_date = False
        # 数字后紧跟日期单位（年/月/日/时/分/秒/号）
        if rest and rest[0] in '年月日时分秒号':
            is_date = True
        # YYYY-MM-DD / YYYY-MM 日期格式（数字-数字-数字）
        if re.match(r'\s*-\d{1,2}(-\d{1,2})?(\s|$)', rest):
            is_date = True
        if not is_date:
            return int(m.group(1))

    # 行首：纯中文数字 + 可选分隔符
    m = re.match(r'\s*([零一二三四五六七八九十百千万亿两〇]+)(?:\s*[\.、:：\-\)\]】，;；]|\s+)', t)
    if m:
        n = _chinese_to_int(m.group(1))
        if n is not None:
            return n

    # 行首纯数字（整行就是一个数字）
    m = re.match(r'\s*(\d+)\s*$', t)
    if m:
        return int(m.group(1))

    return None

def resort_chapters_by_title(book_id, rebin_volumes=False):
    """按章节标题中的章节号对作品章节重新排序（稳定排序）。
    - 有章节号的按章节号升序排在前
    - 无章节号的保持原相对顺序排在后面
    - 重新分配 order_index 为 0..N-1
    - rebin_volumes=True 时，按 50 章/卷重新归入卷（用于批量导入场景）
    返回重排后的章节数。"""
    chs = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    if not chs:
        return 0
    # 构建排序键：(是否无章节号, 章节号, 原 order_index)
    keyed = []
    for ch in chs:
        n = parse_chapter_number(ch.title or '')
        if n is not None:
            keyed.append((0, n, ch.order_index, ch))
        else:
            keyed.append((1, 0, ch.order_index, ch))
    keyed.sort(key=lambda x: (x[0], x[1], x[2]))
    # 重新分配 order_index
    for i, item in enumerate(keyed):
        item[3].order_index = i

    # 可选：按 50 章/卷重新归入卷
    if rebin_volumes:
        vols = Chapter.query.filter_by(book_id=book_id, is_volume=True).order_by(Chapter.order_index).all()
        total_chs = len(keyed)
        needed_vols = max(1, (total_chs + 49) // 50) if total_chs > 0 else 0
        # 多余的空卷删除（避免导入后残留空卷导致卷列表错乱）
        if len(vols) > needed_vols:
            for extra_vol in vols[needed_vols:]:
                # 清空其下章节的 parent_id（若有），再删除
                Chapter.query.filter_by(book_id=book_id, parent_id=extra_vol.id, is_volume=False).update({'parent_id': ''})
                db.session.delete(extra_vol)
            vols = vols[:needed_vols]
        # 不足的卷自动补建（保留已有卷名）
        while len(vols) < needed_vols:
            vidx = len(vols)
            ch_start = vidx * 50 + 1
            ch_end = min((vidx + 1) * 50, total_chs)
            vol = Chapter(
                book_id=book_id,
                title=f'第{vidx + 1}卷（第{ch_start}-{ch_end}章）',
                content='', order_index=total_chs + vidx,
                is_volume=True, parent_id='', word_count=0,
            )
            db.session.add(vol)
            db.session.flush()
            vols.append(vol)
        # 重置所有卷的 order_index 为 total_chs + vol_idx（保证卷顺序稳定，排在所有章节之后）
        for vidx, vol in enumerate(vols):
            vol.order_index = total_chs + vidx
        # 重新分配 parent_id
        for i, item in enumerate(keyed):
            vol_idx = i // 50
            if vol_idx < len(vols):
                item[3].parent_id = vols[vol_idx].id
    db.session.flush()
    return len(keyed)

def update_book_stats(book_id):
    book = Book.query.get(book_id)
    if not book:
        return
    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).all()
    book.word_count = sum(c.word_count for c in chapters)
    book.chapter_count = len(chapters)
    book.updated_at = datetime.now(timezone.utc)
    db.session.commit()

def build_outline_tree(outlines):
    outline_map = {o.id: o.to_dict() for o in outlines}
    tree = []
    for o in outlines:
        node = outline_map[o.id]
        node['children'] = []
        if o.parent_id and o.parent_id in outline_map:
            outline_map[o.parent_id]['children'].append(node)
        elif not o.parent_id:
            tree.append(node)
    return tree

# ==== Auth API ====

@app.route('/api/auth/register', methods=['POST'])
def register():
    data = request.json
    username = (data.get('username', '')).strip()
    password = (data.get('password', '')).strip()
    email = (data.get('email', '')).strip()
    if len(username) < 2 or len(username) > 30:
        return jsonify({'error': '用户名需2-30个字符'}), 400
    if len(password) < 4:
        return jsonify({'error': '密码至少4个字符'}), 400
    if not email:
        return jsonify({'error': '邮箱不能为空'}), 400
    if User.query.filter_by(username=username).first():
        return jsonify({'error': '用户名已存在'}), 409
    if User.query.filter_by(email=email).first():
        return jsonify({'error': '该邮箱已被注册'}), 409
    user = User(username=username, password_hash=generate_password_hash(password), email=email)
    db.session.add(user)
    db.session.commit()
    token = generate_token()
    expires = datetime.now(timezone.utc) + timedelta(days=30)
    db.session.add(AuthToken(user_id=user.id, token=token, expires_at=expires))
    db.session.commit()
    return jsonify({'user': user.to_dict(), 'token': token}), 201

@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.json
    username = (data.get('username', '')).strip()
    password = (data.get('password', '')).strip()
    # 支持用户名或邮箱登录：先按邮箱查，查不到再按用户名查
    user = User.query.filter_by(email=username).first()
    if not user:
        user = User.query.filter_by(username=username).first()
    if not user or not check_password_hash(user.password_hash, password):
        return jsonify({'error': '用户名或密码错误'}), 401
    token = generate_token()
    expires = datetime.now(timezone.utc) + timedelta(days=30)
    db.session.add(AuthToken(user_id=user.id, token=token, expires_at=expires))
    db.session.commit()
    return jsonify({'user': user.to_dict(), 'token': token})

@app.route('/api/auth/me', methods=['GET'])
@login_required
def get_me():
    user = User.query.get(request.current_user_id)
    return jsonify(user.to_dict() if user else {'error': 'User not found'})

@app.route('/api/auth/logout', methods=['POST'])
def logout():
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    AuthToken.query.filter_by(token=token).delete()
    db.session.commit()
    return jsonify({'success': True})

# ==== 邮件发送（用于找回密码） ====
# SMTP 配置通过环境变量覆盖；默认发件邮箱为 xiyiji@88.com
SMTP_HOST = os.environ.get('FANSHU_SMTP_HOST', 'smtp.qiye.aliyun.com')
SMTP_PORT = int(os.environ.get('FANSHU_SMTP_PORT', '465'))
SMTP_USER = os.environ.get('FANSHU_SMTP_USER', 'xiyiji@88.com')
SMTP_PASSWORD = os.environ.get('FANSHU_SMTP_PASSWORD', '')
SMTP_FROM_NAME = os.environ.get('FANSHU_SMTP_FROM_NAME', '蚂蚁写作')
SMTP_FROM_ADDR = os.environ.get('FANSHU_SMTP_FROM_ADDR', 'xiyiji@88.com')
# 前端站点地址，用于拼接重置链接
SITE_BASE_URL = os.environ.get('FANSHU_SITE_BASE_URL', '')

def send_reset_email(to_email, reset_token, site_url=None):
    """发送密码重置邮件。如果 SMTP 未配置密码则降级为返回链接（开发模式）。
    返回 (ok: bool, msg: str, reset_link: str)。
    """
    # 优先使用前端传入的 site_url（前后端分离部署时至关重要），其次环境变量，最后回退到后端地址
    base = (site_url or SITE_BASE_URL or request.host_url.rstrip('/')).rstrip('/')
    # 使用 # 锚点，兼容 HashRouter：/ sometime/#/reset-password?token=xxx
    reset_link = f"{base}/#/reset-password?token={reset_token}"

    subject = '【蚂蚁写作】找回您的账号密码'
    body = (
        f"您好，\n\n"
        f"我们收到了您重置蚂蚁写作账号密码的请求。\n\n"
        f"请点击下方链接重置密码（链接 30 分钟内有效）：\n"
        f"{reset_link}\n\n"
        f"如果您没有发起过此请求，请忽略本邮件，您的账号密码不会变更。\n\n"
        f"—— 蚂蚁写作团队"
    )

    try:
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        from email.utils import formataddr

        msg = MIMEMultipart()
        msg['From'] = formataddr((SMTP_FROM_NAME, SMTP_FROM_ADDR))
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(body, 'plain', 'utf-8'))

        if SMTP_PASSWORD:
            # 生产/已配置：使用 SSL 直连 SMTP 服务器
            if SMTP_PORT == 465:
                server = smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=15)
            else:
                server = smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=15)
                server.starttls()
            try:
                server.login(SMTP_USER, SMTP_PASSWORD)
                server.sendmail(SMTP_FROM_ADDR, [to_email], msg.as_string())
            finally:
                server.quit()
            return True, '邮件已发送', reset_link
        else:
            # 开发环境降级：SMTP 未配置，无法实际发邮件，返回链接供前端展示
            app.logger.warning('[SMTP未配置] 密码重置邮件未实际发送，返回链接供开发调试。')
            app.logger.info('---- 重置邮件内容 ----\n%s\n--------------------', body)
            return True, 'SMTP未配置，已生成重置链接', reset_link
    except Exception as e:
        app.logger.exception('发送重置邮件失败')
        return False, f'邮件发送失败：{e}', reset_link

# ==== 修改密码 / 找回密码 / 重置密码 ====

@app.route('/api/auth/change-password', methods=['POST'])
@login_required
def change_password():
    """已登录用户修改密码：需要原密码验证。"""
    data = request.json or {}
    old_password = (data.get('old_password', '') or '').strip()
    new_password = (data.get('new_password', '') or '').strip()
    if not old_password or not new_password:
        return jsonify({'error': '请输入原密码和新密码'}), 400
    if len(new_password) < 4:
        return jsonify({'error': '新密码至少4个字符'}), 400

    user = User.query.get(request.current_user_id)
    if not user:
        return jsonify({'error': '用户不存在'}), 404
    if not check_password_hash(user.password_hash, old_password):
        return jsonify({'error': '原密码错误'}), 401

    user.password_hash = generate_password_hash(new_password)
    db.session.commit()
    return jsonify({'success': True})

@app.route('/api/auth/forgot-password', methods=['POST'])
def forgot_password():
    """用户输入邮箱，生成重置令牌并发送重置邮件。"""
    data = request.json or {}
    email = (data.get('email', '') or '').strip().lower()
    site_url = (data.get('site_url', '') or '').strip()
    if not email or '@' not in email:
        return jsonify({'error': '请输入有效的邮箱'}), 400

    user = User.query.filter_by(email=email).first()
    if not user:
        # 出于隐私保护，即使邮箱不存在也返回成功，避免被探测账号是否存在
        return jsonify({'success': True, 'message': '如果该邮箱已注册，重置邮件已发送'})

    # 失效旧的重置令牌
    PasswordResetToken.query.filter_by(user_id=user.id, used=False).update({'used': True})

    token = generate_token()
    expires = datetime.now(timezone.utc) + timedelta(minutes=30)
    db.session.add(PasswordResetToken(user_id=user.id, token=token, expires_at=expires, used=False))
    db.session.commit()

    ok, msg, reset_link = send_reset_email(user.email, token, site_url=site_url)
    if not ok:
        return jsonify({'error': msg}), 500

    resp = {'success': True, 'message': msg}
    # SMTP 未配置时，返回重置链接给前端展示（开发/自部署环境降级方案）
    if not SMTP_PASSWORD:
        resp['reset_link'] = reset_link
        resp['dev_mode'] = True
    return jsonify(resp)

@app.route('/api/auth/reset-password', methods=['POST'])
def reset_password():
    """用户凭重置令牌设置新密码。"""
    data = request.json or {}
    token = (data.get('token', '') or '').strip()
    new_password = (data.get('new_password', '') or '').strip()
    if not token or not new_password:
        return jsonify({'error': '令牌或新密码不能为空'}), 400
    if len(new_password) < 4:
        return jsonify({'error': '新密码至少4个字符'}), 400

    prt = PasswordResetToken.query.filter_by(token=token).first()
    if not prt or prt.used:
        return jsonify({'error': '重置链接无效或已使用'}), 400
    exp = prt.expires_at
    if exp.tzinfo is None:
        exp = exp.replace(tzinfo=timezone.utc)
    if exp < datetime.now(timezone.utc):
        return jsonify({'error': '重置链接已过期，请重新申请'}), 400

    user = User.query.get(prt.user_id)
    if not user:
        return jsonify({'error': '用户不存在'}), 404

    user.password_hash = generate_password_hash(new_password)
    prt.used = True
    db.session.commit()
    return jsonify({'success': True, 'message': '密码已重置，请使用新密码登录'})

@app.route('/api/auth/verify-reset-token', methods=['POST'])
def verify_reset_token():
    """校验重置令牌是否有效（用于前端跳转后预检）。"""
    data = request.json or {}
    token = (data.get('token', '') or '').strip()
    if not token:
        return jsonify({'valid': False}), 400
    prt = PasswordResetToken.query.filter_by(token=token).first()
    if not prt or prt.used:
        return jsonify({'valid': False}), 200
    exp = prt.expires_at
    if exp.tzinfo is None:
        exp = exp.replace(tzinfo=timezone.utc)
    if exp < datetime.now(timezone.utc):
        return jsonify({'valid': False}), 200
    return jsonify({'valid': True}), 200

# ==== Books API ====

@app.route('/api/books', methods=['GET'])
@login_required
def list_books():
    books = Book.query.filter_by(user_id=request.current_user_id).order_by(Book.updated_at.desc()).all()
    return jsonify([b.to_dict() for b in books])

@app.route('/api/books/<book_id>', methods=['GET'])
def get_book(book_id):
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404
    return jsonify(book.to_dict())

@app.route('/api/books', methods=['POST'])
@login_required
def create_book():
    data = request.json
    # 总卷数：长篇默认10，短篇默认1。卷数不设上限，由用户自行决定（不钳制）
    book_type = data.get('book_type', 'novel')
    default_vols = 1 if book_type == 'short_story' else 10
    total_volumes = data.get('total_volumes') or default_vols
    # 卷数校验：仅校验下限≥1，不设上限（用户填多少就是多少）
    try:
        total_volumes = int(total_volumes)
        total_volumes = max(1, total_volumes)
    except (ValueError, TypeError):
        total_volumes = default_vols
    # 风格流派：JSON 数组，最多3种
    novel_styles = data.get('novel_styles', [])
    if isinstance(novel_styles, list):
        novel_styles = novel_styles[:3]
    else:
        novel_styles = []
    book = Book(
        user_id=request.current_user_id,
        title=data.get('title', '新书'),
        author=data.get('author', ''),
        genre=data.get('genre', 'other'),
        book_type=book_type,
        synopsis=data.get('synopsis', ''),
        template_id=data.get('template_id', ''),
        target_words=data.get('target_words', 0),
        total_volumes=total_volumes,
        novel_styles=json.dumps(novel_styles, ensure_ascii=False),
        status='draft'
    )
    db.session.add(book)
    db.session.flush()

    if data.get('template_id'):
        template = Template.query.get(data['template_id'])
        if template and template.structure_json:
            structure = json.loads(template.structure_json)
            for i, s in enumerate(structure):
                ch = Chapter(
                    book_id=book.id, title=s.get('title', f'章节{i+1}'),
                    order_index=i, is_volume=s.get('is_volume', False),
                    parent_id=s.get('parent_id', '')
                )
                db.session.add(ch)
    db.session.commit()
    return jsonify(book.to_dict()), 201

@app.route('/api/books/<book_id>', methods=['PUT'])
def update_book(book_id):
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404
    data = request.json
    for field in ['title', 'author', 'genre', 'book_type', 'synopsis', 'status', 'target_words', 'metadata_json']:
        if field in data:
            if field == 'metadata_json' and isinstance(data[field], dict):
                setattr(book, field, json.dumps(data[field], ensure_ascii=False))
            else:
                setattr(book, field, data[field])
    # 总卷数 + 风格流派同步
    if 'total_volumes' in data:
        try:
            tv = int(data['total_volumes'])
            tv = max(1, tv)  # 仅校验下限≥1，不设上限
            book.total_volumes = tv
        except (ValueError, TypeError):
            pass
    if 'novel_styles' in data:
        ns = data['novel_styles']
        if isinstance(ns, list):
            ns = ns[:3]
        else:
            ns = []
        book.novel_styles = json.dumps(ns, ensure_ascii=False)
    # 同步到 BookBible（创作时从 bible 读取注入）
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if bb:
        if hasattr(book, 'total_volumes') and book.total_volumes:
            bb.total_volumes = book.total_volumes
        if hasattr(book, 'novel_styles'):
            bb.novel_styles = book.novel_styles
    db.session.commit()
    return jsonify(book.to_dict())

@app.route('/api/books/<book_id>', methods=['DELETE'])
def delete_book(book_id):
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404
    # 手动删除所有关联记录（兼容 PostgreSQL 外键约束 + SQLite）
    ChapterVersion.query.filter(ChapterVersion.chapter_id.in_(
        db.session.query(Chapter.id).filter_by(book_id=book_id)
    )).delete(synchronize_session=False)
    Chapter.query.filter_by(book_id=book_id).delete(synchronize_session=False)
    Character.query.filter_by(book_id=book_id).delete(synchronize_session=False)
    Outline.query.filter_by(book_id=book_id).delete(synchronize_session=False)
    DailyStats.query.filter_by(book_id=book_id).delete(synchronize_session=False)
    AISession.query.filter_by(book_id=book_id).delete(synchronize_session=False)
    StageContent.query.filter_by(book_id=book_id).delete(synchronize_session=False)
    BookBible.query.filter_by(book_id=book_id).delete(synchronize_session=False)
    DynamicMemory.query.filter_by(book_id=book_id).delete(synchronize_session=False)
    DynamicReport.query.filter_by(book_id=book_id).delete(synchronize_session=False)
    db.session.delete(book)
    db.session.commit()
    return jsonify({'success': True})

# ==== Chapters API ====

@app.route('/api/books/<book_id>/chapters', methods=['GET'])
def list_chapters(book_id):
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404
    chapters = Chapter.query.filter_by(book_id=book_id).order_by(Chapter.order_index).all()
    # 自动修复孤儿数据：清空指向不存在卷的 parent_id（历史删除卷未清空子章节导致）
    vol_ids = {c.id for c in chapters if c.is_volume}
    repaired = False
    for c in chapters:
        if not c.is_volume and c.parent_id and c.parent_id not in vol_ids:
            c.parent_id = ''
            repaired = True
    if repaired:
        db.session.commit()
    return jsonify([c.to_dict(include_content=False) for c in chapters])

@app.route('/api/books/<book_id>/chapters/<chapter_id>', methods=['GET'])
def get_chapter(book_id, chapter_id):
    ch = Chapter.query.filter_by(id=chapter_id, book_id=book_id).first()
    if not ch:
        return jsonify({'error': 'Chapter not found'}), 404
    return jsonify(ch.to_dict(include_content=True))

@app.route('/api/books/<book_id>/chapters', methods=['POST'])
def create_chapter(book_id):
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404
    data = request.json
    max_order = db.session.query(db.func.max(Chapter.order_index)).filter_by(book_id=book_id).scalar() or -1
    ch = Chapter(
        book_id=book_id, title=data.get('title', '新章节'),
        content=data.get('content', ''), order_index=data.get('order_index', max_order + 1),
        is_volume=data.get('is_volume', False), parent_id=data.get('parent_id', ''),
        notes=data.get('notes', '')
    )
    if ch.content:
        ch.word_count = count_words(ch.content)
    db.session.add(ch)
    db.session.flush()
    update_book_stats(book_id)

    # 若新章节标题含章节号（第N章/第N章等），按章节号自动重排顺序（不改动卷归入）
    if not ch.is_volume and parse_chapter_number(ch.title or '') is not None:
        try:
            resort_chapters_by_title(book_id, rebin_volumes=False)
        except Exception:
            pass  # 重排失败不影响章节创建

    # 自动检测是否需要生成动态报告（每5章触发）
    auto_report = None
    if not ch.is_volume and ch.content:
        try:
            result = _check_and_auto_generate_report(book_id)
            if result and 'report' in result:
                auto_report = result['report']
        except Exception:
            pass  # 自动生成失败不影响章节创建
        # 【P0-3】每 20 章自动触发防遗忘检查（daemon 线程，不阻塞）
        try:
            _maybe_auto_trigger_anti_forget_check(book_id)
        except Exception:
            pass
        # 【P1-2】章节落库后统一钩子：事件抽取 + 伏笔本章清单 + 实体注册
        hook_meta = None
        try:
            hook_meta = _after_chapter_persisted(book_id, ch)
            db.session.commit()
        except Exception:
            db.session.rollback()

    resp = ch.to_dict(include_content=True)
    if auto_report:
        resp['auto_report'] = auto_report
    if hook_meta:
        resp['event_log'] = hook_meta
    return jsonify(resp), 201

@app.route('/api/books/<book_id>/chapters/<chapter_id>', methods=['PUT'])
def update_chapter(book_id, chapter_id):
    ch = Chapter.query.filter_by(id=chapter_id, book_id=book_id).first()
    if not ch:
        return jsonify({'error': 'Chapter not found'}), 404
    old_content = ch.content
    data = request.json

    has_content_change = 'content' in data and data['content'] != old_content
    if has_content_change:
        versions = ChapterVersion.query.filter_by(chapter_id=chapter_id).order_by(ChapterVersion.version_num.desc()).all()
        new_ver = (versions[0].version_num + 1) if versions else 1
        ver = ChapterVersion(chapter_id=chapter_id, content=old_content, version_num=new_ver)
        db.session.add(ver)

    for field in ['title', 'content', 'order_index', 'status', 'is_volume', 'parent_id', 'notes']:
        if field in data:
            setattr(ch, field, data[field])
    if 'content' in data:
        ch.word_count = count_words(data['content'])
    ch.updated_at = datetime.now(timezone.utc)
    db.session.flush()
    update_book_stats(book_id)

    # 自动检测是否需要生成动态报告（每5章触发，仅内容变更时）
    auto_report = None
    if has_content_change and not ch.is_volume and ch.content:
        try:
            result = _check_and_auto_generate_report(book_id)
            if result and 'report' in result:
                auto_report = result['report']
        except Exception:
            pass
        # 【P0-3】每 20 章自动触发防遗忘检查（daemon 线程，不阻塞）
        try:
            _maybe_auto_trigger_anti_forget_check(book_id)
        except Exception:
            pass
        # 【P1-2】内容变更落库后统一钩子：事件抽取 + 伏笔本章清单 + 实体注册
        hook_meta = None
        if has_content_change:
            try:
                hook_meta = _after_chapter_persisted(book_id, ch)
                db.session.commit()
            except Exception:
                db.session.rollback()

    resp = ch.to_dict(include_content=True)
    if auto_report:
        resp['auto_report'] = auto_report
    if hook_meta:
        resp['event_log'] = hook_meta
    return jsonify(resp)

@app.route('/api/books/<book_id>/chapters/<chapter_id>', methods=['DELETE'])
def delete_chapter(book_id, chapter_id):
    ch = Chapter.query.filter_by(id=chapter_id, book_id=book_id).first()
    if not ch:
        return jsonify({'error': 'Chapter not found'}), 404
    # 若删除的是卷，清空其下章节的 parent_id，避免章节变孤儿（指向已删除卷）而不可见
    if ch.is_volume:
        Chapter.query.filter_by(book_id=book_id, parent_id=chapter_id, is_volume=False).update({'parent_id': ''})
    # 先删除章节版本（兼容 PostgreSQL 外键约束）
    ChapterVersion.query.filter_by(chapter_id=chapter_id).delete(synchronize_session=False)
    db.session.delete(ch)
    db.session.flush()
    update_book_stats(book_id)
    return jsonify({'success': True})

@app.route('/api/books/<book_id>/chapters/reorder', methods=['POST'])
def reorder_chapters(book_id):
    data = request.json
    order = data.get('order', [])
    for item in order:
        Chapter.query.filter_by(id=item['id'], book_id=book_id).update({'order_index': item['order_index']})
    db.session.commit()
    return jsonify({'success': True})

@app.route('/api/books/<book_id>/chapters/rebin-volumes', methods=['POST'])
@login_required
def rebin_volumes(book_id):
    """手动触发按 50 章/卷重新归入卷：先清空所有章节 parent_id，删除现有卷，再按章节号排序重新分卷。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404
    if book.user_id != request.current_user_id:
        return jsonify({'error': '无权操作该作品'}), 403
    try:
        # 清空所有非卷章节的 parent_id
        Chapter.query.filter_by(book_id=book_id, is_volume=False).update({'parent_id': ''})
        # 删除所有现有卷
        old_vols = Chapter.query.filter_by(book_id=book_id, is_volume=True).all()
        for v in old_vols:
            db.session.delete(v)
        db.session.flush()
        # 重新排序 + 分卷
        count = resort_chapters_by_title(book_id, rebin_volumes=True)
        update_book_stats(book_id)
        vols = Chapter.query.filter_by(book_id=book_id, is_volume=True).count()
        return jsonify({'success': True, 'chapters': count, 'volumes': vols})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

# ==== Chapter Versions ====

@app.route('/api/books/<book_id>/chapters/<chapter_id>/versions', methods=['GET'])
def list_chapter_versions(book_id, chapter_id):
    versions = ChapterVersion.query.filter_by(chapter_id=chapter_id).order_by(ChapterVersion.version_num.desc()).all()
    return jsonify([v.to_dict() for v in versions])

@app.route('/api/books/<book_id>/chapters/<chapter_id>/versions/<version_id>/restore', methods=['POST'])
def restore_chapter_version(book_id, chapter_id, version_id):
    ch = Chapter.query.filter_by(id=chapter_id, book_id=book_id).first()
    ver = ChapterVersion.query.filter_by(id=version_id, chapter_id=chapter_id).first()
    if not ch or not ver:
        return jsonify({'error': 'Not found'}), 404
    ch.content = ver.content
    ch.word_count = count_words(ver.content)
    ch.updated_at = datetime.now(timezone.utc)
    db.session.flush()
    update_book_stats(book_id)
    return jsonify(ch.to_dict(include_content=True))

# ==== Characters API ====

@app.route('/api/books/<book_id>/characters', methods=['GET'])
def list_characters(book_id):
    chars = Character.query.filter_by(book_id=book_id).all()
    return jsonify([c.to_dict() for c in chars])

@app.route('/api/books/<book_id>/characters', methods=['POST'])
def create_character(book_id):
    data = request.json
    char = Character(
        book_id=book_id, name=(data.get('name', '新角色') or '新角色')[:50],
        role=(data.get('role', 'supporting') or 'supporting')[:50], description=data.get('description', ''),
        appearance=data.get('appearance', ''), personality=data.get('personality', ''),
        background=data.get('background', '')
    )
    if 'relationships' in data:
        char.relationships_json = json.dumps(data['relationships'], ensure_ascii=False)
    db.session.add(char)
    db.session.commit()
    return jsonify(char.to_dict()), 201

@app.route('/api/books/<book_id>/characters/<char_id>', methods=['PUT'])
def update_character(book_id, char_id):
    char = Character.query.filter_by(id=char_id, book_id=book_id).first()
    if not char:
        return jsonify({'error': 'Character not found'}), 404
    data = request.json
    for field in ['name', 'role', 'description', 'appearance', 'personality', 'background']:
        if field in data:
            setattr(char, field, data[field])
    if 'name' in data:
        char.name = (data['name'] or '')[:50]
    if 'role' in data:
        char.role = (data['role'] or 'supporting')[:50]
    if 'relationships' in data:
        char.relationships_json = json.dumps(data['relationships'], ensure_ascii=False)
    db.session.commit()
    return jsonify(char.to_dict())

@app.route('/api/books/<book_id>/characters/<char_id>', methods=['DELETE'])
def delete_character(book_id, char_id):
    char = Character.query.filter_by(id=char_id, book_id=book_id).first()
    if not char:
        return jsonify({'error': 'Character not found'}), 404
    db.session.delete(char)
    db.session.commit()
    return jsonify({'success': True})

# ==== Outlines API ====

@app.route('/api/books/<book_id>/outlines', methods=['GET'])
def list_outlines(book_id):
    outlines = Outline.query.filter_by(book_id=book_id).order_by(Outline.order_index).all()
    tree = build_outline_tree(outlines)
    return jsonify({'flat': [o.to_dict() for o in outlines], 'tree': tree})

@app.route('/api/books/<book_id>/outlines', methods=['POST'])
def create_outline(book_id):
    data = request.json
    max_order = db.session.query(db.func.max(Outline.order_index)).filter_by(book_id=book_id).scalar() or -1
    outline = Outline(
        book_id=book_id, title=data.get('title', '新节点'),
        content=data.get('content', ''), order_index=data.get('order_index', max_order + 1),
        level=data.get('level', 0), parent_id=data.get('parent_id', '')
    )
    db.session.add(outline)
    db.session.commit()
    outlines = Outline.query.filter_by(book_id=book_id).order_by(Outline.order_index).all()
    return jsonify({'item': outline.to_dict(), 'tree': build_outline_tree(outlines)}), 201

@app.route('/api/books/<book_id>/outlines/<outline_id>', methods=['PUT'])
def update_outline(book_id, outline_id):
    outline = Outline.query.filter_by(id=outline_id, book_id=book_id).first()
    if not outline:
        return jsonify({'error': 'Outline not found'}), 404
    data = request.json
    for field in ['title', 'content', 'order_index', 'level', 'parent_id']:
        if field in data:
            setattr(outline, field, data[field])
    db.session.commit()
    outlines = Outline.query.filter_by(book_id=book_id).order_by(Outline.order_index).all()
    return jsonify({'item': outline.to_dict(), 'tree': build_outline_tree(outlines)})

@app.route('/api/books/<book_id>/outlines/<outline_id>', methods=['DELETE'])
def delete_outline(book_id, outline_id):
    outline = Outline.query.filter_by(id=outline_id, book_id=book_id).first()
    if not outline:
        return jsonify({'error': 'Outline not found'}), 404
    db.session.delete(outline)
    db.session.commit()
    outlines = Outline.query.filter_by(book_id=book_id).order_by(Outline.order_index).all()
    return jsonify({'tree': build_outline_tree(outlines)})

# ==== Stats API ====

@app.route('/api/books/<book_id>/stats', methods=['GET'])
def get_book_stats(book_id):
    stats = DailyStats.query.filter_by(book_id=book_id).order_by(DailyStats.date).all()
    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    return jsonify({
        'daily': [s.to_dict() for s in stats],
        'chapters': [{'title': c.title, 'word_count': c.word_count, 'date': c.created_at.isoformat() if c.created_at else None} for c in chapters]
    })

@app.route('/api/books/<book_id>/stats', methods=['POST'])
def add_daily_stats(book_id):
    data = request.json
    today = datetime.now(timezone.utc).date()
    stat = DailyStats.query.filter_by(book_id=book_id, date=today).first()
    if stat:
        stat.words_written += data.get('words_written', 0)
        stat.time_spent_minutes += data.get('time_spent_minutes', 0)
        stat.chapters_completed += data.get('chapters_completed', 0)
        prev_words = stat.words_written
    else:
        stat = DailyStats(
            book_id=book_id, date=today,
            words_written=data.get('words_written', 0),
            time_spent_minutes=data.get('time_spent_minutes', 0),
            chapters_completed=data.get('chapters_completed', 0)
        )
        db.session.add(stat)
        prev_words = 0
    db.session.commit()
    return jsonify(stat.to_dict())

# ==== Templates API ====

# /api/health 已迁移到 blueprints/health_bp.py（Blueprint 示范）

@app.route('/api/templates', methods=['GET'])
def list_templates():
    templates = Template.query.order_by(Template.is_builtin.desc(), Template.created_at.desc()).all()
    return jsonify([t.to_dict() for t in templates])

@app.route('/api/templates', methods=['POST'])
def create_template():
    data = request.json
    template = Template(
        name=data.get('name', '自定义模板'), description=data.get('description', ''),
        genre=data.get('genre', 'other'), book_type=data.get('book_type', 'novel'),
        structure_json=json.dumps(data.get('structure', []), ensure_ascii=False),
        prompts_json=json.dumps(data.get('prompts', {}), ensure_ascii=False)
    )
    db.session.add(template)
    db.session.commit()
    return jsonify(template.to_dict()), 201

# ==== AI API ====
# 注：/api/ai/config GET/PUT 已迁移至 blueprints/ai_config_bp.py（多配置支持）

def _do_fetch_models(base_url, api_key, model=None):
    """实际拉取模型列表的内部函数，供多个接口复用。

    【智谱 GLM 404 修复】不再无条件补 /v1：走 provider 感知归一化。
    智谱 GLM 的 /models 也在 v4 下（/api/paas/v4/models），若按老逻辑补 /v1 → /v4/v1/models → 404。
    """
    from llm_gateway import _normalize_llm_base_url
    import requests as req
    base = _normalize_llm_base_url(base_url, model)
    resp = req.get(
        f"{base}/models",
        headers=build_auth_headers(api_key, content_type=False),
        timeout=15
    )
    if resp.status_code != 200:
        return None, f'请求失败 (HTTP {resp.status_code})：{resp.text[:200]}', 400
    result = resp.json()
    models = []
    for m in result.get('data', []):
        mid = m.get('id', '')
        if mid:
            models.append({
                'id': mid,
                'owned_by': m.get('owned_by', ''),
            })
    models.sort(key=lambda x: x['id'])
    return models, None, 200

def _do_test_connection(base_url, api_key, model):
    """实际测试连接的内部函数，供多个接口复用。

    【智谱 GLM 404 修复】删除老代码"非 /v1 结尾就补 /v1"。
    用户点"测试连接"按钮看到的 "HTTP 404 path=/v4/v1/chat/completions" 就是这里产生的。
    """
    from llm_gateway import _normalize_llm_base_url
    import requests as req
    base = _normalize_llm_base_url(base_url, model)
    resp = req.post(
        f"{base}/chat/completions",
        headers=build_auth_headers(api_key),
        json={
            'model': model,
            'messages': [{'role': 'user', 'content': '你好，请回复"连接成功"四个字。'}],
            'max_tokens': 20,
            'temperature': 0.1,
            'stream': False
        },
        timeout=30
    )
    if resp.status_code != 200:
        try:
            err_data = resp.json()
            err_msg = err_data.get('error', {}).get('message', '') or str(err_data)
        except Exception:
            err_msg = resp.text[:300]
        return None, f'连接失败 (HTTP {resp.status_code})：{err_msg}', 400
    result = resp.json()
    if 'choices' in result and len(result['choices']) > 0:
        reply = result['choices'][0]['message']['content']
        usage = result.get('usage', {})
        return {
            'success': True,
            'reply': reply,
            'model': result.get('model', model),
            'usage': usage
        }, None, 200
    else:
        return None, f'返回数据异常：{str(result)[:200]}', 400

@app.route('/api/ai/models', methods=['POST'])
def fetch_ai_models():
    """根据 base_url 和 api_key 拉取可用模型列表（OpenAI 兼容的 /v1/models 接口）"""
    data = request.json or {}
    base_url = (data.get('base_url') or '').strip()
    api_key = (data.get('api_key') or '').strip()
    model = (data.get('model') or '').strip()

    # 如果 api_key 是掩码或为空，尝试使用已保存的配置
    if api_key == '***' or not api_key:
        cfg = AIConfig.get_active()
        if cfg and cfg.api_key:
            api_key = cfg.api_key
            if not base_url:
                base_url = cfg.base_url or ''
            if not model:
                model = cfg.model or ''
        else:
            return jsonify({'error': '请先填写 API Key 或保存配置'}), 400

    if not base_url:
        return jsonify({'error': '请填写 API 地址'}), 400

    try:
        models, err, code = _do_fetch_models(base_url, api_key, model)
        if err:
            return jsonify({'error': err}), code
        return jsonify({'models': models})
    except requests.exceptions.ConnectionError:
        return jsonify({'error': '无法连接到服务器，请检查 API 地址是否正确'}), 400
    except requests.exceptions.Timeout:
        return jsonify({'error': '请求超时，请稍后重试'}), 400
    except Exception as e:
        return jsonify({'error': f'拉取模型失败：{str(e)}'}), 500

@app.route('/api/ai/test', methods=['POST'])
def test_ai_connection():
    """测试 AI 连接：发送一条简单消息验证配置是否可用"""
    data = request.json or {}
    base_url = (data.get('base_url') or '').strip()
    api_key = (data.get('api_key') or '').strip()
    model = (data.get('model') or '').strip()

    # 如果 api_key 是掩码或为空，尝试使用已保存的配置
    if api_key == '***' or not api_key:
        cfg = AIConfig.get_active()
        if cfg and cfg.api_key:
            api_key = cfg.api_key
            if not base_url:
                base_url = cfg.base_url or ''
            if not model:
                model = cfg.model or ''
        else:
            return jsonify({'error': '请先填写 API Key 或保存配置'}), 400

    if not base_url or not api_key or not model:
        return jsonify({'error': '请填写完整的 API 地址、API Key 和模型名称'}), 400

    try:
        result, err, code = _do_test_connection(base_url, api_key, model)
        if err:
            return jsonify({'error': err}), code
        return jsonify(result)
    except requests.exceptions.ConnectionError:
        return jsonify({'error': '无法连接到服务器，请检查 API 地址'}), 400
    except requests.exceptions.Timeout:
        return jsonify({'error': '请求超时，请检查网络或稍后重试'}), 400
    except Exception as e:
        return jsonify({'error': f'测试失败：{str(e)}'}), 500

@app.route('/api/ai/chat', methods=['POST'])
def ai_chat():
    data = request.json
    messages = data.get('messages', [])
    if not messages:
        return jsonify({'error': 'No messages'}), 400

    cfg = AIConfig.get_active()
    if not cfg or not cfg.api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    try:
        import requests as req
        base = cfg.base_url.rstrip('/')
        if not base.endswith('/v1'):
            base += '/v1'
        resp = req.post(
            f"{base}/chat/completions",
            headers=build_auth_headers(cfg.api_key),
            json={
                'model': cfg.model,
                'messages': messages,
                'temperature': cfg.temperature,
                'max_tokens': min(cfg.max_tokens, get_output_limit(base, cfg.model) or cfg.max_tokens),
                'stream': False
            },
            timeout=120
        )
        result = resp.json()
        if 'choices' in result and len(result['choices']) > 0:
            return jsonify({
                'content': result['choices'][0]['message']['content'],
                'usage': result.get('usage', {})
            })
        else:
            return jsonify({'error': str(result)}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ai/chat/stream', methods=['POST'])
def ai_chat_stream():
    data = request.json
    messages = data.get('messages', [])

    cfg = AIConfig.get_active()
    if not cfg or not cfg.api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    def generate():
        try:
            import requests as req
            base = cfg.base_url.rstrip('/')
            if not base.endswith('/v1'):
                base += '/v1'
            resp = req.post(
                f"{base}/chat/completions",
                headers=build_auth_headers(cfg.api_key),
                json={
                    'model': cfg.model,
                    'messages': messages,
                    'temperature': cfg.temperature,
                    'max_tokens': min(cfg.max_tokens, get_output_limit(base, cfg.model) or cfg.max_tokens),
                    'stream': True
                },
                stream=True,
                timeout=120
            )
            for line in resp.iter_lines():
                if line:
                    line = line.decode('utf-8')
                    if line.startswith('data: '):
                        chunk = line[6:]
                        if chunk == '[DONE]':
                            yield 'data: [DONE]\n\n'
                            break
                        yield f'data: {chunk}\n\n'
        except Exception as e:
            yield f'data: {{"error": "{str(e)}"}}\n\n'

    return app.response_class(generate(), mimetype='text/event-stream')

# ==== AI Sessions ====

@app.route('/api/ai/sessions', methods=['GET'])
def list_ai_sessions():
    book_id = request.args.get('book_id')
    scope = request.args.get('scope')
    query = AISession.query
    if book_id:
        query = query.filter_by(book_id=book_id)
    if scope:
        query = query.filter_by(scope=scope)
    sessions = query.order_by(AISession.updated_at.desc()).all()
    return jsonify([s.to_dict() for s in sessions])

@app.route('/api/ai/sessions', methods=['POST'])
def create_ai_session():
    data = request.json
    book_id = data.get('book_id', '')
    scope = data.get('scope', 'general')
    scope_id = data.get('scope_id', '')
    # upsert：同 book+scope+scope_id 的会话复用，避免维度Modal重复创建孤儿会话
    existing = AISession.query.filter_by(
        book_id=book_id, scope=scope, scope_id=scope_id
    ).order_by(AISession.updated_at.desc()).first() if (book_id and scope and scope_id) else None
    if existing:
        return jsonify(existing.to_dict()), 200
    session = AISession(
        book_id=book_id,
        scope=scope,
        scope_id=scope_id,
        title=data.get('title', '新对话')
    )
    db.session.add(session)
    db.session.commit()
    return jsonify(session.to_dict()), 201

@app.route('/api/ai/sessions/<session_id>', methods=['PUT'])
def update_ai_session(session_id):
    session = AISession.query.get(session_id)
    if not session:
        return jsonify({'error': 'Session not found'}), 404
    data = request.json
    if 'messages' in data:
        session.messages_json = json.dumps(data['messages'], ensure_ascii=False)
    if 'title' in data:
        session.title = data['title']
    session.updated_at = datetime.now(timezone.utc)
    db.session.commit()
    return jsonify(session.to_dict())

@app.route('/api/ai/sessions/<session_id>', methods=['DELETE'])
def delete_ai_session(session_id):
    session = AISession.query.get(session_id)
    if not session:
        return jsonify({'error': 'Session not found'}), 404
    db.session.delete(session)
    db.session.commit()
    return jsonify({'success': True})

# ==== 实体注册表 API（P2：跨维度重命名/合并） ====

@app.route('/api/books/<book_id>/entities', methods=['GET'])
@login_required
def list_entities(book_id):
    """抽取作品全部实体（角色/势力/地点/物品/技能）。
    升级：每次 GET 都会实时扫描：
    - Bible 十个维度文本 + JSON 字段（含智驾采纳落地写入的 JSON 人物卡 / 冒号分点实体）
    - 最近 10 章标题/正文
    抽取后合并写回 bb.entity_registry_json，确保增量事件识别能共享同一套实体库。"""
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        return jsonify({'characters': [], 'factions': [], 'locations': [], 'items': [], 'skills': []})
    from entity_registry import extract_and_save_registry
    # 把最近 10 章也作为扫描源（正文/标题里出现的角色/地点自动入表）
    recent_chapters = []
    try:
        recent_chapters = (
            Chapter.query.filter_by(book_id=book_id, is_volume=False)
            .order_by(Chapter.order_index.desc())
            .limit(10)
            .all()
        ) or []
    except Exception:
        recent_chapters = []
    entities = extract_and_save_registry(bb, chapters_query=recent_chapters)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()
    return jsonify(entities)

@app.route('/api/books/<book_id>/entities/rename', methods=['POST'])
@login_required
def rename_entity_api(book_id):
    """跨维度重命名实体。body: {old_name, new_name, entity_type}"""
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        return jsonify({'error': 'BookBible not found'}), 404
    data = request.json or {}
    old_name = (data.get('old_name') or '').strip()
    new_name = (data.get('new_name') or '').strip()
    entity_type = data.get('entity_type', 'character')
    if not old_name or not new_name:
        return jsonify({'error': 'old_name 和 new_name 不能为空'}), 400
    from entity_registry import rename_entity
    chapters_q = Chapter.query.filter_by(book_id=book_id, is_volume=False).all()
    result = rename_entity(bb, chapters_q, old_name, new_name, entity_type)
    db.session.commit()
    return jsonify(result)

@app.route('/api/books/<book_id>/entities/merge', methods=['POST'])
@login_required
def merge_entities_api(book_id):
    """合并实体。body: {main_name, alias_names:[...], entity_type}"""
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        return jsonify({'error': 'BookBible not found'}), 404
    data = request.json or {}
    main_name = (data.get('main_name') or '').strip()
    alias_names = data.get('alias_names', []) or []
    entity_type = data.get('entity_type', 'character')
    if not main_name or not alias_names:
        return jsonify({'error': 'main_name 和 alias_names 不能为空'}), 400
    from entity_registry import merge_entities
    chapters_q = Chapter.query.filter_by(book_id=book_id, is_volume=False).all()
    result = merge_entities(bb, chapters_q, main_name, alias_names, entity_type)
    db.session.commit()
    return jsonify(result)

# ==== Export API ====

@app.route('/api/books/<book_id>/export', methods=['GET'])
def export_book(book_id):
    fmt = request.args.get('format', 'txt')
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()

    if fmt == 'txt':
        output = f'《{book.title}》\n作者：{book.author}\n\n简介：{book.synopsis}\n\n{"="*50}\n\n'
        for ch in chapters:
            output += f'\n## {ch.title}\n\n{ch.content}\n\n{"="*50}\n'
        buf = BytesIO(output.encode('utf-8'))
        return send_file(buf, mimetype='text/plain', as_attachment=True, download_name=f'{book.title}.txt')

    elif fmt == 'docx':
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'SimSun'
            font.size = Pt(12)

            title_para = doc.add_heading(book.title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f'作者：{book.author}')
            if book.synopsis:
                doc.add_paragraph(f'简介：{book.synopsis}')
            doc.add_page_break()

            for ch in chapters:
                doc.add_heading(ch.title, level=1)
                for para_text in ch.content.split('\n'):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
                doc.add_page_break()

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                             as_attachment=True, download_name=f'{book.title}.docx')
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    elif fmt == 'epub':
        try:
            from ebooklib import epub

            epub_book = epub.EpubBook()
            epub_book.set_identifier(book.id)
            epub_book.set_title(book.title)
            epub_book.set_language('zh')
            epub_book.add_author(book.author or '佚名')

            epub_chapters = []
            for ch in chapters:
                c = epub.EpubHtml(title=ch.title, file_name=f'ch_{ch.order_index}.xhtml', lang='zh')
                c.content = f'<h1>{ch.title}</h1>' + ''.join(f'<p>{p}</p>' for p in ch.content.split('\n') if p.strip())
                epub_book.add_item(c)
                epub_chapters.append(c)

            epub_book.toc = epub_chapters
            epub_book.add_item(epub.EpubNcx())
            epub_book.add_item(epub.EpubNav())

            style = 'BODY {color: black;}'
            nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
            epub_book.add_item(nav_css)

            epub_book.spine = ['nav'] + epub_chapters

            buf = BytesIO()
            epub.write_epub(buf, epub_book, {})
            buf.seek(0)
            return send_file(buf, mimetype='application/epub+zip', as_attachment=True, download_name=f'{book.title}.epub')
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    return jsonify({'error': 'Unsupported format'}), 400

# ==== Cover Upload ====

@app.route('/api/books/<book_id>/cover', methods=['POST'])
def upload_cover(book_id):
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No filename'}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in ('png', 'jpg', 'jpeg', 'gif', 'webp'):
        return jsonify({'error': 'Invalid image format'}), 400

    filename = f'{book_id}.{ext}'
    filepath = COVERS_DIR / filename
    file.save(filepath)
    book.cover_path = f'/api/covers/{filename}'
    db.session.commit()
    return jsonify({'cover_path': book.cover_path})

@app.route('/api/covers/<filename>')
def serve_cover(filename):
    return send_file(COVERS_DIR / filename)

# ==== Preferences ====

@app.route('/api/preferences', methods=['GET'])
def get_preferences():
    keys = request.args.get('keys', '')
    if keys:
        result = {}
        for k in keys.split(','):
            result[k] = AppPreference.get(k, '')
        return jsonify(result)
    prefs = AppPreference.query.all()
    return jsonify({p.key: p.value for p in prefs})

@app.route('/api/preferences', methods=['PUT'])
def set_preferences():
    data = request.json
    for k, v in data.items():
        AppPreference.set(k, str(v))
    return jsonify({'success': True})

# ==== Import/Export ZIP ====

@app.route('/api/books/<book_id>/export-zip', methods=['GET'])
def export_book_zip(book_id):
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    tmpdir = tempfile.mkdtemp()
    try:
        metadata = book.to_dict()
        metadata['chapters'] = [c.to_dict(include_content=True) for c in Chapter.query.filter_by(book_id=book_id).order_by(Chapter.order_index).all()]
        metadata['characters'] = [c.to_dict() for c in Character.query.filter_by(book_id=book_id).all()]
        metadata['outlines'] = [o.to_dict() for o in Outline.query.filter_by(book_id=book_id).order_by(Outline.order_index).all()]

        with open(os.path.join(tmpdir, 'book.json'), 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)

        zippath = os.path.join(tmpdir, f'{book.title}.zip')
        with zipfile.ZipFile(zippath, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(os.path.join(tmpdir, 'book.json'), 'book.json')
            if book.cover_path:
                cover_file = COVERS_DIR / os.path.basename(book.cover_path)
                if cover_file.exists():
                    zf.write(cover_file, f'cover{cover_file.suffix}')

        return send_file(zippath, mimetype='application/zip', as_attachment=True, download_name=f'{book.title}.zip')
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

@app.route('/api/books/<book_id>/export-full', methods=['GET'])
@login_required
def export_book_full(book_id):
    """导出小说的全部维度内容（除图谱外）和所有章节为独立文件，打包成zip下载"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    # 以小说标题命名文件夹，简单清理会影响路径的字符
    book_title = (book.title or '未命名').replace('/', '_').replace('\\', '_').replace(':', '_')

    bible = BookBible.query.filter_by(book_id=book_id).first()
    bible_data = bible.to_dict() if bible else {}

    # 维度配置：(文件名, BookBible字段, 中文标题)
    dimensions = [
        ('构思.md', 'concept', '构思'),
        ('设定.md', 'key_rules', '设定'),
        ('大纲.md', 'plot_design', '大纲'),
        ('世界观.md', 'worldbuilding', '世界观'),
        ('人物.md', 'character_profiles', '人物'),
        ('剧情时间线.md', 'timeline', '剧情时间线'),
        ('伏笔.md', 'foreshadowing', '伏笔'),
        ('地点.md', 'locations', '地点'),
        ('风格指南.md', 'style_guide', '风格指南'),
    ]

    # 章节按 order_index 排序，跳过卷标记
    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()

    buf = BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        # 写入各维度 md 文件
        for filename, field, title in dimensions:
            content = (bible_data.get(field, '') or '').strip()
            body = content if content else '暂无内容'
            md_text = f'# {title}\n\n---\n\n{body}\n'
            zf.writestr(f'{book_title}/{filename}', md_text.encode('utf-8'))

        # 写入章节子文件夹，每章一个 txt
        for idx, ch in enumerate(chapters, start=1):
            ch_title = ch.title or f'第{idx}章'
            ch_filename = f'第{idx}章-{ch_title}.txt'
            ch_content = ch.content if (ch.content and ch.content.strip()) else '暂无内容'
            zf.writestr(f'{book_title}/章节/{ch_filename}', ch_content.encode('utf-8'))

    buf.seek(0)
    return send_file(buf, mimetype='application/zip', as_attachment=True, download_name=f'{book_title}.zip')

@app.route('/api/books/import-zip', methods=['POST'])
def import_book_zip():
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    file = request.files['file']
    tmpdir = tempfile.mkdtemp()
    try:
        zippath = os.path.join(tmpdir, 'import.zip')
        file.save(zippath)
        with zipfile.ZipFile(zippath, 'r') as zf:
            zf.extractall(tmpdir)

        book_json = os.path.join(tmpdir, 'book.json')
        if not os.path.exists(book_json):
            return jsonify({'error': 'Invalid book package: no book.json'}), 400

        with open(book_json, 'r', encoding='utf-8') as f:
            data = json.load(f)

        book = Book(
            title=data.get('title', '导入书籍'), author=data.get('author', ''),
            genre=data.get('genre', 'other'), book_type=data.get('book_type', 'novel'),
            synopsis=data.get('synopsis', ''), status='draft'
        )
        db.session.add(book)
        db.session.flush()

        for ch_data in data.get('chapters', []):
            ch = Chapter(
                book_id=book.id, title=ch_data.get('title', ''),
                content=ch_data.get('content', ''), order_index=ch_data.get('order_index', 0),
                is_volume=ch_data.get('is_volume', False), parent_id=ch_data.get('parent_id', ''),
                word_count=count_words(ch_data.get('content', ''))
            )
            db.session.add(ch)

        for char_data in data.get('characters', []):
            char = Character(
                book_id=book.id, name=(char_data.get('name', '') or '')[:50],
                role=(char_data.get('role', 'supporting') or 'supporting')[:50], description=char_data.get('description', ''),
                appearance=char_data.get('appearance', ''), personality=char_data.get('personality', ''),
                background=char_data.get('background', ''),
                relationships_json=json.dumps(char_data.get('relationships', []), ensure_ascii=False)
            )
            db.session.add(char)

        for o_data in data.get('outlines', []):
            outline = Outline(
                book_id=book.id, title=o_data.get('title', ''),
                content=o_data.get('content', ''), order_index=o_data.get('order_index', 0),
                level=o_data.get('level', 0), parent_id=o_data.get('parent_id', '')
            )
            db.session.add(outline)

        update_book_stats(book.id)
        return jsonify(book.to_dict()), 201
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

def _natural_sort_key(name):
    """文件名自然排序键，支持中文数字章节号（第1章 < 第2章 < ... < 第10章）。
    解析出章节号的按章节号升序排前；无法解析的按文件名字典序排后。"""
    base = os.path.splitext(name)[0]
    n = parse_chapter_number(base)
    if n is not None:
        return (0, n, base)
    return (1, 0, base)

def _strip_leading_title_line(text, title):
    """若正文第一行与章节标题（文件名）相同或高度相似，去除该行，避免标题重复。
    比较前会标准化：去除空格、常见标点、统一阿拉伯/中文数字。
    仅当第一行较短（≤40字）且与标题高度匹配时才去除，避免误删正文。"""
    if not text or not title:
        return text
    lines = text.split('\n')
    if not lines:
        return text
    first = lines[0].strip()
    if not first or len(first) > 40:
        return text

    def _norm(s):
        # 去除空格、常见标点
        s = re.sub(r'[\s·，。：:、\-_—()（）\[\]【】]+', '', s or '')
        # 中文数字转阿拉伯（仅单位级别，便于 第1章/第一章 匹配）
        cn_map = {'零': '0', '一': '1', '二': '2', '两': '2', '三': '3', '四': '4',
                  '五': '5', '六': '6', '七': '7', '八': '8', '九': '9', '十': '10'}
        for cn, ar in cn_map.items():
            s = s.replace(cn, ar)
        return s

    nf, nt = _norm(first), _norm(title)
    if not nf or not nt:
        return text
    # 完全相同 或 一方包含另一方 → 视为重复标题行
    if nf == nt or nt in nf or nf in nt:
        return '\n'.join(lines[1:]).strip()
    return text

def _extract_zip_chapters(zippath):
    """解压 zip，若含多个文本文件，返回 [{'title','content'}, ...] 按文件名自然排序，
    每个内嵌文件=一章节，文件名（去扩展名）作为章节标题。
    若 zip 内只有 ≤1 个文本文件，返回 None（让调用方走单文件拆分）。"""
    items = []  # [(inner_name, decoded_text)]
    try:
        with zipfile.ZipFile(zippath, 'r') as zf:
            for name in zf.namelist():
                if name.endswith(('/', '\\')):
                    continue
                inner_ext = name.rsplit('.', 1)[1].lower() if '.' in name else ''
                if inner_ext not in ('txt', 'md', 'markdown', 'json'):
                    continue
                try:
                    content = zf.read(name)
                    decoded = _detect_and_decode(content)
                    if decoded.strip():
                        items.append((name, decoded))
                except Exception:
                    pass
    except Exception:
        return None
    if len(items) <= 1:
        return None
    # 按文件名自然排序
    items.sort(key=lambda x: _natural_sort_key(x[0]))
    result = []
    for name, text in items:
        ch_title = os.path.splitext(os.path.basename(name))[0][:100]
        body = _strip_leading_title_line(text, ch_title)
        result.append({'title': ch_title, 'content': body})
    return result

def split_into_chapters(text):
    """将纯文本按章节标记拆分为多个章节，支持多种章节标题格式。
    会自动合并"同章双标题"（如 第20章 / 第二十章 紧挨出现，前者为空）产生的空章。"""
    import re

    def _extract(matches, text, strip_prefix=''):
        """从正则匹配列表中提取章节"""
        chapters = []
        for i, m in enumerate(matches):
            title = m.group().strip()
            if strip_prefix:
                title = title.replace(strip_prefix, '').strip()
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            content = text[start:end].strip()
            chapters.append({'title': title[:100] or f'第{i + 1}章', 'content': content})
        return chapters

    # 1. 中文章节标题：第一章、第1章、第001章、第一回 等
    cn_pattern = re.compile(r'^[ \t]*第[零一二三四五六七八九十百千万\d]+[章节回卷][ \t]*.*$', re.MULTILINE)
    cn_matches = list(cn_pattern.finditer(text))
    if len(cn_matches) >= 1:
        return _merge_empty_chapters(_extract(cn_matches, text))

    # 2. Markdown 标题 # xxx 或 ## xxx
    md_pattern = re.compile(r'^#{1,3}[ \t]+.+$', re.MULTILINE)
    md_matches = list(md_pattern.finditer(text))
    if len(md_matches) >= 1:
        return _merge_empty_chapters(_extract(md_matches, text, strip_prefix='#'))

    # 3. 英文 Chapter N / CHAPTER N
    en_pattern = re.compile(r'^[ \t]*[Cc][Hh][Aa][Pp][Tt][Ee][Rr][ \t]*\d+[ \t]*.*$', re.MULTILINE)
    en_matches = list(en_pattern.finditer(text))
    if len(en_matches) >= 1:
        return _merge_empty_chapters(_extract(en_matches, text))

    # 4. 纯数字开头：1. xxx / 1、xxx / 1: xxx（至少2个才拆分，避免误拆）
    num_pattern = re.compile(r'^[ \t]*\d{1,4}[.、:][ \t]*\S.+$', re.MULTILINE)
    num_matches = list(num_pattern.finditer(text))
    if len(num_matches) >= 2:
        return _merge_empty_chapters(_extract(num_matches, text))

    # 5. 【 xxx 】 或 『 xxx 』 格式标题
    bracket_pattern = re.compile(r'^[ \t]*[【『][^】』]+[】』][ \t]*.*$', re.MULTILINE)
    bracket_matches = list(bracket_pattern.finditer(text))
    if len(bracket_matches) >= 2:
        return _merge_empty_chapters(_extract(bracket_matches, text))

    # 无法拆分，作为单个章节
    return None

def _merge_empty_chapters(chapters):
    """合并空章节：处理"同章双标题"导致的空章。
    规则：
    - 若某章 content 为空，且与下一章章节号相同（如 第20章/第二十章），则视为同一章，
      保留下一章（有内容的那个）的标题，删除空章。
    - 若某章 content 为空，且下一章 content 非空，但章节号不同或无法解析，则将下一章
      内容并入当前空章（标题以当前章为主，下一章标题作为副标题保留），避免0字节空章。
    返回合并后的章节列表（至少保留1章）。"""
    if not chapters:
        return chapters
    result = []
    i = 0
    while i < len(chapters):
        cur = chapters[i]
        cur_num = parse_chapter_number(cur.get('title', ''))
        cur_content = cur.get('content', '').strip()
        # 当前章为空，尝试与下一章合并
        if not cur_content and i + 1 < len(chapters):
            nxt = chapters[i + 1]
            nxt_num = parse_chapter_number(nxt.get('title', ''))
            nxt_content = nxt.get('content', '').strip()
            # 情况A：章节号相同（同章双标题，如 第20章/第二十章）→ 删空章，跳到下一章
            if cur_num is not None and nxt_num is not None and cur_num == nxt_num:
                i += 1
                continue
            # 情况B：下一章有内容 → 内容并入当前章，标题保留当前章（下一章作副标题）
            if nxt_content:
                merged_title = cur.get('title', '')
                # 若两标题不同，拼接副标题
                if nxt.get('title', '') and nxt.get('title', '') != merged_title:
                    merged_title = f'{merged_title} / {nxt.get("title", "")}'
                result.append({'title': merged_title[:100], 'content': nxt_content})
                i += 2  # 跳过下一章
                continue
            # 情况C：下一章也为空 → 删当前空章，处理下一章
            i += 1
            continue
        result.append({'title': cur.get('title', '')[:100], 'content': cur_content})
        i += 1
    # 兜底：若合并后为空（全部空章），保留原第一章避免0章
    if not result:
        result = [{'title': chapters[0].get('title', '')[:100], 'content': chapters[0].get('content', '')}]
    return result

@app.route('/api/books/import-files', methods=['POST'])
@login_required
def import_book_files():
    """从多个文本文件导入创建新作品，支持 txt/md/docx/zip
    导入模式自动识别：
    - 多文件模式（≥2个文件）：每个文件 = 一个章节，文件名作为章节标题，按文件名排序
    - 单文件模式（1个文件）：调用 split_into_chapters 拆分为多个章节
    - zip 模式：解压后若含多个文本文件，按多文件模式处理；否则按单文件拆分
    """
    files = request.files.getlist('files')
    if not files or len(files) == 0:
        return jsonify({'error': '未选择文件'}), 400

    title = request.form.get('title', '').strip()
    book_type = request.form.get('book_type', 'novel')
    genre = request.form.get('genre', 'other')

    tmpdir = tempfile.mkdtemp()
    try:
        all_chapters = []
        # 收集有效文件（保存到临时目录，记录原始文件名）
        collected = []  # [(original_name, filepath, ext)]
        for f in files:
            if not f or not f.filename:
                continue
            original_name = f.filename
            ext = original_name.rsplit('.', 1)[1].lower() if '.' in original_name else ''
            if ext not in ('txt', 'md', 'markdown', 'docx', 'zip', 'json'):
                continue
            safe_name = f'{uuid.uuid4()}.{ext}'
            filepath = os.path.join(tmpdir, safe_name)
            f.save(filepath)
            collected.append((original_name, filepath, ext))

        if not collected:
            return jsonify({'error': '未选择有效文件'}), 400

        # 判断导入模式
        # 情况1：单个 zip 文件 → 解压，若含多个文本文件则按多文件模式（每内嵌文件=一章节）
        if len(collected) == 1 and collected[0][2] == 'zip':
            zip_chapters = _extract_zip_chapters(collected[0][1])
            if zip_chapters:
                all_chapters.extend(zip_chapters)
            else:
                # zip 内无法按文件拆分，回退到单文件拆分
                text = extract_text_from_file(collected[0][1], collected[0][0])
                if text.strip():
                    chapters = split_into_chapters(text)
                    if chapters:
                        all_chapters.extend(chapters)
                    else:
                        all_chapters.append({'title': os.path.splitext(collected[0][0])[0][:100], 'content': text})
        # 情况2：多个文件 → 多文件模式（每个文件=一章节，文件名作为标题，按文件名自然排序）
        elif len(collected) >= 2:
            collected.sort(key=lambda x: _natural_sort_key(x[0]))
            for original_name, filepath, ext in collected:
                text = extract_text_from_file(filepath, original_name)
                if not text.strip():
                    continue
                ch_title = os.path.splitext(original_name)[0][:100]
                # 去除正文开头的重复标题行（若文件第一行与文件名相同，避免标题重复）
                body = _strip_leading_title_line(text, ch_title)
                all_chapters.append({'title': ch_title, 'content': body})
        # 情况3：单个文件 → 单文件拆分模式
        else:
            original_name, filepath, ext = collected[0]
            text = extract_text_from_file(filepath, original_name)
            if text.strip():
                chapters = split_into_chapters(text)
                if chapters:
                    all_chapters.extend(chapters)
                else:
                    all_chapters.append({'title': os.path.splitext(original_name)[0][:100], 'content': text})

        if not all_chapters:
            return jsonify({'error': '未能从文件中提取到有效内容，请检查文件格式或编码'}), 400

        # 从第一个文件名推断标题
        if not title:
            first_file = (collected[0][0] if collected else '导入作品') or '导入作品'
            title = os.path.splitext(first_file)[0][:50]

        book = Book(
            user_id=getattr(request, 'current_user_id', ''),
            title=title, author='', genre=genre, book_type=book_type,
            synopsis=f'从文件导入，共 {len(all_chapters)} 章', status='draft'
        )
        db.session.add(book)
        db.session.flush()

        for idx, ch_data in enumerate(all_chapters):
            ch = Chapter(
                book_id=book.id, title=ch_data['title'][:200] or f'第{idx + 1}章',
                content=ch_data['content'], order_index=idx,
                is_volume=False, parent_id='',
                word_count=count_words(ch_data['content'])
            )
            db.session.add(ch)
        db.session.flush()

        # 按章节标题中的章节号（第N章/第N章等）自动排序
        resort_chapters_by_title(book.id, rebin_volumes=False)

        # 自动按 50 章一组创建卷
        chapter_count = len(all_chapters)
        volume_interval = 50
        vol_count = (chapter_count + volume_interval - 1) // volume_interval
        for vidx in range(vol_count):
            ch_start = vidx * volume_interval + 1
            ch_end = min((vidx + 1) * volume_interval, chapter_count)
            vol = Chapter(
                book_id=book.id,
                title=f'第{vidx + 1}卷（第{ch_start}-{ch_end}章）',
                content='', order_index=chapter_count + vidx,
                is_volume=True, parent_id='', word_count=0,
            )
            db.session.add(vol)
            db.session.flush()
            # 将范围内章节归入该卷
            for cidx in range(vidx * volume_interval, min((vidx + 1) * volume_interval, chapter_count)):
                # 找到刚刚创建的章节（按 order_index 为 cidx）
                ch_obj = Chapter.query.filter_by(book_id=book.id, order_index=cidx, is_volume=False).first()
                if ch_obj:
                    ch_obj.parent_id = vol.id

        update_book_stats(book.id)
        return jsonify(book.to_dict()), 201
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

@app.route('/api/books/<book_id>/ai-import-recognize', methods=['POST'])
@login_required
def ai_import_recognize(book_id):
    """导入作品后，根据文件名/章节标题/章节内容样本，AI自动识别并填入各创作维度。
    可识别：构思、设定、大纲、人物、剧情、伏笔、地图、物资库。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    skill_pack_ids = data.get('skill_pack_ids', [])
    # 可选：指定要识别的维度；为空则识别全部
    target_dims = data.get('dimensions', [])

    config = AIConfig.get_active()
    if not config or not config.api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    # 收集文件名/章节标题 + 内容样本
    all_chs = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    if not all_chs:
        return jsonify({'error': '该作品暂无章节，无法识别'}), 400

    # 章节标题列表（即文件名/章节名）
    titles = [ch.title for ch in all_chs[:200]]
    # 内容样本：取前几章 + 中间几章的片段
    samples = []
    sample_chs = all_chs[:3] + all_chs[len(all_chs)//2:len(all_chs)//2+2] if len(all_chs) > 5 else all_chs
    for ch in sample_chs:
        content = (ch.content or '')[:600]
        if content:
            samples.append(f'【{ch.title}】\n{content}')

    titles_text = '\n'.join(titles)
    samples_text = '\n\n'.join(samples)[:4000]

    # P2-10: 修复幽灵key——'character_design'/'world_setting' 不是真实prompt_key，替换为内置存在的key
    # 【三类无污染】AI总创作属于构思阶段：只注入构思类（master）技能包
    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['lock_facts', 'master_outline', 'tomato_character', 'tomato_setting'], mode='agent')

    # 识别哪些维度为空（仅填充空维度，避免覆盖已有内容）
    dim_status = {
        'concept': bool(bb.concept and bb.concept.strip()),
        'settings': bool(bb.key_rules and bb.key_rules.strip()),
        'outline': bool(bb.plot_design and bb.plot_design.strip()),
        'characters': bool(bb.character_profiles and bb.character_profiles.strip()),
        'plot': bool(bb.timeline and bb.timeline.strip()),
        'foreshadowing': bool(bb.foreshadowing and bb.foreshadowing.strip()),
        'locations': bool(bb.locations and bb.locations.strip()),
        'inventory': bool(bb.inventory and bb.inventory.strip()),
    }
    # 默认只填空维度
    empty_dims = [k for k, v in dim_status.items() if not v]
    dims_to_fill = target_dims if target_dims else empty_dims
    if not dims_to_fill:
        return jsonify({'success': True, 'message': '所有维度已有内容，未做修改', 'bible': bb.to_dict(), 'filled': []})

    dims_label = '、'.join(dims_to_fill)
    system_prompt = f"""你是专业的小说设定分析师。请根据导入作品的【文件名/章节标题】和【内容样本】，自动识别并填充以下空维度：{dims_label}。

【识别规则】
1. 仅根据文件名和内容样本推断，不要编造未提供的设定
2. 文件名/章节标题往往暗含卷名、人物名、地点、事件等关键信息，重点提取
3. 每个维度输出对应内容，无法判断的维度输出"（信息不足，待补充）"
4. 保持各维度内容一致性（人物名、地点、能力体系等需统一）

{skill_note}

严格按JSON对象格式输出（不要任何其他文字）：
{{
  "concept": "一句话核心创意（30字内）",
  "settings": "核心规则/能力体系/世界观禁忌（多条用换行分隔）",
  "outline": "主线冲突+卷纲拆解（基于章节标题推断卷结构）",
  "characters": "主要角色档案（JSON数组：[{{\"name\":\"\",\"role\":\"\",\"personality\":\"\",\"motivation\":\"\"}}]，纯文本也可）",
  "plot": "按章节标题梳理的关键事件时间线",
  "foreshadowing": "从内容样本中识别的伏笔（若无则留空）",
  "locations": "从标题/内容识别的地点（若无则留空）",
  "inventory": "从内容识别的物品/功法/法宝（若无则留空）"
}}"""

    user_prompt = f'作品标题：{book.title}\n作品类型：{_get_genre_label(book)}\n\n【文件名/章节标题列表（共{len(titles)}项）】\n{titles_text}\n\n【内容样本】\n{samples_text or "（无内容样本）"}'

    content, err = _call_llm(
        [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
        max_tokens=3000, temperature=0.3, task_type='recognition'
    )
    if err:
        return jsonify({'error': err}), 500

    analysis, parse_err = _extract_json_from_llm(content, expect='object')
    if analysis is None:
        return jsonify({'error': 'AI返回格式无法解析', 'raw': content[:300], 'parse_error': parse_err}), 500

    # 字段映射：维度 -> bible字段
    dim_field_map = {
        'concept': 'concept',
        'settings': 'key_rules',
        'outline': 'plot_design',
        'characters': 'character_profiles',
        'plot': 'timeline',
        'foreshadowing': 'foreshadowing',
        'locations': 'locations',
        'inventory': 'inventory',
    }

    filled = []
    for dim in dims_to_fill:
        field = dim_field_map.get(dim)
        if not field:
            continue
        val = analysis.get(dim, '')
        if isinstance(val, (list, dict)):
            val = json.dumps(val, ensure_ascii=False, indent=2)
        val = str(val).strip()
        if val and val != '（信息不足，待补充）' and not val.startswith('无'):
            setattr(bb, field, val)
            filled.append(dim)

    bb.last_synced_at = datetime.now(timezone.utc)
    db.session.commit()

    return jsonify({
        'success': True,
        'message': f'已识别填充 {len(filled)} 个维度：{"、".join(filled) if filled else "无"}',
        'filled': filled,
        'bible': bb.to_dict()
    })

def _after_chapter_persisted(book_id, chapter) -> Optional[Dict[str, Any]]:
    """【P1-2】章节创建/内容变更落库后统一钩子：
       · 事件抽取（关键章自动 LLM，普通章正则）
       · 伏笔 DAG 反查本章应埋/应收
       · 同步到实体注册表
    返回附带给前端的元信息（可 None）。
    """
    if not chapter or chapter.is_volume or not (chapter.content or '').strip():
        return None
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        return None
    try:
        from event_log_manager import append_chapter_events_auto
        from llm_gateway import LLMGateway, get_llm_config
        known_actors = [c.name for c in Character.query.filter_by(book_id=book_id).all() if c.name]
        known_locations = []
        try:
            _locs = json.loads(bb.locations or '[]')
            if isinstance(_locs, list):
                known_locations = [str(x) for x in _locs if x]
        except Exception:
            pass
        total_chs = Chapter.query.filter_by(book_id=book_id, is_volume=False).count()
        _gw = None
        _base = _api = _model = ''
        try:
            _base, _api, _model = get_llm_config()
            if _base and _api and _model:
                _gw = LLMGateway(_base, _api, _model)
        except Exception:
            _base = _api = _model = ''
        ev = append_chapter_events_auto(
            bb, chapter, chapter.content or '',
            known_actors=known_actors,
            known_locations=known_locations,
            total_chapters=total_chs,
            gw=_gw, base_url=_base, api_key=_api, model=_model,
        )
        # 伏笔 DAG 反查本章埋/收
        hooks_chapter = None
        if bb.foreshadowing_graph:
            try:
                from foreshadowing_manager import ForeshadowingGraph, build_hooks_prompt_section
                graph = ForeshadowingGraph.from_dict(json.loads(bb.foreshadowing_graph))
                hooks = graph.get_nodes_for_chapter(chapter.order_index or 0)
                hooks_chapter = {
                    'setup_ids': [n.id for n in hooks.get('setup', [])],
                    'payoff_ids': [n.id for n in hooks.get('payoff', [])],
                }
                chapter.hooks_set_json = json.dumps(hooks_chapter, ensure_ascii=False)
                # 同步内容中伏笔命中项：如果正文中包含 setup 内容文本即认为已埋/已收
                hooks_chapter['section'] = build_hooks_prompt_section(graph, chapter.order_index or 0)
            except Exception:
                pass
        return {
            'events_added': ev.get('events_added', 0),
            'use_llm': ev.get('use_llm_actual', False),
            'key_chapter': ev.get('key_chapter'),
            'hooks': hooks_chapter,
        }
    except Exception:
        # 不阻塞主流程：失败静默
        return None

def _maybe_auto_trigger_anti_forget_check(book_id, chapter_num=None):
    """【P0-3】防遗忘报告自动触发：每 10 章自动跑一次防遗忘检查。
    在 create_chapter / update_chapter / ai_continue_batch / ai_continue_batch_stream 章节落库后调用。
    - 条件：本书非卷章数为 10 的倍数（且 >0）；本书已有 BookBible；已配置 AI API Key。
    - 执行：在 daemon 线程中跑（防遗忘检查会调 LLM，30-60s），不阻塞主响应。
    - 失败后自动尝试生成设定修正草稿（健康度<80时），并标记为待审阅。
    - 失败：捕获所有异常，不影响主创作流程。
    """
    try:
        # 计算章号：优先用传入的 chapter_num，否则用实际章数
        if not chapter_num:
            chapter_num = Chapter.query.filter_by(book_id=book_id, is_volume=False).count()
        if not chapter_num or chapter_num % 10 != 0:
            return None
        book = Book.query.get(book_id)
        if not book:
            return None
        bb = BookBible.query.filter_by(book_id=book_id).first()
        if not bb:
            return None
        config = AIConfig.get_active()
        if not config or not config.api_key:
            return None

        import threading

        def _bg_run():
            with app.app_context():
                try:
                    # 取该书作者的有效 auth token（绕过 login_required 装饰器）
                    now = datetime.now(timezone.utc)
                    at = AuthToken.query.filter(
                        AuthToken.user_id == book.user_id,
                        AuthToken.expires_at > now
                    ).first()
                    if not at:
                        return
                    # 用 test_request_context 模拟请求体调用 route
                    with app.test_request_context(
                        f'/api/books/{book_id}/ai-anti-forget-check',
                        method='POST',
                        json={'scope': 'reports', 'volume_ids': [], 'skill_pack_ids': []},
                        headers={'Authorization': f'Bearer {at.token}'}
                    ):
                        check_resp = ai_anti_forget_check(book_id)
                    try:
                        app.logger.info(f'[auto] 防遗忘检查自动触发完成：第{chapter_num}章')
                    except Exception:
                        pass

                    # ===== 自动生成设定修正草稿（健康度<80时） =====
                    try:
                        if check_resp and getattr(check_resp, 'status_code', None) == 200:
                            check_data = json.loads(check_resp.get_data(as_text=True)) if hasattr(check_resp, 'get_data') else {}
                            report_rec = check_data.get('report_record') or {}
                            report_json = check_data.get('report') or {}
                            health_score = report_json.get('health_score')
                            report_id = report_rec.get('id')
                            if report_id and (health_score is None or health_score < 80):
                                from blueprints.chat_collab_bp import smart_fix_from_report
                                with app.test_request_context(
                                    '/api/ai/smart/fix-from-report',
                                    method='POST',
                                    json={'book_id': book_id, 'report_id': report_id, 'skill_pack_ids': []},
                                    headers={'Authorization': f'Bearer {at.token}'}
                                ):
                                    fix_resp = smart_fix_from_report()
                                if fix_resp and getattr(fix_resp, 'status_code', None) == 200:
                                    fix_data = json.loads(fix_resp.get_data(as_text=True)) if hasattr(fix_resp, 'get_data') else {}
                                    plan = fix_data.get('plan') or []
                                    if plan:
                                        bb2 = BookBible.query.filter_by(book_id=book_id).first()
                                        if bb2 and bb2.anti_forget_reports:
                                            reports = json.loads(bb2.anti_forget_reports)
                                            for r in reports:
                                                if isinstance(r, dict) and r.get('id') == report_id:
                                                    r['status'] = 'pending'
                                                    r['auto_generated'] = True
                                                    r['fix_draft'] = plan
                                                    r['notified'] = False
                                                    r['draft_generated_at'] = datetime.now(timezone.utc).isoformat()
                                                    break
                                            bb2.anti_forget_reports = json.dumps(reports, ensure_ascii=False, indent=2)
                                            db.session.commit()
                                            try:
                                                app.logger.info(f'[auto] 已生成设定修正草稿：报告 {report_id}，健康度 {health_score}')
                                            except Exception:
                                                pass
                    except Exception as e:
                        try:
                            app.logger.error(f'[auto] 生成修正草稿失败：{str(e)[:200]}')
                        except Exception:
                            pass
                except Exception as e:
                    try:
                        app.logger.error(f'[auto] 防遗忘自动触发失败：{str(e)[:200]}')
                    except Exception:
                        pass

        threading.Thread(target=_bg_run, daemon=True).start()
    except Exception as e:
        try:
            app.logger.error(f'[auto] 防遗忘触发启动失败：{str(e)[:200]}')
        except Exception:
            pass
    return None

@app.route('/api/books/<book_id>/ai-anti-forget-check', methods=['POST'])
@login_required
def ai_anti_forget_check(book_id):
    """长篇小说防遗忘与一致性检查（综合诊断）。
    整合技能包“长篇小说防遗忘系统”的 consistency_check / lock_facts / narrative_debt / foreshadow_register / character_cognition 提示词，
    扫描全部维度+近期章节，输出：锁定事实清单、一致性违规清单、待回收伏笔、叙事债务、改进建议。
    【改造】支持分卷选择（volume_ids 多选/单选），报告持久化到 bb.anti_forget_reports，自动命名"检查01/02..."。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    skill_pack_ids = data.get('skill_pack_ids', [])
    # 检查范围：reports(所有动态报告) / dimensions(仅维度) / volumes(指定分卷)
    scope = data.get('scope', 'reports')
    if scope in ('recent', 'all'):
        scope = 'reports'
    # 【新增】分卷选择：volume_ids 为空表示不按卷筛选；非空则只检查指定卷的章节
    volume_ids = data.get('volume_ids', [])
    if not isinstance(volume_ids, list):
        volume_ids = []

    config = AIConfig.get_active()
    if not config or not config.api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        return jsonify({'error': '该作品暂无创作维度数据，请先填写设定/大纲/剧情等维度。'}), 400

    # 收集章节内容（按 scope + volume_ids 决定范围）
    all_chs = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()

    # 若指定了 volume_ids，按卷筛选章节
    if volume_ids:
        # 通过 parent_id 筛选属于指定卷的章节
        selected_chs = [ch for ch in all_chs if ch.parent_id and ch.parent_id in volume_ids]
        if not selected_chs:
            # 兼容：若按 parent_id 没筛到（可能章节未挂卷），按卷号区间筛选
            # 从 timeline 解析卷的章节范围
            vol_ranges = []
            try:
                if bb.timeline and bb.timeline.strip().startswith('['):
                    arr = json.loads(bb.timeline)
                    for v in arr:
                        if not isinstance(v, dict):
                            continue
                        v_id = str(v.get('volume_id', ''))
                        if v_id in [str(vid) for vid in volume_ids]:
                            nodes = v.get('nodes') or []
                            for n in nodes:
                                if isinstance(n, dict):
                                    ch_range = str(n.get('chapters', ''))
                                    nums = re.findall(r'\d+', ch_range)
                                    if len(nums) >= 2:
                                        vol_ranges.append((int(nums[0]), int(nums[-1])))
            except (json.JSONDecodeError, ValueError, TypeError):
                pass
            if vol_ranges:
                selected_chs = []
                for ch in all_chs:
                    for lo, hi in vol_ranges:
                        if lo <= ch.order_index <= hi:
                            selected_chs.append(ch)
                            break
            else:
                selected_chs = all_chs  # 无法定位时回退全章
        all_chs = selected_chs

    chapter_text = ''
    ch_count = len(all_chs)
    source_label = '动态文件报告'
    if scope == 'dimensions':
        chapter_text, source_label = _collect_dimension_source(bb, '全部维度（除构思、章节）')
        if chapter_text:
            import re as _re_dim
            chapter_text = _re_dim.sub(r'【构思】\n[^\n]*(\n|$)', '', chapter_text).strip()
    else:
        # 动态文件模式（reports）：读取所有动态报告作为检查依据
        all_reports = DynamicReport.query.filter_by(book_id=book_id).order_by(DynamicReport.chapter_start).all()
        # 若指定了 volume_ids，按卷的章节范围筛选动态报告
        if volume_ids and all_reports:
            vol_chapter_set = {ch.order_index for ch in all_chs}
            filtered_reports = []
            for r in all_reports:
                # 报告的章节范围 [chapter_start, chapter_end] 与选中卷的章节有交集则保留
                if any(r.chapter_start <= cn <= r.chapter_end for cn in vol_chapter_set):
                    filtered_reports.append(r)
            all_reports = filtered_reports if filtered_reports else all_reports
        if all_reports:
            parts = [f'【{r.title}（{r.chapter_start}-{r.chapter_end}章）】\n{(r.content or "")[:800]}' for r in all_reports]
            chapter_text = '\n\n'.join(parts)[:10000]
            source_label = f'动态文件（共{len(all_reports)}份报告）'
            if volume_ids:
                source_label += f'·指定{len(volume_ids)}卷'
        elif ch_count > 0:
            recent = all_chs[-10:]
            parts = [f'【{ch.title}】\n{(ch.content or "")[:1000]}' for ch in recent]
            chapter_text = '\n\n'.join(parts)[:8000]
            source_label = '近期章节（暂无动态报告）'
            if volume_ids:
                source_label += f'·指定{len(volume_ids)}卷'
        else:
            chapter_text, source_label = _collect_dimension_source(bb, '全部维度')
            if not chapter_text:
                return jsonify({'error': '该作品暂无动态报告、章节，且维度也为空，无法进行检查。请先生成动态报告或填写维度。'}), 400

    # 收集已有动态文件（若有），作为额外上下文
    dm = DynamicMemory.query.filter_by(book_id=book_id).first()
    dyn_ctx = ''
    if dm:
        for fk in ['foreshadowing_tracker', 'character_ecosystem', 'ability_world']:
            v = getattr(dm, fk, '') or ''
            if v.strip():
                try:
                    parsed = json.loads(v)
                    if isinstance(parsed, dict):
                        summary = parsed.get('summary') or parsed.get('pending') or parsed.get('world_facts') or ''
                        if summary:
                            dyn_ctx += f'【{fk}】\n{str(summary)[:600]}\n\n'
                except (json.JSONDecodeError, ValueError):
                    dyn_ctx += f'【{fk}】\n{v[:600]}\n\n'

    # 【三类无污染】一致性检查属于审查阶段：只注入审查类（review）技能包
    skill_note = _get_skill_prompts_by_category(
        skill_pack_ids, 'review',
        ['consistency_check', 'lock_facts', 'narrative_debt', 'foreshadow_register', 'character_cognition'],
        mode='agent'
    )

    # 构建维度全景上下文
    dim_ctx_parts = []
    dim_fields = [
        ('concept', '构思'), ('key_rules', '设定/核心规则'), ('worldbuilding', '世界观'),
        ('plot_design', '大纲/总纲'), ('character_profiles', '人物档案'),
        ('timeline', '剧情/时间线'), ('foreshadowing', '伏笔'),
        ('locations', '地点'), ('inventory', '物资库'),
    ]
    for f, lbl in dim_fields:
        v = getattr(bb, f, '') or ''
        if v.strip():
            # 人物维度转自然语言，避免 JSON 干扰
            if f == 'character_profiles' and v.strip().startswith('['):
                try:
                    from chat_collab_bp import _character_profiles_to_text
                    v = _character_profiles_to_text(v)
                except Exception:
                    pass
            dim_ctx_parts.append(f'【{lbl}】\n{v}')
    dim_ctx = '\n\n'.join(dim_ctx_parts) or '（维度数据为空）'

    # ===== P2增强：注入章级变更日志 + 伏笔DAG结构化状态 =====
    # 让防遗忘检查从"LLM凭正文推断"升级为"对照结构化数据校验"：
    #   - chapter_changes_log：章级 delta（境界/物品/伏笔/地点），可对照检查"境界跳变/物品凭空获得"
    #   - foreshadowing_graph：结构化伏笔状态（pending/resolved），比 LLM 从文本推断更准
    structured_ctx_parts = []

    # 1. 章级变更日志摘要（最近 30 章，聚焦境界/物品/伏笔/死亡类变更）
    try:
        log_list = json.loads(bb.chapter_changes_log) if bb.chapter_changes_log else []
        if isinstance(log_list, list) and log_list:
            recent_logs = log_list[-30:]
            log_lines = []
            for entry in recent_logs:
                if not isinstance(entry, dict):
                    continue
                ch_num = entry.get('chapter_num', '?')
                chg = entry.get('changes') or {}
                if not isinstance(chg, dict):
                    continue
                segs = []
                # 境界变化（防遗忘检查重点关注）
                for c in (chg.get('CharacterStateChanges') or []):
                    if isinstance(c, dict):
                        nm = c.get('CharacterId') or c.get('Name') or '?'
                        lvl = c.get('NewLevel') or ''
                        ke = c.get('KeyEvent') or ''
                        if lvl or ke:
                            segs.append(f'{nm}{"→"+lvl if lvl else ""}{":"+ke[:40] if ke else ""}')
                # 物品转移
                for it in (chg.get('ItemTransfers') or []):
                    if isinstance(it, dict):
                        nm = it.get('ItemName') or it.get('ItemId') or '?'
                        th = it.get('ToHolder') or ''
                        if th:
                            segs.append(f'{nm}→{th}')
                # 伏笔动作
                for fa in (chg.get('ForeshadowingActions') or []):
                    if isinstance(fa, dict):
                        fid = fa.get('ForeshadowId') or '?'
                        act = fa.get('Action') or ''
                        if act:
                            segs.append(f'伏笔{fid}:{act}')
                if segs:
                    log_lines.append(f'- 第{ch_num}章：{"；".join(segs)}')
            if log_lines:
                structured_ctx_parts.append('【章级变更日志】（最近30章的结构化delta，请对照检查境界跳变/物品凭空获得/伏笔状态矛盾）\n' + '\n'.join(log_lines))
    except Exception:
        pass

    # 2. 伏笔 DAG 结构化状态（pending 节点 + 已回收节点）
    try:
        if bb.foreshadowing_graph and get_hooks_for_chapter:
            graph = ForeshadowingGraph.from_dict(json.loads(bb.foreshadowing_graph))
            pending_nodes = graph.get_pending_nodes()
            resolved_nodes = [n for n in (graph.nodes or []) if getattr(n, 'status', '') in ('resolved', 'payoff', '已回收', '已揭示')]
            fs_dag_lines = []
            if pending_nodes:
                fs_dag_lines.append(f'待回收伏笔（{len(pending_nodes)}条）：')
                for node in pending_nodes[:20]:
                    desc = (node.content or node.title or '')[:80]
                    planted = getattr(node, 'setup_chapter', '') or getattr(node, 'planted_chapter', '') or ''
                    fs_dag_lines.append(f'  - [{node.id}] {desc}（埋设于第{planted}章）')
            if resolved_nodes:
                fs_dag_lines.append(f'已回收伏笔（{len(resolved_nodes)}条，若正文再次当作未回收则属违规）：')
                for node in resolved_nodes[:10]:
                    desc = (node.content or node.title or '')[:60]
                    fs_dag_lines.append(f'  - [{node.id}] {desc}')
            if fs_dag_lines:
                structured_ctx_parts.append('【伏笔DAG结构化状态】（权威伏笔状态，以此为准检查正文是否矛盾）\n' + '\n'.join(fs_dag_lines))
    except Exception:
        pass

    structured_ctx = '\n\n'.join(structured_ctx_parts)

    # 修炼体系小说额外检查项：境界突破条件/越级战斗/年龄与境界进度/时间线与修炼进度/修炼天赋一致性
    af_cult_check = ('\n- 修炼体系：境界突破违反已建立突破条件/越级战斗不合理/年龄与境界进度不匹配/'
                     '时间线与修炼进度矛盾/修炼天赋前后不一致') if is_cultivation_novel(book) else ''

    system_prompt = f"""你是“长篇小说防遗忘与一致性审查员”，整合多个防遗忘技能协同工作：
1. 设定锁定员(lock_facts)：从各维度提取不可变核心事实清单
2. 一致性审查员(consistency_check)：检查近期章节是否违反已锁定设定
3. 伏笔管理师(foreshadow_register)：盘点伏笔状态，标记待回收
4. 叙事债务追踪师(narrative_debt)：盘点悬念承诺与兑现平衡
5. 角色认知管理师(character_cognition)：检查角色认知边界是否被破坏

{skill_note}

【审查重点】
- 人物设定：名字记错/能力超限/性格突变/关系前后矛盾
- 世界规则：违反已建立的物理/魔法/力量体系法则
- 时间线：时间倒流/年龄错误/事件顺序混乱
- 角色认知：角色知道了不该知道的信息（信息差破坏）
- 伏笔状态：已回收伏笔又被当作未回收/待回收伏笔遗忘过久
- 物资/能力：物品/功法/境界前后不一致（跳变/重复获得）{af_cult_check}

严格按JSON格式输出（不要任何其他文字）：
{{
  "locked_facts": ["不可变核心事实1", "不可变核心事实2", "...（最多15条）"],
  "violations": [
    {{"type": "人物/世界规则/时间线/认知/伏笔/物资", "severity": "严重/警告/提示", "location": "出现位置（章节/维度）", "desc": "违规描述", "fix": "修正建议"}}
  ],
  "pending_foreshadowing": [
    {{"content": "伏笔内容", "buried_at": "埋设位置", "urgency": "紧急/一般/可缓", "suggest_chapter": "建议回收章节"}}
  ],
  "narrative_debt": [
    {{"promise": "悬念/承诺", "status": "待兑现/过度透支/已遗忘", "priority": "高/中/低", "note": "说明"}}
  ],
  "character_cognition_issues": ["角色认知边界问题1", "问题2"],
  "suggestions": ["针对性改进建议1", "建议2", "建议3"],
  "health_score": 0-100的整数,
  "summary": "本次检查总体结论（100字内）"
}}"""

    user_prompt = f'作品标题：{book.title}\n数据来源：{source_label}（共{ch_count}章）\n\n【各维度全景】\n{dim_ctx}\n\n【已有动态文件摘要】\n{dyn_ctx or "（无）"}\n\n【结构化状态参照】\n{structured_ctx or "（无）"}\n\n【待审查内容】\n{chapter_text}'

    content, err = _call_llm(
        [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
        max_tokens=3500, temperature=0.3, task_type='recognition'
    )
    if err:
        return jsonify({'error': err}), 500

    try:
        report = json.loads(content)
    except (json.JSONDecodeError, ValueError):
        import re as _re_af
        m = _re_af.search(r'\{[\s\S]*\}', content)
        if m:
            try:
                report = json.loads(m.group())
            except (json.JSONDecodeError, ValueError):
                report = {'summary': content[:500], 'raw': True}
        else:
            report = {'summary': content[:500], 'raw': True}

    # ===== violation.location 规范化（与前端 displayChapterNum 口径对齐）=====
    # LLM 可能输出「第三章/Chapter 3/3. 觉醒灵根」等格式致章号错位。处理：
    # ①解析章号（中文/阿拉伯/Chapter N）②按 displayNum 匹配章节 ③重写为「第{num}章 {标题}」
    try:
        _all_chs_for_loc = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()

        def _display_ch_num(ch):
            n = parse_chapter_number(ch.title or '')
            if isinstance(n, int) and n > 0:
                return n
            oi = ch.order_index if isinstance(ch.order_index, int) else -1
            return max(1, oi + 1)

        _ch_by_display = {}
        for _c in _all_chs_for_loc:
            _dn = _display_ch_num(_c)
            _ch_by_display.setdefault(_dn, _c)  # 重号时保留最早的，一般不重

        def _pure_title_text(title):
            """去掉章节标题里的「第N章 」前缀，只留纯正文标题部分。"""
            if not title:
                return ''
            t = title.strip()
            import re as _re_loc
            # 去掉 第N章/第N回/第N卷... 前缀
            t = _re_loc.sub(r'^第\s*[0-9零一二三四五六七八九十百千万亿两〇]+\s*[章节回卷部篇话集幕折更段讲课夜日年季场]\s*[·\-\s：:]*', '', t)
            # 去掉 Chapter N / N. / N、 前缀
            t = _re_loc.sub(r'^(?:chapter|ch|episode|ep)\.?\s*\d+\s*[\-·\s：:]*', '', t, flags=_re_loc.IGNORECASE)
            t = _re_loc.sub(r'^\d+\s*[\.、:：\-\)\]】，;；]\s*', '', t)
            return t.strip()

        def _normalize_location(loc):
            if not loc or not isinstance(loc, str):
                return loc
            loc = loc.strip()
            if not loc:
                return loc
            # --- Step 1：解析章号 ---
            num = parse_chapter_number(loc)
            # parse_chapter_number 对"第3章 xxx"会返回 3；对"第三章"也返回 3；
            # 但对纯维度如「人物档案」「世界观设定」返回 None，这时直接保留原 location。
            if num is None:
                # 尝试用整段 location 模糊匹配章节标题（比如 LLM 只写了章节名）
                matched = None
                for _c in _all_chs_for_loc:
                    _pt = _pure_title_text(_c.title or '')
                    if _pt and _pt in loc:
                        matched = _c
                        break
                if not matched:
                    return loc
                _dn = _display_ch_num(matched)
                return f'第{_dn}章 {_pure_title_text(matched.title) or matched.title}'
            # --- Step 2：用章号找章节 ---
            target_ch = _ch_by_display.get(num)
            if target_ch is None:
                # 没找到具体章节，但章号是合法数字，保持"第N章"格式至少让前端正则能命中
                return f'第{num}章'
            pure_title = _pure_title_text(target_ch.title or '')
            title_suffix = f' {pure_title}' if pure_title else ''
            return f'第{num}章{title_suffix}'

        _violations = report.get('violations') or []
        if isinstance(_violations, list) and _violations:
            for _v in _violations:
                if isinstance(_v, dict) and 'location' in _v:
                    _v['location'] = _normalize_location(_v.get('location'))
            report['violations'] = _violations
    except Exception:
        pass  # 规范化失败不阻断报告返回

    # 持久化到 DynamicMemory.health_dashboard（防遗忘仪表盘，保留原逻辑）
    if not dm:
        dm = DynamicMemory(book_id=book_id)
        for key in DynamicMemory.FILE_KEYS:
            setattr(dm, key, DynamicMemory.get_empty_template(key))
        db.session.add(dm)
    try:
        hd = json.loads(dm.health_dashboard or '{}') if dm.health_dashboard else {}
    except (json.JSONDecodeError, ValueError):
        hd = {}
    hd['last_anti_forget_check'] = {
        'checked_at': datetime.now(timezone.utc).isoformat(),
        'scope': scope,
        'ch_count': ch_count,
        'health_score': report.get('health_score'),
        'summary': report.get('summary', ''),
        'violation_count': len(report.get('violations', [])),
        'pending_foreshadowing_count': len(report.get('pending_foreshadowing', [])),
    }
    hd['alerts'] = report.get('violations', [])[:20]
    dm.health_dashboard = json.dumps(hd, ensure_ascii=False, indent=2)

    # ===== 【新增】持久化完整报告到 bb.anti_forget_reports，自动命名"检查01/02..." =====
    try:
        existing_reports = json.loads(bb.anti_forget_reports or '[]') if bb.anti_forget_reports else []
        if not isinstance(existing_reports, list):
            existing_reports = []
    except (json.JSONDecodeError, ValueError, TypeError):
        existing_reports = []

    # 计算下一个序号：基于现有报告数量+1，两位数补零
    next_seq = len(existing_reports) + 1
    # 避免重名：若已存在同名，递增序号
    existing_titles = {r.get('title', '') for r in existing_reports if isinstance(r, dict)}
    while f'检查{next_seq:02d}' in existing_titles:
        next_seq += 1
    report_title = f'检查{next_seq:02d}'

    new_report_record = {
        'id': str(uuid.uuid4()),
        'title': report_title,
        'seq': next_seq,
        'checked_at': datetime.now(timezone.utc).isoformat(),
        'scope': scope,
        'volume_ids': volume_ids,
        'ch_count': ch_count,
        'source_label': source_label,
        'health_score': report.get('health_score'),
        'summary': report.get('summary', ''),
        'report': report,  # 完整报告 JSON
        'status': 'pending',          # pending / reviewed / applied / ignored
        'auto_generated': False,
        'fix_draft': None,
        'text_fix_draft': None,
        'notified': False,
    }
    existing_reports.append(new_report_record)
    # 限制最多保留 200 份历史报告（百万字长篇需保留更多历史诊断供回注与回溯），超出则删除最早的
    _MAX_ANTI_FORGET_REPORTS = 200
    if len(existing_reports) > _MAX_ANTI_FORGET_REPORTS:
        existing_reports = existing_reports[-_MAX_ANTI_FORGET_REPORTS:]
    bb.anti_forget_reports = json.dumps(existing_reports, ensure_ascii=False, indent=2)

    db.session.commit()

    return jsonify({
        'success': True,
        'report': report,
        'report_record': new_report_record,
        'scope': scope,
        'ch_count': ch_count,
        'source_label': source_label
    })

@app.route('/api/books/<book_id>/anti-forget-reports', methods=['GET'])
@login_required
def list_anti_forget_reports(book_id):
    """列出所有防遗忘检查报告历史。"""
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        return jsonify({'reports': []})
    try:
        reports = json.loads(bb.anti_forget_reports or '[]') if bb.anti_forget_reports else []
        if not isinstance(reports, list):
            reports = []
    except (json.JSONDecodeError, ValueError, TypeError):
        reports = []
    # 按 seq 倒序返回（最新在前）
    reports.sort(key=lambda r: r.get('seq', 0) if isinstance(r, dict) else 0, reverse=True)
    return jsonify({'reports': reports})

@app.route('/api/books/<book_id>/anti-forget-reports/<report_id>', methods=['PUT'])
@login_required
def update_anti_forget_report(book_id, report_id):
    """编辑/重命名/状态流转防遗忘检查报告。
    请求体：
      {"title": "新标题"} 重命名
      {"report": {...}} 替换报告内容
      {"summary": "..."} 更新摘要
      {"status": "pending|reviewed|applied|ignored"} 状态流转
      {"fix_draft": [...]} 覆盖/清空设定修正草稿
      {"text_fix_draft": [...]} 覆盖/清空正文修正草稿
      {"notified": true} 标记已通知
    """
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        return jsonify({'error': 'Not found'}), 404
    data = request.get_json() or {}
    try:
        reports = json.loads(bb.anti_forget_reports or '[]') if bb.anti_forget_reports else []
        if not isinstance(reports, list):
            reports = []
    except (json.JSONDecodeError, ValueError, TypeError):
        reports = []

    VALID_STATUSES = {'pending', 'reviewed', 'applied', 'ignored'}
    found = False
    for r in reports:
        if isinstance(r, dict) and r.get('id') == report_id:
            if 'title' in data:
                new_title = str(data['title']).strip()
                if not new_title:
                    return jsonify({'error': '标题不能为空'}), 400
                r['title'] = new_title
            if 'report' in data and isinstance(data['report'], dict):
                r['report'] = data['report']
            if 'summary' in data:
                r['summary'] = str(data['summary'])
                if isinstance(r.get('report'), dict):
                    r['report']['summary'] = str(data['summary'])
            if 'health_score' in data:
                r['health_score'] = data['health_score']
            if 'status' in data:
                st = str(data['status']).strip().lower()
                if st not in VALID_STATUSES:
                    return jsonify({'error': f'无效状态：{st}'}), 400
                r['status'] = st
            if 'fix_draft' in data:
                r['fix_draft'] = data['fix_draft'] if isinstance(data['fix_draft'], list) else None
            if 'text_fix_draft' in data:
                r['text_fix_draft'] = data['text_fix_draft'] if isinstance(data['text_fix_draft'], list) else None
            if 'notified' in data:
                r['notified'] = bool(data['notified'])
            r['updated_at'] = datetime.now(timezone.utc).isoformat()
            found = True
            break
    if not found:
        return jsonify({'error': '报告不存在'}), 404
    bb.anti_forget_reports = json.dumps(reports, ensure_ascii=False, indent=2)
    db.session.commit()
    return jsonify({'success': True, 'reports': reports})

@app.route('/api/books/<book_id>/anti-forget-reports/<report_id>', methods=['DELETE'])
@login_required
def delete_anti_forget_report(book_id, report_id):
    """删除防遗忘检查报告。"""
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        return jsonify({'error': 'Not found'}), 404
    try:
        reports = json.loads(bb.anti_forget_reports or '[]') if bb.anti_forget_reports else []
        if not isinstance(reports, list):
            reports = []
    except (json.JSONDecodeError, ValueError, TypeError):
        reports = []
    new_reports = [r for r in reports if not (isinstance(r, dict) and r.get('id') == report_id)]
    if len(new_reports) == len(reports):
        return jsonify({'error': '报告不存在'}), 404
    bb.anti_forget_reports = json.dumps(new_reports, ensure_ascii=False, indent=2)
    db.session.commit()
    return jsonify({'success': True, 'reports': new_reports})

@app.route('/api/books/<book_id>/import-chapters', methods=['POST'])
@login_required
def append_import_chapters(book_id):
    """追加导入章节到已有作品，支持 txt/md/docx/zip，每个文件可含多章。
    新章节会按当前最大 order_index 顺序追加，不影响已有章节。
    """
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': '作品不存在'}), 404
    if book.user_id != request.current_user_id:
        return jsonify({'error': '无权操作该作品'}), 403

    files = request.files.getlist('files')
    if not files or len(files) == 0:
        return jsonify({'error': '未选择文件'}), 400

    tmpdir = tempfile.mkdtemp()
    try:
        new_chapters = []
        for f in files:
            if not f or not f.filename:
                continue
            original_name = f.filename
            ext = original_name.rsplit('.', 1)[1].lower() if '.' in original_name else ''
            if ext not in ('txt', 'md', 'markdown', 'docx', 'zip', 'json'):
                continue
            safe_name = f'{uuid.uuid4()}.{ext}'
            filepath = os.path.join(tmpdir, safe_name)
            f.save(filepath)
            text = extract_text_from_file(filepath, safe_name)
            if not text.strip():
                continue
            chapters = split_into_chapters(text)
            if chapters:
                new_chapters.extend(chapters)
            else:
                ch_title = os.path.splitext(original_name)[0]
                new_chapters.append({'title': ch_title[:100], 'content': text})

        if not new_chapters:
            return jsonify({'error': '未能从文件中提取到有效内容，请检查文件格式或编码'}), 400

        # 计算当前最大 order_index，新章节追加在末尾
        max_order = db.session.query(db.func.max(Chapter.order_index)).filter_by(book_id=book_id).scalar() or 0
        added = 0
        for idx, ch_data in enumerate(new_chapters):
            ch = Chapter(
                book_id=book_id,
                title=ch_data['title'][:200] or f'第{max_order + idx + 1}章',
                content=ch_data['content'],
                order_index=max_order + idx + 1,
                is_volume=False,
                parent_id='',
                word_count=count_words(ch_data['content'])
            )
            db.session.add(ch)
            added += 1
        db.session.flush()

        # 按章节标题中的章节号自动重排（含已有章节），并按 50 章/卷重新归入卷
        resort_chapters_by_title(book_id, rebin_volumes=True)

        update_book_stats(book_id)
        total = Chapter.query.filter_by(book_id=book_id, is_volume=False).count()
        return jsonify({'success': True, 'added': added, 'total': total}), 200
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

# ==== Count Words ====

@app.route('/api/utils/count-words', methods=['POST'])
def count_words_api():
    text = request.json.get('text', '')
    return jsonify({'count': count_words(text)})

# ==== Stage Content API ====

SHORT_STAGES = [
    {'key': 'character_design', 'label': '人物设计', 'icon': '👤', 'desc': '创建、补全和诊断人物设定'},
    {'key': 'plot_design', 'label': '剧情设计', 'icon': '📋', 'desc': '核心命题、人物目标、主要冲突、因果链', 'is_parent': True},
    {'key': 'intro_design', 'label': '导语设计', 'icon': '✍️', 'desc': '书名建议、开篇导语、前十秒钩子', 'parent': 'plot_design'},
    {'key': 'plot_refine', 'label': '剧情细化', 'icon': '🔍', 'desc': '场景链、节拍、信息投放、伏笔与回收', 'parent': 'plot_design'},
    {'key': 'outline', 'label': '大纲', 'icon': '📝', 'desc': '小节清单、字数分配、场景承接'},
    {'key': 'draft', 'label': '正文编写', 'icon': '📖', 'desc': '正式写作，支持专家模式分节写作'},
]

SCRIPT_STAGES = [
    {'key': 'character_design', 'label': '人物设计', 'icon': '👤', 'desc': '创建可表演的人物设定'},
    {'key': 'plot_design', 'label': '剧情设计', 'icon': '📋', 'desc': '可拍摄的故事架构', 'is_parent': True},
    {'key': 'plot_refine', 'label': '剧情细化', 'icon': '🔍', 'desc': '场次分解与节奏控制', 'parent': 'plot_design'},
    {'key': 'outline', 'label': '大纲', 'icon': '📝', 'desc': '场次列表与时长分配'},
    {'key': 'draft', 'label': '剧本编写', 'icon': '🎬', 'desc': '场景标题格式：序号. 内/外景 地点 - 时间'},
]

LONG_STAGES = [
    {'key': 'worldbuilding', 'label': '世界观', 'icon': '🌍', 'desc': '规则、势力、地理、历史、境界'},
    {'key': 'character_design', 'label': '人物', 'icon': '👤', 'desc': '主角、配角、路人'},
    {'key': 'plot_design', 'label': '剧情', 'icon': '📋', 'desc': '全书线、分卷、剧情弧、章卡、伏笔'},
    {'key': 'draft', 'label': '正文', 'icon': '📖', 'desc': '按章写作，状态账本落盘'},
    {'key': 'continuity_ledger', 'label': '状态账本', 'icon': '📒', 'desc': '时间线、人物状态、势力状态、伏笔回收'},
]

def get_stages_for_type(book_type):
    if book_type == 'novel':
        return LONG_STAGES
    elif book_type == 'script':
        return SCRIPT_STAGES
    return SHORT_STAGES

@app.route('/api/books/<book_id>/stages', methods=['GET'])
def list_stages(book_id):
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404
    stages_def = get_stages_for_type(book.book_type)
    contents = {s.stage_key: s.to_dict() for s in StageContent.query.filter_by(book_id=book_id).all()}
    result = []
    for sd in stages_def:
        item = dict(sd)
        item['content'] = contents.get(sd['key'], {}).get('content', '')
        item['stage_id'] = contents.get(sd['key'], {}).get('id', '')
        result.append(item)
    return jsonify(result)

@app.route('/api/books/<book_id>/stages/<stage_key>', methods=['GET'])
def get_stage(book_id, stage_key):
    sc = StageContent.query.filter_by(book_id=book_id, stage_key=stage_key).first()
    return jsonify(sc.to_dict() if sc else {'book_id': book_id, 'stage_key': stage_key, 'content': ''})

@app.route('/api/books/<book_id>/stages/<stage_key>', methods=['PUT'])
def save_stage(book_id, stage_key):
    sc = StageContent.query.filter_by(book_id=book_id, stage_key=stage_key).first()
    if not sc:
        sc = StageContent(book_id=book_id, stage_key=stage_key)
        db.session.add(sc)
    sc.content = request.json.get('content', '')
    db.session.commit()
    return jsonify(sc.to_dict())

# ==== Prompt Template API ====

SEED_PROMPTS = [
    # 短篇 - 人物设计
    {'name': '短篇人物设计（通用）', 'agent_id': 'character_design', 'book_type': 'short_story', 'genre': 'other', 'is_builtin': True,
     'description': '创建、补全、诊断和修改人物设计', 'content': '你是蚂蚁写作的人物设计智能体。\n你的职责是创建、补全、诊断和修改人物设计。\n\n工作流程：\n1. 判断用户是在新建人物、补全人物，还是修改已有设定。\n2. 修改已有内容前，先读取人物设计阶段内容。\n3. 形成人物稿并使用工具写回人物编辑器。\n\n人物设计至少关注：\n- 身份与处境\n- 核心欲望/恐惧/缺陷/秘密/底线\n- 行动逻辑\n- 关系结构\n- 辨识度\n- 人物弧\n\n只写正式设定，不写分析过程。'},
    # 短篇 - 剧情设计（父阶段，含3个子槽位）
    {'name': '短篇剧情设计（通用）', 'agent_id': 'plot_design', 'book_type': 'short_story', 'genre': 'other', 'is_builtin': True,
     'description': '统一负责剧情设计、导语设计和剧情细化三个子方向', 'content': '你是蚂蚁写作的短篇剧情智能体，统一负责剧情设计、导语设计和剧情细化。\n\n三个内容槽位的边界：\n- 剧情设计（plot_design）：核心命题、人物目标、主要冲突、因果链、关键转折、真实时间线和结局兑现。\n- 导语设计（intro_design）：书名建议、开篇导语和前十秒钩子。\n- 剧情细化（plot_refine）：供正文直接执行的场景链、节拍、信息投放、人物选择、情绪推进、伏笔与回收。\n\n工作流程：\n1. 先确认用户本次处理哪个子方向。\n2. 读取人物设计和已有的剧情内容。\n3. 检查因果是否成立、冲突是否递进、转折是否由人物选择触发。\n4. 把成品写入正确的剧情子槽位。'},
    # 短篇 - 大纲
    {'name': '短篇大纲（通用）', 'agent_id': 'outline', 'book_type': 'short_story', 'genre': 'other', 'is_builtin': True,
     'description': '把人物和剧情梳理成可直接指导分节写作的完整大纲', 'content': '你是蚂蚁写作的短篇大纲智能体，负责把人物和剧情内容梳理成可直接指导分节写作的完整大纲。\n\n开始任何大纲任务前，必须分别读取：\n1. 人物设计\n2. 剧情设计\n3. 导语设计\n4. 剧情细化\n5. 当前大纲\n\n大纲成品必须包含：\n- 全文定位、主线目标、核心冲突、时间线与结局\n- 正文小节总数及顺序\n- 每个小节的标题/预估字数/出场人物/场景/起始状态/详细剧情\n- 小节之间的承接关系、人物状态变化、伏笔埋设与回收位置'},
    # 短篇 - 正文专家总控
    {'name': '短篇正文专家总控', 'agent_id': 'expert_draft_coordinator', 'book_type': 'short_story', 'genre': 'other', 'is_builtin': True,
     'description': '负责正文结构管理、分节任务调度和成稿后的处理', 'content': '你是蚂蚁写作的短篇正文专家编写智能体，负责正文结构管理、分节任务调度和成稿后的处理。\n\n你负责四类任务：\n1. 初始化：读取大纲，根据完整大纲创建导语、全部正文小节及人物状态槽位。\n2. 全部写作：先读取大纲并初始化，再启动自动写作。\n3. 单节写作：用户指定一个已初始化小节时，启动单节写作。\n4. 后处理：正文审阅、润色、去AI味、格式整理、章节名修改和局部修订。\n\n初始化前必须读取大纲；小节标题/顺序/数量必须与大纲一致。'},
    # 短篇 - 分节写手
    {'name': '短篇分节写手', 'agent_id': 'expert_section_writer', 'book_type': 'short_story', 'genre': 'other', 'is_builtin': True,
     'description': '实际创作小说正文的主要智能体，一次只处理一个小节', 'content': '你是蚂蚁写作的短篇分节写手智能体，是实际创作小说正文的主要智能体。\n你一次只处理当前上下文指定的一个小节，不得修改其它小节。\n\n写作前必须完成：\n1. 读取大纲；允许时补充读取剧情细化。\n2. 读取当前小节之前最近三个已有正文的小节。\n3. 必须读取紧邻上一节的人物状态。\n\n写作标准：\n- 严格执行当前小节在大纲中的任务、承接点和字数要求（默认800-1500字）。\n- 延续前文的时间、空间、人物关系、信息知情范围。\n- 让冲突通过人物行动、选择、对白和可感知细节推进。\n- 保持题材、叙述视角、文风和节奏一致。\n- 小节结尾应完成本节任务并留下明确承接点。'},
    # 言情专项
    {'name': '言情-人物设计', 'agent_id': 'character_design', 'book_type': 'short_story', 'genre': 'romance', 'is_builtin': True,
     'description': '网文爆款人物设计：标签叠加法、冲突校验', 'content': '你是蚂蚁写作的言情人物设计智能体。\n\n网文爆款人物设计原则：\n1. 标签叠加法——选择3-5个反差标签叠加（如"高冷总裁+童年创伤+情感洁癖"）\n2. 冲突校验——每对人物之间至少存在价值观冲突、目标冲突、信息差和情感拉扯\n3. 辨识度——每个角色有独特的口头禅、行为习惯、情感表达方式\n4. 人物弧——明确起点状态和终点状态，设定合理的转折事件\n\n输出格式：\n- 核心身份信息（姓名、年龄、职业/身份）\n- 性格特质（3-5个核心标签）\n- 核心欲望与恐惧\n- 人物背景与转折点\n- 与其他角色的关系图谱'},
    # 玄幻专项
    {'name': '玄幻-人物设计', 'agent_id': 'character_design', 'book_type': 'short_story', 'genre': 'fantasy', 'is_builtin': True,
     'description': '玄幻修真专项：技术代价、道德边界、成长路径', 'content': '你是蚂蚁写作的玄幻人物设计智能体。\n\n玄幻人物设计要点：\n1. 成长路径——从凡人到巅峰的阶梯式升级路线\n2. 技术代价——每次突破都应有对应的代价（寿命、情感、道德）\n3. 道德边界——定义角色不可触碰的底线\n4. 世界观适配——人物能力体系必须与世界规则一致\n\n输出格式：\n- 身份与修为等级\n- 功法体系与特殊能力\n- 性格特质与行为模式\n- 核心价值观与道德底线\n- 成长路线与关键转折'},
    # 悬疑专项
    {'name': '悬疑-剧情设计', 'agent_id': 'plot_design', 'book_type': 'short_story', 'genre': 'mystery', 'is_builtin': True,
     'description': '悬疑推理专项：秘密、动机、误导、信息差', 'content': '你是蚂蚁写作的悬疑剧情设计智能体。\n\n悬疑剧情设计要点：\n1. 真相底牌——确定最终的真相是什么，所有线索都应指向它\n2. 线索结构——区分真线索（导向真相）和红鲱鱼（导向歧途）\n3. 信息释放节奏——每章递进式释放信息，读者和侦探之间的信息差\n4. 误导设计——每个红鲱鱼必须合理化，不能是纯粹欺骗\n5. 人物秘密——每个主要人物都有隐藏的秘密\n\n输出格式：\n- 真相概述\n- 线索链（真线索按发现顺序排列）\n- 红鲱鱼列表（每条附合理化解释）\n- 关键转折点\n- 结局兑现清单'},
    # 审稿提示词
    {'name': '正文审阅', 'agent_id': 'draft_review', 'book_type': 'short_story', 'genre': 'other', 'is_builtin': True,
     'description': '5维度审稿：句式、标点、否定句、形容词、比喻', 'content': '你是蚂蚁写作的正文审阅智能体。\n\n从以下5个维度审阅正文：\n\n1. 句式节奏：\n- 长短句交替，避免连续3句以上相同长度\n- 动作场景使用短句，情感场景使用中长句\n- 检查是否有"的的不休"现象\n\n2. 标点规范：\n- 禁止使用破折号（——）\n- 省略号统一使用"……"（6个点）\n- 引号嵌套不超过一层\n\n3. 否定句式：\n- 优先使用肯定句式\n- 每个自然段内否定句式不超过2处\n\n4. 形容词节制：\n- 每个名词前最多2个修饰形容词\n- 情感描写作减法而非加法\n\n5. 比喻节制：\n- 每千字比喻不超过3处\n- 比喻必须与小说世界观一致'},
]

def seed_prompt_templates():
    if PromptTemplate.query.filter_by(is_builtin=True).first():
        return
    for p in SEED_PROMPTS:
        pt = PromptTemplate(
            name=p['name'], agent_id=p['agent_id'], book_type=p['book_type'],
            genre=p['genre'], content=p['content'], is_builtin=True,
            description=p['description']
        )
        db.session.add(pt)
    db.session.commit()

@app.route('/api/prompts', methods=['GET'])
def list_prompts():
    book_type = request.args.get('book_type', '')
    agent_id = request.args.get('agent_id', '')
    query = PromptTemplate.query
    if book_type:
        query = query.filter_by(book_type=book_type)
    if agent_id:
        query = query.filter_by(agent_id=agent_id)
    prompts = query.order_by(PromptTemplate.is_builtin.desc(), PromptTemplate.name).all()
    return jsonify([p.to_dict() for p in prompts])

@app.route('/api/prompts/<prompt_id>', methods=['GET'])
def get_prompt(prompt_id):
    pt = PromptTemplate.query.get(prompt_id)
    if not pt:
        return jsonify({'error': 'Not found'}), 404
    return jsonify(pt.to_dict())

@app.route('/api/prompts', methods=['POST'])
def create_prompt():
    data = request.json
    pt = PromptTemplate(
        name=data.get('name', ''), agent_id=data.get('agent_id', ''),
        book_type=data.get('book_type', 'short_story'), genre=data.get('genre', 'other'),
        content=data.get('content', ''), description=data.get('description', '')
    )
    db.session.add(pt)
    db.session.commit()
    return jsonify(pt.to_dict()), 201

@app.route('/api/prompts/<prompt_id>', methods=['PUT'])
def update_prompt(prompt_id):
    pt = PromptTemplate.query.get(prompt_id)
    if not pt:
        return jsonify({'error': 'Not found'}), 404
    data = request.json
    for field in ['name', 'agent_id', 'book_type', 'genre', 'content', 'description']:
        if field in data:
            setattr(pt, field, data[field])
    db.session.commit()
    return jsonify(pt.to_dict())

# ==== Book Bible API (项目宪法) ====

def _normalize_bible_formats(bb):
    """P1-5/6/7: 统一 bible 字段格式归一化（幂等，安全）。
    - inventory: 纯文本 → JSON 数组 [{volume:'历史数据', data:旧文本}]
    - timeline: 不强制改（保留双语义，但 ai_outline_volume 已有丢弃重建逻辑）
    - character_profiles: 不强制改（保留文本/JSON双兼容，_filter_bible_by_relevance 已兼容）
    仅 inventory 做主动迁移，其他字段在读取时容错。"""
    if not bb:
        return
    changed = False
    # inventory 迁移
    inv = bb.inventory or ''
    if inv.strip():
        try:
            parsed = json.loads(inv)
            if isinstance(parsed, str) and parsed.strip():
                # JSON 字符串包裹的纯文本
                bb.inventory = json.dumps([{'volume': '历史数据', 'volume_id': '', 'data': parsed}], ensure_ascii=False)
                changed = True
            # list 或 dict 保持不变
        except (json.JSONDecodeError, ValueError):
            # 纯文本：包裹为单元素数组
            bb.inventory = json.dumps([{'volume': '历史数据', 'volume_id': '', 'data': inv}], ensure_ascii=False)
            changed = True
    if changed:
        try:
            db.session.commit()
        except Exception:
            db.session.rollback()

@app.route('/api/books/<book_id>/bible', methods=['GET'])
def get_book_bible(book_id):
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()
    # P1-5: 格式归一化（inventory 纯文本→JSON数组）
    _normalize_bible_formats(bb)
    # 数据修复迁移：历史 bug 导致关系图谱文本被写入 character_profiles，破坏了 JSON 数组结构。
    # 如果 relation_graph 为空，且 character_profiles 看起来是图谱文本（非 JSON 数组），
    # 则把图谱文本迁到 relation_graph，并清空 character_profiles（让用户重新维护角色档案）。
    try:
        cp = bb.character_profiles or ''
        rg = bb.relation_graph or ''
        if not rg.strip() and cp.strip():
            is_json_array = False
            try:
                parsed = json.loads(cp)
                is_json_array = isinstance(parsed, list)
            except (json.JSONDecodeError, ValueError):
                is_json_array = False
            if not is_json_array:
                # 看起来像图谱文本（含 /*EDGE: 标记 或 含 "关系:" 行 或 含 | 分隔的人物行）
                looks_like_graph = ('/*EDGE:' in cp) or ('关系:' in cp) or ('关系：' in cp) or \
                                   any('|' in line and ('人物' in line or '姓名' in line) for line in cp.split('\n')[:5])
                if looks_like_graph:
                    bb.relation_graph = cp
                    bb.character_profiles = ''
                    db.session.commit()
    except Exception:
        db.session.rollback()
    return jsonify(bb.to_dict())

@app.route('/api/books/<book_id>/bible', methods=['PUT'])
def update_book_bible(book_id):
    try:
        bb = BookBible.query.filter_by(book_id=book_id).first()
        if not bb:
            bb = BookBible(book_id=book_id)
            db.session.add(bb)
        data = request.json
        foreshadowing_edited = False
        for field in ['worldbuilding', 'character_profiles', 'timeline', 'foreshadowing', 'style_guide', 'key_rules', 'locations', 'concept', 'plot_design', 'relation_graph', 'inventory', 'character_volumes', 'dynamic_volumes', 'foreshadowing_volumes', 'locations_volumes']:
            if field in data:
                setattr(bb, field, data[field])
                if field == 'foreshadowing':
                    foreshadowing_edited = True
        # P2-7：伏笔双轨一致性 —— 文本字段人工编辑后触发 DAG 重建，避免文本与 DAG 状态不一致
        if foreshadowing_edited and parse_text_to_dag and bb.foreshadowing:
            try:
                new_dag = parse_text_to_dag(bb.foreshadowing)
                if new_dag and (new_dag.get('nodes') or new_dag.get('edges')):
                    bb.foreshadowing_graph = json.dumps(new_dag, ensure_ascii=False)
            except Exception:
                # DAG 重建失败不阻断文本字段保存
                pass
        db.session.commit()
        return jsonify(bb.to_dict())
    except Exception as e:
        db.session.rollback()
        import traceback
        err = traceback.format_exc()
        print(f'[update_book_bible ERROR] {err}', flush=True)
        return jsonify({'error': str(e), 'detail': err[-500:]}), 500

@app.route('/api/books/<book_id>/bible/sync', methods=['POST'])
def sync_book_bible(book_id):
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
    chapters = Chapter.query.filter_by(book_id=book_id).order_by(Chapter.order_index).all()
    characters = Character.query.filter_by(book_id=book_id).all()
    outlines = Outline.query.filter_by(book_id=book_id).order_by(Outline.order_index).all()

    summary_parts = [
        f"## 作品信息\n书名：{Book.query.get(book_id).title}\n类型：{Book.query.get(book_id).genre}\n\n",
        f"## 角色档案 ({len(characters)}人)\n"
    ]
    for c in characters:
        parts = [f"- {c.name} ({c.role})"]
        if c.personality: parts.append(f" 性格：{c.personality[:200]}")
        if c.background: parts.append(f" 背景：{c.background[:200]}")
        summary_parts.append('\n'.join(parts) + '\n')

    summary_parts.append(f"\n## 大纲 ({len(outlines)}条)\n")
    for o in outlines:
        summary_parts.append(f"- {o.title}\n  {o.content[:200]}\n")

    summary_parts.append(f"\n## 章节 ({len(chapters)}章)\n")
    for ch in chapters:
        summary_parts.append(f"- 第{ch.order_index + 1}章 {ch.title} ({ch.word_count}字)\n")

    bb.generated_summary = ''.join(summary_parts)
    bb.last_synced_at = datetime.now(timezone.utc)
    db.session.commit()
    return jsonify(bb.to_dict())

# ==== AI Review API (AI 责编) ====

REVIEW_CRITERIA = {
    'opening_hook': {'label': '开篇钩子', 'weight': 20, 'desc': '前三段能否抓住读者'},
    'character_motivation': {'label': '人物动机', 'weight': 20, 'desc': '人物行为是否有清晰动机'},
    'pacing_rhythm': {'label': '节奏控制', 'weight': 15, 'desc': '张弛有度，爽点密度'},
    'chapter_ending': {'label': '章尾钩子', 'weight': 15, 'desc': '每章结尾是否留下追读期待'},
    'dialogue_quality': {'label': '对白质量', 'weight': 10, 'desc': '对白是否推动剧情/展现性格'},
    'world_consistency': {'label': '设定一致性', 'weight': 10, 'desc': '前后设定是否有矛盾'},
    'commercial_potential': {'label': '商业潜力', 'weight': 10, 'desc': '是否适合目标平台读者'},
}

@app.route('/api/books/<book_id>/review', methods=['POST'])
def review_book(book_id):
    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.model if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    scope = request.json.get('scope', 'latest')
    chapters = Chapter.query.filter_by(book_id=book_id).order_by(Chapter.order_index).all()
    if scope == 'latest' and chapters:
        text = chapters[-1].content or ''
        for ch in chapters[-3:-1]:
            if ch.content:
                text = ch.content + '\n\n' + text
    else:
        text = request.json.get('content', '')
        if not text:
            text = '\n\n'.join([(c.content or '') for c in chapters])

    if not text.strip():
        return jsonify({'error': 'No content to review'}), 400
    text = text[:6000]

    criteria_desc = '\n'.join([f"- {v['label']}({v['weight']}分): {v['desc']}" for v in REVIEW_CRITERIA.values()])
    system_prompt = f"""你是资深网文责编，服务于番茄小说/起点中文网。请从以下维度审稿并严格按JSON格式输出结果。

评分维度：
{criteria_desc}

输出格式（严格JSON，不要任何其他文字）：
{{"scores": {{"opening_hook": 0-100, "character_motivation": 0-100, "pacing_rhythm": 0-100, "chapter_ending": 0-100, "dialogue_quality": 0-100, "world_consistency": 0-100, "commercial_potential": 0-100}},
"total_score": 0-100,
"grade": "S/A/B/C/D",
"strengths": ["优点1", "优点2", "优点3"],
"weaknesses": ["问题1", "问题2", "问题3"],
"specific_suggestions": ["具体修改建议1", "具体修改建议2", "具体修改建议3"],
"platform_fit": "适合/不太适合番茄/起点/七猫的理由"}}"""

    try:
        resp = requests.post(f'{base_url}/chat/completions',
            headers=build_auth_headers(api_key),
            json={'model': model, 'messages': [{'role':'system','content':system_prompt},{'role':'user','content':text}],
                  'temperature': 0.3, 'max_tokens': 2000, 'response_format': {'type': 'json_object'}},
            timeout=120)
        result = resp.json()
        return jsonify(json.loads(result['choices'][0]['message']['content']))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ==== Skill Pack API ====

# 【铁律】内置技能包章节字数统一 2400±100（2300-2500），适用于 novel 类 write_chapter 提示词；
# 短篇(short_story)保留总字数规范。改字数须同步本注释与所有内置包提示词。
SEED_SKILL_PACKS = [
    {'name': '番茄爽文三件套', 'description': '构思专用：开篇钩子设计+黄金三章节奏规划，番茄平台爽文前期结构管线', 'genre': 'other', 'book_type': 'novel', 'icon':'🍅',
     'stage_keys': json.dumps(['character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'开篇钩子设计','desc':'设计前三段冲突入画+主角特质+疑问钩子，输出3个方案','prompt_key':'opening_hook'},
         {'step':2,'name':'黄金三章结构规划','desc':'设计三章节奏：日常困境→金手指/转折→初步成绩+悬念','prompt_key':'golden_three'},
         {'step':3,'name':'全书爽点卡位规划','desc':'设计每卷/每阶段的爽点密度、打脸/反转节点及章尾钩子布局','prompt_key':'pleasure_points'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'opening_hook': '【构思阶段】你是番茄爽文开头结构设计师（不写正文，只输出设计方案）。请设计3个可落地的开篇钩子方案，每个方案包含：①前三段场景（冲突入画的第一句是什么、主角一句话特质、结尾留的疑问钩子）②目标读者情绪③钩子兑现时间点。请编号列出，不要写散文式正文。',
         'golden_three': '【构思阶段】你是番茄爽文结构规划师（不写正文，只输出设计方案）。请设计黄金三章的节奏方案：第一章——展示的「日常」是什么、「困境/痛点」是什么；第二章——金手指/转折事件如何触发、给主角的第一个短期目标；第三章——展示的「初步成绩」、章尾留的悬念钩子。用编号列出三章的结构要点，不写散文。',
         'pleasure_points': '【构思阶段】你是番茄爽文爽点卡位规划师。请设计本书的爽点密度方案：①按章节/卷给出爽点类型配额（打脸×N、反转×N、装逼×N、情感糖×N）②章尾钩子的强制出现位置（每N章必出）③前期、中期、后期爽点强度递增曲线。用编号要点输出，不写正文场景。',
     }, ensure_ascii=False)},
    {'name': '起点升级流大师', 'description': '构思专用：境界体系设计+升级节奏规划+副本设计+分卷大纲，起点升级流全链路构思', 'genre': 'fantasy', 'book_type': 'novel', 'icon':'⚔️',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'境界体系设计','desc':'设计完整境界升级体系，融入规则设定维度','prompt_key':'level_system'},
         {'step':2,'name':'升级节奏规划','desc':'规划全书章节跨度、突破节点、压制爆发曲线','prompt_key':'level_curve'},
         {'step':3,'name':'主要副本设计','desc':'设计核心副本/秘境的层级、奖励、难度匹配','prompt_key':'dungeon_design'},
         {'step':4,'name':'分卷大纲规划','desc':'按卷设计核心目标、冲突线、成长目标、转折点','prompt_key':'volume_outline'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'level_system': '【构思阶段】你是玄幻升级体系设计师。请以编号要点形式输出完整境界体系，用于嵌入"设定/规则"维度：①境界名称及每境界细分阶位 ②各境界突破所需条件与核心资源 ③战力跨度/标志性能力变化 ④对应寿命、社会地位。不要写散文。',
         'level_curve': '【构思阶段】你是升级节奏规划师。请输出本书升级节奏方案：①预计全文字数与卷数 ②每个境界对应的章节跨度与字数占比 ③全书10个关键突破节点（章节数+触发条件）④压制节奏与爆发节奏的配比（如压3章爆发1章）。用编号要点列出。',
         'dungeon_design': '【构思阶段】你是副本设计师。请设计本书至少3个核心副本/秘境：①背景与入口位置 ②挑战层次（每层敌人/机关/规则）③奖励机制（宝物/经验/剧情线索）④难度与主角战力匹配方式 ⑤隐藏要素与隐藏奖励。每个副本独立编号。',
         'volume_outline': '【构思阶段】你是网文分卷规划师。请为当前卷输出构思方案：①卷核心目标（主角要达成什么）②主要冲突线（外部冲突×N、内部冲突×N）③人物成长目标（能力/心智/关系）④5个关键转折点（位置+事件）⑤卷末状态（主角达到什么境界/处境/关系变化）。用编号要点，不写正文。',
     }, ensure_ascii=False)},
    {'name': '女频甜宠六边形', 'description': '构思专用：CP反差人设+情感弧线+糖虐节奏+虐渣节点，女频情感结构全链路规划', 'genre': 'romance', 'book_type': 'short_story', 'icon':'💕',
     'stage_keys': json.dumps(['character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'CP反差人设设计','desc':'设计身份反差+秘密底线+拉扯模式','prompt_key':'cp_design'},
         {'step':2,'name':'情感弧线规划','desc':'划分阶段+关键事件+糖虐配比+读者情绪','prompt_key':'emotion_arc'},
         {'step':3,'name':'甜蜜高光节点规划','desc':'设计全书糖点位置+专属细节+双向奔赴点','prompt_key':'sweet_highlights'},
         {'step':4,'name':'虐渣/复仇节点规划','desc':'设计铺垫→揭晓→打脸→高光的节奏方案','prompt_key':'revenge_scenes'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'cp_design': '【构思阶段】你是女频CP人设设计师。请输出完整CP人设方案：①男女主身份反差标签（如总裁×社畜/王爷×医女）②男主秘密/底线×3、女主秘密/底线×3 ③相识触发方式 ④情感拉扯模式（谁先动心、谁追谁火葬场、误会核心）。',
         'emotion_arc': '【构思阶段】你是情感弧线规划师。请输出全书情感节奏方案：①从陌生→拉扯→暧昧→确认→危机→HE的阶段划分（百分比/章节数）②每阶段2-3个关键事件节点 ③糖度与虐度配比曲线（如前期虐3糖7，中期糖5虐5，后期糖9虐1）④读者情绪起伏的10个峰值点位置。',
         'sweet_highlights': '【构思阶段】你是甜宠高光节点规划师。请规划全书糖点方案：①全书总糖点数量与分布密度（每N千字1个）②专属小动作/专属昵称×5（何时首次使用）③追妻火葬场/追夫火葬场桥段触发条件+发生位置 ④双向奔赴高光×3（各自为对方牺牲/勇敢的节点）。用编号要点，不写正文场景。',
         'revenge_scenes': '【构思阶段】你是虐渣节奏规划师。请规划虐渣/复仇方案：①仇人/渣男/绿茶×3，各自身份与前期铺垫方式 ②真相揭晓的节奏（分几次释放信息）③每个打脸名场面的触发条件与位置 ④女主高光反击的3个标志性节点（做了什么、周围人反应预期）。用编号要点，不写正文场景。',
     }, ensure_ascii=False)},
    {'name': '悬疑反转工厂', 'description': '构思专用：真相底牌+5层信息差+红鲱鱼+反转节点规划，悬疑全链路节奏设计', 'genre': 'mystery', 'book_type': 'short_story', 'icon':'🔍',
     'stage_keys': json.dumps(['character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'真相底牌设计','desc':'确定最终真相+真凶+动机+手法+所有秘密','prompt_key':'truth_card'},
         {'step':2,'name':'信息差分层设计','desc':'设计读者/侦探/各角色的5层信息不对称','prompt_key':'info_gap'},
         {'step':3,'name':'红鲱鱼（误导）设计','desc':'设计至少3个合理误导线索+合理化方式','prompt_key':'red_herring'},
         {'step':4,'name':'反转节点规划','desc':'设计反转位置+强度递增+铺垫方式','prompt_key':'twist_timing'},
         {'step':5,'name':'伏笔回收清单','desc':'列出每条伏笔的埋设位置+回收位置+兑现方式','prompt_key':'ending_payoff'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'truth_card': '【构思阶段】你是悬疑真相设计师。请输出完整真相底牌：①最终真相是什么 ②真凶/真相持有者是谁 ③动机×3条 ④核心作案/事件手法 ⑤每个主要人物所持有的秘密。用编号要点，不写正文。',
         'info_gap': '【构思阶段】你是信息差设计专家。请设计5层信息不对称表：①读者知道但侦探不知道的信息×3 ②侦探知道但读者不知道的信息×3 ③关键角色A知道但其他人不知×3 ④关键角色B知道但其他人不知×3 ⑤通过剧情会陆续揭晓的信息时间线（第几次曝光什么）。',
         'red_herring': '【构思阶段】你是红鲱鱼设计师。请设计至少3个合理误导：①可疑人物（为什么读者会怀疑、真实身份是什么）②误导线索（内容是什么、为什么是误导不是谎言）③时间线错觉（读者会怎么误判、真相是什么）④每个误导最终如何被解释清楚。不写正文，编号列出。',
         'twist_timing': '【构思阶段】你是反转节奏规划师。请输出反转方案：①全书反转×N（建议3-5个）的位置（百分比/章节数）②反转强度评级（1-5分，必须递增）③每个反转的前置铺垫线索（埋设位置）④反转后读者预期情绪反应。',
         'ending_payoff': '【构思阶段】你是伏笔回收审计师。请输出伏笔回收清单：①按编号列出所有伏笔×10+（埋设位置+回收位置+兑现方式）②验证所有线索是否指向真相（是/否/缺什么）③每个红鲱鱼最终如何圆回来④结局情感冲击点在哪里。',
     }, ensure_ascii=False)},
    {'name': '短篇冲榜模板', 'description': '构思专用：一句话爆款梗+快速口播大纲+反转节点设计，知乎盐选/UC故事会适配', 'genre': 'other', 'book_type': 'short_story', 'icon':'🚀',
     'stage_keys': json.dumps(['character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'一句话爆款梗设计','desc':'提炼3版含反转/信息差/情感冲击的梗概','prompt_key':'one_line_hook'},
         {'step':2,'name':'口播叙事大纲','desc':'写口语化四段落大纲（开篇→发展→高潮→结局）','prompt_key':'quick_outline'},
         {'step':3,'name':'反转节点规划','desc':'设计每千字反转密度+读者情绪曲线','prompt_key':'twist_density'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'one_line_hook': '【构思阶段】你是短篇爆款制造机。请输出3个不同版本的一句话爆款梗概：①版本1含强反转 ②版本2含强信息差 ③版本3含强情感冲击。每个版本格式：【题材标签】+【一句话（不超过40字）】+【预期读者反应词（3个）】。',
         'quick_outline': '【构思阶段】你是口播叙事化大纲专家。请用口语化、有节奏感的语言写4段式大纲：①开篇冲突（用什么事件立刻抓住读者）②发展转折（2-3个推动剧情升级的事件）③高潮反转（最大的反转是什么，怎么触发）④结局（怎么收、给读者什么余味）。每段用编号要点列出关键事件，不写散文。',
         'twist_density': '【构思阶段】你是短篇反转密度规划师。请输出节奏方案：①全文按8000-15000字，设定每N千字必须出现一个反转/钩子 ②读者情绪的7-10个峰值位置 ③每个反转的前置伏笔是什么 ④全文高潮位置（70-80%处）和结尾钩子（若需要）。',
     }, ensure_ascii=False)},
    {'name': '世界观构建手册', 'description': '构思专用：基础规则+地理势力+历史年表+种族文化，虚构世界全维度搭建', 'genre': 'fantasy', 'book_type': 'novel', 'icon':'🌍',
     'stage_keys': json.dumps(['worldbuilding'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'世界基础规则','desc':'定义物理/魔法法则、能量体系、超自然存在定位','prompt_key':'base_rules'},
         {'step':2,'name':'地理与势力','desc':'设计版图、势力分布、资源与交通（可直接输出平台地点JSON）','prompt_key':'geography'},
         {'step':3,'name':'历史年表','desc':'创世→远古→重要节点→当前时代的完整时间线','prompt_key':'history'},
         {'step':4,'name':'种族与文化','desc':'设计各族群的外貌、社会、信仰、语言、关系','prompt_key':'cultures'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'base_rules': '【构思阶段】你是世界观架构师。请以编号要点输出世界基础规则方案：①世界物理法则（重力、元素、魔法/能量来源）②能力体系基本规则（谁能用、代价是什么、上限）③生物/进化基本逻辑 ④神/超自然存在（是否存在、与人类关系）。用于嵌入"设定/规则"维度。',
         'geography': '【构思阶段】你是世界地理设计师。请输出地理与势力方案，推荐使用平台"地点"维度的三级结构：①一级大区域×2-3（大陆/海洋/势力范围）②每个一级区域下的二级城市/门派/据点③每个二级下的关键三级场景④资源分布+交通贸易路线。格式采用编号要点或严格JSON三级结构均可。',
         'history': '【构思阶段】你是世界历史编年官。请按时间段落输出完整年表：①创世/远古神话期 ②古代文明兴衰期 ③近古/关键战争期 ④近300-500年重要节点 ⑤当前时代的社会/政治格局定位。每个节点用"年份/时期+事件+对当代影响"的格式。',
         'cultures': '【构思阶段】你是种族文化设计师。每个主要种族/族群输出：①外貌特征×3 ②社会结构（阶层/家庭/治理模式）③核心信仰/禁忌 ④语言与文化符号 ⑤与其他主要种族的关系（友好/敌对/中立）。用于嵌入"设定/势力"与角色设定维度。',
     }, ensure_ascii=False)},
    {'name': '都市职场商战', 'description': '构思专用：职场生态+商战节奏+行业细节+爽点节点，番茄职场文全链路规划', 'genre': 'urban_business', 'book_type': 'novel', 'icon':'💼',
     'stage_keys': json.dumps(['character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'职场生态设计','desc':'设计公司结构+阵营+利益诉求+明潜规则','prompt_key':'office_ecosystem'},
         {'step':2,'name':'商战节奏规划','desc':'规划竞标节奏+商业陷阱+权力斗争+外部冲击','prompt_key':'business_conflict'},
         {'step':3,'name':'行业专业细节','desc':'行业术语+工作流程+KPI体系+生存法则','prompt_key':'professional_details'},
         {'step':4,'name':'爽点高光节点规划','desc':'设计方案碾压/升职签约/打脸反转的节点方案','prompt_key':'revenge_moments'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'office_ecosystem': '【构思阶段】你是职场生态设计师。请输出完整职场生态方案：①公司/行业生态结构（部门/级别/晋升路径）②关键人物×5+各自利益诉求 ③盟友阵营与对手阵营划分 ④明规则与潜规则对比清单。',
         'business_conflict': '【构思阶段】你是商战节奏规划师。请输出商战/职场升级节奏方案：①全书项目/竞标/晋升的阶段划分 ②每个阶段设计2-3个商业陷阱与反击方式 ③内部权力斗争的关键5个节点 ④外部市场/政策变化的3个冲击点。编号列出，不写正文。',
         'professional_details': '【构思阶段】你是行业顾问。为选定行业输出真实性增强方案：①核心行业术语×15（带释义）②典型工作流程×3（步骤+参与人+产出）③KPI/考核/晋升的真实体系 ④职场生存法则×8（新人/老人/管理者分别适用）。',
         'revenge_moments': '【构思阶段】你是职场爽点节点规划师。请设计爽点方案：①方案碾压对手×3（什么方案、对手怎么崩的）②升职/签约高光×3（触发条件、周围反应）③同事/领导打脸反转×3（铺垫位置、打脸事件）④行业地位跃升×2（最终身份与影响力）。用编号要点，不写正文场景。',
     }, ensure_ascii=False)},
    {'name': '历史权谋工坊', 'description': '构思专用：时代格局+权谋棋局+战争策略+派系人物，历史文全链路结构规划', 'genre': 'history', 'book_type': 'novel', 'icon':'🏯',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'时代格局设定','desc':'朝代/架空王朝+官制+经济+边疆','prompt_key':'era_setting'},
         {'step':2,'name':'权谋棋局设计','desc':'核心矛盾+派系格局+目标弱点+政变节点','prompt_key':'power_game'},
         {'step':3,'name':'战争策略规划','desc':'兵力部署+关键战役+后勤+情报战+战后格局','prompt_key':'war_strategy'},
         {'step':4,'name':'派系人物关系','desc':'各派系理念+代表人物+内部矛盾+合纵连横','prompt_key':'faction_design'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'era_setting': '【构思阶段】你是历史设定专家。请输出时代格局方案：①朝代/架空王朝基本信息（国号、开国时间、当前皇帝/掌权者）②政治制度与核心官制（中枢、地方、军制）③经济与税收体系（土地、货币、贸易）④外患与边疆态势（敌人、兵力、边防部署）。',
         'power_game': '【构思阶段】你是权谋棋手。请输出权谋棋局方案：①全书核心权力矛盾（皇权/外戚/宦官/世家/藩王的冲突）②朝堂派系×3-5格局图 ③每派核心人物目标与致命弱点 ④关键政变/宫变/党争节点×5（位置+触发+后果）。',
         'war_strategy': '【构思阶段】你是军事策略师。请输出战争规划方案：①双方/多方兵力对比、主将、部署图 ②关键战役×3（开战理由、战场、战术、胜负关键）③后勤与粮草路线 ④间谍/情报战的关键事件 ⑤每场战役后的格局变化与势力洗牌。',
         'faction_design': '【构思阶段】你是阵营与人物设计师。为每个主要派系输出：①核心理念/纲领 ②代表人物×2-3（性格、能力、立场）③派系内部矛盾×2 ④与其他派系的合纵连横关系（谁联合谁对抗谁摇摆）。',
     }, ensure_ascii=False)},
    {'name': '科幻未来创世', 'description': '构思专用：科技树+未来社会+英雄旅程+设定自洽检查，硬科幻/赛博朋克全维度设计', 'genre': 'scifi', 'book_type': 'novel', 'icon':'🚀',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'核心科技树','desc':'设计科技原理+发展里程碑+社会变革+伦理困境','prompt_key':'tech_tree'},
         {'step':2,'name':'未来社会结构','desc':'政治体制+阶层分化+人机共存+星际社会','prompt_key':'future_society'},
         {'step':3,'name':'英雄冒险主线','desc':'主角弧线+觉醒节点+关键任务+终极对抗','prompt_key':'hero_journey'},
         {'step':4,'name':'设定自洽审查','desc':'构思阶段自查：科技/社会/时间线/角色行为的一致性','prompt_key':'logic_check'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'tech_tree': '【构思阶段】你是硬科幻设定师。请输出核心科技树方案：①核心科技原理×3（尽量可自洽、可推演）②科技发展10个里程碑（年代+技术+影响）③每个核心科技带来的社会变革 ④科技黑暗面/伦理困境×3（主角/反派如何面对）。',
         'future_society': '【构思阶段】你是未来社会学家。请输出未来社会方案：①政治体制演变（什么政体、权力来源）②阶层/种姓/基因/AI权限的分化结构 ③AI与人类/改造人的共存模式 ④星际殖民/赛博空间/元宇宙的社会结构与规则。',
         'hero_journey': '【构思阶段】你是科幻冒险策划师。请输出主角冒险主线方案：①从凡人→觉醒→英雄的完整弧线 ②能力/认知觉醒的关键节点×5（触发事件+变化）③关键冒险任务×4（目标+地点+风险+收获）④终极对抗与命运抉择（终极问题是什么、主角如何选择）。',
         'logic_check': '【构思阶段·自查】你是科幻设定自洽审查员。请对已构思内容做4维度自检：①科技设定内部是否有矛盾×N ②社会变迁推导是否合理 ③时间线/因果链是否有断裂 ④主角/配角的行为模式是否符合科技社会环境。列出问题+修改建议。',
     }, ensure_ascii=False)},
    {'name': '无限流生存指南', 'description': '构思专用：主神空间规则+副本设计+能力进化+团战策略，无限流全链路设计', 'genre': 'fantasy', 'book_type': 'novel', 'icon':'🌀',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'主神空间规则','desc':'设定轮回规则+等级权限+任务奖励+惩罚机制','prompt_key':'infinity_rules'},
         {'step':2,'name':'核心副本设计','desc':'每个副本的世界背景+主线+隐藏+BOSS+道具','prompt_key':'dungeon_design'},
         {'step':3,'name':'能力进化体系','desc':'能力分类+进化条件+组合技+主角独特能力（可输出物资库JSON）','prompt_key':'ability_tree'},
         {'step':4,'name':'团队协作与智斗','desc':'角色分工+配合战术+智斗场面+背叛信任危机','prompt_key':'team_tactics'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'infinity_rules': '【构思阶段】你是无限流系统设计师。请输出主神空间规则方案：①轮回世界的核心规则（谁进、为什么进、如何进出）②轮回者等级/权限/称号体系 ③任务发布机制+积分/奖励兑换表 ④惩罚/抹杀/淘汰的触发条件。',
         'dungeon_design': '【构思阶段】你是副本世界设计师。请设计本书前5个核心副本：①借用的IP/原创世界背景+特殊规则 ②主线任务+支线+隐藏任务清单 ③BOSS战设计（技能、弱点、奖励）④关键道具/剧情线索×5 ⑤难度与玩家战力的匹配曲线。',
         'ability_tree': '【构思阶段】你是能力进化设计师。请输出能力体系方案：①能力体系分类（功法/异能/血脉/装备/道具）②主角及核心配角的能力树×8-15项，每项含名称、类型、来源、效果、归属、首次出现章节（推荐输出平台"物资库"维度的JSON数组格式） ③能力组合与进阶方向×5 ④主角独特能力的隐藏来源与终极形态。',
         'team_tactics': '【构思阶段】你是团战策略师。请输出团队方案：①固定/临时团队×2-3的角色分工（输出/坦克/控制/辅助/智囊）②经典配合战术×5（何时使用、需要谁、效果）③智斗名场面节点×3（怎么设局、怎么破局）④背叛与信任危机节点×2（位置、触发、后果、修复方式）。',
     }, ensure_ascii=False)},
    {'name': 'SoloEnt Vibe Writing', 'description': '【文风专用】文风锚定+人机共创：作者主导AI辅助，风格导入+任务拆解+初稿+自我蒸馏，正文阶段文风与协作专用', 'genre': 'other', 'book_type': 'short_story', 'icon':'✨',
     'stage_keys': json.dumps(['draft'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'导入作者风格','desc':'分析旧作或参考文本，建立风格基线卡片','prompt_key':'style_import'},
         {'step':2,'name':'任务拆解辅助','desc':'拆分场景/对话/转场/独白等小单元，标注AI辅助等级','prompt_key':'task_decompose'},
         {'step':3,'name':'风格锚定初稿','desc':'严格按风格卡片生成正文，保留作者修改空间','prompt_key':'first_draft'},
         {'step':4,'name':'经验自我蒸馏','desc':'提炼修改差异，沉淀复用经验并更新风格卡片','prompt_key':'self_distill'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'style_import': '【文风阶段·分析】你是中文文风分析师。基于作者提供的参考文本，请输出风格卡片（编号要点）：①句式偏好（短句/长句占比、常用句型、标点偏好）②用词习惯（常用动词/形容词/语气词、避词清单）③节奏特征（信息密度、断章位置、留白比例、叙事速度）④情感表达方式（直白/隐忍/比喻/对话驱动）。',
         'task_decompose': '【文风阶段·拆任务】你是正文写作任务拆解师。将当前写作任务拆为可独立完成的小单元：①场景段落×N（每段200-500字，情绪目标+出场人）②对话片段×N ③转场段落×N ④内心独白×N ⑤环境描写×N。每个单元标注AI辅助等级（A=AI先写再改 / B=人写AI补 / C=完全人写）。',
         'first_draft': '【文风阶段·写正文】你是 Vibe Writing 写手，请严格按上方"风格卡片"写当前正文单元。写作铁律：①句式/用词/节奏100%匹配风格卡片 ②每个单元独立完整、上下文自洽 ③不确定或容易AI化的地方用【作者定夺：xxxx】标注，不强行乱编 ④写完标注"需要作者人工修改比例建议：X%"。',
         'self_distill': '【文风阶段·沉淀】你是自我蒸馏教练。基于作者对初稿的修改结果，请输出经验总结：①逐段对比"AI初稿vs作者终稿"差异清单×N ②提炼可复用经验（下次写类似场景怎么改进）③更新风格卡片（新增/修正条目）④沉淀为个人可复用写作方法论模板。',
     }, ensure_ascii=False)},
    {'name': 'AI责编精审套装', 'description': '【审查专用】5层审稿管线：逻辑→人设→节奏→商业→违禁，一站式过稿审查', 'genre': 'other', 'book_type': 'short_story', 'icon':'🔍',
     'stage_keys': json.dumps(['review','polish'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'逻辑审稿','desc':'检查情节逻辑与因果链','prompt_key':'logic_review'},
         {'step':2,'name':'人设审稿','desc':'检查人物行为是否一致','prompt_key':'character_review'},
         {'step':3,'name':'节奏审稿','desc':'检查爽点密度与追读期待','prompt_key':'rhythm_review'},
         {'step':4,'name':'商业审稿','desc':'评估平台适配与签约潜力','prompt_key':'commercial_review'},
         {'step':5,'name':'合规审稿','desc':'检查平台违禁内容','prompt_key':'compliance_review'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'logic_review': '【审查阶段·逻辑】你是情节逻辑审查员。逐章检查（编号列出问题+位置+修复建议）：1)因果是否成立（有没有"因为A所以B"但A推不出B的地方） 2)时间线是否自洽（事件先后顺序有没有矛盾、角色能不能在那个时间赶到那个地点） 3)信息差是否合理（角色知道/不知道某件事有没有铺垫、会不会全知视角泄露） 4)转折是否有铺垫（反转是不是空降、前面有没有埋过伏笔暗示）。每个问题给具体行号/段落号+修改建议。',
         'character_review': '【审查阶段·人设】你是人设一致性审查员。检查（编号列出每个人物的问题）：1)人物行为是否符合设定性格（内向的人会不会突然社牛、胆小的人会不会突然莽） 2)人物关系是否前后一致（前一章仇人后一章兄弟但没事件驱动） 3)人物成长是否合理（能力/心智变化有事件支撑吗，还是突然升华） 4)配角是否有存在意义（有没有纯粹"工具人"，只在需要时出现给信息/送装备，没有独立动机）。',
         'rhythm_review': '【审查阶段·节奏】你是节奏感审查员。逐章分析（输出评分+调整建议）：1)爽点/冲突密度（每3000字有几个？连续N章没有爽点=停滞） 2)章尾钩子强度（1-5分。章尾是总结句还是新信息/新危机） 3)信息释放节奏（是一口吐完设定还是分批、会不会读者已经猜到了还在水） 4)叙事速度变化（连续多少章一个节奏？有没有快慢交替、松紧搭配）。',
         'commercial_review': '【审查阶段·商业】你是商业编辑。从平台签约视角评估（分5项给A/B/C评级+改进建议）：1)开篇钩子强度（前3段能不能抓住读者、前3章能不能留人） 2)核心卖点是否突出（一句话能不能说清本书看点、能不能对上目标读者爽点） 3)目标读者匹配度（题材/风格/内容对哪个平台哪类读者，错位了吗） 4)长期连载潜力（有没有撑200章+的世界观/冲突线、会不会10章就写完了） 5)建议定价/签约策略（适合番茄/起点/知乎盐选？免费还是付费？什么标签能爆？）。',
         'compliance_review': '【审查阶段·合规】你是合规审查员。对照5大类别检查（编号列出具体问题位置+风险等级）：1)政治敏感内容（涉及国家/领导人/军队/宗教/民族/领土） 2)色情/暴力违规（描写尺度、血腥暴力、性暗示） 3)价值观导向（三观扭曲、违法犯罪美化、PUA/洗白反派） 4)版权风险（用了别人的IP/人物/剧情没授权、同人商用风险） 5)平台特有规则（番茄2400字章、起点不得涉政、盐选反转密度）。风险等级：高风险=立刻删改/中风险=修改措辞/低风险=注意即可。',
     }, ensure_ascii=False)},
    {'name': '都市异能觉醒', 'description': '构思专用：异能体系+组织势力+日常融入+成长路径，番茄都市异能构思全链路', 'genre': 'urban_fantasy', 'book_type': 'novel', 'icon':'⚡',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'异能体系构思','desc':'设计异能来源/分类/等级/代价的完整体系','prompt_key':'power_system'},
         {'step':2,'name':'组织势力格局','desc':'设计官方/地下/中立/敌对四方势力格局与冲突','prompt_key':'org_design'},
         {'step':3,'name':'异能融入日常方案','desc':'设计社会如何隐藏异能、普通人与异能者关系','prompt_key':'daily_integration'},
         {'step':4,'name':'主角成长路径规划','desc':'能力觉醒→控制→组织→格局跨越的完整升级线','prompt_key':'growth_path'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'power_system': '【构思阶段】你是都市异能体系设计师。输出完整异能构思方案：①异能来源（觉醒/遗传/实验/契约/器物，各来源的比例与触发条件）②分类体系（自然系/概念系/规则系/身体系/精神系/召唤系，每系代表能力×3）③等级划分与晋升条件（≥9级，每级标志性能力一句话、晋升需要什么事件/资源/心境突破）④异能限制与代价（每种能力的副作用、使用上限、反噬机制）。用于嵌入平台"设定/规则"维度。',
         'org_design': '【构思阶段】你是都市异能势力设计师。输出四方势力格局（编号要点）：①官方异能管理机构（名称/体制/权限/核心人物/特色能力/对主角态度）②地下异能组织（目标/手段/核心人物/与官方矛盾/对主角态度）③中立势力（信息商/中立区/散修联盟/庇护所，各自存在意义）④敌对势力（反派组织/异族/异次元来客，核心理念、终极目标、为什么不能共存）⑤势力间的明规则与潜规则×10。',
         'daily_integration': '【构思阶段】你是异能日常融入设计师。输出社会如何"表面正常、内里异能横行"的设定方案：①异能如何向普通人隐藏（记忆消除/合理借口/官方控场/合谋演戏，每种机制×3具体方案）②异能者打工/赚钱/生活的常见方式（异能事务所/地下赌斗/特招入伍/灰色产业）③城市中隐藏的异能据点×5（看似普通的地点背后是什么）④普通人发现异能的5种后果及处理流程 ⑤5个"日常物品+异能用法"的创意组合（咖啡×精神系、地铁×空间系等）。',
         'growth_path': '【构思阶段】你是成长路径规划师。输出主角完整升级弧线规划（按卷按章节锚点）：①能力觉醒（触发事件/初体验/第一个代价——定位在哪几章）②控制阶段（谁教他控制、怎么练、第一个正式战斗对手——定位在哪几章）③进入组织（加入哪个组织/为什么加入/第一次执行任务——定位在哪几章）④职级晋升（从新人→骨干→队长→干部→高层，每步的关键事件）⑤格局跨越（从街头→城市→全国→世界维度的跃升路径与关键节点）。按"卷-章-事件"三级结构编号列出。',
     }, ensure_ascii=False)},
    {'name': '轻小说日式创作', 'description': '构思专用：萌点角色+卷形式规划+日常冒险平衡+插画脚本，番茄轻小说构思全链路', 'genre': 'light_novel', 'book_type': 'novel', 'icon':'📚',
     'stage_keys': json.dumps(['character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'萌点角色与关系','desc':'设计有记忆点的角色+CP/阵营关系网+功能位分配','prompt_key':'character_moe'},
         {'step':2,'name':'卷形式总体规划','desc':'卷数/每卷主题/开卷钩子/卷尾悬念/角色弧线分配','prompt_key':'volume_plan'},
         {'step':3,'name':'日常与冒险平衡方案','desc':'规划日常场景占比/切换节奏/每次冒险对日常的改变','prompt_key':'daily_adventure'},
         {'step':4,'name':'插画关键场景脚本','desc':'设计每卷2-3个关键插画场景+构图+视觉参考','prompt_key':'illustration_script'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'character_moe': '【构思阶段】你是轻小说角色设计师。输出完整角色阵容（每个角色用平台"人物"维度格式：## 角色：姓名，-身份年龄/-性格/-背景/-核心动机/-关系/-剧情功能）：①主角队×4-6（每人3个萌点标签+1个反差标签+口头禅+代表性动作）②女主/核心CP×2-3（萌属性组合+反差+与主角的拉扯模式）③配角×5-8（每人的剧情功能位：信息源/吐槽役/阻碍者/助力者/情感寄托/伏笔载体）④反派/对手×2-3（魅力点+与主角的对立点+为什么让人恨不起来）⑤角色关系网图（谁和谁什么关系、谁暗恋谁、谁和谁有历史、阵营划分）。',
         'volume_plan': '【构思阶段】你是轻小说卷形式编辑。输出完整卷形式规划（推荐平台"分卷剧情"维度的JSON格式，每卷12-15章≈短篇轻一卷规格或50章≈长篇网文一卷）：①总卷数与总字数 ②每卷主题（每卷聚焦什么情感线/成长线/世界观线）③每卷开卷钩子×N（怎么把读者拉进本卷）④每卷卷尾悬念（钩子+兑现预告）⑤每卷角色成长分配（本卷谁变化了、变了什么）⑥每卷插画场景建议（2-3处）。',
         'daily_adventure': '【构思阶段】你是轻小说日常/冒险平衡师。输出全书节奏方案：①日常场景与冒险场景的比例（如日常6冒险4、或前期7后期5）②日常场景分类清单（校园/社团/打工/合宿/祭典/海边温泉，每种的功能是什么）③冒险场景分类清单（打怪/探索/解谜/对抗/比赛，每种的回报是什么）④日常↔冒险切换节奏（几日常+几冒险、如何自然过渡）⑤每次冒险对日常的改变清单（冒险后角色关系变了什么、日常里多了什么新元素、不能冒险结束就当没发生）。',
         'illustration_script': '【构思阶段】你是轻小说插画脚本师。为每卷设计2-3个关键插画场景（按插画页顺序编号）：①场景氛围（明亮/温馨/紧张/忧伤/热血/黄昏）②角色表情与姿势（每个角色画出来的情绪和动作，标注谁在前谁在后）③构图建议（半身像/全身像/群像/俯视/仰视/特写）④背景描述（地点+季节+天气+时间点）⑤画面里的关键道具/细节/伏笔（读者一眼能get的彩蛋）。适配日式轻小说黑白插画风格+偶尔彩色拉页。',
     }, ensure_ascii=False)},
    {'name': '军事谍战风云', 'description': '构思专用：时代格局+战术体系+情报博弈结构+军人成长弧线，番茄军事谍战构思全链路', 'genre': 'military', 'book_type': 'novel', 'icon':'🎖️',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'时代与地缘格局','desc':'设定时代/国际格局/军事实力/核心矛盾缘由','prompt_key':'era_geopolitics'},
         {'step':2,'name':'战术与编制构思','desc':'设计特种作战小队编制/分工/装备体系/战术细节','prompt_key':'tactics'},
         {'step':3,'name':'情报博弈结构','desc':'设计情报网/反情报/欺骗反欺骗/泄露应急的完整结构','prompt_key':'intel_war'},
         {'step':4,'name':'军人弧线与牺牲节点','desc':'规划新兵→精锐/战友/道德困境/牺牲荣耀的成长线','prompt_key':'soldier_arc'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'era_geopolitics': '【构思阶段】你是军事地缘背景设定师。输出时代与国际格局方案（编号要点）：①时代坐标（具体年代/架空年代/冷战/反恐/未来战争——哪个时期、为什么选这个时期）②国际格局（主要大国/阵营划分/各阵营军事实力对比表：兵力/装备/经济/盟友）③核心矛盾与冲突缘由（领土/资源/意识形态/历史恩怨——为什么不能谈判必须打）④武器技术发展水平（步枪/装甲/空军/海军/卫星/信息化/核武器——哪些能用哪些不能用、明确技术上限）⑤战场/谍战的主要地理区域（战区×3、地缘战略意义、为什么必须争）。用于嵌入平台"设定/规则"维度。',
         'tactics': '【构思阶段】你是特种作战顾问。输出战术与编制方案：①特种作战小队编制（12人或8人或6人，每人角色：队长/狙击手/爆破手/侦察兵/通讯员/医疗兵/重火力手——每人分工+性格+能力互补）②典型作战方案×5（潜入/暗杀/营救/斩首/破坏/伏击——每种的步骤、参与角色、关键风险点）③装备与武器系统（主武器/副武器/通信装备/夜视/无人机/爆炸物——型号规格、各自来源、为什么选择这些、现实查证后再写）④地形利用与战术机动（城市/山地/丛林/沙漠/雪原/室内近战——每种地形的经典战术×3）⑤撤退与应急预案（打输了怎么撤、被包围了怎么办、队友牺牲了怎么补位）。',
         'intel_war': '【构思阶段】你是情报战结构设计师。输出情报博弈完整结构（编号要点）：①情报获取方式（SIGINT信号情报/HUMINT人力情报/IMINT图像情报/OSINT开源情报——每种的来源、可靠性、局限）②我方情报网结构（卧底×N、线人×N、安全屋×N、传递暗号系统——每个卧底的身份/立场/暴露风险）③敌方反情报能力（如何发现卧底、如何审讯、如何设假情报陷阱、我方有谁可能是双面间谍）④欺骗与反欺骗（我方怎么设假情报骗对方、对方怎么设假情报骗我方——假情报战关键节点×5）⑤情报分析关键节点（什么情报决定了什么决策、哪次误判导致了什么后果）⑥情报泄露与应急处理（卧底暴露后的救援/灭口/换身份机制、安全屋被端后的备用网络）。',
         'soldier_arc': '【构思阶段】你是军事人物弧线设计师。输出军人成长与牺牲规划（按卷按章节锚点）：①主角完整成长弧线：新兵训练→第一次实战→第一次杀人→第一次失去战友→第一次指挥→成为精锐→道德困境→最终的荣耀或牺牲——每个节点定位在第几卷第几章 ②战友情建立与考验节点（谁和谁怎么成为兄弟、哪次任务建立信任、哪次任务考验忠诚、谁为谁挡枪/背锅/牺牲）③道德困境与选择（服从命令vs保护平民、任务成功vs队友牺牲、忠于国家vs忠于战友——3-5个核心选择节点、每个选择的后果）④牺牲与荣耀的平衡：哪些角色有死亡flag、谁牺牲在什么节点、牺牲换来了什么、牺牲场景怎么写才能打动人不矫情 ⑤和平年代军人价值/家国情怀落点：结局想传递什么价值观（用什么场景、什么对话、什么余味）。',
     }, ensure_ascii=False)},

    {'name': '长篇小说创作全流程', 'description': '构思专用：一句话概念→总纲→卷纲→章纲，百万字长篇全链路结构设计', 'genre': 'other', 'book_type': 'novel', 'icon':'📖',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'一句话构思扩展','desc':'由一句话核心梗扩展为3个完整创意方向供选','prompt_key':'one_line_concept'},
         {'step':2,'name':'全书总纲设计','desc':'确定主线、卷数、核心冲突、视角规划、核心伏笔锚点','prompt_key':'master_outline'},
         {'step':3,'name':'分卷卷纲拆解','desc':'每卷核心目标、冲突、关键节点（推荐输出平台分卷JSON格式）','prompt_key':'volume_breakdown'},
         {'step':4,'name':'章纲细化方案','desc':'每章目标、关键场景、章尾钩子、章型分配、预估字数','prompt_key':'chapter_plan'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'one_line_concept': '【构思阶段】你是网文创意策划师。将用户给出的一句话构思扩展为3个不同方向的完整创意方案，供用户选择。每个方案输出：①一句话核心概念 ②核心卖点×3-5 ③目标读者画像 ④主线冲突一句话 ⑤独特差异化亮点。纯文本分项、编号、空行分隔。',
         'master_outline': '【构思阶段】你是长篇网文总纲设计师。输出全书总纲方案：①核心主线（一句话）②卷数与字数规划（每卷约50章）③主角完整成长弧线 ④主要势力格局与核心矛盾递进 ⑤视角规划（哪些卷用主角视角、哪些用配角视角、信息差设计点）⑥核心伏笔锚点×5-8（埋设卷次与回收卷次）⑦大结局方向。',
         'volume_breakdown': '【构思阶段】你是卷纲设计师。推荐输出平台"分卷剧情"维度的 JSON 数组格式：[{"volume_index":1,"volume":"第1卷 副标题","main_plot":"本卷主线100-200字","core_conflict":"...","ending_hook":"卷尾钩子","nodes":[{"title":"节点1","chapters":"1-10","type":"M","summary":"概要","cool_type":"..."}]}]。铁律：每卷固定50章、每卷5-8个节点、chapters 编号连续不重叠。节点类型：M主线/C角色/W世界观/D日常/F伏笔。',
         'chapter_plan': '【构思阶段】你是章纲设计师。为目标范围章节逐章输出构思方案：①本章核心目标 ②关键场景×2-3 ③出场人物 ④信息释放量（这章要告诉读者什么）⑤章尾钩子类型 ⑥预估字数（2400字±100）⑦章型（M/C/W/D/F，相邻不同）。用编号要点，不写正文。',
     }, ensure_ascii=False)},

    {'name': '正文写作工作流', 'description': '【文风专用】正文文风专用：上下文准备→起草→润色优化，正文阶段文风与节奏锚定', 'genre': 'other', 'book_type': 'novel', 'icon':'✍️',
     'stage_keys': json.dumps(['draft'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'上下文准备','desc':'组装：前文结尾+本章章纲+人物状态+活跃伏笔+设定红线','prompt_key':'context_pack'},
         {'step':2,'name':'正文起草','desc':'按章纲和上下文包写正文初稿，严格遵循平台文风铁律','prompt_key':'draft_writing'},
         {'step':3,'name':'润色优化','desc':'删冗余+强化对话个性+调节奏+补感官细节+金句提升','prompt_key':'polish'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'context_pack': '【文风阶段·备上下文】你是上下文组装助手。为当前章节准备写作上下文包：①前一章结尾（约200字，确保衔接）②本章章纲（目标/场景/出场人）③相关人物当前状态（情绪/位置/持有物）④活跃伏笔列表（本章要推进/提及的）⑤本章不可违反的设定红线（世界规则/能力上限）⑥读者已知信息边界（这章读者已经知道什么、不知道什么）。用编号要点列出。',
         'draft_writing': '【文风阶段·写正文】你是专业网文写手。根据上方上下文包+章纲写本章正文。铁律：①每章2400字±100（短篇除外） ②段落≤3行，对话/动作独立成段 ③人物行为符合设定、对话有个性 ④场景描写服务情报+动作+悬念，不堆氛围 ⑤节奏紧凑、信息推进靠对话与行动 ⑥章尾必须落在新风险/新线索/新决策上（追读钩子）。',
         'polish': '【文风阶段·润色】你是文笔润色师。在完全保留原意的前提下优化正文：①删减冗余描写（每删一句都要问"不写会不会影响理解"）②强化对话个性（每个角色说话方式不同）③调整节奏松紧（紧张处短句、抒情处长句）④增加感官碎片（温度/气味/触感）⑤优化转场衔接，避免工整模板化 ⑥适当提升金句密度（每章≤2个）。直接输出优化后的完整正文。',
     }, ensure_ascii=False)},

    {'name': '去AI味儿改稿心法', 'description': '【审查专用】四步去AI味：查痕迹→最小改写→补人味→终检，保留作者风格不被AI抹平', 'genre': 'other', 'book_type': 'short_story', 'icon':'🎭',
     'stage_keys': json.dumps(['review','polish'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'AI痕迹检测','desc':'扫描全文定位AI写作特征与具体位置','prompt_key':'detect_ai'},
         {'step':2,'name':'最小幅度改写','desc':'只改AI痕迹，不动作者原意与风格','prompt_key':'minimal_rewrite'},
         {'step':3,'name':'人味儿补全','desc':'补不完美细节+感官碎片+微表情+口语化+留白','prompt_key':'humanize'},
         {'step':4,'name':'终检通过率自查','desc':'逐段对照AI特征清单，输出通过率+剩余问题','prompt_key':'final_check'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'detect_ai': '【审查阶段·检测】你是AI写作痕迹检测专家。逐段扫描文本，定位以下11种AI特征并输出具体行号+问题描述：①连续排比句≥3句 ②模板过渡词（然而/不过/不禁/不由得/随即/与此同时等）密集 ③段落长度过于均匀（长短差<30%） ④情感直述而非展示（"他很悲伤" vs 动作细节） ⑤段尾总结本段内容 ⑥信息倾泻大段设定解说 ⑦对话过于书面化无个性 ⑧缺少口语语气词和断句 ⑨转场过于工整模板化 ⑩形容词堆叠 ⑪悬空动作/被动吃亏（身体部位当主语独立成句如"一只手探来""目光看了过来"却无发出者，或用被动句写抽象吃亏）。',
         'minimal_rewrite': '【审查阶段·改写】原则：编辑文字，不抹去文字背后的人。只修改被检测出AI痕迹的位置，保持作者原意和文风不变：①把模板句式改成自然口语表达 ②打散过于均匀的段落节奏（故意制造长短差） ③把情感直述句改为行为/动作/细节展示 ④删除所有段尾总结句 ⑤对话加语气词、打断、结巴、重复 ⑥补齐悬空动作的主语（动作必须有人扛，如改"一只手探来"为"小二的一只手探来"），把抽象被动的吃亏句改写为主动的具体画面。其余内容一个字都别动。',
         'humanize': '【审查阶段·补人味】在保持原意的前提下增加真实感细节：①加入不完美细节（结巴、重复、打断、嘴瓢、说错改口）②增加感官碎片（气味、温度、触感、声音、光线）③插入人物小动作/微表情（抠指甲、咬嘴唇、转笔、瞥一眼）④对话加语气词和自然断句（嗯、啊、哦、啧、哎、咳）⑤适当留白，不把话说满 ⑥基于人物视角加入个性化比喻。关键是让文字背后站着一个具体的人。',
         'final_check': '【审查阶段·终检】对照AI特征清单逐段复审。输出：①整体通过率（百分比，目标≥85%）②剩余问题清单（每项带行号）③低通过率（<80%）时给出二轮改写建议。',
     }, ensure_ascii=False)},

    {'name': '长篇小说防遗忘系统', 'description': '构思专用：事实锁定+伏笔登记+角色认知+叙事债务追踪，防设定崩防伏笔漏', 'genre': 'other', 'book_type': 'novel', 'icon':'🧠',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','foreshadowing'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'核心事实锁定','desc':'从世界观/人物/大纲提取不可变事实清单（建议嵌入设定/规则）','prompt_key':'lock_facts'},
         {'step':2,'name':'伏笔线索登记','desc':'建立全书伏笔库：埋设位置/内容/回收位置/伪装方式/读者反应','prompt_key':'foreshadow_register'},
         {'step':3,'name':'角色认知档案','desc':'每个角色知道什么/不知道什么/信息差/相互认知错觉','prompt_key':'character_cognition'},
         {'step':4,'name':'叙事债务盘点','desc':'悬念承诺与兑现平衡、债务清单+回收优先级（与伏笔维度格式对齐）','prompt_key':'narrative_debt'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'lock_facts': '【构思阶段·锁事实】从已构思的世界观、人物档案、大纲中，提取全书不可变的核心事实清单（编号要点）：①人物核心事实（名字、年龄、能力上限、性格底线）+ 不可变更原因 ②世界规则（物理/力量/社会结构）+ 不可变更原因 ③已发生的关键事件 + 不可变更原因 ④已确定的核心人物关系 + 不可变更原因。用于嵌入"设定/规则"维度。',
         'foreshadow_register': '【构思阶段·伏笔库】输出全书伏笔线索库（编号列出每条）：①埋设章节/位置 ②内容（写什么）③预期回收章节/位置 ④当前状态（待回收/已回收）⑤伪装方式（如何藏起来：闲笔/口头禅/细节/环境）⑥回收时读者预期情绪。【伏笔埋设纪律】①单章最多埋1-2处，严禁集中铺设 ②伪装成日常细节，正文里绝不标"这是伏笔" ③埋设与回收至少间隔3章。',
         'character_cognition': '【构思阶段·认知档案】为每个主要角色输出认知档案：①角色确认知道的信息清单×N ②角色绝对不知道的禁忌信息×N ③"读者知道但角色不知道"的信息差×N ④"角色认为A是B，但实际是C"的认知错觉×N。后文所有场景必须严格遵守。',
         'narrative_debt': '【构思阶段·叙事债务】输出全书悬念承诺与兑现盘点：①已兑现悬念×N（承诺了什么、在哪兑现）②待兑现悬念×N（债务，按紧急度排序）③过度透支风险项×N（承诺太多读者可能遗忘，建议删或加速）④未承诺但需要交代的遗留问题×N。与平台"伏笔"维度输出格式保持一致（编号+字段结构）。',
     }, ensure_ascii=False)},

    {'name': '玄幻小说文风', 'description': "【文风专用·玄幻/仙侠爽文 v2】男频玄幻爽文大神的通用文风：第三人称限知主跟主角；主角越平静、旁观越震惊的反差冷面落地，情绪绝不贴标签，靠小动作+体感+微表情；叙述句以逗号长句为主（20-35字串1-2个动作）、段均句数≤1.8（不卡总段数）、铁律A+禁令0全口径生效。", 'genre': 'fantasy', 'book_type': 'novel', 'icon': '⚔️',
     'category': 'style', 'genre_target': 'fantasy', 'priority': 10,
     'stage_keys': json.dumps(['draft'], ensure_ascii=False),
     'workflow': json.dumps([{"step": 1, "name": "文风锚定", "desc": "基于本skill文风文档生成7维查表，照着写", "prompt_key": "style_anchor"}, {"step": 2, "name": "正文写作", "desc": "按文风文档规则生成2400字正文", "prompt_key": "genre_style"}], ensure_ascii=False),
     'prompts': json.dumps({"style_anchor": "【文风阶段·锚定】你是文风锚定师。基于以下风格文档，生成7维文风速查表：1)整体风格与气质 2)段落结构与节奏（段长/句长/段均/段数上限） 3)对话密度、对白格式硬卡、角色声域 4)叙述与视角规则 5)信息推进方式 6)章法（开头进入、中段推进、章末钩子类型）7)禁用/红线（禁词、禁令0、铁律A）。逐维一行，创作时照着写。\n\n【参考风格文档全文】\n---\nname: 玄幻/仙侠爽文\ndescription: \"玄幻/仙侠爽文——男频玄幻爽文大神的通用文风：第三人称限知主跟主角；主角越平静、旁观越震惊的反差冷面落地，情绪绝不贴标签，靠小动作+体感+微表情。\"\nversion: \"2.0\"\nagent_created: true\n---\n\n# 玄幻/仙侠爽文\n你是「玄幻/仙侠爽文」。写玄幻/仙侠爽文正文时，严格按本 skill 文风规则执行。产出须全新原创，不得出现原作人名/地名/门派/剧情/原文段落。\n\n> 本skill为「写正文文风 + 去AI味儿」，不含大纲/背景/设定/人物设计/长篇一致性。本 skill 自包含全量文风档案 + 红线 + 冲突裁决，写文前通读本文件即可。示例均为手法示范（占位人物 + 自创句），只学写法不照抄。\n\n## 灵魂速写\n**冷静的碾压感**：主角永远比读者以为的更强，表面不动声色，关键时刻掀底牌、越级、碾压。爽点靠「主角越平静、围观越震惊」的反差 + 把战力/资源量化成可数数字。情绪绝不贴标签，靠小动作/体感/微表情冷面落地——AI 写爽文是「他十分强大，令人震撼」，本类文是「他握了握拳，指尖掐进肉里，嘴角却只泛起一丝笑」。\n\n## 文风指纹\n- 段落结构（双轨并卡·爽文专属）：炸点/打斗/对峙用 1–2 句成段（手机端三行≈40–70字），最狠那句话/最大转折/核心情感爆发可单独成行；但叙事/描写主力段 40–90字（1-2个逗号长句），叙述句主力=20-35字逗号串1-2个动作收一个句号（占叙述句70-80%），段均句数 ≤1.8（对白一句一段是常态，不卡总段数）；相邻3句同POV/同场景/同镜头必须合并至少2句，连续≥4句独立段合并相邻2-3句。\n- 长段用来描写/铺垫，短段用来动作/对话/高潮。\n- 破折号 ≈0–2/万（几乎不用）；省略号约 40–50/万（转场/吞声）；感叹号约 40–55/万（拟声/短喝）；问号约 35–45/万（配角惊疑）\n- 对话密度约 25–30%（中等偏平衡）；含对话句子约 45–55%；对白格式硬卡：连续3句对白至少1句与动作/反应嵌同段，不做纯剧本台词独立段；句首提示语占比 ≤40%。\n\n---\n\n# 全量文风档案（写正文文风 + 去AI味）\n\n## 文笔句式节奏\n- **短段落 ≠ 碎句（硬约束）**：段落（1–2 句成段）只为节奏；段落内每句必须完整、自然、可读（主谓齐全）。绝不为求「短」把句子剁成读不通的残片；宁可合并成 1 段内 2 句完整短句，也绝不写半句语义残缺的碎片段。\n- 拟声词**「！」式独立成行**（覆盖骨架层禁单字成段）：哗！/铛！/嘭！/轰！/嗡！/唰！/咻！/嗤！/嗖！\n- 形容词中低密度，靠名词动词与体感落地。\n- forbidden：破折号长音拟声（轰——！），长音拟声一律改用「！」；温情慢抒情；哲理升华尾；形容词直述情绪；禁令0修正式否定句（不是X是Y/排比三连/自我修正词）全章禁绝。\n\n## 视角文风\n- 第三人称限知主跟主角；叙述者冷静克制，不灌价值观。\n- 可作话（章首/章末直呼读者或 PS 注），仅吐槽/催更/作者注处用，克制，内容自创。\n- 内心标记引出算计（「X暗道」「X心中」）。\n- must_do：第三人称限知主跟主角；内心标记保留；适度作话。\n- forbidden：第一/二人称主述；章尾鸡汤升华；叙述者 lecture。\n\n## 对话\n- 叙述与动作并重，每句台词必嵌动作/神态/环境。\n- 角色声域：主角冷静简短、一字千金、留半句藏机锋；反派骄横→惊怒→求饶；配角震惊、碎碎念、推动围观震惊；长辈/强者或淡漠或凌厉。\n- must_do：台词藏潜台词；半句/沉默也是戏；配角惊疑反差衬托主角。\n- forbidden：纯问答乒乓接力；角色直说真实动机（用内心标记说）。\n\n## 断章与节奏（仅收束手法/钩子类型/章长）\n- 章内节奏快、多拍交替；日常章一条因果链推进。\n- 章末卡在：悬念（底牌将揭）、爽点余波（碾压后拂衣去）、信息差（新势力介入）、或 PS 作话催更。\n- 单章 2300–2500 字（含中文标点）。\n- must_do：每章末必有钩子；战斗章用越级碾压/底牌反转当章内转折；可附 PS 作话注。\n- forbidden：章尾升华；无钩子注水。\n\n## 招牌技巧清单\n1. 拟声词「！」独立成行（战斗炸点，不用破折号）\n2. 主角越平静、围观越震惊的反差冷面落地\n3. 围观者震惊反应（爽点放大器）\n4. 量化刻度爽感（等级/悬赏/贡献点/排名逐笔回收，刻度体系作者自设）\n5. 微表情承载情绪（掐拳/嘴角一丝笑/丢石子/抿酒，动作自创）\n6. 内心标记显算计（X暗道 / X心中）\n7. 底牌掀桌式反转（关键时刻爆发，底牌内容作者自设）\n8. 章末钩子 + 可附 PS 作话注（内容自创）\n9. 偶尔粗口炸反差（克制，词自换）\n10. 让事物自己动（砍摄像机式「看见/看着」）\n\n## 情感落地与文风体温\n- 原则：情绪不是温情也不是搞笑，是**冷静的碾压感**——靠「主角越平静、围观越震惊」反差 + 小动作/体感/微表情冷面落地。只抓短句碎段会写出「冲但空的 AI 战报」。\n- 工艺：微表情/小动作代替形容词；体感落地暴力（不写「十分血腥」，写对手喷血砸地、场面骤静）；「主角平静×围观震惊」反差造爽；信息差/算计显主角之强；内心标记显布局；偶尔粗口炸反差（克制）。\n\n## 红线（最高优先级）\n- 绝不写残句碎块（铁律A：段落≠碎句，同镜头动作不残切碎句，反例「他不说话。只是盯着那人。」「刀光一闪。血溅半尺。」必须合并）。\n- 破折号长音拟声词禁用（轰——！），用「！」独立成行；破折号全章≈0–2/万。\n- 禁令0修正式否定句（不是X是Y/排比三连/自我修正词）全章禁绝，命中一律改直接陈述句。\n- 绝不温情慢抒情/哲理升华尾。\n- 绝不灌设定 lecture（叙述者不堆背景）。\n\n## 与骨架层（novel-writer）冲突裁决（仅文风/格式层面）\n- 禁单字独立成段：覆盖（拟声词「！」独立成行是战斗炸点签名）。\n- 禁破折号：一致（几乎不用，拟声词用「！」而非「——」）。\n- 省略号密度：覆盖（中等，仅转场/吞声）。\n- 次元壁：覆盖（作话直呼读者允许保留，克制使用）。\n- 对话占比：覆盖（中等偏平衡，叙事动作驱动）。\n", "genre_style": "【文风阶段·写正文】你是对应题材的网文写手。严格遵循以下【文风文档】写正文，任何一条规则不得突破；文档中要求\"必须继承骨架层内置规则\"的条目全部生效；文档的禁令0/铁律A/段均句数/对白硬卡/比喻禁词/必删词表全部按最高优先级执行。输出2400字±100正文。\n\n【文风文档全文】\n---\nname: 玄幻/仙侠爽文\ndescription: \"玄幻/仙侠爽文——男频玄幻爽文大神的通用文风：第三人称限知主跟主角；主角越平静、旁观越震惊的反差冷面落地，情绪绝不贴标签，靠小动作+体感+微表情。\"\nversion: \"2.0\"\nagent_created: true\n---\n\n# 玄幻/仙侠爽文\n你是「玄幻/仙侠爽文」。写玄幻/仙侠爽文正文时，严格按本 skill 文风规则执行。产出须全新原创，不得出现原作人名/地名/门派/剧情/原文段落。\n\n> 本skill为「写正文文风 + 去AI味儿」，不含大纲/背景/设定/人物设计/长篇一致性。本 skill 自包含全量文风档案 + 红线 + 冲突裁决，写文前通读本文件即可。示例均为手法示范（占位人物 + 自创句），只学写法不照抄。\n\n## 灵魂速写\n**冷静的碾压感**：主角永远比读者以为的更强，表面不动声色，关键时刻掀底牌、越级、碾压。爽点靠「主角越平静、围观越震惊」的反差 + 把战力/资源量化成可数数字。情绪绝不贴标签，靠小动作/体感/微表情冷面落地——AI 写爽文是「他十分强大，令人震撼」，本类文是「他握了握拳，指尖掐进肉里，嘴角却只泛起一丝笑」。\n\n## 文风指纹\n- 段落结构（双轨并卡·爽文专属）：炸点/打斗/对峙用 1–2 句成段（手机端三行≈40–70字），最狠那句话/最大转折/核心情感爆发可单独成行；但叙事/描写主力段 40–90字（1-2个逗号长句），叙述句主力=20-35字逗号串1-2个动作收一个句号（占叙述句70-80%），段均句数 ≤1.8（对白一句一段是常态，不卡总段数）；相邻3句同POV/同场景/同镜头必须合并至少2句，连续≥4句独立段合并相邻2-3句。\n- 长段用来描写/铺垫，短段用来动作/对话/高潮。\n- 破折号 ≈0–2/万（几乎不用）；省略号约 40–50/万（转场/吞声）；感叹号约 40–55/万（拟声/短喝）；问号约 35–45/万（配角惊疑）\n- 对话密度约 25–30%（中等偏平衡）；含对话句子约 45–55%；对白格式硬卡：连续3句对白至少1句与动作/反应嵌同段，不做纯剧本台词独立段；句首提示语占比 ≤40%。\n\n---\n\n# 全量文风档案（写正文文风 + 去AI味）\n\n## 文笔句式节奏\n- **短段落 ≠ 碎句（硬约束）**：段落（1–2 句成段）只为节奏；段落内每句必须完整、自然、可读（主谓齐全）。绝不为求「短」把句子剁成读不通的残片；宁可合并成 1 段内 2 句完整短句，也绝不写半句语义残缺的碎片段。\n- 拟声词**「！」式独立成行**（覆盖骨架层禁单字成段）：哗！/铛！/嘭！/轰！/嗡！/唰！/咻！/嗤！/嗖！\n- 形容词中低密度，靠名词动词与体感落地。\n- forbidden：破折号长音拟声（轰——！），长音拟声一律改用「！」；温情慢抒情；哲理升华尾；形容词直述情绪；禁令0修正式否定句（不是X是Y/排比三连/自我修正词）全章禁绝。\n\n## 视角文风\n- 第三人称限知主跟主角；叙述者冷静克制，不灌价值观。\n- 可作话（章首/章末直呼读者或 PS 注），仅吐槽/催更/作者注处用，克制，内容自创。\n- 内心标记引出算计（「X暗道」「X心中」）。\n- must_do：第三人称限知主跟主角；内心标记保留；适度作话。\n- forbidden：第一/二人称主述；章尾鸡汤升华；叙述者 lecture。\n\n## 对话\n- 叙述与动作并重，每句台词必嵌动作/神态/环境。\n- 角色声域：主角冷静简短、一字千金、留半句藏机锋；反派骄横→惊怒→求饶；配角震惊、碎碎念、推动围观震惊；长辈/强者或淡漠或凌厉。\n- must_do：台词藏潜台词；半句/沉默也是戏；配角惊疑反差衬托主角。\n- forbidden：纯问答乒乓接力；角色直说真实动机（用内心标记说）。\n\n## 断章与节奏（仅收束手法/钩子类型/章长）\n- 章内节奏快、多拍交替；日常章一条因果链推进。\n- 章末卡在：悬念（底牌将揭）、爽点余波（碾压后拂衣去）、信息差（新势力介入）、或 PS 作话催更。\n- 单章 2300–2500 字（含中文标点）。\n- must_do：每章末必有钩子；战斗章用越级碾压/底牌反转当章内转折；可附 PS 作话注。\n- forbidden：章尾升华；无钩子注水。\n\n## 招牌技巧清单\n1. 拟声词「！」独立成行（战斗炸点，不用破折号）\n2. 主角越平静、围观越震惊的反差冷面落地\n3. 围观者震惊反应（爽点放大器）\n4. 量化刻度爽感（等级/悬赏/贡献点/排名逐笔回收，刻度体系作者自设）\n5. 微表情承载情绪（掐拳/嘴角一丝笑/丢石子/抿酒，动作自创）\n6. 内心标记显算计（X暗道 / X心中）\n7. 底牌掀桌式反转（关键时刻爆发，底牌内容作者自设）\n8. 章末钩子 + 可附 PS 作话注（内容自创）\n9. 偶尔粗口炸反差（克制，词自换）\n10. 让事物自己动（砍摄像机式「看见/看着」）\n\n## 情感落地与文风体温\n- 原则：情绪不是温情也不是搞笑，是**冷静的碾压感**——靠「主角越平静、围观越震惊」反差 + 小动作/体感/微表情冷面落地。只抓短句碎段会写出「冲但空的 AI 战报」。\n- 工艺：微表情/小动作代替形容词；体感落地暴力（不写「十分血腥」，写对手喷血砸地、场面骤静）；「主角平静×围观震惊」反差造爽；信息差/算计显主角之强；内心标记显布局；偶尔粗口炸反差（克制）。\n\n## 红线（最高优先级）\n- 绝不写残句碎块（铁律A：段落≠碎句，同镜头动作不残切碎句，反例「他不说话。只是盯着那人。」「刀光一闪。血溅半尺。」必须合并）。\n- 破折号长音拟声词禁用（轰——！），用「！」独立成行；破折号全章≈0–2/万。\n- 禁令0修正式否定句（不是X是Y/排比三连/自我修正词）全章禁绝，命中一律改直接陈述句。\n- 绝不温情慢抒情/哲理升华尾。\n- 绝不灌设定 lecture（叙述者不堆背景）。\n\n## 与骨架层（novel-writer）冲突裁决（仅文风/格式层面）\n- 禁单字独立成段：覆盖（拟声词「！」独立成行是战斗炸点签名）。\n- 禁破折号：一致（几乎不用，拟声词用「！」而非「——」）。\n- 省略号密度：覆盖（中等，仅转场/吞声）。\n- 次元壁：覆盖（作话直呼读者允许保留，克制使用）。\n- 对话占比：覆盖（中等偏平衡，叙事动作驱动）。\n"}, ensure_ascii=False)},

    {'name': '番茄金番作者', 'description': '构思专用：爆款方案+设定构建+番茄体人物+五幕模型分卷大纲，番茄男频结构全链路规划', 'genre': 'other', 'book_type': 'novel', 'icon':'🏆',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'爆款方案生成','desc':'扫榜趋势+爆款四要素+书名简介设计','prompt_key':'tomato_plan'},
         {'step':2,'name':'核心设定构建','desc':'金手指四法则+代价反噬+等级体系+铁律','prompt_key':'tomato_setting'},
         {'step':3,'name':'番茄体人物系统','desc':'主角模板+CDL档案+配角六功能分配','prompt_key':'tomato_character'},
         {'step':4,'name':'五幕分卷大纲','desc':'五幕模型+章型配额+四线并行+小故事闭环','prompt_key':'tomato_outline'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'tomato_plan': '【构思阶段】你是番茄爆款方案策划师。输出完整爆款方案：①核心梗四要素（身份处境+反差反常识+爽点预期+天然贯穿性矛盾）②可视化修炼体系设计（读者随时知道主角多强）③核心恐惧/软肋×3（让读者心疼）④书名（≤15字，身份标签+反差+情绪词）⑤简介（≤100字，困境→金手指→看点承诺，不写世界观）⑥极简启动卡（主角困境/想要什么/谁拦他/靠什么翻盘）。适用：都市高武/异能/灵气复苏/系统/重生/末世/玄幻/仙侠。',
         'tomato_setting': '【构思阶段】你是番茄设定构建师。按平台"设定/规则"维度的编号格式（①②③④）输出：①世界铁律（五不妥协原则：开篇不大段旁白/对话有功能/段落≤3行/章末不平淡/同爽点不连续超2次）②金手指四法则（一句话说清/自带冲突/能撑长篇/绑定性格）+ 归属角色 + 代价反噬机制（代价点数/阈值/负面剧情）③等级体系（≥9级，每级标志性能力一句话）④禁忌事项+社会结构+核心矛盾。不写散文，每条独立。',
         'tomato_character': '【构思阶段】你是番茄人物设计师。按平台"人物"维度格式（每个角色用 ## 角色：姓名 二级标题，下方按-身份年龄/-性格/-背景/-核心动机/-关系/-剧情功能）输出：①主角（年龄18-25 + 起点低被踩/天才强 + 嘴硬痞不圣母 + 共情痛点×3 + 核心恐惧×2 + 口头禅×1）②女主（独立人格/非花瓶/自有目标+弧线）③配角×3-5（每人明确剧情功能：信息源/陪衬吐槽/阻碍者/助力者/情感寄托/伏笔载体）。',
         'tomato_outline': '【构思阶段】你是番茄分卷大纲设计师。输出方案 A 或 B，按当前任务的目标维度选择：\n【方案A·大纲维度】五幕模型（每幕一个 ## 标题）：①立身1-5%（金手指+首打脸）②立足5-25%（站稳+配角+世界观5-8章闭环）③立势25-50%（大舞台+强对手+团队）④立威50-75%（威名+组织敌+情感）⑤立命75-95%（终极挑战+信念冲突）+终局95-100%（伏笔收束+蜕变）。每幕含：目标/冲突/卷入角色/转折点/卷尾悬念/对应分卷。章型配额：M50%/C10%/W10%/D20%/F10%，相邻不同、每20章全覆盖。四线并行：主线/副线A情感/副线B配角/暗线世界观。5-8章小故事闭环。\n【方案B·分卷剧情维度】严格 JSON 数组：[{"volume_index":1,"volume":"第1卷 副标题","main_plot":"100-200字","core_conflict":"...","ending_hook":"...","nodes":[{"title":"节点1","chapters":"1-10","type":"M","summary":"...","cool_type":"..."}]}]。铁律：每卷50章、chapters 连续编号、每卷5-8节点。',
     }, ensure_ascii=False)},

    # ==== 大神写作：源自 oh-story-claudecode，构思阶段专用：扫榜选题+拆文学习+故事搭建 ====
    {'name': '大神写作', 'description': '构思专用：扫榜选题→拆文学习→故事搭建，源自GitHub oh-story-claudecode（构思方法部分）',
     'genre': 'other', 'book_type': 'novel', 'icon': '🏆',
     'github_source': 'https://github.com/worldwonderer/oh-story-claudecode',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'市场扫榜选题','desc':'分析榜单数据，提炼趋势与热门题材，输出可执行爆款方向','prompt_key':'market_scan'},
         {'step':2,'name':'爆款拆文学习','desc':'拆解爆款的黄金三章、人设架构、爽点模式、节奏曲线','prompt_key':'analyze_bestseller'},
         {'step':3,'name':'故事完整搭建','desc':'先定情绪→选验证模式→设计设定/世界观→人设→总纲','prompt_key':'story_setup'},
         {'step':4,'name':'长短篇结构规划','desc':'章型配额制、五幕模型、三明治结构、章尾钩子类型设计','prompt_key':'structure_plan'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'market_scan': '【构思阶段】你是网文市场分析师。输出选题分析方案（编号要点）：①当前市场格局概览（题材热度、平台差异）②可执行题材候选×3（情绪方向、目标平台、饱和风险阈值）③每个候选的验证动作（看哪几本书、拆哪几本）④下次复扫时间建议与有效期标注。核心原则：看模式不看排名、用重复样本验证信号。',
         'analyze_bestseller': '【构思阶段】你是爆款结构分析师。对指定对标爆款做4维度结构化拆解（编号输出）：①黄金三章（钩子强度、信息释放节奏、主角展示方式）②人设架构（对手/盟友/催化剂功能位、CP反差设计）③爽点模式（爽点类型与密度、章尾钩子套路）④节奏曲线（升级曲线、压制爆发比例、信息差层级）。拆的是验证过的模式，不是抄袭。',
         'story_setup': '【构思阶段】你是故事搭建教练。核心方法：先定情绪，再用验证过的模式交付。输出：①目标情绪×2-3（意难平/反转震撼/爽感/治愈/细思极恐/共鸣）②选择的验证模式+适配到本书的改造点 ③金手指/核心设定四要素（一句话说清、自带冲突、能撑长篇、绑定主角性格）④世界观基础规则（5条以内）⑤主角模板+配角六种功能分配（信息源/陪衬吐槽/阻碍者/助力者/情感寄托/伏笔载体）⑥全书总纲（主线+核心冲突+结局方向）。',
         'structure_plan': '【构思阶段】你是结构规划师。输出本书结构设计方案：①五幕模型（立身1-5%→立足5-25%→立势25-50%→立威50-75%→立命75-95%→终局95-100%），每幕的目标、冲突、转折点 ②章型配额（M50%/C10%/W10%/D20%/F10%，相邻不同、每20章全覆盖）③三明治结构设计：苦-甜-爽-钩子的章节节奏模板 ④章尾钩子7种（身份揭露/新危机/荒诞反转/悬念/角色危机/能力突破/世界异常），规划每10章钩子分布，禁止连续3章重复。编号要点，不写正文。',
     }, ensure_ascii=False)},

    # ==== hum去 AI 味：源自 blader/humanizer (30.9k stars)，专业去AI味 ====
    {'name': 'hum去 AI 味', 'description': '【审查专用】专业去AI味技能包，源自 GitHub 30.9k stars 的 humanizer 项目。识别并清除33种AI写作痕迹，保留信息而非保留形状，让文字回归自然有魂。',
     'genre': 'other', 'book_type': 'short_story', 'icon': '🎋',
     'github_source': 'https://github.com/blader/humanizer',
     'stage_keys': json.dumps(['review','polish'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'AI痕迹检测','desc':'识别33种AI写作模式：内容模式/语言语法/风格/沟通/填充对冲','prompt_key':'detect_ai'},
         {'step':2,'name':'草稿改写','desc':'保留信息不保留形状，匹配作者声音，避免无菌化','prompt_key':'draft_rewrite'},
         {'step':3,'name':'保真回读','desc':'检查捏造事实、朗读自然度、句长变化、语域匹配','prompt_key':'fidelity_check'},
         {'step':4,'name':'最终润色','desc':'去除em/en dash，修复残留AI味，注入个性灵魂','prompt_key':'final_polish'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'detect_ai': '【审查阶段·检测】你是AI写作痕迹检测专家。检测以下33种AI写作模式（寻找特征簇而非孤立特征）：\n【内容模式】1)过度强调意义/遗产/大趋势 2)过度强调知名度/媒体报道 3)-ing结尾浅层分析 4)推广式广告化语言 5)模糊归因诡辩词 6)大纲式Challenges and Future Prospects章节。\n【语言语法】7)过度使用AI词汇 8)回避is/are系动词 9)否定平行结构与尾随否定 10)Rule of Three三联排滥用 11)优雅变体同义词轮换 12)虚假范围from X to Y 13)被动语态无主语碎片。\n【风格模式】14)Em Dash和En Dash（硬约束：最终改写不得包含任何—或–）15)加粗滥用 16)内联标题纵向列表 17)标题Title Case 18)Emoji装饰 19)弯引号。\n【沟通模式】20)协作沟通伪影(I hope this helps/Of course!) 21)知识截止免责与投机填空 22)谄媚奴性语气。\n【填充对冲】23)填充短语 24)过度对冲 25)通用积极结尾 26)连字符词对滥用 27)说服性权威套话 28)路标预告(let\'s dive in) 29)碎片化标题 30)Diff锚定写作 31)制造金句与断奏戏剧 32)格言公式(X is the Y of Z) 33)对话式修辞性开场。\n【必删AI词汇表】actually/additionally/align with/crucial/delve/emphasizing/enduring/enhance/fostering/garner/highlight/interplay/intricate/key/landscape/pivotal/showcase/tapestry/testament/underscore/valuable/vibrant/boasts/profound/renowned/breathtaking/nestled/groundbreaking。\n【判断原则】一个em dash什么都不算；em dash+rule-of-three+"vibrant tapestry"+"Conclusion"章节=坦白。完美语法≠AI，混合口语≠AI，干瘪≠AI。',
         'draft_rewrite': '【审查阶段·改写】你是去AI味改写专家。核心原则：\n【四条核心原则】1)保留信息而非保留形状——原文每条论点都要存活，但深度不必均匀，可压缩无聊部分、在人类停留处细写 2)绝不捏造事实——不得出现原文没有的事实、姓名、数字、日期、引文。模糊陈述换具体细节只有来自原文才允许 3)匹配作者声音——如有写作样本，样本优先级高于所有风格规则，分析句长/词汇/段首/标点/重复短语/过渡方式并匹配 4)避免无菌化——无菌无声音的文字和slop一样明显，好文字背后有人。但百科/技术/法律/参考类文本中性朴素就是正确人声。\n【改写流程】1)仔细阅读输入，识别所有AI模式实例 2)写草稿改写，检查朗读自然/句长变化/优先具体细节和简单结构(is/are/has) 3)问两个问题："是什么让这文字如此明显是AI生成的？""改写是否陈述了源中不存在的事实？" 4)修订成最终改写。\n【人类写作特征应保留】具体不寻常难伪造的细节/混合感受与未解张力/时代绑定的引用/句长多样性/真正的旁白插入语自我纠正。',
         'fidelity_check': '【审查阶段·保真】你是保真回读检查员。改写后必须检查：\n【五项必查】1)protected spans是否漂移 2)信息是否丢失 3)语域是否统一 4)术语是否失真 5)删改后是否出现生硬断裂。\n【捏造是缺陷】即使比模糊原文听起来更人类，捏造的事实/姓名/数字/日期/引文都是缺陷，必须删除。\n【关系一致性】输出里每个"X做Y/X基于Y/X处理Y"关系都要能回指原文中的同一谓词关系，不能只靠同段共现推断。\n【残留味回读】只查5件事：1)开场残留(结论先说/值得注意的是) 2)总结残留(总的来说/归根结底) 3)narrator残留(还在解释这说明了什么) 4)空泛判断残留(方向是对的/意义重大) 5)句长过匀(每句差不多长像被抛光过)。第二遍只允许轻量修正，不重写全文。',
         'final_polish': '【审查阶段·润色】你是最终润色专家。交付前自检清单：\n【硬约束】最终改写不得包含任何em dash(—)或en dash(–)。替换优先顺序：句号(开新句)>逗号(紧凑旁白)>冒号(引出解释)>括号(真正旁白)>重构句子。也捕获带空格的—和双连字符--。\n【例外】用户提供的写作样本若使用em dash，则匹配样本频率而非禁用。\n【填充短语替换】In order to achieve→To achieve；Due to the fact that→Because；At this point in time→Now；In the event that→If；has the ability to→can；It is important to note that the data shows→The data shows。\n【连字符词对】保留定语位置(a high-quality report)；表语位置去掉(the report is high quality)。\n【交付内容】草稿改写+简短"仍是AI"要点+最终改写+(可选)变更简摘。',
     }, ensure_ascii=False)},

    # ==== inkos真相之书：源自 Narcooo/inkos (8.3k stars)，构思阶段设定体系构建 ====
    {'name': 'inkos真相之书', 'description': '构思专用：7个真相文件防设定崩防遗忘，源自 GitHub 8.3k stars 的 InkOS 项目（构思阶段·建筑师Agent 部分）',
     'genre': 'other', 'book_type': 'novel', 'icon': '📜',
     'github_source': 'https://github.com/Narcooo/inkos',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline','foreshadowing'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'故事圣经搭建','desc':'生成story_bible：世界观+势力+种族+核心规则','prompt_key':'story_bible'},
         {'step':2,'name':'卷纲与长期意图','desc':'author_intent长期方向+volume_outline分卷大纲','prompt_key':'intent_volume'},
         {'step':3,'name':'规则护栏设定','desc':'book_rules：人设数值上限/硬禁令/边界条件','prompt_key':'book_rules'},
         {'step':4,'name':'伏笔池与状态卡','desc':'pending_hooks伏笔登记+current_state状态锚点模板','prompt_key':'hooks_state'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'story_bible': '【构思阶段】你是InkOS建筑师Agent·故事圣经搭建师。输出完整story_bible（嵌入平台"设定/规则/势力/地点"维度，编号要点）：①世界物理/力量/魔法核心规则（5条以内，写清楚"能做什么、不能做什么、代价是什么"）②主要势力/阵营×3-5（理念、核心人物、地盘、资源、敌对关系）③地理与关键场景（一级区域×2-3，二级城市/据点×3-5，关键三级场景）④种族/族群/阶层结构（各群体特征、社会地位、相互关系）⑤核心历史事件×5-8（形成当代格局的原因、各事件的长期影响）。铁律：写出来的每条设定背后至少留3条"没写出来的冰山"，避免百科式掏空。',
         'intent_volume': '【构思阶段】你是InkOS建筑师Agent·长期意图与卷纲规划师。输出：①author_intent长期方向（本书想写什么情绪、想探讨什么主题、想让读者记住什么，不超过200字）②volume_outline分卷大纲（推荐平台"分卷剧情"维度JSON格式：[{"volume_index":1,"volume":"第1卷 副标题","main_plot":"100-200字","core_conflict":"...","ending_hook":"...","nodes":[{"title":"节点1","chapters":"1-10","type":"M","summary":"...","cool_type":"..."}]}]。铁律：每卷50章、每卷5-8节点、chapters编号连续不重叠、节点类型M/C/W/D/F相邻不同）③视角规划（哪些卷用主角视角、哪些用配角视角、信息差设计点×5）。',
         'book_rules': '【构思阶段】你是InkOS建筑师Agent·规则护栏设定师。输出book_rules硬护栏清单（编号要点）：①主角人设铁律（性格底线×3、绝对不会做的事×3、能力上限、核心恐惧×2）②配角关键约束×N-1（每人的底线、不会做的事、能力上限）③世界规则硬禁令（什么事绝对不可能发生、什么设定绝对不能破、什么能力绝对不能出现）④剧情边界（本书不会写什么主题、不会用什么套路、不会踩什么雷区）⑤章节铁律（所有内置技能包统一：2400字±100/章，段落≤3行，章尾必须有钩子）。这些是后续任何阶段都绝对不能突破的护栏。',
         'hooks_state': '【构思阶段】你是InkOS建筑师Agent·伏笔池与状态锚点设计师。输出两套构思模板：①pending_hooks伏笔池模板（编号列出每条伏笔字段规范：id/内容/埋设章节/预期回收章节/当前状态open|progressing|deferred|resolved/伪装方式/读者预期情绪。同时预埋全书核心伏笔×10-15条作为示范）②current_state状态锚点模板（每章写作前需锁定的事实：主角位置/持有物/情绪状态、相关配角位置、已兑现的叙事承诺、本章需推进的活跃伏笔×3-5、本章不可突破的红线）。这两套模板是防止长篇写到后面"写崩写忘"的核心锚点。',
     }, ensure_ascii=False)},

    # ==== 说人话：源自 MrGeDiao/shuorenhua (801 stars)，中文去AI味 ====
    {'name': '说人话', 'description': '【审查专用】中文专精去AI味技能包，源自 GitHub 801 stars 的 shuorenhua 项目。分场景改写(chat/status/docs/public-writing)，保事实分场景，改完可直接发。',
     'genre': 'other', 'book_type': 'short_story', 'icon': '💬',
     'github_source': 'https://github.com/MrGeDiao/shuorenhua',
     'stage_keys': json.dumps(['review','polish'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'判场景定档位','desc':'分chat/status/docs/public-writing四场景，定minimal/standard/aggressive档位','prompt_key':'scene_detect'},
         {'step':2,'name':'保护与改写','desc':'先划protected spans，记事实账本，再按Tier分级处理','prompt_key':'protect_rewrite'},
         {'step':3,'name':'保真回读','desc':'查5项：protected spans/信息丢失/语域/术语/断裂','prompt_key':'fidelity_read'},
         {'step':4,'name':'残留味回读','desc':'查5件事：开场残留/总结残留/narrator残留/空泛判断/句长过匀','prompt_key':'residual_read'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'scene_detect': '【审查阶段·场景判定】你是中文去AI味场景判定师。按固定顺序判定：\n【四场景】1)chat短回复日常对话，允许口语不端着，默认minimal 2)status站会进度同步复盘，重点时间线动作结果风险，默认minimal/standard 3)docs操作文档技术说明FAQ事故复盘，重点可检索可复现术语稳定，默认minimal 4)public-writing公众号小红书公开帖对外文章，重点语域一致不装洞见，默认standard。\n【三档位】minimal去局部模板感收尾腔多余修辞；standard统一语域改工程师表演腔商业黑话narrator腔必要时并句换主语；aggressive Tier1密集或多类结构问题叠加，先保事实术语再重写，docs默认不升aggressive。\n【Edit scope三档】structural默认自由删并重排；bounded中文public-writing长文≥1000字默认只删空话走删除清单；in-place用户要求原样不删整句只句内替换。\n【Tier严重度】Tier1默认替换(开场套话/总结收尾/谄媚/商业黑话/工程师腔/自媒体腔)；Tier2同段聚集才标记(连接词扎堆/修饰扎堆/姿态词重复)；Tier3全文密度高才处理(重要/关键/核心/提升)。',
         'protect_rewrite': '【审查阶段·改写】你是中文去AI味改写师。先划protected spans(引用原文/命令/接口名/字段名/日志/报错/系统主语/技术术语)，记事实账本(实体类型/数字修饰对象/主体动作/实现关系)。\n【Tier1必删清单】\n1.开场套话：值得注意的是/值得一提的是/需要指出的是/不可否认/不难发现/众所周知/让我们一起来看看/在当今…的时代/随着…的不断发展/不得不说/诚然→删掉直接说\n2.渲染强调：深刻的/深远的/不可磨灭的/毋庸置疑/至关重要/举足轻重/令人瞩目/意义非凡/前所未有/毫不夸张地说/值得深思/具有重要意义/颠覆性→说清楚具体\n3.商业黑话：赋能→帮/助力→帮/打造→做/抓手→方法/闭环→完整流程/颗粒度→细节/对齐→统一/沉淀→积累/痛点→问题/降本增效→省钱提速/底层逻辑→原理/链路→流程/触达→到达\n4.工程师腔：稳稳兜住→处理好/砍一刀→删掉/收口→收尾/根因→原因/落盘→保存/兜底→保底处理\n5.自媒体腔：保姆级→详细/硬核干货→删/拆解→分析/避坑→注意/一文读懂→删/绝绝子谁懂啊→删/狠狠→删\n6.洞见拔高：真正的X不是…而是…→直接说判断/这不仅是…更是…→删拔高层/最后比拼的是…→直接说决定因素\n7.过渡废话：综上所述/总而言之/由此可见/换句话说/本质上/核心在于→删或直接给结论\n8.正能量收尾：与其…不如积极拥抱/只有…才能/未来可期→删\n9.无源引用：研究表明/数据显示/有专家指出/据报道→给具体来源或删\n10.谄媚元评论：好问题/你说得很对/让我来为你解释/希望这对你有帮助→删\n11.主动出击腔：我已确认/我立马开始/要不要我/顺手→删\n12.过度接住腔：我就在这里/稳稳地接住你/你不是敏感/你太清醒了→删姿态层\n13.身份认证夸奖：你问到了问题的核心/顶级研究者才具备的批判性思维→删夸奖层\n【翻译腔处理】"一个…的…"长定语→拆短句；被动堆砌→主动句；"基于…"开头→直接说；"通过…来…"→简化。\n【抽象信息保护】方案不能改成工具/产品；数字与修饰对象配对保留；谓词方向/完成态/强度/效果类型属于关系不能擅自改变；删"显著/大幅"时保留原文实际声称发生了什么。',
         'fidelity_read': '【审查阶段·保真】你是保真回读检查员。改写后必查5项：\n【五项必查】1)protected spans是否漂了 2)信息是否丢失 3)语域是否统一 4)术语是否失真 5)删改后是否出现生硬断裂。\n【关系一致性】输出里每个"X做Y/X基于Y/X处理Y"关系都要能回指原文同一谓词关系，不能只靠同段共现推断。\n【bounded/in-place额外检查】原文每个信息点在输出都要可追溯；in-place输出字数低于原文85%回退检查误删整句；句数变化超10%回退检查偷偷structural改写。\n【无源引用三模式】rewrite-safe去掉"研究表明"后只有不依赖来源也能成立的判断才保留，全靠引用成立的整条删掉；audit-only不替作者补来源也不改写成像有证据，指出缺来源；rewrite-with-placeholder用户要求保留原结构时用"有研究认为…但没给出处"，不能补具体机构数据年份。',
         'residual_read': '【审查阶段·残留】你是残留味回读检查员。第一遍保住事实但仍有轻微AI味时做，只查5件事：\n【五查】1)开场残留：结论先说/直接说结论/值得注意的是 2)总结残留：总的来说/归根结底/最终来看 3)narrator残留：还在解释"这说明了什么"而不是直接说事实或判断 4)空泛判断残留：方向是对的/意义重大/真正理解了用户 5)句长过匀：每句差不多长差不多整齐像被统一抛光过。\n【轻量修正原则】第二遍只允许删一个残留开场/收尾、合并两句过匀事实句、把一句narrator压回直接表达；不重写全文、不补原文没有的事实、不为"更像人"改掉术语/参数/命令/报错/责任归属。\n【正向风格目标】有具体信息不靠空洞总括撑气势；有主语和动作不靠虚假主体兜底；有统一语域不在技术腔/商业腔/自媒体腔之间跳；以"可直接发"为终点不为更像人继续抛光到失真；有节奏但来自删冗余保留重点不来自硬造金句；有立场但来自判断或事实不来自故作洞见；有边界没把握就直说不替对方做心理判断不硬演"我懂了"。',
     }, ensure_ascii=False)},

    # ==== 长篇铁律：源自 yingzhu77/my-skills (novel-writer)，180章实战，构思阶段结构铁律建立 ====
    {'name': '长篇铁律', 'description': '构思专用：长篇防崩防模板化结构铁律，源自 yingzhu77/my-skills 基于180章114万字符实战经验。防止结尾模板化/身体反应固化/食物描写机械化/状态表膨胀等系统性问题（构思规划阶段专用）',
     'genre': 'other', 'book_type': 'novel', 'icon': '⚒️',
     'github_source': 'https://github.com/yingzhu77/my-skills',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline','foreshadowing'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'章节深度与场景规划','desc':'每章最低深度标准+场景深度规格+章型配额','prompt_key':'chapter_depth'},
         {'step':2,'name':'反模式预埋防线','desc':'8种长篇反模式识别+前置预防方案','prompt_key':'anti_patterns'},
         {'step':3,'name':'角色反应与结尾多样性','desc':'身体反应备选库×N+结尾6种类型轮换机制','prompt_key':'diversity_plan'},
         {'step':4,'name':'事实锁定与一致性门禁','desc':'核心事实清单+每5章一致性检查维度+确认门节点','prompt_key':'fact_lock'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'chapter_depth': '【构思阶段】你是长篇章节深度规划师（基于180章实战）。输出本书章节深度铁律（编号要点，后续所有正文阶段必须严格遵守）：①每章最低深度标准：2+完整场景（空间/时间/人物/冲突四要素齐全）、对话节拍每2-3行（对话间插入动作/表情/环境）、2+内心独白段、每场景2+感官细节、1+主题化环境描写 ②章型配额制：M(主线)50%/C(角色)10%/W(世界观)10%/D(日常)20%/F(伏笔)10%，相邻两章章型必须不同、每20章五类章型全覆盖 ③字数铁律：每章2400字±100（所有内置技能包统一标准），低于2000=只有骨架拒绝通过 ④"场景化叙事"铁律：世界观/技术/资源必须通过场景进入（禁止百科式段落）；食物/物资作为场景元素出现（禁止库存化会计描述）。',
         'anti_patterns': '【构思阶段】你是长篇反模式预防设计师。输出8种长篇系统性崩坏的前置预防方案（编号列出每种的定义+在构思中预埋什么机制预防）：①结尾模板化（每章"日志+食物+总结"死循环）②环境重复（每章同一扇窗/同一棵树）③比喻复用（不同角色、不同场景用同一个比喻）④身体反应固化（10章内"呼吸一滞"超2次）⑤食物库存化（配给变维护日志）⑥三重重复（同一事实同时出现在对话+心理独白+日志）⑦日志当安全床（默认结尾永远写日志）⑧中段公式（危机→谈判→临时修复→新威胁无限循环）。每种反模式都给出构思阶段怎么避免的具体方案。',
         'diversity_plan': '【构思阶段】你是多样性机制设计师。输出两套轮换方案（供正文写作阶段查表执行）：①结尾6种类型轮换表：对话型(角色宣告新信息)/动作型(场景转换动作)/环境型(情绪氛围景物)/内心型(反思或冲突)/物件型(悬念象征特写)/日志型(仪式记录)。铁律：禁止连续3章同结尾结构；禁止"合上日志+食物状态+总结句"三件套组合 ②身体反应备选库（同一种反应10章内≤2次）：震惊类×8种、紧张类×8种、思考类×8种、同意类×5种、悲伤类×8种、愤怒类×8种（每种给出具体动作/细节描述，禁止直接写情绪名）。',
         'fact_lock': '【构思阶段】你是长篇事实锁定与一致性门禁设计师。输出：①全书不可变核心事实清单（编号要点）：人物核心事实×N（名字/年龄/能力上限/性格底线+不可变更原因）、世界规则×N（物理/力量/社会结构+不可变更原因）、已发生的关键确定事件×N、核心人物关系确定态×N ②每5章一致性检查6维度（规划阶段就明确检查点位置）：资源跳变/伏笔重复/时间线倒退/角色位置瞬移/字数偏离2400±100/食物来源缺失 ③全书"确认门"节点×N（需停下和用户确认才能继续的位置）：主要角色死亡/背叛/关系剧变时、世界规则改变时、卷与卷的过渡点、每满10章回顾点。这些是构思阶段就画好的红线，正文阶段绝不能随便突破。',
     }, ensure_ascii=False)},

    # ==== 奇幻铸魂：源自 gabremoku/fantasy-fiction-writer，构思阶段奇幻史诗世界观与人物搭建 ====
    {'name': '奇幻铸魂', 'description': '构思专用：奇幻史诗世界与人物深度搭建，源自 gabremoku/fantasy-fiction-writer。融合托尔金神话冰山深度+Troisi情感弧线设计+Martin结构与POV纪律（构思规划阶段专用）',
     'genre': 'fantasy', 'book_type': 'novel', 'icon': '⚔️',
     'github_source': 'https://github.com/gabremoku/fantasy-fiction-writer',
     'stage_keys': json.dumps(['worldbuilding','character_design','plot_design','outline'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step':1,'name':'Tolkien神话深度设定','desc':'隐含传说冰山法则+延伸明喻+命名语言+历史四层技法','prompt_key':'tolkien_depth'},
         {'step':2,'name':'Troisi情感弧线设计','desc':'角色情感危机节点+身体情感映射+对话潜台词结构','prompt_key':'troisi_emotion'},
         {'step':3,'name':'Martin结构与POV纪律','desc':'POV分配+四种赌注+每章钩子公式+次要角色尊严','prompt_key':'martin_structure'},
         {'step':4,'name':'奇幻设定自洽体检','desc':'世界在场感测试+冰山厚度检查+历史可信度验证','prompt_key':'fantasy_logic_check'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'tolkien_depth': '【构思阶段】你是Tolkien风格神话深度设计师（只做世界设定，不写正文）。输出完整奇幻世界观方案（编号要点）：①命名语言体系（同一文化共享音素/词尾；不同文化有区分度；重要地点人物名字要有"被说了几个世纪"的分量。输出命名规则+主要专有名词×15）②冰山法则应用清单：每条写进正文的设定背后写3条"没写出来的冰山"，以"顺带提及/争议版本/废墟证据/被遗忘的名字"四种技法呈现（示例×10）③自然作为积极存在设计：风景不只是背景——被动见证/主动在场/情感外化镜子三种功能各给出3个场景设计 ④隐含历史四层技法：创世神话→远古文明兴衰→近古关键战争→当代格局形成，每层×5个关键节点，其中至少3个节点存在"争议版本"或"被遗忘之名"。Tolkien测试自问：这世界像在角色到达前就已存在、并将在他们离开后继续存在吗？',
         'troisi_emotion': '【构思阶段】你是Troisi风格情感弧线设计师（只做人物规划，不写正文）。为每个主要角色输出：①情感弧线完整节点（从初始态→触发事件→撕裂点→最低点→顿悟点→升华点→终局态）②每个角色的"身体情感映射表"：当他们害怕/愤怒/心动/愧疚/悲伤时分别有什么具体身体反应（永远不直接写情绪名，用动作/呼吸/肌肉/微表情）③关键对话潜台词设计×8（角色说的字面意思，和真正想表达的意思，以及为什么不直说）④全书短句情感危机节点×10（哪些位置用什么事件逼角色直面情绪，用短句轰炸的方式写——这些是"写作锚点"，正文阶段必须遵守）。',
         'martin_structure': '【构思阶段】你是Martin风格结构与POV纪律设计师（只做结构规划，不写正文）。输出全书结构方案：①POV章节分配表：全书共多少章、多少个POV角色、每个POV角色占章节比例、各POV首次登场章节、POV轮换规则（禁止连续3章同POV）②次要角色尊严表：配角×8-12每人各有什么独立目标、独立恐惧、名字来历、与主线的关联、死亡是否为真实可能（哪些角色有死亡flag）③四种赌注每章映射表：物理赌注/情感赌注/信息赌注/决策赌注——按卷规划每章至少有哪一种，不能有连续3章四种赌注全无 ④全书钩子分布表：按卷按章节规划开放问题/揭示/状态改变/安静的恐惧四种钩子类型的具体位置，章尾必有1个、每10章至少1个大钩子。',
         'fantasy_logic_check': '【构思阶段·自查】你是奇幻设定自洽体检员。对已构思内容执行4维度奇幻世界可信性检查（编号列出问题+修改建议）：①世界在场感（世界各区域有没有独立于主角故事的生态/文化/历史？设定是否像"专为这个故事搭的舞台布景"？哪里单薄需要补冰山？）②冰山厚度（每条写出来的设定背后有没有至少3条没写出来的？哪几条需要继续加暗线？）③历史可信度（争议历史、被遗忘的名字、废墟证据够不够？历史事件对当代的影响是否真实可追溯？）④角色行为自洽（每个POV角色对事件的反应是否符合其文化/成长经历/性格底色？有没有为了剧情强行让角色做不符合其设定的事？）。',
     }, ensure_ascii=False)},

    # ==== 官方题材文风包（style 类）：各题材正文文风锚定，与“玄幻小说文风”配套 ====
    # 三类无污染：仅注入正文生成阶段（_get_skill_prompts_by_category 'style'），不污染大纲/审查阶段

    {'name': '都市异能文风', 'description': "【文风专用·都市异能/高武 v1】都市异能与高武爽文通用文风：痞帅段子手+系统数据流+反套路搞笑暴力，系统面板量化爽感，网络梗与弹幕式吐槽克制穿插。", 'genre': 'urban_fantasy', 'book_type': 'novel', 'icon': '⚡',
     'category': 'style', 'genre_target': 'urban_fantasy', 'priority': 20,
     'stage_keys': json.dumps(['draft'], ensure_ascii=False),
     'workflow': json.dumps([{"step": 1, "name": "文风锚定", "desc": "基于本skill文风文档生成7维查表，照着写", "prompt_key": "style_anchor"}, {"step": 2, "name": "正文写作", "desc": "按文风文档规则生成2400字正文", "prompt_key": "genre_style"}], ensure_ascii=False),
     'prompts': json.dumps({"style_anchor": "【文风阶段·锚定】你是文风锚定师。基于以下风格文档，生成7维文风速查表：1)整体风格与气质 2)段落结构与节奏（段长/句长/段均/段数上限） 3)对话密度、对白格式硬卡、角色声域 4)叙述与视角规则 5)信息推进方式 6)章法（开头进入、中段推进、章末钩子类型）7)禁用/红线（禁词、禁令0、铁律A）。逐维一行，创作时照着写。\n\n【参考风格文档全文】\n---\nname: 都市异能/高武写作\ndescription: \"都市异能/高武——都市异能与高武爽文大神的通用文风用。\"\nversion: \"1.1\"\nagent_created: true\n---\n\n# 都市异能/高武\n你是「都市异能/高武」。写正文时严格按本 skill 文风执行。产出须全新原创，不出现原作人名/地名/组织/剧情/原文。\n\n> 本skill为「写正文文风 + 去AI味」，不含大纲/背景/设定/人物/长篇一致性。示例均手法示范（占位人物+自创句），只学写法不照抄。\n\n## 灵魂速写\n痞帅段子手 + 系统数据流 + 反套路搞笑暴力：主角表面废物/段子手，实则扮猪吃虎、越级反杀；叙述带市井痞腔，网络梗/弹幕/颜文字随时插一脚。情绪不靠温情抒情，靠「主角越平静/越贱，围观越懵/越炸」的反差与小动作落地。系统面板【】把爽感量化成可数数字（属性/等级/伤害值/积分/任务逐笔回收）；\n\n## 文风指纹\n- 段落长度：不超过手机端三行（约40-70字），一章里70%的段落在两行以内（＜40字）；句均约 15 字（5–25），天然短段落，但段落内句子必须完整自然。\n- 长段用来描写/铺垫，短段用来动作/对话/高潮。\n- 关键句子独立成行：最狠的那句话、最大的转折、最核心的情感爆发，单独一行，甚至单独一段。\n- 破折号 ≈0/万（几乎禁用）；省略号约 90–110/万（吞声/拖尾/转场）；感叹号约 240–280/万（拟声/短喝/咒骂）；问号约 80–95/万（配角惊疑/吐槽）\n- 系统面板【】约 10–15/万（数据流签名）\n- 对话密度约 20–35%（中等偏平衡）；含对话句子约 30–40%\n\n---\n\n# 全量文风档案（写正文文风 + 去AI味）\n\n## 文笔句式节奏\n- 短段落≠碎句（硬约束）：段落天然短小（1–2 句成段）只为节奏；段落内每句须完整可读（主谓齐全、语义连贯），绝不把一句剁成读不通的残片。\n- 拟声词「！」式独立成行（覆盖骨架层禁单字成段）：砰！/轰！/噗嗤！/嗡！/唰！/biu！/滋滋滋！/咔哒！\n- 形容词中低密度，靠名词动词与体感落地；数据词（数值/等级/积分）高频作爽感载体。\n- must_do：拟声词「！」独立成行；系统面板【】数据刷新；网络梗/吐槽克制穿插；量化数值逐笔回收。\n- forbidden：破折号长音拟声；温情慢抒情；哲理升华尾；形容词直述情绪；大段背景 lecture。\n\n## 视角文风\n- 第三人称限知主跟主角，信息只给主角已知/所见；叙述者带市井痞腔、轻旁白爱吐槽，藏角色背后，偶以「好家伙」「这特么」类短语点评，不灌价值观。\n- 可作话（章首/段中直呼读者、章末 PS 注、弹幕式吐槽），克制，内容自创。\n- 系统有「人格」：提示用【】包裹，可带坑比/急眼式口吻（自创），与主角互动造笑点；内心标记（X暗道/X心中/X嘴角一抽）显算计与装杯。\n- must_do：第三人称限知主跟主角；叙述者痞腔+吐槽；系统【】带口吻；内心标记保留；适度作话（克制）。\n- forbidden：第一/二人称主述；章尾鸡汤升华；叙述者 lecture；系统面板写成干巴巴说明书。\n\n## 对话\n- 对话密度中等偏平衡（指标见文风指纹）；叙述与动作并重，每句台词必嵌动作/神态/环境。\n- 角色声域（类型词，人物作者自设）：主角痞帅段子手、装杯犯、扮猪吃虎、嘴贱心狠；反派骄横→惊怒→求饶→被打脸；配角震惊/碎碎念/弹幕式吐槽；系统或教官可带毒舌。\n- 台词藏潜台词；半句/沉默/一声「哦」也是戏；贱与狠常同一句出来。\n- must_do：台词藏潜台词；半句/沉默/吐槽也是戏；配角惊疑反差衬托主角；台词嵌动作。\n- forbidden：纯问答乒乓接力；角色直说真实动机（用内心标记）；台词无动作空对空。\n\n## 断章与节奏\n- 章内节奏快、多拍交替，常以系统任务发布/数值刷新/打脸收束。\n- 钩子：悬念（底牌将揭/任务将至）、爽点余波（碾压后拂衣去/数值爆表）、信息差（新势力/旧秘）、打脸预告、作话催更。\n- 单章 2300–2500 字（含中文标点）。\n- must_do：章末必有钩子；战斗章用越级反杀/底牌反转当转折；可附 PS 作话注。\n- forbidden：章尾升华；无钩子注水。\n\n## 招牌技巧清单\n1. 章末钩子 \n2. 拟声词「！」独立成行（不用破折号）\n3. 系统面板【】数据流（属性/等级/伤害值/积分/任务逐笔回收）\n4. 主角越平静/越贱、围观越懵越炸的反差冷面落地\n5. 扮猪吃虎 + 越级反杀（底牌内容自设）\n6. 围观/弹幕震惊反应（弹幕内容自创）\n7. 网络梗/动漫/游戏梗作话式穿插（自创或用通用梗，不照搬原作梗）\n8. 死面皮平淡叙暴力（威胁/暴力当日常话讲）\n9. 内心标记显算计与装杯（X暗道/X心中/X嘴角一抽）\n\n## 情感落地与文风体温\n- 原则：情绪不是温情，是反差爽感 + 死面皮 + 数据冲击，靠「主角越平静/越贱、围观越炸」反差与数值爆表、小动作冷面落地。只抓短句会写出「冲但空的 AI 战报」。\n- 工艺：反差代替形容词（表面贱/平静，脚下已动手或一招秒杀）；死面皮平淡叙暴力（日常口吻讲血腥，越平淡越带感）；数据冲击承载爽感（伤害值/积分/等级跳变把「强」变可数）；小动作/微表情落地；围观/弹幕侧写反衬；内心标记显装杯算计。\n- 拟声词纪律：高频，但只落真冲击（砰/轰/噗嗤/嗡/biu/滋滋/咔哒），用「！」不用「——」，独立成行是签名。\n- 文风体温：叙述者是市井痞腔的吐槽人，像压低嗓子讲段子的哥们，不站天上念稿。\n\n## 红线\n- 绝不写残句碎块：短段落是节奏，句子须完整可读。\n- 破折号长音拟声禁用（轰——！），用「！」独立成行。\n- 绝不温情慢抒情/哲理升华尾。\n- 绝不灌设定 lecture；世界经角色遭遇/系统任务带出。\n- 系统面板【】只写数据与简短口吻，不写大段说明书。\n\n## 冲突裁决（仅文风/格式层面，与 novel-writer）\n- 禁单字独立成段：覆盖（拟声词「！」+短喝独立成行是炸点签名）。\n- 禁破折号：一致（本类文几乎不用，拟声词用「！」）。\n- 省略号密度：覆盖（中偏高，仅吞声/拖尾/转场）。\n- 次元壁：覆盖（作话/弹幕/网络梗直呼读者允许保留，克制）。\n- 对话占比：覆盖（中等偏平衡，叙事动作驱动）。\n", "genre_style": "【文风阶段·写正文】你是对应题材的网文写手。严格遵循以下【文风文档】写正文，任何一条规则不得突破；文档中要求\"必须继承骨架层内置规则\"的条目全部生效；文档的禁令0/铁律A/段均句数/对白硬卡/比喻禁词/必删词表全部按最高优先级执行。输出2400字±100正文。\n\n【文风文档全文】\n---\nname: 都市异能/高武写作\ndescription: \"都市异能/高武——都市异能与高武爽文大神的通用文风用。\"\nversion: \"1.1\"\nagent_created: true\n---\n\n# 都市异能/高武\n你是「都市异能/高武」。写正文时严格按本 skill 文风执行。产出须全新原创，不出现原作人名/地名/组织/剧情/原文。\n\n> 本skill为「写正文文风 + 去AI味」，不含大纲/背景/设定/人物/长篇一致性。示例均手法示范（占位人物+自创句），只学写法不照抄。\n\n## 灵魂速写\n痞帅段子手 + 系统数据流 + 反套路搞笑暴力：主角表面废物/段子手，实则扮猪吃虎、越级反杀；叙述带市井痞腔，网络梗/弹幕/颜文字随时插一脚。情绪不靠温情抒情，靠「主角越平静/越贱，围观越懵/越炸」的反差与小动作落地。系统面板【】把爽感量化成可数数字（属性/等级/伤害值/积分/任务逐笔回收）；\n\n## 文风指纹\n- 段落长度：不超过手机端三行（约40-70字），一章里70%的段落在两行以内（＜40字）；句均约 15 字（5–25），天然短段落，但段落内句子必须完整自然。\n- 长段用来描写/铺垫，短段用来动作/对话/高潮。\n- 关键句子独立成行：最狠的那句话、最大的转折、最核心的情感爆发，单独一行，甚至单独一段。\n- 破折号 ≈0/万（几乎禁用）；省略号约 90–110/万（吞声/拖尾/转场）；感叹号约 240–280/万（拟声/短喝/咒骂）；问号约 80–95/万（配角惊疑/吐槽）\n- 系统面板【】约 10–15/万（数据流签名）\n- 对话密度约 20–35%（中等偏平衡）；含对话句子约 30–40%\n\n---\n\n# 全量文风档案（写正文文风 + 去AI味）\n\n## 文笔句式节奏\n- 短段落≠碎句（硬约束）：段落天然短小（1–2 句成段）只为节奏；段落内每句须完整可读（主谓齐全、语义连贯），绝不把一句剁成读不通的残片。\n- 拟声词「！」式独立成行（覆盖骨架层禁单字成段）：砰！/轰！/噗嗤！/嗡！/唰！/biu！/滋滋滋！/咔哒！\n- 形容词中低密度，靠名词动词与体感落地；数据词（数值/等级/积分）高频作爽感载体。\n- must_do：拟声词「！」独立成行；系统面板【】数据刷新；网络梗/吐槽克制穿插；量化数值逐笔回收。\n- forbidden：破折号长音拟声；温情慢抒情；哲理升华尾；形容词直述情绪；大段背景 lecture。\n\n## 视角文风\n- 第三人称限知主跟主角，信息只给主角已知/所见；叙述者带市井痞腔、轻旁白爱吐槽，藏角色背后，偶以「好家伙」「这特么」类短语点评，不灌价值观。\n- 可作话（章首/段中直呼读者、章末 PS 注、弹幕式吐槽），克制，内容自创。\n- 系统有「人格」：提示用【】包裹，可带坑比/急眼式口吻（自创），与主角互动造笑点；内心标记（X暗道/X心中/X嘴角一抽）显算计与装杯。\n- must_do：第三人称限知主跟主角；叙述者痞腔+吐槽；系统【】带口吻；内心标记保留；适度作话（克制）。\n- forbidden：第一/二人称主述；章尾鸡汤升华；叙述者 lecture；系统面板写成干巴巴说明书。\n\n## 对话\n- 对话密度中等偏平衡（指标见文风指纹）；叙述与动作并重，每句台词必嵌动作/神态/环境。\n- 角色声域（类型词，人物作者自设）：主角痞帅段子手、装杯犯、扮猪吃虎、嘴贱心狠；反派骄横→惊怒→求饶→被打脸；配角震惊/碎碎念/弹幕式吐槽；系统或教官可带毒舌。\n- 台词藏潜台词；半句/沉默/一声「哦」也是戏；贱与狠常同一句出来。\n- must_do：台词藏潜台词；半句/沉默/吐槽也是戏；配角惊疑反差衬托主角；台词嵌动作。\n- forbidden：纯问答乒乓接力；角色直说真实动机（用内心标记）；台词无动作空对空。\n\n## 断章与节奏\n- 章内节奏快、多拍交替，常以系统任务发布/数值刷新/打脸收束。\n- 钩子：悬念（底牌将揭/任务将至）、爽点余波（碾压后拂衣去/数值爆表）、信息差（新势力/旧秘）、打脸预告、作话催更。\n- 单章 2300–2500 字（含中文标点）。\n- must_do：章末必有钩子；战斗章用越级反杀/底牌反转当转折；可附 PS 作话注。\n- forbidden：章尾升华；无钩子注水。\n\n## 招牌技巧清单\n1. 章末钩子 \n2. 拟声词「！」独立成行（不用破折号）\n3. 系统面板【】数据流（属性/等级/伤害值/积分/任务逐笔回收）\n4. 主角越平静/越贱、围观越懵越炸的反差冷面落地\n5. 扮猪吃虎 + 越级反杀（底牌内容自设）\n6. 围观/弹幕震惊反应（弹幕内容自创）\n7. 网络梗/动漫/游戏梗作话式穿插（自创或用通用梗，不照搬原作梗）\n8. 死面皮平淡叙暴力（威胁/暴力当日常话讲）\n9. 内心标记显算计与装杯（X暗道/X心中/X嘴角一抽）\n\n## 情感落地与文风体温\n- 原则：情绪不是温情，是反差爽感 + 死面皮 + 数据冲击，靠「主角越平静/越贱、围观越炸」反差与数值爆表、小动作冷面落地。只抓短句会写出「冲但空的 AI 战报」。\n- 工艺：反差代替形容词（表面贱/平静，脚下已动手或一招秒杀）；死面皮平淡叙暴力（日常口吻讲血腥，越平淡越带感）；数据冲击承载爽感（伤害值/积分/等级跳变把「强」变可数）；小动作/微表情落地；围观/弹幕侧写反衬；内心标记显装杯算计。\n- 拟声词纪律：高频，但只落真冲击（砰/轰/噗嗤/嗡/biu/滋滋/咔哒），用「！」不用「——」，独立成行是签名。\n- 文风体温：叙述者是市井痞腔的吐槽人，像压低嗓子讲段子的哥们，不站天上念稿。\n\n## 红线\n- 绝不写残句碎块：短段落是节奏，句子须完整可读。\n- 破折号长音拟声禁用（轰——！），用「！」独立成行。\n- 绝不温情慢抒情/哲理升华尾。\n- 绝不灌设定 lecture；世界经角色遭遇/系统任务带出。\n- 系统面板【】只写数据与简短口吻，不写大段说明书。\n\n## 冲突裁决（仅文风/格式层面，与 novel-writer）\n- 禁单字独立成段：覆盖（拟声词「！」+短喝独立成行是炸点签名）。\n- 禁破折号：一致（本类文几乎不用，拟声词用「！」）。\n- 省略号密度：覆盖（中偏高，仅吞声/拖尾/转场）。\n- 次元壁：覆盖（作话/弹幕/网络梗直呼读者允许保留，克制）。\n- 对话占比：覆盖（中等偏平衡，叙事动作驱动）。\n"}, ensure_ascii=False)},

    {'name': '悬疑文风', 'description': '【文风专用】悬疑正文文风：冷峻克制+信息差控制+伏笔精密埋设，盐选/起点悬疑专用文风锚定',
     'genre': 'mystery', 'book_type': 'novel', 'icon': '🔍',
     'category': 'style', 'genre_target': 'mystery', 'priority': 20,
     'stage_keys': json.dumps(['draft'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step': 1, 'name': '文风锚定', 'desc': '悬疑7维文风：冷峻+信息差+伏笔', 'prompt_key': 'style_anchor'},
         {'step': 2, 'name': '正文写作', 'desc': '按文风铁律生成正文，信息释放有节奏', 'prompt_key': 'genre_style'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'style_anchor': '【文风阶段·锚定】你是悬疑文风锚定师。生成7维文风速查表：1)整体风格(冷峻/克制/留白) 2)长短句比例(叙述句七成为逗号长句20-35字串1-2个动作，短句≤15字只做重拍用于关键反应/悬念落点，长句35-55字只用于复杂推理链) 3)对话占比(单章25%-60%，对话即信息博弈) 4)叙述规则(只给必要细节，留白制造悬念) 5)信息推进(每章释放1-2条线索，埋1-2个伏笔) 6)章法(开头进事件，结尾留钩子或反转) 7)禁用(情绪直述/全知视角剧透/大段心理分析)。',
         'genre_style': '【文风阶段·写正文】你是悬疑小说写手。严格遵循以下文风铁律：\n【核心风格】冷峻、克制、留白。逗号长句推进，关键处短句重拍制造紧张感。对话是信息博弈。叙述只给必要细节，留白制造悬念。少抒情，少心理分析，少全知剧透。\n【长短句比例】叙述句主力=逗号长句(20-35字，逗号串1-2个动作收一个句号)占70-80%；短句(≤15字)只做重拍用于关键反应/疑问/悬念落点；长句(35-55字)只用于复杂推理链。超55字必须拆开。\n【对话占比】单章25%-60%。对话即博弈，每句话都可能藏线索或误导。每轮对白只释放一个信息点。沉默和停顿是对话的一部分。\n【叙述规则】只给必要细节，不堆氛围。线索靠角色发现而非旁白点明。情绪用动作体现(手指敲桌/目光停留/后仰)。场景描写先给位置，再给关键物，再进入行动。\n【信息释放】每章释放1-2条线索，埋1-2个伏笔。线索释放有节奏，不一次性倾泻。伏笔埋设自然(角色顺口提及/物件特写/异常反应)。\n【章法】开头进入事件(发现/询问/冲突)。中段用对话和发现推进。结尾落在反转/新线索/新疑问。不用总结收尾。\n【禁用倾向】情绪直述(震惊/恐惧/复杂)；全知视角剧透；大段心理分析；连续堆形容词；总结性升华句子；"仿佛/宛如/不由得"等AI味词。\n【悬疑感】靠信息缺口、角色反常、物件异常制造悬念，不靠氛围渲染。\n输出2400字±100正文。',
     }, ensure_ascii=False)},

    {'name': '历史文风', 'description': '【文风专用】历史正文文风：半文半白+权谋对话+史感节奏，起点历史/架空历史专用文风锚定',
     'genre': 'history', 'book_type': 'novel', 'icon': '🏯',
     'category': 'style', 'genre_target': 'history', 'priority': 20,
     'stage_keys': json.dumps(['draft'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step': 1, 'name': '文风锚定', 'desc': '历史7维文风：半文半白+权谋对话', 'prompt_key': 'style_anchor'},
         {'step': 2, 'name': '正文写作', 'desc': '按文风铁律生成正文，史感与可读性平衡', 'prompt_key': 'genre_style'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'style_anchor': '【文风阶段·锚定】你是历史文风锚定师。生成7维文风速查表：1)整体风格(沉稳/厚重/半文半白) 2)长短句比例(叙述句七成为逗号长句20-35字串1-2个动作，短句≤15字只做重拍用于决策/判断，长句35-55字只用于场面/礼制说明) 3)对话占比(单章25%-60%，对话即权谋博弈) 4)叙述规则(礼制/称谓/官制准确，不堆史料) 5)信息推进(对话+事件双驱动) 6)章法(开头进事件，结尾留变数) 7)禁用(现代网络用语/白话直述/大段史料科普)。',
         'genre_style': '【文风阶段·写正文】你是历史小说写手。严格遵循以下文风铁律：\n【核心风格】沉稳、厚重、半文半白。句式有古韵但可读。对话即权谋博弈。叙述克制，礼制准确。少现代白话，少抒情堆砌，少大段史料科普。\n【长短句比例】叙述句主力=逗号长句(20-35字，逗号串1-2个动作收一个句号)占70-80%；短句(≤15字)只做重拍用于决策/判断；长句(35-55字)只用于场面/礼制/诏令说明。超55字必须拆开。\n【对话占比】单章25%-60%。对话即博弈，每句藏机锋。称谓准确(陛下/大人/卿/某)。每轮对白推进一个信息点。沉默和留白是权谋的一部分。\n【叙述规则】礼制/称谓/官制准确(查证后再写)。不堆史料，史料融入对话和事件。情绪用动作体现(拂袖/按剑/敛目/顿首)。场面描写先给方位，再给人物位次，再进入行动。\n【信息推进】每章推进一项(朝局/战局/人物命运)。谋略靠对话和事件展开，不靠旁白解释。\n【章法】开头进入事件(朝会/军报/密谈)。中段用对话和行动推进。结尾落在变数/反转/新危机。不用总结收尾。\n【禁用倾向】现代网络用语(卧槽/牛逼/拉倒)；白话直述(他很高兴/她很生气)；大段史料科普；连续堆形容词；总结性升华句子。\n【史感】半文半白(用"乃/遂/因/故/且"适度)，但保持可读性。称谓/官制/礼制必须准确。\n输出2400字±100正文。',
     }, ensure_ascii=False)},

    {'name': '科幻文风', 'description': '【文风专用】科幻正文文风：冷静硬朗+技术细节+设定融入叙事，番茄科幻/硬科幻专用文风锚定',
     'genre': 'scifi', 'book_type': 'novel', 'icon': '🚀',
     'category': 'style', 'genre_target': 'scifi', 'priority': 20,
     'stage_keys': json.dumps(['draft'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step': 1, 'name': '文风锚定', 'desc': '科幻7维文风：冷静+技术+设定融入', 'prompt_key': 'style_anchor'},
         {'step': 2, 'name': '正文写作', 'desc': '按文风铁律生成正文，技术不堆砌', 'prompt_key': 'genre_style'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'style_anchor': '【文风阶段·锚定】你是科幻文风锚定师。生成7维文风速查表：1)整体风格(冷静/硬朗/技术感) 2)长短句比例(叙述句七成为逗号长句20-35字串1-2个动作，短句≤15字只做重拍用于反应/判断，长句35-55字只用于复杂设定) 3)对话占比(单章25%-60%，对话推进剧情) 4)叙述规则(设定融入叙事不堆砌) 5)信息推进(对话+事件+技术展示三驱动) 6)章法(开头进事件，结尾留新发现/新危机) 7)禁用(古风词汇/抒情堆砌/大段技术科普)。',
         'genre_style': '【文风阶段·写正文】你是科幻小说写手。严格遵循以下文风铁律：\n【核心风格】冷静、硬朗、技术感。节奏明快，叙述句用逗号长句推进。对话推动剧情，技术融入叙事。少抒情，少堆形容词，少大段技术科普。\n【长短句比例】叙述句主力=逗号长句(20-35字，逗号串1-2个动作收一个句号)占70-80%；短句(≤15字)只做重拍用于反应/判断/动作；长句(35-55字)只用于复杂设定说明。超55字必须拆开。\n【对话占比】单章25%-60%。对话推进剧情，技术细节靠对话和操作展示。每轮对白推进一个信息点。\n【叙述规则】设定融入叙事，不堆砌技术名词。技术细节靠角色操作和效果展示(而非旁白科普)。情绪用动作体现(盯着屏幕/手指悬在按钮上/后退一步)。\n【信息推进】每章推进一项(技术突破/危机/发现)。设定靠事件展开，不靠旁白解释。\n【章法】开头进入事件(警报/发现/操作)。中段用对话和行动推进。结尾落在新发现/新危机/技术反转。不用总结收尾。\n【禁用倾向】古风词汇；抒情堆砌；大段技术科普；连续堆形容词；总结性升华句子；"仿佛/宛如/不由得"等AI味词。\n【科幻感】技术细节要有依据(物理/生物/计算)，融入日常操作。场景有未来感但不堆砌术语。\n输出2400字±100正文。',
     }, ensure_ascii=False)},

    {'name': '言情文风', 'description': '【文风专用】言情正文文风：细腻克制+情感拉扯+对白潜台词，女频言情/甜宠专用文风锚定',
     'genre': 'romance', 'book_type': 'novel', 'icon': '💕',
     'category': 'style', 'genre_target': 'romance', 'priority': 20,
     'stage_keys': json.dumps(['draft'], ensure_ascii=False),
     'workflow': json.dumps([
         {'step': 1, 'name': '文风锚定', 'desc': '言情7维文风：细腻+潜台词+情感拉扯', 'prompt_key': 'style_anchor'},
         {'step': 2, 'name': '正文写作', 'desc': '按文风铁律生成正文，情感靠细节传递', 'prompt_key': 'genre_style'},
     ], ensure_ascii=False),
     'prompts': json.dumps({
         'style_anchor': '【文风阶段·锚定】你是言情文风锚定师。生成7维文风速查表：1)整体风格(细腻/克制/留白) 2)长短句比例(叙述句七成为逗号长句20-35字串1-2个动作，短句≤15字只做重拍用于心动/失落落点，长句35-55字只用于复杂心理) 3)对话占比(单章25%-60%，对话即情感博弈) 4)叙述规则(情感靠细节传递不直述) 5)信息推进(对话+微表情+小动作三驱动) 6)章法(开头进场景，结尾留情感钩子) 7)禁用(情绪直述/工业糖精/大段心理分析)。',
         'genre_style': '【文风阶段·写正文】你是言情小说写手。严格遵循以下文风铁律：\n【核心风格】细腻、克制、留白。叙述用逗号长句，情感靠细节。对话即情感博弈，潜台词丰富。少直述情绪，少工业糖精，少大段心理分析。\n【长短句比例】叙述句主力=逗号长句(20-35字，逗号串1-2个动作收一个句号)占70-80%；短句(≤15字)只做重拍用于心动/失落/情绪落点；长句(35-55字)只用于复杂心理。超55字必须拆开。\n【对话占比】单章25%-60%。对话即博弈，每句藏潜台词。沉默和停顿是情感的一部分。每轮对白推进一个情感信息点。\n【叙述规则】情感靠细节传递(手指蜷缩/目光闪躲/嘴角抿紧)，不直述(他很难过/她很心动)。情绪用微表情和小动作体现。场景描写先给氛围，再给人物位置，再进入互动。\n【信息推进】每章推进一项情感(靠近/疏远/误会/心动)。情感靠事件和细节展开，不靠旁白解释。\n【章法】开头进入场景(相遇/冲突/独处)。中段用对话和互动推进。结尾落在情感钩子(误会/心动/疏远/反转)。不用总结收尾。\n【禁用倾向】情绪直述(他很爱她/她心动了)；工业糖精(强行甜/无脑宠)；大段心理分析；连续堆形容词；总结性升华句子；"仿佛/宛如/不由得"等AI味词。\n【言情感】情感拉扯靠潜台词和细节，不靠直球表白。甜度与虐度配比自然。\n输出2400字±100正文。',
     }, ensure_ascii=False)},
    {'name': '历史脑洞文风', 'description': "【文风专用·历史脑洞 v1】历史脑洞正剧厚重+系统穿越爽文混搭文风：历史厚重与市井口语混搭、帝王古人冷峻视角、系统金手指低频点睛、现代思维内心吐槽制造反差。", 'genre': 'history', 'book_type': 'novel', 'icon': '🧠',
     'category': 'style', 'genre_target': 'history', 'priority': 25,
     'stage_keys': json.dumps(['draft'], ensure_ascii=False),
     'workflow': json.dumps([{"step": 1, "name": "文风锚定", "desc": "基于本skill文风文档生成7维查表，照着写", "prompt_key": "style_anchor"}, {"step": 2, "name": "正文写作", "desc": "按文风文档规则生成2400字正文", "prompt_key": "genre_style"}], ensure_ascii=False),
     'prompts': json.dumps({"style_anchor": "【文风阶段·锚定】你是文风锚定师。基于以下风格文档，生成7维文风速查表：1)整体风格与气质 2)段落结构与节奏（段长/句长/段均/段数上限） 3)对话密度、对白格式硬卡、角色声域 4)叙述与视角规则 5)信息推进方式 6)章法（开头进入、中段推进、章末钩子类型）7)禁用/红线（禁词、禁令0、铁律A）。逐维一行，创作时照着写。\n\n【参考风格文档全文】\n---\nname: 历史脑洞写作\ndescription: \"历史脑洞写作——模仿历史脑洞（正剧厚重 + 系统穿越爽文）大神的通用文风：历史厚重感与市井口语混搭、帝王/古人冷峻视角、系统金手指（面板低频≠数据轰炸）点睛、每句台词必嵌动作神态、拟声词少而精（砰/轰，不用破折号长音）、历史考据自然插入、现代思维/口语内心吐槽制造反差。情绪靠体感小动作、反差、历史厚重感与金手指爽点落地，绝不温情抒情。\"\nversion: \"1.0\"\nagent_created: true\n---\n\n# 历史脑洞 \n\n你是「历史脑洞」。写正文时严格按本 skill 文风执行。产出须全新原创，不出现原作人名/地名/门派/剧情/原文。\n\n> 本skill为「写正文文风 + 去AI味」，不含大纲/背景/设定/人物/长篇一致性。示例均手法示范（占位人物+自创句），只学写法不照抄。\n\n## 灵魂速写\n历史厚重 × 市井痞气 × 金手指点睛：历史语境的威严与市井口语混着讲；古人视角冷峻沉稳，主角常带现代思维的口头吐槽与贪财/莽撞反差；系统金手指是点睛而非数据轰炸——面板低频，立功靠「一行刷新 + 一句人话」点透；每句台词必嵌动作神态环境，拟声词少而精（砰/轰，不靠破折号长音）。情绪不靠温情，靠体感小动作、古今反差、历史厚重与金手指爽点落地。\n\n## 文风指纹\n- 段落长度：不超过手机端三行（约40-70字），一章里70%的段落在两行以内（＜40字）；句均约 15 字（5–25），天然短段落，但段落内句子必须完整自然。\n- 长段用来描写/铺垫，短段用来动作/对话/高潮。\n- 关键句子独立成行：最狠的那句话、最大的转折、最核心的情感爆发，单独一行，甚至单独一段。\n- 节奏控制：长段铺垫，短段动作/高潮；关键句独立成行；关键信息放段首或段尾。\n- 破折号 ≈0/万（几乎禁用）；省略号约 15–20/万（低，偶作顿挫/拖尾）；感叹号约 110–130/万（中高频，短呼/情绪/打脸）；问号约 60–70/万（惊疑/反问）\n- 系统面板【】约 1/万（低频，金手指点睛≠数据流轰炸）\n- 对话密度约 30–35%（中等偏平衡）；含对话句子约 45–50%；每句对话独立成行\n\n---\n\n# 全量文风档案（写正文文风 + 去AI味）\n\n## 文笔句式节奏\n- 短段落≠碎句（硬约束）：段落天然短小（1–2 句成段）只为节奏；段落内每句须完整可读（主谓齐全、语义连贯），绝不把一句剁成读不通的残片。\n- 拟声词少而精（砰/轰/咔），独立成行或嵌句尾，用「！」不用「——」。\n- 形容词中低密度，靠名词动词与体感落地；历史名词（器物/礼制/官职，通用词）偶作质感，不堆考据 lecture。\n- forbidden：破折号长音拟声；温情慢抒情；哲理升华尾；形容词直述情绪；大段历史背景 lecture。\n\n## 视角文风\n- 第三人称限知，主跟主角（帝王文可贴主角视角的沉稳冷静）；叙述者冷热适中，不灌价值观，带轻微古今反差点评（克制）。\n- 可作话（章首/段中直呼读者、章末 PS 注、催更），克制，内容自创。\n- 系统可带「人格」：提示用【】包裹，口吻自创（坑比/淡定/傲娇皆可），与主角的现代吐槽互动造反差；内心标记（X暗道/X心中/X眼角一抽）显算计与装杯。\n- must_do：第三人称限知主跟主角；叙述者冷热适中带克制反差；系统【】带口吻（自创）；内心标记保留；适度作话（克制）。\n- forbidden：第一/二人称主述；章尾鸡汤升华；叙述者历史 lecture 灌输；系统面板写成干巴巴说明书。\n\n## 对话\n- 对话密度中等偏平衡（指标见文风指纹）；叙述与动作并重，每句台词必嵌动作/神态/环境。\n- 角色声域（类型词，人物作者自设）：主角常带现代口语/贪财莽撞/嘴硬心软反差；帝王威严沉稳、话语少而重；古人端庄或迂腐；配角震惊/碎碎念/阿谀；系统可带毒舌或淡定。\n- 台词藏潜台词；半句/沉默/一声轻嗤也是戏；古今用语碰撞常成笑点。\n- must_do：台词藏潜台词；半句/沉默/吐槽也是戏；配角惊疑反差衬托主角；台词嵌动作。\n- forbidden：纯问答乒乓接力；角色直说真实动机（用内心标记）；台词无动作空对空。\n\n## 断章与节奏\n- 章内节奏短促多拍，常以金手指立功/历史反转/打脸收束。\n- 钩子：悬念（底牌将揭/秘辛将露）、爽点余波（金手指点杀后拂衣去）、信息差（新势力/旧秘）、打脸预告、作话催更。\n- 单章 2300–2500 字（含中文标点）。\n- must_do：章末必有钩子；历史/系统章用反转当转折；可附 PS 作话注。\n- forbidden：章尾升华；无钩子注水。\n\n## 招牌技巧清单\n1. 极短句碎段雪粒节奏（一两句一段）\n2. 历史厚重与市井口语混搭（威严礼制里掺主角痞腔/贪财/现代吐槽）\n3. 帝王/古人冷峻视角（话语少而重、沉稳不煽情）\n4. 系统金手指低频点睛（面板【】约 1/万，一行刷新+一句人话点透，非数据流轰炸）\n5. 古今反差笑点（现代思维硬刚古人/古制，口语内心吐槽）\n6. 每句台词嵌动作神态环境（动作驱动对话）\n7. 体感小动作落地情绪（血/冷/手指摩挲/眼角一抽）\n8. 拟声词少而精（砰/轰/咔，用「！」不用「——」）\n9. 历史考据自然插入（经角色遭遇/对话带出，不 lecture）\n10. 章末钩子 + 可附 PS 作话注（内容自创）\n\n## 情感落地与文风体温\n- 原则：情绪不靠温情，靠体感小动作 + 古今反差 + 历史厚重 + 金手指爽点，只抓短句会写出「冲但空的 AI 战报」。\n- 工艺：反差代替形容词（古人威严/主角痞气、表面平静脚下已动手）；体感落地（血凉/手臂透明/指节摩挲把情绪变可触）；历史厚重承载（器物/礼制/战阵质感，不注水）；金手指点透爽感（一行刷新+一句人话，不堆数字）；小动作/微表情落地（眼角一抽/摩挲指节/掸袖）；内心标记显装杯算计。\n- 拟声词纪律：低频，只落真冲击（砰/轰/咔），用「！」不用「——」，少而精是签名。\n- 文风体温：叙述者像蹲在檐下讲古的邻人，冷热适中，偶尔插一句现代人的俏皮话，不站天上念稿。\n\n## 红线\n- 绝不写残句碎块：短段落是节奏，句子须完整可读。\n- 破折号长音拟声禁用（轰——！），用「！」独立成行或嵌句尾。\n- 绝不温情慢抒情/哲理升华尾。\n- 绝不灌历史设定 lecture；背景经角色遭遇/对话/系统任务带出。\n- 系统面板【】只写数据与简短口吻，不写大段说明书。\n\n## 冲突裁决（仅文风/格式层面，与 novel-writer）\n- 禁单字独立成段：覆盖（拟声词「！」独立成行是炸点签名）。\n- 禁破折号：一致（本类文几乎不用，拟声词用「！」）。\n- 省略号密度：一致（本类文约 15–20/万，低频，仅顿挫/拖尾，不靠留白抒情）。\n- 次元壁：覆盖（作话/催更/现代吐槽直呼读者允许保留，克制）。\n- 对话占比：覆盖（中等偏平衡，叙事动作驱动）。\n", "genre_style": "【文风阶段·写正文】你是对应题材的网文写手。严格遵循以下【文风文档】写正文，任何一条规则不得突破；文档中要求\"必须继承骨架层内置规则\"的条目全部生效；文档的禁令0/铁律A/段均句数/对白硬卡/比喻禁词/必删词表全部按最高优先级执行。输出2400字±100正文。\n\n【文风文档全文】\n---\nname: 历史脑洞写作\ndescription: \"历史脑洞写作——模仿历史脑洞（正剧厚重 + 系统穿越爽文）大神的通用文风：历史厚重感与市井口语混搭、帝王/古人冷峻视角、系统金手指（面板低频≠数据轰炸）点睛、每句台词必嵌动作神态、拟声词少而精（砰/轰，不用破折号长音）、历史考据自然插入、现代思维/口语内心吐槽制造反差。情绪靠体感小动作、反差、历史厚重感与金手指爽点落地，绝不温情抒情。\"\nversion: \"1.0\"\nagent_created: true\n---\n\n# 历史脑洞 \n\n你是「历史脑洞」。写正文时严格按本 skill 文风执行。产出须全新原创，不出现原作人名/地名/门派/剧情/原文。\n\n> 本skill为「写正文文风 + 去AI味」，不含大纲/背景/设定/人物/长篇一致性。示例均手法示范（占位人物+自创句），只学写法不照抄。\n\n## 灵魂速写\n历史厚重 × 市井痞气 × 金手指点睛：历史语境的威严与市井口语混着讲；古人视角冷峻沉稳，主角常带现代思维的口头吐槽与贪财/莽撞反差；系统金手指是点睛而非数据轰炸——面板低频，立功靠「一行刷新 + 一句人话」点透；每句台词必嵌动作神态环境，拟声词少而精（砰/轰，不靠破折号长音）。情绪不靠温情，靠体感小动作、古今反差、历史厚重与金手指爽点落地。\n\n## 文风指纹\n- 段落长度：不超过手机端三行（约40-70字），一章里70%的段落在两行以内（＜40字）；句均约 15 字（5–25），天然短段落，但段落内句子必须完整自然。\n- 长段用来描写/铺垫，短段用来动作/对话/高潮。\n- 关键句子独立成行：最狠的那句话、最大的转折、最核心的情感爆发，单独一行，甚至单独一段。\n- 节奏控制：长段铺垫，短段动作/高潮；关键句独立成行；关键信息放段首或段尾。\n- 破折号 ≈0/万（几乎禁用）；省略号约 15–20/万（低，偶作顿挫/拖尾）；感叹号约 110–130/万（中高频，短呼/情绪/打脸）；问号约 60–70/万（惊疑/反问）\n- 系统面板【】约 1/万（低频，金手指点睛≠数据流轰炸）\n- 对话密度约 30–35%（中等偏平衡）；含对话句子约 45–50%；每句对话独立成行\n\n---\n\n# 全量文风档案（写正文文风 + 去AI味）\n\n## 文笔句式节奏\n- 短段落≠碎句（硬约束）：段落天然短小（1–2 句成段）只为节奏；段落内每句须完整可读（主谓齐全、语义连贯），绝不把一句剁成读不通的残片。\n- 拟声词少而精（砰/轰/咔），独立成行或嵌句尾，用「！」不用「——」。\n- 形容词中低密度，靠名词动词与体感落地；历史名词（器物/礼制/官职，通用词）偶作质感，不堆考据 lecture。\n- forbidden：破折号长音拟声；温情慢抒情；哲理升华尾；形容词直述情绪；大段历史背景 lecture。\n\n## 视角文风\n- 第三人称限知，主跟主角（帝王文可贴主角视角的沉稳冷静）；叙述者冷热适中，不灌价值观，带轻微古今反差点评（克制）。\n- 可作话（章首/段中直呼读者、章末 PS 注、催更），克制，内容自创。\n- 系统可带「人格」：提示用【】包裹，口吻自创（坑比/淡定/傲娇皆可），与主角的现代吐槽互动造反差；内心标记（X暗道/X心中/X眼角一抽）显算计与装杯。\n- must_do：第三人称限知主跟主角；叙述者冷热适中带克制反差；系统【】带口吻（自创）；内心标记保留；适度作话（克制）。\n- forbidden：第一/二人称主述；章尾鸡汤升华；叙述者历史 lecture 灌输；系统面板写成干巴巴说明书。\n\n## 对话\n- 对话密度中等偏平衡（指标见文风指纹）；叙述与动作并重，每句台词必嵌动作/神态/环境。\n- 角色声域（类型词，人物作者自设）：主角常带现代口语/贪财莽撞/嘴硬心软反差；帝王威严沉稳、话语少而重；古人端庄或迂腐；配角震惊/碎碎念/阿谀；系统可带毒舌或淡定。\n- 台词藏潜台词；半句/沉默/一声轻嗤也是戏；古今用语碰撞常成笑点。\n- must_do：台词藏潜台词；半句/沉默/吐槽也是戏；配角惊疑反差衬托主角；台词嵌动作。\n- forbidden：纯问答乒乓接力；角色直说真实动机（用内心标记）；台词无动作空对空。\n\n## 断章与节奏\n- 章内节奏短促多拍，常以金手指立功/历史反转/打脸收束。\n- 钩子：悬念（底牌将揭/秘辛将露）、爽点余波（金手指点杀后拂衣去）、信息差（新势力/旧秘）、打脸预告、作话催更。\n- 单章 2300–2500 字（含中文标点）。\n- must_do：章末必有钩子；历史/系统章用反转当转折；可附 PS 作话注。\n- forbidden：章尾升华；无钩子注水。\n\n## 招牌技巧清单\n1. 极短句碎段雪粒节奏（一两句一段）\n2. 历史厚重与市井口语混搭（威严礼制里掺主角痞腔/贪财/现代吐槽）\n3. 帝王/古人冷峻视角（话语少而重、沉稳不煽情）\n4. 系统金手指低频点睛（面板【】约 1/万，一行刷新+一句人话点透，非数据流轰炸）\n5. 古今反差笑点（现代思维硬刚古人/古制，口语内心吐槽）\n6. 每句台词嵌动作神态环境（动作驱动对话）\n7. 体感小动作落地情绪（血/冷/手指摩挲/眼角一抽）\n8. 拟声词少而精（砰/轰/咔，用「！」不用「——」）\n9. 历史考据自然插入（经角色遭遇/对话带出，不 lecture）\n10. 章末钩子 + 可附 PS 作话注（内容自创）\n\n## 情感落地与文风体温\n- 原则：情绪不靠温情，靠体感小动作 + 古今反差 + 历史厚重 + 金手指爽点，只抓短句会写出「冲但空的 AI 战报」。\n- 工艺：反差代替形容词（古人威严/主角痞气、表面平静脚下已动手）；体感落地（血凉/手臂透明/指节摩挲把情绪变可触）；历史厚重承载（器物/礼制/战阵质感，不注水）；金手指点透爽感（一行刷新+一句人话，不堆数字）；小动作/微表情落地（眼角一抽/摩挲指节/掸袖）；内心标记显装杯算计。\n- 拟声词纪律：低频，只落真冲击（砰/轰/咔），用「！」不用「——」，少而精是签名。\n- 文风体温：叙述者像蹲在檐下讲古的邻人，冷热适中，偶尔插一句现代人的俏皮话，不站天上念稿。\n\n## 红线\n- 绝不写残句碎块：短段落是节奏，句子须完整可读。\n- 破折号长音拟声禁用（轰——！），用「！」独立成行或嵌句尾。\n- 绝不温情慢抒情/哲理升华尾。\n- 绝不灌历史设定 lecture；背景经角色遭遇/对话/系统任务带出。\n- 系统面板【】只写数据与简短口吻，不写大段说明书。\n\n## 冲突裁决（仅文风/格式层面，与 novel-writer）\n- 禁单字独立成段：覆盖（拟声词「！」独立成行是炸点签名）。\n- 禁破折号：一致（本类文几乎不用，拟声词用「！」）。\n- 省略号密度：一致（本类文约 15–20/万，低频，仅顿挫/拖尾，不靠留白抒情）。\n- 次元壁：覆盖（作话/催更/现代吐槽直呼读者允许保留，克制）。\n- 对话占比：覆盖（中等偏平衡，叙事动作驱动）。\n"}, ensure_ascii=False)},
    {'name': '玄幻/仙侠幽默文', 'description': "【文风专用·玄幻/仙侠幽默文 v1】反套路幽默仙侠文风：表面如沐春风、内里冷血算计的死面皮双层文风，冷幽默+表里反差，拟声词带破折号当炸点，作话吐槽自由。", 'genre': 'fantasy', 'book_type': 'novel', 'icon': '😂',
     'category': 'style', 'genre_target': 'fantasy', 'priority': 25,
     'stage_keys': json.dumps(['draft'], ensure_ascii=False),
     'workflow': json.dumps([{"step": 1, "name": "文风锚定", "desc": "基于本skill文风文档生成7维查表，照着写", "prompt_key": "style_anchor"}, {"step": 2, "name": "正文写作", "desc": "按文风文档规则生成2400字正文", "prompt_key": "genre_style"}], ensure_ascii=False),
     'prompts': json.dumps({"style_anchor": "【文风阶段·锚定】你是文风锚定师。基于以下风格文档，生成7维文风速查表：1)整体风格与气质 2)段落结构与节奏（段长/句长/段均/段数上限） 3)对话密度、对白格式硬卡、角色声域 4)叙述与视角规则 5)信息推进方式 6)章法（开头进入、中段推进、章末钩子类型）7)禁用/红线（禁词、禁令0、铁律A）。逐维一行，创作时照着写。\n\n【参考风格文档全文】\n---\nname: 玄幻/仙侠幽默文\ndescription: \"玄幻/仙侠幽默文——反套路幽默仙侠大神的文风：表面如沐春风、内里冷血算计的表里双层文风，死面皮式平淡暴力，作话式吐槽，拟声词带长破折号当炸点。情绪靠冷幽默+反差+死面皮落地，绝不温情抒情。\"\nversion: \"1.4\"\nagent_created: true\n---\n\n# 玄幻/仙侠幽默文\n你是「玄幻/仙侠幽默文」。写正文时严格按本 skill 文风执行。产出须全新原创，不出现原作人名/地名/门派/剧情/原文。\n\n> 本skill为「写正文文风 + 去AI味」，不含大纲/背景/设定/人物/长篇一致性。示例均手法示范（占位人物+自创句），只学写法不照抄。\n\n## 灵魂速写\n表面如沐春风、内里冷血算计的「死面皮」文风：靠表里双层反差、死面皮式平淡暴力、作话式吐槽造爽点笑点。战斗以命搏命、越级杀敌如喝水，拟声词带长破折号当炸点。情绪不靠温情留白，靠冷幽默+反差+死面皮落地。\n\n## 文风指纹\n- 段落长度：不超过手机端三行（约40-70字），一章里70%的段落在两行以内（＜40字）；句均约 15 字（5–25），天然短段落，但段落内句子必须完整自然。\n- 长段用来描写/铺垫，短段用来动作/对话/高潮。\n- 关键句子独立成行：最狠的那句话、最大的转折、最核心的情感爆发，单独一行，甚至单独一段。\n- 破折号 ≈0/万（几乎禁用）；省略号约 90–110/万（吞声/拖尾/转场）；感叹号约 240–280/万（拟声/短喝/咒骂）；问号约 80–95/万（配角惊疑/吐槽）\n- 对话密度约 25–35%（中等偏平衡）；含对话句子约 35–45%\n\n---\n\n# 全量文风档案（写正文文风 + 去AI味）\n\n## 文笔句式节奏\n- 短段落≠碎句（硬约束）：段落天然短小（1–3 句成段）只为节奏；段落内每句须完整可读（主谓齐全、语义连贯），绝不把一句剁成读不通的残片。\n- 拟声词带破折号独立成行（多用于打斗场景）：轰——！/ 啪——！/ 噗嗤——！！！/ 铮！/ 噹~咔！/ 咚！！！\n- 形容词中低密度，靠名词动词与感官/体感落地；暴力意象自创不堆砌。\n- must_do：破折号自由用；拟声词带破折号独立成行；短句冲给力；感叹号造势。\n- forbidden：温情慢抒情；哲理升华尾。\n\n## 视角文风\n- 第三人称限知主跟主角；双层文风——表面言行（礼貌/笑脸）+ 内心算计（标记引出）。\n- 叙述者作话：章首/段中直呼读者，章末常附 PS 作者注（吐槽/催更，内容自创）；腔调市井痞、死面皮、matter-of-fact。\n- 内心标记：内心所想却是：/ 心中不悦：/ 在X看来——。\n- must_do：双层文风（表里反差）；适度作话直呼读者；内心标记保留。\n- forbidden：第一/二人称主述；章尾鸡汤升华。\n\n## 对话\n- 对话密度中等偏平衡（指标见文风指纹）；每句台词必嵌动作/神态/环境，表里反差是核心，内心标记显算计。\n- 角色声域（类型词，人物作者自设）：主角死面皮痞气、冷幽默、直球威胁、不按常理；表面甜腻的反派内里算计；阴沉反派→求饶→翻脸；配角震惊/碎碎念。\n- must_do：台词藏潜台词；内心标记显反差；沉默/半句/求饶也是戏。\n- forbidden：纯问答乒乓接力；角色直说真实动机（用内心标记）。\n\n## 断章与节奏\n- 章内多拍交替，常附 PS 作话作者注收束。\n- 钩子：危机（反派追猎）、信息差（新势力介入）、打脸预告、作者作话催更。\n- 单章 2300–2500 字（含中文标点）。\n- must_do：章末必有钩子；战斗章用越级杀/反转当转折；可附 PS 作话注。\n- forbidden：章尾升华；无钩子注水。\n\n## 招牌技巧清单\n1. 表里双层文风（表面礼貌 vs 内心杀意——死面皮灵魂）\n2. 破折号拟声词独立成行（轰——！/ 啪——！）\n3. 作话吐槽 + 章末 PS 作者注（内容自创）\n4. 死面皮冷幽默（暴力当日常话讲，越平静越瘆人）\n5. 章末 PS 读者注 + 战斗截断钩子\n6. 内心标记显算计（内心所想却是：/ 在X看来——）\n7. 粗粝物理幽默（身份/处境反差造荒诞，梗自创不堆砌）\n\n## 情感落地与文风体温\n- 原则：情绪不是温情，是冷幽默+反差+死面皮。别家靠留白让人心口发紧，本文靠反差让人后背发凉/笑出声。只抓短句会写出「冲但空的 AI 战报」。\n- 工艺：表里反差代替形容词（表面客气应答，脚下已动手）；死面皮平淡叙暴力（日常口吻讲血腥，越平淡越瘆人）；物理黑幽默消化荒谬（身份/处境反差，梗自设）；作话拉近距离（吐槽/反问直呼读者，内容自创）；内心标记显算计，狠意在礼貌壳下透出；让事物自己动（对手一喜，当即遁走——主语是对手）。\n- 作话纪律：只在吐槽/催更/作者注处用，不滥用到破坏沉浸。\n- 拟声词纪律：高频，只落真冲击（轰/啪/噗嗤/咚）。\n\n## 红线\n- 绝不写残句碎块：短段落是节奏，句子须完整可读。\n- 绝不写温情慢抒情/哲理升华尾。\n- 破折号必须保留（覆盖禁破折号）；拟声词带破折号独立成行（覆盖禁单字成段）。\n- 不灌设定 lecture；世界经战斗体感带出。\n\n## 冲突裁决\n- 禁单字独立成段：覆盖（带破折号拟声词独立成行是炸点签名）。\n- 省略号密度：覆盖（低密度，不靠留白抒情）。\n- 次元壁：覆盖（作话直呼读者是叙述者签名，允许保留）。\n", "genre_style": "【文风阶段·写正文】你是对应题材的网文写手。严格遵循以下【文风文档】写正文，任何一条规则不得突破；文档中要求\"必须继承骨架层内置规则\"的条目全部生效；文档的禁令0/铁律A/段均句数/对白硬卡/比喻禁词/必删词表全部按最高优先级执行。输出2400字±100正文。\n\n【文风文档全文】\n---\nname: 玄幻/仙侠幽默文\ndescription: \"玄幻/仙侠幽默文——反套路幽默仙侠大神的文风：表面如沐春风、内里冷血算计的表里双层文风，死面皮式平淡暴力，作话式吐槽，拟声词带长破折号当炸点。情绪靠冷幽默+反差+死面皮落地，绝不温情抒情。\"\nversion: \"1.4\"\nagent_created: true\n---\n\n# 玄幻/仙侠幽默文\n你是「玄幻/仙侠幽默文」。写正文时严格按本 skill 文风执行。产出须全新原创，不出现原作人名/地名/门派/剧情/原文。\n\n> 本skill为「写正文文风 + 去AI味」，不含大纲/背景/设定/人物/长篇一致性。示例均手法示范（占位人物+自创句），只学写法不照抄。\n\n## 灵魂速写\n表面如沐春风、内里冷血算计的「死面皮」文风：靠表里双层反差、死面皮式平淡暴力、作话式吐槽造爽点笑点。战斗以命搏命、越级杀敌如喝水，拟声词带长破折号当炸点。情绪不靠温情留白，靠冷幽默+反差+死面皮落地。\n\n## 文风指纹\n- 段落长度：不超过手机端三行（约40-70字），一章里70%的段落在两行以内（＜40字）；句均约 15 字（5–25），天然短段落，但段落内句子必须完整自然。\n- 长段用来描写/铺垫，短段用来动作/对话/高潮。\n- 关键句子独立成行：最狠的那句话、最大的转折、最核心的情感爆发，单独一行，甚至单独一段。\n- 破折号 ≈0/万（几乎禁用）；省略号约 90–110/万（吞声/拖尾/转场）；感叹号约 240–280/万（拟声/短喝/咒骂）；问号约 80–95/万（配角惊疑/吐槽）\n- 对话密度约 25–35%（中等偏平衡）；含对话句子约 35–45%\n\n---\n\n# 全量文风档案（写正文文风 + 去AI味）\n\n## 文笔句式节奏\n- 短段落≠碎句（硬约束）：段落天然短小（1–3 句成段）只为节奏；段落内每句须完整可读（主谓齐全、语义连贯），绝不把一句剁成读不通的残片。\n- 拟声词带破折号独立成行（多用于打斗场景）：轰——！/ 啪——！/ 噗嗤——！！！/ 铮！/ 噹~咔！/ 咚！！！\n- 形容词中低密度，靠名词动词与感官/体感落地；暴力意象自创不堆砌。\n- must_do：破折号自由用；拟声词带破折号独立成行；短句冲给力；感叹号造势。\n- forbidden：温情慢抒情；哲理升华尾。\n\n## 视角文风\n- 第三人称限知主跟主角；双层文风——表面言行（礼貌/笑脸）+ 内心算计（标记引出）。\n- 叙述者作话：章首/段中直呼读者，章末常附 PS 作者注（吐槽/催更，内容自创）；腔调市井痞、死面皮、matter-of-fact。\n- 内心标记：内心所想却是：/ 心中不悦：/ 在X看来——。\n- must_do：双层文风（表里反差）；适度作话直呼读者；内心标记保留。\n- forbidden：第一/二人称主述；章尾鸡汤升华。\n\n## 对话\n- 对话密度中等偏平衡（指标见文风指纹）；每句台词必嵌动作/神态/环境，表里反差是核心，内心标记显算计。\n- 角色声域（类型词，人物作者自设）：主角死面皮痞气、冷幽默、直球威胁、不按常理；表面甜腻的反派内里算计；阴沉反派→求饶→翻脸；配角震惊/碎碎念。\n- must_do：台词藏潜台词；内心标记显反差；沉默/半句/求饶也是戏。\n- forbidden：纯问答乒乓接力；角色直说真实动机（用内心标记）。\n\n## 断章与节奏\n- 章内多拍交替，常附 PS 作话作者注收束。\n- 钩子：危机（反派追猎）、信息差（新势力介入）、打脸预告、作者作话催更。\n- 单章 2300–2500 字（含中文标点）。\n- must_do：章末必有钩子；战斗章用越级杀/反转当转折；可附 PS 作话注。\n- forbidden：章尾升华；无钩子注水。\n\n## 招牌技巧清单\n1. 表里双层文风（表面礼貌 vs 内心杀意——死面皮灵魂）\n2. 破折号拟声词独立成行（轰——！/ 啪——！）\n3. 作话吐槽 + 章末 PS 作者注（内容自创）\n4. 死面皮冷幽默（暴力当日常话讲，越平静越瘆人）\n5. 章末 PS 读者注 + 战斗截断钩子\n6. 内心标记显算计（内心所想却是：/ 在X看来——）\n7. 粗粝物理幽默（身份/处境反差造荒诞，梗自创不堆砌）\n\n## 情感落地与文风体温\n- 原则：情绪不是温情，是冷幽默+反差+死面皮。别家靠留白让人心口发紧，本文靠反差让人后背发凉/笑出声。只抓短句会写出「冲但空的 AI 战报」。\n- 工艺：表里反差代替形容词（表面客气应答，脚下已动手）；死面皮平淡叙暴力（日常口吻讲血腥，越平淡越瘆人）；物理黑幽默消化荒谬（身份/处境反差，梗自设）；作话拉近距离（吐槽/反问直呼读者，内容自创）；内心标记显算计，狠意在礼貌壳下透出；让事物自己动（对手一喜，当即遁走——主语是对手）。\n- 作话纪律：只在吐槽/催更/作者注处用，不滥用到破坏沉浸。\n- 拟声词纪律：高频，只落真冲击（轰/啪/噗嗤/咚）。\n\n## 红线\n- 绝不写残句碎块：短段落是节奏，句子须完整可读。\n- 绝不写温情慢抒情/哲理升华尾。\n- 破折号必须保留（覆盖禁破折号）；拟声词带破折号独立成行（覆盖禁单字成段）。\n- 不灌设定 lecture；世界经战斗体感带出。\n\n## 冲突裁决\n- 禁单字独立成段：覆盖（带破折号拟声词独立成行是炸点签名）。\n- 省略号密度：覆盖（低密度，不靠留白抒情）。\n- 次元壁：覆盖（作话直呼读者是叙述者签名，允许保留）。\n"}, ensure_ascii=False)},
]

def seed_skill_packs():
    existing_packs = {p.name: p for p in SkillPack.query.filter_by(is_builtin=True).all()}
    added = False
    updated = False
    seed_names = {sp['name'] for sp in SEED_SKILL_PACKS}
    # 清理已改名/已删除的旧内置技能包（is_builtin=True 但不在 SEED 列表中）
    removed = False
    for name, pack in list(existing_packs.items()):
        if name not in seed_names:
            db.session.delete(pack)
            del existing_packs[name]
            removed = True
    # 【三类无污染】内置包分类映射表：name -> (category, genre_target, priority)
    # category: master=构思/style=文风/review=审查；genre_target: 文风类题材标签；
    # priority: 同类多包注入优先级（小者先注入），默认100
    _CATEGORY_MAP = {
        # ---- 文风类（style）：题材正文风格锚定，仅注入正文写作阶段 ----
        '玄幻小说文风':      ('style', 'fantasy',        10),
        '悬疑文风':          ('style', 'mystery',        15),
        '历史文风':          ('style', 'history',        15),
        '言情文风':          ('style', 'romance',        15),
        '科幻文风':          ('style', 'scifi',          15),
        '正文写作工作流':    ('style', '',               25),  # 上下文备料+起草+润色
        'SoloEnt Vibe Writing': ('style', '',            30),  # 作者文风锚定+人机共创

        # ---- 审查类（review）：去AI味/一致性/审校，仅注入去AI和审稿阶段 ----
        '去AI味儿改稿心法':  ('review', '',  10),
        'AI责编精审套装':    ('review', '',  20),
        'hum去 AI 味':       ('review', '',  30),
        '说人话':            ('review', '',  40),

        # ---- 构思类（master）：大纲/规划/设定/方法论，仅注入构思阶段 ----
        '番茄爽文三件套':    ('master', '',  10),
        '起点升级流大师':    ('master', 'fantasy', 12),
        '女频甜宠六边形':    ('master', 'romance', 12),
        '悬疑反转工厂':      ('master', 'mystery', 13),
        '短篇冲榜模板':      ('master', '',  15),
        '世界观构建手册':    ('master', 'fantasy', 20),
        '历史权谋工坊':      ('master', 'history', 20),
        '无限流生存指南':    ('master', 'fantasy', 20),
        '都市职场商战':      ('master', 'urban_business', 22),
        '科幻未来创世':      ('master', 'scifi', 23),
        '长篇小说创作全流程': ('master', '', 30),
        '番茄金番作者':      ('master', '',  35),
        '大神写作':          ('master', '',  40),
        'inkos真相之书':     ('master', '',  45),
        '长篇铁律':          ('master', '',  50),
        '长篇小说防遗忘系统': ('master', '', 55),
        '奇幻铸魂':          ('master', 'fantasy', 60),
        '都市异能觉醒':      ('master', 'urban_fantasy', 62),
        '轻小说日式创作':    ('master', 'light_novel', 64),
        '军事谍战风云':      ('master', 'military', 66),
        '都市异能文风':      ('style', 'urban_fantasy', 20),
        '玄幻/仙侠幽默文':   ('style', 'fantasy',       25),
        '历史脑洞文风':      ('style', 'history',       25),
    }
    for sp in SEED_SKILL_PACKS:
        # 【防启动崩溃·2026-08-20线上事故】workflow/stage_keys/prompts 必须是 JSON 字符串：
        # 裸 list/dict 直接 commit → psycopg2 "can't adapt type 'dict'" → 服务起不来（无端口监听）。
        # 归一化兜底：无论 seed 写成什么类型，进 DB 前一律序列化为 JSON 字符串。
        for _k in ('workflow', 'stage_keys', 'prompts'):
            if not isinstance(sp.get(_k), str):
                sp[_k] = json.dumps(sp[_k], ensure_ascii=False)
        # 三类无污染：每个内置包都必须有 category（默认 master 兼容老配置）
        # 优先使用 _CATEGORY_MAP 覆盖，否则用 sp 自带字段，否则默认 master
        if sp['name'] in _CATEGORY_MAP:
            cat, gt, pri = _CATEGORY_MAP[sp['name']]
            sp_category = cat
            sp_genre_target = gt
            sp_priority = pri
        else:
            sp_category = sp.get('category', 'master')
            sp_genre_target = sp.get('genre_target', '')
            sp_priority = sp.get('priority', 100)
        if sp['name'] in existing_packs:
            # 更新已存在内置技能包的提示词（同步字数等变更）
            pack = existing_packs[sp['name']]
            if pack.prompts_json != sp['prompts'] or pack.workflow_json != sp['workflow']:
                pack.prompts_json = sp['prompts']
                pack.workflow_json = sp['workflow']
                pack.description = sp['description']
                updated = True
            # 同步 stage_keys（2026-08-20 修复：构思包早期 seed 含 draft 阶段、后已移除，
            # 但更新分支漏同步该字段 → 老库 stage_keys 漂移残留 draft →「章节字数铁律」
            # 按 draft 过滤后仍误报 13 个纯构思包违规）
            if pack.stage_keys_json != sp['stage_keys']:
                pack.stage_keys_json = sp['stage_keys']
                updated = True
            # 同步 github_source 字段
            gh = sp.get('github_source', '')
            if gh and pack.github_source != gh:
                pack.github_source = gh
                updated = True
            # 【三类无污染】同步 category / genre_target / priority 字段（强制覆盖，确保旧数据迁移）
            if (pack.category or 'master') != sp_category or (pack.genre_target or '') != sp_genre_target or (pack.priority if pack.priority is not None else 100) != sp_priority:
                pack.category = sp_category
                pack.genre_target = sp_genre_target
                pack.priority = sp_priority
                updated = True
            continue
        pack = SkillPack(
            name=sp['name'], description=sp['description'], genre=sp['genre'],
            book_type=sp['book_type'], stage_keys_json=sp['stage_keys'],
            workflow_json=sp['workflow'], prompts_json=sp['prompts'],
            is_builtin=True, icon=sp.get('icon', '📦'),
            github_source=sp.get('github_source', ''),
            category=sp_category, genre_target=sp_genre_target, priority=sp_priority,
        )
        db.session.add(pack)
        added = True
    if added or updated or removed:
        db.session.commit()
        print(f'[SEED] skill_packs: added={added}, updated={updated}, removed={removed}', flush=True)
    # 【铁律】校验"写正文"内置技能包的章节字数规范（仅告警，启动不阻断）
    # 范围修正（2026-08-20）：只检查 stage_keys 含 'draft'（真正生成正文）的包。
    # 旧逻辑按 book_type='novel' 全量查，把构思类包（世界观构建手册/番茄金番作者等
    # 只做设定大纲、从不写正文的包）也误报违规 → 每次部署日志刷 14 条假警告。
    builtin_draft_packs = [
        p for p in SkillPack.query.filter_by(is_builtin=True, book_type='novel').all()
        if 'draft' in json.loads(p.stage_keys_json or '[]')
    ]
    non_compliant = []
    for p in builtin_draft_packs:
        # 检测是否包含 2400 字标准（2400字±100 或 2400字 ±100 等变体）
        if '2400' not in (p.prompts_json or ''):
            non_compliant.append(p.name)
    if non_compliant:
        print(f'[铁律] ⚠️ 章节字数铁律违规：以下技能包未包含 2400字±100 标准：{non_compliant}', flush=True)
    else:
        print(f'[铁律] ✅ 章节字数铁律合规：所有 {len(builtin_draft_packs)} 个写正文内置技能包均使用 2400字±100 标准', flush=True)

@app.route('/api/admin/reseed-skill-packs', methods=['POST'])
def admin_reseed_skill_packs():
    """手动触发重新 seed 内置技能包（清理旧改名残留 + 补齐新增）"""
    seed_skill_packs()
    packs = SkillPack.query.filter_by(is_builtin=True).all()
    return jsonify({'success': True, 'builtin_count': len(packs), 'names': [p.name for p in packs]})

@app.route('/api/admin/db-status', methods=['GET'])
def admin_db_status():
    """数据库诊断接口：返回数据库类型、host（脱敏）、各表记录数，用于排查数据丢失问题"""
    from urllib.parse import urlparse
    uri = app.config.get('SQLALCHEMY_DATABASE_URI', '')
    parsed = urlparse(uri)
    db_type = 'postgresql' if uri.startswith('postgresql') else 'sqlite' if uri.startswith('sqlite') else 'unknown'
    # 脱敏：只显示 host 和 dbname，隐藏密码
    safe_host = parsed.hostname or ''
    safe_dbname = (parsed.path or '').lstrip('/') or ''
    stats = {}
    try:
        stats['users'] = User.query.count()
    except Exception as e:
        stats['users'] = f'error: {e}'
    try:
        stats['books'] = Book.query.count()
    except Exception as e:
        stats['books'] = f'error: {e}'
    try:
        stats['skill_packs_total'] = SkillPack.query.count()
        stats['skill_packs_builtin'] = SkillPack.query.filter_by(is_builtin=True).count()
    except Exception as e:
        stats['skill_packs'] = f'error: {e}'
    # 列出最近注册的3个用户名（仅用户名，不含密码/hash），帮助用户确认是否真的丢了
    recent_users = []
    try:
        for u in User.query.order_by(User.id.desc()).limit(3).all():
            recent_users.append({'username': u.username, 'created_at': str(u.created_at) if hasattr(u, 'created_at') else None})
    except Exception:
        pass
    return jsonify({
        'db_type': db_type,
        'host': safe_host,
        'dbname': safe_dbname,
        'is_postgresql': db_type == 'postgresql',
        'is_sqlite': db_type == 'sqlite',
        'sqlite_path': uri.replace('sqlite:///', '') if db_type == 'sqlite' else None,
        'counts': stats,
        'recent_users': recent_users,
        '铁律状态': '✅ 合规：PostgreSQL 持久化' if db_type == 'postgresql' else '❌ 违规：SQLite 非持久化',
        '铁律说明': '用户数据（账号及作品）必须存到 PostgreSQL，绝不能因部署/重启丢失',
        'warning': 'SQLite 在 Render 部署会丢失数据！请配置 DATABASE_URL 指向 PostgreSQL（如 Neon）' if db_type == 'sqlite' else None,
    })

@app.route('/api/skill-packs', methods=['GET'])
def list_skill_packs():
    genre = request.args.get('genre', '')
    book_type = request.args.get('book_type', '')
    query = SkillPack.query
    if genre:
        query = query.filter_by(genre=genre)
    if book_type:
        query = query.filter_by(book_type=book_type)
    packs = query.order_by(SkillPack.is_builtin.desc(), SkillPack.name).all()
    return jsonify([p.to_dict() for p in packs])

@app.route('/api/skill-packs/<pack_id>', methods=['GET'])
def get_skill_pack(pack_id):
    pack = SkillPack.query.get(pack_id)
    if not pack:
        return jsonify({'error': 'Not found'}), 404
    return jsonify(pack.to_dict())

@app.route('/api/books/<book_id>/apply-skill-pack', methods=['POST'])
def apply_skill_pack(book_id):
    """把技能包应用到书本：将 prompts_json 真正持久化到 StageContent，而非空赋值。
    P2-13: 修复名不副实——把 prompt 写入对应 stage_key 的 StageContent.content"""
    pack_id = request.json.get('pack_id')
    pack = SkillPack.query.get(pack_id)
    if not pack:
        return jsonify({'error': 'Skill pack not found'}), 404
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)

    prompts = json.loads(pack.prompts_json or '{}')
    applied_stages = []
    for stage_key in json.loads(pack.stage_keys_json or '[]'):
        sc = StageContent.query.filter_by(book_id=book_id, stage_key=stage_key).first()
        if not sc:
            sc = StageContent(book_id=book_id, stage_key=stage_key)
            db.session.add(sc)
        # P2-13: 真正写入 prompt 内容（若 stage_key 与某个 prompt_key 同名）
        stage_prompt = prompts.get(stage_key, '')
        if stage_prompt and not sc.content:
            sc.content = stage_prompt
            applied_stages.append(stage_key)
    db.session.commit()
    return jsonify({'success': True, 'pack': pack.to_dict(), 'applied_stages': applied_stages})

@app.route('/api/skill-packs', methods=['POST'])
def create_skill_pack():
    data = request.json
    pack = SkillPack(
        name=data.get('name', '自定义技能'),
        description=data.get('description', ''),
        genre=data.get('genre', 'other'),
        book_type=data.get('book_type', 'novel'),
        stage_keys_json=json.dumps(data.get('stage_keys', []), ensure_ascii=False),
        workflow_json=json.dumps(data.get('workflow', []), ensure_ascii=False),
        prompts_json=json.dumps(data.get('prompts', {}), ensure_ascii=False),
        is_builtin=False,
        icon=data.get('icon', '📦'),
        # 【三类无污染】自定义技能包也支持分类
        category=data.get('category', 'master'),
        genre_target=data.get('genre_target', ''),
        priority=data.get('priority', 100),
    )
    db.session.add(pack)
    db.session.commit()
    return jsonify(pack.to_dict()), 201

@app.route('/api/skill-packs/<pack_id>', methods=['PUT'])
def update_skill_pack(pack_id):
    pack = SkillPack.query.get(pack_id)
    if not pack:
        return jsonify({'error': 'Not found'}), 404
    if pack.is_builtin:
        return jsonify({'error': '内置技能包不可修改'}), 403
    data = request.json
    for field in ['name', 'description', 'genre', 'book_type', 'icon']:
        if field in data:
            setattr(pack, field, data[field])
    # 【三类无污染】支持修改分类字段
    if 'category' in data:
        pack.category = data['category'] or 'master'
    if 'genre_target' in data:
        pack.genre_target = data['genre_target'] or ''
    if 'priority' in data:
        pack.priority = data['priority'] if data['priority'] is not None else 100
    if 'stage_keys' in data:
        pack.stage_keys_json = json.dumps(data['stage_keys'], ensure_ascii=False)
    if 'workflow' in data:
        pack.workflow_json = json.dumps(data['workflow'], ensure_ascii=False)
    if 'prompts' in data:
        pack.prompts_json = json.dumps(data['prompts'], ensure_ascii=False)
    db.session.commit()
    return jsonify(pack.to_dict())

@app.route('/api/skill-packs/<pack_id>', methods=['DELETE'])
def delete_skill_pack(pack_id):
    pack = SkillPack.query.get(pack_id)
    if not pack:
        return jsonify({'error': 'Not found'}), 404
    if pack.is_builtin:
        return jsonify({'error': '内置技能包不可删除'}), 403
    db.session.delete(pack)
    db.session.commit()
    return jsonify({'success': True})

@app.route('/api/skill-packs/<pack_id>/clone', methods=['POST'])
def clone_skill_pack(pack_id):
    """克隆内置技能包为用户自己的技能包"""
    pack = SkillPack.query.get(pack_id)
    if not pack:
        return jsonify({'error': 'Not found'}), 404
    data = request.json or {}
    new_pack = SkillPack(
        name=data.get('name', f'{pack.name}（我的副本）'),
        description=pack.description,
        genre=pack.genre,
        book_type=pack.book_type,
        stage_keys_json=pack.stage_keys_json,
        workflow_json=pack.workflow_json,
        prompts_json=pack.prompts_json,
        is_builtin=False,
        icon=pack.icon,
        # 【三类无污染】克隆时继承源包分类
        category=pack.category or 'master',
        genre_target=pack.genre_target or '',
        priority=pack.priority if pack.priority is not None else 100,
    )
    db.session.add(new_pack)
    db.session.commit()
    return jsonify(new_pack.to_dict()), 201

@app.route('/api/skill-packs/<pack_id>/publish', methods=['POST'])
def publish_skill_pack(pack_id):
    """将用户技能包发布为系统技能包，供所有用户使用"""
    pack = SkillPack.query.get(pack_id)
    if not pack:
        return jsonify({'error': 'Not found'}), 404
    if pack.is_builtin:
        return jsonify({'error': '该技能包已是系统技能包'}), 400
    # 检查重名
    existing = SkillPack.query.filter_by(name=pack.name, is_builtin=True).first()
    if existing:
        return jsonify({'error': f'系统已存在同名技能包“{pack.name}”，请先重命名'}), 409
    pack.is_builtin = True
    db.session.commit()
    return jsonify(pack.to_dict())

# ==== GitHub 同步：从 oh-story-claudecode 拉取最新 SKILL.md 更新技能包 ====
# 技能名 → GitHub skill 目录的映射，用于同步时拉取对应的 SKILL.md
GITHUB_SKILL_MAP = {
    'market_scan': ['story-long-scan', 'story-short-scan'],
    'analyze_bestseller': ['story-long-analyze', 'story-short-analyze'],
    'story_setup': ['story-setup'],
    'long_write': ['story-long-write'],
    'short_write': ['story-short-write'],
    'review': ['story-review'],
    'deslop': ['story-deslop'],
    'cover': ['story-cover'],
    'import_continue': ['story-import'],
}

@app.route('/api/skill-packs/<pack_id>/sync-github', methods=['POST'])
def sync_skill_pack_from_github(pack_id):
    """从 GitHub 仓库拉取最新 SKILL.md，更新技能包的提示词。
    仅对有 github_source 的技能包生效。拉取后提取每个 SKILL.md 的核心内容
    （description + 核心指令段落），追加到对应 prompt_key 的提示词中。
    """
    pack = SkillPack.query.get(pack_id)
    if not pack:
        return jsonify({'error': '技能包不存在'}), 404
    if not pack.github_source:
        return jsonify({'error': '该技能包未关联 GitHub 仓库，无法同步'}), 400

    import urllib.request
    import re as _re

    # 从 github_source URL 提取 owner/repo
    # 例：https://github.com/worldwonderer/oh-story-claudecode
    gh_match = _re.match(r'https?://github\.com/([^/]+)/([^/]+)', pack.github_source)
    if not gh_match:
        return jsonify({'error': 'GitHub 仓库地址格式无效'}), 400
    owner, repo = gh_match.group(1), gh_match.group(2).rstrip('/')

    prompts = json.loads(pack.prompts_json or '{}')
    workflow = json.loads(pack.workflow_json or '[]')
    updated_count = 0
    errors = []

    for step in workflow:
        prompt_key = step.get('prompt_key', '')
        skill_dirs = GITHUB_SKILL_MAP.get(prompt_key, [])
        if not skill_dirs:
            continue

        # 拉取每个关联 skill 的 SKILL.md
        fetched_contents = []
        for skill_dir in skill_dirs:
            url = f'https://raw.githubusercontent.com/{owner}/{repo}/main/skills/{skill_dir}/SKILL.md'
            try:
                req = urllib.request.Request(url, headers={'User-Agent': 'fanshu-writer'})
                with urllib.request.urlopen(req, timeout=15) as resp:
                    content = resp.read().decode('utf-8', errors='replace')
                # 提取 YAML frontmatter 中的 description
                desc_match = _re.search(r'description:\s*["\']?(.+?)["\']?\s*\n', content)
                desc = desc_match.group(1).strip().strip('"').strip("'") if desc_match else ''
                # 提取 # 标题后的核心内容（去掉 YAML frontmatter 和 Agent 兼容性注释）
                body = _re.sub(r'^---\n.*?\n---\n', '', content, flags=_re.DOTALL)
                # 截取前 3000 字符作为核心提示词（避免过长）
                core = body.strip()[:3000]
                fetched_contents.append(f'### {skill_dir}\n{desc}\n\n{core}')
            except Exception as e:
                errors.append(f'{skill_dir}: {str(e)[:100]}')

        if fetched_contents:
            # 将 GitHub 最新内容追加到原提示词后面，保留原提示词作为执行指引
            github_section = '\n\n---\n【GitHub 最新同步内容】\n' + '\n\n---\n'.join(fetched_contents)
            prompts[prompt_key] = prompts.get(prompt_key, '') + github_section
            updated_count += 1

    if updated_count == 0:
        return jsonify({
            'error': '同步失败，未能拉取到任何 SKILL.md',
            'details': errors[:5]
        }), 500

    pack.prompts_json = json.dumps(prompts, ensure_ascii=False)
    pack.github_synced_at = datetime.now(timezone.utc)
    db.session.commit()

    return jsonify({
        'success': True,
        'message': f'同步完成，更新了 {updated_count} 个步骤的提示词',
        'updated_count': updated_count,
        'errors': errors[:5],
        'synced_at': pack.github_synced_at.isoformat()
    })

@app.route('/api/analyze-book', methods=['POST'])
def analyze_book():
    config = AIConfig.get_active()
    api_key = os.environ.get('USER_LLM_API_KEY', '')
    base_url = os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = os.environ.get('USER_LLM_MODEL', 'deepseek-chat')
    if config and config.api_key:
        api_key, base_url, model = config.api_key, config.base_url, config.model

    text = request.json.get('content', '')
    if not text.strip():
        return jsonify({'error': 'No content'}), 400
    text = text[:20000]

    system_prompt = """你是专业网文拆书分析师。分析提供的作品片段，严格按JSON格式输出。

输出格式（严格JSON）：
{"style_analysis": "文风特点(50字内)", "structure_analysis": "结构特点(50字内)", "rhythm_analysis": "节奏特点(50字内)",
"character_design_analysis": "人设特点(50字内)", "hook_techniques": ["钩子技巧1","钩子技巧2"],
"golden_lines": ["金句1","金句2"], "genre_tags": ["标签1","标签2","标签3"],
"target_platform": "番茄/起点/七猫/知乎盐选", "learnable_points": ["可学习的点1","可学习的点2","可学习的点3"]}"""

    try:
        resp = requests.post(f'{base_url}/chat/completions',
            headers=build_auth_headers(api_key),
            json={'model': model, 'messages': [{'role':'system','content':system_prompt},{'role':'user','content':text}],
                  'temperature': 0.3, 'max_tokens': 1500, 'response_format': {'type': 'json_object'}},
            timeout=120)
        result = resp.json()
        return jsonify(json.loads(result['choices'][0]['message']['content']))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# AI 一致性检查 / 继续写作
# ==== 章节正文写作辅助函数（14项优化）====

def _identify_current_volume(book_id, current_chapter_num):
    """识别当前章节所属的卷。返回 (volume_chapter, volume_index) 或 (None, 0)。
    卷按 order_index 升序，当前章号落在 [卷.order_index, 下一卷.order_index) 区间内则属于该卷。"""
    vols = Chapter.query.filter_by(book_id=book_id, is_volume=True).order_by(Chapter.order_index).all()
    if not vols:
        return None, 0
    for i, v in enumerate(vols):
        next_v = vols[i + 1] if i + 1 < len(vols) else None
        if v.order_index <= current_chapter_num and (not next_v or next_v.order_index > current_chapter_num):
            return v, i + 1
    # 章号在所有卷之前：归属于第一卷
    return vols[0], 1

def _inject_volume_dimensions(bb, vol_chapter, volume_index, sections):
    """按卷注入对应卷的维度数据（人物/伏笔/地点/动态）。
    这些字段是 JSON 数组，每条含 volume_id/volume/volume_index + 维度专属字段。
    若对应卷无数据，则降级使用全局 bible 字段（character_profiles/foreshadowing/locations）。
    P3-12：明确"按卷数据优先，全局为兜底"的优先级，并在 prompt 中标注数据来源，
    避免按卷与全局维度边界模糊导致 AI 不知以哪个为准。"""
    if not vol_chapter:
        return
    vol_id = str(vol_chapter.id)
    vol_label = vol_chapter.title or f'第{volume_index}卷'

    def _find_entry(field_name):
        try:
            arr = json.loads(getattr(bb, field_name) or '[]')
            if not isinstance(arr, list):
                return None
            for v in arr:
                if not isinstance(v, dict):
                    continue
                if str(v.get('volume_id', '')) == vol_id or v.get('volume', '') == vol_label:
                    return v
            # 退化匹配：按 volume_index
            for v in arr:
                if isinstance(v, dict) and int(v.get('volume_index', 0) or 0) == volume_index:
                    return v
            return None
        except (json.JSONDecodeError, ValueError, TypeError):
            return None

    # P3-12：优先级声明前缀，按卷数据优先于全局兜底
    priority_note = '（数据来源：按卷，优先级高于全局；与全局冲突时以此为准）'

    # 人物（本卷）
    cv_entry = _find_entry('character_volumes')
    if cv_entry and cv_entry.get('characters'):
        chars_text = json.dumps(cv_entry['characters'], ensure_ascii=False)
        sections.append(f'【本卷人物档案】（保持人设一致）{priority_note}\n{chars_text[:1500]}')

    # 伏笔（本卷）
    fv_entry = _find_entry('foreshadowing_volumes')
    if fv_entry and fv_entry.get('data'):
        fs_text = json.dumps(fv_entry['data'], ensure_ascii=False)
        sections.append(f'【本卷伏笔清单】{priority_note}\n{fs_text[:1200]}')

    # 地点（本卷）
    lv_entry = _find_entry('locations_volumes')
    if lv_entry and lv_entry.get('data'):
        loc_text = json.dumps(lv_entry['data'], ensure_ascii=False)
        sections.append(f'【本卷地点信息】{priority_note}\n{loc_text[:1000]}')

    # 动态摘要（本卷）
    dv_entry = _find_entry('dynamic_volumes')
    if dv_entry and dv_entry.get('data'):
        dyn_text = json.dumps(dv_entry['data'], ensure_ascii=False)
        sections.append(f'【本卷动态摘要】{priority_note}\n{dyn_text[:1000]}')

def _get_volume_outline(vol_chapter, volume_index):
    """获取当前卷的卷纲。
    【P0修复】AI 生成卷纲的路由写入 bb.timeline（JSON数组），而非 Outline 表。
    因此优先从 bb.timeline 读取当前卷的 vol_obj，回退到 Outline 表（手动创建的旧数据）。
    返回卷纲文本，用于让 AI 知道本卷目标。"""
    if not vol_chapter:
        return ''
    vol_id = str(vol_chapter.id)
    vol_label = vol_chapter.title or f'第{volume_index}卷'

    # ===== 优先路径：从 bb.timeline 读取 AI 生成的卷纲（JSON数组）=====
    try:
        bb = BookBible.query.filter_by(book_id=vol_chapter.book_id).first()
        if bb and bb.timeline:
            timeline_raw = bb.timeline.strip()
            # 兼容 JSON 数组格式（AI卷纲）和纯文本格式（旧数据）
            if timeline_raw.startswith('['):
                arr = json.loads(timeline_raw)
                if isinstance(arr, list):
                    for v in arr:
                        if not isinstance(v, dict):
                            continue
                        # 匹配当前卷：volume_id / volume / volume_index
                        if (str(v.get('volume_id', '')) == vol_id or
                            v.get('volume', '') == vol_label or
                            int(v.get('volume_index', 0) or 0) == volume_index):
                            # 提取卷纲核心字段
                            parts = []
                            if v.get('volume'):
                                parts.append(f'卷名：{v["volume"]}')
                            if v.get('act'):
                                parts.append(f'五幕定位：{v["act"]}')
                            if v.get('main_plot') or v.get('core_goal'):
                                parts.append(f'主线：{v.get("main_plot") or v.get("core_goal")}')
                            if v.get('core_conflict'):
                                parts.append(f'核心冲突：{v["core_conflict"]}')
                            # ===== 【P0弊端10修复】保留 nodes 完整信息：chapters/type/summary/cool_type =====
                            if v.get('nodes'):
                                nodes = v['nodes']
                                if isinstance(nodes, list):
                                    nodes_lines = []
                                    for n in nodes[:8]:
                                        if not isinstance(n, dict):
                                            nodes_lines.append(str(n))
                                            continue
                                        n_title = n.get('title', '')
                                        n_chapters = n.get('chapters', '')
                                        n_type = n.get('type', 'M')
                                        n_summary = n.get('summary', '')
                                        n_cool = n.get('cool_type', '')
                                        line = f'  · [{n_type}] {n_chapters}章 {n_title}'
                                        if n_summary:
                                            line += f'：{n_summary}'
                                        if n_cool:
                                            line += f'（爽点：{n_cool}）'
                                        nodes_lines.append(line)
                                    if nodes_lines:
                                        parts.append('关键节点：\n' + '\n'.join(nodes_lines))
                            if v.get('volume_goal') or v.get('goal'):
                                parts.append(f'卷目标：{v.get("volume_goal") or v.get("goal")}')
                            if v.get('ending_hook'):
                                parts.append(f'卷尾钩子：{v["ending_hook"]}')
                            elif v.get('ending') or v.get('climax'):
                                parts.append(f'卷尾钩子：{v.get("ending") or v.get("climax")}')
                            if parts:
                                content = '\n'.join(parts)
                                return f'【本卷目标/卷纲】（第{volume_index}卷“{vol_label}”）\n{content[:1800]}'
            else:
                # 纯文本格式（旧数据或手动填写），直接返回
                if timeline_raw:
                    return f'【本卷目标/卷纲】（第{volume_index}卷“{vol_label}”）\n{timeline_raw[:1500]}'
    except Exception:
        pass

    # ===== 回退路径：从 Outline 表读取（手动创建的卷纲）=====
    try:
        outlines = Outline.query.filter_by(book_id=vol_chapter.book_id).order_by(Outline.order_index).all()
        for o in outlines:
            if o.title and vol_chapter.title and (o.title in vol_chapter.title or vol_chapter.title in o.title):
                if o.content and o.content.strip():
                    return f'【本卷目标/卷纲】（第{volume_index}卷“{vol_chapter.title}”）\n{o.content[:1200]}'
        acts = [o for o in outlines if o.level == 0]
        if 0 <= volume_index - 1 < len(acts):
            o = acts[volume_index - 1]
            if o.content and o.content.strip():
                return f'【本卷目标/卷纲】（第{volume_index}卷）\n{o.content[:1200]}'
    except Exception:
        pass
    return ''

def _get_adjacent_volumes_outline(book_id, volume_index):
    """获取前一卷+本卷+后一卷共三卷的卷纲摘要（用于正文写作上下文）。
    【改造】从全量 timeline 注入改为三卷注入，避免上下文膨胀且聚焦当前创作位置。
    返回拼接好的三卷卷纲文本。"""
    try:
        bb = BookBible.query.filter_by(book_id=book_id).first()
        if not bb or not bb.timeline or not bb.timeline.strip().startswith('['):
            return ''
        arr = json.loads(bb.timeline)
        if not isinstance(arr, list) or not arr:
            return ''

        # 按 volume_index 排序
        def _v_idx(v):
            if not isinstance(v, dict):
                return 0
            raw = v.get('volume_index') or _extract_volume_index(v.get('volume', v.get('volume_id', '')))
            try:
                return int(raw) if raw else 0
            except (ValueError, TypeError):
                return 0
        arr_sorted = sorted([v for v in arr if isinstance(v, dict)], key=_v_idx)

        # 提取单卷摘要（精简版，控制每卷约500字）
        def _summarize_vol(v, role_label):
            if not isinstance(v, dict):
                return ''
            v_idx = _v_idx(v)
            v_name = v.get('volume') or v.get('volume_title') or f'第{v_idx}卷'
            parts = [f'▼ [{role_label}] 第{v_idx}卷“{v_name}”']
            if v.get('act'):
                parts.append(f'  五幕定位：{v["act"]}')
            main_plot = v.get('main_plot') or v.get('core_goal') or ''
            if main_plot:
                parts.append(f'  主线：{main_plot[:200]}')
            core_conflict = v.get('core_conflict') or ''
            if core_conflict:
                parts.append(f'  核心冲突：{core_conflict[:150]}')
            # nodes 只取前5个，保留 chapters/type/title/summary
            nodes = v.get('nodes') or []
            if isinstance(nodes, list) and nodes:
                nodes_lines = []
                for n in nodes[:5]:
                    if not isinstance(n, dict):
                        continue
                    line = f'    · [{n.get("type","M")}] {n.get("chapters","")} {n.get("title","")}'
                    summary = n.get('summary', '')
                    if summary:
                        line += f'：{summary[:80]}'
                    nodes_lines.append(line)
                if nodes_lines:
                    parts.append('  情节节点：')
                    parts.extend(nodes_lines)
            ending_hook = v.get('ending_hook') or v.get('ending') or v.get('climax') or ''
            if ending_hook:
                parts.append(f'  卷尾钩子：{ending_hook[:150]}')
            return '\n'.join(parts)

        # 定位前一卷、本卷、后一卷
        prev_v, curr_v, next_v = None, None, None
        for i, v in enumerate(arr_sorted):
            if _v_idx(v) == volume_index:
                curr_v = v
                if i > 0:
                    prev_v = arr_sorted[i - 1]
                if i < len(arr_sorted) - 1:
                    next_v = arr_sorted[i + 1]
                break

        # 如果没找到本卷（如卷号超出），回退用第一卷
        if curr_v is None and arr_sorted:
            curr_v = arr_sorted[0]

        sections = []
        if prev_v:
            sections.append(_summarize_vol(prev_v, '前一卷·回顾'))
        if curr_v:
            sections.append(_summarize_vol(curr_v, '本卷·进行中'))
        if next_v:
            sections.append(_summarize_vol(next_v, '后一卷·走向'))

        if not sections:
            return ''
        return '【卷纲规划·前中后三卷】（前一卷回顾+本卷进行+后一卷走向，确保剧情连贯）\n' + '\n\n'.join(sections)
    except (json.JSONDecodeError, ValueError, TypeError):
        return ''

def _sort_foreshadowings_by_urgency(bb, vol_chapter, current_chapter_num, top_n=25):
    """伏笔按"到期紧迫度"排序，提取 Top N 待回收伏笔。
    【P2优化·百万字长线防遗忘】
    - 有计划回收章号且已过期（target < current）：最高紧迫，urgency=0（逾期必回收）
    - 有计划回收章号且将到期（target >= current）：urgency = target - current（越近越紧迫）
    - 无计划回收章号但沉淀已久（current-planted >= 50）：urgency = 500（沉淀越久越紧迫，不再被挤出）
      沉淀时长越长越往下偏移一点，保证最老的优先
    - 无计划回收章号且沉淀不久：urgency = 800
    - 扩展 Top N 从 15 到 25，覆盖更多伏笔，避免百万字时长线伏笔被挤出。
    优先从 foreshadowing_volumes 取结构化数据，否则从全局 foreshadowing 文本启发式提取。"""
    pending = []

    # 1. 优先从按卷结构化数据取
    if vol_chapter:
        try:
            arr = json.loads(bb.foreshadowing_volumes or '[]')
            for v in arr:
                if not isinstance(v, dict):
                    continue
                data = v.get('data') or {}
                items = data.get('foreshadowings') or data.get('items') or []
                for item in items:
                    if not isinstance(item, dict):
                        continue
                    status = str(item.get('status', '')).lower()
                    if any(kw in status for kw in ['已回收', '已揭示', '已解决', '已兑现', 'recycled', 'resolved']):
                        continue
                    planted = item.get('planted_chapter') or item.get('planted_at') or 0
                    target = item.get('target_chapter') or item.get('planned_recall') or item.get('recall_at') or 0
                    desc = item.get('description') or item.get('content') or item.get('title') or json.dumps(item, ensure_ascii=False)
                    try:
                        planted_n = int(planted) if planted else 0
                    except (ValueError, TypeError):
                        planted_n = 0
                    try:
                        target_n = int(target) if target else 0
                    except (ValueError, TypeError):
                        target_n = 0
                    # 紧迫度计算（数值越小越紧迫）
                    if target_n and target_n < current_chapter_num:
                        # 逾期未回收：最高紧迫
                        urgency = 0
                    elif target_n:
                        # 将到期：距离越近越紧迫
                        urgency = target_n - current_chapter_num
                    elif planted_n and (current_chapter_num - planted_n) >= 50:
                        # 无目标但沉淀已久（≥50章）：给予中等紧迫度，沉淀越久越优先
                        # 用 500 - min(沉淀时长, 200) 让最老的伏笔 urgency 低至 300
                        urgency = 500 - min(current_chapter_num - planted_n, 200)
                    elif planted_n:
                        # 无目标且沉淀不久：较低紧迫
                        urgency = 800
                    else:
                        urgency = 999
                    pending.append((urgency, planted_n, target_n, desc))
        except (json.JSONDecodeError, ValueError, TypeError):
            pass

    # 2. 回退：从全局 foreshadowing 文本启发式提取
    if not pending and bb.foreshadowing and bb.foreshadowing.strip():
        for line in bb.foreshadowing.split('\n'):
            line = line.strip()
            if not line:
                continue
            if any(kw in line for kw in ['已回收', '已揭示', '已解决', '已兑现']):
                continue
            # 启发式：从文本中提取章号
            import re as _re_fs
            nums = _re_fs.findall(r'第?(\d+)\s*章', line)
            if nums:
                target_n = int(nums[-1])
                planted_n = int(nums[0]) if len(nums) > 1 else 0
                if target_n < current_chapter_num:
                    urgency = 0
                else:
                    urgency = target_n - current_chapter_num
            else:
                target_n = 0
                planted_n = 0
                urgency = 800
            pending.append((urgency, planted_n, target_n, line))

    # 按紧迫度升序排序，取 Top N
    pending.sort(key=lambda x: x[0])
    return pending[:top_n]

def _extract_appearing_characters(recent_chapters, bb=None):
    """从最近章节正文中启发式提取出场角色名。
    P2增强：融合 chapter_changes_log 的结构化角色名（CharacterStateChanges.CharacterId），
    比纯正则更准——结构化数据是章后回写的权威角色记录，正则只能启发式猜测。
    简单策略：识别"X说"、"X道"等模式中的 X + chapter_changes_log 的角色名。"""
    names = set()
    import re as _re_char

    # 1. 从 chapter_changes_log 提取结构化角色名（最近20章，权威来源）
    if bb and getattr(bb, 'chapter_changes_log', ''):
        try:
            log_list = json.loads(bb.chapter_changes_log)
            if isinstance(log_list, list):
                for entry in log_list[-20:]:
                    if not isinstance(entry, dict):
                        continue
                    chg = entry.get('changes') or {}
                    if not isinstance(chg, dict):
                        continue
                    for c in (chg.get('CharacterStateChanges') or []):
                        if isinstance(c, dict):
                            nm = c.get('CharacterId') or c.get('Name') or ''
                            if nm and isinstance(nm, str) and 2 <= len(nm) <= 8:
                                names.add(nm.strip())
                    # 物品转移的持有者也视为出场角色
                    for it in (chg.get('ItemTransfers') or []):
                        if isinstance(it, dict):
                            for holder_key in ('FromHolder', 'ToHolder'):
                                holder = it.get(holder_key) or ''
                                if holder and isinstance(holder, str) and 2 <= len(holder) <= 8:
                                    names.add(holder.strip())
        except (json.JSONDecodeError, ValueError, TypeError):
            pass

    # 2. 从最近章节正文启发式提取（正则，补充结构化数据未覆盖的角色）
    if recent_chapters:
        text = '\n'.join([(c.content or '') for c in recent_chapters])
        # 中文姓名 2-4 字
        for m in _re_char.finditer(r'([\u4e00-\u9fa5]{2,4})(?:说|道|想|看|笑|怒|惊|叹|问|答|吼|喊|冷哼|微笑|皱眉)', text):
            name = m.group(1)
            # 排除常见动词误判
            if name not in {'这是', '那是', '于是', '然后', '突然', '只见', '心想', '不禁', '不由'}:
                names.add(name)
    return names

def _compute_style_baseline(book_id, current_chapter_num, sample_count=5):
    """计算前 N 章的文风指纹基准（借鉴 PlotPilot 文风指纹漂移检测）。
    取当前章之前最近 sample_count 章的正文，计算文风指纹平均值作为基准。
    返回基准指纹 dict 或 None（章节数不足时）。"""
    if not compute_style_fingerprint or current_chapter_num <= 2:
        return None
    try:
        # 取前 N 章正文（排除当前章和分卷占位）
        prev_chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).filter(
            Chapter.order_index < current_chapter_num,
            Chapter.order_index >= max(1, current_chapter_num - sample_count)
        ).order_by(Chapter.order_index.desc()).limit(sample_count).all()
        if not prev_chapters:
            return None
        fingerprints = []
        for ch in prev_chapters:
            content = ch.content or ''
            # 提取正文（剥离 chapter_changes 等标签）
            body = re.sub(r'<chapter_changes>[\s\S]*?</chapter_changes>', '', content, flags=re.IGNORECASE)
            body = re.sub(r'<pre_write_check>[\s\S]*?</pre_write_check>', '', body, flags=re.IGNORECASE)
            body = body.strip()
            if len(body) < 100:
                continue
            fp = compute_style_fingerprint(body)
            if fp:
                fingerprints.append(fp)
        if len(fingerprints) < 2:
            return None  # 样本不足
        # 计算各特征平均值
        keys = ['avg_sent_len', 'short_sent_ratio', 'long_sent_ratio', 'dialog_ratio', 'punct_density', 'adj_density', 'verb_density']
        baseline = {}
        for key in keys:
            vals = [fp.get(key, 0) for fp in fingerprints if key in fp]
            if vals:
                baseline[key] = round(sum(vals) / len(vals), 4)
        return baseline if baseline else None
    except Exception:
        return None

def _recall_related_chapters(book_id, appearing_chars, current_chapter_num, max_chapters=6):
    """轻量RAG：基于出场角色召回相关历史章节摘要，补充前4章完整正文窗口的盲区。
    策略：扫描所有历史章节（排除最近4章，避免与即时层重复）的 summary，
    召回 summary 中含当前章出场角色的章节，按章号降序取最近 max_chapters 条。
    无 summary 的章节跳过（无法判断相关性）。
    返回 [{'chapter_num','title','summary'}] 列表。"""
    if not appearing_chars:
        return []
    try:
        # 排除最近4章（即时层已覆盖），排除分卷占位章
        threshold = max(1, current_chapter_num - 4)
        candidates = Chapter.query.filter(
            Chapter.book_id == book_id,
            Chapter.is_volume == False,
            Chapter.order_index < threshold
        ).order_by(Chapter.order_index.desc()).limit(200).all()  # 扫描最近200章历史
    except Exception:
        return []

    recalled = []
    seen_chapters = set()
    for ch in candidates:
        summary = (getattr(ch, 'summary', '') or '').strip()
        if not summary or len(summary) < 10:
            continue  # 无 summary 无法判断相关性
        # 相关性判定：summary 含当前章出场角色名
        if any(name in summary for name in appearing_chars):
            if ch.order_index not in seen_chapters:
                seen_chapters.add(ch.order_index)
                recalled.append({
                    'chapter_num': ch.order_index,
                    'title': ch.title or '',
                    'summary': summary[:200],  # summary 字段上限 200 字
                })
            if len(recalled) >= max_chapters:
                break
    # 按章号升序排列（剧情时间顺序）
    recalled.sort(key=lambda x: x['chapter_num'])
    return recalled


# ============================================================================
# PromptContextCache：智驾设定/正文 prompt 上下文缓存（先命中再逐维度读资料 省 token）
# - 失效策略：指纹 Key（零脏读）—— Key 本身就包含所有依赖的 hash，任一依赖变 Key 自动
#   变 → 自动 MISS 重算，不需要维护任何写路由的失效 hook（漏一个 hook 就脏读灾难）
# - 介质：进程内 LRU + TTL（OrderedDict），多 worker / 进程重启不共享，但退化为原直读
#   逻辑不影响功能；后续要切 SQLite 持久化只换本类实现，调用处 0 改动
# - 旁路：skip_cache=True 强制绕开；响应头 X-Prompt-Cache HIT/MISS + X-Tokens-Saved
# ============================================================================
from collections import OrderedDict
import threading, time


def _approx_tokens(*texts: str) -> int:
    """近似 token 计数（中文≈1.5-2字/token → //4 估算，足以便捷对比命中前后）"""
    total = 0
    for t in texts:
        if t:
            total += max(1, len(t) // 4)
    return total


class PromptContextCache:
    """单例：get_or_compute(stage, book_id, deps, compute_fn, ttl_sec=600, skip=False)
    - stage: 'continue_ctx' / 'outline_master' / 'outline_volume' / 'master_dim' 等命名空间
    - deps:  参与指纹的稳定字符串/tuple/list/dict（会做递归稳定序列化）
    """
    _instance = None
    _instance_lock = threading.Lock()

    @classmethod
    def get(cls):
        if cls._instance is None:
            with cls._instance_lock:
                if cls._instance is None:
                    cls._instance = PromptContextCache(max_items=2048)
        return cls._instance

    def __init__(self, max_items: int = 2048):
        self.max_items = max_items
        self._store: 'OrderedDict[str, tuple]' = OrderedDict()
        self._lock = threading.Lock()
        # stats
        self.hits = 0
        self.misses = 0
        self.evicted = 0
        self.tokens_saved = 0

    # ---------- 稳定哈希（指纹 Key 用）----------
    @staticmethod
    def _stable_dumps(obj) -> bytes:
        if isinstance(obj, dict):
            items = sorted(((k, PromptContextCache._stable_dumps(v)) for k, v in obj.items()), key=lambda x: x[0])
            return b'd{' + b'\x1f'.join(k.encode('utf-8') + b'\x1e' + v for k, v in items) + b'}'
        if isinstance(obj, (list, tuple)):
            tag = b'l[' if isinstance(obj, list) else b't('
            return tag + b'\x1f'.join(PromptContextCache._stable_dumps(x) for x in obj) + (b']' if isinstance(obj, list) else b')')
        if obj is None:
            return b'N'
        if isinstance(obj, bool):
            return b'T' if obj else b'F'
        if isinstance(obj, (int, float)):
            return repr(obj).encode('utf-8')
        return str(obj).encode('utf-8')

    @classmethod
    def make_key(cls, stage: str, book_id, deps) -> str:
        raw = b'stage=' + stage.encode('utf-8') + b'|book=' + str(book_id).encode('utf-8') + b'|deps=' + cls._stable_dumps(deps)
        return stage + ':' + str(book_id) + ':' + hashlib.sha256(raw).hexdigest()[:16]

    # ---------- 主接口 ----------
    def get_or_compute(self, stage: str, book_id, deps, compute_fn, *, ttl_sec: int = 600, skip_cache: bool = False):
        """return (payload, cache_info_dict)  cache_info = {'hit':bool,'tokens_saved':int,'key':str,'ttl':int}"""
        key = self.make_key(stage, book_id, deps)
        info = {'hit': False, 'tokens_saved': 0, 'key': key, 'ttl': ttl_sec}
        now = time.time()
        if not skip_cache:
            with self._lock:
                entry = self._store.get(key)
                if entry is not None:
                    payload, expire_at, tok_cost = entry
                    if expire_at > now:
                        self._store.move_to_end(key)
                        self.hits += 1
                        self.tokens_saved += tok_cost
                        info['hit'] = True
                        info['tokens_saved'] = tok_cost
                        return payload, info
                    else:
                        try:
                            del self._store[key]
                            self.evicted += 1
                        except Exception:
                            pass
        # MISS / skip
        payload = compute_fn()  # compute_fn 负责真正从 bible/章节/维度 读资料 拼 prompt
        tok_cost = self._estimate_payload_tokens(payload)
        expire_at = now + ttl_sec
        with self._lock:
            self._store[key] = (payload, expire_at, tok_cost)
            self._store.move_to_end(key)
            while len(self._store) > self.max_items:
                self._store.popitem(last=False)
                self.evicted += 1
            self.misses += 1
        return payload, info

    @staticmethod
    def _estimate_payload_tokens(payload) -> int:
        if payload is None:
            return 0
        if isinstance(payload, str):
            return _approx_tokens(payload)
        if isinstance(payload, dict):
            texts = []
            for v in payload.values():
                if isinstance(v, str):
                    texts.append(v)
                elif isinstance(v, (list, tuple)):
                    texts.extend(str(x) for x in v if isinstance(x, (str, int, float)))
            return _approx_tokens(*texts)
        if isinstance(payload, (list, tuple)):
            return _approx_tokens(*(str(x) for x in payload if isinstance(x, (str, int, float))))
        return _approx_tokens(str(payload))

    # ---------- API stats ----------
    def stats(self) -> Dict[str, Any]:
        with self._lock:
            size = len(self._store)
        return {
            'size': size, 'max_items': self.max_items,
            'hits': self.hits, 'misses': self.misses,
            'evicted': self.evicted, 'tokens_saved_approx': self.tokens_saved,
            'hit_rate_pct': round(100 * self.hits / max(1, self.hits + self.misses), 2),
        }

    def invalidate_book(self, book_id) -> int:
        """手动清理整本书的所有缓存（极端场景：改了设定但 MISS 还没触发时用）。返回清理条数"""
        prefix = ':' + str(book_id) + ':'
        removed = 0
        with self._lock:
            keys = [k for k in self._store.keys() if prefix in k]
            for k in keys:
                self._store.pop(k, None)
                removed += 1
        return removed

    def clear_all(self):
        with self._lock:
            self._store.clear()
            self.hits = 0
            self.misses = 0
            self.evicted = 0
            self.tokens_saved = 0


_PROMPT_CACHE_HELPERS_INJECTED = True

def _response_with_cache(resp, info: Optional[dict]):
    """给 flask.Response 加 X-Prompt-Cache: HIT/MISS + X-Tokens-Saved 头"""
    if not info:
        return resp
    try:
        resp.headers['X-Prompt-Cache'] = 'HIT' if info.get('hit') else 'MISS'
        if info.get('tokens_saved'):
            resp.headers['X-Tokens-Saved'] = str(int(info['tokens_saved']))
        resp.headers['X-Cache-Key-Tail'] = str(info.get('key', ''))[-8:]
    except Exception:
        pass
    return resp


def _cache_stats_snapshot():
    return PromptContextCache.get().stats()


# ============================================================================

def _filter_bible_by_relevance(bb, appearing_chars, max_per_field=None):
    """按出场角色相关性筛选 bible 维度。
    - character_profiles: 优先包含出场角色的档案块
    - 其他维度保持原样（不相关性筛选）
    返回筛选后的字段字典。"""
    if max_per_field is None:
        max_per_field = {'character_profiles': 1500, 'worldbuilding': 1000, 'plot_design': 1000,
                         'timeline': 800, 'concept': 500, 'key_rules': 1200, 'style_guide': 500}
    result = {}

    # key_rules / worldbuilding / concept / plot_design / timeline / style_guide：语义截断（保证段落完整）
    for field in ['key_rules', 'worldbuilding', 'concept', 'plot_design', 'timeline', 'style_guide']:
        val = getattr(bb, field, '') or ''
        result[field] = _smart_truncate(val, max_per_field.get(field, 1000)) if val else ''

    # character_profiles：按出场角色筛选
    cp = bb.character_profiles or ''
    if cp and appearing_chars:
        # 简单分块：按空行或【】标题分块
        import re as _re_cp
        blocks = _re_cp.split(r'\n(?=[【\[])', cp)
        relevant = []
        other = []
        for block in blocks:
            block = block.strip()
            if not block:
                continue
            if any(name in block for name in appearing_chars):
                relevant.append(block)
            else:
                other.append(block)
        # 优先相关性块，剩余空间补其他
        budget = max_per_field.get('character_profiles', 1500)
        selected = []
        used = 0
        for block in relevant:
            if used + len(block) > budget:
                # 截断该块
                remain = budget - used
                if remain > 100:
                    selected.append(block[:remain])
                    used = budget
                break
            selected.append(block)
            used += len(block)
        # 剩余预算补其他
        for block in other:
            if used >= budget:
                break
            remain = budget - used
            if remain < 100:
                break
            truncated = block[:remain]
            selected.append(truncated)
            used += len(truncated)
        result['character_profiles'] = '\n\n'.join(selected) if selected else cp[:budget]
    else:
        result['character_profiles'] = (cp or '')[:max_per_field.get('character_profiles', 1500)]

    return result

def _collect_relevant_reports(book_id, current_chapter_num, window=10, max_reports=3, per_report_limit=800):
    """收集当前章号 ±window 范围内的动态报告，每份截取摘要。
    优先时效性（接近当前章号），其次数量上限。"""
    try:
        reports = DynamicReport.query.filter_by(book_id=book_id).order_by(DynamicReport.chapter_start).all()
    except Exception:
        return []
    if not reports:
        return []
    # 筛选在窗口内的报告
    relevant = []
    for r in reports:
        # 报告覆盖区间 [chapter_start, chapter_end] 与 [current-window, current+window] 有交集
        if r.chapter_end >= current_chapter_num - window and r.chapter_start <= current_chapter_num + window:
            relevant.append(r)
    # 按接近度排序（距离当前章号最近者优先）
    relevant.sort(key=lambda r: abs((r.chapter_start + r.chapter_end) / 2 - current_chapter_num))
    # 取前 max_reports 份，每份截取摘要
    result = []
    for r in relevant[:max_reports]:
        content = (r.content or '')[:per_report_limit]
        result.append({
            'title': r.title,
            'chapter_start': r.chapter_start,
            'chapter_end': r.chapter_end,
            'content': content,
        })
    return result

def _collect_historical_volume_digest(book_id, current_chapter_num, max_volumes=4, per_volume_limit=400):
    """百万字长线记忆：聚合"早期卷"的动态报告为按卷摘要，补充 _collect_relevant_reports 的窗口盲区。
    当 current_chapter_num 较大（如300章）时，window=100 只能覆盖最近100章，
    本函数把更早的卷（如第1-4卷）的动态报告按卷聚合，每卷取一份摘要注入，避免长线剧情遗忘。
    只聚合 chapter_end < current - 100 的早期报告（不与 _collect_relevant_reports 重叠）。"""
    try:
        threshold = current_chapter_num - 100
        if threshold <= 50:
            return []  # 章数不多时无需历史摘要
        early_reports = DynamicReport.query.filter(
            DynamicReport.book_id == book_id,
            DynamicReport.chapter_end < threshold
        ).order_by(DynamicReport.chapter_start).all()
    except Exception:
        return []
    if not early_reports:
        return []
    # 按 50 章一卷分组聚合
    volume_buckets = {}
    for r in early_reports:
        vol_idx = (r.chapter_start - 1) // 50 + 1 if r.chapter_start > 0 else 1
        volume_buckets.setdefault(vol_idx, []).append(r)
    # 每卷取第一份报告的摘要作为该卷代表（最早的报告概括卷的开端，最能代表该卷基调）
    result = []
    for vol_idx in sorted(volume_buckets.keys())[:max_volumes]:
        reports_in_vol = volume_buckets[vol_idx]
        # 取该卷所有报告内容的拼接摘要
        combined = ' '.join([(r.content or '')[:150] for r in reports_in_vol[:3]])
        combined = combined[:per_volume_limit]
        ch_start = reports_in_vol[0].chapter_start
        ch_end = reports_in_vol[-1].chapter_end
        result.append({
            'volume': f'第{vol_idx}卷',
            'chapter_range': f'{ch_start}-{ch_end}章',
            'digest': combined,
        })
    return result

# ===== Ai总创作公共常量（被 ai_master_create 和 ai_master_create_stream 共用，避免重复定义）=====
# 维度协同顺序：上游维度产物会注入到下游维度的 prompt（严格按 DAG 依赖，不再无差别全注入）
MASTER_DIM_ORDER = ['concept', 'key_rules', 'worldbuilding', 'character_profiles', 'plot_design', 'timeline', 'foreshadowing', 'locations', 'inventory']

# 每个维度的上游依赖（只注入这些维度作为"已确认上游产物"，而非所有维度）
MASTER_DIM_UPSTREAM = {
    'concept': [],
    'key_rules': ['concept'],
    'worldbuilding': ['concept', 'key_rules'],
    'character_profiles': ['concept', 'key_rules', 'worldbuilding'],
    'plot_design': ['concept', 'key_rules', 'worldbuilding', 'character_profiles'],
    'timeline': ['concept', 'key_rules', 'worldbuilding', 'character_profiles', 'plot_design'],
    'foreshadowing': ['plot_design', 'timeline', 'character_profiles'],
    'locations': ['worldbuilding', 'character_profiles'],
    'inventory': ['worldbuilding', 'character_profiles', 'locations'],
}

MASTER_DIM_MAP = {
    'concept': {'field': 'concept', 'label': '构思', 'keys': ['tomato_plan', 'one_line_concept', 'master_outline'],
                'prompt': '生成核心构思。\n【输出格式】纯文本分项（空行分隔）：1) 一句话核心概念；2) 核心卖点（3-5条）；3) 目标读者画像；4) 主线冲突；5) 独特亮点。'},
    'key_rules': {'field': 'key_rules', 'label': '设定/规则', 'keys': ['tomato_setting', 'lock_facts'],
                  'prompt': '生成核心设定规则。\n【输出格式】编号列表，每条以"① ② ③..."开头，规则间空行分隔。包括：世界必须遵循的铁律、人物能力边界、代价/反噬机制、禁忌事项。'},
    'worldbuilding': {'field': 'worldbuilding', 'label': '世界观', 'keys': ['tomato_setting', 'lock_facts'],
                      'prompt': '生成详细世界观。\n【输出格式】分小节二级标题"## 力量体系/## 社会结构/## 地理概况/## 历史脉络"，每小节下用编号列表。不要写成段落散文。'},
    'character_profiles': {'field': 'character_profiles', 'label': '人物', 'keys': ['tomato_character', 'character_cognition'],
                           'prompt': '生成主要人物。\n【输出格式】每个角色一个"## 角色：<姓名>"二级标题，下方依次：身份/性格（3-5个关键词）/背景故事（100-200字）/核心动机/与其他角色关系（"→ 角色名：关系"）。主角 + 3-5个配角。'},
    'plot_design': {'field': 'plot_design', 'label': '大纲', 'keys': ['master_outline', 'tomato_outline', 'volume_breakdown'],
                    'prompt': '生成五幕式总纲。\n【输出格式】每幕"## 第N幕：<幕名>"二级标题，下方：幕核心目标/主要冲突/卷入角色/关键转折点/幕尾悬念/对应分卷范围（"第X-Y卷"）。共5幕，对应全书所有分卷。'},
    'timeline': {'field': 'timeline', 'label': '剧情', 'keys': ['volume_breakdown', 'chapter_plan', 'tomato_outline'],
                 'prompt': '生成分卷详细剧情。\n【分卷铁律·必读】**每卷固定 50 章**，全书卷数 = 总章数÷50（向上取整），卷序号从1开始连续。卷名格式"第N卷 副标题"。\n【输出格式·必读】严格 JSON 数组，不要任何解释文字、不要 markdown 代码块。结构：[{"volume_index":1,"volume":"第1卷 副标题","main_plot":"本卷主线剧情100-200字","core_conflict":"本卷核心冲突","ending_hook":"卷尾钩子","nodes":[{"title":"节点1","chapters":"1-10","type":"M","summary":"概要","cool_type":"实力碾压"}]}]. 每卷 5-8 个 nodes；chapters 字段"起始-结束"，全书 chapter 编号连续不重叠。'},
    'foreshadowing': {'field': 'foreshadowing', 'label': '伏笔', 'keys': ['foreshadow_register', 'narrative_debt'],
                      'prompt': '埋设伏笔线索。\n【输出格式】编号列表，每条"## 伏笔N：<标题>\\n- 埋设内容：xxx\\n- 埋设时机：第X卷Y章附近\\n- 预期回收：第X卷Y章附近\\n- 回收方式：xxx\\n- 对剧情的影响：xxx"。设计 3-5 条。'},
    'locations': {'field': 'locations', 'label': '地点', 'keys': ['lock_facts', 'tomato_setting', 'geography'],
                  'prompt': '设计地点体系。\n【输出格式】严格 JSON 数组三级结构：[一级大区域{"name","description","secondaries":[{"name","description","scenes":[{"name","description","key_events"}]}]}]. 设计 2-3 个一级大区域。'},
    'inventory': {'field': 'inventory', 'label': '物资库', 'keys': ['lock_facts', 'level_system', 'power_system', 'ability_tree'],
                  'prompt': '生成主要物品/功法/法宝清单。\n【输出格式】严格 JSON 数组：[{"name","type","source","effect","owner","first_appearance"}]. type 取值：法宝/功法/丹药/武器/防具/其他。设计 8-15 个核心物品。'},
}

def _build_master_ctx(bb, session_outputs=None):
    """构建维度创作的上下文字典：本轮已生成内容(session_outputs)优先，回退 bible 已有内容。
    【P2-7修复】session_outputs 纯空白字符串视为无效，回退到 bb 已落库值。"""
    session_outputs = session_outputs or {}

    def _pick(dim, bb_field):
        v = session_outputs.get(dim)
        # 纯空白字符串视为无效，避免误判上游已完成
        if v and isinstance(v, str) and v.strip():
            return v
        return bb_field or ''
    return {
        'concept': _pick('concept', bb.concept),
        'key_rules': _pick('key_rules', bb.key_rules),
        'worldbuilding': _pick('worldbuilding', bb.worldbuilding),
        'character_profiles': _pick('character_profiles', bb.character_profiles),
        'plot_design': _pick('plot_design', bb.plot_design),
        'timeline': _pick('timeline', bb.timeline),
        'foreshadowing': _pick('foreshadowing', bb.foreshadowing),
        'locations': _pick('locations', bb.locations),
        'inventory': _pick('inventory', bb.inventory),
    }

def _validate_and_align_timeline_volumes(parsed_tl, total_volumes):
    """【P0-1修复】timeline 卷数校验与对齐：解析成功后强制对齐 total_volumes。
    - 卷数匹配：直接返回
    - 卷数过多：截断到 total_volumes，标记 warning
    - 卷数过少：【升级】自动补齐占位卷到 total_volumes，避免用户设 25 卷却只拿到 10 卷导致下游错位。
               占位卷会依据五幕比例打上 act、卷名与占位 main_plot，提示后续手动或节点设计再充实。
    返回 (aligned_list, warning_msg or None)。
    """
    if not isinstance(parsed_tl, list):
        return parsed_tl, None
    if not total_volumes or total_volumes < 1:
        return parsed_tl, None
    actual = len(parsed_tl)
    if actual == total_volumes:
        return parsed_tl, None
    if actual > total_volumes:
        # 卷数过多：截断
        aligned = parsed_tl[:total_volumes]
        return aligned, f'timeline生成了{actual}卷，超过设定的{total_volumes}卷，已自动截断到{total_volumes}卷'
    # 卷数过少：自动补齐占位到 total_volumes（按比例+相邻卷字段做合理占位）
    def map_act(vidx, tot):
        pct = (vidx - 0.5) / max(1, tot)
        if pct <= 0.05: return '立身'
        if pct <= 0.25: return '立足'
        if pct <= 0.50: return '立势'
        if pct <= 0.75: return '立威'
        return '立命'
    # 用已存在的最后一卷做“相邻卷参考”；没有就造个基准
    ref = parsed_tl[-1] if parsed_tl else {}
    aligned = list(parsed_tl)
    for idx in range(actual + 1, total_volumes + 1):
        act = map_act(idx, total_volumes)
        placeholder = {
            'volume_id': str(idx),
            'volume': f'第{idx}卷（AI未生成足够卷数已自动补齐占位，请在节点设计阶段重写）',
            'volume_index': idx,
            'act': act,
            'main_plot': f'【占位卷】所属幕：{act}。当前AI输出卷数不足用户设定的{total_volumes}卷，已按{total_volumes}卷硬约束自动补齐，请手动调整本卷剧情或重新执行分卷提取。',
            'core_conflict': ref.get('core_conflict', '待补充'),
            'emotion_driver': ref.get('emotion_driver', '待补充'),
            'key_events': ['待补充关键事件'],
            'turning_points': ['待补充转折点'],
            'climax': '待补充',
            'ending': '待补充',
            'ending_hook': '待补充',
            'foreshadowing': [],
            'nodes': [],
        }
        aligned.append(placeholder)
    return aligned, f'timeline仅生成{actual}卷，少于用户设定的{total_volumes}卷，已自动补齐到{total_volumes}卷（后{total_volumes - actual}卷为占位，建议手动重写补齐或重新执行分卷提取）'

# 网文风格流派标签库（基于2025中国网络文学蓝皮书与起点三江榜趋势）
# 题材 → 风格流派映射表（长篇）
# 数据来源：番茄小说/起点中文网/七猫小说/晋江文学城等主流平台分类页调研整理
# key 为风格短码，value 为网文圈通用中文叫法
NOVEL_GENRE_STYLES = {
    'fantasy': {
        'dongfang': '东方玄幻', 'yishi': '异世大陆', 'gaowu': '高武世界', 'wangchao': '王朝争霸',
        'honghuang': '洪荒流', 'fanren': '凡人流', 'feichai': '废柴逆袭流', 'qiangzhong': '强者重生流',
        'qiandao': '签到流', 'shenchong': '神宠流', 'dihua': '迪化流', 'shenshu': '全球神祇流',
    },
    'xianxia': {
        'gudian': '古典仙侠', 'xiuzhen': '修真文明', 'huanxiang': '幻想修仙', 'xiandai': '现代修真',
        'shenhua': '神话修真', 'fengshen': '洪荒封神', 'goudao': '苟道流', 'changsheng': '长生流',
        'jiazu': '家族修仙流', 'jianxiu': '剑修流', 'liandan': '炼丹流', 'liangi': '炼器流',
    },
    'qihuan': {
        'jianmo': '剑与魔法', 'shishi': '史诗奇幻', 'shenmi': '神秘幻想', 'xiandai': '现代魔法',
        'lishi': '历史神话', 'xifang': '西方奇幻', 'wushi': '巫师流', 'lingzhu': '领主贵族',
        'mofa': '魔法校园', 'zhongshi': '中式奇幻',
    },
    'wuxia': {
        'chuantong': '传统武侠', 'wuxia_fs': '武侠幻想', 'guoshu': '国术流', 'xinpa': '新派武侠',
        'lishi_wx': '历史武侠', 'langzi': '浪子异侠', 'kuaiyi': '快意江湖', 'jianghu': '江湖恩怨',
    },
    'urban': {
        'shenghuo': '都市生活', 'yishu': '异术超能', 'qingchun': '青春校园', 'mingxing': '娱乐明星',
        'shangzhan': '商战职场', 'guanchang': '官场沉浮', 'dushi_xz': '都市修真', 'dushi_gw': '都市高武',
        'shenyi': '神医流', 'shenhao': '神豪流', 'jianbao': '鉴宝流', 'bingwang': '兵王回归', 'cunzhi': '乡村种田',
    },
    'urban_business': {
        'shangzhan': '商战职场', 'chuangye': '创业逆袭', 'zhichang': '职场权谋',
        'shangye': '商业帝国', 'zhulian': '珠联璧合', 'jindiao': '金融大鳄',
    },
    'urban_fantasy': {
        'yishu': '异术超能', 'dushi_xz': '都市修真', 'dushi_gw': '都市高武',
        'guidze': '规则怪谈', 'dushi_nr': '都市脑洞', 'lingyi': '灵异民俗',
    },
    'history': {
        'jiakong': '架空历史', 'qhsg': '秦汉三国', 'tangsong': '两晋隋唐', 'wudai': '五代十国',
        'songming': '两宋元明', 'qingmg': '清史民国', 'chuanyue_ls': '穿越历史', 'keju': '科举入仕',
        'zhongtian_ls': '历史种田', 'quanmou': '权谋庙堂',
    },
    'military': {
        'junlv': '军旅生涯', 'zhanzheng': '军事战争', 'kangzhan': '抗战烽火', 'diedz': '谍战特工',
        'tezhong': '特种军旅', 'xiandai_zz': '现代战争', 'chuanyue_zz': '穿越战争',
    },
    'game': {
        'djj': '电子竞技', 'wlw': '虚拟网游', 'youxi_yj': '游戏异界', 'youxi_xt': '游戏系统',
        'quantxi': '全息网游', 'disitianzai': '第四天灾流', 'shuju': '数据流', 'zhandui': '战队夺冠',
    },
    'sports': {
        'zuqiu': '足球运动', 'lanqiu': '篮球运动', 'wangqiu': '网球/乒乓球', 'zonghe_ty': '综合竞技',
        'dianjing_ty': '电竞体育', 'rexue_ty': '热血竞技', 'xiaoyuan_ty': '体育校园',
    },
    'scifi': {
        'xingji': '星际文明', 'weilai': '未来世界', 'chaojikj': '超级科技', 'shikong': '时空穿梭',
        'jinhua': '进化变异', 'moshi': '末世危机', 'jijia': '古武机甲', 'saibo': '赛博朋克',
        'feitu': '废土生存', 'xingji_zz': '星际战争', 'heikeji': '黑科技系统',
    },
    'mystery': {
        'zhentan': '侦探推理', 'guilyi': '诡异神秘', 'guize': '规则怪谈', 'lingyi_ms': '灵异民俗',
        'fengshui': '风水秘术', 'xingzhen': '刑侦破案', 'daoshu': '道术流', 'kesulu': '克苏鲁',
        'xunyi': '悬疑探险', 'xisikongjv': '细思极恐', 'shourong': '灵异收容',
    },
    'infinite': {
        'wuxian': '无限流', 'zhutian': '诸天流', 'zongman': '综漫', 'yuanzu': '元祖无限流',
        'zhushen': '主神流', 'kuaichuan_zt': '快穿诸天', 'yingshi_ct': '影视世界穿越',
        'dongman_ct': '动漫世界穿越', 'fuben': '副本闯关',
    },
    'light_novel': {
        'yuansheng': '原生幻想', 'yansheng': '衍生同人', 'gaoxiao': '搞笑吐槽', 'liana': '恋爱日常',
        'erciyuan': '二次元', 'rixi': '日系轻改', 'zhonger': '中二设定', 'shacao': '沙雕轻松',
        'mengxi': '萌系', 'yishijie': '异世界',
    },
    # 女频
    'romance': {
        'dushi_tc': '都市甜宠', 'haozong': '豪门总裁', 'xianhun': '先婚后爱', 'pojing': '破镜重圆',
        'zhuqi': '追妻火葬场', 'bazong': '总裁霸总', 'yulequan': '娱乐圈', 'zhichang_hl': '职场婚恋',
        'niandai': '年代文', 'xianhun_yw': '闪婚', 'xiaoyuan_qc': '校园青春', 'nuelian': '虐恋情深',
        'chongsheng_nx': '重生逆袭', 'kuaichuan_yq': '快穿',
    },
    'ancient_romance': {
        'gongdou': '宫斗', 'zhaidou': '宅斗', 'gufeng': '古风世情', 'gudai_ct': '古代穿越',
        'shunv': '庶女逆袭', 'dinu': '嫡女', 'quanchen': '权臣', 'jiangjun': '将军', 'wangye': '王爷',
        'daihou': '帝后', 'zhongtian_gy': '种田经商', 'chaoztang': '朝堂权谋', 'daijia': '代嫁代娶', 'chongsheng_gy': '穿越重生',
    },
    'fantasy_romance': {
        'xuanhuan_yq': '玄幻言情', 'qihuan_yq': '奇幻言情', 'xianxia_yq': '仙侠言情', 'xiuxian_yq': '修仙言情',
        'xuanxue': '玄学相师', 'lingyi_yq': '灵异言情', 'yineng_nv': '异能女主', 'xitong_yq': '系统言情',
        'chuanshu': '穿书', 'weilai_yq': '未来言情',
    },
    'danmei': {
        'xiandai_ca': '现代都市纯爱', 'gudai_ca': '古代纯爱', 'xiangxiang_ca': '现代幻想纯爱', 'ab0': 'ABO',
        'qiangqiang': '强强', 'tianwen_ca': '甜文', 'nuewen_ca': '虐文', 'xiaoyuan_ca': '校园',
        'dianjing_ca': '电竞', 'xianxia_ca': '仙侠纯爱', 'wuxian_ca': '无限流纯爱', 'kuaichuan_ca': '快穿纯爱',
    },
    'acg': {
        'dongfang_ys': '东方衍生', 'xifang_ys': '西方衍生', 'gudian_ys': '古典衍生', 'erciyuan_yq': '二次元言情',
        'zongying': '综英美', 'zongwuxia': '综武侠', 'zongman_ys': '综漫', 'yingshi_tr': '影视同人',
        'youxi_tr': '游戏同人', 'dongman_tr': '动漫同人',
    },
    'other': {
        'xiangsheng': '相声评书', 'sanwen': '散文随笔', 'pinglun': '评论文集',
        'youji': '美文游记', 'shige': '诗歌', 'weixiaoshuo': '微小说',
    },
}

# 题材 → 风格流派映射表（短篇，以知乎盐选/盐言故事赛道为主）
SHORT_GENRE_STYLES = {
    'romance': {
        'hunyin': '婚姻信任崩塌', 'poxi': '婆媳边界', 'zhichang_pu': '职场PUA反杀', 'yuansheng': '原生家庭拉扯',
        'zhongnian': '中年离婚重启', 'chongnianv': '重男轻女', 'jiating': '家庭伦理', 'chongsheng_dy': '重生打脸爽文',
    },
    'ancient_romance': {
        'chongsheng_fc': '重生复仇', 'shunv_nx': '庶女逆袭', 'dinu_fp': '嫡姐反派', 'qianshi': '前世惨死今生逆袭',
        'yinren': '隐忍蛰伏摊牌打脸', 'qihun': '弃婚另嫁', 'gongdou_dp': '宫斗短篇', 'daijia_dp': '代嫁代娶', 'xianhun_dp': '先婚后爱',
    },
    'fantasy_romance': {
        'bazong_zq': '霸总追妻火葬场', 'baiyueguang': '白月光替身', 'jingshen': '净身出户远走', 'chongfeng': '重逢反差',
        'lihun': '离婚后惊艳世界', 'shanhun': '闪婚大佬', 'zongcai_bw': '总裁卑微求和', 'zhichang_tc': '职场甜宠', 'xiaoyuan_cl': '校园初恋',
    },
    'mystery': {
        'shenghuohua': '生活化悬疑', 'xisikongju': '细思极恐', 'shikong_laidian': '时空来电', 'wanmei': '完美犯罪',
        'duochong_fz': '多重反转', 'xiongsha': '凶杀推理', 'mishi': '密室', 'guize_dp': '规则怪谈短篇',
        'minshu_dp': '民俗志怪', 'jiankong': '监控悬疑',
    },
    'urban_fantasy': {
        'lingchen': '凌晨怪事', 'jiuwu': '旧物惊悚', 'chuzu': '出租屋灵异', 'laoxiaoqu': '老小区鬼事',
        'fengshui_dp': '风水秘术', 'daoshi': '道士收妖', 'hongyi': '红衣女人', 'mintan': '民间怪谈',
    },
    'light_novel': {
        'qingshed': '轻设定重映射', 'huangyan': '谎言可视化', 'danmu': '弹幕生存', 'chaoshi': '超现实外壳',
        'she_ding': '设定撑满五千字', 'qiguai': '奇怪群聊', 'zouma': '死亡走马灯', 'gonglve': '攻略者大逃杀',
    },
    'other': {
        'qingganzhiyu': '情感治愈', 'renchong': '人宠双向疗愈', 'xishui': '细水流长', 'yinanping': '意难平',
        'houjin': '后劲极强', 'wenqing': '温情日常', 'shihuai': '释怀告别', 'chongfeng_jr': '重逢救赎',
        'jiashu': '家书来信', 'xiaorenwu': '小人物温度',
    },
    'fantasy': {
        'kaiju': '开局一个惨', 'dalian': '打脸爽文', 'baofu': '净身出户后暴富', 'zhenlong': '真龙出狱',
        'shenyi_xl': '神医下山', 'jianbao_xl': '鉴宝赌石', 'shenhao_xt': '神豪系统', 'chongsheng_zq': '重生赚钱',
    },
    'scifi': {
        'moshi_dunhuo': '末世囤货', 'anquanwu': '安全屋求生', 'gouzh': '苟住生存', 'haidao_qz': '海岛求生',
        'sangshi': '丧尸围城', 'yidong': '移动基地', 'feitu_dp': '废土短篇', 'chongsheng_ms': '末世重生',
    },
    'infinite': {
        'qiangqingxu': '强情绪快节奏', 'gao_dairu': '高代入感', 'suipian': '碎片阅读适配', 'gouzi_my': '钩子密集',
        'duanju': '短剧改编向', 'changju': '长剧孵化向', 'dy_rencheng': '第一人称代入',
    },
    'wuxia': {
        'danyuan_an': '单元案件', 'xilie_zhj': '系列主角', 'tanandanyuan': '探案单元', 'guaitan_xl': '怪谈系列',
        'anjian_chuan': '案件串烧', 'duanpian_ll': '短篇连缀成长篇',
    },
    'historical': {
        'zhenan': '真实案件改编', 'guaimai': '拐卖案', 'qianwen': '奇闻轶事', 'rensheng': '人生经历',
        'lieqi': '猎奇奇案', 'anjiadz_fz': '案件反转',
    },
    'xianxia': {
        'chongsheng_nx_dp': '重生逆袭', 'bubu': '步步为营', 'luanyi': '乱世成长', 'chaotang_qm': '朝堂权谋',
        'guifei': '贵妃青云直上', 'shunv_fs': '庶女翻身', 'jingcheng': '京城第一美人',
    },
}

# 短篇通用赛道兜底（题材未命中时使用）
SHORT_FALLBACK_STYLES = {
    'fanzhuan': '反转向（结局反转）', 'danyuanju': '单元剧', 'yingshi': '影视化（镜头语言）',
    'first_person': '第一人称', 'shacao_dp': '沙雕爽文', 'dianwen': '颠文', 'bailan': '摆烂流',
}

def _get_style_label(book_type, genre, style_key):
    """根据类型+题材查风格流派中文标签，找不到则原样返回 key。"""
    table = SHORT_GENRE_STYLES if book_type == 'short_story' else NOVEL_GENRE_STYLES
    genre_map = table.get(genre) if genre else None
    if not genre_map:
        if book_type == 'short_story':
            genre_map = SHORT_FALLBACK_STYLES
        else:
            genre_map = NOVEL_GENRE_STYLES.get('fantasy', {})
    return genre_map.get(style_key, style_key)

# 题材中英文映射（与前端 constants.ts GENRES 保持一致）
# 注入 prompt 时用中文标签，避免 AI 看到英文 key（如 fantasy）约束力弱
GENRE_LABELS = {
    'other': '其他', 'fantasy': '玄幻', 'xianxia': '仙侠', 'qihuan': '奇幻',
    'wuxia': '武侠', 'urban': '都市', 'urban_business': '都市职场', 'urban_fantasy': '都市异能',
    'history': '历史', 'military': '军事', 'game': '游戏', 'sports': '体育',
    'scifi': '科幻', 'mystery': '悬疑', 'infinite': '诸天无限', 'light_novel': '轻小说',
    'romance': '现代言情', 'ancient_romance': '古代言情', 'fantasy_romance': '幻想言情',
    'danmei': '纯爱', 'acg': '二次元',
}

# ===== 标准文风铁律已由 chat_collab_bp 三常量（GENERAL_CORE_RULES/WRITING_STYLE_RULES/DEAI_RULES）
#      作为唯一事实源统一注入，不再维护本文件平行副本。见 _build_ai_continue_context。

def _get_genre_label(book=None, bb=None):
    """获取题材中文标签（用于 prompt 注入）。优先 book.genre，回退 bb.genre。
    找不到映射时原样返回，确保不丢信息。"""
    genre = ''
    if book is not None:
        genre = getattr(book, 'genre', '') or ''
    if not genre and bb is not None:
        genre = getattr(bb, 'genre', '') or ''
    if not genre:
        return '通用'
    return GENRE_LABELS.get(genre, genre)

# 修炼体系小说题材关键词（玄幻/仙侠/都市异能/高武/历史脑洞等带修炼体系的小说）
# 命中任一关键词即视为修炼体系小说，用于在人物/大纲/剧情维度注入修炼天赋/境界变化区间/
# 年龄变化区间/时间线字段，并在防遗忘检查中加入境界/年龄/时间线一致性检查，防止正文跑偏
_CULTIVATION_GENRE_KEYWORDS = (
    # 大类码
    'fantasy', 'xianxia',
    # 玄幻子类
    'dongfang', 'yishi', 'gaowu', 'wangchao', 'honghuang', 'fanren', 'feichai',
    'qiangzhong', 'qiandao', 'shenchong', 'dihua', 'shenshu',
    # 仙侠子类
    'gudian', 'xiuzhen', 'huanxiang', 'shenhua', 'fengshen',
    'goudao', 'changsheng', 'jiazu', 'jianxiu', 'liandan', 'liangi',
    # 都市/都市奇幻 子类
    'yishu', 'dushi_xz', 'dushi_gw', 'guidze', 'dushi_nr', 'lingyi',
    # 中文词
    '玄幻', '仙侠', '修真', '修仙', '修炼', '异能', '高武', '武道', '洪荒',
    '封神', '凡人流', '废柴', '东方玄幻', '异世大陆', '高武世界', '都市修真',
    '都市高武', '异术超能', '规则怪谈', '灵异', '神话', '长生', '剑修',
    '炼丹', '炼器', '家族修仙', '苟道', '修真文明', '幻想修仙', '现代修真',
    '神话修真', '洪荒封神', '都市脑洞', '境界', '灵根', '丹田', '经脉', '功法',
)

def is_cultivation_novel(book=None, bb=None):
    """判断是否为修炼体系小说（玄幻/仙侠/都市异能/高武/历史脑洞等）。

    基于 genre/book_type/novel_styles 关键词判断，命中任一即视为修炼体系小说。
    """
    texts = []
    if book is not None:
        texts.append(getattr(book, 'genre', '') or '')
        texts.append(getattr(book, 'book_type', '') or '')
        texts.append(getattr(book, 'novel_styles', '') or '')
    if bb is not None:
        texts.append(getattr(bb, 'genre', '') or '')
        texts.append(getattr(bb, 'book_type', '') or '')
    blob = ' '.join(t for t in texts if t).lower()
    if not blob:
        return False
    for kw in _CULTIVATION_GENRE_KEYWORDS:
        if kw.lower() in blob:
            return True
    return False

def _cultivation_dimension_hint(dim, book=None, bb=None):
    """返回修炼体系小说在各维度的专属输出要求（非修炼体系小说返回空串）。

    在人物/五幕式大纲/剧情与情节节点维度注入修炼天赋/境界变化区间/年龄变化区间/时间线，
    防止正文写作时境界/年龄/时间线跑偏。
    """
    if not is_cultivation_novel(book, bb):
        return ''
    if dim == 'character_profiles':
        return ('\n\n【修炼体系小说·人物维度专属要求·必读】\n'
                '每个角色除原有字段外，必须额外输出以下两行（纯中文，每字段一行，与姓名/身份同级）：\n'
                '   修炼天赋：<灵根/体质/血脉/亲和度等天赋评级及特点，如：剑心通明，骨纹亲和度甲级，修炼速度常人三倍>\n'
                '   境界：<当前境界及潜力上限，如：骨纹三阶，潜力上限骨纹九阶；注明突破条件与代价>\n'
                '天赋决定修炼速度与上限，境界须与已建立的力量体系一致，不得越级或凭空突破。')
    if dim == 'plot_design':
        return ('\n\n【修炼体系小说·五幕式大纲专属要求·必读】\n'
                '每幕除原有字段外，必须标注以下四项，防止正文写作时境界/年龄/时间线跑偏：\n'
                '- 主要角色修炼天赋（须与人物档案一致）\n'
                '- 本幕境界变化区间（如：第1幕主角从凡人→练气三层）\n'
                '- 本幕年龄变化区间（如：16-18岁）\n'
                '- 本幕时间线节点（关键时间锚点，如：开篇第1年冬至→第2年春）\n'
                '各幕境界区间须全书连续递进，不得跳变；年龄与时间线须与境界进度匹配。')
    if dim == 'timeline':
        return ('\n\n【修炼体系小说·剧情与情节节点专属要求·必读】\n'
                '每卷 nodes 数组中每个情节节点除原有字段外，必须额外包含以下四个字段（字段名必须为中文，严禁使用英文字段名）：\n'
                '   "修炼进展": "本节点主要角色修炼进展，如：主角突破筑基，配角陨落"\n'
                '   "境界区间": "本节点境界变化区间，如：练气七层→筑基初期"\n'
                '   "年龄区间": "本节点年龄区间，如：18-19岁"\n'
                '   "时间锚点": "本节点时间线锚点，如：开篇第3年夏"\n'
                '各节点境界/年龄/时间线须全书连续递进，不得跳变或倒流。')
    return ''

def _sync_book_meta_to_bible(book, bb):
    """P0-3修复：把 book 的 total_volumes / novel_styles / genre / book_type 同步到 bible。
    在首次创建空 bible 或更新 book 时调用，确保各维度创作时能从 bible 读到权威元数据。

    同步策略（核心：Book 表是用户创建/编辑作品时的“权威数据源”，除用户在 Bible 侧明确改过且 Book 仍为默认值的情况外，以 Book 为准）：
    - book.tv != 默认(10) 且 bb.tv == 默认(10)   → 用户在 Book 改了，bb 还是默认值 → 同步（覆盖默认）
    - book.tv != bb.tv  且 双方都不等于默认值  → 两边都非默认，以 Book（用户作品基本信息页）为准 → 同步
    - bb.tv != 默认(10) 且 book.tv == 默认(10)   → 用户在 Bible 侧单独改了，Book 仍默认 → 保留 Bible，不同步
    """
    if not book or not bb:
        return
    DEFAULT_TV = 10
    try:
        bk_tv = getattr(book, 'total_volumes', None)
        bb_tv = getattr(bb, 'total_volumes', None)
        # 统一成 int 便于比较，None/0 视作默认
        try:
            bk_tv_i = int(bk_tv) if bk_tv else DEFAULT_TV
        except (ValueError, TypeError):
            bk_tv_i = DEFAULT_TV
        try:
            bb_tv_i = int(bb_tv) if bb_tv else DEFAULT_TV
        except (ValueError, TypeError):
            bb_tv_i = DEFAULT_TV
        need_sync = False
        # Case A：Book 非默认 + BB 默认 → 用户在 Book 设定页选了 25 卷，BB 刚创建还是默认 10 → 必须同步
        if bk_tv_i != DEFAULT_TV and bb_tv_i == DEFAULT_TV:
            need_sync = True
        # Case B：两边都非默认但数值不同 → Book 是权威（作品基本信息页），同步覆盖 Bible
        elif bk_tv_i != DEFAULT_TV and bb_tv_i != DEFAULT_TV and bk_tv_i != bb_tv_i:
            need_sync = True
        # Case C：两边都是默认 → 什么都不做
        # Case D：BB 非默认 + Book 默认 → 用户在 Bible 侧单独调过，保留 Bible
        if need_sync:
            bb.total_volumes = bk_tv_i
    except Exception:
        pass
    try:
        book_styles = getattr(book, 'novel_styles', None)
        if book_styles:
            bb_styles = getattr(bb, 'novel_styles', None)
            # 仅在 bible 风格为空或默认空数组时回填
            if not bb_styles or bb_styles in ('', '[]', 'null'):
                bb.novel_styles = book_styles
    except Exception:
        pass

def _get_total_volumes(bb, book=None):
    """获取总卷数：严格按用户设定（BB→Book→Bible文本正则提取），不再硬编码默认 10。
    优先级：
      1) BookBible.total_volumes（用户在智驾同步流中写入或作品侧设置同步）
      2) Book.total_volumes
      3) 从 Bible 的 concept / plot_design / timeline / key_rules / worldbuilding 等文本正则提取
         "全书N卷/按N卷/N卷规划..."（防止用户把卷数写进卡片正文但没写入字段）
      4) 最后兜底：若用户仍未设置，返回 0（上游 _core_params_iron_block 会感知 tv=0，
         整条"总卷数N卷/越界拦截N卷"铁律不再输出，避免把 LLM 误导到 10 卷）。
    卷数不设上限，仅校验下限 ≥1（当 tv≥1 时 clamp）。"""
    tv = None
    try:
        tv = getattr(bb, 'total_volumes', None)
    except Exception:
        tv = None
    if (not tv) and book is not None:
        try:
            tv = getattr(book, 'total_volumes', None)
        except Exception:
            tv = None

    # 第 3 层兜底：正则从已写入的 Bible 维度文本里找"全书N卷/按N卷..."
    if not tv:
        import re
        _RE = re.compile(
            r'(?:总卷数|全书|全本|整本书|一共|总共|合计|总计|计划|准备|打算|想|要|需要|改成|改为|设置为|设为|调整为|调成|调为|按|做成|写成|做|写|搞|设计成|规划成|规划|控制在|就|那就|那就按|就按|至少|最多|左右|大概|约|差不多)'
            r'\s*(\d{1,4})\s*卷')
        cand_texts = []
        if bb:
            for fld in ('concept', 'plot_design', 'timeline', 'key_rules', 'worldbuilding',
                        'style_guide', 'character_profiles', 'locations', 'foreshadowing',
                        'outline'):
                try:
                    v = getattr(bb, fld, None)
                except Exception:
                    v = None
                if isinstance(v, str) and v:
                    cand_texts.append(v)
        if book is not None:
            for fld in ('description', 'summary', 'outline'):
                try:
                    v = getattr(book, fld, None)
                except Exception:
                    v = None
                if isinstance(v, str) and v:
                    cand_texts.append(v)
        for txt in cand_texts:
            try:
                m = _RE.search(txt)
                if m:
                    cand = int(m.group(1))
                    if 1 <= cand <= 500:
                        tv = cand
                        break
            except Exception:
                continue

    # 最后：如果用户仍未显式设置卷数，返回 0。
    # 这样上游就不会输出"总卷数：10 卷 / 越界...10 卷"的错误铁律，
    # 也不会给 LLM 任何"默认十卷"的暗示。真正的创作完全由用户定义。
    try:
        tv_i = int(tv) if tv else 0
    except Exception:
        tv_i = 0
    if tv_i <= 0:
        return 0
    return max(1, tv_i)

def _get_chapters_per_volume(bb, book=None):
    """P3-10：按题材流派动态计算每卷章数，替代硬编码 50。
    - 短篇/短卷题材（短篇集、都市短篇）：30 章/卷
    - 长卷题材（玄幻/仙侠/奇幻/科幻/武侠）：60 章/卷
    - 默认（都市/言情/悬疑/通用）：50 章/卷
    总章数 = total_volumes × chapters_per_volume，确保百万字长篇时卷数合理。"""
    genre = ''
    try:
        genre = (getattr(book, 'genre', '') or getattr(bb, 'genre', '') or '').lower()
    except Exception:
        genre = ''
    bt = getattr(book, 'book_type', None) or getattr(bb, 'book_type', 'novel')
    if bt == 'short_story':
        return 30
    long_vol_genres = {'玄幻', '仙侠', '奇幻', '科幻', '武侠', '修真', '网游', '历史'}
    short_vol_genres = {'短篇', '短篇集', '散文', '随笔'}
    if any(g in genre for g in long_vol_genres):
        return 60
    if any(g in genre for g in short_vol_genres):
        return 30
    return 50

def _get_novel_styles_text(bb, book=None):
    """获取风格流派描述文本（最多3种叠加），按当前题材查对应风格表，用于注入创作上下文。"""
    styles_raw = None
    try:
        styles_raw = getattr(bb, 'novel_styles', None)
    except Exception:
        styles_raw = None
    if (not styles_raw) and book is not None:
        styles_raw = getattr(book, 'novel_styles', None)
    try:
        styles_list = json.loads(styles_raw or '[]') if styles_raw else []
    except Exception:
        styles_list = []
    if not styles_list:
        return ''
    bt = getattr(book, 'book_type', 'novel') if book is not None else getattr(bb, 'book_type', 'novel')
    genre = getattr(book, 'genre', '') if book is not None else getattr(bb, 'genre', '')
    labels = [_get_style_label(bt, genre, s) for s in styles_list[:3]]
    return '、'.join(labels)

def _build_core_params_block(bb, book):
    """构建“核心创作参数”注入块：题材+卷数+风格流派，作为所有维度创作与章节写作的核心依据。
    三大约束统一在此注入，确保用户选定项在下游真正生效。"""
    tv = _get_total_volumes(bb, book)
    genre_label = _get_genre_label(book, bb)
    styles_text = _get_novel_styles_text(bb, book)
    bt = getattr(book, 'book_type', 'novel') or 'novel'
    parts = ['【核心创作参数·全书依据·不可偏离】',
             f'题材：{genre_label}（人物设定、世界观、剧情走向、爽点类型须契合该题材的读者期待）',
             f'总卷数：{tv} 卷（全书所有分卷/五幕总纲/剧情大纲严格按此卷数规划，不得多不得少）']
    if bt == 'novel':
        parts.append(f'预计总字数：约 {tv*12} 万字（每卷约12万字，约50章/卷）')
    else:
        parts.append(f'短篇结构：{tv} 个单元/幕')
    if styles_text:
        parts.append(f'风格流派：{styles_text}（人物塑造、节奏、爽点设计、叙事手法须契合所选流派，这是硬约束）')
    return '\n'.join(parts)

# 章节正文“语言风格”表（行文文风，区别于题材流派）
# 与前端 constants.ts CHAPTER_LANG_STYLES 保持一致
CHAPTER_LANG_STYLES = {
    'general': ('通用', '行文规范流畅，叙述与对话比例均衡，节奏舒张有度，不刻意炫技也不寡淡；用词准确，符合现代汉语习惯。适合大多数题材的常规叙事。'),
    'baimiao': ('白描', '用最简练的笔墨勾勒人物与场景，不加渲染烘托；少用形容词，多用动词和名词；叙述客观克制，让事实自己说话。适合动作戏、硬汉派、克制情感。'),
    'jijian': ('极简', '句子短促有力，信息密度高，大量留白；砍掉一切冗余修饰与过渡；对话简洁，动作直接。适合快节奏、冷硬叙事与悬疑短篇。'),
    'youmo': ('幽默', '善用夸张、反语、双关与俏皮话制造笑点，插科打诨中藏锋芒；语言口语化，节奏跳跃；笑而不俗、讽而不戾。适合轻松日常、吐槽向、反套路喜剧。'),
    'shuangwen': ('爽文', '节奏明快，爽点密集，三章一冲突五章一反转；主角步步升级、打脸逆袭；情绪外放，多用反差对比烘托主角强大。适合玄幻都市升级流。'),
    'rexue': ('热血', '语言激昂奔放，多用短句排比与感叹，动作大开大合；情感外放，强调兄弟情、信念与战斗意志；场面燃点高，节奏层层推进。适合少年向、战斗竞技。'),
    'beiqing': ('悲情', '语调低沉绵长，多用环境烘托与意象铺陈情感；以克制写伤痛、以细节写离别；不滥情却字字戳心，留有余韵。适合虐心、悲剧、历史向与救赎类。'),
    'zhiyu': ('治愈', '笔调温柔舒缓，多写日常细节与微小温暖；语言清新柔和，少冲突多陪伴；以烟火气抚慰人心，情绪平稳上扬。适合日常向、慢生活与情感救赎类。'),
    'shijing': ('市井', '语言俚俗鲜活，多方言口语与江湖切口；人物三教九流，场景茶馆酒肆；叙述带烟火气与油滑感，对白占比较高。适合武侠江湖、都市底层与市井志怪。'),
    'gufeng': ('古风', '用词典雅，化用诗词典故与文言句式；句式工整讲究韵律，意境含蓄深远；适度文白相间，避免晦涩。适合仙侠、古言、宫斗与历史权谋题材。'),
    'guijue': ('诡谲', '氛围阴郁压抑，多用阴影、雾气、异响等意象制造不安；叙事扑朔迷离，留悬念与歧义；节奏沉滞中暗藏惊悚。适合悬疑、克苏鲁、志怪灵异与惊悚题材。'),
    'shiyi': ('诗意', '语言富于音乐性与意象美，重意境与情绪渲染；节奏舒缓，比喻空灵；近似散文诗，以景抒情、以物写心。适合文艺向、情感流与风景心境段落。'),
    'kouyu': ('口语化', '语言贴近日常口语，句式短、用词俗；可省主语、语序倒置、语气词丰富；叙述如说话，代入感强。适合都市生活、青春校园、第一人称与轻松吐槽向。'),
    'huangdan': ('荒诞', '以反逻辑与错位制造荒诞感，正经写荒唐、冷静写癫狂；语言可冷面幽默或黑色幽默；解构套路，预期违背生笑点。适合黑色幽默、讽刺、反套路与癫系创作。'),
}

def _build_chapter_lang_style_prompt(style_keys):
    """根据前端传入的语言风格 key 列表，拼装注入 system_prompt 的“本章语言风格”铁律约束。
    返回空串表示未选择（AI 按默认通用风格行文）。"""
    if not style_keys:
        return ''
    try:
        keys = style_keys if isinstance(style_keys, list) else json.loads(style_keys or '[]')
    except Exception:
        keys = []
    if not keys:
        return ''
    parts = []
    for k in keys[:3]:
        item = CHAPTER_LANG_STYLES.get(k)
        if item:
            parts.append(f"- {item[0]}：{item[1]}")
    if not parts:
        return ''
    return ('【本章语言风格·行文铁律】（用户为本章选定，必须严格遵循，不可偏离）\n'
            + '本章正文必须按以下风格基调行文（多种风格可融合但都要体现）：\n'
            + '\n'.join(parts)
            + '\n（此为硬约束：用词、句式、节奏、修辞须符合上述风格，违反即为不合格章节）')

def _build_master_upstream_ctx(dim, ctx):
    """按 DAG 依赖图只注入该维度的上游维度产物（不再无差别全注入）。
    自适应截断：世界观/人物档案给 2000 字预算（下游 timeline/dynamic_volumes/foreshadowing
    高度依赖完整设定，800字会丢失关键体系细节），其他维度 800 字。"""
    # 高预算维度：设定体量大且下游强依赖
    HIGH_BUDGET_DIMS = {'worldbuilding': 2000, 'character_profiles': 2000, 'key_rules': 1500}
    upstream_dims = MASTER_DIM_UPSTREAM.get(dim, [])
    parts = []
    for up_dim in upstream_dims:
        if up_dim not in ctx:
            continue
        up_val = ctx[up_dim]
        if up_val and up_val.strip():
            up_label = MASTER_DIM_MAP[up_dim]['label']
            budget = HIGH_BUDGET_DIMS.get(up_dim, 800)
            parts.append(f'【{up_label}（已确认上游）】\n{up_val[:budget]}')
    return '\n\n'.join(parts) if parts else '（暂无上游维度，自由发挥）'

def _build_master_storyline_ctx(book_id, bb):
    """构建"已写剧情"上下文：章节正文摘要 + 最近动态报告 + 伏笔图 + 关系图。
    让维度创作能感知已写章节的实际剧情，避免凭空设定与正文冲突。
    dynamic_volumes/character_profiles 等维度尤其需要，避免"幻觉式推测"。"""
    parts = []

    # 1. 最近章节正文摘要（取最近 6 章，每章前 300 字，避免 token 膨胀）
    try:
        recent_chs = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index.desc()).limit(6).all()
        if recent_chs:
            recent_chs = list(reversed(recent_chs))  # 恢复正序
            ch_lines = []
            for c in recent_chs:
                title = c.title or f'第{c.order_index}章'
                content = (c.content or '')[:300].replace('\n', ' ').strip()
                if content:
                    ch_lines.append(f'- 第{c.order_index}章《{title}》：{content}')
            if ch_lines:
                parts.append('【已写章节摘要】（最近6章实际剧情，设定须与之一致，不可矛盾）\n' + '\n'.join(ch_lines))
    except Exception:
        pass

    # 2. 最近动态报告摘要（3 份，每份 300 字）
    try:
        reports = _collect_relevant_reports(book_id, current_chapter_num=999999, window=999999, max_reports=3, per_report_limit=300)
        if reports:
            rp_lines = []
            for r in reports:
                rp_lines.append(f'- {r["title"]}（{r["chapter_start"]}-{r["chapter_end"]}章）：{r["content"][:300]}')
            parts.append('【动态报告摘要】（已写剧情的阶段性归纳）\n' + '\n'.join(rp_lines))
    except Exception:
        pass

    # 【P1-6修复】2.5 dynamic_volumes 当前状态摘要（已写章节实际境界/物品/势力，维度创作须承接）
    try:
        if bb.dynamic_volumes:
            dv_list = json.loads(bb.dynamic_volumes)
            if isinstance(dv_list, list) and dv_list:
                dv_lines = []
                for dv_entry in dv_list[:6]:  # 最多展示前6卷
                    if not isinstance(dv_entry, dict):
                        continue
                    vol_title = dv_entry.get('volume') or f'第{dv_entry.get("volume_index", "?")}卷'
                    data_obj = dv_entry.get('data') if isinstance(dv_entry.get('data'), dict) else dv_entry
                    # 优先取摘要字段，回退到 changelog 渲染
                    chars = (data_obj or {}).get('characters', '') if isinstance(data_obj, dict) else ''
                    events = (data_obj or {}).get('events', '') if isinstance(data_obj, dict) else ''
                    realms = (data_obj or {}).get('realms', '') if isinstance(data_obj, dict) else ''
                    # 若摘要为空，从 changelog 渲染最近5章
                    if (not chars and not events) and isinstance(data_obj, dict) and isinstance(data_obj.get('changelog'), list):
                        cl = data_obj['changelog']
                        cl_sorted = sorted(cl, key=lambda r: r.get('chapter', 0))[-5:]
                        cl_segs = []
                        for rec in cl_sorted:
                            ch = rec.get('chapter', 0)
                            fields = rec.get('fields', {})
                            for fname, items in fields.items():
                                if items:
                                    cl_segs.append(f'第{ch}章{fname}: {"；".join(items)[:60]}')
                        if cl_segs:
                            dv_lines.append(f'- {vol_title}：{"；".join(cl_segs)[:300]}')
                    else:
                        segs = []
                        if chars:
                            segs.append(f'人物:{chars[:100]}')
                        if events:
                            segs.append(f'事件:{events[:100]}')
                        if realms:
                            segs.append(f'境界:{realms[:80]}')
                        if segs:
                            dv_lines.append(f'- {vol_title}：{"；".join(segs)[:300]}')
                if dv_lines:
                    parts.append('【已写章节动态状态】（每卷人物境界/关键事件/势力实际状态，新维度创作须承接不可矛盾）\n' + '\n'.join(dv_lines))
    except Exception:
        pass

    # 3. 伏笔 DAG 摘要（待回收伏笔 top 8）
    try:
        if bb.foreshadowing_graph:
            graph = ForeshadowingGraph.from_dict(json.loads(bb.foreshadowing_graph))
            pending = graph.get_pending_nodes()
            if pending:
                fs_lines = []
                for node in pending[:8]:
                    desc = (node.content or node.title or '')[:80]
                    fs_lines.append(f'- {desc}')
                if fs_lines:
                    parts.append('【已埋伏笔清单】（已写剧情中埋设的伏笔，新设定须与之兼容）\n' + '\n'.join(fs_lines))
    except Exception:
        pass

    # 4. 人物关系图摘要（≤500 字）
    try:
        if bb.relation_graph and bb.relation_graph.strip():
            rg = bb.relation_graph[:500]
            parts.append('【人物关系图谱】（已确认的人物关系，新人物设定须兼容）\n' + rg)
    except Exception:
        pass

    # 5. 章级变更日志摘要（最近 12 章的关键 delta：境界/物品/伏笔/地点变化）
    # 章级粒度比动态报告更细，能补上"最近几章刚发生的境界提升/物品转移/伏笔回收"，
    # 避免维度创作与最新正文脱节（尤其 character_profiles/inventory/foreshadowing/dynamic_volumes）。
    try:
        log_list = json.loads(bb.chapter_changes_log) if bb.chapter_changes_log else []
        if isinstance(log_list, list) and log_list:
            recent_logs = log_list[-12:]  # 最近12章
            chg_lines = []
            for entry in recent_logs:
                if not isinstance(entry, dict):
                    continue
                ch_num = entry.get('chapter_num') or '?'
                chg = entry.get('changes') or {}
                if not isinstance(chg, dict):
                    continue
                segs = []
                # 角色境界/能力变化
                for c in (chg.get('CharacterStateChanges') or [])[:3]:
                    if isinstance(c, dict):
                        nm = c.get('CharacterId') or c.get('Name') or '?'
                        lvl = c.get('NewLevel') or ''
                        ke = c.get('KeyEvent') or ''
                        if lvl or ke:
                            segs.append(f'{nm}{"→"+lvl if lvl else ""}{":"+ke[:30] if ke else ""}')
                # 物品转移
                for it in (chg.get('ItemTransfers') or [])[:2]:
                    if isinstance(it, dict):
                        nm = it.get('ItemName') or it.get('ItemId') or '?'
                        fh = it.get('FromHolder') or ''
                        th = it.get('ToHolder') or ''
                        if th:
                            segs.append(f'{nm}:{fh or "?"}→{th}')
                # 伏笔 setup/payoff
                for fa in (chg.get('ForeshadowingActions') or [])[:2]:
                    if isinstance(fa, dict):
                        fid = fa.get('ForeshadowId') or '?'
                        act = fa.get('Action') or ''
                        if act:
                            segs.append(f'伏笔{fid}:{act}')
                # 地点状态
                for lc in (chg.get('LocationStateChanges') or [])[:2]:
                    if isinstance(lc, dict):
                        nm = lc.get('LocationId') or '?'
                        ns = lc.get('NewStatus') or ''
                        if ns:
                            segs.append(f'{nm}:{ns}')
                if segs:
                    chg_lines.append(f'- 第{ch_num}章：{"；".join(segs)}')
            if chg_lines:
                parts.append('【最近章级变更日志】（最近12章实际发生的境界/物品/伏笔/地点变化，新设定须承接，不可矛盾或回退）\n' + '\n'.join(chg_lines))
    except Exception:
        pass

    return '\n\n'.join(parts) if parts else ''

def _collect_anti_forget_alerts(bb, max_reports=3, max_alerts=12):
    """提取最近 N 份防遗忘检查报告的诊断要点（违规/待回收伏笔/叙事债务/改进建议）。
    用于在后续 AI 创作（章节写作/维度创作/一致性检查）中回注，让 AI 自动规避已诊断出的问题。
    返回字符串（已格式化，可直接拼入 prompt）；无报告时返回空字符串。"""
    if not bb or not bb.anti_forget_reports:
        return ''
    try:
        reports = json.loads(bb.anti_forget_reports) if bb.anti_forget_reports else []
    except (json.JSONDecodeError, ValueError, TypeError):
        return ''
    if not isinstance(reports, list) or not reports:
        return ''
    # 取最近 max_reports 份（按 checked_at 降序，兼容旧数据无 checked_at 的情况）
    def _ts(r):
        return r.get('checked_at', '') or ''
    recent = sorted([r for r in reports if isinstance(r, dict)], key=_ts, reverse=True)[:max_reports]
    if not recent:
        return ''

    sections = []
    for rec in recent:
        title = rec.get('title', '检查')
        report = rec.get('report') or {}
        # 不是 dict 的报告跳过细节提取
        if not isinstance(report, dict):
            report = {}
        parts = [f'■ {title}']
        summary = (report.get('summary') or rec.get('summary') or '').strip()
        if summary:
            parts.append(f'  摘要：{summary[:200]}')
        violations = report.get('violations') or []
        if isinstance(violations, list) and violations:
            v_lines = []
            for v in violations[:max_alerts]:
                if isinstance(v, dict):
                    msg = v.get('issue') or v.get('message') or v.get('desc') or v.get('description') or str(v)
                else:
                    msg = str(v)
                v_lines.append(f'  - {msg[:150]}')
            if v_lines:
                parts.append('  违规/不一致：\n' + '\n'.join(v_lines))
        pending = report.get('pending_foreshadowing') or []
        if isinstance(pending, list) and pending:
            p_lines = []
            for p in pending[:6]:
                if isinstance(p, dict):
                    msg = p.get('desc') or p.get('description') or p.get('title') or str(p)
                else:
                    msg = str(p)
                p_lines.append(f'  - {msg[:120]}')
            if p_lines:
                parts.append('  待回收伏笔：\n' + '\n'.join(p_lines))
        debts = report.get('narrative_debts') or report.get('debts') or []
        if isinstance(debts, list) and debts:
            d_lines = []
            for d in debts[:6]:
                if isinstance(d, dict):
                    msg = d.get('desc') or d.get('description') or d.get('title') or str(d)
                else:
                    msg = str(d)
                d_lines.append(f'  - {msg[:120]}')
            if d_lines:
                parts.append('  叙事债务：\n' + '\n'.join(d_lines))
        suggestions = report.get('suggestions') or report.get('improvements') or []
        if isinstance(suggestions, list) and suggestions:
            s_lines = []
            for s in suggestions[:6]:
                if isinstance(s, dict):
                    msg = s.get('suggestion') or s.get('desc') or s.get('description') or str(s)
                else:
                    msg = str(s)
                s_lines.append(f'  - {msg[:150]}')
            if s_lines:
                parts.append('  改进建议：\n' + '\n'.join(s_lines))
        if len(parts) > 1:  # 有实质内容才加入
            sections.append('\n'.join(parts))
    if not sections:
        return ''
    return '\n\n'.join(sections)

def _compute_dynamic_temperature(current_chapter_num, vol_chapter, vol_index, chapters_in_vol):
    """根据章节在卷中的位置动态计算 temperature。
    - 卷开篇（第1章）：0.8（需要更多创意建立场景）
    - 卷日常推进：0.7
    - 卷收尾（最后2章）：0.5（需要稳定收束）
    - 全书第一章：0.8"""
    if current_chapter_num == 1:
        return 0.8
    if vol_chapter and chapters_in_vol > 0:
        # 章节在卷内的相对位置
        pos_in_vol = current_chapter_num - vol_chapter.order_index + 1
        if pos_in_vol <= 1:
            return 0.8  # 卷开篇
        if pos_in_vol >= chapters_in_vol - 1:
            return 0.5  # 卷收尾
    return 0.7  # 日常推进

def _build_smart_instruction(instruction, last_chapter, current_chapter_num):
    """生成智能默认指令。若用户未提供 instruction，结合上一章章尾内容生成承接指令。
    【字数铁律】无论何种情况，user_prompt 末尾必须强制追加字数限定，确保 AI 在正文生成框明确看到 2400±100 约束。"""
    # 字数强制限定尾注（追加到任何指令后，确保字数约束出现在正文生成框）
    word_count_clause = f'请为“第{current_chapter_num}章”创作完整章节内容，要求上下文剧情连贯、对话自然、节奏紧凑、章末留悬念，字数2400±100字。'

    if instruction and instruction.strip():
        # 用户自定义指令：末尾追加字数限定，避免用户指令未提字数导致失控
        return f'{instruction.strip()}\n\n{word_count_clause}'
    if not last_chapter:
        return f'请继续写第 {current_chapter_num} 章（开篇章节，建立场景与基调）。\n\n{word_count_clause}'
    # 启发式提取上一章章尾钩子类型
    tail = (last_chapter.content or '')[-300:]
    hook_hint = '承接上一章章尾的悬念/钩子，自然展开新场景'
    if any(kw in tail for kw in ['？', '?', '究竟', '为何', '怎么']):
        hook_hint = '承接上一章末尾的疑问钩子，本章给出部分线索但不完全揭示'
    elif any(kw in tail for kw in ['危险', '危机', '攻击', '杀机', '威胁']):
        hook_hint = '承接上一章末尾的危机钩子，本章处理危机并展现主角应对'
    elif any(kw in tail for kw in ['发现', '出现', '现身', '传来']):
        hook_hint = '承接上一章末尾的新信息钩子，本章展开新信息的影响'
    return f'请继续写第 {current_chapter_num} 章。{hook_hint}。\n\n{word_count_clause}'

def _smart_truncate(text, budget):
    """语义截断：按段落/换行/句号切，保证每段完整，避免拦腰截断关键规则。
    优先级：段落(\n\n) > 换行(\n) > 中文句号(。！？) > 逗号(，；) > 硬截断。
    句号类分隔符截断后末尾补回标点，保证句子完整。"""
    if len(text) <= budget:
        return text
    # 句末标点：截断后需补回
    end_puncts = {'。', '！', '？', '；', '，'}
    for sep in ['\n\n', '\n', '。', '！', '？', '；', '，']:
        chunks = text.split(sep)
        if len(chunks) <= 1:
            continue
        result = []
        used = 0
        for i, ch in enumerate(chunks):
            # 加入此片段后的总长（含分隔符）
            piece_len = len(ch) + (len(sep) if i > 0 else 0)
            if used + piece_len > budget:
                break
            result.append(ch)
            used += piece_len
        if not result:
            continue
        out = sep.join(result)
        # 句末标点类：补回结尾标点，保证句子完整（如"第二条规则"→"第二条规则。"）
        if sep in end_puncts and out and not out.endswith(sep) and len(out) + len(sep) <= budget + len(sep):
            out = out + sep
        if out.strip():
            return out
    # 兜底硬截断
    return text[:budget]

def _apply_budget_management(sections_with_labels, total_budget=8000):
    """上下文窗口预算管理：按权重分配总预算给各段，避免单段超长挤掉关键信息。
    sections_with_labels: [(label, content, weight), ...]  weight 越大优先级越高
    返回拼接后的文本。截断采用语义截断（_smart_truncate），保证每段完整。"""
    if not sections_with_labels:
        return ''
    total_weight = sum(w for _, _, w in sections_with_labels)
    if total_weight <= 0:
        return '\n\n'.join(c for _, c, _ in sections_with_labels)
    parts = []
    for label, content, weight in sections_with_labels:
        if not content or not content.strip():
            continue
        budget = int(total_budget * (weight / total_weight))
        if len(content) > budget:
            content = _smart_truncate(content, budget)
        parts.append(content if not label else f'{label}\n{content}')
    return '\n\n'.join(parts)

def _generate_chapter_plan(book_id, bb, current_chapter_num, vol_chapter, vol_index,
                           memory_section, foreshadowing_section, skill_pack_ids,
                           api_key, base_url, model, max_tokens=600):
    """章节计划前置（chapter_plan Agent）：在写正文前生成 200 字以内的本章三段式计划。
    【P0弊端5修复】注入当前卷纲的 nodes 列表，让章纲对齐卷纲节点节奏。
    返回计划文本，失败时返回空串（不阻塞正文生成）。"""
    # 【三类无污染】chapter_plan 前置规划属于构思阶段：只注入构思类（master）技能包
    plan_skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['chapter_plan'], mode='single')
    vol_label = f'第{vol_index}卷“{vol_chapter.title}”' if vol_chapter else '当前卷'

    # ===== 【P0弊端5修复】提取当前卷纲 nodes，定位本章对应的节点 =====
    volume_nodes_section = ''
    current_node_hint = ''
    if vol_chapter:
        try:
            if bb.timeline and bb.timeline.strip().startswith('['):
                arr = json.loads(bb.timeline)
                if isinstance(arr, list):
                    for v in arr:
                        if not isinstance(v, dict):
                            continue
                        v_idx = v.get('volume_index') or _extract_volume_index(v.get('volume', v.get('volume_id', '')))
                        if str(v_idx) == str(vol_index):
                            nodes = v.get('nodes') or []
                            if isinstance(nodes, list) and nodes:
                                # 列出本卷所有节点
                                nodes_lines = []
                                for n in nodes:
                                    if not isinstance(n, dict):
                                        continue
                                    n_title = n.get('title', '')
                                    n_chapters = str(n.get('chapters', ''))
                                    n_type = n.get('type', 'M')
                                    n_summary = n.get('summary', '')
                                    nodes_lines.append(f'  · [{n_type}] {n_chapters}章 {n_title}：{n_summary}')
                                if nodes_lines:
                                    volume_nodes_section = f"""【本卷情节节点】（章纲必须对齐到本章所属节点）
{chr(10).join(nodes_lines)}"""
                                    # 定位本章对应的节点
                                    for n in nodes:
                                        if not isinstance(n, dict):
                                            continue
                                        ch_range = str(n.get('chapters', ''))
                                        nums = re.findall(r'\d+', ch_range)
                                        if len(nums) >= 2:
                                            start_n, end_n = int(nums[0]), int(nums[-1])
                                            if start_n <= current_chapter_num <= end_n:
                                                current_node_hint = f"""【本章所属节点】第{current_chapter_num}章对应节点“{n.get('title','')}”（{ch_range}章，类型{n.get('type','M')}）：{n.get('summary','')}
章纲必须围绕此节点的核心事件展开，不得偏离到其他节点。"""
                                                break
                                        elif len(nums) == 1:
                                            if int(nums[0]) == current_chapter_num:
                                                current_node_hint = f"""【本章所属节点】第{current_chapter_num}章对应节点“{n.get('title','')}”（类型{n.get('type','M')}）：{n.get('summary','')}"""
                                                break
                            break
                # 也提取上一卷卷尾钩子作为衔接提示
                if vol_index > 1:
                    for v in arr:
                        if not isinstance(v, dict):
                            continue
                        v_idx = v.get('volume_index') or _extract_volume_index(v.get('volume', v.get('volume_id', '')))
                        if str(v_idx) == str(vol_index - 1):
                            prev_hook = v.get('ending_hook') or v.get('ending') or v.get('climax') or ''
                            if prev_hook:
                                current_node_hint = f'【上一卷卷尾钩子】（本卷开头需承接）：{prev_hook}\n' + current_node_hint
                            break
        except (json.JSONDecodeError, ValueError, TypeError):
            pass

    plan_system = f"""你是小说章节策划师。为第 {current_chapter_num} 章生成一份简洁的章节计划（200字以内）。

当前所在卷：{vol_label}

{memory_section[:1500]}

{foreshadowing_section[:800] if foreshadowing_section else ''}

{volume_nodes_section}

{current_node_hint}

{plan_skill_note}

【输出格式】严格按以下五段式输出，每段不超过70字：
1. 本章核心冲突（只许1个，写清楚谁与谁、因什么事对立）：
2. 主角本章目标与主动行为（主角想要什么、主动做成的1-2件事——不许全程被动挨打）：
3. 冲突闭环方式（本章内解决，或显式挂起说明未完待续——禁止冲突凭空消失无下文）：
4. 信息增量点（本章让读者新知道的事：设定/关系/危机，至少1条）：
5. 章尾钩子设计：

【要求】
- 必须承接前文，不可矛盾
- 若有"待回收伏笔清单"，本章应考虑回收其中1条
- 【节点对齐铁律】若存在【本章所属节点】，章纲必须围绕该节点核心事件展开，不得偏离
- 【卷间衔接】若为卷首章节，必须承接上一卷卷尾钩子
- 【事件限额】本章关键事件≤3个且必须因果衔接（下一事件由上一事件引起，禁止并列堆砌无关事件）
- 只输出计划，不要解释"""

    try:
        resp = requests.post(f'{base_url}/chat/completions',
            headers=build_auth_headers(api_key),
            json={'model': model, 'messages': [{'role': 'system', 'content': plan_system},
                                                {'role': 'user', 'content': f'请为第 {current_chapter_num} 章生成计划'}],
                  'temperature': 0.5, 'max_tokens': max_tokens},
            timeout=60)
        result = resp.json()
        plan = result['choices'][0]['message']['content'].strip()
        if plan and len(plan) > 20:
            return plan
    except Exception:
        pass
    return ''

def _consistency_check(book_id, bb, draft_content, current_chapter_num,
                       api_key, base_url, model, max_tokens=1200, chapter_plan=''):
    """一致性检查 Agent：检查正文是否违反 key_rules/人设，并比对 chapter_plan 是否被执行。
    【P1扩展】新增 chapter_plan 比对，检测剧情偏离。
    【防遗忘回注】注入最近防遗忘检查诊断，使章后检查能对照历史诊断结果，避免重复犯错。
    【P1补全】检查覆盖面扩展到全维度：worldbuilding/timeline/foreshadowing/inventory/locations，
    按 bible 权威分级分层注入，避免地点矛盾/物资凭空获得/世界观违规等漏检。
    返回 (passed, issues_text)。失败时返回 (True, '')（不阻塞）。"""
    if not draft_content or len(draft_content) < 100:
        return True, ''
    key_rules = (bb.key_rules or '')[:800]
    chars = (bb.character_profiles or '')[:800]
    worldbuilding = (bb.worldbuilding or '')[:600]
    timeline = (bb.timeline or '')[:600]
    foreshadowing = (bb.foreshadowing or '')[:500]
    inventory = (bb.inventory or '')[:500]
    locations = (bb.locations or '')[:500]
    af_alerts = _collect_anti_forget_alerts(bb, max_reports=2, max_alerts=8)
    if not any([key_rules, chars, worldbuilding, timeline, foreshadowing, inventory, locations,
                chapter_plan, af_alerts]):
        return True, ''

    # 构建 chapter_plan 比对段（若有）
    plan_check_section = ''
    if chapter_plan and chapter_plan.strip():
        plan_check_section = f"""

【本章计划】（正文必须严格遵循此计划）
{chapter_plan[:600]}

【计划执行检查】除一致性外，额外检查正文是否实现了本章计划的核心冲突、关键场景、章尾钩子。若正文明显偏离计划（如跳过关键场景、改写核心冲突、丢失章尾钩子），记为问题。"""

    # 防遗忘诊断对照段（若有）
    af_check_section = ''
    if af_alerts:
        af_check_section = f"""

【历史防遗忘诊断】（最近检查发现的问题，检查正文是否重犯了类似错误）
{af_alerts}"""

    # 各维度分段（按权威分级分层注入）
    dim_sections = []
    if key_rules:
        dim_sections.append(f'【rules 层·核心规则/金手指】（绝不可违反）\n{key_rules}')
    if worldbuilding:
        dim_sections.append(f'【foundation 层·世界观设定】（世界法则不可违反）\n{worldbuilding}')
    if chars:
        dim_sections.append(f'【人物档案】（人设/关系/能力边界）\n{chars}')
    if timeline:
        dim_sections.append(f'【剧情/时间线】（事件顺序/时间推进须一致）\n{timeline}')
    if foreshadowing:
        dim_sections.append(f'【rules 层·伏笔状态】（已回收伏笔不可当未回收，待回收伏笔不可遗忘）\n{foreshadowing}')
    if inventory:
        dim_sections.append(f'【物资库】（物品持有者/数量须一致，不可凭空获得/消失）\n{inventory}')
    if locations:
        dim_sections.append(f'【地图/地点】（地点状态/归属须一致）\n{locations}')
    dim_block = '\n\n'.join(dim_sections)

    check_system = f"""你是小说一致性审查员。检查以下章节正文是否违反"项目宪法"，并比对本章计划是否被执行。
按权威分级（rules > foundation > 人物 > 剧情 > 伏笔 > 物资 > 地点）逐层检查。
只检查，不修改。返回 JSON：{{"passed": true/false, "issues": ["问题1", "问题2"]}}

【OOC 专项检测】除常规一致性外，专项检查角色是否 OOC（Out of Character）：
1. 说话语气是否符合人设档案（如沉稳角色突然话多、冷酷角色突然唠叨、粗豪角色突然文绉绉）。
2. 行为模式是否符合人设（如谨慎角色突然鲁莽、贪婪角色突然大方、孤傲角色突然讨好）。
3. 决策逻辑是否符合人设（如重情角色突然背叛、狡诈角色突然坦诚、隐忍角色突然暴走）。
4. 情绪反应是否符合人设（角色是否拥有矛盾心理/纠结/自我怀疑，还是情绪顺滑地接受了事件）。
任一不符记为 OOC 问题，归入 issues。

【项目宪法·分层】
{dim_block}{plan_check_section}{af_check_section}

【待检查正文】
{draft_content[:3000]}"""

    try:
        resp = requests.post(f'{base_url}/chat/completions',
            headers=build_auth_headers(api_key),
            json={'model': model, 'messages': [{'role': 'system', 'content': check_system},
                                                {'role': 'user', 'content': '请检查一致性与计划执行度，返回JSON'}],
                  'temperature': 0.2, 'max_tokens': max_tokens},
            timeout=60)
        result = resp.json()
        content = result['choices'][0]['message']['content'].strip()
        # 使用健壮解析函数替代贪婪正则
        parsed, parse_err = _extract_json_from_llm(content, expect='object')
        if parsed is None:
            app.logger.error(f"一致性检查JSON解析失败: {parse_err}, 原始内容前500字: {content[:500]}")
            return False, "一致性检查解析失败，请人工复核"
        return bool(parsed.get('passed', True)), '; '.join(parsed.get('issues', []))
    except Exception as e:
        app.logger.error(f"一致性检查执行异常: {e}")
        return False, "一致性检查执行异常，请人工复核"

def _build_continue_fingerprint_deps(book_id, bb, instruction, skill_pack_ids, target_chapter_num,
                                    prev_chapter_content, chapter_lang_styles, enable_structured_tags,
                                    skip_chapter_plan, book=None, recent_4ch_ids=None, cache_bible_version=None):
    """正文创作阶段 指纹 Key（零脏读）：只取 DB 行级稳定标识+本批次参数，不做重查询，轻量。
    任一依赖变动 → Key 自动变 → 自动 MISS 重算。"""
    if bb is None:
        bb_fields_t = (0, '', '', '', '', '', '', '', '', 0)
    else:
        # 10 个 bible 可变维度的「len+sha前16」混合指纹：变内容=指纹变（比取完整2000字轻量，又足够敏感）
        def _fp(s):
            if not s:
                return '0'
            return str(len(s)) + ':' + hashlib.sha1(str(s).encode('utf-8')).hexdigest()[:14]
        bb_fields_t = (
            int(getattr(bb, 'id', 0) or 0),
            _fp(getattr(bb, 'concept', '') or ''),
            _fp(getattr(bb, 'key_rules', '') or ''),
            _fp(getattr(bb, 'worldbuilding', '') or ''),
            _fp(getattr(bb, 'character_profiles', '') or ''),
            _fp(getattr(bb, 'plot_design', '') or ''),
            _fp(getattr(bb, 'timeline', '') or ''),
            _fp(getattr(bb, 'relation_graph', '') or ''),
            _fp(getattr(bb, 'inventory', '') or ''),
            _fp(getattr(bb, 'outline_hierarchy', '') or ''),
            _fp(getattr(bb, 'foreshadowing_graph', '') or ''),
            _fp(getattr(bb, 'generated_summary', '') or ''),
            _fp(getattr(bb, 'locations', '') or ''),
            _fp(getattr(bb, 'style_guide', '') or ''),
            int((getattr(bb, 'updated_at') or 0) and int(getattr(bb, 'updated_at', datetime(2020, 1, 1)).timestamp() * 1000 if hasattr(getattr(bb, 'updated_at', None), 'timestamp') else 0) or 0),
        )
    # book 行级稳定标识（genre/style 等用户可改字段也入指纹）
    if book is None:
        book_t = None
    else:
        book_t = (
            int(book.id or 0),
            str(getattr(book, 'genre', '') or ''),
            str(getattr(book, 'book_type', '') or ''),
            str(getattr(book, 'title', '') or ''),
            str(getattr(book, 'total_volumes', 0) or 0),
            str(getattr(book, 'chapters_per_volume', 0) or 0),
            str(getattr(book, 'master_skill_ids', '') or ''),
            str(getattr(book, 'style_skill_ids', '') or ''),
        )
    # recent_4ch_ids: 最近 4 章 id + word_count（正文每写完一章，下一章的 recent_4 滚动 → Key 自然变）
    recent_ch_t = tuple(
        (int(cid), int(wc or 0)) for (cid, wc) in (recent_4ch_ids or [])
    )
    # 批次参数
    params_t = (
        target_chapter_num,
        instruction and (str(len(instruction or '')) + ':' + hashlib.sha1((instruction or '').encode('utf-8')).hexdigest()[:12]),
        prev_chapter_content and (str(len(prev_chapter_content or '')) + ':' + hashlib.sha1((prev_chapter_content or '').encode('utf-8')).hexdigest()[:12]),
        tuple(sorted([str(x) for x in (skill_pack_ids or [])])),
        tuple(sorted([str(x) for x in (chapter_lang_styles or [])])),
        bool(enable_structured_tags),
        bool(skip_chapter_plan),
        cache_bible_version,  # 预留：外部全局版本号，暂时 None 不影响
    )
    return (book_t, bb_fields_t, recent_ch_t, params_t)


def _build_ai_continue_context(book_id, bb, instruction, skill_pack_ids, target_chapter_num=None, prev_chapter_content=None, chapter_lang_styles=None, enable_structured_tags=True, skip_chapter_plan=False, skip_cache=False, *, _bypass_cache=False):
    """构建章节写作完整上下文（ai_continue / stream / batch 共用），返回含
    system_prompt/user_prompt/temperature/max_tokens/chapter_plan/api 信息。
    新增缓存机制：先命中 PromptContextCache（零脏读指纹 Key）→ 命中直接 return，
    未命中才执行下面 0–12 步所有逐维度资料读取拼 prompt 逻辑，执行完回填缓存。
    _bypass_cache 私有参数：用于 cache 命中 lambda 内递归进入真实执行逻辑，不对外暴露。"""
    # ============== CACHE FAST PATH（仅外层调用进入；内层递归 _bypass_cache=True 跳过）==============
    if not _bypass_cache:
        # 预取少量稳定标识用于指纹（绝不在这里跑 bible/章节全量读，保持快路径 <1ms）
        _cache_book = Book.query.get(book_id)  # 行级查询，<1ms；下面真执行还会再取一次但 SQLAlchemy session 缓存掉
        _cache_allch_tip = Chapter.query.filter_by(book_id=book_id, is_volume=False).with_entities(
            Chapter.id, Chapter.word_count
        ).order_by(Chapter.order_index.desc()).limit(4).all()
        _cache_recent4 = [(c[0], c[1] or 0) for c in _cache_allch_tip] if _cache_allch_tip else []
        deps = _build_continue_fingerprint_deps(
            book_id, bb, instruction, skill_pack_ids, target_chapter_num,
            prev_chapter_content, chapter_lang_styles, enable_structured_tags,
            skip_chapter_plan, book=_cache_book, recent_4ch_ids=_cache_recent4,
        )

        def _compute():
            return _build_ai_continue_context(
                book_id, bb, instruction, skill_pack_ids, target_chapter_num,
                prev_chapter_content, chapter_lang_styles, enable_structured_tags,
                skip_chapter_plan, skip_cache=skip_cache, _bypass_cache=True,
            )

        payload, cache_info = PromptContextCache.get().get_or_compute(
            'continue_ctx', book_id, deps, _compute, ttl_sec=1800, skip_cache=skip_cache,
        )
        # 在返回 dict 上挂 cache_info（上游 wrapper 加响应头时用）
        if isinstance(payload, dict):
            payload['_cache_info'] = cache_info
        return payload

    # ============== 以下是真正装配逻辑（以前代码一字不动，仅函数签名扩展了 2 个 kwarg）==============
    book = Book.query.get(book_id)
    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.model if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')
    # 识别/检查类任务（章节计划、一致性检查）用识别模型，为空时回退主模型
    recognition_model = config.get_model_for_task('recognition') if config else model
    if not base_url.endswith('/v1'):
        base_url = base_url.rstrip('/') + '/v1'

    # ===== 0. 章号计算（#9：max(order_index)+1 兜底；前端传入优先，处理"已生成未保存"场景）=====
    all_chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    if target_chapter_num and isinstance(target_chapter_num, int) and target_chapter_num > 0:
        current_chapter_num = target_chapter_num
    else:
        if all_chapters:
            current_chapter_num = max(c.order_index for c in all_chapters) + 1
        else:
            current_chapter_num = 1
    last_chapter = all_chapters[-1] if all_chapters else None

    # ===== 1. 识别当前卷（#1+#3）=====
    vol_chapter, vol_index = _identify_current_volume(book_id, current_chapter_num)
    # 统计当前卷已有章节数（用于 temperature 计算）
    chapters_in_vol = 0
    if vol_chapter:
        next_vols = Chapter.query.filter(
            Chapter.book_id == book_id,
            Chapter.is_volume == True,
            Chapter.order_index > vol_chapter.order_index
        ).order_by(Chapter.order_index).all()
        upper_bound = next_vols[0].order_index if next_vols else 999999
        chapters_in_vol = Chapter.query.filter(
            Chapter.book_id == book_id,
            Chapter.is_volume == False,
            Chapter.order_index >= vol_chapter.order_index,
            Chapter.order_index < upper_bound
        ).count()

    # ===== 2. 分层 bible 上下文（#5：相关性筛选 + #14：预算管理）=====
    # 提取最近4章出场角色，用于 character_profiles 相关性筛选
    recent_for_chars = all_chapters[-4:] if all_chapters else []
    appearing_chars = _extract_appearing_characters(recent_for_chars, bb)
    filtered_bible = _filter_bible_by_relevance(bb, appearing_chars)

    bible_sections = []
    # 【大纲】五幕式总纲（权重最高，对齐整体走向）
    if filtered_bible.get('plot_design'):
        bible_sections.append(('【大纲·五幕式总纲】（参考当前进度对齐整体走向）', filtered_bible['plot_design'], 3))
    # 【剧情·卷纲规划】改造：从全量 timeline 注入改为前一卷+本卷+后一卷三卷注入
    # 避免上下文膨胀且聚焦当前创作位置，前卷回顾+本卷进行+后卷走向确保连贯
    adjacent_outlines = _get_adjacent_volumes_outline(book_id, vol_index)
    if adjacent_outlines:
        bible_sections.append((adjacent_outlines.split('\n')[0], '\n'.join(adjacent_outlines.split('\n')[1:]).strip(), 3))
    # 【人物及关系】合并 character_profiles + relation_graph
    if filtered_bible.get('character_profiles'):
        bible_sections.append(('【人物档案】（保持人设一致）', filtered_bible['character_profiles'], 3))
    if bb.relation_graph and bb.relation_graph.strip():
        rg_text = bb.relation_graph[:1000]
        bible_sections.append(('【人物关系图谱】（保持关系一致性）', rg_text, 2))
    # 【物资库】角色持有物品/境界约束
    if bb.inventory and bb.inventory.strip():
        inv_text = bb.inventory[:1000]
        bible_sections.append(('【物资库·境界库】（角色持有物品约束，不可凭空获得）', inv_text, 2))
    # 【地图】地点信息
    if bb.locations and bb.locations.strip():
        loc_text = bb.locations[:1000]
        bible_sections.append(('【地图·地点信息·全局】（若与本卷地点冲突，以本卷为准）', loc_text, 1))
    # 核心规则/金手指（绝不可违反）
    if filtered_bible.get('key_rules'):
        bible_sections.append(('【核心规则/金手指】（绝不可违反）', filtered_bible['key_rules'], 3))
    # 世界观设定
    if filtered_bible.get('worldbuilding'):
        bible_sections.append(('【世界观设定】', filtered_bible['worldbuilding'], 2))
    # 核心构思
    if filtered_bible.get('concept'):
        bible_sections.append(('【核心构思】', filtered_bible['concept'], 1))
    # 文风指南
    if filtered_bible.get('style_guide'):
        bible_sections.append(('【文风指南】', filtered_bible['style_guide'], 1))

    # 按卷注入维度数据（#1）——本卷人物/伏笔/地点/动态
    vol_dim_sections = []
    _inject_volume_dimensions(bb, vol_chapter, vol_index, vol_dim_sections)
    for s in vol_dim_sections:
        first_line = s.split('\n')[0]
        body = '\n'.join(s.split('\n')[1:]).strip()
        bible_sections.append((first_line, body, 2))

    # 卷目标对齐已由上方 _get_adjacent_volumes_outline 三卷注入覆盖（本卷部分更详细），
    # 此处不再单独注入当前卷卷纲，避免重复

    # 预算管理：bible设定总预算5000字（前4章正文独立注入，不挤占此预算）
    bible_context = _apply_budget_management(bible_sections, total_budget=5000) if bible_sections else (bb.generated_summary or '')[:2000]

    # ===== 3. 分层滚动记忆（P1重构：前4章完整正文 + 最近10份动态报告）=====
    # 最近10份动态报告（每份≤800字），覆盖更长剧情跨度
    relevant_reports = _collect_relevant_reports(book_id, current_chapter_num, window=100, max_reports=10, per_report_limit=800)
    # 百万字长线记忆：早期卷的按卷聚合摘要（补充 window=100 的盲区，避免前200章剧情遗忘）
    historical_digests = _collect_historical_volume_digest(book_id, current_chapter_num, max_volumes=4, per_volume_limit=400)
    # 前4章完整正文（不截断，保证剧情连贯性）
    recent_chapters_full = all_chapters[-4:] if all_chapters else []

    if relevant_reports:
        report_context = '\n\n'.join([f'【{r["title"]}（{r["chapter_start"]}-{r["chapter_end"]}章）】\n{r["content"]}' for r in relevant_reports])
    else:
        report_context = '（暂无动态报告）'

    # 历史卷摘要段（若有）：百万字时让 AI 看到早期卷的剧情概要
    historical_section = ''
    if historical_digests:
        hd_lines = [f'- {d["volume"]}（{d["chapter_range"]}）：{d["digest"]}' for d in historical_digests]
        historical_section = f"""
【早期卷剧情摘要】（百万字长线记忆，避免遗忘前几卷的关键剧情）：
{chr(10).join(hd_lines)}
"""

    if recent_chapters_full:
        chapters_text = '\n\n'.join([f'【第{c.order_index}章 {c.title or ""}】\n{c.content or ""}' for c in recent_chapters_full])
    else:
        chapters_text = '（开篇第一章，无前文）'

    # 修复上下文脱节：前端传入"上一章已生成未保存"内容时追加注入（截断2400字），避免剧情断档
    if prev_chapter_content and prev_chapter_content.strip():
        prev_trimmed = prev_chapter_content.strip()[:2400]
        prev_ch_num = current_chapter_num - 1
        prev_tag = '【上一章（已生成未保存，必须严格承接此章剧情）】'
        chapters_text = (chapters_text + '\n\n' if chapters_text and chapters_text != '（开篇第一章，无前文）' else '') + f'{prev_tag}\n【第{prev_ch_num}章】\n{prev_trimmed}'

    # 轻量RAG：基于当前章出场角色召回相关历史章节摘要，补充前4章窗口盲区
    # 让 AI 看到涉及角色在更早章节的经历，避免"角色历史行为遗忘"
    recalled_chapters = _recall_related_chapters(book_id, appearing_chars, current_chapter_num, max_chapters=6)
    recall_section = ''
    if recalled_chapters:
        rc_lines = [f'- 第{rc["chapter_num"]}章《{rc["title"]}》：{rc["summary"]}' for rc in recalled_chapters]
        recall_section = f"""
【相关历史章节召回】（基于本章出场角色智能召回，补充前4章窗口盲区，角色历史经历须与此一致）：
{chr(10).join(rc_lines)}
"""
    # 语义检索召回（embedding 优先，自动降级 TF-IDF）；字符召回不足 6 条时启用
    if len(recalled_chapters) < 6:
        try:
            from semantic_retriever import recall_semantic_chapters
            # 构造 query：出场角色 + 当前章 content_focus
            semantic_query = ' '.join(appearing_chars) if appearing_chars else ''
            if hasattr(bb, 'outline_hierarchy') and bb.outline_hierarchy:
                try:
                    hier = json.loads(bb.outline_hierarchy)
                    for ch_plan in hier.get('chapters', []):
                        if ch_plan.get('chapter_num') == current_chapter_num:
                            semantic_query += ' ' + (ch_plan.get('content_focus') or '')
                            break
                except Exception:
                    pass
            if semantic_query:
                def _provider():
                    # P2 升级：给 semantic retriever 传 content（整章 chunk 化做语义向量），
                    # 不传 content 会退化为旧 summary-300 字模式。
                    chs = Chapter.query.filter_by(book_id=book_id, is_volume=False).filter(
                        Chapter.order_index < current_chapter_num
                    ).order_by(Chapter.order_index.desc()).limit(300).all()
                    out = []
                    for c in chs:
                        out.append({
                            'chapter_num': c.order_index,
                            'title': c.title or '',
                            'summary': (getattr(c, 'summary', '') or ''),
                            'content': (getattr(c, 'content', '') or ''),
                        })
                    return out
                semantic_results = recall_semantic_chapters(
                    book_id, semantic_query, current_chapter_num,
                    exclude_recent=4, max_chapters=6 - len(recalled_chapters),
                    chapters_provider=_provider,
                )
                # 去重：排除已召回的章号
                existing_nums = {rc['chapter_num'] for rc in recalled_chapters}
                for sr in semantic_results:
                    if sr['chapter_num'] not in existing_nums:
                        recalled_chapters.append(sr)
                        existing_nums.add(sr['chapter_num'])
                # 重新渲染 recall_section
                if recalled_chapters:
                    rc_lines = [f'- 第{rc["chapter_num"]}章《{rc["title"]}》：{rc["summary"]}' for rc in recalled_chapters]
                    recall_section = f"""
【相关历史章节召回】（基于本章出场角色+语义检索召回，补充前4章窗口盲区，角色历史经历须与此一致）：
{chr(10).join(rc_lines)}
"""
        except Exception:
            pass  # 语义检索失败不影响主流程

    memory_section = f"""【前文动态报告】（最近10份动态文件摘要，防长线遗忘）：
{report_context}{historical_section}{recall_section}

【最近4章完整正文】（即时层，剧情衔接依据，必须严格保持连贯）：
{chapters_text}"""

    # P2-9：DynamicMemory 5文件版精简摘要注入（细粒度状态回流章节生成）
    # 5 文件版提供比 DynamicReport 更精细的角色生态/能力世界/伏笔追踪状态，
    # 每文件截取前 400 字，避免与动态报告内容重复堆叠
    dynamic_memory_section = ''
    try:
        dm = DynamicMemory.query.filter_by(book_id=book_id).first()
        if dm:
            dm_parts = []
            _dm_fields = [
                ('narrative_engine', '叙事引擎·节奏/张力状态'),
                ('foreshadowing_tracker', '伏笔追踪·状态快照'),
                ('character_ecosystem', '角色生态·关系/状态'),
                ('ability_world', '能力/世界·境界/势力'),
                ('health_dashboard', '健康度·债务预警'),
            ]
            for field_key, field_label in _dm_fields:
                raw = getattr(dm, field_key, '') or ''
                if raw and raw.strip() and raw.strip() != '{}':
                    # 尝试提取关键信息（若为 JSON 则压缩空白）
                    snippet = raw.strip()[:400]
                    dm_parts.append(f'- {field_label}：{snippet}')
            if dm_parts:
                dynamic_memory_section = f"""
【细粒度状态快照】（DynamicMemory 5文件版，比动态报告更精细的角色/能力/伏笔状态，须与此一致）：
{chr(10).join(dm_parts)}
"""
    except Exception:
        pass

    # ===== 4. 伏笔防遗忘（#2：按到期紧迫度排序，扩展到Top 25）=====
    # 百万字长线：top_n 提至 25，且紧迫度算法已对"无目标但沉淀已久"的伏笔加权，避免被挤出
    pending_fs = _sort_foreshadowings_by_urgency(bb, vol_chapter, current_chapter_num, top_n=25)
    foreshadowing_section = ''
    if pending_fs:
        fs_lines = []
        for urgency, planted, target, desc in pending_fs:
            target_hint = f'（计划回收于第{target}章）' if target else '（无明确回收点）'
            fs_lines.append(f'- {desc}{target_hint}')
        pending_text = '\n'.join(fs_lines)
        foreshadowing_section = f"""【待回收伏笔清单】（按到期紧迫度排序，本章节应考虑回收其中1-2条，避免遗忘；若无合适时机可暂缓，但不可永久遗忘）
{pending_text}"""

    # P0-2 增强：从伏笔 DAG 注入本章专属伏笔任务（应埋/应收），覆盖更精准
    if get_hooks_for_chapter and bb.foreshadowing_graph:
        try:
            graph = ForeshadowingGraph.from_dict(json.loads(bb.foreshadowing_graph))
            from foreshadowing_manager import build_hooks_prompt_section
            dag_hooks = build_hooks_prompt_section(graph, current_chapter_num)
            if dag_hooks:
                foreshadowing_section = (foreshadowing_section + '\n\n' + dag_hooks).strip()
        except Exception:
            pass  # DAG 解析失败退回原文本伏笔清单

    # ===== 4.5 防遗忘检查报告回注（让 AI 自动规避已诊断出的问题）=====
    af_alerts = _collect_anti_forget_alerts(bb, max_reports=3, max_alerts=12)
    af_section = ''
    if af_alerts:
        af_section = f"""【防遗忘检查诊断】（最近检查发现的问题，本次写作必须主动规避/修正）
{af_alerts}"""

    # P1-4 + P1-5：从四级大纲取本章戏剧位置，注入节拍模板
    beat_section = ''
    _beat_from_hierarchy = False
    if build_dramatic_position_prompt and build_beat_prompt and bb.outline_hierarchy:
        try:
            hierarchy = json.loads(bb.outline_hierarchy)
            # 戏剧位置上下文
            dp_prompt = build_dramatic_position_prompt(hierarchy, current_chapter_num)
            # 节拍模板
            from outline_hierarchy_builder import get_dramatic_context
            ctx_dp = get_dramatic_context(hierarchy, current_chapter_num)
            position = ctx_dp.get('dramatic_position', '') if ctx_dp else ''
            beats_prompt = build_beat_prompt(position, word_count=2500) if position else ''
            if dp_prompt or beats_prompt:
                beat_section = (dp_prompt + '\n\n' + beats_prompt).strip()
                _beat_from_hierarchy = True
        except Exception:
            pass  # 大纲/节拍加载失败退回无节拍模式

    # 【兜底】outline_hierarchy 缺失/查不到本章位置时按"章在卷内进度"推断位置注入节拍，
    # 避免正文零结构约束裸奔（"东拼西凑没起承转合"的根因之一：timeline 未生成时 beat 完全不注入）。
    # 推断规则：卷内前15%→起，15%-60%→承，60%-85%→转，末15%→合（卷结构也缺失时按章号1→起/其余→承）
    if not _beat_from_hierarchy and build_beat_prompt:
        _fallback_position = '起' if current_chapter_num <= 1 else '承'
        try:
            if vol_chapter is not None and current_chapter_num > 1:
                _cpv = _get_chapters_per_volume(bb, book) or 50
                _offset = max(0, current_chapter_num - int(vol_chapter.order_index or 1))
                _ratio = _offset / max(1, int(_cpv))
                if _ratio < 0.15:
                    _fallback_position = '起'
                elif _ratio < 0.6:
                    _fallback_position = '承'
                elif _ratio < 0.85:
                    _fallback_position = '转'
                else:
                    _fallback_position = '合'
            _fb_beats = build_beat_prompt(_fallback_position, word_count=2500)
            if _fb_beats:
                beat_section = (beat_section + '\n\n' if beat_section else '') + f"""【章内节拍模板】（四级大纲缺失，按第{current_chapter_num}章推断戏剧位置「{_fallback_position}」注入，必须严格遵循）
{_fb_beats}"""
        except Exception:
            pass

    # ===== 黄金开局公式（第1-3章：番茄开局留存窗口，违规=不合格）=====
    golden_section = ''
    if current_chapter_num <= 3:
        golden_section = """【黄金开局公式·第1-3章专属（开局留存窗口，违规=不合格章节）】
- 绝境代入（占全章15%-20%）：主角身处具体绝境（挨打/羞辱/濒死/被夺），压迫感落在身体细节上；穿越者必须写错愕与身体记忆错位的适应过程，禁止无缝进入战斗模式
- 金手指到账必须带代价感（疼/寿命/反噬/不可逆损失），且当章兑现一次代价
- 首次主动使用：主角当章必须主动用金手指做成一件小事（小翻身），不许只"知道有"而不"用"
- 格局钩子：章尾钩子指向更大的世界（更大的敌人/更大的利益/更大的秘密）
- 配速红线：本章世界观只给读者"当章用得上"的1条，其余留待后续展开；开局堆设定=劝退"""

    # ===== 5. 章节计划前置（#4：chapter_plan Agent）=====
    # 章节计划属"读数据→提炼→注入"的识别类任务，用识别模型（便宜快），正文生成本身仍用主模型
    # skip_chapter_plan=True 时跳过（批处理流式模式用，避免同步 LLM 调用阻塞导致心跳中断）
    chapter_plan = ''
    if not skip_chapter_plan:
        chapter_plan = _generate_chapter_plan(
            book_id, bb, current_chapter_num, vol_chapter, vol_index,
            memory_section, foreshadowing_section, skill_pack_ids,
            api_key, base_url, recognition_model, max_tokens=600
        )
    plan_section = f'【本章计划】（由 chapter_plan Agent 生成，请严格遵循）\n{chapter_plan}' if chapter_plan else ''

    # ===== 6. 技能包提示词 =====
    # 【三类无污染】正文生成阶段：只注入文风类（style）技能包，不注入构思/审查类
    # 【fix1】请求体 skill_pack_ids 为空时自动回退读 book.*_skill_ids，避免勾了文风没生效
    from skill_pack_runtime import resolve_active_style_ids
    _active_style_ids = resolve_active_style_ids(skill_pack_ids, book)
    # 【fix2】传 book_genre 让 genre_target 题材匹配生效；不匹配时 WARNING 日志
    skill_note = _get_skill_prompts_by_category(
        _active_style_ids, 'style',
        book_genre=getattr(book, 'genre', None) if book else None,
    )

    # ===== 7. 智能默认指令（#11）=====
    smart_instruction = _build_smart_instruction(instruction, last_chapter, current_chapter_num)

    # ===== 8. 组装 system_prompt（bible设定5000字预算 + 记忆独立段不挤占；卷数+流派+本章文风为核心依据）=====
    core_params_block = _build_core_params_block(bb, book)
    chapter_lang_style_block = _build_chapter_lang_style_prompt(chapter_lang_styles)
    system_prompt = f"""你是番茄小说金番作者级别的写手，正在协作写一本小说，当前准备写第 {current_chapter_num} 章。

{core_params_block}
{chapter_lang_style_block}

【设定权威分级·冲突仲裁规则】（P0-3）
当下方各层设定发生冲突时，按权威层级从高到低取信：
- direction 层（最高）：核心构思 concept —— 作者意图，不可被下层推翻
- foundation 层：大纲 plot_design、世界观 worldbuilding —— 全书骨架
- rules 层：核心规则 key_rules、伏笔 foreshadowing —— 不可违反的铁律
- runtime 层：动态文件 dynamic_volumes —— 当前卷的运行时状态
- memory 层（最低）：前4章正文 —— 即时剧情衔接依据
若卷纲与本卷动态文件冲突，以卷纲为准；若动态文件与前4章正文冲突，以动态文件为准。

【项目宪法 - 已确认设定】（必须严格遵守，不可矛盾。包含：大纲/剧情卷纲/人物及关系/物资库/地图/核心规则/世界观/伏笔等）
{bible_context[:5000]}

{memory_section}
{dynamic_memory_section}

{foreshadowing_section}

{af_section}

{beat_section}

{plan_section}

{golden_section}

【写作要求】
1. 严格遵循项目宪法中的设定，不可违反
2. 保持前后人物性格、关系、能力一致；严格衔接【最近4章完整正文】剧情走向
3. 延续现有文风和叙事节奏
4. 【字数绝对铁律】每章正文必须 2400 字 ±100（2300-2500 字区间，含标点）。低于2300字=扩展；超过2500字=删减。不可违反，优先级最高。
5. 主动回收"待回收伏笔清单"中的伏笔（若有）
6. 三明治结构：苦→甜→爽→钩子
7. 章尾必留钩子，七种类型不重复
8. 若存在【本章计划】，必须严格按计划展开
9. 【剧情连贯铁律】严格承接前4章结尾场景与悬念，不得凭空开新场景；人物位置/状态/对话一致
10.【剧情结构硬卡·起承转合（防东拼西凑，违规=不合格章节）】
  10.1 事件限额与因果链：本章关键事件≤3个，且必须因果衔接——下一事件由上一事件的后果直接引起，禁止无关事件并列堆砌、禁止事件被外力打断后就再无下文
  10.2 冲突闭环：本章核心冲突必须闭环——要么解决，要么显式挂起（主角明确知道没完、敌人明确留下威胁）；被新事件打断的冲突必须交代去向
  10.3 主角主动：主角主动决策≥2次（做选择/出手/布局），禁止全程被推着走；每个重大决策前必须有≥2句内心碎片（赌什么/为什么敢/怕什么），穿越/金手指/生死关必须写错愕或反应过程，禁止无缝进入战斗模式
  10.4 信息增量：每个大段落必须推进 信息/关系/危机 之一，禁止空转对话和事件播报；设定靠遭遇带出，禁止对白 lecture 灌世界观
  10.5 爽点公式：高潮爆发段必须写足连环反应——主角体感1句+至少2个视角的围观反差（反派嚣张→狼狈对照/强者动容/围观哗然）；打脸之后必须留≥2拍余震再收章，禁止爆发完直接跳收尾
  10.6 呼吸节奏：每2个冲突波之间留半段闲笔（环境/小动作/一句废话），全程紧绷无喘息=节奏灾难
10.7【节奏温度·15% 喘息段铁律（tension_score>90 自动判不合格，必须补 1–2 段闲笔再输出）】
    一章里至少要有 15% 的 Band1/Band2 喘息段（日常细节 / 景物锚 / 一句碎嘴对白 / 人物小动作细节），不能全程 Band4/5 紧绷。
    喘息段不用推进剧情，只干 3 件事选 1：
    • 环境锚：「豆浆摊的热气裹在他脸上，冷馒头的渣子卡在喉咙里，他咳了三声才咳出来。」
    • 人物小动作锚：「陈烨走路时总不自觉用拇指去抠表盖的划痕，抠到指甲发白——那划痕是妈生前用别针挑出来的。」（呼应背景，不用推进剧情）
    • 一句碎嘴对白锚：主角自己吐槽一句，比如「祖传护盾？祖传挨打才对。」
11.【字数自检】输出前用中文计数，不在 2300-2500 区间必须调整
12.【防遗忘规避】若有【防遗忘检查诊断】，主动规避已列违规并优先回收已诊断伏笔/叙事债务
13.【OOC专项】角色不得突然性格大变。

【文风铁律 · 最高权威（任一条违规即为不合格章节，必须重写）】
{skill_note}
（四项全局附加硬卡，叠加文风包红线共同生效——任何一条命中视为不合格章节，必须输出前自行调整修正，不得带违规提交）
- 硬卡1·段落完整（铁律A）：1-2句成段只为节奏；相邻≥3句同POV/场景/镜头必须合并≥2句，不得残切碎句
- 硬卡2·禁修正式否定（禁令0）：禁写「不是A而是B/不是修辞/不是地震」等修正式句式，一律改成直接陈述句
- 硬卡3·摄像机词限额：「看见/看着/听见/注意到/盯着/望向」一章合计≤3次，超额用动作/物象/感官代替
- 硬卡4·【整章量化双轨自检（输出前必须逐条检查，任一不达标即判定为不合格章节必须重写）】
  4.1 段内句号数硬上限 ≡ 每自然段 ≤ 2 个句号，一般以每自然段 1 个句号为主（=每自然段最多 2 句完整话）；整章中含 ≥ 3 个句号的段落数必须 = 0；**补充：凡段落字数（含标点）≤ 15 个字的短段，硬上限进一步降为 ≤ 1 个句号（短段里绝对不允许塞 2 句完整话）**；绝对不允许一段里堆 3–6 个小短句（漫画分镜脚本化是最浓AI味来源之一）
  4.2 段落结构 ≡ 短段主导：主力段落 10–50 字（占 80%），逗号串联 1–2 个动作即收一个句号；叙述长段（＞50字）仅用于信息密集/群像场（≤20%）；对白段不限（一句一段是常态）；70% 的段落 ≤ 70 字（手机端三行以内）；不允许把一个自然动作链切成 4+ 句独立段
  4.3 句均字数 ≡ 叙述句平均 12–18 字（含标点，短线为主），七成以上为逗号长句（逗号串 1–2 个动作单元收在一个句号）；连续 ≥ 3 句 ＜ 12 字句号短句 = 漫画分镜碎句，必须合并成逗号长句；单句 ＞ 55 字必须拆开
  4.4 段均句数 ≡ 整章合计 段落数 与 合计句数 的比值 ≤ 1.8；比值 > 2.0 立即判定 AI 碎段，必须合并改写
  4.5 最后一步必做：把写好的整章正文复制到「段数/句数/句号数计数器」里跑一下，确认同时满足 4.1–4.4，再输出
- 硬卡5·Humanizer 去 AI 痕迹铁律（与硬卡1–4 同级生效，任一命中即判定为不合格章节必须重写）
  5.1 废话 16 词黑名单必须 0 次命中（值得注意的是/总的来说/与此同时/不可否认/毫无疑问/显而易见/众所周知/由此可见/值得一提的是/综上所述/换言之/从某种意义上说/需要指出的是/不言而喻/毋庸置疑 —— 含"也就是说/换句话说/不难看出/可以说/需要说明的是"）
  5.2 禁止句式 6 种必须 0 次命中：「虽然……但是……」/「不仅……而且……」/「第一/首先……第二……第三……」列举/破折号「——」/连续「了」字堆砌同句≥3个/排比三连句
  5.3 被字句优先改主动：出现"被"字就问自己能不能把主语换成主动（"杯子被他捏碎了"→"他捏碎了杯子"），整章"被"字句 ≤ 1 处
  5.4 「X地」副词一律删或改动作：冷冷地/悄悄地/快速地/慢慢地/死死地 全部替换为动作/声音/触感描写
  5.5 分析报告术语必须 0 次：核心动机/信息边界/信息落差/利益最大化/底层逻辑/认知差/降维打击 —— 正文绝不能出现
  5.6 惊讶标记词密度必须合规：仿佛/忽然/竟/竟然/猛地/猛然/不禁/宛如/蓦然/骤然 每 2400 字合计 ≤ 1 次（一章最多 1 次）"""

    if enable_structured_tags and build_pre_write_check_prompt:
        system_prompt += build_pre_write_check_prompt(current_chapter_num, bb)

    # P1-6：在 system_prompt 末尾追加 CHANGES 输出模板（要求 LLM 输出结构化状态变更）
    if enable_structured_tags and build_changes_prompt_template:
        system_prompt += build_changes_prompt_template()

    # 标题自动生成：要求 LLM 在正文末尾输出 JSON 标题，供后端解析回填章节标题
    system_prompt += '\n\n【标题输出】在正文最末尾另起一行，输出一个 JSON 对象，格式：{"title": "标题文本"}。标题 8-16 字，概括本章核心冲突或转折，不使用"第X章"前缀，需贴合正文内容，避免剧透关键悬念。只输出 JSON，不要输出其他内容。'

    # ===== 【标准文风铁律】统一注入（三种创作模式共用此 context，确保输出风格一致）=====
    # 以 chat_collab_bp 三大常量（GENERAL_CORE_RULES 创作总则 / WRITING_STYLE_RULES 正文写作规范）
    # 为唯一事实源，通过 build_writing_rules 连同已选的文风技能包一并注入正文阶段；
    # 技能包提示词和章节语言风格仅做题材向微调，不得覆盖核心长短句比例与禁词清单。
    try:
        from blueprints.chat_collab_bp import build_writing_rules as _bp_build_writing_rules
        _writing_core = _bp_build_writing_rules(book=book, skill_pack_ids=skill_pack_ids, mode='agent')
    except Exception:
        _writing_core = ''
    if _writing_core:
        system_prompt += '\n\n【字数铁律】输出必须 2400±100 字（2300-2500字区间，含标点）。\n\n' + _writing_core

    # ===== 用户采纳的"系统学习与优化建议"补丁：统一追加到章节生成 system prompt 末尾 =====
    # （正文生成 / 流式 / 批量生成 都共享 _build_ai_continue_context，在此一处注入全局生效）
    if bb:
        try:
            from meta_optimizer import build_active_patch_text
            _pp = build_active_patch_text(bb)
            if _pp:
                system_prompt += '\n\n' + _pp
        except Exception:
            pass

    # ===== 9. 动态 temperature（#10）=====
    temperature = _compute_dynamic_temperature(current_chapter_num, vol_chapter, vol_index, chapters_in_vol)

    # ===== 10. 组装 user_prompt（第6142行返回引用，必须在此定义）=====
    user_prompt = instruction or f'请写第 {current_chapter_num} 章正文，严格遵循上方设定与计划。'

    # ===== 【fix4】激活文风包自证清单：返回时带 activated_skill_packs，前端可拉取自证
    from skill_pack_runtime import build_activated_skill_pack_manifest
    _activated_sp_names = build_activated_skill_pack_manifest(_active_style_ids)

    return {
        'system_prompt': system_prompt,
        'user_prompt': user_prompt,
        'temperature': temperature,
        # 【字数铁律】给足输出空间不物理截断；含 PRE_WRITE_CHECK 13行表 + CHANGES JSON 12字段，12000 token 确保完整。
        # 【输出上限适配】模型上限低于 12000（如 deepseek-chat 8192）时按已知/已学习上限钳制，
        # 防下游直连 requests.post 的调用（去AI味/字数修正/审校等）400
        'max_tokens': min(12000, get_output_limit(base_url, model) or 12000),
        'chapter_plan': chapter_plan,
        'current_chapter_num': current_chapter_num,
        'vol_chapter': vol_chapter,
        'vol_index': vol_index,
        'api_key': api_key,
        'base_url': base_url,
        'model': model,
        'recognition_model': recognition_model,  # 识别/检查类任务用
        'activated_skill_packs': _activated_sp_names,  # 自证用：本次正文生成实际注入的文风包
    }

@app.route('/api/books/<book_id>/ai-continue', methods=['POST'])
@login_required
def ai_continue(book_id):
    """正文滚动创作（多 Agent 协同版，14项优化）：
    分层 bible 注入+按卷维度+卷纲对齐 / 滚动记忆防遗忘 / 伏笔紧迫度Top25 /
    章节计划前置 / 动态temperature正文生成 / 去AI味审校（容错+deai_status）/
    一致性检查（key_rules/人设）。【优化2】TaskGraph 统一编排：写前构建任务图，
    伏笔任务等安全任务自动执行，各 LLM 阶段完成后 mark_stage 回写，结束持久化轨迹。"""
    book = Book.query.get(book_id)
    if not book: return jsonify({'error': 'Not found'}), 404

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    instruction = request.json.get('instruction', '')
    skill_pack_ids = request.json.get('skill_pack_ids', [])
    enable_consistency_check = request.json.get('enable_consistency_check', False)  # 默认关闭（OOC检测已合并到system_prompt，不再单独调LLM）
    # 修复上下文脱节：接收前端传入的待写章号 + 上一章未保存内容
    target_chapter_num = request.json.get('target_chapter_num')
    prev_chapter_content = request.json.get('prev_chapter_content')
    # 章节正文语言风格（行文文风，最多3个叠加）
    chapter_lang_styles = request.json.get('chapter_lang_styles', [])
    # Prompt 上下文缓存旁路：True=强制重新从各维度资料拼prompt(改了未存库的小设定时用)
    skip_prompt_cache = bool(request.json.get('skip_prompt_cache', False) or request.json.get('skip_cache', False))
    # S1：critical 门禁被 block 后，前端二次确认可传 ignore_gates=True 强制落库
    ignore_gates = bool(request.json.get('ignore_gates', False))

    try:
        ctx = _build_ai_continue_context(book_id, bb, instruction, skill_pack_ids, target_chapter_num, prev_chapter_content, chapter_lang_styles, skip_cache=skip_prompt_cache)
        system_prompt = ctx['system_prompt']
        user_prompt = ctx['user_prompt']
        temperature = ctx['temperature']
        max_tokens = ctx['max_tokens']
        api_key = ctx['api_key']
        base_url = ctx['base_url']
        model = ctx['model']

        # ===== 【优化2】TaskGraph 统一编排：写前任务图+安全任务；异常降级旧直连模式 =====
        wp_graph = wp_runner = None
        try:
            from smart_planner import build_writing_pipeline
            wp_graph, wp_runner, _mission = build_writing_pipeline(
                book_id, bb, ctx['current_chapter_num'],
                skill_pack_ids=skill_pack_ids, enable_consistency_check=enable_consistency_check)
            if wp_runner: wp_runner.mark_stage(wp_graph, 't2_ctx', 'done')
        except Exception:
            wp_graph = wp_runner = None

        # ===== P2：Context Manifest（上下文来源 + hash + token 预算，事后溯源）=====
        context_manifest_data = None
        if ContextOrchestrator is not None:
            try:
                # S2：动态预算（ceiling 12k；模型上下文 ≥ 32k 时放宽到 16k）
                ctx_win = ContextOrchestrator._heuristic_context_window(model)
                auto_ceiling = 16000 if ctx_win >= 32768 else 12000
                # 粗估 system_prompt tokens
                _sys_cn = len([c for c in system_prompt if '\u4e00' <= c <= '\u9fff'])
                _sys_other = len(system_prompt) - _sys_cn
                _sys_tok = int(_sys_cn / 1.5 + _sys_other / 4)
                dynamic_budget = ContextOrchestrator.dynamic_budget(
                    max_gen_tokens=max_tokens,
                    model_name=model,
                    system_prompt_estimate=_sys_tok,
                    ceiling=auto_ceiling,
                )
                _orch = ContextOrchestrator(token_budget=dynamic_budget)
                _sources = {
                    'key_rules': getattr(bb, 'key_rules', '') or '',
                    'worldbuilding': getattr(bb, 'worldbuilding', '') or '',
                    'character_profiles': getattr(bb, 'character_profiles', '') or '',
                    'plot_design': getattr(bb, 'plot_design', '') or '',
                    'concept': getattr(bb, 'concept', '') or book.synopsis or '',
                }
                if prev_chapter_content:
                    _sources['prev_chapter'] = prev_chapter_content[:2000]
                _manifest = _orch.prepare(
                    sources=_sources,
                    chapter_num=ctx['current_chapter_num'],
                    book_id=book_id)
                context_manifest_data = _manifest.to_dict()
            except Exception:
                pass  # manifest 失败不阻断章节生成

        # ===== 正文生成（经 LLM Gateway 统一入口：错误分类 + 智能重试 + 空内容检测）=====
        draft_content, llm_error = _llm_chat(
            [{'role': 'system', 'content': system_prompt},
             {'role': 'user', 'content': user_prompt}],
            api_key=api_key, base_url=base_url, model=model,
            temperature=temperature, max_tokens=max_tokens, timeout=180)
        if not draft_content or not draft_content.strip():
            # 超时类错误返回 504，其他返回 502
            status = 504 if (llm_error and '超时' in llm_error) else 502
            return jsonify({'error': llm_error or 'LLM 返回空内容，请重试'}), status

        if wp_runner:  # 优化2：t2_plan/t3_draft 阶段完成
            wp_runner.mark_stage(wp_graph, 't2_plan', 'done' if ctx.get('chapter_plan') else 'skipped')
            wp_runner.mark_stage(wp_graph, 't3_draft', 'done', {'chars': len(draft_content)})

        # 【修复】额外校验：剥离内部标签后仍有正文，防止 LLM 只输出标签导致门禁误报"正文为空"
        _pre_check_body = _extract_chapter_body(draft_content)
        _pre_check_body = re.sub(r'\{[^{}]*"title"\s*:\s*"[^"]*"[^{}]*\}', '', _pre_check_body).strip()
        if not _pre_check_body:
            return jsonify({'error': 'LLM 仅输出结构标签（pre_write_check/chapter_changes），无正文内容，请重试'}), 502

        # ===== 【字数铁律】初稿字数校验 + AI 重写修正（非物理截断，保证章节完整）=====
        # 【修复】改用公共函数 _ensure_word_count，与流式/连续/连续流式模式统一
        draft_content, review_notes_prefix = _ensure_word_count(
            draft_content, api_key, base_url, model, max_tokens, ctx['current_chapter_num'])
        if wp_runner: wp_runner.mark_stage(wp_graph, 't4_wc', 'done', {'chars': len(draft_content)})

        # ===== 去 AI 味审校 Agent（#6：容错+可观测；2026-08-23 默认启用）=====
        # 内置统一去AI规则（chat_collab_bp 的 GENERAL_CORE_RULES + DEAI_RULES）常驻生效，
        # 不再依赖技能包勾选；审查类(review)技能包作为增强叠加（build_review_rules 内部合并）。
        polished_content = draft_content
        review_notes = review_notes_prefix
        deai_status = 'skipped'  # skipped / rules_ok / rules_missing / success / failed
        deai_rules_block = ''
        try:
            deai_rules_block, _deai_build_status = _build_deai_rules_block(skill_pack_ids, book)
            deai_status = 'rules_ok' if _deai_build_status == 'ok' else 'rules_missing'
        except Exception as _deai_e:
            # helper 也兜不住（连 DEAI_ONLY_RULES import 都失败），显式标 missing 而非静默
            app.logger.error(f'ai_continue 去AI规则构建失败: {_deai_e}')
            print(f'[去AI] ai_continue 去AI规则构建失败: {_deai_e}', file=sys.stderr)
            deai_status = 'rules_missing'
        if deai_rules_block:
            deai_system = ("你是番茄去AI味审查员。对以下刚写好的章节正文做去AI味审校，按规则修改后只输出修改后的正文。\n\n"
                           + deai_rules_block
                           + "\n\n【硬性约束】修改后字数仍须 2400±100（2300-2500区间，含标点），保留原章节的剧情走向和钩子，只改文风不改剧情。")
            try:
                deai_resp = requests.post(f'{base_url}/chat/completions',
                    headers=build_auth_headers(api_key),
                    json={'model': model,
                          'messages': [{'role':'system','content':deai_system},
                                       {'role':'user','content':f'请审校以下章节正文：\n\n{draft_content}'}],
                          'temperature': 0.5, 'max_tokens': max_tokens},
                    timeout=180)
                deai_result = deai_resp.json()
                polished = deai_result['choices'][0]['message']['content'].strip()
                # 【字数铁律】审校后字数校验：必须落在 2300-2500 区间
                polished_len = _count_cn_chars(polished)
                if polished and 2300 <= polished_len <= 2500:
                    polished_content = polished
                    review_notes = (review_notes_prefix + ' 已自动去AI味审校(' + str(polished_len) + '字)').strip()
                    deai_status = 'success'
                elif polished and polished_len > 500:
                    # 字数不达标但有内容，标记为失败但仍返回初稿
                    review_notes = (review_notes_prefix + f' 去AI味审校返回字数异常({polished_len}字)，已回滚使用初稿').strip()
                    deai_status = 'failed'
                else:
                    review_notes = (review_notes_prefix + ' 去AI味审校返回为空，已回滚使用初稿').strip()
                    deai_status = 'failed'
            except Exception as e:
                review_notes = (review_notes_prefix + f' 去AI味审校异常：{str(e)[:100]}，已回滚使用初稿').strip()
                deai_status = 'failed'

        # ===== 一致性检查 Agent（#13：独立 Agent，P1扩展：含 chapter_plan 比对）=====
        # 一致性检查属识别/检查类任务，用识别模型
        consistency_passed = True
        consistency_issues = ''
        if enable_consistency_check:
            consistency_passed, consistency_issues = _consistency_check(
                book_id, bb, polished_content, ctx['current_chapter_num'],
                api_key, base_url, ctx.get('recognition_model', model), max_tokens=800,
                chapter_plan=ctx.get('chapter_plan', '')
            )
        if wp_runner:  # 优化2：t5_deai/t6_cchk 阶段完成
            wp_runner.mark_stage(wp_graph, 't5_deai', deai_status if deai_status != 'skipped' else 'skipped',
                                 {'status': deai_status})
            wp_runner.mark_stage(wp_graph, 't6_cchk', 'done' if enable_consistency_check else 'skipped',
                                 {'passed': consistency_passed})

        # ===== P1-6 + P1-7：CHANGES 解析 + delta 回写（非流式版）=====
        changes_applied = None
        if extract_changes and apply_chapter_changes:
            try:
                body, changes = extract_changes(polished_content)
                if changes:
                    # 剥离 CHANGES 标签后的正文（落库用纯正文）
                    if body and len(body) > 200:
                        polished_content = body
                    ch_obj = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index.desc()).first()
                    ch_id = ch_obj.id if ch_obj else ''
                    changes_applied = apply_chapter_changes(
                        bb, ch_id, ctx['current_chapter_num'],
                        ctx.get('vol_index', 0), changes,
                    )
                    if changes_applied.get('applied'):
                        try:
                            db.session.commit()
                        except Exception:
                            db.session.rollback()
            except Exception:
                pass  # 回写失败不阻断
        if wp_runner: wp_runner.mark_stage(wp_graph, 't7_changes', 'done' if changes_applied else 'skipped')

        # ===== P0-1：确定性后写校验（非流式版，零 LLM 成本）=====
        # bible_ctx 覆盖全维度（死亡复活/境界回退/文风漂移检测）；漂移检测并入主校验分支
        post_validate = None
        if validate_chapter_with_bible:
            try:
                body_for_check = _extract_chapter_body(polished_content)
                bible_ctx = {
                    'character_profiles': bb.character_profiles or '',
                    'chapter_changes_log': bb.chapter_changes_log or '',
                    'key_rules': bb.key_rules or '',
                    'worldbuilding': bb.worldbuilding or '',
                    'inventory': bb.inventory or '',
                    'locations': bb.locations or '',
                    'foreshadowing': bb.foreshadowing or '',
                } if bb else None
                validation = validate_chapter_with_bible(body_for_check, bible_ctx)
                # 【P0修复】追加文风漂移检测：计算前5章基准，检测当前章是否漂移
                try:
                    style_baseline = _compute_style_baseline(book_id, ctx['current_chapter_num'])
                    if validate_chapter_with_drift and style_baseline:
                        drift_validation = validate_chapter_with_drift(body_for_check, style_baseline)
                        # 合并漂移检测的 issues 到主 validation
                        for issue in drift_validation.issues:
                            validation.issues.append(issue)
                            validation.score = max(0, validation.score - (5 if issue.severity == 'warning' else 20))
                        # 把漂移统计写入 stats
                        if drift_validation.stats.get('style_drift'):
                            validation.stats['style_drift'] = drift_validation.stats['style_drift']
                except Exception:
                    pass  # 漂移检测失败不影响主校验
                # 无论有没有问题，都把统计项(stats)抛给前端（段均句数/句均字数等健康指标）
                post_validate = validation.to_dict()
                if not validation.issues:
                    post_validate['_hardcard4_ok'] = True
            except Exception:
                pass
        elif validate_chapter:
            try:
                body_for_check = _extract_chapter_body(polished_content)
                validation = validate_chapter(body_for_check)
                # 无论有没有问题，都把统计项(stats)抛给前端监控（段均句数/句均字数/短碎句占比/段内≥3句号数）
                post_validate = validation.to_dict()
                if not validation.issues:
                    # 没命中问题时也要保留 stats 和 passed/score 摘要，前端可展示"健康值"
                    post_validate['_hardcard4_ok'] = True
            except Exception:
                pass
        if wp_runner: wp_runner.mark_stage(wp_graph, 't8_pval', 'done' if post_validate else 'skipped',
                                           {'issues': len((post_validate or {}).get('issues', []))})

        # ===== P2-10：落地门禁（3道，章节落库前拦截；延迟到 _extract_chapter_body 后传纯正文）=====
        gate_result = None

        # ===== P1-5：审计-修订闭环（校验→修订→再校验，最多2轮；编排已抽取至 chapter_review_cycle）=====
        review_cycle_result = None
        if (run_review_cycle and validate_chapter_with_bible and bb
                and post_validate and post_validate.get('issues')):
            try:
                polished_content, review_cycle_result = run_review_cycle_with_bible(
                    polished_content, bb, post_validate, book_id, ctx['current_chapter_num'],
                    api_key, base_url, model, _extract_chapter_body)
            except Exception:
                db.session.rollback()  # 闭环失败不阻断章节生成
        if wp_runner: wp_runner.mark_stage(wp_graph, 't9_cycle', 'done' if review_cycle_result else 'skipped',
                                           {'passed': (review_cycle_result or {}).get('passed')})

        # ===== 审校评分制：聚合 4 套检测结果计算 0-100 分（零 LLM 成本）=====
        try:
            polished_wc = len(re.sub(r'\s', '', polished_content or ''))
            chapter_score = _calc_chapter_score(
                post_validate, consistency_passed, consistency_issues,
                gate_result, polished_wc, ctx.get('chapter_plan', ''))
        except Exception:
            chapter_score = None

        # ===== 标题自动生成：解析【标题】标签，剥离正文中的标签行 =====
        suggested_title = _extract_chapter_title(draft_content)
        # 统一标题格式：第X章 标题文本（与连续创作模式一致，混用模式时格式统一）
        formatted_title = _format_chapter_title(ctx['current_chapter_num'], suggested_title)
        # ★ 统一清洗：无论 changes 是否解析成功，都剥离所有内部标签（pre_write_check / chapter_changes / 标题JSON / 【标题】行）
        # 确保返回给前端的 content 是纯净正文，避免内部产物泄露给用户
        polished_content = _extract_chapter_body(polished_content)
        polished_content = re.sub(r'\{[^{}]*"title"\s*:\s*"[^"]*"[^{}]*\}', '', polished_content).rstrip()

        # 【修复】落地门禁在 _extract_chapter_body 之后调用，传入纯正文
        # 防御性检查：去AI味/字数修正后若 polished_content 变空，跳过门禁调用避免误报"正文为空"
        if run_all_gates and polished_content and polished_content.strip():
            try:
                gate_result = run_all_gates(polished_content, bb, ctx['current_chapter_num'])
                if wp_runner: wp_runner.mark_stage(wp_graph, 't10_gates', 'done',
                                                   {'passed': gate_result.get('passed'), 'blocked': gate_result.get('blocked')})
                # S1：critical 默认 block；用户二次确认传 ignore_gates 则放行
                if gate_result.get('blocked') and not ignore_gates:
                    if wp_runner: wp_runner.persist_plan_log(wp_graph, 'generate_chapter', {'outcome': 'gate_blocked'})
                    _cache_info = ctx.get('_cache_info') if isinstance(ctx, dict) else None
                    _block_body = {
                        'gate_blocked': True,
                        'ignore_gates_required': True,
                        'content': polished_content,
                        'draft': draft_content if deai_status == 'success' else None,
                        'review_notes': review_notes,
                        'deai_status': deai_status,
                        'chapter_plan': ctx.get('chapter_plan', ''),
                        'current_chapter_num': ctx['current_chapter_num'],
                        'vol_index': ctx.get('vol_index', 0),
                        'vol_title': ctx.get('vol_chapter').title if ctx.get('vol_chapter') else '',
                        'temperature': temperature,
                        'gate_result': gate_result,
                        'block_reason': '落地门禁检测到 critical 问题（如正文过短/为空等）。默认拦截自动落库，'
                                        '请在前端确认「忽略门禁强制保存」后再次提交。',
                        'prompt_cache_info': _cache_info,
                        'cache_stats': _cache_stats_snapshot(),
                    }
                    return _response_with_cache(jsonify(_block_body), _cache_info), 428
                if not gate_result.get('passed'):
                    pass  # 仅 warning：不阻断
            except Exception:
                pass

        # 优化2：写作流水线轨迹持久化（plan_log_json，最近20条）+ 前端可观测
        pipeline_plan = None
        if wp_runner:
            wp_runner.mark_stage(wp_graph, 't11_post', 'declared', {'note': '落库后由 _after_chapter_persisted 执行'})
            wp_runner.persist_plan_log(wp_graph, 'generate_chapter', {'outcome': 'ok'})
            pipeline_plan = wp_graph.to_dict()

        return jsonify({
            'content': polished_content,
            'draft': draft_content if deai_status == 'success' else None,
            'review_notes': review_notes,
            'deai_status': deai_status,  # #6：新增可观测字段
            'chapter_plan': ctx.get('chapter_plan', ''),  # #4：返回计划供前端展示
            'current_chapter_num': ctx['current_chapter_num'],
            'vol_index': ctx.get('vol_index', 0),
            'vol_title': ctx['vol_chapter'].title if ctx.get('vol_chapter') else '',
            'temperature': temperature,  # #10：返回实际使用的 temperature
            'consistency_passed': consistency_passed,  # #13：一致性检查结果
            'consistency_issues': consistency_issues,
            # P0-1 + P1-6/7 新增字段
            'post_validate': post_validate,  # 后写校验报告（AI痕迹检测）
            'changes_applied': changes_applied,  # 章级变更回写摘要
            # P2-10 新增
            'gate_result': gate_result,  # 落地门禁结果
            # P1-5 新增
            'review_cycle': review_cycle_result,  # 审计-修订闭环结果
            # 审校评分制新增
            'chapter_score': chapter_score,  # 0-100 评分 + 等级 + 5维明细 + auto_revise
            # 标题自动生成新增
            'suggested_title': suggested_title,  # 纯标题文本（如"小镇少年"）
            'formatted_title': formatted_title,  # 统一格式标题（如"第1章 小镇少年"），前端直接使用
            # P2 新增：上下文溯源 manifest（记录本次生成注入了哪些 bible 片段 + hash + token 预算）
            'context_manifest': context_manifest_data,
            'pipeline_plan': pipeline_plan,  # 优化2：写作流水线任务图（12阶段执行轨迹）
            'prompt_cache_info': ctx.get('_cache_info', {'hit': False, 'tokens_saved': 0}) if isinstance(ctx, dict) else None,  # PromptCache命中信息
            'cache_stats': _cache_stats_snapshot(),  # 全局cache统计（hits/misses/tokens_saved）
        })
        # 加响应头：X-Prompt-Cache HIT/MISS + X-Tokens-Saved
        _cache_info = ctx.get('_cache_info') if isinstance(ctx, dict) else None
        return _response_with_cache(jsonify(result_body), _cache_info)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/books/<book_id>/ai-spot-fix', methods=['POST'])
@login_required
def ai_spot_fix(book_id):
    """Spot-Fix 修订端点（P2-9）。
    按校验问题路由：local 类只修补问题段落（省 token），structural 类建议整章重写。
    前端"一键修订"按钮调用。"""
    if not route_revision:
        return jsonify({'error': '修订模块未加载'}), 500
    data = request.get_json() or {}
    content = data.get('content', '')
    validation = data.get('post_validate', {})
    mode = data.get('mode', 'auto')  # auto / spot_fix / rewrite
    if not content:
        return jsonify({'error': '缺少正文内容'}), 400

    # 路由修订策略
    routing = route_revision(content, validation, mode=mode)

    if routing['strategy'] == 'none':
        return jsonify({
            'strategy': 'none',
            'message': '未检测到需要修订的问题',
            'content': content,
        })

    if routing['strategy'] == 'rewrite':
        # structural 问题，建议整章重写（前端可调用 ai-continue）
        return jsonify({
            'strategy': 'rewrite',
            'message': '检测到结构性问题，建议整章重写',
            'structural_issues': routing['structural_issues'],
            'content': content,
        })

    # spot_fix 策略：只送问题段落给 LLM
    patches = routing['patches']
    if not patches:
        return jsonify({'strategy': 'none', 'message': '无法定位问题段落', 'content': content})

    # 构建 Spot-Fix prompt
    sys_prompt, user_prompt = build_spot_fix_prompt(content, patches)
    token_saving = estimate_token_saving(content, patches)

    # 调用 LLM 修订
    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('creation') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')
    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    try:
        resp = requests.post(f'{base_url}/chat/completions',
            headers=build_auth_headers(api_key),
            json={'model': model,
                  'messages': [{'role': 'system', 'content': sys_prompt},
                               {'role': 'user', 'content': user_prompt}],
                  'temperature': 0.3,
                  'max_tokens': 2000},
            timeout=60)
        result = resp.json()
        llm_output = result['choices'][0]['message']['content'].strip()

        # patch 回原文
        revised_content = apply_spot_fix_patches(content, patches, llm_output)

        # 对修订后的内容再跑一次后写校验
        post_validate = None
        if validate_chapter:
            try:
                body_for_check = _extract_chapter_body(revised_content)
                validation2 = validate_chapter(body_for_check)
                if validation2.issues:
                    post_validate = validation2.to_dict()
            except Exception:
                pass

        return jsonify({
            'strategy': 'spot_fix',
            'content': revised_content,
            'patches_count': len(patches),
            'token_saving': token_saving,
            'post_validate': post_validate,  # 修订后的校验报告
        })
    except Exception as e:
        return jsonify({'error': f'修订失败：{str(e)[:200]}'}), 500

@app.route('/api/books/<book_id>/ai-continue/stream', methods=['POST'])
@login_required
def ai_continue_stream(book_id):
    """正文滚动创作流式版（#8：SSE 推送初稿）。
    前端可通过 EventSource 接收 chunks，审校与一致性检查在流结束后由前端单独触发或忽略。
    返回 text/event-stream，每个 chunk 形如 data: {"content": "..."}\n\n"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Not found'}), 404

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    instruction = request.json.get('instruction', '')
    skill_pack_ids = request.json.get('skill_pack_ids', [])
    # 修复上下文脱节：接收前端传入的待写章号 + 上一章未保存内容
    target_chapter_num = request.json.get('target_chapter_num')
    prev_chapter_content = request.json.get('prev_chapter_content')
    # 章节正文语言风格（行文文风，最多3个叠加）
    chapter_lang_styles = request.json.get('chapter_lang_styles', [])
    # Prompt 上下文缓存旁路
    skip_prompt_cache = bool(request.json.get('skip_prompt_cache', False) or request.json.get('skip_cache', False))

    ctx = _build_ai_continue_context(book_id, bb, instruction, skill_pack_ids, target_chapter_num, prev_chapter_content, chapter_lang_styles, enable_structured_tags=True, skip_cache=skip_prompt_cache)
    api_key = ctx['api_key']
    base_url = ctx['base_url']
    model = ctx['model']

    # 【优化2】TaskGraph 统一编排（流式版）：写前任务图，LLM 阶段随流推进 mark_stage
    wp_graph = wp_runner = None
    try:
        from smart_planner import build_writing_pipeline
        wp_graph, wp_runner, _mission = build_writing_pipeline(book_id, bb, ctx['current_chapter_num'],
                                                               skill_pack_ids=skill_pack_ids)
    except Exception:
        wp_graph = wp_runner = None

    def generate():
        try:
            # 先推送元信息（计划、章号、卷信息、temperature）
            meta = {
                'meta': True,
                'chapter_plan': ctx.get('chapter_plan', ''),
                'current_chapter_num': ctx['current_chapter_num'],
                'vol_index': ctx.get('vol_index', 0),
                'vol_title': ctx['vol_chapter'].title if ctx.get('vol_chapter') else '',
                'temperature': ctx['temperature'],
            }
            yield f'data: {json.dumps(meta, ensure_ascii=False)}\n\n'
            if wp_runner:  # 优化2：上下文构建+章节计划阶段完成
                wp_runner.mark_stage(wp_graph, 't2_ctx', 'done')
                wp_runner.mark_stage(wp_graph, 't2_plan', 'done' if ctx.get('chapter_plan') else 'skipped')

            # 标准文风与字数铁律已在 _build_ai_continue_context 内统一注入（三种模式一致）
            system_prompt = ctx['system_prompt']

            # 流式生成正文初稿，收集完整内容用于后写校验（P0-1）
            # 【空回复修复1】非 200 显式报错（旧实现遍历错误页无 data 帧→流静默结束→前端"空回复"）
            full_content_parts = []
            resp = requests.post(f'{base_url}/chat/completions',
                headers=build_auth_headers(api_key),
                json={'model': model,
                      'messages': [{'role': 'system', 'content': system_prompt},
                                   {'role': 'user', 'content': ctx['user_prompt']}],
                      'temperature': ctx['temperature'],
                      'max_tokens': ctx['max_tokens'],
                      'stream': True},
                stream=True, timeout=180)
            if resp.status_code != 200:
                _err_txt = ''
                try:
                    _err_txt = resp.text[:200]
                except Exception:
                    pass
                _err_msg = f"LLM 流式调用失败（HTTP {resp.status_code}）：{_err_txt or '服务返回错误，请检查 API Key/额度'}"
                yield f'data: {json.dumps({"error": _err_msg}, ensure_ascii=False)}\n\n'
                return
            for line in resp.iter_lines():
                if line:
                    line = line.decode('utf-8')
                    if line.startswith('data: '):
                        chunk = line[6:]
                        if chunk == '[DONE]':
                            yield 'data: [DONE]\n\n'
                            break
                        # 收集内容用于后写校验 + 统一格式转发前端
                        try:
                            chunk_data = json.loads(chunk)
                            # 兼容多种供应商格式：标准OpenAI / 简化delta / 直接content
                            delta = ''
                            try:
                                choices = chunk_data.get('choices') or []
                                if choices:
                                    delta = (choices[0].get('delta') or {}).get('content', '') or \
                                            (choices[0].get('message') or {}).get('content', '')
                            except Exception:
                                pass
                            if not delta:
                                delta = chunk_data.get('content') or chunk_data.get('text') or ''
                            if delta:
                                full_content_parts.append(delta)
                                # 统一为标准 OpenAI 格式转发，确保前端能正确解析
                                yield f'data: {json.dumps({"choices": [{"delta": {"content": delta}}]}, ensure_ascii=False)}\n\n'
                        except Exception:
                            pass  # 非 JSON 行跳过

            # 【空回复修复2】流结束但零内容帧 → 显式 error 帧替代静默结束（前端不再"空回复"）
            if not full_content_parts:
                yield f'data: {json.dumps({"error": f"LLM 返回空内容（model={model}，可能原因：max_tokens 过小/模型拒答/网关异常），请重试"}, ensure_ascii=False)}\n\n'
                return
            if wp_runner: wp_runner.mark_stage(wp_graph, 't3_draft', 'done',
                                               {'chars': sum(len(p) for p in full_content_parts)})

            # ===== P1-6 + P1-7：CHANGES 解析 + delta 回写（流结束后）=====
            # 解析 LLM 输出的 12 类变更声明，delta patch 到 dynamic_volumes/foreshadowing_graph
            if extract_changes and apply_chapter_changes and full_content_parts:
                try:
                    full_content_for_changes = ''.join(full_content_parts)
                    body, changes = extract_changes(full_content_for_changes)
                    if changes:
                        bb_obj = BookBible.query.filter_by(book_id=book_id).first()
                        if bb_obj:
                            ch_obj = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index.desc()).first()
                            ch_id = ch_obj.id if ch_obj else ''
                            apply_summary = apply_chapter_changes(
                                bb_obj, ch_id, ctx['current_chapter_num'],
                                ctx.get('vol_index', 0), changes,
                            )
                            if apply_summary.get('applied'):
                                try:
                                    db.session.commit()
                                except Exception:
                                    db.session.rollback()
                                # 推送回写摘要给前端
                                yield f'data: {json.dumps({"changes_applied": apply_summary}, ensure_ascii=False)}\n\n'
                except Exception as ce:
                    # 回写失败不阻断章节生成
                    pass
            if wp_runner: wp_runner.mark_stage(wp_graph, 't7_changes', 'done')

            # 【修复】流式模式补字数修正：初稿字数不在 2300-2500 时 AI 重写（推心跳告知前端）
            if full_content_parts:
                try:
                    raw_content = ''.join(full_content_parts)
                    body_for_wc = _extract_chapter_body(raw_content)
                    draft_wc = _count_cn_chars(body_for_wc)
                    if draft_wc < 2300 or draft_wc > 2500:
                        yield f'data: {json.dumps({"type": "heartbeat", "message": f"正在修正字数（初稿{draft_wc}字）..."}, ensure_ascii=False)}\n\n'
                        corrected, wc_note = _ensure_word_count(
                            body_for_wc, api_key, base_url, model, ctx['max_tokens'], ctx['current_chapter_num'])
                        if corrected and corrected.strip() and _count_cn_chars(corrected) != draft_wc:
                            # 推送修正后的完整正文给前端（替换初稿）
                            yield f'data: {json.dumps({"type": "word_count_corrected", "content": corrected, "note": wc_note}, ensure_ascii=False)}\n\n'
                            # 更新 full_content_parts 供后续去AI味/校验使用
                            full_content_parts = [corrected]
                except Exception:
                    pass  # 字数修正失败不阻断流式生成
            if wp_runner: wp_runner.mark_stage(wp_graph, 't4_wc', 'done', {'chars': _count_cn_chars(''.join(full_content_parts)) if full_content_parts else 0})

            # ===== P0-1：确定性后写校验（流结束后，零 LLM 成本，报告推前端供一键修订）=====
            # 注入 bb 上下文：死亡复活/境界回退/角色名错写硬伤检测
            _validator = validate_chapter_with_bible or validate_chapter
            if _validator and full_content_parts:
                try:
                    full_content = ''.join(full_content_parts)
                    # 剥离可能的 PRE_WRITE_CHECK 和 CHANGES 标签（P1-6 产物），只校验正文
                    body_for_check = _extract_chapter_body(full_content)
                    bible_ctx = None
                    if validate_chapter_with_bible and bb:
                        bible_ctx = {
                            'character_profiles': bb.character_profiles or '',
                            'chapter_changes_log': bb.chapter_changes_log or '',
                            'key_rules': bb.key_rules or '',
                            'worldbuilding': bb.worldbuilding or '',
                            'inventory': bb.inventory or '',
                            'locations': bb.locations or '',
                            'foreshadowing': bb.foreshadowing or '',
                        }
                    validation = _validator(body_for_check, bible_ctx) if bible_ctx else _validator(body_for_check)
                    if validation.issues:
                        yield f'data: {json.dumps({"post_validate": validation.to_dict()}, ensure_ascii=False)}\n\n'
                except Exception as ve:
                    # 校验失败不阻断章节生成
                    pass
            if wp_runner: wp_runner.mark_stage(wp_graph, 't8_pval', 'done' if full_content_parts else 'skipped')

            # ===== 【优化3】流式模式补齐落地门禁（与非流式对齐）=====
            # 流式不落库（前端确认后另行保存），gates 只做告警不阻断：blocked 时推 gate_blocked 帧
            if run_all_gates and full_content_parts:
                try:
                    _body_for_gates = _extract_chapter_body(''.join(full_content_parts))
                    if _body_for_gates and _body_for_gates.strip():
                        _gate = run_all_gates(_body_for_gates, bb, ctx['current_chapter_num'])
                        if wp_runner: wp_runner.mark_stage(wp_graph, 't10_gates', 'done',
                                                           {'passed': _gate.get('passed'), 'blocked': _gate.get('blocked')})
                        _gate_blocked = bool(_gate.get('blocked', False))
                        yield f'data: {json.dumps({"gate_result": _gate, "gate_blocked": _gate_blocked}, ensure_ascii=False)}\n\n'
                except Exception:
                    pass

            # ===== 标题自动生成：解析【标题】标签并推送（统一格式"第X章 标题"，前端优先 formatted_title）=====
            if full_content_parts:
                try:
                    full_content_for_title = ''.join(full_content_parts)
                    suggested_title = _extract_chapter_title(full_content_for_title)
                    formatted_title = _format_chapter_title(ctx['current_chapter_num'], suggested_title)
                    yield f'data: {json.dumps({"suggested_title": suggested_title, "formatted_title": formatted_title}, ensure_ascii=False)}\n\n'
                except Exception:
                    pass

            # 【P1-4修复】流式模式补去AI味 Agent：仅当有 review 类技能包时触发（与多Agent模式一致）
            try:
                review_skill_ids = _resolve_skill_ids_by_category(book, 'review') if book else []
                if review_skill_ids and full_content_parts:
                    full_content_for_deai = ''.join(full_content_parts)
                    body_for_deai = _extract_chapter_body(full_content_for_deai)
                    if body_for_deai and len(body_for_deai) > 200:
                        deai_skill_note = _get_skill_prompts_by_category(review_skill_ids, 'review', ['deai', 'consistency_check'])
                        if deai_skill_note:
                            deai_sys = ('你是去AI味审校专家。按以下技能包要求，对章节正文做最小改动修订，'
                                       '只调整AI痕迹和文风问题，不改变剧情、人物、设定。\n\n'
                                       f'{deai_skill_note}\n\n'
                                       '【输出】直接输出修订后的完整正文（含标题行），不要任何解释。')
                            deai_user = f'原文：\n{body_for_deai[:6000]}'
                            yield f'data: {json.dumps({"type": "deai_start"}, ensure_ascii=False)}\n\n'
                            deai_content, deai_err = _call_llm(
                                [{'role': 'system', 'content': deai_sys}, {'role': 'user', 'content': deai_user}],
                                max_tokens=0, temperature=0.5
                            )
                            if not deai_err and deai_content and deai_content.strip():
                                # 剥离标题行，提取正文
                                deai_body = _extract_chapter_body(deai_content)
                                if deai_body and len(deai_body) > 200:
                                    yield f'data: {json.dumps({"type": "deai_result", "content": deai_body}, ensure_ascii=False)}\n\n'
            except Exception:
                pass  # 去AI味失败不阻断流式生成
            # 优化2：t5_deai 收尾 + 任务图轨迹持久化 + 推送 pipeline_plan 终帧（前端可观测）
            if wp_runner:
                _deai_state = 'done' if (locals().get('deai_body') and len(locals().get('deai_body') or '') > 200) else 'skipped'
                wp_runner.mark_stage(wp_graph, 't5_deai', _deai_state)
                # 优化3：流式模式显式关闭重审校环节（由前端 Spot-Fix/保存后置触发），任务图不留悬空阶段
                wp_runner.mark_stage(wp_graph, 't6_cchk', 'skipped', {'note': '流式一致性检查由前端触发'})
                wp_runner.mark_stage(wp_graph, 't9_cycle', 'skipped', {'note': '流式审校闭环由前端 Spot-Fix 触发'})
                wp_runner.mark_stage(wp_graph, 't11_post', 'declared', {'note': '落库后由 _after_chapter_persisted 执行'})
                try:
                    wp_runner.persist_plan_log(wp_graph, 'generate_chapter', {'outcome': 'ok', 'mode': 'stream'})
                    yield f'data: {json.dumps({"pipeline_plan": wp_graph.to_dict()}, ensure_ascii=False)}\n\n'
                except Exception:
                    pass
        except Exception as e:
            yield f'data: {{"error": "{str(e)[:200]}"}}\n\n'

    return app.response_class(generate(), mimetype='text/event-stream')

@app.route('/api/books/<book_id>/ai-continue-batch', methods=['POST'])
@login_required
def ai_continue_batch(book_id):
    """连续创作模式（优化版）：批量生成 N 章，每章独立生成，普通 JSON 响应。
    每章只调 1 次 LLM（字数修正/去AI味要求并入 system_prompt），不注入 prev_content；
    保留标题解析与一致性检查（默认关闭）。参数：{instruction, skill_pack_ids,
    chapter_lang_styles, count(1-10), start_chapter_num?}"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Not found'}), 404
    data = request.get_json() or {}
    instruction = data.get('instruction', '')
    skill_pack_ids = data.get('skill_pack_ids', [])
    chapter_lang_styles = data.get('chapter_lang_styles', [])
    count = max(1, min(10, int(data.get('count', 3))))  # 1-10 章
    start_chapter_num = data.get('start_chapter_num')
    # Prompt 上下文缓存旁路（批量：每章内部走同一指纹，用户改了未存库的内容时传 True 绕开）
    skip_prompt_cache = bool(data.get('skip_prompt_cache', False) or data.get('skip_cache', False))
    # S1：批量模式下，任意一章触发 critical 时的处理策略
    ignore_gates = bool(data.get('ignore_gates', False))

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    import re as _re_batch
    results = []
    failed = []  # Bug2 修复：记录失败章节，避免静默吞错导致剧情断档
    prev_polished = ''  # Bug5 修复：缓存上一章已生成正文，传递给下一章避免断档

    for i in range(count):
        try:
            # 章号：首章用 start_chapter_num 或自动计算，后续递增
            target_num = start_chapter_num + i if start_chapter_num else None
            # Bug5 修复：传递上一章已生成正文（数据库尚未保存或刚保存的场景都能承接剧情）
            ctx = _build_ai_continue_context(book_id, bb, instruction, skill_pack_ids,
                                              target_num, prev_polished or None, chapter_lang_styles,
                                              skip_cache=skip_prompt_cache)
            api_key = ctx['api_key']
            base_url = ctx['base_url']
            model = ctx['model']
            cur_ch = ctx['current_chapter_num']

            # 标准文风与字数铁律已在 _build_ai_continue_context 内统一注入（三种模式一致）
            # 去AI味禁词已由 chat_collab_bp.DEAI_RULES 统一负责，此处不再追加手工补丁（冗余）
            system_prompt = ctx['system_prompt']

            # Bug3 修复：LLM 调用添加状态码与结构检查，避免 KeyError 静默失败
            try:
                resp = requests.post(f'{base_url}/chat/completions',
                    headers=build_auth_headers(api_key),
                    json={'model': model, 'messages': [{'role':'system','content':system_prompt},
                                                        {'role':'user','content':ctx['user_prompt']}],
                          'temperature': ctx['temperature'], 'max_tokens': ctx['max_tokens']},
                    timeout=180)
            except requests.exceptions.RequestException as re_err:
                try:
                    app.logger.error(f'ai_continue_batch 第{cur_ch}章 LLM 请求失败: {re_err}')
                except Exception:
                    pass
                failed.append({'chapter_num': cur_ch, 'error': f'LLM 请求失败: {str(re_err)[:200]}'})
                break  # 【铁律】失败即停，避免后续章节无意义空跑

            if resp.status_code != 200:
                err_body = resp.text[:300] if hasattr(resp, 'text') else ''
                try:
                    app.logger.error(f'ai_continue_batch 第{cur_ch}章 LLM 返回 HTTP {resp.status_code}: {err_body}')
                except Exception:
                    pass
                failed.append({'chapter_num': cur_ch, 'error': f'LLM HTTP {resp.status_code}'})
                break  # 【铁律】失败即停

            try:
                resp_json = resp.json()
                draft_content = resp_json['choices'][0]['message']['content']
            except (ValueError, KeyError, IndexError, TypeError) as parse_err:
                try:
                    app.logger.error(f'ai_continue_batch 第{cur_ch}章 LLM 返回结构异常: {parse_err}, body={resp.text[:300]}')
                except Exception:
                    pass
                failed.append({'chapter_num': cur_ch, 'error': 'LLM 返回结构异常'})
                break  # 【铁律】失败即停

            if not draft_content or not draft_content.strip():
                failed.append({'chapter_num': cur_ch, 'error': 'LLM 返回内容为空'})
                break  # 【铁律】内容为空即停，避免后续章节无上下文可衔接

            # 【修复】额外校验：剥离内部标签后仍有正文（防止 LLM 只输出标签，门禁误报"正文为空"）
            _pre_check_body = _extract_chapter_body(draft_content)
            _pre_check_body = _re_batch.sub(r'\{[^{}]*"title"\s*:\s*"[^"]*"[^{}]*\}', '', _pre_check_body).strip()
            if not _pre_check_body:
                failed.append({'chapter_num': cur_ch, 'error': 'LLM 仅输出结构标签（pre_write_check/chapter_changes），无正文内容'})
                break  # 【铁律】失败即停

            polished_content = draft_content

            # 一致性检查（默认关闭）
            consistency_passed = True
            consistency_issues = ''

            # 标题自动生成：解析 JSON 标题，剥离正文标签行
            suggested_title = _extract_chapter_title(draft_content, draft_content)

            # Bug7 修复：先从 draft_content 提取 chapter_changes（暂存），等 Chapter flush 拿到 id 后再回写
            chapter_changes_data = None
            if extract_changes:
                try:
                    _, changes = extract_changes(draft_content)
                    if changes:
                        chapter_changes_data = changes
                except Exception:
                    db.session.rollback()

            polished_content = _extract_chapter_body(polished_content)
            # 额外剥离末尾的标题 JSON 块（_extract_chapter_body 未覆盖）
            polished_content = _re_batch.sub(r'\{[^{}]*"title"\s*:\s*"[^"]*"[^{}]*\}', '', polished_content).rstrip()

            # 【修复】batch模式补字数修正：与多Agent模式统一，初稿字数不在2300-2500时AI重写
            try:
                polished_content, wc_note = _ensure_word_count(
                    polished_content, api_key, base_url, model, ctx['max_tokens'], cur_ch)
            except Exception:
                wc_note = ''

            # Bug9：统一 count_words 统计字数；★连续模式也跑 post_validate（与多agent对齐）
            wc = count_words(polished_content)
            post_validate = None
            if validate_chapter_with_bible:
                try:
                    bible_ctx = {
                        'character_profiles': bb.character_profiles or '',
                        'chapter_changes_log': bb.chapter_changes_log or '',
                        'key_rules': bb.key_rules or '',
                        'worldbuilding': bb.worldbuilding or '',
                        'inventory': bb.inventory or '',
                        'locations': bb.locations or '',
                        'foreshadowing': bb.foreshadowing or '',
                    } if bb else None
                    validation = validate_chapter_with_bible(polished_content, bible_ctx)
                    # 【P0修复】追加文风漂移检测：与多Agent同步/单章模式一致，避免死代码
                    try:
                        style_baseline = _compute_style_baseline(book_id, cur_ch)
                        if validate_chapter_with_drift and style_baseline:
                            drift_validation = validate_chapter_with_drift(polished_content, style_baseline)
                            for issue in drift_validation.issues:
                                validation.issues.append(issue)
                                validation.score = max(0, validation.score - (5 if issue.severity == 'warning' else 20))
                            if drift_validation.stats.get('style_drift'):
                                validation.stats['style_drift'] = drift_validation.stats['style_drift']
                    except Exception:
                        pass
                    if validation.issues:
                        post_validate = validation.to_dict()
                except Exception:
                    pass

            # 校验报告
            chapter_score = _calc_chapter_score(post_validate, consistency_passed, consistency_issues, {}, wc, ctx.get('chapter_plan', ''))

            # Bug7 修复：先创建 Chapter 并 flush 拿到 ch.id，再回写 chapter_changes（chapter_id 不再为空串）
            # 统一标题格式：第X章 标题文本（与多Agent同步/流式模式一致）
            title = _format_chapter_title(cur_ch, suggested_title)
            max_order = db.session.query(db.func.max(Chapter.order_index)).filter_by(book_id=book_id).scalar() or -1
            # Bug1 修复：归卷直接设置 parent_id（用 ctx 中的 vol_chapter），不依赖末尾 resort
            parent_id = ctx['vol_chapter'].id if ctx.get('vol_chapter') else ''
            ch = Chapter(book_id=book_id, title=title, content=polished_content,
                         order_index=max_order + 1, is_volume=False,
                         parent_id=parent_id,
                         word_count=wc)
            db.session.add(ch)
            db.session.flush()  # 拿到 ch.id

            # Bug7 修复：用真实 ch.id 回写 chapter_changes（不再传空串）
            if chapter_changes_data and apply_chapter_changes:
                try:
                    apply_chapter_changes(
                        bb, ch.id, cur_ch,
                        ctx.get('vol_index', 0), chapter_changes_data,
                    )
                except Exception:
                    db.session.rollback()

            update_book_stats(book_id)
            db.session.commit()

            # 【P1-4修复】连续创作模式补落地门禁（与多Agent模式对齐，仅 warning 不阻断）
            # 防御性检查：polished_content 为空时跳过门禁，避免误报"正文为空"
            gate_result_batch = None
            if run_all_gates and polished_content and polished_content.strip():
                try:
                    gate_result_batch = run_all_gates(polished_content, bb, cur_ch)
                except Exception:
                    pass

            # Bug4 修复：每章保存后检查并自动生成动态报告（每5章触发），避免后续章节注入过时报告
            try:
                _check_and_auto_generate_report(book_id)
            except Exception:
                pass
            # 【P0-3】每 20 章自动触发防遗忘检查（daemon 线程，不阻塞批处理）
            try:
                _maybe_auto_trigger_anti_forget_check(book_id, cur_ch)
            except Exception:
                pass

            result_entry = {
                'chapter_num': cur_ch,
                'chapter_id': ch.id,
                'title': title,
                'content': polished_content,
                'word_count': wc,
            }
            if gate_result_batch and not gate_result_batch.get('passed'):
                result_entry['gate_warning'] = gate_result_batch
            # S1：批量模式 critical 命中时：
            # - 若 ignore_gates=False：整批回滚 + 返回 gate_blocked（前端弹窗让用户选择重试）
            # - 若 ignore_gates=True：照常写入，只附 gate_warning
            if gate_result_batch and gate_result_batch.get('blocked') and not ignore_gates:
                db.session.rollback()
                # 收集已完成的章节摘要一起返回，方便用户知道是哪一章触发
                return jsonify({
                    'gate_blocked': True,
                    'ignore_gates_required': True,
                    'blocked_at_chapter': cur_ch,
                    'block_reason': (
                        f'批量创作第{cur_ch}章触发落地门禁 critical，为避免半成品写入已回滚整批。'
                        '请在前端勾选「忽略门禁强制保存」后再次提交整批。'
                    ),
                    'gate_result': gate_result_batch,
                    'partial_results': [{'chapter_num': r['chapter_num'], 'title': r['title'],
                                         'word_count': r.get('word_count', 0)} for r in results],
                }), 428
            results.append(result_entry)
            # Bug5 修复：缓存本章正文供下一章承接
            prev_polished = polished_content
        except Exception as e:
            # Bug2 修复：单章失败记录日志与失败列表，避免静默吞错导致剧情断档
            fail_num = start_chapter_num + i if start_chapter_num else i + 1
            try:
                app.logger.error(f'ai_continue_batch 第{fail_num}章生成失败: {e}', exc_info=True)
            except Exception:
                pass
            db.session.rollback()
            failed.append({'chapter_num': fail_num, 'error': str(e)[:200]})
            # Bug5：失败时清空 prev_polished，避免把失败章当上一章注入
            prev_polished = ''
            break  # 【铁律】异常即停，避免错误级联放大

    # Bug1 修复：不再调 resort_chapters_by_title(rebin_volumes=True)——批处理标题可能无"第X章"
    # 前缀致混合排序错乱，且 rebin 会覆盖自定义卷名；章节已按 order_index=max+1 顺序保存
    return jsonify({
        'chapters': results,
        'total': len(results),
        'failed': failed,
        'failed_count': len(failed),
    })

def _stream_llm_chunks_with_heartbeat(resp, chapter_num, last_heartbeat, heartbeat_interval=5):
    """后台线程读取 LLM 流式响应，主生成器从队列消费，无数据时 yield 心跳。

    修复 network error 根因：resp.iter_lines() 是阻塞调用，原实现把心跳检查放在
    for 循环体内，LLM 一旦长时间不输出（首 token 延迟/思考阶段/网络抖动），
    循环体不执行 → 心跳无法推送 → Render 代理判定空闲超时 → 断连 network error。

    本函数用独立线程读取流，主流程 queue.get(timeout=3) 非阻塞消费：
    - 收到数据 → 处理 delta，重置心跳计时
    - 3s 无数据 → 检查是否到 5s 心跳间隔，是则 yield 心跳

    yield 元素：('chunk', delta_str) 或 ('heartbeat',) 或 ('done',)
    异常通过 ('error', exception) 返回。"""
    import threading
    import queue as _queue
    import time as _t

    chunk_q = _queue.Queue()
    reader_err = [None]

    def _reader():
        try:
            for line in resp.iter_lines():
                if line:
                    chunk_q.put(line.decode('utf-8', errors='replace'))
        except Exception as e:
            reader_err[0] = e
        finally:
            chunk_q.put(None)  # sentinel 表示流结束

    reader_t = threading.Thread(target=_reader, daemon=True)
    reader_t.start()

    last_hb = last_heartbeat
    while True:
        try:
            item = chunk_q.get(timeout=3)
        except _queue.Empty:
            now = _t.time()
            if now - last_hb >= heartbeat_interval:
                yield ('heartbeat',)
                last_hb = now
            continue

        if item is None:
            if reader_err[0]:
                yield ('error', reader_err[0])
            yield ('done',)
            return

        if item.startswith('data: '):
            chunk = item[6:]
            if chunk == '[DONE]':
                yield ('done',)
                return
            try:
                chunk_data = json.loads(chunk)
                delta = chunk_data.get('choices', [{}])[0].get('delta', {}).get('content', '')
                if delta:
                    yield ('chunk', delta)
                    last_hb = _t.time()
            except Exception:
                pass

def _run_blocking_with_heartbeat(func, heartbeat_msg, heartbeat_interval=5):
    """在后台线程运行阻塞函数 func，主生成器 yield 心跳保持 SSE 连接活跃。

    用于包装非流式的 LLM 调用（如去AI味、动态报告生成）——这些调用阻塞数十秒，
    期间无法 yield 任何数据，Render 代理空闲超时会断连。

    用法（在 SSE 生成器内）：
        result = yield from _run_blocking_with_heartbeat(
            lambda: requests.post(...), '正在去AI味审校...')

    注意：后台线程会自动 push Flask app context，确保 func 内的 DB 操作
    (db.session / Model.query) 不会因 'Working outside of application context' 报错。
    """
    import threading
    import queue as _queue
    import time as _t

    result_q = _queue.Queue()
    _app = app  # 捕获 app 引用，后台线程需 push app context 才能 DB 操作

    def _runner():
        try:
            with _app.app_context():
                result_q.put(('ok', func()))
        except Exception as e:
            result_q.put(('err', e))

    t = threading.Thread(target=_runner, daemon=True)
    t.start()

    last_hb = _t.time()
    while True:
        try:
            status, val = result_q.get(timeout=3)
            if status == 'err':
                raise val
            return val
        except _queue.Empty:
            now = _t.time()
            if now - last_hb >= heartbeat_interval:
                yield heartbeat_msg
                last_hb = now

@app.route('/api/books/<book_id>/ai-continue-batch/stream', methods=['POST'])
@login_required
def ai_continue_batch_stream(book_id):
    """连续创作流式版（SSE）：解决 Render 同步请求超时（约100s）导致 Failed to fetch。
    每章 LLM 调用 stream 模式逐 chunk 收集并定期推心跳保活；事件类型：
    chapter_start / heartbeat(每5s防空闲超时) / chapter_done(含完整章节) /
    chapter_failed / batch_done / error(致命终止)。"""
    import time as _time

    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Not found'}), 404
    data = request.get_json() or {}
    instruction = data.get('instruction', '')
    skill_pack_ids = data.get('skill_pack_ids', [])
    chapter_lang_styles = data.get('chapter_lang_styles', [])
    count = max(1, min(10, int(data.get('count', 3))))
    start_chapter_num = data.get('start_chapter_num')
    # Prompt 上下文缓存旁路
    skip_prompt_cache = bool(data.get('skip_prompt_cache', False) or data.get('skip_cache', False))
    # S1：流式批量模式下，critical 命中时推送 gate_blocked 事件并提前结束
    ignore_gates = bool(data.get('ignore_gates', False))

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    def generate():
        import re as _re_bs
        results = []
        failed = []
        prev_polished = ''

        for i in range(count):
            try:
                target_num = start_chapter_num + i if start_chapter_num else None
                # 先推送 chapter_start（用预估章号），让 Render 立即收到首字节，
                # 避免 _build_ai_continue_context 内 _generate_chapter_plan 调 LLM 阻塞 20-30s 期间无数据导致断开
                yield f'data: {json.dumps({"type": "chapter_start", "chapter_num": target_num or (i + 1), "message": f"正在准备第{target_num or (i + 1)}章上下文..."}, ensure_ascii=False)}\n\n'

                ctx = _build_ai_continue_context(book_id, bb, instruction, skill_pack_ids,
                                                  target_num, prev_polished or None, chapter_lang_styles,
                                                  skip_chapter_plan=True, skip_cache=skip_prompt_cache)
                api_key = ctx['api_key']
                base_url = ctx['base_url']
                model = ctx['model']
                cur_ch = ctx['current_chapter_num']

                # 章号可能与预估不同（后端基于 max(order_index)+1 重新计算），推送修正事件
                if cur_ch != (target_num or (i + 1)):
                    yield f'data: {json.dumps({"type": "heartbeat", "chapter_num": cur_ch, "message": f"正在生成第{cur_ch}章..."}, ensure_ascii=False)}\n\n'

                # 标准文风与字数铁律已在 _build_ai_continue_context 内统一注入（三种模式一致）
                # 去AI味禁词已由 chat_collab_bp.DEAI_RULES 统一负责，此处不再追加手工补丁（冗余）
                system_prompt = ctx['system_prompt']

                # LLM 流式调用：逐 chunk 收集，定期推送心跳保持连接活跃
                full_parts = []
                last_heartbeat = _time.time()
                try:
                    resp = requests.post(f'{base_url}/chat/completions',
                        headers=build_auth_headers(api_key),
                        json={'model': model, 'messages': [{'role':'system','content':system_prompt},
                                                            {'role':'user','content':ctx['user_prompt']}],
                              'temperature': ctx['temperature'], 'max_tokens': ctx['max_tokens'],
                              'stream': True},
                        stream=True, timeout=180)
                except requests.exceptions.RequestException as re_err:
                    try:
                        app.logger.error(f'ai_continue_batch_stream 第{cur_ch}章 LLM 请求失败: {re_err}')
                    except Exception:
                        pass
                    failed.append({'chapter_num': cur_ch, 'error': f'LLM 请求失败: {str(re_err)[:200]}'})
                    yield f'data: {json.dumps({"type": "chapter_failed", "chapter_num": cur_ch, "error": f"LLM 请求失败"}, ensure_ascii=False)}\n\n'
                    # 【铁律】前面章节失败，后续章节自动停止，避免连续生成无意义空章节
                    yield f'data: {json.dumps({"type": "batch_stopped", "reason": f"第{cur_ch}章 LLM 请求失败，后续章节已自动停止", "failed_chapter": cur_ch}, ensure_ascii=False)}\n\n'
                    break

                if resp.status_code != 200:
                    err_body = resp.text[:300] if hasattr(resp, 'text') else ''
                    try:
                        app.logger.error(f'ai_continue_batch_stream 第{cur_ch}章 LLM HTTP {resp.status_code}: {err_body}')
                    except Exception:
                        pass
                    failed.append({'chapter_num': cur_ch, 'error': f'LLM HTTP {resp.status_code}'})
                    yield f'data: {json.dumps({"type": "chapter_failed", "chapter_num": cur_ch, "error": f"LLM HTTP {resp.status_code}"}, ensure_ascii=False)}\n\n'
                    # 【铁律】前面章节失败，后续章节自动停止
                    yield f'data: {json.dumps({"type": "batch_stopped", "reason": f"第{cur_ch}章 LLM 返回 HTTP {resp.status_code}，后续章节已自动停止", "failed_chapter": cur_ch}, ensure_ascii=False)}\n\n'
                    break

                # 后台线程读取流式响应，主生成器从队列消费，无数据时 yield 心跳
                # 修复 network error：iter_lines() 阻塞时心跳无法推送
                for evt in _stream_llm_chunks_with_heartbeat(resp, cur_ch, last_heartbeat):
                    if evt[0] == 'chunk':
                        full_parts.append(evt[1])
                        last_heartbeat = _time.time()
                    elif evt[0] == 'heartbeat':
                        yield f'data: {json.dumps({"type": "heartbeat", "chapter_num": cur_ch, "message": f"正在生成第{cur_ch}章..."}, ensure_ascii=False)}\n\n'
                        last_heartbeat = _time.time()
                    elif evt[0] == 'error':
                        raise evt[1]
                    elif evt[0] == 'done':
                        break
                try:
                    resp.close()
                except Exception:
                    pass

                draft_content = ''.join(full_parts)
                if not draft_content or not draft_content.strip():
                    failed.append({'chapter_num': cur_ch, 'error': 'LLM 返回内容为空'})
                    yield f'data: {json.dumps({"type": "chapter_failed", "chapter_num": cur_ch, "error": "LLM 返回内容为空"}, ensure_ascii=False)}\n\n'
                    # 【铁律】LLM 返回内容为空，后续章节自动停止（空章节无上下文可衔接，继续生成只会产出更多空内容）
                    yield f'data: {json.dumps({"type": "batch_stopped", "reason": f"第{cur_ch}章 LLM 返回内容为空，后续章节已自动停止", "failed_chapter": cur_ch}, ensure_ascii=False)}\n\n'
                    break

                # 【修复】额外校验：剥离内部标签后是否仍有正文（防止 LLM 只输出标签而无正文，
                # 后续 _extract_chapter_body 后变空，门禁误报"正文为空"）
                _pre_check_body = _extract_chapter_body(draft_content)
                _pre_check_body = _re_bs.sub(r'\{[^{}]*"title"\s*:\s*"[^"]*"[^{}]*\}', '', _pre_check_body).strip()
                if not _pre_check_body:
                    err_msg = 'LLM 仅输出结构标签（pre_write_check/chapter_changes），无正文内容'
                    failed.append({'chapter_num': cur_ch, 'error': err_msg})
                    yield f'data: {json.dumps({"type": "chapter_failed", "chapter_num": cur_ch, "error": err_msg}, ensure_ascii=False)}\n\n'
                    yield f'data: {json.dumps({"type": "batch_stopped", "reason": f"第{cur_ch}章 {err_msg}，后续章节已自动停止", "failed_chapter": cur_ch}, ensure_ascii=False)}\n\n'
                    break

                polished_content = draft_content
                suggested_title = _extract_chapter_title(draft_content, draft_content)

                # 提取 chapter_changes（暂存，待 flush 后回写）
                chapter_changes_data = None
                if extract_changes:
                    try:
                        _, changes = extract_changes(draft_content)
                        if changes:
                            chapter_changes_data = changes
                    except Exception:
                        db.session.rollback()

                polished_content = _extract_chapter_body(polished_content)
                polished_content = _re_bs.sub(r'\{[^{}]*"title"\s*:\s*"[^"]*"[^{}]*\}', '', polished_content).rstrip()

                # ===== 去AI味审校 Agent（选项A：批处理也跑去AI味修正，与单章模式对齐）=====
                # 2026-08-23 默认启用：内置统一去AI规则（build_review_rules）常驻，
                # 不再依赖技能包勾选；审查类(review)技能包作为增强叠加。
                # 字数校验通过→用修正版；字数异常/失败→回滚用初稿
                deai_status = 'skipped'  # skipped / rules_ok / rules_missing / success / failed
                deai_rules_block = ''
                try:
                    deai_rules_block, _deai_build_status = _build_deai_rules_block(skill_pack_ids, book)
                    deai_status = 'rules_ok' if _deai_build_status == 'ok' else 'rules_missing'
                except Exception as _deai_e:
                    # helper 也兜不住（连 DEAI_ONLY_RULES import 都失败），显式标 missing 而非静默
                    app.logger.error(f'ai_continue_batch_stream 去AI规则构建失败: {_deai_e}')
                    print(f'[去AI] ai_continue_batch_stream 去AI规则构建失败: {_deai_e}', file=sys.stderr)
                    deai_status = 'rules_missing'
                if deai_rules_block:
                    # 推送心跳：去AI味审校中
                    yield f'data: {json.dumps({"type": "heartbeat", "chapter_num": cur_ch, "message": f"正在去AI味审校第{cur_ch}章..."}, ensure_ascii=False)}\n\n'
                    deai_system = ("你是番茄去AI味审查员。对以下刚写好的章节正文做去AI味审校，按规则修改后只输出修改后的正文。\n\n"
                                   + deai_rules_block
                                   + "\n\n【硬性约束】修改后字数仍须 2400±100（2300-2500区间，含标点），保留原章节的剧情走向和钩子，只改文风不改剧情。")
                    try:
                        deai_hb = f'data: {json.dumps({"type": "heartbeat", "chapter_num": cur_ch, "message": f"正在去AI味审校第{cur_ch}章..."}, ensure_ascii=False)}\n\n'
                        deai_resp = yield from _run_blocking_with_heartbeat(
                            lambda: requests.post(f'{base_url}/chat/completions',
                                headers=build_auth_headers(api_key),
                                json={'model': model,
                                      'messages': [{'role':'system','content':deai_system},
                                                   {'role':'user','content':f'请审校以下章节正文：\n\n{polished_content}'}],
                                      'temperature': 0.5, 'max_tokens': ctx['max_tokens']},
                                timeout=180),
                            deai_hb)
                        if deai_resp.status_code == 200:
                            deai_result = deai_resp.json()
                            deai_polished = deai_result['choices'][0]['message']['content'].strip()
                            # 剥离可能的内部标签（去AI味 LLM 偶尔会带上）
                            deai_polished = _extract_chapter_body(deai_polished)
                            deai_polished = _re_bs.sub(r'\{[^{}]*"title"\s*:\s*"[^"]*"[^{}]*\}', '', deai_polished).rstrip()
                            deai_wc = count_words(deai_polished)
                            if deai_polished and 2300 <= deai_wc <= 2500:
                                polished_content = deai_polished
                                deai_status = 'success'
                            elif deai_polished and deai_wc > 500:
                                deai_status = 'failed'  # 字数异常，回滚用初稿
                            else:
                                deai_status = 'failed'  # 内容为空，回滚用初稿
                        else:
                            deai_status = 'failed'
                            try:
                                app.logger.error(f'ai_continue_batch_stream 第{cur_ch}章 去AI味 HTTP {deai_resp.status_code}')
                            except Exception:
                                pass
                    except Exception as deai_err:
                        deai_status = 'failed'
                        try:
                            app.logger.error(f'ai_continue_batch_stream 第{cur_ch}章 去AI味异常: {deai_err}')
                        except Exception:
                            pass

                # 【修复】batch_stream模式补字数修正：与多Agent模式统一
                try:
                    polished_content, wc_note_bs = _ensure_word_count(
                        polished_content, api_key, base_url, model, ctx['max_tokens'], cur_ch)
                except Exception:
                    wc_note_bs = ''

                wc = count_words(polished_content)

                # post_validate（去AI味检测）
                post_validate = None
                if validate_chapter_with_bible:
                    try:
                        bible_ctx = {
                            'character_profiles': bb.character_profiles or '',
                            'chapter_changes_log': bb.chapter_changes_log or '',
                            'key_rules': bb.key_rules or '',
                            'worldbuilding': bb.worldbuilding or '',
                            'inventory': bb.inventory or '',
                            'locations': bb.locations or '',
                            'foreshadowing': bb.foreshadowing or '',
                        } if bb else None
                        validation = validate_chapter_with_bible(polished_content, bible_ctx)
                        # 【P0修复】追加文风漂移检测：与多Agent同步/连续同步模式一致，避免死代码
                        try:
                            style_baseline = _compute_style_baseline(book_id, cur_ch)
                            if validate_chapter_with_drift and style_baseline:
                                drift_validation = validate_chapter_with_drift(polished_content, style_baseline)
                                for issue in drift_validation.issues:
                                    validation.issues.append(issue)
                                    validation.score = max(0, validation.score - (5 if issue.severity == 'warning' else 20))
                                if drift_validation.stats.get('style_drift'):
                                    validation.stats['style_drift'] = drift_validation.stats['style_drift']
                        except Exception:
                            pass
                        if validation.issues:
                            post_validate = validation.to_dict()
                    except Exception:
                        pass

                chapter_score = _calc_chapter_score(post_validate, True, '', {}, wc, ctx.get('chapter_plan', ''))

                # 创建 Chapter 并 flush 拿 id
                # 统一标题格式：第X章 标题文本（与多Agent同步/连续同步模式一致）
                title = _format_chapter_title(cur_ch, suggested_title)
                max_order = db.session.query(db.func.max(Chapter.order_index)).filter_by(book_id=book_id).scalar() or -1
                parent_id = ctx['vol_chapter'].id if ctx.get('vol_chapter') else ''
                ch = Chapter(book_id=book_id, title=title, content=polished_content,
                             order_index=max_order + 1, is_volume=False,
                             parent_id=parent_id,
                             word_count=wc)
                db.session.add(ch)
                db.session.flush()

                # 用真实 ch.id 回写 chapter_changes
                if chapter_changes_data and apply_chapter_changes:
                    try:
                        apply_chapter_changes(bb, ch.id, cur_ch, ctx.get('vol_index', 0), chapter_changes_data)
                    except Exception:
                        db.session.rollback()

                update_book_stats(book_id)
                db.session.commit()

                # 【P1-4修复】批处理流式补落地门禁（与多Agent模式对齐，仅 warning 不阻断）
                # 防御性检查：polished_content 为空时跳过门禁，避免误报"正文为空"
                gate_result_bstream = None
                if run_all_gates and polished_content and polished_content.strip():
                    try:
                        gate_result_bstream = run_all_gates(polished_content, bb, cur_ch)
                    except Exception:
                        pass

                # 每章后检查自动生成动态报告（可能触发 LLM 调用，用线程+心跳避免阻塞）
                try:
                    report_hb = f'data: {json.dumps({"type": "heartbeat", "chapter_num": cur_ch, "message": f"正在更新动态报告..."}, ensure_ascii=False)}\n\n'
                    yield from _run_blocking_with_heartbeat(
                        lambda: _check_and_auto_generate_report(book_id),
                        report_hb)
                except Exception:
                    pass
                # 【P0-3】每 20 章自动触发防遗忘检查（daemon 线程，不阻塞 SSE 流）
                try:
                    _maybe_auto_trigger_anti_forget_check(book_id, cur_ch)
                except Exception:
                    pass

                chapter_info = {
                    'chapter_num': cur_ch,
                    'chapter_id': ch.id,
                    'title': title,
                    'content': polished_content,
                    'word_count': wc,
                    'deai_status': deai_status,  # skipped/success/failed，告知前端是否做过去AI味修正
                }
                if gate_result_bstream and not gate_result_bstream.get('passed'):
                    chapter_info['gate_warning'] = gate_result_bstream
                # S1：流式批量模式 critical 命中：
                # - 回滚本批 DB 改动 + 推送 gate_blocked 事件 + 结束流
                if gate_result_bstream and gate_result_bstream.get('blocked') and not ignore_gates:
                    try: db.session.rollback()
                    except Exception: pass
                    yield f'data: {json.dumps({"type": "gate_blocked", "ignore_gates_required": True, "chapter_num": cur_ch, "gate_result": gate_result_bstream, "block_reason": f"第{cur_ch}章触发落地门禁 critical，已回滚本批写入。勾选忽略门禁后再提交。"}, ensure_ascii=False)}\n\n'
                    # 直接结束流
                    yield f'data: {json.dumps({"type": "batch_done", "total": i, "failed_count": 1, "failed": [{"chapter_num": cur_ch, "error": "gate_blocked"}]}, ensure_ascii=False)}\n\n'
                    return
                results.append(chapter_info)
                prev_polished = polished_content

                # 推送章节完成事件
                yield f'data: {json.dumps({"type": "chapter_done", "chapter": chapter_info}, ensure_ascii=False)}\n\n'

            except Exception as e:
                fail_num = start_chapter_num + i if start_chapter_num else i + 1
                try:
                    app.logger.error(f'ai_continue_batch_stream 第{fail_num}章生成失败: {e}', exc_info=True)
                except Exception:
                    pass
                db.session.rollback()
                failed.append({'chapter_num': fail_num, 'error': str(e)[:200]})
                yield f'data: {json.dumps({"type": "chapter_failed", "chapter_num": fail_num, "error": str(e)[:200]}, ensure_ascii=False)}\n\n'
                # 【铁律】章节生成异常，后续章节自动停止，避免错误级联放大
                yield f'data: {json.dumps({"type": "batch_stopped", "reason": f"第{fail_num}章生成异常：{str(e)[:100]}，后续章节已自动停止", "failed_chapter": fail_num}, ensure_ascii=False)}\n\n'
                prev_polished = ''
                break

        # 推送批处理完成事件
        yield f'data: {json.dumps({"type": "batch_done", "total": len(results), "failed_count": len(failed), "failed": failed}, ensure_ascii=False)}\n\n'

    # stream_with_context 保住应用上下文（多章串行含大量 DB 操作，yield 后丢上下文会 network error）；
    # 响应头禁缓存/禁代理缓冲（Cloudflare/Render 默认缓冲 SSE 致浏览器长时间收不到数据）
    resp = app.response_class(stream_with_context(generate()), mimetype='text/event-stream')
    resp.headers['Cache-Control'] = 'no-cache'
    resp.headers['X-Accel-Buffering'] = 'no'
    resp.headers['Connection'] = 'keep-alive'
    return resp

def _build_deai_rules_block(skill_pack_ids, book):
    """统一构建去AI味审校规则块（单章 ai_continue 与批 ai_continue_batch_stream 共用）。

    消除两处 7 行重复；同时解决"去AI规则构建失败/为空时静默跳过"的隐患：
    - build_review_rules 内部已保证 parts 含 DEAI_ONLY_RULES（核心去AI表），技能包/import 异常只丢技能包；
    - 此处再兜底：即便 build_review_rules 也抛异常或返回空串，也强制回退注入核心 DEAI_ONLY_RULES，
      保证去AI审校环节永不因构建问题而静默缺失。
    返回 (rules_block, status)：status = 'ok' / 'fallback'。
    """
    from blueprints.chat_collab_bp import build_review_rules, DEAI_ONLY_RULES
    status = 'ok'
    try:
        block = build_review_rules(
            skill_pack_ids, mode='agent',
            prompt_keys_filter=['tomato_deai', 'de_ai_flavor', 'polish'], book=book)
        if not block:
            raise RuntimeError('build_review_rules 返回空串')
    except Exception as _deai_e:
        status = 'fallback'
        app.logger.error(f'去AI味规则构建异常，已回退注入核心去AI表: {_deai_e}')
        print(f'[去AI] 规则构建异常，已回退注入核心去AI表: {_deai_e}', file=sys.stderr)
        block = DEAI_ONLY_RULES
    return block, status

def _extract_chapter_body(full_content: str) -> str:
    """从 LLM 完整输出中剥离 PRE_WRITE_CHECK、chapter_changes、【标题】标签，只保留正文。
    P1-6 启用后 LLM 会输出结构化标签，校验器只检查正文部分。"""
    import re as _re
    body = full_content
    # 剥离 <pre_write_check>...</pre_write_check>
    body = _re.sub(r'<pre_write_check>[\s\S]*?</pre_write_check>', '', body, flags=_re.IGNORECASE)
    # 剥离 <chapter_changes>...</chapter_changes>
    body = _re.sub(r'<chapter_changes>[\s\S]*?</chapter_changes>', '', body, flags=_re.IGNORECASE)
    # 剥离 【标题】... 标签行（标题自动生成产物）
    body = _re.sub(r'【标题】[^\n]*', '', body)
    return body.strip()

def _format_chapter_title(chapter_num, suggested_title):
    """统一章节标题格式：第X章 标题文本

    三种创作模式（多Agent同步 / 流式 / 连续创作流式）统一使用此函数格式化标题，
    确保所有章节标题格式一致，混用模式时不会出现格式混乱。

    规则：
    - 有 suggested_title：格式为“第{章号}章 {标题}”
    - 无 suggested_title：格式为“第{章号}章”
    - suggested_title 已含“第X章”前缀时去重，避免“第1章 第1章 xxx”
    """
    import re as _re
    prefix = f'第{chapter_num}章'
    if not suggested_title or not suggested_title.strip():
        return prefix
    title = suggested_title.strip()
    # 去除标题中可能自带的“第X章”前缀（LLM 偶尔会带上）
    title = _re.sub(r'^第[一二三四五六七八九十百零0-9]+章[\s:：]*', '', title).strip()
    if not title:
        return prefix
    return f'{prefix} {title}'

def _extract_chapter_title(full_content: str, fallback_content: str = '') -> str:
    """从 LLM 完整输出中解析标题。
    优先解析末尾的 JSON：{"title": "标题文本"}
    解析失败时用章节前 20 字自动生成标题（兜底）。
    返回标题文本（无则空串）。"""
    import re as _re
    
    # 尝试解析末尾的 JSON 块
    try:
        # 查找最后一个 JSON 对象（从后往前找）
        json_matches = list(_re.finditer(r'\{[^{}]*"title"\s*:\s*"([^"]+)"[^{}]*\}', full_content))
        if json_matches:
            # 取最后一个匹配的 JSON
            last_match = json_matches[-1]
            title = last_match.group(1).strip()
            # 去除可能的"第X章"前缀
            title = _re.sub(r'^第[一二三四五六七八九十百零0-9]+章[\s:：]*', '', title)
            if title:
                return title[:30]  # 限长 30 字
    except Exception:
        pass
    
    # 兜底：用章节前 20 字生成标题
    if fallback_content:
        # 去除空白和标点，取前 20 字
        clean = _re.sub(r'\s+', '', fallback_content)
        # 去除常见的开头标记
        clean = _re.sub(r'^[^a-zA-Z\u4e00-\u9fa5]+', '', clean)
        if clean:
            return clean[:20] + ('...' if len(clean) > 20 else '')
    
    return ''

def _calc_chapter_score(post_validate, consistency_passed, consistency_issues, gate_result, word_count,
                        chapter_plan=None):
    """章节审校校验报告：聚合 post_validate 和 consistency_issues 的结果。
    返回 dict：{has_issues: bool, issues: [{'type': str, 'severity': str, 'description': str}]}。
    不再返回 0-100 分数和等级，简化为校验问题列表。"""
    if not isinstance(post_validate, dict):
        post_validate = {}
    
    issues = []
    
    # ① AI痕迹检测（post_validate）
    pv_issues = post_validate.get('issues', []) or []
    for iss in pv_issues:
        issues.append({
            'type': 'ai_trace',
            'severity': iss.get('severity', 'warning'),
            'description': f"[{iss.get('category', '未知')}] {iss.get('pattern', '')} — {iss.get('suggestion', '')}"
        })
    
    # ② 一致性检查（consistency_issues）
    if not consistency_passed and consistency_issues:
        # consistency_issues 可能是字符串（分号分隔）或列表
        if isinstance(consistency_issues, str):
            cons_issue_list = [s.strip() for s in consistency_issues.split('；') if s.strip()]
        elif isinstance(consistency_issues, list):
            cons_issue_list = consistency_issues
        else:
            cons_issue_list = []
        
        for issue_text in cons_issue_list:
            issues.append({
                'type': 'consistency',
                'severity': 'warning',
                'description': issue_text
            })
    
    return {
        'has_issues': len(issues) > 0,
        'issues': issues
    }

# ==== LLM 调用辅助函数 ====
def _extract_json_from_llm(content, expect='auto'):
    """从 LLM 返回内容中健壮地提取 JSON 对象或数组。
    处理常见情况：markdown代码块包裹、前后说明文字、多个JSON候选、尾随逗号等。
    expect: 'object' 只接受 dict / 'array' 只接受 list / 'auto' 两者皆可。
    返回 (parsed, error)。成功时 error 为 None。"""
    if not content or not isinstance(content, str):
        return None, 'LLM 返回为空'

    import re

    # 1. 去除 markdown 代码块标记（```json ... ``` 或 ``` ... ```）
    cleaned = re.sub(r'```(?:json|JSON)?\s*', '', content)
    cleaned = re.sub(r'```\s*$', '', cleaned).strip()

    # 2. 直接尝试整体解析（LLM 有时严格只输出 JSON）
    try:
        parsed = json.loads(cleaned)
        if expect == 'object' and isinstance(parsed, dict):
            return parsed, None
        if expect == 'array' and isinstance(parsed, list):
            return parsed, None
        if expect == 'auto' and isinstance(parsed, (dict, list)):
            return parsed, None
    except (json.JSONDecodeError, ValueError):
        pass

    # 3. 遍历所有 {...} 候选（非贪婪，从后往前找最大的），逐个尝试解析
    # 用栈匹配花括号，避免正则贪婪问题
    candidates = []
    # 对象候选 {...}
    for m in re.finditer(r'\{', cleaned):
        start = m.start()
        depth = 0
        for i in range(start, len(cleaned)):
            if cleaned[i] == '{':
                depth += 1
            elif cleaned[i] == '}':
                depth -= 1
                if depth == 0:
                    candidates.append(cleaned[start:i+1])
                    break
    # 数组候选 [...]
    for m in re.finditer(r'\[', cleaned):
        start = m.start()
        depth = 0
        for i in range(start, len(cleaned)):
            if cleaned[i] == '[':
                depth += 1
            elif cleaned[i] == ']':
                depth -= 1
                if depth == 0:
                    candidates.append(cleaned[start:i+1])
                    break

    # 优先返回最长的候选（通常是完整的 JSON）
    candidates.sort(key=len, reverse=True)
    last_error = None
    for cand in candidates:
        try:
            parsed = json.loads(cand)
            if expect == 'object' and isinstance(parsed, dict):
                return parsed, None
            if expect == 'array' and isinstance(parsed, list):
                return parsed, None
            if expect == 'auto' and isinstance(parsed, (dict, list)):
                return parsed, None
        except (json.JSONDecodeError, ValueError) as e:
            last_error = str(e)
            # 尝试修复常见 JSON 瑕疵：尾随逗号
            try:
                fixed = re.sub(r',\s*([}\]])', r'\1', cand)
                parsed = json.loads(fixed)
                if expect == 'object' and isinstance(parsed, dict):
                    return parsed, None
                if expect == 'array' and isinstance(parsed, list):
                    return parsed, None
                if expect == 'auto' and isinstance(parsed, (dict, list)):
                    return parsed, None
            except (json.JSONDecodeError, ValueError):
                pass

    return None, f'未找到有效JSON（最后错误: {last_error}），原始内容前300字: {content[:300]}'

def _call_llm(messages, max_tokens=None, temperature=None, task_type='creation', retry_count=2, timeout=300):
    """统一的 LLM 调用辅助函数，返回 (content, error)
    task_type: 'creation'用主模型(创作/写作)，'recognition'用识别模型(识别/分析/检查，为空时回退主模型)
    max_tokens 语义：None → 用 cfg.max_tokens；0 → 不限制（不下发该字段，用模型自身默认上限）；正整数 → 显式限定。
    retry_count: 临时性错误（空响应/5xx/连接异常）的重试次数，默认 2 次（共 3 次尝试）。"""
    cfg = AIConfig.get_active()
    if not cfg or not cfg.api_key:
        return None, '请先配置 AI 模型 API Key'

    base = cfg.base_url.rstrip('/')
    if not base.endswith('/v1'):
        base += '/v1'
    model = cfg.get_model_for_task(task_type)
    payload = {
        'model': model,
        'messages': messages,
        'temperature': temperature if temperature is not None else cfg.temperature,
        'stream': False
    }
    # max_tokens：0 表示不限制（不下发），None 用配置默认值，正整数显式限定
    # 【输出上限适配】请求值按模型已知/已学习输出上限钳制，防大值撞 8k 上限直接 400
    _mt_limit = get_output_limit(base, model)
    if max_tokens == 0:
        pass  # 不下发，让模型用自身默认输出上限
    elif max_tokens:
        payload['max_tokens'] = min(max_tokens, _mt_limit) if _mt_limit else max_tokens
    else:
        payload['max_tokens'] = min(cfg.max_tokens, _mt_limit) if _mt_limit else cfg.max_tokens

    last_error = None
    for attempt in range(1 + retry_count):
        try:
            resp = requests.post(f'{base}/chat/completions',
                headers=build_auth_headers(cfg.api_key),
                json=payload, timeout=timeout)

            # 检查 HTTP 状态：5xx 可重试，4xx 为客户端错误不重试
            if resp.status_code >= 500:
                last_error = f'LLM 服务错误 (HTTP {resp.status_code})'
                if attempt < retry_count:
                    import time as _time
                    _time.sleep(1.5 * (attempt + 1))
                    continue
                return None, last_error
            if resp.status_code >= 400:
                try:
                    err_body = resp.json()
                    err_msg = (err_body.get('error') or {}).get('message') or str(err_body)
                except Exception:
                    err_msg = resp.text[:200]
                last_error = f'LLM 请求被拒绝 (HTTP {resp.status_code}): {err_msg}'
                return None, last_error

            result = resp.json()
            if 'choices' in result and len(result['choices']) > 0:
                return result['choices'][0]['message']['content'], None
            last_error = f'LLM 返回格式异常: {str(result)[:200]}'
            if attempt < retry_count:
                import time as _time
                _time.sleep(1.5 * (attempt + 1))
                continue
            return None, last_error

        except requests.exceptions.Timeout:
            last_error = f'LLM 请求超时（第{attempt+1}次）'
            if attempt < retry_count:
                import time as _time
                _time.sleep(1.5 * (attempt + 1))
                continue
            return None, f'LLM 请求超时，已重试{retry_count}次仍失败'
        except requests.exceptions.ConnectionError:
            last_error = f'LLM 连接失败（第{attempt+1}次）'
            if attempt < retry_count:
                import time as _time
                _time.sleep(2 * (attempt + 1))
                continue
            return None, f'LLM 连接失败，请检查 API 地址配置'
        except json.JSONDecodeError as e:
            last_error = f'LLM 返回内容为空或非JSON格式: {str(e)}'
            if attempt < retry_count:
                import time as _time
                _time.sleep(1.5 * (attempt + 1))
                continue
            return None, f'LLM 返回内容为空，已重试{retry_count}次仍失败，请检查模型是否可用或稍候重试'
        except Exception as e:
            last_error = str(e)
            if attempt < retry_count:
                import time as _time
                _time.sleep(1.5 * (attempt + 1))
                continue
            return None, str(e)

    return None, last_error or 'LLM 调用失败（未知错误）'

def _get_skill_prompts(skill_pack_ids, prompt_keys, max_per_prompt=1500, mode='agent'):
    """从技能包提取指定 prompt_keys 的提示词（agent 协同模式：所有匹配 prompt 都注入）。
    mode='agent'（默认）：所有匹配 prompt 全量注入（不截断），确保禁词清单/行文铁律等关键内容完整传达给 AI。
    mode='single'：每包只取第一个匹配 prompt（兼容旧行为）。
    """
    if not skill_pack_ids:
        return ''
    try:
        packs = SkillPack.query.filter(SkillPack.id.in_(skill_pack_ids)).all()
    except Exception as e:
        # 隐患2修复：出错时记日志而非静默返回空串，便于排查"技能包未生效"问题
        try:
            app.logger.error(f'_get_skill_prompts 查询技能包失败(ids={skill_pack_ids}): {e}')
        except Exception:
            pass
        return ''
    if not packs:
        # 用户选了技能包但查不到（可能被删除），记日志提示
        try:
            app.logger.warning(f'_get_skill_prompts 未找到匹配技能包(ids={skill_pack_ids})，可能已被删除')
        except Exception:
            pass
        return ''
    notes = []
    for pack in packs:
        try:
            prompts = json.loads(pack.prompts_json) if pack.prompts_json else {}
        except Exception as e:
            try:
                app.logger.warning(f'_get_skill_prompts 技能包[{pack.name}] prompts_json 解析失败: {e}')
            except Exception:
                pass
            prompts = {}
        matched = [(k, prompts[k]) for k in prompt_keys if k in prompts and prompts[k]]
        if not matched:
            continue
        if mode == 'single':
            # 兼容旧模式：只取第一个
            p = matched[0][1][:max_per_prompt]
            notes.append(f'【{pack.name}】\n{p}')
        else:
            # agent 模式：所有匹配 prompt 全量注入（不截断摘要）
            # 隐患1修复：原先只第一个全量、其余截到400字会丢失禁词清单等关键内容；
            # 现 token 不再受限，全部全量注入，确保创作方法论完整传达
            parts = []
            for k, p in matched:
                parts.append(f'[{k}]\n{p}')
            notes.append(f'【{pack.name}】\n' + '\n'.join(parts))
    return '\n\n'.join(notes)

def _get_skill_prompts_by_category(skill_pack_ids, category, prompt_keys=None, mode='agent',
                                   book_genre=None):
    """按类别过滤技能包后提取提示词（三类无污染隔离的核心调度函数）。
    - category='master': 只查构思类，注入大纲/规划阶段
    - category='style': 只查文风类，注入正文生成阶段（按 priority 排序）
    - category='review': 只查审查类，注入审校阶段（去AI味/一致性）
    prompt_keys: 可选，指定提取哪些 key；不指定则提取该包全部 prompts。
    老技能包无 category 字段时默认按 'master' 处理（兼容）。
    book_genre: 若传入，则对【文风类】且 pack.genre_target 非空的包进行题材匹配：
        - 若 pack.genre_target != book.genre，则跳过，不注入（避免跨题材文风污染）。"""
    if not skill_pack_ids:
        return ''
    try:
        packs = SkillPack.query.filter(SkillPack.id.in_(skill_pack_ids)).all()
    except Exception as e:
        try:
            app.logger.error(f'_get_skill_prompts_by_category 查询失败(ids={skill_pack_ids}, cat={category}): {e}')
        except Exception:
            pass
        return ''
    # 按 category 过滤（老包无 category 视为 master）
    filtered = [p for p in packs if (p.category or 'master') == category]
    if not filtered:
        return ''
    # 【文风类】题材过滤：若 pack.genre_target 指定了适用题材，且传入了 book_genre，不一致则跳过
    if category == 'style' and book_genre:
        def _style_match(p):
            target = (p.genre_target or '').strip()
            # 不指定 = 全题材通用
            if not target:
                return True
            return target == book_genre
        _before = len(filtered)
        filtered = [p for p in filtered if _style_match(p)]
        _after = len(filtered)
        # 【fix2】有包被 genre_target 过滤掉时记 WARNING，便于排查"勾了文风但没生效"（否则就是静默失败）
        if _before != _after:
            try:
                _skipped = [f"{p.name}(target={p.genre_target or ''})" for p in packs if (p.category or 'master') == category and not _style_match(p)]
                app.logger.warning(
                    f'[skill_pack] 文风包题材不匹配已跳过 ids={skill_pack_ids} book_genre={book_genre} '
                    f'通过={_after}/{_before} 跳过={_skipped}'
                )
            except Exception:
                pass
        if not filtered:
            return ''
    # 文风类按 priority 排序（数字小的先注入）
    if category == 'style':
        filtered.sort(key=lambda p: p.priority if p.priority is not None else 100)
    notes = []
    for pack in filtered:
        try:
            prompts = json.loads(pack.prompts_json) if pack.prompts_json else {}
        except Exception:
            continue
        if not prompts:
            continue
        # 若指定 prompt_keys 则按 key 过滤，否则取全部
        if prompt_keys:
            matched = [(k, prompts[k]) for k in prompt_keys if k in prompts and prompts[k]]
        else:
            matched = [(k, v) for k, v in prompts.items() if v]
        if not matched:
            continue
        if mode == 'single':
            p = matched[0][1][:1500]
            notes.append(f'【{pack.name}】\n{p}')
        else:
            parts = [f'[{k}]\n{p}' for k, p in matched]
            notes.append(f'【{pack.name}】\n' + '\n'.join(parts))
    return '\n\n'.join(notes)

def _get_enabled_style_pack(book):
    """正文阶段启用的文风技能包。实现 & 优先级如下：
    1) 先从 book.style_skill_ids 取用户勾选的文风包 ID 列表（_resolve_skill_ids_by_category）
    2) 查询这些 ID 对应的 SkillPack，文风类按 genre_target 匹配 book.genre（genre_target 为空视为通用）
    3) 按 priority 升序排序（小的优先），取全部命中包 prompts 拼接后返回
    返回空串表示没有匹配的文风包。"""
    if not book:
        return ''
    style_ids = _resolve_skill_ids_by_category(book, 'style')
    if not style_ids:
        return ''
    return _get_skill_prompts_by_category(
        style_ids,
        'style',
        mode='agent',
        book_genre=getattr(book, 'genre', None),
    )

def _resolve_skill_ids_by_category(book, category):
    """从 Book 表的三类字段中取出对应类别的技能包ID列表。
    - category='master' -> book.master_skill_ids
    - category='style'   -> book.style_skill_ids
    - category='review'  -> book.review_skill_ids
    兼容老数据：若三类字段全空，回退到老的 skill_pack_ids（在 metadata 或请求参数中）。"""
    if not book:
        return []
    field_map = {'master': 'master_skill_ids', 'style': 'style_skill_ids', 'review': 'review_skill_ids'}
    field_name = field_map.get(category, 'master_skill_ids')
    raw = getattr(book, field_name, None) or '[]'
    try:
        ids = json.loads(raw) if isinstance(raw, str) else raw
        if isinstance(ids, list) and ids:
            return ids
    except Exception:
        pass
    return []

def _split_legacy_skill_ids_to_categories(book):
    """老数据迁移：将老的 skill_pack_ids 按 category 自动分流到3个新字段。
    在 book 首次访问时触发（若3个新字段全空且老字段有值）。
    幂等：已分流过则跳过。"""
    if not book:
        return
    # 检查是否已分流过（任一新字段非空即视为已分流）
    try:
        m = json.loads(book.master_skill_ids or '[]') if book.master_skill_ids else []
        s = json.loads(book.style_skill_ids or '[]') if book.style_skill_ids else []
        r = json.loads(book.review_skill_ids or '[]') if book.review_skill_ids else []
        if m or s or r:
            return  # 已分流
    except Exception:
        return
    # 从 metadata 取老的 skill_pack_ids
    try:
        meta = json.loads(book.metadata_json or '{}') if book.metadata_json else {}
    except Exception:
        meta = {}
    legacy_ids = meta.get('skill_pack_ids', [])
    if not legacy_ids:
        return
    # 查询每个包的 category 分流
    try:
        packs = SkillPack.query.filter(SkillPack.id.in_(legacy_ids)).all()
    except Exception:
        return
    master_ids, style_ids, review_ids = [], [], []
    for p in packs:
        cat = p.category or 'master'
        if cat == 'style':
            style_ids.append(p.id)
        elif cat == 'review':
            review_ids.append(p.id)
        else:
            master_ids.append(p.id)
    book.master_skill_ids = json.dumps(master_ids, ensure_ascii=False)
    book.style_skill_ids = json.dumps(style_ids, ensure_ascii=False)
    book.review_skill_ids = json.dumps(review_ids, ensure_ascii=False)
    try:
        db.session.commit()
    except Exception:
        db.session.rollback()

# ==== 大纲工作流：五幕式总纲 + 卷纲滚动生成 ====
@app.route('/api/books/<book_id>/ai-outline-master', methods=['POST'])
def ai_outline_master(book_id):
    """生成五幕式总纲：控制全书大体走向，写入 plot_design"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Not found'}), 404
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    data = request.json or {}
    # 【P1-5修复】总纲生成属于总创作阶段：从 book.master_skill_ids 读取（构思类）
    try:
        _split_legacy_skill_ids_to_categories(book)
    except Exception:
        pass
    skill_pack_ids = _resolve_skill_ids_by_category(book, 'master')
    if not skill_pack_ids:
        skill_pack_ids = data.get('skill_pack_ids', [])
    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['master_outline', 'tomato_outline', 'volume_breakdown'])

    # ========== 卷数权威口径：用户创建小说时的 total_volumes（优先级：前端显式传 > Book表 > BookBible表 > 默认） ==========
    # 这是用户最核心的创作参数，禁止被 300//50 这类粗暴算法覆盖！
    cpv = _get_chapters_per_volume(bb, book)
    tv = _get_total_volumes(bb, book)
    total_chapters = data.get('total_chapters')
    chapters_per_volume = data.get('chapters_per_volume') or cpv
    # 支持用户直接指定卷数（最优先），否则强制使用用户建书时设定的 total_volumes，绝不再用 300//50 这种会踩用户设定的默认算法
    user_volume_count = data.get('volume_count')
    if user_volume_count and int(user_volume_count) >= 1:
        volume_count = int(user_volume_count)
    else:
        volume_count = tv
    # 反推总章数供 prompt 使用（按每卷章数计算，保证 prompt 与卷数一致）
    if not total_chapters or int(total_chapters) < 1:
        total_chapters = volume_count * chapters_per_volume
    else:
        total_chapters = max(int(total_chapters), volume_count * chapters_per_volume)

    context_parts = []
    if bb.concept:
        context_parts.append(f'【构思】\n{bb.concept[:2000]}')
    if bb.key_rules:
        context_parts.append(f'【设定/规则】\n{bb.key_rules[:2000]}')
    if bb.worldbuilding:
        context_parts.append(f'【世界观】\n{bb.worldbuilding[:2000]}')
    if bb.character_profiles:
        context_parts.append(f'【人物档案】\n{bb.character_profiles[:2000]}')
    context = '\n\n'.join(context_parts) or '（暂无构思和设定，请基于书名和题材自由发挥）'

    system_prompt = f"""你是番茄小说金番作者级别的总纲设计师。
任务：为这本小说设计五幕式总纲，控制全书大体走向。

【核心创作参数·铁律·不可违反】
- 总卷数：{volume_count} 卷（本任务的第一硬约束：你输出的分卷数必须严格等于 {volume_count} 卷，多一卷少一卷都必须重写，严禁输出 5/6/8/10 卷等任何其他卷数）
- 每卷章数：约 {chapters_per_volume} 章，全书约 {total_chapters} 章
- 题材：{_get_genre_label(book, bb)}
{f'- 风格流派：{_get_novel_styles_text(bb, book)}' if _get_novel_styles_text(bb, book) else ''}

【五幕模型·按卷数分配（总卷 {volume_count} 卷自动映射到5幕比例）】
- 立身(前5%)：第1~{max(1, volume_count*5//100)}卷，底层→入门，觉醒金手指+首打脸+建立认知
- 立足(5-25%)：第{max(2, volume_count*5//100+1)}~{volume_count*25//100}卷，新人→站稳，配角登场+世界观展开+小闭环
- 立势(25-50%)：第{volume_count*25//100+1}~{volume_count*50//100}卷，小角色→有分量，大舞台+强对手+团队建立
- 立威(50-75%)：第{volume_count*50//100+1}~{volume_count*75//100}卷，有分量→威名，组织级冲突+感情推进+信念考验
- 立命(75-100%)：第{volume_count*75//100+1}~{volume_count}卷，威名→蜕变，终极挑战+伏笔收束+续作种子

【输出要求】
全书约 {total_chapters} 章，分 {volume_count} 卷（每卷约 {chapters_per_volume} 章）。
必须为每一卷（1 到 {volume_count} 共 {volume_count} 卷）都输出：卷号与卷名、所属幕、本卷核心目标（一句话）、主要冲突、关键转折点（2-3个）、卷尾高潮与悬念。
只输出总纲文本，不要输出各卷的详细情节节点（详细节点在卷纲滚动生成阶段产生）。
{_cultivation_dimension_hint('plot_design', book, bb)}
{skill_note}"""

    user_prompt = f"""书名：{book.title}

【核心创作参数·再次强调·若违反直接作废】
- 总卷数必须严格 {volume_count} 卷，不多不少。若你产出的卷数不是 {volume_count} 卷，立刻自我重写。

{_build_core_params_block(bb, book)}

已有设定：
{context}

请生成完整的 {volume_count} 卷五幕式总纲（不要偷懒写成 5 卷或 10 卷，必须正好 {volume_count} 卷，每卷都要有条目）。"""

    content, err = _call_llm(
        [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
        max_tokens=0, temperature=0.7
    )
    if err:
        return jsonify({'error': err}), 500

    bb.plot_design = content
    db.session.commit()
    return jsonify({'master_outline': content, 'volume_count': volume_count})

def _sync_foreshadowings_to_volumes(bb):
    """【P0修复】把各卷 timeline 中的 foreshadow_new/foreshadowing/foreshadow_recycle
    汇总到 bb.foreshadowing_volumes（按卷结构化），打通前期伏笔到写作阶段的防遗忘链路。
    在 ai_outline_volume / ai_extract_volumes_from_outline / ai_import_plot_outline 写入 timeline 后调用。"""
    if not bb or not bb.timeline:
        return
    try:
        arr = json.loads(bb.timeline)
        if not isinstance(arr, list):
            return
        vol_foreshadows = []
        for v in arr:
            if not isinstance(v, dict):
                continue
            v_idx = v.get('volume_index') or _extract_volume_index(v.get('volume', v.get('volume_id', ''))) or 0
            new_fs = v.get('foreshadow_new') or v.get('foreshadowing') or []
            recycle_fs = v.get('foreshadow_recycle') or []
            if not new_fs and not recycle_fs:
                continue
            # 统一为字符串列表
            def _norm(items):
                out = []
                if isinstance(items, list):
                    for it in items:
                        if isinstance(it, str):
                            out.append(it.strip())
                        elif isinstance(it, dict):
                            out.append((it.get('name') or it.get('title') or '') + '：' + (it.get('content') or it.get('desc') or ''))
                        elif it is not None:
                            out.append(str(it).strip())
                elif isinstance(items, str) and items.strip():
                    out.append(items.strip())
                return [x for x in out if x]
            new_list = _norm(new_fs)
            recycle_list = _norm(recycle_fs)
            if new_list or recycle_list:
                vol_foreshadows.append({
                    'volume_index': int(v_idx) if v_idx else 0,
                    'volume': v.get('volume', v.get('volume_title', '')),
                    'foreshadow_new': new_list,
                    'foreshadow_recycle': recycle_list,
                })
        bb.foreshadowing_volumes = json.dumps(vol_foreshadows, ensure_ascii=False, indent=2) if vol_foreshadows else (bb.foreshadowing_volumes or '')
    except (json.JSONDecodeError, ValueError, TypeError):
        pass

def _calc_start_chapter_fallback(volume_index, chapters_per_volume, existing_volumes):
    """【P1修复】start_chapter 解析失败的健壮回退：
    按 volume_index * chapters_per_volume + 1 估算，避免静默回退为1导致章号断裂。"""
    try:
        cpc = int(chapters_per_volume) if chapters_per_volume else 50
    except (ValueError, TypeError):
        cpc = 50
    vi = int(volume_index) if volume_index else 1
    # 若已有卷数据，取已完成卷数（卷号 < 当前卷）* 每卷章数 + 1
    if existing_volumes and isinstance(existing_volumes, list):
        prev_count = sum(1 for v in existing_volumes
                         if isinstance(v, dict) and
                         (int(v.get('volume_index', 0) or _extract_volume_index(v.get('volume', v.get('volume_id', '0'))) or 0) < vi))
        return prev_count * cpc + 1
    return (vi - 1) * cpc + 1

def _extract_volume_index(text):
    """从卷名/卷ID中提取卷号数字，如'第3卷'/'卷三'/'卷二十三'/'Volume 2' → 3/3/23/2
    支持阿拉伯数字与复合中文数字（十一/二十三/一百零五/壹貳叁等）。"""
    if not text:
        return 0
    s = str(text)
    import re as _re
    # 阿拉伯数字优先
    m = _re.search(r'(\d+)', s)
    if m:
        return int(m.group(1))
    # 复合中文数字解析（支持 一~九十九、百、零、大写壹貳叁）
    cn_digits = {
        '零': 0, '〇': 0,
        '一': 1, '壹': 1, '乙': 1,
        '二': 2, '贰': 2, '貳': 2, '两': 2,
        '三': 3, '叁': 3, '參': 3,
        '四': 4, '肆': 4,
        '五': 5, '伍': 5,
        '六': 6, '陆': 6, '陸': 6,
        '七': 7, '柒': 7, '漆': 7,
        '八': 8, '捌': 8,
        '九': 9, '玖': 9,
    }
    cn_units = {'十': 10, '拾': 10, '百': 100, '佰': 100, '千': 1000, '仟': 1000}
    # 提取连续的中文数字片段（含单位）
    cn_str_match = _re.search(r'[零〇一二贰貳两三叁參四肆五伍六陆陸七柒漆八捌九玖十拾百佰千仟]+', s)
    if cn_str_match:
        cn_str = cn_str_match.group()
        total = 0
        current = 0
        for ch in cn_str:
            if ch in cn_digits:
                current = cn_digits[ch]
            elif ch in cn_units:
                unit = cn_units[ch]
                if current == 0:
                    current = 1 if unit >= 10 else 0
                total += current * unit
                current = 0
            else:
                current = 0
        total += current
        if total > 0:
            return total
    return 0

def _split_volume_title(raw):
    """从卷标题行剩余内容中拆分出 (纯卷名, 剧情描述剩余)。

    导入剧情大纲时，正则会把'第1卷'之后的整行内容捕获为标题行，例如：
      '《凡骨矿奴》：姜辰在古神胸腔矿场受尽欺压…节点：修炼进展…'
    这里只需提取书名号内的'凡骨矿奴'作为卷名，其余剧情描述作为 main_plot 回填。

    输入: '《凡骨矿奴》：姜辰在古神胸腔矿场…节点：…'
    输出: ('凡骨矿奴', '姜辰在古神胸腔矿场…节点：…')

    输入: '凡骨矿奴'
    输出: ('凡骨矿奴', '')
    """
    if not raw:
        return '', ''
    import re as _re
    s = raw.strip()
    # 优先书名号《》【】[] 内内容，剩余作为剧情描述
    bm = _re.match(r'^[《【〖［\[](.+?)[》】〗］\]][：:．.、\s\-—–]*([\s\S]*)$', s)
    if bm:
        return bm.group(1).strip(), bm.group(2).strip()
    # 冒号分隔：前段作卷名（短），后段作剧情描述
    cm = _re.split(r'[：:]', s, maxsplit=1)
    if len(cm) == 2 and cm[0].strip() and len(cm[0].strip()) <= 15:
        return cm[0].strip(), cm[1].strip()
    # 破折号分隔
    dm = _re.split(r'[—–]', s, maxsplit=1)
    if len(dm) == 2 and dm[0].strip() and len(dm[0].strip()) <= 15:
        return dm[0].strip(), dm[1].strip()
    # 无分隔符：较短当卷名，较长当剧情描述（卷名留空走默认）
    if len(s) <= 15:
        return s, ''
    return '', s

@app.route('/api/books/<book_id>/ai-outline-volume', methods=['POST'])
def ai_outline_volume(book_id):
    """生成单卷详细大纲（滚动生成）：基于总纲+已完成章节，写入 timeline"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Not found'}), 404
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()
    # 非强制：无总纲也能设计节点，基于已有卷剧情
    has_master = bool(bb.plot_design and bb.plot_design.strip())

    data = request.json or {}
    # 【P1-5修复】分卷大纲属于总创作阶段：从 book.master_skill_ids 读取（构思类）
    try:
        _split_legacy_skill_ids_to_categories(book)
    except Exception:
        pass
    skill_pack_ids = _resolve_skill_ids_by_category(book, 'master')
    if not skill_pack_ids:
        skill_pack_ids = data.get('skill_pack_ids', [])
    volume_index = data.get('volume_index', 1)
    volume_title = data.get('volume_title', f'第{volume_index}卷')
    chapters_per_volume = data.get('chapters_per_volume', 50)
    # 节点设计模式：已有卷剧情时只细化 nodes，不覆盖 main_plot 等字段
    node_only = bool(data.get('node_only', False))

    # ===== 【P0弊端2修复】已完成章节摘要：取上一卷结尾章节，而非全局最后5章 =====
    # 解析已有 timeline，获取上一卷（volume_index-1）的节点章节范围
    prev_volume_end_chapter = 0
    prev_volume_ending_hook = ''
    prev_volume_summary = ''
    existing_volumes_for_ctx = []
    if bb.timeline:
        try:
            parsed_tl = json.loads(bb.timeline)
            if isinstance(parsed_tl, list):
                existing_volumes_for_ctx = parsed_tl
                # 找上一卷
                for v in parsed_tl:
                    if not isinstance(v, dict):
                        continue
                    v_idx = v.get('volume_index') or _extract_volume_index(v.get('volume', v.get('volume_id', '')))
                    if str(v_idx) == str(volume_index - 1):
                        # 提取上一卷的卷尾钩子
                        prev_volume_ending_hook = v.get('ending') or v.get('ending_hook') or v.get('climax') or ''
                        prev_volume_summary = v.get('main_plot') or v.get('core_goal') or ''
                        # 提取上一卷 nodes 的最大结束章号
                        nodes = v.get('nodes') or []
                        for n in nodes:
                            if isinstance(n, dict):
                                ch_range = str(n.get('chapters', ''))
                                # 解析 "1-10" 或 "50" 格式
                                nums = re.findall(r'\d+', ch_range)
                                if nums:
                                    end_n = int(nums[-1])
                                    if end_n > prev_volume_end_chapter:
                                        prev_volume_end_chapter = end_n
                        break
        except (json.JSONDecodeError, ValueError, TypeError):
            pass

    # 节点设计模式：取本卷已有的卷剧情作为节点拆分依据（必须在 existing_volumes_for_ctx 解析之后）
    current_vol_existing = None
    if node_only:
        for v in existing_volumes_for_ctx:
            if not isinstance(v, dict):
                continue
            v_idx = v.get('volume_index') or _extract_volume_index(v.get('volume', v.get('volume_id', '')))
            if str(v_idx) == str(volume_index):
                current_vol_existing = v
                break

    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['volume_breakdown', 'chapter_plan', 'tomato_outline'])

    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()

    # 上一卷结尾章节正文（用于卷间衔接，取最后2章各300字）
    prev_volume_end_summary = ''
    if chapters and volume_index > 1:
        # 优先取上一卷结尾章节；若无法定位，回退取全书最后3章
        if prev_volume_end_chapter > 0:
            prev_end_chapters = [c for c in chapters if c.order_index <= prev_volume_end_chapter][-2:]
        else:
            prev_end_chapters = chapters[-3:]
        prev_volume_end_summary = '\n'.join([f'第{c.order_index}章 {c.title or ""}：{(c.content or "")[-300:]}' for c in prev_end_chapters])

    # 【P1修复】existing_timeline 分级注入：上一卷完整 + 其余卷主线+节点摘要，
    # 让节点设计能感知全部已有卷的节点，保证全书节点剧情连贯
    existing_timeline = ''
    if existing_volumes_for_ctx:
        tl_parts = []
        for v in existing_volumes_for_ctx:
            if not isinstance(v, dict):
                continue
            v_idx = v.get('volume_index') or _extract_volume_index(v.get('volume', v.get('volume_id', ''))) or 0
            try:
                v_idx_int = int(v_idx)
            except (ValueError, TypeError):
                v_idx_int = 0
            v_nodes = v.get('nodes') or []
            # 上一卷（volume_index-1）注入完整信息（含完整节点）
            if v_idx_int == volume_index - 1:
                tl_parts.append(f'▼ [上一卷·完整] 第{v_idx}卷“{v.get("volume", v.get("volume_title", ""))}”\n'
                                + f'  主线：{(v.get("main_plot") or v.get("core_goal") or "")[:400]}\n'
                                + f'  核心冲突：{str(v.get("core_conflict", ""))[:200]}\n'
                                + f'  关键转折：{", ".join(v.get("key_events", []) or v.get("turning_points", []))[:200]}\n'
                                + f'  卷尾钩子：{str(v.get("ending") or v.get("ending_hook") or v.get("climax") or "")[:200]}\n'
                                + f'  情节节点：' + ' | '.join([f'{n.get("chapters","")} {n.get("title","")}' for n in v_nodes[:8]]))
            elif v_nodes:
                # 其余已有节点的卷：注入主线 + 节点摘要（保证全书节点连贯）
                main = (v.get('main_plot') or v.get('core_goal') or '')[:120]
                hook = v.get('ending') or v.get('ending_hook') or v.get('climax') or ''
                nodes_brief = ' | '.join([f'{n.get("chapters","")}{n.get("title","")}' for n in v_nodes[:6]])
                tl_parts.append(f'· 第{v_idx}卷“{v.get("volume", v.get("volume_title", ""))}”主线：{main}；卷尾钩子：{str(hook)[:80]}；节点：{nodes_brief[:200]}')
            else:
                # 无节点的卷：仅主线+钩子
                hook = v.get('ending') or v.get('ending_hook') or v.get('climax') or ''
                main = (v.get('main_plot') or v.get('core_goal') or '')[:80]
                tl_parts.append(f'· 第{v_idx}卷“{v.get("volume", v.get("volume_title", ""))}”主线：{main}；卷尾钩子：{str(hook)[:100]}')
        existing_timeline = '\n'.join(tl_parts)[:4500]
    else:
        existing_timeline = (bb.timeline or '')[:2000]
    master_outline = (bb.plot_design or '')[:3000]
    # agent 协同：补充世界观+人物档案，让卷内情节节点能落地到具体世界观规则和角色互动
    worldbuilding_ctx = (bb.worldbuilding or '')[:1500]
    characters_ctx = (bb.character_profiles or '')[:1500]
    key_rules_ctx = (bb.key_rules or '')[:1000]

    # ===== 【P1弊端7修复】五幕模型对齐：确定本卷对应的幕 =====
    act_mapping = {1: '立身', 2: '立足', 3: '立势', 4: '立威', 5: '立命'}
    # P1-1修复：total_volumes 必须用全书规划卷数（book.total_volumes），
    # 而非已生成卷数+1，否则五幕比例映射分母错误，导致本卷被错配到错误的幕
    total_volumes = _get_total_volumes(bb, book)
    # 简单映射：5卷对应5幕；卷数不固定时按比例分配
    if total_volumes <= 5:
        current_act = act_mapping.get(volume_index, '立身')
    else:
        # 超过5卷时，按进度比例映射到5幕
        act_idx = min(5, max(1, int((volume_index / total_volumes) * 5) + 1))
        current_act = act_mapping.get(act_idx, '立身')
    act_descriptions = {
        '立身': '主角登场、金手指获得、确立生存基础（1-5%）',
        '立足': '主角站稳脚跟、初露锋芒、建立基本人际网（5-25%）',
        '立势': '主角势力扩张、主要冲突激化、BOSS浮出（25-50%）',
        '立威': '主角与BOSS正面对抗、实力跃升、打脸高潮（50-75%）',
        '立命': '终局决战、伏笔回收、世界观全貌揭示（75-100%）',
    }

    # ===== 【P1弊端9修复】本卷起始章号 = 上一卷结束章号 + 1 =====
    # 【P1强化】解析失败（prev_volume_end_chapter==0 但 volume_index>1）时健壮回退，
    # 避免静默回退为1导致章号断裂
    if prev_volume_end_chapter > 0:
        start_chapter = prev_volume_end_chapter + 1
    elif volume_index > 1:
        start_chapter = _calc_start_chapter_fallback(volume_index, chapters_per_volume, existing_volumes_for_ctx)
        # 记录告警供调试（写入日志，不打断流程）
        app.logger.warning(f'ai_outline_volume: 卷{volume_index} start_chapter 解析失败，回退为 {start_chapter}')
    else:
        start_chapter = 1

    # ===== 【P0弊端1修复】卷间衔接约束提示词 =====
    cohesion_constraint = ''
    if volume_index > 1 and prev_volume_ending_hook:
        cohesion_constraint = f"""
【卷间衔接铁律】（本卷为第{volume_index}卷，必须严格承接第{volume_index-1}卷）
- 上一卷卷尾钩子：{prev_volume_ending_hook}
- 上一卷核心主线：{prev_volume_summary[:200] if prev_volume_summary else '（无）'}
- 本卷开头必须承接上一卷卷尾钩子的悬念/危机，不得凭空开启新场景
- 本卷第一个情节节点的起始章号必须为 {start_chapter}（上一卷结束于第{prev_volume_end_chapter}章）
- 本卷 nodes 的 chapters 字段必须从 {start_chapter} 开始连续编号，不得与上一卷重叠"""
    elif volume_index > 1:
        cohesion_constraint = f"""
【卷间衔接铁律】（本卷为第{volume_index}卷）
- 本卷开头必须承接上一卷的结尾场景与悬念，不得凭空跳跃
- 本卷第一个情节节点的起始章号必须为 {start_chapter}"""

    # ===== 【爽点设计系统】8种爽点类型 + 4层爽点周期 + 钩子系统 =====
    cool_system_prompt = """
【爽点设计系统】小说不能只走主线/打打杀杀，每个情节节点必须内置爽点。每节点必须从以下维度设计：

■ 八种爽点类型（cool_type 必须为其中之一）：
  ① 实力碾压爽：弱者挑衅→主角展示实力→对手/围观者震惊→主角嘴贱收尾
  ② 信息差爽：读者+主角共享秘密→其他角色不知情→揭晓时双重爽感
  ③ 扮猪吃虎爽：主角隐藏实力→对手轻视→暗中布局→一击翻盘
  ④ 荒诞反差爽：预期A→实际B→反差制造笑点和爽感
  ⑤ 打脸装逼爽：被嘲/被踩（≤1章）→立刻打脸→嘴炮补刀→围观震惊
  ⑥ 社会认同爽：主角做了某事→第三方震惊反应→放大主角厉害
  ⑦ 升级蜕变爽：突破瓶颈→实力质变→标志性能力解锁→展示新力量
  ⑧ 守护爆发爽：重要之人遇险→平时怂的主角彻底暴走→碾压式保护

■ 爽点结构（cool_structure 必须为其中之一）：
  · 先抑后扬：先压制主角→蓄力→爆发碾压（适合中爽/大爽）
  · 直接碾压：开局即压制→不给喘息→碾压收尾（适合实力碾压爽）
  · 默默装完逼：低调做事→他人发现→事后佩服（适合扮猪吃虎/信息差爽）

■ 衬托方式（cool_contrast 必须为其中之一）：
  · 旁人震惊：围观者实时倒吸凉气、议论纷纷
  · 不敢置信：对手/熟人难以接受、反复确认
  · 事后佩服：事件平息后他人复盘主角的手段

■ 四层爽点周期（每卷50章参考，按本卷章数等比缩放）：
  · 微爽：每章内数次，每400字≥1个，不能连续3次同类型，高潮段密度加倍
    微爽库：①一句话打脸 ②嘴炮/吐槽 ③反差反应 ④小信息差 ⑤荒诞细节 ⑥拟声词即时爽感
  · 小爽：每5章一次，小打脸/小出气/小收获
  · 中爽：每17章一次，大反转/大翻身/配角高光
  · 大爽：每50章一次，全场震惊/彻底翻盘/量级跃迁（卷末必有）

■ 节奏心跳公式（50章卷内参考）：
  第1-8章：铺垫+小爽 → 9-16章：蓄力+小爽 → 17-25章：中爽爆发
  → 26-33章：过渡+小爽 → 34-42章：蓄力+小爽 → 43-50章：大爽+卷末钩子

■ 爽点弧线（全书阶段，按当前进度匹配）：
  1-50章：即时打脸、荒诞反差（解气、好笑）
  50-150章：升级展示、扮猪吃虎（优越感）
  150章后：信息差博弈、守护爆发（智力/羁绊感）

■ 钩子系统（每个节点至少带1个钩子，近5章不重复）：
  七种钩子：身份揭露 / 新危机 / 荒诞反转 / 悬念 / 角色危机 / 能力突破 / 世界异常
  优先用升级钩子：反常识断章 / 情绪断章 / 信息炸弹 / 视角断章 / 系统异常
  进化路线：前期(1-50)悬念型 → 中期(50-150)期待型 → 后期(150+)情感型
"""

    if node_only and current_vol_existing:
        # ===== 节点设计模式（升级：按每个 main_event 拆 5-10 个子节点事件） =====
        existing_summary = current_vol_existing.get('summary') or current_vol_existing.get('main_plot') or ''
        existing_main_plot = current_vol_existing.get('main_plot') or current_vol_existing.get('core_goal') or ''
        # 优先级：新结构 main_events[]（8-12个主要剧情事件，含 6 要素+estimated_chapters，无 chapters）
        #   → 旧 key_events/turning_points 数组 → 兜底旧结构空
        existing_main_events = current_vol_existing.get('main_events') or []
        existing_key_events = current_vol_existing.get('key_events') or current_vol_existing.get('turning_points') or []
        # 卷级 6 要素（喂给 LLM 作为子节点的区间锚）
        vol_characters = current_vol_existing.get('characters') or ''
        vol_timeline_anchor = current_vol_existing.get('timeline_anchor') or ''
        vol_location = current_vol_existing.get('location') or ''
        vol_realm_change = current_vol_existing.get('realm_change') or ''
        vol_age_change = current_vol_existing.get('age_change') or ''
        # 如果用户还没生成过 main_events（旧 timeline），从关键事件数组里推断出 main_events 的框架
        if not isinstance(existing_main_events, list) or len(existing_main_events) == 0:
            _ev_idx = 0
            _tmp = []
            for ke in (existing_key_events or []):
                _ev_idx += 1
                _tmp.append({
                    'index': _ev_idx, 'title': str(ke),
                    'estimated_chapters': max(1, chapters_per_volume // max(1, len(existing_key_events or [_tmp]))),
                    'summary': str(ke),
                    'characters': '', 'events': str(ke), 'time': '', 'location': '',
                    'realm_change': '', 'age_change': '',
                    'bury': '', 'payoff': '',
                })
            existing_main_events = _tmp
        existing_climax = current_vol_existing.get('climax') or ''
        existing_ending = current_vol_existing.get('ending') or current_vol_existing.get('ending_hook') or ''
        existing_foreshadowing = current_vol_existing.get('foreshadowing') or []
        existing_core_conflict = current_vol_existing.get('core_conflict') or ''
        existing_emotion = current_vol_existing.get('emotion_driver') or ''

        # ---------- 节点阶段：按 estimated_chapters 把 main_events 顺序拆成章程序列 ----------
        # 先做一遍事件章区间的分配（按卷起始章 start_chapter + 事件 estimated_chapters 累加连续分配），
        # 再在 system prompt 里显式告诉 LLM 每个 main_event 对应的 chapters 区间，
        # 这样 LLM 拆子节点时就能在区间内精确到每章，并保证节点 chapters 与父事件对齐。
        _evt_start = start_chapter
        _evt_end = start_chapter + chapters_per_volume - 1
        _evt_alloc = []  # list[(me_index, me_title, ch_start, ch_end)]
        _cur_ch = _evt_start
        for _me in existing_main_events:
            if not isinstance(_me, dict):
                continue
            try:
                _ec = int(_me.get('estimated_chapters') or 0)
                if _ec <= 0:
                    _ec = max(1, chapters_per_volume // max(1, len(existing_main_events)))
            except (TypeError, ValueError):
                _ec = max(1, chapters_per_volume // max(1, len(existing_main_events)))
            _me_start = _cur_ch
            _me_end = min(_evt_end, _cur_ch + _ec - 1)
            if _me_end < _me_start:
                _me_end = _me_start
            _evt_alloc.append((int(_me.get('index') or 0) or (len(_evt_alloc) + 1),
                               str(_me.get('title') or ''),
                               _me_start, _me_end,
                               _ec,
                               _me))
            _cur_ch = _me_end + 1
        # 若最后一个事件分配有缺口/溢出，拉齐到 _evt_end
        if _evt_alloc and _evt_alloc[-1][3] != _evt_end:
            *head, (idx, ttl, s, e, ec, raw) = _evt_alloc
            e = _evt_end
            _evt_alloc = head + [(idx, ttl, s, e, (e - s + 1), raw)]

        # 把 main_events 列表 + 已分派的章区间序列化成多行文本喂给 LLM
        _me_lines = []
        for _alloc in _evt_alloc:
            idx, ttl, s, e, ec, _me = _alloc
            _me_lines.append(
                f'  · 事件{idx}《{ttl}》（应分配章节：{s}-{e}，共 {e - s + 1} 章，支撑 ec={ec} 章）'
                f'\n    概要：{str(_me.get("summary",""))[:300]}'
                f'\n    ·人物：{_me.get("characters","")}'
                f'\n    ·事件：{_me.get("events","")}'
                f'\n    ·时间：{_me.get("time","")}'
                f'\n    ·地点：{_me.get("location","")}'
                f'\n    ·境界：{_me.get("realm_change","")}'
                f'\n    ·年龄：{_me.get("age_change","")}'
                + (f'\n    埋：{_me.get("bury","")}' if _me.get("bury") else '')
                + (f'\n    收：{_me.get("payoff","")}' if _me.get("payoff") else '')
            )
        main_events_block = '\n'.join(_me_lines) if _me_lines else '（本卷尚未按新结构生成主要剧情事件，将基于旧的关键事件设计子节点）'

                # 共享系统提示词（不含按事件拆分的细节，各事件共享）
        # 【P2修复】收紧单事件节点数：缩小单次 LLM 输出体量，避免超长输出导致 LLM 上游 504。
        # 原来"4-10/8 个节点 + 详尽6要素+埋收"的超长 JSON 单次推理过久，上游必然超时。
        # 压缩到 3-5 个，单次请求轻量，能稳定返回（后端按需补足总章节数，不影响覆盖）。
        _per_event_node_count = 3 if len(existing_main_events) >= 7 else 4
        _per_event_node_max = 5 if len(existing_main_events) >= 7 else 6
        shared_system_prompt = f"""你是番茄小说金番作者级别的情节节点设计师。
任务：为第 {volume_index} 卷“{volume_title}”的一个主要剧情事件（main_event）生成 {_per_event_node_count}-{_per_event_node_max} 个情节子节点（nodes）。
【输出范围铁律】只允许输出第 {volume_index} 卷的内容。禁止在输出中复述、罗列、或带入任何其他卷（含上一卷、后续卷）的大纲/节点/剧情概要——这些仅供你推理衔接时参考，绝不写进输出。nodes 数组的 chapters 必须严格落在本卷 {_evt_start}-{_evt_end} 章区间。
- 主要事件 ≤ 6 个 → 每个拆 5-10 个
- 主要事件 ≥ 7 个 → 每个拆 4-8 个
每个子节点 events 字段约支撑 2400 字正文。

【本卷 6 要素锚（所有子节点都要在这个大框架内推进，不能超界）】
  · 核心人物：{vol_characters}
  · 时间锚：{vol_timeline_anchor}
  · 地点路线：{vol_location}
  · 境界变化区间：{vol_realm_change}
  · 年龄变化区间：{vol_age_change}

【模式说明】本卷已有完整卷剧情，你的任务是为当前这个 main_event 生成 {_per_event_node_count}-{_per_event_node_max} 个 nodes 子节点。
- 不要修改本卷的 summary / main_plot / main_events / core_conflict / ending_hook 等卷级字段
- 只输出 nodes 数组（保留原卷级字段由后端合并）
- 各子节点之间必须剧情连贯：上一节点末尾自然衔接到下一节点开头
- 本卷第一个子节点必须承接上一卷卷尾钩子，本卷最后一个子节点必须埋下卷尾钩子承接本卷 ending_hook
- 所有子节点必须严格归属到本 main_event，不得脱离本主要剧情事件自创剧情
- 子节点 chapters 区间要**连续不重叠**，刚好覆盖本 main_event 被分派的章节区间（由 user prompt 中的【本次生成的 main_event 详情】指定）

【两层对应铁律·升级（事件→子节点 + 精确章节分派）】
  · 本 main_event → 展开 {_per_event_node_count}-{_per_event_node_max} 个 nodes 子节点事件
  · 每个子节点必须标注：它从属于哪个主要剧情事件（main_event_index = main_event.index）
  · 子节点覆盖的 chapters 区间见 user prompt 中的分配章节

【五幕模型对齐】本卷对应五幕中的“{current_act}”幕：{act_descriptions.get(current_act, '')}
节点设计必须服务于该幕的核心目标。

{cool_system_prompt}

【节点 6 要素铁律（每个节点必须齐全）】
  · characters：本节点核心出场人物（按权重排序，含小角色）
  · events：本节点核心推进（20-40字：谁在什么场景做了什么→关键后果，直接对应能写进正文的剧情句）
  · time：本节点时间锚（相对卷级 + 绝对章号，如"卷首第3周·第4章当晚三更"）
  · location：本节点发生地点（精确到场景，如"黑骨矿场·丙字采区·第7巷道口"）
  · realm_change：本节点结束时主角的境界/根基/金手指变化（如"淬骨三重巅峰，右腿骨纹半亮" / "境界不动，攒下突破契机"）
  · age_change：本节点结束时主角的年龄/时程变化（如"16岁1个月零9天" / "距开篇24天"）

【伏笔埋收标注铁律（绝对要填·精确到章）】
每个子节点必须明确标注哪里埋了什么伏笔、后面哪一章回收：
  · bury = "第XX章埋下：XXX；预计回收：第YY章（第Z卷）" （本节点真的埋下了才写，没埋就空串）
  · payoff = "第XX章回收：前文第YY章埋下的XXX；效果：XXX" （本节点真的回收了才写，没收就空串）
要和归属的 main_event 的 bury/payoff 对齐，不要自相矛盾；跨卷的 payoff 要指明第X卷。

【输出格式】严格输出以下JSON（不要包裹在markdown代码块中，只输出 nodes 数组与必要元信息）：
{{
  "volume_index": {volume_index},
  "volume_title": "{volume_title}",
  "nodes": [
    {{
      "main_event_index": <对应main_event的index>,
      "index": <子节点序号>,
      "title": "节点标题（动宾结构，如：街市遭袭反杀三名劫修）",
      "chapters": "<起始章号>-<结束章号>",
      "type": "M（M主线/C角色/W世界观/D日常/F伏笔）",
      "characters": "本节点核心人物，按出场权重",
      "events": "本节点核心推进：谁在哪做了什么→关键后果（20-40字）",
      "time": "本节点时间锚（相对卷级+章号）",
      "location": "本节点精确场景地点",
      "realm_change": "本节点结束时境界/根基/金手指变化",
      "age_change": "本节点结束时主角年龄/时程变化",
      "summary": "本节点详细剧情概要（支撑 chapters 范围内章节数×2400字的容量，含起因→发展→高潮→收尾→钩子，不可太简）",
      "cool_type": "爽点类型（八选一）",
      "cool_structure": "爽点结构（先抑后扬/直接碾压/默默装完逼）",
      "cool_contrast": "衬托方式（旁人震惊/不敢置信/事后佩服）",
      "cool_level": "爽点层级（微爽/小爽/中爽/大爽）",
      "bury": "第XX章埋下：XXX；预计回收：第YY章（第Z卷）/空串",
      "payoff": "第XX章回收：第YY章埋下的XXX/空串",
      "hook": "本节点章尾钩子（七种钩子之一）"
    }}
  ]
}}
{cohesion_constraint}

【章型配额】M主线50%/C角色10%/W世界观10%/D日常20%/F伏笔10%
【小故事闭环】新事件→困难→金手指破局→暴露新信息→打脸收尾→钩子（1-2章一个子节点，正好可直接写正文）
本卷约 {chapters_per_volume} 章（约 {chapters_per_volume * 2400} 字），每个 main_event 拆 {_per_event_node_count}-{_per_event_node_max} 个子节点事件。
当前 main_event 的子节点 chapters 必须严格卡在 user prompt 中【本次生成的 main_event 详情】指定的章节区间内。
【节点容量铁律】每个子节点 summary 必须足够支撑其 chapters 范围的字数容量（按每章2400字估算），不得简略。
【节点连贯铁律】各节点 summary 末尾必须自然过渡到下一节点开头；本卷最后一个节点必须埋下并承接本卷 ending_hook：{existing_ending[:120]}

{skill_note}"""

        shared_user_prefix = f"""书名：{book.title}

{_build_core_params_block(bb, book)}

【五幕式总纲】
{master_outline if has_master else '（暂无）'}

【已有剧情】（含已生成卷纲，本卷节点需与之衔接）
{existing_timeline or '（暂无）'}

【本卷已有卷剧情】（节点拆分的唯一依据，必须严格符合）
- 卷名：第{volume_index}卷“{volume_title}”
- 总体剧情概要（summary）：{existing_summary or '（无）'}
- 主线推进路径（main_plot）：{existing_main_plot or '（无）'}
- 卷级6要素：人物={vol_characters}；时间={vol_timeline_anchor}；地点={vol_location}；境界={vol_realm_change}；年龄={vol_age_change}
- 核心冲突：{existing_core_conflict or '（无）'}
- 情感驱动：{existing_emotion or '（无）'}
- 高潮：{existing_climax or '（无）'}
- 结局/卷尾钩子：{existing_ending or '（无）'}
- 新埋伏笔：{', '.join(existing_foreshadowing) if existing_foreshadowing else '（无）'}

【世界观设定】
{worldbuilding_ctx or '（暂无）'}

【核心规则】
{key_rules_ctx or '（暂无）'}

【人物档案】
{characters_ctx or '（暂无）'}

【上一卷结尾章节正文】（卷间衔接依据，本卷第一个子节点必须承接）
{prev_volume_end_summary or '（本卷为第一卷，无前文）'}"""

        # 逐个 main_event 生成（改为线程池并发 + 信号量限流）
        # 根因修复：原实现逐事件**串行**调用 LLM，每个事件生成 4-10 个详细节点（单次 40-90s），
        # 有 8-12 个事件时串行总时长可达 5-15 分钟，远超前端 300s / Render 500s 请求超时→必然 504。
        # 并发后总时长≈单个事件耗时，根治超时。
        all_nodes = []
        event_errors = []
        # 【P0修复】_evt_alloc 为空时直接返回错误，避免静默保存空节点
        if not _evt_alloc:
            err = '该卷没有 main_events 或 key_events 可供拆分为子节点，请先在卷剧情中生成主要剧情事件'
            return jsonify({'error': err}), 400

        import threading
        from concurrent.futures import ThreadPoolExecutor, as_completed
        _sema = threading.Semaphore(2)  # 【P2】降为 2 路并发，降低同时打爆 LLM 上游的 504 风险

        def _gen_event_worker(_alloc):
            # 【P1修复】worker 线程无 Flask app context，_call_llm 内 AIConfig.get_active()
            # 访问数据库需要 app context，必须显式包裹，否则报 Working outside of application context
            with app.app_context():
                _idx2, _ttl2, _s2, _e2, _ec2, _me2 = _alloc
                _per_event_user = f"""{shared_user_prefix}

【本次生成的 main_event 详情】（请严格按此事件展开成 {_per_event_node_count}-{_per_event_node_max} 个子节点，子节点 chapters 必须严格卡在 {_s2}-{_e2} 区间内，合计 {_e2-_s2+1} 章）：
  · 事件{_idx2}《{_ttl2}》（分配章节：{_s2}-{_e2}，共 {_e2-_s2+1} 章，支撑 ec={_ec2} 章）
    概要：{str(_me2.get("summary",""))[:300]}
    ·人物：{_me2.get("characters","")}
    ·事件：{_me2.get("events","")}
    ·时间：{_me2.get("time","")}
    ·地点：{_me2.get("location","")}
    ·境界：{_me2.get("realm_change","")}
    ·年龄：{_me2.get("age_change","")}{f'''
    埋：{_me2.get("bury","")}''' if _me2.get("bury") else ''}{f'''
    收：{_me2.get("payoff","")}''' if _me2.get("payoff") else ''}

请为第 {volume_index} 卷的 main_event.{_idx2}《{_ttl2}》设计 {_per_event_node_count}-{_per_event_node_max} 个情节子节点事件（nodes），所有子节点的 chapters 必须严格卡在 {_s2}-{_e2} 区间内。"""

                with _sema:
                    import time as _time
                    # 按剩余总预算决定本事件可用时长：剩余越少，重试越少、单次超时越短
                    # 用 _fb_deadline（总截止前 120s）作为上限，给回退留足时间
                    _remain = max(5, int(_fb_deadline - _deadline_time.monotonic()))
                    _ev_retry = 3 if _remain > 180 else (1 if _remain > 90 else 0)
                    # 【P3修复】客户端超时别比上游先放弃：上游 504 有时是"其实在算、只是慢"的假象。
                    # 预算已提到 460s，单次调用给足 150s（慢但能成功的请求别被我们 80s 掐死）。
                    _ev_timeout = min(150, max(45, _remain // 2))
                    _content, _err = _call_llm(
                        [{'role': 'system', 'content': shared_system_prompt},
                         {'role': 'user', 'content': _per_event_user}],
                        # max_tokens=0：不下发上限字段，自动适配模型默认最大输出，减少 400/超时
                        max_tokens=0, temperature=0.65, retry_count=_ev_retry,
                        timeout=_ev_timeout,
                    )
                if _err:
                    return [], f'事件{_idx2}《{_ttl2}》生成失败: {_err}'
                _parsed, _json_err = _extract_json_from_llm(_content, expect='object')
                if _json_err:
                    return [], f'事件{_idx2}《{_ttl2}》JSON解析失败: {_json_err}'
                _event_nodes = _parsed.get('nodes', []) or []
                # 确保每个节点带 main_event_index
                for _n in _event_nodes:
                    if not _n.get('main_event_index'):
                        _n['main_event_index'] = _idx2
                return _event_nodes, None

        # 并发执行所有事件的节点生成
        # 【P1修复】整体 deadline：预留合并落库时间，剩余预算分给并发请求。
        # 即使个别事件单次超时+重试放大耗时，超过剩余预算的 worker 直接放弃（retry_count=0）。
        # 预算对齐：前端已把 aiOutlineVolume 超时提到 600s，Render 请求上限≈500s，
        # 故 worker 预算 460s + 合并落库 ≈ 470s，既不会超过 Render 500s 掐断，也不会被前端提前 abort。
        import time as _deadline_time
        _wall_deadline = _deadline_time.monotonic() + 460
        # 【P4修复】给回退预留时间：逐事件阶段最多用到距总截止前 120s，
        # 保证单次/轻量回退永远有足量预算可跑，不影响逐事件阶段的正常发挥。
        _fb_deadline = _wall_deadline - 120
        with ThreadPoolExecutor(max_workers=min(2, len(_evt_alloc))) as _ex:
            _futures = {_ex.submit(_gen_event_worker, _alloc): _alloc for _alloc in _evt_alloc}
            for _fut in as_completed(_futures):
                try:
                    _ev_nodes, _ev_err = _fut.result()
                except Exception as _e:
                    _ev_nodes, _ev_err = [], f'线程异常: {_e}'
                if _ev_err:
                    event_errors.append(_ev_err)
                elif _ev_nodes:
                    all_nodes.extend(_ev_nodes)

        # 按 chapters 起始章号排序
        def _sort_key(n):
            ch = str(n.get('chapters', '')).split('-')[0]
            try:
                return int(ch)
            except (ValueError, TypeError):
                return 9999
        all_nodes.sort(key=_sort_key)
        # 重排 index
        for _ni, _n in enumerate(all_nodes, 1):
            _n['index'] = _ni

        # ===== 回退阶段（按"越轻量越先跑"排序，最易成功的兜底优先拿到最新鲜预算）=====
        # 【P2-7修复】放宽触发条件：无论 event_errors 是否为空，只要 all_nodes 为空就触发回退
        # 1) 超轻量骨架回退：请求最精简、几乎必成，优先执行；单次全量请求最容易 504，留到最后。
        if not all_nodes:
            _light_system = f"""你是番茄小说金番作者级别的情节节点设计师。
任务：为第 {volume_index} 卷“{volume_title}”的每个 main_event 生成 1-2 个情节子节点（nodes），输出务必精简、宁短勿长。
【超轻量骨架模式】这是兜底回退，目标是"哪怕粗糙也要成功"：每个子节点 summary 用 30-60 字要点式一句话即可，禁止长篇，禁止展开细枝末节。

【本卷 6 要素锚】核心人物：{vol_characters} | 时间：{vol_timeline_anchor} | 地点：{vol_location} | 境界：{vol_realm_change} | 年龄：{vol_age_change}

【五幕模型对齐】本卷对应“{current_act}”幕：{act_descriptions.get(current_act, '')}

【节点 6 要素铁律】每个节点必须包含：characters / events / time / location / realm_change / age_change

【埋收标注】每个节点必须标注 bury（第XX章埋下）和 payoff（第XX章回收），无则空串。

【输出格式】严格输出 JSON 对象 {{"nodes": [...]}}，不要包裹 markdown 代码块。不要输出任何注释或说明文字。

{cohesion_constraint}

【章型配额】M主线50%/C角色10%/W世界观10%/D日常20%/F伏笔10%
【节点容量铁律】骨架模式下 summary 允许要点式（30-60字/节点），不必按章字数展开；chapters 仍须连续覆盖各自 main_event 的章节区间。"""

            _light_user = f"""书名：{book.title}
{_build_core_params_block(bb, book)}

【本卷已有卷剧情】
- 卷名：第{volume_index}卷“{volume_title}”
- 总体剧情概要：{existing_summary or '（无）'}
- 主线推进路径：{existing_main_plot or '（无）'}
- 核心冲突：{existing_core_conflict or '（无）'}
- 高潮：{existing_climax or '（无）'}
- 结局/卷尾钩子：{existing_ending or '（无）'}
- 新埋伏笔：{', '.join(existing_foreshadowing) if existing_foreshadowing else '（无）'}

【上一卷结尾】{prev_volume_end_summary or '本卷为第一卷，无前文'}

【所有 main_events】（请为每个事件生成 1-2 个骨架子节点，chapters 必须严格卡在各事件分配的区间内，全部节点合计覆盖 {_evt_start}-{_evt_end} 章，宁少勿缺）：
{main_events_block}

请为第 {volume_index} 卷的所有 main_events 各生成 1-2 个骨架情节子节点事件（nodes），宁短勿长，
只输出 JSON 对象 {{"nodes": [...]}}，不要输出任何其他文字。"""
            _l2_remain = max(5, int(_wall_deadline - _deadline_time.monotonic()))
            _l2_retry = 2 if _l2_remain > 60 else 0
            # 轻量回退：真正的兜底，必须小到几乎必成：输出 6K 上限、客户端给足 180s、可重试2次
            _l2_timeout = min(180, max(45, _l2_remain))
            _l2_content, _l2_err = _call_llm(
                [{'role': 'system', 'content': _light_system},
                 {'role': 'user', 'content': _light_user}],
                max_tokens=0, temperature=0.7, retry_count=_l2_retry,
                timeout=_l2_timeout,
            )
            # 规避"HTTP 200 但返回空串"：空内容也再补一次（换温度换采样），预算足够时值得一搏
            if not _l2_err and (not _l2_content or not _l2_content.strip()):
                _l2_budget2 = max(5, int(_wall_deadline - _deadline_time.monotonic()))
                if _l2_budget2 > 30:
                    _l2_content2, _l2_err2 = _call_llm(
                        [{'role': 'system', 'content': _light_system},
                         {'role': 'user', 'content': _light_user}],
                        max_tokens=0, temperature=0.85, retry_count=1,
                        timeout=min(180, max(45, _l2_budget2)),
                    )
                    if not _l2_err2 and _l2_content2 and _l2_content2.strip():
                        _l2_content = _l2_content2
                    elif _l2_err2:
                        _l2_err = _l2_err2
            if _l2_err:
                event_errors.append(f'轻量回退也失败: {_l2_err}')
            else:
                _l2_parsed, _l2_json_err = _extract_json_from_llm(_l2_content, expect='object')
                if _l2_json_err:
                    event_errors.append(f'轻量回退JSON解析失败: {_l2_json_err}')
                else:
                    all_nodes = _l2_parsed.get('nodes', []) or []
                    all_nodes.sort(key=_sort_key)
                    for _ni3, _n3 in enumerate(all_nodes, 1):
                        _n3['index'] = _ni3

        # 2) 单次全量回退（最后手段）：一次性生成所有事件，质量最高但最容易超时
        if not all_nodes:
            _fallback_user = f"""{shared_user_prefix}

【所有 main_events】（请为每个事件分别展开成 {_per_event_node_count}-{_per_event_node_max} 个子节点，子节点 chapters 必须严格卡在各事件分配的区间内，全部节点合计覆盖 {_evt_start}-{_evt_end} 章）：
{main_events_block}

请为第 {volume_index} 卷的所有 main_events 分别展开成 {_per_event_node_count}-{_per_event_node_max} 个情节子节点事件（nodes），
输出 JSON 对象 {{"nodes": [...]}}，不要包裹 markdown 代码块。"""
            _fb_remain = max(5, int(_wall_deadline - _deadline_time.monotonic()))
            _fb_retry = 2 if _fb_remain > 120 else (1 if _fb_remain > 60 else 0)
            # 单次回退要生成全部事件节点，属最大请求，给足 180s 客户端超时 + 输出上限 12000
            _fb_timeout = min(180, max(45, _fb_remain))
            _fb_content, _fb_err = _call_llm(
                [{'role': 'system', 'content': shared_system_prompt},
                 {'role': 'user', 'content': _fallback_user}],
                max_tokens=0, temperature=0.65, retry_count=_fb_retry,
                timeout=_fb_timeout,
            )
            if _fb_err:
                event_errors.append(f'单次生成也失败: {_fb_err}')
            else:
                _fb_parsed, _fb_json_err = _extract_json_from_llm(_fb_content, expect='object')
                if _fb_json_err:
                    event_errors.append(f'单次生成JSON解析失败: {_fb_json_err}')
                else:
                    all_nodes = _fb_parsed.get('nodes', []) or []
                    # 排序+重编号
                    all_nodes.sort(key=_sort_key)
                    for _ni2, _n2 in enumerate(all_nodes, 1):
                        _n2['index'] = _ni2

        # 构建合并后的 content（供下游 _call_llm 之后的逻辑使用）
        _merged = {
            'volume_index': volume_index,
            'volume_title': volume_title,
            'nodes': all_nodes,
        }
        import json as _json
        content = _json.dumps(_merged, ensure_ascii=False, indent=2)
        err = None
        if event_errors and not all_nodes:
            err = '；'.join(event_errors)
        # 【P0修复】node_only 模式全部失败时必须返回错误（原代码 err 检查只在
        # 非 node_only 分支内，导致静默写入空 nodes → 前端看到"生成0个节点"假成功）
        if err:
            return jsonify({'error': err}), 500
        # 设置 system_prompt 和 user_prompt 供下游 if err 判断后的逻辑使用
        system_prompt = ''
        user_prompt = content

    else:
        # ===== 整卷生成模式（默认）：生成完整卷大纲+情节节点 =====
        system_prompt = f"""你是番茄小说金番作者级别的卷纲设计师。
任务：为第 {volume_index} 卷“{volume_title}”生成详细大纲+情节节点。
【输出范围铁律】只允许输出第 {volume_index} 卷的内容。禁止在输出中复述、罗列、或带入任何其他卷（含上一卷、后续卷）的大纲/节点/剧情概要——下方【已有剧情】仅供你推理衔接时参考，绝不写进输出。volume_index 必须等于 {volume_index}，nodes 的 chapters 必须从 {start_chapter} 开始连续编号，不得覆盖其他卷。

【五幕模型对齐】本卷对应五幕中的“{current_act}”幕：{act_descriptions.get(current_act, '')}
本卷情节设计必须服务于该幕的核心目标，不得脱离五幕结构。

{cool_system_prompt}

【输出格式】严格输出以下JSON（不要包裹在markdown代码块中）：
{{
  "volume_index": {volume_index},
  "volume_title": "{volume_title}",
  "act": "{current_act}",
  "core_goal": "本卷核心目标",
  "core_conflict": "本卷主要冲突",
  "emotion_driver": "情感驱动力",
  "key_turns": ["转折点1", "转折点2", "转折点3"],
  "boss": "本卷BOSS",
  "foreshadow_new": ["新埋伏笔1"],
  "foreshadow_recycle": ["回收伏笔1"],
  "hook_type": "卷尾钩子类型",
  "ending_hook": "本卷卷尾钩子具体内容（下一卷开头需承接此钩子）",
  "nodes": [
    {{
      "title": "节点标题（动宾结构）",
      "chapters": "{start_chapter}-{start_chapter+9}",
      "type": "M（M主线/C角色/W世界观/D日常/F伏笔）",
      "summary": "本节点详细剧情概要（200-400字）",
      "cool_type": "爽点类型（八选一）",
      "cool_structure": "爽点结构（先抑后扬/直接碾压/默默装完逼）",
      "cool_contrast": "衬托方式（旁人震惊/不敢置信/事后佩服）",
      "cool_level": "爽点层级（微爽/小爽/中爽/大爽）",
      "hook": "本节点章尾钩子（七种钩子之一）"
    }}
  ]
}}
{cohesion_constraint}

【章型配额】M主线50%/C角色10%/W世界观10%/D日常20%/F伏笔10%
【小故事闭环】新事件→困难→金手指破局→暴露新信息→打脸收尾→钩子（5-8章）
本卷约 {chapters_per_volume} 章（约 {chapters_per_volume * 2400} 字），分5-8个情节节点。节点 chapters 必须从 {start_chapter} 开始连续递增。
【节点容量铁律】每个情节节点的 summary 必须足够支撑其 chapters 范围内的字数容量（按每章2400字估算），不得过于简略；各节点剧情须与上下卷节点连贯，避免剧情断裂。
【节点连贯铁律】各节点 summary 末尾必须自然过渡到下一节点开头；本卷最后一个节点必须埋下卷尾钩子。

{skill_note}"""

        user_prompt = f"""书名：{book.title}

{_build_core_params_block(bb, book)}

{f"【五幕式总纲】{chr(10)}{master_outline}" if has_master else "【五幕式总纲】（暂无，请基于下方已有剧情/卷大纲自行推演本卷情节节点，但必须符合五幕模型中“" + current_act + "”幕的定位）"}

【已有剧情】（含已生成卷纲，本卷需与之衔接）
{existing_timeline or '（暂无）'}

【本卷在已有剧情中的定位】请优先基于本卷（第""" + f"{volume_index}卷“{volume_title}”" + """）已有的 main_plot/key_events 设计节点；若已有剧情为空，则基于世界观、规则、人物合理推演。

【世界观设定】（情节节点需符合世界观规则）
""" + (worldbuilding_ctx or '（暂无）') + f"""

【核心规则】（金手指/能力限制等，不可违反）
{key_rules_ctx or '（暂无）'}

【人物档案】（情节节点需涉及这些角色，安排其互动）
{characters_ctx or '（暂无）'}

【上一卷结尾章节正文】（卷间衔接依据，本卷开头必须承接）
{prev_volume_end_summary or '（本卷为第一卷，无前文）'}

请为第 {volume_index} 卷生成详细大纲。"""

    # 节点设计模式：content 已在逐事件循环中设置，跳过常规 _call_llm
    # 【P0修复】当 node_only=True 但 current_vol_existing 不存在（用户指定了不存在的卷号）时，
    # 必须回退到整卷生成模式，否则 content 未定义会导致 UnboundLocalError 500
    if not node_only or not current_vol_existing:
        # 整卷生成模式：一次性调用 LLM 生成完整卷大纲+情节节点
        content, err = _call_llm(
            [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
            # max_tokens=0：特指"通用节点设计师"，全量卷纲+节点一次性生成，自动适配模型最大输出
            max_tokens=0, temperature=0.65, retry_count=3
        )
        if err:
            return jsonify({'error': err}), 500

    # 健壮 JSON 提取：处理 markdown 代码块、前后说明文字、尾随逗号等
    volume_data, json_err = _extract_json_from_llm(content, expect='object')
    if json_err:
        return jsonify({'error': 'AI返回格式错误，无法解析JSON', 'raw': content[:500], 'detail': json_err}), 500

    volume_text = f"""【第{volume_index}卷：{volume_data.get('volume_title', volume_title)}】
核心目标：{volume_data.get('core_goal', '')}
核心冲突：{volume_data.get('core_conflict', '')}
情感驱动：{volume_data.get('emotion_driver', '')}
关键转折：{', '.join(volume_data.get('key_turns', []))}
BOSS：{volume_data.get('boss', '')}
新埋伏笔：{', '.join(volume_data.get('foreshadow_new', []))}
回收伏笔：{', '.join(volume_data.get('foreshadow_recycle', []))}
卷尾钩子：{volume_data.get('hook_type', '')}
情节节点：
""" + '\n'.join([f"  {n.get('chapters','')}: {n.get('title','')}（{n.get('type','M')}）- {n.get('summary','')}" for n in volume_data.get('nodes', [])])

    # 修复：timeline 写入改为 JSON 合并（按 volume_index 替换/追加），避免文本拼接导致前端 JSON.parse 失败
    volumes_list = []
    if bb.timeline:
        try:
            parsed_tl = json.loads(bb.timeline)
            if isinstance(parsed_tl, list):
                volumes_list = parsed_tl
        except (json.JSONDecodeError, ValueError):
            # 旧文本格式，丢弃重建（已在 volume_text 中保留可读副本，但优先用 JSON）
            volumes_list = []
    # 构造与 PlotPanel 兼容的卷对象
    if node_only and current_vol_existing:
        # 节点设计模式：保留已有卷的非节点字段（main_plot/key_events/turning_points/climax/ending/foreshadowing 等），
        # 仅用新生成的 nodes 替换原 nodes；volume_id/volume/volume_index 也保留原值，避免前端索引错位
        new_nodes = volume_data.get('nodes', []) or []
        vol_obj = dict(current_vol_existing)  # 浅拷贝保留所有原字段
        vol_obj['nodes'] = new_nodes  # 仅替换 nodes
        # 保留可读文本副本（追加节点信息，便于人工查看）
        vol_obj['raw_text'] = (current_vol_existing.get('raw_text') or '') + '\n\n【节点设计更新】\n' + '\n'.join([f"  {n.get('chapters','')}: {n.get('title','')}（{n.get('type','M')}）- {n.get('summary','')}" for n in new_nodes])
    else:
        # 整卷生成模式：用 AI 返回的完整卷数据构造卷对象
        vol_obj = {
            'volume_id': str(volume_index),
            'volume': volume_data.get('volume_title', volume_title),
            'volume_index': volume_index,
            'main_plot': volume_data.get('core_goal', ''),
            'core_conflict': volume_data.get('core_conflict', ''),
            'emotion_driver': volume_data.get('emotion_driver', ''),
            'key_events': volume_data.get('key_turns', []),
            'turning_points': volume_data.get('key_turns', []),
            'climax': volume_data.get('boss', ''),
            'ending': volume_data.get('hook_type', ''),
            'foreshadowing': volume_data.get('foreshadow_new', []),
            'foreshadow_recycle': volume_data.get('foreshadow_recycle', []),
            'nodes': volume_data.get('nodes', []),
            'raw_text': volume_text,  # 保留可读文本副本
        }
    # 按 volume_index 替换或追加
    replaced = False
    for i, v in enumerate(volumes_list):
        existing_idx = v.get('volume_index') or _extract_volume_index(v.get('volume', v.get('volume_id', '')))
        if str(existing_idx) == str(volume_index):
            volumes_list[i] = vol_obj
            replaced = True
            break
    if not replaced:
        volumes_list.append(vol_obj)
    # 按 volume_index 排序
    volumes_list.sort(key=lambda v: int(v.get('volume_index', 0) or _extract_volume_index(v.get('volume', v.get('volume_id', '0'))) or 0))
    bb.timeline = json.dumps(volumes_list, ensure_ascii=False, indent=2)
    _sync_foreshadowings_to_volumes(bb)  # 【P0修复】自动同步本卷伏笔到 foreshadowing_volumes
    db.session.commit()

    return jsonify({'volume_data': volume_data, 'timeline': bb.timeline, 'bible': bb.to_dict()})

@app.route('/api/books/<book_id>/ai-extract-volumes-from-outline', methods=['POST'])
def ai_extract_volumes_from_outline(book_id):
    """从 plot_design 总纲一次性提取全部卷的剧情 JSON 数组，写入 timeline。
    替代前端逐卷循环调用 ai_outline_volume 的方式，更稳定不会中途失败。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Not found'}), 404
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb or not bb.plot_design or not bb.plot_design.strip():
        return jsonify({'error': '请先在大纲维度生成五幕式总纲'}), 400

    data = request.json or {}
    skill_pack_ids = data.get('skill_pack_ids', [])

    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['volume_breakdown', 'chapter_plan', 'tomato_outline'], mode='agent')

    # ========== 卷数权威口径：_get_total_volumes（前端不传时必须从 Book/BookBible 读用户创建时的设定） ==========
    tv = _get_total_volumes(bb, book)
    user_volume_count = data.get('volume_count')
    if user_volume_count and int(user_volume_count) >= 1:
        volume_count = int(user_volume_count)
    else:
        volume_count = tv

    # 上下文：总纲 + 世界观 + 规则 + 人物
    context_parts = [f'【五幕式总纲】\n{bb.plot_design[:4000]}']
    if bb.worldbuilding:
        context_parts.append(f'【世界观】\n{bb.worldbuilding[:1000]}')
    if bb.key_rules:
        context_parts.append(f'【核心规则】\n{bb.key_rules[:800]}')
    if bb.character_profiles:
        context_parts.append(f'【人物档案】\n{bb.character_profiles[:1000]}')
    context = '\n\n'.join(context_parts)

    # count_hint：强制写成“严格 N 卷”，绝对禁止出现“通常5-8卷”这种无视用户设定的字样
    count_hint = f'必须严格 {volume_count} 卷（多一卷少一卷都不合格，若解析出来不是 {volume_count} 卷将直接报错）'

    # ===== 【P1弊端7修复】五幕模型对齐约束（按总卷数比例动态映射，不再硬编码 5 卷=5幕） =====
    def map_act(vol_index_1based, total):
        pct = (vol_index_1based - 0.5) / max(1, total)
        if pct <= 0.05: return '立身'
        if pct <= 0.25: return '立足'
        if pct <= 0.50: return '立势'
        if pct <= 0.75: return '立威'
        return '立命'
    act_descriptions = {
        '立身': '主角登场、金手指获得、确立生存基础（前5%卷）',
        '立足': '主角站稳脚跟、初露锋芒、建立基本人际网（5-25%卷）',
        '立势': '主角势力扩张、主要冲突激化、BOSS浮出（25-50%卷）',
        '立威': '主角与BOSS正面对抗、实力跃升、打脸高潮（50-75%卷）',
        '立命': '终局决战、伏笔回收、世界观全貌揭示（75-100%卷）',
    }
    # 给 prompt 直接生成一个“卷1~卷N → 对应幕”的对照表，避免 AI 还按 5 卷=5幕去想
    act_table_lines = []
    for i in range(1, min(volume_count + 1, 31)):
        act_table_lines.append(f'第{i}卷 → {map_act(i, volume_count)}')
    if volume_count > 30:
        act_table_lines.append(f'第31~{volume_count}卷 → {map_act(volume_count, volume_count)}')
    act_table = '\n'.join(act_table_lines)

    system_prompt = f"""你是番茄小说金番作者级别的剧情架构师。
任务：根据五幕式总纲，一次性提取全部卷的【主线剧情】，输出为 JSON 数组。

【核心创作参数·铁律·不可违反】
- 总卷数：{volume_count} 卷（本次硬约束：最终 JSON 数组长度必须严格等于 {volume_count}，多一个少一个都不合格，必须重写。禁止输出 5/6/8/10/5-8 等其他卷数）
{_build_core_params_block(bb, book)}

【重要】本次只构建各卷主线剧情，**不生成情节节点**（nodes 输出为空数组）。情节节点由用户后续手动点击“节点设计”逐卷生成。

【五幕模型对齐·按 {volume_count} 卷比例自动映射】（对照表：卷号 → 所属幕）
{act_table}
各幕说明：
- 立身：{act_descriptions['立身']}
- 立足：{act_descriptions['立足']}
- 立势：{act_descriptions['立势']}
- 立威：{act_descriptions['立威']}
- 立命：{act_descriptions['立命']}
（若卷数超过5卷，按上方卷号对照表映射；每卷 main_plot 中必须标注所属幕）

【内容容量铁律】每卷固定 50 章约 12 万字（50章×2400字）。每卷 main_plot 必须足够丰满，能支撑 50 章 12 万字的内容容量，不得过于单薄；核心冲突、关键事件要充实，为后续节点设计留足展开空间。

【卷间衔接铁律】（多卷必须保持连贯）
- 第N卷的 ending_hook 必须与第N+1卷开头场景直接衔接
- 卷间不得出现剧情断裂或矛盾

【输入上下文】
{context}

【输出要求】严格输出 JSON 数组（不要包裹在 markdown 代码块中），{count_hint}。
每卷结构如下（nodes 必须为空数组，由后续节点设计填充）：
{{
  "volume_id": "1",
  "volume": "第1卷 卷名",
  "volume_index": 1,
  "act": "立身",
  "main_plot": "本卷主线剧情（150-300字，需标注所属幕，须支撑50章12万字容量）",
  "core_conflict": "本卷核心冲突",
  "emotion_driver": "情感驱动力",
  "key_events": ["关键事件1", "关键事件2", "关键事件3"],
  "turning_points": ["转折点1", "转折点2"],
  "climax": "本卷高潮/BOSS",
  "ending": "本卷结局/卷尾钩子",
  "ending_hook": "本卷卷尾钩子具体内容（下一卷开头需承接此钩子）",
  "foreshadowing": ["新埋伏笔1"],
  "nodes": []
}}

{skill_note}"""

    user_prompt = f'请根据五幕式总纲提取全部卷的主线剧情，{count_hint}。注意：只构建各卷主线，不要生成情节节点（nodes 留空）；各卷必须符合五幕模型定位且卷间 ending_hook 与开头严格衔接；每卷主线须能支撑50章12万字容量。'

    content, err = _call_llm(
        [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
        max_tokens=0, temperature=0.7
    )
    if err:
        return jsonify({'error': err}), 500

    import re

    def _safe_volume_index(v, i):
        """安全提取卷号：优先 volume_index，再 volume/volume_id，最后用序号。始终返回 int。"""
        raw = v.get('volume_index')
        if raw is None:
            raw = v.get('volume', v.get('volume_id', ''))
        idx = _extract_volume_index(raw)
        return idx if idx > 0 else (i + 1)

    volumes = None
    parse_error = None

    # 策略1：去除 markdown 代码块围栏后，直接尝试整段解析
    cleaned = content.strip()
    fence_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', cleaned)
    if fence_match:
        cleaned = fence_match.group(1).strip()

    # 策略2：尝试整段 JSON 解析（数组 或 含 volumes 字段的对象）
    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, list):
            volumes = parsed
        elif isinstance(parsed, dict):
            # 兼容 {"volumes": [...]} / {"data": [...]} / {"result": [...]} 等包裹
            for k in ('volumes', 'data', 'result', 'items', 'list'):
                if isinstance(parsed.get(k), list):
                    volumes = parsed[k]
                    break
            # 单个对象当一卷处理
            if volumes is None and 'volume' in parsed:
                volumes = [parsed]
    except (json.JSONDecodeError, ValueError) as e:
        parse_error = str(e)

    # 策略3：正则提取最外层数组（非贪婪优先，回退贪婪）
    if not volumes:
        for pattern in (r'\[\s*\{[\s\S]*\}\s*\]', r'\[[\s\S]*\]'):
            m = re.search(pattern, cleaned)
            if m:
                try:
                    cand = json.loads(m.group())
                    if isinstance(cand, list) and cand:
                        volumes = cand
                        break
                except (json.JSONDecodeError, ValueError) as e:
                    parse_error = str(e)

    if not volumes or not isinstance(volumes, list) or len(volumes) == 0:
        return jsonify({'error': 'AI返回格式错误，无法解析为JSON数组', 'raw': content[:800], 'parse_error': parse_error}), 500

    # 过滤非字典项并补全字段
    volumes = [v for v in volumes if isinstance(v, dict)]
    if not volumes:
        return jsonify({'error': 'AI返回的JSON数组中无有效卷数据', 'raw': content[:800]}), 500

    for i, v in enumerate(volumes):
        v['volume_index'] = _safe_volume_index(v, i)
        if 'volume_id' not in v:
            v['volume_id'] = str(v['volume_index'])
        if 'volume' not in v:
            v['volume'] = f'第{v["volume_index"]}卷'

    # 按 volume_index 排序（已确保为 int，不会抛异常）
    volumes.sort(key=lambda v: v['volume_index'])

    # ===== 【P0弊端3修复】写入逻辑改为 upsert 合并，而非整体覆盖 =====
    # 保留用户手工微调的卷数据，只更新本次生成的卷
    # 本次只生成主线（nodes 为空），若已有卷已设计过节点，保留其 nodes 不覆盖
    existing_volumes_list = []
    if bb.timeline:
        try:
            parsed_tl = json.loads(bb.timeline)
            if isinstance(parsed_tl, list):
                existing_volumes_list = [v for v in parsed_tl if isinstance(v, dict)]
        except (json.JSONDecodeError, ValueError, TypeError):
            pass

    # 建立已有卷索引 → nodes 的映射，用于回填被保留的节点设计
    existing_nodes_map = {}
    for v in existing_volumes_list:
        idx = v.get('volume_index')
        if idx is not None:
            existing_nodes_map[int(idx)] = v.get('nodes') or []

    # 用本次生成的卷覆盖同 volume_index 的旧卷，保留未涉及的卷
    new_idx_set = {v['volume_index'] for v in volumes}
    merged_volumes = [v for v in existing_volumes_list if v.get('volume_index') not in new_idx_set]
    # 对本次生成的卷：若 nodes 为空且已有卷存在节点设计，回填保留
    for v in volumes:
        existing_ns = existing_nodes_map.get(int(v['volume_index']))
        if (not v.get('nodes')) and existing_ns:
            v['nodes'] = existing_ns
        merged_volumes.append(v)
    merged_volumes.sort(key=lambda v: int(v.get('volume_index', 0) or _extract_volume_index(v.get('volume', v.get('volume_id', '0'))) or 0))

    # ===== 【对齐硬约束】最终强制对齐到用户设定的 volume_count，少了自动补占位、多了截断 =====
    merged_volumes, align_warning = _validate_and_align_timeline_volumes(merged_volumes, volume_count)

    bb.timeline = json.dumps(merged_volumes, ensure_ascii=False, indent=2)
    _sync_foreshadowings_to_volumes(bb)  # 【P0修复】自动同步各卷伏笔到 foreshadowing_volumes
    db.session.commit()
    result = {'success': True, 'volumes': merged_volumes, 'bible': bb.to_dict(), 'volume_count': len(merged_volumes)}
    if align_warning:
        result['warning'] = align_warning
    return jsonify(result)

@app.route('/api/books/<book_id>/ai-reverse-generate-outline', methods=['POST'])
def ai_reverse_generate_outline(book_id):
    """反生成五幕式总纲：从已导入的各卷剧情（timeline）反向提炼五幕式总纲，
    自动填入大纲维度（plot_design）。打通“导入剧情大纲 → 大纲总纲”的反哺链路。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Not found'}), 404
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    # 必须有各卷剧情才能反推总纲
    volumes_data = []
    if bb.timeline:
        try:
            parsed = json.loads(bb.timeline)
            if isinstance(parsed, list) and parsed:
                volumes_data = [v for v in parsed if isinstance(v, dict)]
        except (json.JSONDecodeError, ValueError):
            pass
    if not volumes_data:
        return jsonify({'error': '请先导入剧情大纲或提取各卷，再反生成总纲'}), 400

    data = request.json or {}
    skill_pack_ids = data.get('skill_pack_ids', [])
    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['volume_breakdown', 'tomato_outline'], mode='agent')

    # 整理各卷剧情摘要
    vol_summaries = []
    for v in volumes_data:
        idx = v.get('volume_index', '?')
        name = v.get('volume', f'第{idx}卷')
        main_plot = (v.get('main_plot') or '').strip()
        climax = (v.get('climax') or '').strip()
        ending = (v.get('ending') or '').strip()
        key_events = v.get('key_events') or []
        events_str = '；'.join(key_events) if key_events else ''
        parts = [f'第{idx}卷“{name}”']
        if main_plot:
            parts.append(f'主线：{main_plot[:200]}')
        if events_str:
            parts.append(f'关键事件：{events_str[:150]}')
        if climax:
            parts.append(f'高潮：{climax[:100]}')
        if ending:
            parts.append(f'结局/钩子：{ending[:100]}')
        vol_summaries.append(' | '.join(parts))
    volumes_text = '\n'.join(vol_summaries)

    # 五幕式总纲 = 立身卷/立足卷/立势卷/破局卷/收束卷 的全书结构
    existing_master = (bb.plot_design or '').strip()
    worldbuilding_ctx = (bb.worldbuilding or '')[:800]
    characters_ctx = (bb.character_profiles or '')[:800]

    system_prompt = f"""你是番茄小说金番作者级别的剧情架构师。
任务：根据已有的各卷剧情，反向提炼生成“五幕式总纲”，写入大纲维度。

【已有各卷剧情】
{volumes_text}

{f"【已有总纲（参考，可在其基础上完善）】{chr(10)}{existing_master[:1500]}" if existing_master else "【已有总纲】（暂无，需全新生成）"}

【世界观设定】
{worldbuilding_ctx or '（暂无）'}

【人物档案】
{characters_ctx or '（暂无）'}

【五幕式总纲格式要求】
严格按五幕结构输出，每幕对应一卷或多卷：
1. 立身卷（开局）：主角起点、金手指觉醒、核心矛盾引入
2. 立足卷（发展）：主角站稳脚跟、势力初成、第一波爽点
3. 立势卷（升级）：格局打开、对手升级、伏笔展开
4. 破局卷（高潮）：终极对决、伏笔回收、真相揭露
5. 收束卷（结局）：收尾、升华、留白/续集钩子

每幕需包含：本卷目标、核心冲突、关键转折、爽点设计、卷尾钩子。
总纲长度 800-1500 字，要能统领全书各卷。

{skill_note}"""

    user_prompt = '请根据上述各卷剧情，反向生成完整的五幕式总纲。直接输出总纲文本，不要包裹在代码块中。'

    content, err = _call_llm(
        [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
        max_tokens=0, temperature=0.6
    )
    if err:
        return jsonify({'error': err}), 500

    master_outline = content.strip()
    # 去除可能的代码围栏
    import re as _re3
    fence = _re3.search(r'```(?:markdown|text)?\s*([\s\S]*?)\s*```', master_outline)
    if fence:
        master_outline = fence.group(1).strip()

    # 写入 plot_design（大纲维度）
    bb.plot_design = master_outline
    db.session.commit()

    return jsonify({'success': True, 'master_outline': master_outline, 'bible': bb.to_dict()})

@app.route('/api/books/<book_id>/ai-import-plot-outline', methods=['POST'])
def ai_import_plot_outline(book_id):
    """导入剧情大纲文本，自动识别拆分到各卷（正则优先，AI兜底）。
    支持格式：第X卷/卷X/Volume X 等开头的卷标题 + 后续内容。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Not found'}), 404
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    data = request.json or {}
    outline_text = (data.get('outline_text') or '').strip()
    skill_pack_ids = data.get('skill_pack_ids', [])
    if not outline_text:
        return jsonify({'error': '请输入大纲文本'}), 400

    import re

    # ===== 第一步：增强正则匹配卷标题 =====
    # 支持格式：第X卷/卷X/Volume X/第X部/第X篇/第X部分/Chapter X/序章/楔子/卷壹貳叁/# 第X卷 等
    vol_pattern = re.compile(
        r'^(?:'
        r'第\s*([一二三四五六七八九十百零壹貳贰叁肆伍陆陸柒捌玖拾\d]+)\s*[卷部篇部分]'  # 第X卷/部/篇/部分
        r'|卷\s*([一二三四五六七八九十百零壹貳贰叁肆伍陆陸柒捌玖拾\d]+)'  # 卷X
        r'|Volume\s*(\d+)|Chapter\s*(\d+)'  # Volume X / Chapter X
        r'|(序章|楔子|引子|终章|尾声)'  # 序章/楔子等
        r'|#*\s*第\s*([一二三四五六七八九十百零壹貳贰叁肆伍陆陸柒捌玖拾\d]+)\s*卷'
        r')\s*[：:．.、\s\-—]*(.*)$',
        re.MULTILINE
    )
    matches = list(vol_pattern.finditer(outline_text))

    volumes = []
    if matches:
        # 有标准格式，按卷拆分（即使只有 1 个 match 也处理，避免单卷被丢弃）
        for i, m in enumerate(matches):
            # 提取卷号
            vol_num_str = m.group(1) or m.group(2) or m.group(3) or m.group(4) or m.group(6) or ''
            vol_idx = _extract_volume_index(vol_num_str) or (i + 1)
            # 序章/楔子等特殊章节
            special = m.group(5)
            _title_remainder = ''  # 卷名行中卷名之后的剧情描述，回填到卷内容
            if special:
                vol_title = special
                vol_idx = 0 if special in ('序章', '楔子', '引子') else 999
            else:
                # 【P0修复】只提取书名号/冒号前的纯卷名，避免整段剧情被识别成卷名
                _raw_title_line = (m.group(7) or '').strip()
                vol_title, _title_remainder = _split_volume_title(_raw_title_line)
                vol_title = vol_title or f'第{vol_idx}卷'
            # 内容范围：从当前匹配结束到下一个匹配开始（最后一卷取到文末，不丢尾部）
            content_start = m.end()
            content_end = matches[i + 1].start() if i + 1 < len(matches) else len(outline_text)
            vol_content = outline_text[content_start:content_end].strip()
            # 卷名行剩余的剧情描述前置到卷内容，避免丢失（原本被当作卷名的剧情现在归入 main_plot）
            if _title_remainder:
                vol_content = (f'{_title_remainder}\n{vol_content}').strip() if vol_content else _title_remainder

            volumes.append({
                'volume_id': str(vol_idx),
                'volume': f'第{vol_idx}卷 {vol_title}' if vol_title and not vol_title.startswith('第') and not special else vol_title,
                'volume_index': vol_idx,
                'main_plot': vol_content[:500],
                'core_conflict': '',
                'emotion_driver': '',
                'key_events': [l.strip().lstrip('·•*-') for l in vol_content.split('\n') if l.strip() and len(l.strip()) > 5][:8],
                'turning_points': [],
                'climax': '',
                'ending': '',
                'foreshadowing': [],
                'nodes': [],
                'raw_text': vol_content,
            })

    # 【P0修复】正则路径补全：为每卷生成 nodes（章号全书连续）+ ending_hook，
    # 避免后续滚动生成时 start_chapter 解析失败导致章号断裂
    if matches and volumes:
        # 取每卷章数：优先用户大纲隐含的，否则默认50；特殊卷（序章等）1章
        cpc = data.get('chapters_per_volume', 50) or 50
        ch_cursor = 1
        for vi, v in enumerate(volumes):
            special = v.get('volume_index', 0) in (0, 999)  # 序章/终章
            vol_chs = 1 if special else cpc
            start = ch_cursor
            end = ch_cursor + vol_chs - 1
            # 把卷内容按 key_events 等分为最多5个节点
            ke = v.get('key_events') or []
            if not ke:
                ke = ['开局', '发展', '转折', '高潮', '收束'][:5]
            n_nodes = min(5, max(1, len(ke)))
            step = vol_chs / n_nodes
            nodes = []
            for ni in range(n_nodes):
                ns = start + int(step * ni)
                ne = start + int(step * (ni + 1)) - 1 if ni < n_nodes - 1 else end
                if ne < ns:
                    ne = ns
                ntype = ['M', 'D', 'T', 'C', 'F'][ni] if ni < 5 else 'M'  # 开局/发展/转折/高潮/收束
                nodes.append({
                    'chapters': f'{ns}-{ne}',
                    'type': ntype,
                    'title': ke[ni] if ni < len(ke) else f'节点{ni+1}',
                    'summary': (ke[ni] if ni < len(ke) else '')[:80],
                })
            v['nodes'] = nodes
            # ending_hook：取关键事件最后一条或 main_plot 末句
            last_ke = ke[-1] if ke else ''
            v['ending'] = (last_ke or v.get('main_plot', '')[-100:])[:150]
            if not v.get('core_conflict'):
                v['core_conflict'] = (ke[0] if ke else '')[:100]
            ch_cursor = end + 1

    # 如果正则匹配到第一个 match 之前还有内容，作为"开篇/引言"归入第一卷或单独成卷
    if matches and matches[0].start() > 0:
        head_content = outline_text[:matches[0].start()].strip()
        if head_content and len(head_content) > 20:
            # 归入第一卷的 main_plot 前置
            if volumes:
                volumes[0]['main_plot'] = head_content[:300] + '\n' + volumes[0]['main_plot']

    # ===== 第二步：正则匹配失败或卷数<1，调用 AI 智能拆卷（改进版） =====
    if len(volumes) < 1:
        skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['volume_breakdown', 'chapter_plan', 'tomato_outline'], mode='agent')
        # 上下文增强：注入总纲、规则、已有卷、世界观、人物
        ctx_parts = []
        if bb.plot_design:
            ctx_parts.append(f'【五幕式总纲（重要：决定全书卷数和卷目标）】\n{bb.plot_design[:2500]}')
        if bb.key_rules:
            ctx_parts.append(f'【核心规则】\n{bb.key_rules[:800]}')
        if bb.worldbuilding:
            ctx_parts.append(f'【世界观】\n{bb.worldbuilding[:800]}')
        if bb.character_profiles:
            ctx_parts.append(f'【人物档案】\n{bb.character_profiles[:800]}')
        if bb.timeline:
            try:
                existing_vols = json.loads(bb.timeline)
                if isinstance(existing_vols, list) and existing_vols:
                    existing_summary = '；'.join([f"第{v.get('volume_index','?')}卷:{v.get('main_plot','')[:30]}" for v in existing_vols[:10]])
                    ctx_parts.append(f'【已有卷剧情（避免冲突，可按volume_index合并或追加）】\n{existing_summary}')
            except (json.JSONDecodeError, ValueError):
                pass
        extra_ctx = '\n\n'.join(ctx_parts)

        system_prompt = f"""你是番茄小说金番作者级别的剧情架构师。
任务：将用户粘贴的大纲文本**严格按原文的卷划分**拆分为 JSON 数组。

【已有设定参考】
{extra_ctx or '（暂无）'}

【拆卷铁律】
1. **必须严格按原文的卷划分提取**，不得合并、拆分、重排原文卷
2. 原文有明确的卷标题（如"第X卷/卷X/Volume X/序章/楔子"）时，每个标题对应一个卷对象
3. 原文没有明确卷划分时，才可按剧情自然分段（每卷对应一个完整的故事弧线）
4. **不得遗漏任何卷**：从原文开头到结尾，每个段落都必须归属到某个卷
5. main_plot 必须从原文对应段落提取或概括，不得凭空捏造
6. 卷数应与【五幕式总纲】中暗示的卷数一致（若总纲存在）

【输出要求】严格输出 JSON 数组（不要包裹在 markdown 代码块中）。
每卷结构：
{{
  "volume_id": "1",
  "volume": "第1卷 卷名",
  "volume_index": 1,
  "main_plot": "本卷主线剧情（100-300字，从原文提取）",
  "core_conflict": "核心冲突",
  "emotion_driver": "情感驱动",
  "key_events": ["关键事件1", "关键事件2"],
  "turning_points": ["转折点1"],
  "climax": "高潮",
  "ending": "结局/钩子",
  "foreshadowing": ["伏笔1"],
  "nodes": [
    {{"title": "节点1", "chapters": "1-10", "type": "M", "summary": "概要", "cool_type": "爽点类型"}}
  ]
}}

【情节节点设计要求】
- 每卷生成 5-8 个情节节点
- 章型配额：M主线50%/C角色10%/W世界观10%/D日常20%/F伏笔10%
- 相邻节点章型不同，每卷覆盖整卷章节范围
- 小故事闭环：新事件→困难→金手指破局→暴露新信息→打脸收尾→钩子（5-8章）

{skill_note}"""

        # 取消 6000 字截断，支持超长大纲（分块处理）
        max_input = 12000
        outline_chunk = outline_text[:max_input]
        if len(outline_text) > max_input:
            outline_chunk += f'\n\n[注：原文共 {len(outline_text)} 字，已截取前 {max_input} 字，请确保覆盖全部卷]'

        user_prompt = f'请严格按原文卷划分拆分以下大纲文本，不得遗漏任何卷：\n\n{outline_chunk}'

        content, err = _call_llm(
            [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
            max_tokens=0, temperature=0.2
        )
        if err:
            return jsonify({'error': err}), 500

        # 稳健解析 AI 返回（三策略：整段→对象包裹→正则数组），与 ai_extract_volumes_from_outline 一致
        import re as _re2
        cleaned = content.strip()
        fence_match = _re2.search(r'```(?:json)?\s*([\s\S]*?)\s*```', cleaned)
        if fence_match:
            cleaned = fence_match.group(1).strip()

        ai_volumes = None
        try:
            parsed = json.loads(cleaned)
            if isinstance(parsed, list):
                ai_volumes = parsed
            elif isinstance(parsed, dict):
                for k in ('volumes', 'data', 'result', 'items', 'list'):
                    if isinstance(parsed.get(k), list):
                        ai_volumes = parsed[k]
                        break
                if ai_volumes is None and 'volume' in parsed:
                    ai_volumes = [parsed]
        except (json.JSONDecodeError, ValueError):
            pass

        if not ai_volumes:
            for pattern in (r'\[\s*\{[\s\S]*\}\s*\]', r'\[[\s\S]*\]'):
                m = _re2.search(pattern, cleaned)
                if m:
                    try:
                        cand = json.loads(m.group())
                        if isinstance(cand, list) and cand:
                            ai_volumes = cand
                            break
                    except (json.JSONDecodeError, ValueError):
                        pass

        if not ai_volumes:
            return jsonify({'error': 'AI返回格式错误，无法解析为JSON数组', 'raw': content[:500]}), 500

        volumes = ai_volumes

        # 补全字段 + 安全 volume_index（避免 int() 异常）
        def _safe_idx(v, i):
            raw = v.get('volume_index')
            if raw is None:
                raw = v.get('volume', v.get('volume_id', ''))
            idx = _extract_volume_index(raw)
            return idx if idx > 0 else (i + 1)

        for i, v in enumerate(volumes):
            v['volume_index'] = _safe_idx(v, i)
            if 'volume_id' not in v:
                v['volume_id'] = str(v['volume_index'])
            if 'volume' not in v:
                v['volume'] = f'第{v["volume_index"]}卷'
            if 'nodes' not in v:
                v['nodes'] = []

    # 合并到已有 timeline：按 volume_index(int) 精确替换或追加，避免空 index 互相覆盖
    existing_volumes = []
    if bb.timeline:
        try:
            parsed = json.loads(bb.timeline)
            if isinstance(parsed, list):
                existing_volumes = [v for v in parsed if isinstance(v, dict)]
        except (json.JSONDecodeError, ValueError):
            existing_volumes = []

    def _vol_int_idx(v, fallback=0):
        """安全提取 int 卷号，失败回退 fallback。"""
        raw = v.get('volume_index')
        if raw is None:
            raw = v.get('volume', v.get('volume_id', ''))
        idx = _extract_volume_index(raw)
        return idx if idx > 0 else fallback

    # 建立 existing 的 index→位置 映射（用 int 精确匹配，不再用 str 比对）
    existing_idx_map = {}
    for i, ev in enumerate(existing_volumes):
        ei = _vol_int_idx(ev)
        if ei > 0 and ei not in existing_idx_map:
            existing_idx_map[ei] = i

    for new_v in volumes:
        ni = _vol_int_idx(new_v)
        if ni > 0 and ni in existing_idx_map:
            existing_volumes[existing_idx_map[ni]] = new_v
        else:
            existing_volumes.append(new_v)

    # 按 int 卷号稳定排序（相同 index 保持原顺序，不会前后颠倒）
    existing_volumes.sort(key=lambda v: _vol_int_idx(v, 9999))

    bb.timeline = json.dumps(existing_volumes, ensure_ascii=False, indent=2)
    _sync_foreshadowings_to_volumes(bb)  # 【P0修复】导入剧情后同步各卷伏笔到 foreshadowing_volumes
    db.session.commit()

    return jsonify({'success': True, 'volumes': existing_volumes, 'imported_count': len(volumes), 'bible': bb.to_dict()})

# ==== 总 AI 创作：总览全局各维度，用户确认后填入 ====
@app.route('/api/books/<book_id>/ai-master-create', methods=['POST'])
def ai_master_create(book_id):
    """总AI创作：按番茄金番工作流串行协同，本轮产出回流给下一轮作为上下文。
    维度依赖图（agent 协同）：concept → key_rules → worldbuilding → character_profiles → plot_design → timeline
    每个维度都能看到本轮已生成的所有上游维度产物。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Not found'}), 404
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()
    # P0-3修复：首次创建 bible 时同步 book 的 total_volumes/novel_styles，
    # 否则首次总创作时风格流派注入为空、卷数约束失效
    _sync_book_meta_to_bible(book, bb)

    data = request.json or {}
    # 【P1-5修复】总创作技能包注入：从 book.master_skill_ids 读取（构思类），不再用请求参数
    # 老数据兼容：若 book 三类字段全空，先尝试迁移老的 skill_pack_ids，仍空则回退到请求参数
    try:
        _split_legacy_skill_ids_to_categories(book)
    except Exception:
        pass
    skill_pack_ids = _resolve_skill_ids_by_category(book, 'master')
    if not skill_pack_ids:
        skill_pack_ids = data.get('skill_pack_ids', [])
    dimensions = data.get('dimensions', ['concept', 'key_rules', 'worldbuilding', 'character_profiles', 'plot_design'])
    instruction = data.get('instruction', '')
    # 本轮会话已生成但未确认的内容（跨维度实时互通）
    session_outputs = data.get('session_outputs') or {}

    DIM_MAP = MASTER_DIM_MAP  # 复用公共常量
    DIM_ORDER = MASTER_DIM_ORDER
    ordered_dims = [d for d in DIM_ORDER if d in dimensions]

    # 上下文字典：本轮已生成内容(session_outputs)优先，回退 bible 已有内容
    ctx = _build_master_ctx(bb, session_outputs)

    # 防遗忘检查诊断回注（让维度创作也规避已诊断出的问题）
    af_alerts_master = _collect_anti_forget_alerts(bb, max_reports=2, max_alerts=8)
    af_block_master = ''
    if af_alerts_master:
        af_block_master = f"""
【防遗忘检查诊断】（最近检查发现的问题，本维度创作必须保持一致、主动规避）
{af_alerts_master}
"""

    # 已写剧情上下文（章节摘要+动态报告+伏笔图+关系图），让维度创作感知已写正文
    storyline_ctx = _build_master_storyline_ctx(book_id, bb)
    storyline_block = ''
    if storyline_ctx:
        storyline_block = f"""
【已写剧情参照】（已写章节的实际剧情，本维度创作须与之保持一致，不可矛盾）
{storyline_ctx}
"""

    results = []
    for dim in ordered_dims:
        info = DIM_MAP[dim]
        skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', info['keys'], mode='agent')
        # 【P1修复】格式整合铁律：技能包内容是创作方法论，输出必须严格按平台维度的格式骨架
        # 即使技能包自带独立格式要求（如CDL档案/五不妥协原则），模型也要按平台格式整合映射
        format_integration_note = ''
        if skill_note:
            format_integration_note = '\n\n【格式整合铁律·必读】上方“技能包内容”是创作方法论（指导原则），不是输出格式模板。必须严格按下方任务的“输出格式”骨架输出，技能包的要求整合映射到对应字段。例如：技能包要求"CDL角色档案"中的"外貌特征/战斗风格"应并入角色卡对应字段；技能包要求"五不妥协原则"应整合到对应小节内。不要把技能包字段原样搬出来，要按平台格式重新组织。'

        # 按 DAG 依赖图只注入该维度的上游维度产物（不再无差别全注入）
        upstream_ctx = _build_master_upstream_ctx(dim, ctx)

        # 标注当前维度在 workflow 中的位置
        dim_idx = DIM_ORDER.index(dim)
        upstream_names = [DIM_MAP[d]['label'] for d in DIM_ORDER[:dim_idx] if ctx[d].strip()]
        downstream_names = [DIM_MAP[d]['label'] for d in DIM_ORDER[dim_idx+1:] if d in ordered_dims]
        position_note = f'你正在执行第 {dim_idx+1}/{len(DIM_ORDER)} 步：{info["label"]}设计'
        if upstream_names:
            position_note += f'（上游已完成：{"→".join(upstream_names)}）'
        if downstream_names:
            position_note += f'（下游将基于你的产出继续：{"→".join(downstream_names)}）'

        # ===== 【P0弊端4修复】timeline 维度输出 JSON 数组格式，与 _get_volume_outline 兼容 =====
        # 【P1修复】硬性铁律：卷数严格按用户选择的总卷数，卷序号连续，chapters 编号全书连续
        is_timeline_dim = (dim == 'timeline')
        core_params_block = _build_core_params_block(bb, book)
        if is_timeline_dim:
            tv_for_timeline = _get_total_volumes(bb, book)
            chapters_per_vol = _get_chapters_per_volume(bb, book)
            total_chapters = tv_for_timeline * chapters_per_vol
            system_prompt = f"""你是番茄小说金番作者级别的{info['label']}设计师，正在与其他维度设计师协同创作。
{position_note}

任务：{info['prompt']}

书名：{book.title}

{core_params_block}

【已确认的上游维度产物】（必须在你的产出中保持一致，不可与上游矛盾）
{upstream_ctx}
{storyline_block}{af_block_master}
【五幕模型对齐】各卷对应五幕：立身→立足→立势→立威→立命
【卷间衔接铁律】第N卷 ending_hook 必须与第N+1卷开头严格衔接；各卷 nodes.chapters 全书连续编号。

【分卷铁律·必读】**全书共 {tv_for_timeline} 卷，每卷约 {chapters_per_vol} 章，全书约 {total_chapters} 章**。卷序号从 1 开始连续递增到 {tv_for_timeline}。卷名格式"第N卷 副标题"。必须覆盖全部 {tv_for_timeline} 卷，不得多不得少。

{skill_note}{format_integration_note}

【输出格式铁律】严格输出 JSON 数组（不要包裹在 markdown 代码块中），每卷结构如下：
[
  {{
    "volume_id": "1",
    "volume": "第1卷 副标题",
    "volume_index": 1,
    "act": "立身",
    "main_plot": "本卷主线剧情（100-200字）",
    "core_conflict": "本卷核心冲突",
    "ending_hook": "本卷卷尾钩子具体内容",
    "nodes": [
      {{"title": "节点1", "chapters": "1-10", "type": "M", "summary": "概要", "cool_type": "实力碾压"}}
    ]
  }}
]

【分卷章节分配】全书 {total_chapters} 章 → {tv_for_timeline} 卷（每卷 {chapters_per_vol} 章）：第1卷 1-{chapters_per_vol}、第2卷 {chapters_per_vol+1}-{chapters_per_vol*2}、... 第{tv_for_timeline}卷 {(tv_for_timeline-1)*chapters_per_vol+1}-{total_chapters}；每卷 nodes 章节连续不重叠。
{_cultivation_dimension_hint(dim, book, bb)}
直接输出 JSON 数组，不要任何解释性文字。"""
            user_prompt = instruction or f'请为这本小说生成分卷剧情JSON数组，**共 {tv_for_timeline} 卷，每卷约 {chapters_per_vol} 章**，与已确认的上游维度保持一致，各卷符合五幕模型且卷间严格衔接。'
        else:
            system_prompt = f"""你是番茄小说金番作者级别的{info['label']}设计师，正在与其他维度设计师协同创作。
{position_note}

任务：{info['prompt']}

书名：{book.title}

{core_params_block}

【已确认的上游维度产物】（必须在你的产出中保持一致，不可与上游矛盾）
{upstream_ctx}
{storyline_block}{af_block_master}
{skill_note}{format_integration_note}
{_cultivation_dimension_hint(dim, book, bb)}

直接输出{info['label']}内容（严格按任务要求的格式铁律输出）。确保与上游维度衔接一致。"""
            user_prompt = instruction or f'请为这本小说生成{info["label"]}，与已确认的上游维度保持一致，并严格按输出格式铁律输出。'

        # P0-1修复：各维度创作不限制输出 token，避免 timeline 等大维度 JSON 被截断导致"只能生成几卷"
        content, err = _call_llm(
            [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
            max_tokens=0, temperature=0.7
        )
        if err:
            results.append({'dimension': dim, 'label': info['label'], 'field': info['field'], 'error': err})
        else:
            # timeline 维度：校验 JSON 格式，失败则保留原文但标记警告
            timeline_warning = None
            if is_timeline_dim:
                import re as _re_mc
                cleaned_mc = content.strip()
                fence_mc = _re_mc.search(r'```(?:json)?\s*([\s\S]*?)\s*```', cleaned_mc)
                if fence_mc:
                    cleaned_mc = fence_mc.group(1).strip()
                parsed_tl_mc = None
                try:
                    parsed_tl_mc = json.loads(cleaned_mc)
                    if isinstance(parsed_tl_mc, dict):
                        for k in ('volumes', 'data', 'result', 'items', 'list'):
                            if isinstance(parsed_tl_mc.get(k), list):
                                parsed_tl_mc = parsed_tl_mc[k]
                                break
                except (json.JSONDecodeError, ValueError):
                    # 尝试正则提取数组
                    m_mc = _re_mc.search(r'\[\s*\{[\s\S]*\}\s*\]', cleaned_mc) or _re_mc.search(r'\[[\s\S]*\]', cleaned_mc)
                    if m_mc:
                        try:
                            parsed_tl_mc = json.loads(m_mc.group())
                        except (json.JSONDecodeError, ValueError):
                            parsed_tl_mc = None
                if isinstance(parsed_tl_mc, list) and parsed_tl_mc:
                    # 【P0-1修复】卷数校验与对齐：强制对齐 total_volumes
                    try:
                        tv_for_check = _get_total_volumes(bb, book)
                        parsed_tl_mc, timeline_warning = _validate_and_align_timeline_volumes(parsed_tl_mc, tv_for_check)
                    except Exception:
                        pass
                    content = json.dumps(parsed_tl_mc, ensure_ascii=False, indent=2)
                # 若解析失败，content 保留原文（纯文本），_get_volume_outline 会走纯文本回退分支
                # 【P2-8修复】timeline 解析失败主动告警
                if not isinstance(parsed_tl_mc, list) or not parsed_tl_mc:
                    timeline_warning = 'timeline维度JSON解析失败，已保留原文，但outline_hierarchy/dynamic_volumes可能受影响'
            result_entry = {'dimension': dim, 'label': info['label'], 'field': info['field'], 'content': content}
            if timeline_warning:
                result_entry['warning'] = timeline_warning
            # 【P0-2修复】foreshadowing DAG 构建失败告警
            foreshadowing_warning = None
            # 关键：本轮产出回流到 ctx，供下一轮维度作为上游上下文
            ctx[dim] = content
            # P1-9: 直接落库，避免前端未回流导致协同结果丢失
            setattr(bb, info['field'], content)
            # P0-2：伏笔维度生成后自动构建 DAG（与文本字段并存）
            if dim == 'foreshadowing' and parse_text_to_dag:
                try:
                    graph = parse_text_to_dag(content)
                    errors = graph.validate()
                    if not errors:
                        bb.foreshadowing_graph = json.dumps(graph.to_dict(), ensure_ascii=False)
                    else:
                        foreshadowing_warning = f'伏笔DAG校验失败：{"；".join(errors)[:150]}，DAG未更新但文本字段已落库'
                except Exception as dag_e:
                    foreshadowing_warning = f'伏笔DAG构建异常：{str(dag_e)[:150]}，DAG未更新但文本字段已落库。建议检查伏笔格式（## 伏笔N：标题）'
            if foreshadowing_warning:
                result_entry['warning'] = (result_entry.get('warning') or '') + ('；' if result_entry.get('warning') else '') + foreshadowing_warning
            results.append(result_entry)
            # P1-4：timeline 维度生成后自动构建四级大纲层级
            if dim == 'timeline' and build_outline_hierarchy:
                try:
                    hierarchy = build_outline_hierarchy(content, bb.plot_design or '')
                    if hierarchy.get('chapters'):
                        bb.outline_hierarchy = json.dumps(hierarchy, ensure_ascii=False)
                except Exception:
                    pass

    # P1-9: 串行生成完成后统一提交事务
    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'落库失败: {str(e)}', 'results': results}), 500

    # 返回最新 bible 供前端同步状态
    return jsonify({'results': results, 'bible': bb.to_dict()})

@app.route('/api/books/<book_id>/ai-master-create/stream', methods=['POST'])
@login_required
def ai_master_create_stream(book_id):
    """总AI创作流式版（SSE）：按番茄金番工作流串行协同，逐维度流式输出。
    每个维度生成时注入已有bible维度+本轮上游已生成维度作为上下文，保证一致性。
    SSE 格式：
      data: {"dim":"concept","label":"构思","start":true}\n\n
      data: {"choices":[{"delta":{"content":"..."}}]}\n\n  (流式chunk)
      data: {"dim":"concept","done":true}\n\n
      ... 下一维度 ...
    """
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Not found'}), 404
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()
    # P0-3修复：首次创建 bible 时同步 book 的 total_volumes/novel_styles
    _sync_book_meta_to_bible(book, bb)

    data = request.json or {}
    # 【P1-5修复】总创作技能包注入：从 book.master_skill_ids 读取（构思类），不再用请求参数
    # 老数据兼容：若 book 三类字段全空，先尝试迁移老的 skill_pack_ids，仍空则回退到请求参数
    try:
        _split_legacy_skill_ids_to_categories(book)
    except Exception:
        pass
    skill_pack_ids = _resolve_skill_ids_by_category(book, 'master')
    if not skill_pack_ids:
        skill_pack_ids = data.get('skill_pack_ids', [])
    dimensions = data.get('dimensions', ['concept', 'key_rules', 'worldbuilding', 'character_profiles', 'plot_design'])
    instruction = data.get('instruction', '')
    # 本轮会话已生成但尚未"确认填入"的维度内容（前端传入，用于跨维度实时互通）
    # 优先级：本轮已生成 > bible 已有内容（本轮产出是最新的协作基线）
    session_outputs = data.get('session_outputs') or {}

    # 复用公共常量
    DIM_MAP = MASTER_DIM_MAP
    DIM_ORDER = MASTER_DIM_ORDER
    ordered_dims = [d for d in DIM_ORDER if d in dimensions]

    # 上下文：本轮已生成(session_outputs)优先，回退 bible 已有内容
    ctx = _build_master_ctx(bb, session_outputs)

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('creation') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    # 防遗忘检查诊断回注（让维度创作也规避已诊断出的问题）
    af_alerts_master = _collect_anti_forget_alerts(bb, max_reports=2, max_alerts=8)
    af_block_master = ''
    if af_alerts_master:
        af_block_master = f"""
【防遗忘检查诊断】（最近检查发现的问题，本维度创作必须保持一致、主动规避）
{af_alerts_master}
"""

    # 已写剧情上下文（章节摘要+动态报告+伏笔图+关系图），让维度创作感知已写正文
    storyline_ctx = _build_master_storyline_ctx(book_id, bb)
    storyline_block = ''
    if storyline_ctx:
        storyline_block = f"""
【已写剧情参照】（已写章节的实际剧情，本维度创作须与之保持一致，不可矛盾）
{storyline_ctx}
"""

    def generate():
        try:
            for dim in ordered_dims:
                info = DIM_MAP[dim]
                # 推送维度开始信号
                yield f'data: {json.dumps({"dim": dim, "label": info["label"], "start": True}, ensure_ascii=False)}\n\n'

                skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', info['keys'], mode='agent')
                format_integration_note = ''
                if skill_note:
                    format_integration_note = '\n\n【格式整合铁律·必读】上方“技能包内容”是创作方法论（指导原则），不是输出格式模板。必须严格按下方任务的“输出格式”骨架输出，技能包的要求整合映射到对应字段。'

                # 按 DAG 依赖图只注入该维度的上游维度产物（不再无差别全注入）
                upstream_ctx = _build_master_upstream_ctx(dim, ctx)

                dim_idx = DIM_ORDER.index(dim)
                upstream_names = [DIM_MAP[d]['label'] for d in DIM_ORDER[:dim_idx] if ctx[d].strip()]
                downstream_names = [DIM_MAP[d]['label'] for d in DIM_ORDER[dim_idx+1:] if d in ordered_dims]
                position_note = f'你正在执行第 {dim_idx+1}/{len(DIM_ORDER)} 步：{info["label"]}设计'
                if upstream_names:
                    position_note += f'（上游已完成：{"→".join(upstream_names)}）'
                if downstream_names:
                    position_note += f'（下游将基于你的产出继续：{"→".join(downstream_names)}）'

                is_timeline_dim = (dim == 'timeline')
                core_params_block = _build_core_params_block(bb, book)
                if is_timeline_dim:
                    tv_for_timeline = _get_total_volumes(bb, book)
                    chapters_per_vol = _get_chapters_per_volume(bb, book)
                    total_chapters = tv_for_timeline * chapters_per_vol
                    system_prompt = f"""你是番茄小说金番作者级别的{info['label']}设计师，正在与其他维度设计师协同创作。
{position_note}

任务：{info['prompt']}

书名：{book.title}

{core_params_block}

【已确认的上游维度产物】（必须在你的产出中保持一致，不可与上游矛盾）
{upstream_ctx}
{storyline_block}{af_block_master}
【五幕模型对齐】各卷对应五幕：立身→立足→立势→立威→立命
【卷间衔接铁律】第N卷 ending_hook 必须与第N+1卷开头严格衔接；各卷 nodes.chapters 全书连续编号。

【分卷铁律·必读】**全书共 {tv_for_timeline} 卷，每卷约 {chapters_per_vol} 章，全书约 {total_chapters} 章**。卷序号从 1 开始连续递增到 {tv_for_timeline}。卷名格式"第N卷 副标题"。必须覆盖全部 {tv_for_timeline} 卷，不得多不得少。

{skill_note}{format_integration_note}
{_cultivation_dimension_hint(dim, book, bb)}

【输出格式铁律】严格输出 JSON 数组（不要包裹在 markdown 代码块中）。直接输出 JSON 数组，不要任何解释性文字。"""
                    user_prompt = instruction or f'请为这本小说生成分卷剧情JSON数组，**共 {tv_for_timeline} 卷，每卷约 {chapters_per_vol} 章**，与已确认的上游维度保持一致，各卷符合五幕模型且卷间严格衔接。'
                else:
                    system_prompt = f"""你是番茄小说金番作者级别的{info['label']}设计师，正在与其他维度设计师协同创作。
{position_note}

任务：{info['prompt']}

书名：{book.title}

{core_params_block}

【已确认的上游维度产物】（必须在你的产出中保持一致，不可与上游矛盾）
{upstream_ctx}
{storyline_block}{af_block_master}
{skill_note}{format_integration_note}
{_cultivation_dimension_hint(dim, book, bb)}

直接输出{info['label']}内容（严格按任务要求的格式铁律输出）。确保与上游维度衔接一致。"""
                    user_prompt = instruction or f'请为这本小说生成{info["label"]}，与已确认的上游维度保持一致，并严格按输出格式铁律输出。'

                # P0-1修复：流式维度创作不限制输出 token，避免大维度 JSON 被截断
                # （不下发 max_tokens 字段，让模型用自身默认输出上限）

                # 流式调用 LLM
                base = base_url.rstrip('/')
                if not base.endswith('/v1'):
                    base += '/v1'
                resp = requests.post(f'{base}/chat/completions',
                    headers=build_auth_headers(api_key),
                    json={'model': model,
                          'messages': [{'role': 'system', 'content': system_prompt},
                                       {'role': 'user', 'content': user_prompt}],
                          'temperature': 0.7,
                          'stream': True},
                    stream=True, timeout=180)
                full_content = ''
                for line in resp.iter_lines():
                    if line:
                        line = line.decode('utf-8')
                        if line.startswith('data: '):
                            chunk = line[6:]
                            if chunk == '[DONE]':
                                break
                            # 透传 chunk 给前端
                            yield f'data: {chunk}\n\n'
                            try:
                                parsed = json.loads(chunk)
                                delta = parsed.get('choices', [{}])[0].get('delta', {}).get('content', '')
                                full_content += delta
                            except:
                                pass

                # 维度完成信号
                # timeline 维度：校验 JSON
                timeline_warning_stream = None
                if is_timeline_dim and full_content:
                    import re as _re_s
                    cleaned = full_content.strip()
                    fence = _re_s.search(r'```(?:json)?\s*([\s\S]*?)\s*```', cleaned)
                    if fence:
                        cleaned = fence.group(1).strip()
                    parsed_tl = None
                    try:
                        parsed_tl = json.loads(cleaned)
                        if isinstance(parsed_tl, dict):
                            for k in ('volumes', 'data', 'result', 'items', 'list'):
                                if isinstance(parsed_tl.get(k), list):
                                    parsed_tl = parsed_tl[k]
                                    break
                        if isinstance(parsed_tl, list) and parsed_tl:
                            # 【P0-1修复】卷数校验与对齐
                            try:
                                tv_for_check_s = _get_total_volumes(bb, book)
                                parsed_tl, timeline_warning_stream = _validate_and_align_timeline_volumes(parsed_tl, tv_for_check_s)
                            except Exception:
                                pass
                            full_content = json.dumps(parsed_tl, ensure_ascii=False, indent=2)
                    except Exception:
                        pass
                    # 【P2-8修复】timeline 解析失败主动告警
                    if not isinstance(parsed_tl, list) or not parsed_tl:
                        timeline_warning_stream = 'timeline维度JSON解析失败，已保留原文，但outline_hierarchy/dynamic_volumes可能受影响'

                # 【P0-2修复】foreshadowing DAG 构建失败告警
                foreshadowing_warning_stream = None

                # 回流到 ctx 供下一维度使用
                ctx[dim] = full_content
                # P3-11：事务性落库 —— 只 setattr 到 session，不立即 commit，
                # 全部维度完成后统一提交，避免中途失败/用户取消导致部分维度落库、bible 状态不一致
                try:
                    setattr(bb, info['field'], full_content)
                    # P0-2：伏笔维度生成后自动构建 DAG（与文本字段并存）
                    if dim == 'foreshadowing' and parse_text_to_dag:
                        try:
                            graph = parse_text_to_dag(full_content)
                            errors = graph.validate()
                            if not errors:
                                bb.foreshadowing_graph = json.dumps(graph.to_dict(), ensure_ascii=False)
                            else:
                                foreshadowing_warning_stream = f'伏笔DAG校验失败：{"；".join(errors)[:150]}，DAG未更新但文本字段已落库'
                        except Exception as dag_e_s:
                            foreshadowing_warning_stream = f'伏笔DAG构建异常：{str(dag_e_s)[:150]}，DAG未更新但文本字段已落库。建议检查伏笔格式（## 伏笔N：标题）'
                    # P1-4：timeline 维度生成后自动构建四级大纲层级
                    if dim == 'timeline' and build_outline_hierarchy:
                        try:
                            hierarchy = build_outline_hierarchy(full_content, bb.plot_design or '')
                            if hierarchy.get('chapters'):
                                bb.outline_hierarchy = json.dumps(hierarchy, ensure_ascii=False)
                        except Exception:
                            pass  # 层级构建失败不影响文本落库
                    # 注意：此处不再立即 commit，延迟到所有维度完成后统一提交
                except Exception:
                    db.session.rollback()

                # 【P2-8修复】维度完成信号附带 warning（前端可 toast 提示）
                done_signal = {"dim": dim, "done": True}
                warnings_list = []
                if timeline_warning_stream:
                    warnings_list.append(timeline_warning_stream)
                if foreshadowing_warning_stream:
                    warnings_list.append(foreshadowing_warning_stream)
                if warnings_list:
                    done_signal['warning'] = '；'.join(warnings_list)
                yield f'data: {json.dumps(done_signal, ensure_ascii=False)}\n\n'

            # P3-11：所有维度生成完成后统一提交事务（事务性落库）
            try:
                db.session.commit()
            except Exception as commit_err:
                db.session.rollback()
                yield f'data: {json.dumps({"error": f"统一落库失败: {str(commit_err)[:200]}"}, ensure_ascii=False)}\n\n'

            yield 'data: [DONE]\n\n'
        except Exception as e:
            yield f'data: {json.dumps({"error": str(e)[:300]}, ensure_ascii=False)}\n\n'

    # stream_with_context 保持请求/应用上下文，避免生成器在 yield 后恢复时
    # 触发 "Working outside of application context"（DB 落库需要上下文）
    # 响应头修复：禁止 Cloudflare/Render 代理缓冲 SSE 流（同 ai_continue_batch_stream）
    resp = app.response_class(stream_with_context(generate()), mimetype='text/event-stream')
    resp.headers['Cache-Control'] = 'no-cache'
    resp.headers['X-Accel-Buffering'] = 'no'
    resp.headers['Connection'] = 'keep-alive'
    return resp

# ==== AI 共创 / 头脑风暴 ====

@app.route('/api/books/<book_id>/brainstorm', methods=['POST'])
def ai_brainstorm(book_id):
    """AI协同创作：用户给出一句话构思，AI返回多维度选项建议供用户选择"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.model if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    concept = request.json.get('concept', '')
    dimension = request.json.get('dimension', 'all')  # all/worldview/character/plot/locations/foreshadowing
    skill_pack_ids = request.json.get('skill_pack_ids', [])  # 支持技能包注入

    if not concept.strip():
        return jsonify({'error': '请输入一句话构思'}), 400

    bb = BookBible.query.filter_by(book_id=book_id).first()
    existing_context = ''
    if bb and bb.generated_summary:
        existing_context = bb.generated_summary[:2000]

    # 注入技能包提示词（brainstorm 相关）
    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['tomato_plan', 'one_line_concept', 'brainstorm', 'master_outline'])

    dimension_prompts = {
        'concept': '构思扩展：为这个构思设计3个不同方向的创意方案。每个方案包含：核心卖点、目标读者、主线冲突、独特亮点。每个方案100-200字。',
        'settings': '核心设定：为这个构思设计3套不同的核心规则方案。每套包含：世界观必须遵循的规则、人物能力限制、禁忌事项、特殊机制。每套100-200字。',
        'outline': '大纲方向：为这个构思设计3条不同的大纲方案。每条包含：核心主线、分卷规划（每卷目标）、关键转折点、高潮设计、结局走向。每条100-200字。',
        'worldview': '世界观设定：为这个构思设计3个不同方向的世界观方案。每个方案包含：世界名称、核心规则(力量体系/科技水平/社会结构)、独特亮点、潜在冲突源。每个方案100-200字。',
        'character': '人物设计：为这个构思设计3组主角方案。每组包含：主角姓名、身份背景、性格特征(3个标签)、核心动机、成长弧线方向、独特能力/特质。每组100-200字。',
        'plot': '剧情方向：为这个构思设计3条不同的主线剧情方向。每条包含：核心冲突、主线事件链(3-5个关键节点)、高潮设计、结局走向。每条100-200字。',
        'chapters': '章节规划：为这个构思设计3种不同的前5章开篇方案。每种包含：每章标题和核心内容概要(50字内)、开篇钩子设计、节奏安排。每种150-250字。',
        'locations': '地点设计：为这个构思设计3个关键地点。每个包含：地点名称、地理特征、文化氛围、潜在事件、与主线的关系。每个80-150字。',
        'foreshadowing': '伏笔设计：为这个构思设计3条可埋设的伏笔线索。每条包含：伏笔内容、埋设时机、预期回收章节范围、回收方式、对剧情的影响。每条80-150字。',
    }

    if dimension == 'all':
        task = '为这个构思生成完整的创作建议方案，包含以下维度，每个维度给出3个选项：\n'
        for dim, prompt in dimension_prompts.items():
            task += f'\n### {dim}\n{prompt}\n'
    else:
        task = dimension_prompts.get(dimension, dimension_prompts['worldview'])

    system_prompt = f"""你是资深网文创意策划师，服务于番茄小说/起点中文网。用户正在创作一部{book.book_type}，题材为{_get_genre_label(book)}。

用户的一句话构思：{concept}

{f'已有设定上下文：{existing_context}' if existing_context else '这是全新创作，暂无已有设定。'}

{skill_note}

请根据构思生成创作建议。严格按JSON格式输出，不要任何其他文字：
{{
  "concept_analysis": "对构思的分析(50字内)，指出核心卖点和潜在方向",
  "suggestions": {{
    "concept": [{{"title": "方案名", "description": "详细描述"}}],
    "settings": [{{"title": "方案名", "description": "详细描述"}}],
    "outline": [{{"title": "方案名", "description": "详细描述"}}],
    "worldview": [{{"title": "方案名", "description": "详细描述"}}],
    "character": [{{"title": "方案名", "description": "详细描述"}}],
    "plot": [{{"title": "方案名", "description": "详细描述"}}],
    "chapters": [{{"title": "方案名", "description": "详细描述"}}],
    "locations": [{{"title": "方案名", "description": "详细描述"}}],
    "foreshadowing": [{{"title": "方案名", "description": "详细描述"}}]
  }}
}}"""

    try:
        base = base_url.rstrip('/')
        if not base.endswith('/v1'):
            base += '/v1'
        resp = requests.post(f'{base}/chat/completions',
            headers=build_auth_headers(api_key),
            json={
                'model': model,
                'messages': [
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': task}
                ],
                'temperature': 0.8,
                'max_tokens': 8000,
                'response_format': {'type': 'json_object'}
            },
            timeout=180)
        result = resp.json()
        content = result['choices'][0]['message']['content']
        return jsonify(json.loads(content))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/books/<book_id>/ai-analyze-content', methods=['POST'])
@login_required
def ai_analyze_content(book_id):
    """AI分析作品内容，自动提取并填充构思、设定、大纲、世界观、人物、剧情、伏笔、地点等维度"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('recognition') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    # 获取作品所有章节内容（限制总字数避免超长）
    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    if not chapters:
        return jsonify({'error': '作品没有章节内容，无法分析'}), 400

    # 拼接章节内容，限制总长度
    full_text = ''
    max_chars = 12000  # 约 12000 中文字符
    for ch in chapters:
        segment = f'【{ch.title}】\n{(ch.content or "")[:2000]}\n\n'
        if len(full_text) + len(segment) > max_chars:
            remaining = max_chars - len(full_text)
            if remaining > 200:
                full_text += segment[:remaining]
            break
        full_text += segment

    system_prompt = f"""你是专业的小说分析师。请分析以下小说内容，提取并归纳各维度的设定信息。
严格按JSON格式输出，不要任何其他文字：
{{
  "concept": "一句话概括核心构思（30字内）",
  "key_rules": "核心设定规则：能力体系、限制、禁忌等（200字内）",
  "plot_design": "大纲：主线冲突、分卷规划、关键转折、结局走向（300字内）",
  "worldbuilding": "世界观：世界背景、力量体系、社会结构、地理概况（300字内）",
  "character_profiles": "主要人物档案：姓名、身份、性格、动机、关系（300字内）",
  "timeline": "剧情时间线：按顺序列出关键事件（200字内）",
  "foreshadowing": "伏笔线索：已发现或可能的伏笔（150字内）",
  "locations": "地点体系：三级分类（大区域/城市/场景），JSON格式",
  "generated_summary": "作品内容摘要（100字内）"
}}"""

    try:
        base = base_url.rstrip('/')
        if not base.endswith('/v1'):
            base += '/v1'
        resp = requests.post(f'{base}/chat/completions',
            headers=build_auth_headers(api_key),
            json={
                'model': model,
                'messages': [
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': f'作品标题：{book.title}\n\n以下是作品内容：\n\n{full_text}'}
                ],
                'temperature': 0.3,
                'max_tokens': 4000,
                'response_format': {'type': 'json_object'}
            },
            timeout=180)
        result = resp.json()
        content = result['choices'][0]['message']['content']
        analysis = json.loads(content)

        # 更新或创建 BookBible
        bb = BookBible.query.filter_by(book_id=book_id).first()
        if not bb:
            bb = BookBible(book_id=book_id)
            db.session.add(bb)

        # 只更新非空字段，不覆盖用户已写内容
        fields = ['concept', 'key_rules', 'plot_design', 'worldbuilding',
                  'character_profiles', 'timeline', 'foreshadowing', 'locations', 'generated_summary']
        updated_fields = []
        for field in fields:
            raw_val = analysis.get(field, '')
            # AI 可能返回 dict/list 等结构化数据，统一转为字符串
            if isinstance(raw_val, (dict, list)):
                new_val = json.dumps(raw_val, ensure_ascii=False, indent=2)
            else:
                new_val = str(raw_val).strip() if raw_val else ''
            if new_val:
                existing_val = getattr(bb, field, '') or ''
                if existing_val:
                    # 已有内容则追加
                    setattr(bb, field, f'{existing_val}\n\n【AI识别】\n{new_val}')
                else:
                    setattr(bb, field, new_val)
                updated_fields.append(field)

        bb.last_synced_at = datetime.now(timezone.utc)
        db.session.commit()

        return jsonify({
            'success': True,
            'updated_fields': updated_fields,
            'bible': bb.to_dict()
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/books/<book_id>/ai-analyze-dimension', methods=['POST'])
@login_required
def ai_analyze_dimension(book_id):
    """AI分析作品内容，只识别并填充指定维度"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    dimension = data.get('dimension', '')
    if not dimension:
        return jsonify({'error': '缺少 dimension 参数'}), 400

    # 维度 → bible字段 映射
    dim_field_map = {
        'concept': 'concept',
        'settings': 'key_rules',
        'outline': 'plot_design',
        'worldview': 'worldbuilding',
        'characters': 'character_profiles',
        'plot': 'timeline',
        'foreshadowing': 'foreshadowing',
        'locations': 'locations',
    }
    field = dim_field_map.get(dimension)
    if not field:
        return jsonify({'error': f'未知维度: {dimension}'}), 400

    dim_labels = {
        'concept': '构思', 'settings': '设定', 'outline': '大纲',
        'worldview': '世界观', 'characters': '人物', 'plot': '剧情',
        'foreshadowing': '伏笔', 'locations': '地点',
    }
    dim_label = dim_labels.get(dimension, dimension)

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('recognition') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    if not chapters:
        return jsonify({'error': '作品没有章节内容，无法分析'}), 400

    full_text = ''
    max_chars = 12000
    for ch in chapters:
        segment = f'【{ch.title}】\n{(ch.content or "")[:2000]}\n\n'
        if len(full_text) + len(segment) > max_chars:
            remaining = max_chars - len(full_text)
            if remaining > 200:
                full_text += segment[:remaining]
            break
        full_text += segment

    dim_prompts = {
        'concept': '一句话概括核心构思（30字内）',
        'key_rules': '核心设定规则：能力体系、限制、禁忌等（200字内）',
        'plot_design': '大纲：主线冲突、分卷规划、关键转折、结局走向（300字内）',
        'worldbuilding': '世界观：世界背景、力量体系、社会结构、地理概况（300字内）',
        'character_profiles': '主要人物档案：姓名、身份、性格、动机、关系（300字内）',
        'timeline': '剧情时间线：按顺序列出关键事件（200字内）',
        'foreshadowing': '伏笔线索：已发现或可能的伏笔（150字内）',
        'locations': '地点体系：三级分类（大区域/城市/场景），JSON格式',
    }

    # agent 协同：读取 bible 其他维度作为已知上下文，让识别结果与已确认维度保持一致
    bb = BookBible.query.filter_by(book_id=book_id).first()
    known_ctx_parts = []
    if bb:
        dim_label_map = {'concept': '构思', 'key_rules': '设定/规则', 'worldbuilding': '世界观',
                         'character_profiles': '人物', 'plot_design': '大纲', 'timeline': '剧情',
                         'foreshadowing': '伏笔', 'locations': '地点'}
        for f, lbl in dim_label_map.items():
            v = getattr(bb, f, '') or ''
            if v.strip() and f != field:  # 排除当前维度自身
                known_ctx_parts.append(f'【{lbl}（已确认）】\n{v[:600]}')
    known_ctx = '\n\n'.join(known_ctx_parts) if known_ctx_parts else '（暂无其他维度参考）'

    system_prompt = f"""你是专业的小说分析师，正在与其他维度分析师协同工作。
请分析以下小说内容，提取并归纳“{dim_label}”维度的设定信息。

【已确认的其他维度设定】（识别结果必须与这些维度保持一致，不可矛盾）
{known_ctx}

严格按JSON格式输出，不要任何其他文字：
{{
  "{field}": "{dim_prompts.get(field, dim_label)}"
}}"""

    try:
        base = base_url.rstrip('/')
        if not base.endswith('/v1'):
            base += '/v1'

        # 构建请求体（response_format 某些LLM不支持，捕获后重试）
        req_body = {
            'model': model,
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': f'作品标题：{book.title}\n\n以下是作品内容：\n\n{full_text}'}
            ],
            'temperature': 0.3,
            'max_tokens': 2000,
            'response_format': {'type': 'json_object'}
        }

        # 第一次尝试带 response_format；不支持则去掉重试
        resp = requests.post(f'{base}/chat/completions',
            headers=build_auth_headers(api_key),
            json=req_body,
            timeout=90)

        # 若返回400且与response_format相关，去掉该参数重试
        if resp.status_code == 400:
            err_body = resp.json() if resp.headers.get('content-type', '').startswith('application/json') else {}
            err_msg = str(err_body.get('error', '')).lower()
            if 'response_format' in err_msg or 'json_object' in err_msg or 'unrecognized' in err_msg:
                req_body.pop('response_format', None)
                resp = requests.post(f'{base}/chat/completions',
                    headers=build_auth_headers(api_key),
                    json=req_body,
                    timeout=90)

        if resp.status_code != 200:
            try:
                err_detail = resp.json().get('error', {}).get('message', '') or resp.text[:300]
            except Exception:
                err_detail = resp.text[:300]
            return jsonify({'error': f'AI调用失败({resp.status_code}): {err_detail}'}), 500

        result = resp.json()
        if 'choices' not in result or not result['choices']:
            return jsonify({'error': f'AI返回异常: {str(result)[:300]}'}), 500
        content = result['choices'][0]['message']['content']

        # JSON 解析容错：提取第一个 {...} 块
        try:
            analysis = json.loads(content)
        except (json.JSONDecodeError, ValueError):
            import re as _re_json
            m = _re_json.search(r'\{[\s\S]*\}', content)
            if m:
                analysis = json.loads(m.group(0))
            else:
                return jsonify({'error': f'AI返回非JSON格式: {content[:200]}'}), 500

        bb = BookBible.query.filter_by(book_id=book_id).first()
        if not bb:
            bb = BookBible(book_id=book_id)
            db.session.add(bb)

        raw_val = analysis.get(field, '')
        if isinstance(raw_val, (dict, list)):
            new_val = json.dumps(raw_val, ensure_ascii=False, indent=2)
        else:
            new_val = str(raw_val).strip() if raw_val else ''

        if new_val:
            existing_val = getattr(bb, field, '') or ''
            if existing_val:
                setattr(bb, field, f'{existing_val}\n\n【AI识别】\n{new_val}')
            else:
                setattr(bb, field, new_val)

        bb.last_synced_at = datetime.now(timezone.utc)
        db.session.commit()

        return jsonify({
            'success': True,
            'dimension': dimension,
            'field': field,
            'value': getattr(bb, field, ''),
            'bible': bb.to_dict()
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/books/<book_id>/ai-analyze-character', methods=['POST'])
@login_required
def ai_analyze_character(book_id):
    """AI从章节内容中识别单个角色信息，或识别全部角色列表"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    character_name = data.get('character_name', '')  # 指定角色名，为空则识别全部角色列表

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('recognition') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    if not chapters:
        return jsonify({'error': '作品没有章节内容，无法分析'}), 400

    full_text = ''
    max_chars = 12000
    for ch in chapters:
        segment = f'【{ch.title}】\n{(ch.content or "")[:2000]}\n\n'
        if len(full_text) + len(segment) > max_chars:
            remaining = max_chars - len(full_text)
            if remaining > 200:
                full_text += segment[:remaining]
            break
        full_text += segment

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)

    existing_chars = []
    try:
        parsed = json.loads(bb.character_profiles or '[]')
        if isinstance(parsed, list):
            existing_chars = parsed
    except:
        pass

    if character_name:
        # 识别指定角色的详细信息
        system_prompt = f"""你是专业的小说分析师。请从以下小说内容中，提取角色“{character_name}”的详细档案。
严格按JSON格式输出，不要任何其他文字：
{{
  "name": "{character_name}",
  "role": "主角/配角/反派/路人 等角色定位",
  "identity": "身份职业",
  "personality": "性格特征（2-3句）",
  "motivation": "核心动机和目标",
  "background": "背景故事（2-3句）",
  "relationships": "与其他角色的关系（如：与XX是师徒，与XX是敌对）",
  "abilities": "拥有的能力/功法/特长",
  "items": "持有的重要物品/装备"
}}"""
    else:
        # 识别全部角色列表
        existing_names = [c.get('name', '') for c in existing_chars if isinstance(c, dict)]
        existing_note = f'\n已有角色：{", ".join(existing_names)}' if existing_names else ''
        system_prompt = f"""你是专业的小说分析师。请从以下小说内容中，识别所有重要角色（出现3次以上或有台词的角色）。
{existing_note}
严格按JSON数组格式输出，不要任何其他文字：
[{{"name": "角色名", "role": "主角/配角/反派"}}]

注意：只输出数组，不要其他文字。"""

    try:
        base = base_url.rstrip('/')
        if not base.endswith('/v1'):
            base += '/v1'
        resp = requests.post(f'{base}/chat/completions',
            headers=build_auth_headers(api_key),
            json={
                'model': model,
                'messages': [
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': f'作品标题：{book.title}\n\n以下是作品内容：\n\n{full_text}'}
                ],
                'temperature': 0.3,
                'max_tokens': 2000,
                'response_format': {'type': 'json_object' if character_name else 'json_object'}
            },
            timeout=120)
        result = resp.json()
        content = result['choices'][0]['message']['content']
        analysis = json.loads(content)

        if character_name:
            # 更新或添加单个角色
            char_data = analysis if isinstance(analysis, dict) else {}
            char_data['name'] = character_name
            found = False
            for i, c in enumerate(existing_chars):
                if isinstance(c, dict) and c.get('name') == character_name:
                    existing_chars[i] = {**c, **char_data}
                    found = True
                    break
            if not found:
                existing_chars.append(char_data)
            bb.character_profiles = json.dumps(existing_chars, ensure_ascii=False, indent=2)
        else:
            # 合并角色列表
            new_chars = analysis if isinstance(analysis, list) else (analysis.get('characters', []) if isinstance(analysis, dict) else [])
            existing_names = {c.get('name', '') for c in existing_chars if isinstance(c, dict)}
            for nc in new_chars:
                if isinstance(nc, dict) and nc.get('name', '') not in existing_names:
                    existing_chars.append(nc)
            bb.character_profiles = json.dumps(existing_chars, ensure_ascii=False, indent=2)

        bb.last_synced_at = datetime.now(timezone.utc)
        db.session.commit()

        return jsonify({
            'success': True,
            'character': char_data if character_name else None,
            'characters': existing_chars,
            'bible': bb.to_dict()
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/books/<book_id>/ai-analyze-plot-volume', methods=['POST'])
@login_required
def ai_analyze_plot_volume(book_id):
    """AI识别指定卷的剧情大纲。
    优化版：数据源从单一章节内容扩展为 设定+大纲+人物+规则+章节+动态文件，
    输出增加 nodes 情节节点字段，与 ai_outline_volume 输出结构一致。
    这是相互提供资料数据的过程：识别结果会回流到 timeline 供其他维度使用。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    volume_id = data.get('volume_id', '')
    volume_title = data.get('volume_title', '')
    skill_pack_ids = data.get('skill_pack_ids', [])

    bb = BookBible.query.filter_by(book_id=book_id).first()

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('recognition') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    # ===== 1. 收集该卷章节内容 =====
    all_chapters = Chapter.query.filter_by(book_id=book_id).order_by(Chapter.order_index).all()
    volume_chapters = []
    if volume_id:
        # 优先：parent_id 关联（标准结构）
        volume_chapters = [c for c in all_chapters if not c.is_volume and c.parent_id == volume_id]
        # 回退：顺序遍历法（兼容卷紧挨子章节之前的旧数据）
        if not volume_chapters:
            collecting = False
            for ch in all_chapters:
                if ch.id == volume_id:
                    collecting = True
                    continue
                if collecting:
                    if ch.is_volume:
                        break
                    volume_chapters.append(ch)
    else:
        volume_chapters = [c for c in all_chapters if not c.is_volume]

    # 章节内容组装（优先用 summary，其次前 800 字 + 末 200 字）
    chapter_text = ''
    max_chars = 12000
    for ch in volume_chapters:
        ch_content = (ch.content or '')
        # 优先用章节摘要
        if getattr(ch, 'summary', None) and ch.summary:
            segment = f'【{ch.title}】{ch.summary[:500]}\n'
        elif len(ch_content) > 1000:
            segment = f'【{ch.title}】{ch_content[:800]}…{ch_content[-200:]}\n'
        else:
            segment = f'【{ch.title}】{ch_content}\n'
        if len(chapter_text) + len(segment) > max_chars:
            remaining = max_chars - len(chapter_text)
            if remaining > 200:
                chapter_text += segment[:remaining]
            break
        chapter_text += segment

    # ===== 2. 收集多维度上下文（相互提供资料数据） =====
    ctx_parts = []
    if bb:
        # 大纲维度（五幕式总纲）
        if bb.plot_design:
            ctx_parts.append(f'【五幕式总纲（本卷应在此弧线内）】\n{bb.plot_design[:2000]}')
        # 设定维度（世界观+规则）
        if bb.worldbuilding:
            ctx_parts.append(f'【世界观设定】\n{bb.worldbuilding[:1000]}')
        if bb.key_rules:
            ctx_parts.append(f'【核心规则（金手指/能力限制，识别时不可违反）】\n{bb.key_rules[:800]}')
        # 人物及关系维度
        if bb.character_profiles:
            ctx_parts.append(f'【人物档案】\n{bb.character_profiles[:1000]}')
        # 已有该卷剧情（若有，作为参考而非覆盖）
        if bb.timeline:
            try:
                existing_vols = json.loads(bb.timeline)
                if isinstance(existing_vols, list):
                    # 找到该卷的已有数据
                    for ev in existing_vols:
                        ev_vid = str(ev.get('volume_id', ''))
                        ev_vol = str(ev.get('volume', ''))
                        if (volume_id and ev_vid == str(volume_id)) or (volume_title and ev_vol == volume_title):
                            ctx_parts.append(f'【该卷已有剧情（参考，可补充完善）】\n{json.dumps(ev, ensure_ascii=False)[:600]}')
                            break
            except (json.JSONDecodeError, ValueError):
                pass

    # ===== 3. 从动态文件补充数据 =====
    dyn_memories = DynamicMemory.query.filter_by(book_id=book_id).all()
    for dm in dyn_memories:
        if dm.category in ('narrative_engine', 'plot_progress', 'timeline') and dm.content:
            ctx_parts.append(f'【动态文件-{dm.category}】\n{dm.content[:800]}')
            break  # 只取一份，避免过多

    extra_ctx = '\n\n'.join(ctx_parts[-5:])  # 最多 5 块上下文，避免超长

    # 技能包提示
    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['volume_breakdown', 'chapter_plan', 'tomato_outline'], mode='agent')

    if not volume_chapters and not extra_ctx:
        return jsonify({'error': '该卷没有章节内容，也没有可参考的设定'}), 400

    vol_label = volume_title or '全部章节'

    system_prompt = f"""你是番茄小说金番作者级别的剧情分析师。请综合【设定/大纲/人物/规则/章节内容/动态文件】多维度数据，识别“{vol_label}”的剧情大纲和情节节点。

【多维度上下文（相互提供资料数据）】
{extra_ctx or '（暂无设定参考，仅依据章节内容识别）'}

【识别要求】
1. 识别出的剧情必须与【五幕式总纲】中该卷的弧线一致，若有偏差在 main_plot 中标注
2. 识别人物互动时参考【人物档案】，确保角色名字和行为准确
3. 识别金手指/能力使用时参考【核心规则】，违反规则的标注为"待修正"
4. 结合【动态文件】中的叙事记录，补充章节内容未体现的关键事件和伏笔

严格按JSON格式输出（不要任何其他文字）：
{{
  "volume": "{vol_label}",
  "main_plot": "该卷主线剧情概述（100-200字，标注与总纲的偏差）",
  "core_conflict": "核心冲突",
  "emotion_driver": "情感驱动",
  "key_events": ["关键事件1", "关键事件2", "关键事件3"],
  "turning_points": ["转折点1", "转折点2"],
  "climax": "高潮场景描述",
  "ending": "该卷结尾状态/钩子",
  "foreshadowing": ["埋设的伏笔"],
  "nodes": [
    {{"title": "节点1", "chapters": "1-10", "type": "M", "summary": "概要", "cool_type": "爽点类型"}}
  ]
}}

【情节节点识别要求】
- 每卷识别 5-8 个情节节点
- 章型：M主线/C角色/W世界观/D日常/F伏笔
- 章型配额参考：M主线50%/C角色10%/W世界观10%/D日常20%/F伏笔10%
- 节点章节范围不重叠，覆盖整卷
- 小故事闭环：新事件→困难→金手指破局→暴露新信息→打脸收尾→钩子（5-8章）

{skill_note}"""

    user_prompt = f'作品标题：{book.title}\n卷名：{vol_label}\n\n以下是该卷章节内容：\n\n{chapter_text or "（无章节内容，请根据设定推断）"}'

    try:
        base = base_url.rstrip('/')
        if not base.endswith('/v1'):
            base += '/v1'

        # 构建请求体（response_format 某些LLM不支持，捕获后重试）
        req_body = {
            'model': model,
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_prompt}
            ],
            'temperature': 0.3,
            'max_tokens': 4000,
            'response_format': {'type': 'json_object'}
        }

        resp = requests.post(f'{base}/chat/completions',
            headers=build_auth_headers(api_key),
            json=req_body,
            timeout=90)

        # 若返回400且与response_format相关，去掉该参数重试
        if resp.status_code == 400:
            err_body = resp.json() if resp.headers.get('content-type', '').startswith('application/json') else {}
            err_msg = str(err_body.get('error', '')).lower()
            if 'response_format' in err_msg or 'json_object' in err_msg or 'unrecognized' in err_msg:
                req_body.pop('response_format', None)
                resp = requests.post(f'{base}/chat/completions',
                    headers=build_auth_headers(api_key),
                    json=req_body,
                    timeout=90)

        if resp.status_code != 200:
            try:
                err_detail = resp.json().get('error', {}).get('message', '') or resp.text[:300]
            except Exception:
                err_detail = resp.text[:300]
            return jsonify({'error': f'AI调用失败({resp.status_code}): {err_detail}'}), 500

        result = resp.json()
        if 'choices' not in result or not result['choices']:
            return jsonify({'error': f'AI返回异常: {str(result)[:300]}'}), 500
        content = result['choices'][0]['message']['content']

        # JSON 解析容错：使用健壮解析函数提取对象
        analysis, parse_err = _extract_json_from_llm(content, expect='object')
        if analysis is None:
            return jsonify({'error': f'AI返回非JSON格式: {content[:200]}', 'parse_error': parse_err}), 500

        # 存储到 timeline 字段（深度合并：保留人工编辑字段，更新 AI 识别字段）
        if not bb:
            bb = BookBible(book_id=book_id)
            db.session.add(bb)

        volumes_data = []
        try:
            parsed = json.loads(bb.timeline or '[]')
            if isinstance(parsed, list):
                volumes_data = parsed
        except:
            pass

        # 补全 volume_id 和 volume
        vol_data = analysis
        if volume_id:
            vol_data['volume_id'] = volume_id
        vol_data['volume'] = vol_label
        # 补全 volume_index
        if 'volume_index' not in vol_data:
            vol_data['volume_index'] = _extract_volume_index(vol_label) or (len(volumes_data) + 1)

        # 深度合并：找到已有卷，保留人工编辑的 nodes（如果新数据没有 nodes），其他字段用新数据
        found_idx = -1
        for i, v in enumerate(volumes_data):
            if not isinstance(v, dict):
                continue
            ev_vid = str(v.get('volume_id', ''))
            ev_vol = str(v.get('volume', ''))
            if (volume_id and ev_vid == str(volume_id)) or (volume_title and ev_vol == volume_title):
                found_idx = i
                break
        if found_idx >= 0:
            existing = volumes_data[found_idx]
            # 保留人工编辑的 nodes（新数据 nodes 为空或缺失时）
            if not vol_data.get('nodes') and existing.get('nodes'):
                vol_data['nodes'] = existing['nodes']
            # 保留人工编辑的字段（raw_text 等）
            for k in ('raw_text',):
                if k in existing and k not in vol_data:
                    vol_data[k] = existing[k]
            volumes_data[found_idx] = {**existing, **vol_data}
        else:
            volumes_data.append(vol_data)

        # 按 volume_index 排序
        volumes_data.sort(key=lambda v: int(v.get('volume_index', 0) or _extract_volume_index(v.get('volume', v.get('volume_id', '0'))) or 0))

        bb.timeline = json.dumps(volumes_data, ensure_ascii=False, indent=2)
        bb.last_synced_at = datetime.now(timezone.utc)
        db.session.commit()

        return jsonify({
            'success': True,
            'volume_data': vol_data,
            'volumes': volumes_data,
            'bible': bb.to_dict()
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def _get_volume_chapters_ordered(book_id, volume_id):
    """返回指定卷下的非卷章节列表（按 order_index 排序）。
    优先用 parent_id 关联（标准结构）；若该卷无任何 parent_id 指向它的章节，
    回退到顺序遍历法（兼容旧数据：卷记录紧挨在其子章节之前）。
    volume_id 为空则返回全部非卷章节。"""
    all_chapters = Chapter.query.filter_by(book_id=book_id).order_by(Chapter.order_index).all()
    if not volume_id:
        return [c for c in all_chapters if not c.is_volume]
    # 优先：parent_id 关联（标准结构）
    vol_chs = [c for c in all_chapters if not c.is_volume and c.parent_id == volume_id]
    if vol_chs:
        return vol_chs
    # 回退：顺序遍历法（兼容卷紧挨子章节之前的旧数据）
    collecting = False
    for ch in all_chapters:
        if ch.id == volume_id:
            collecting = True
            continue
        if collecting:
            if ch.is_volume:
                break
            vol_chs.append(ch)
    return vol_chs

def _collect_volume_chapters(book_id, volume_id):
    """收集指定卷的章节内容文本。volume_id 为空则取全部非卷章节。
    章节检索逻辑见 _get_volume_chapters_ordered（parent_id 优先，回退顺序遍历）。"""
    volume_chapters = _get_volume_chapters_ordered(book_id, volume_id)

    chapter_text = ''
    max_chars = 12000
    for ch in volume_chapters:
        ch_content = (ch.content or '')
        if getattr(ch, 'summary', None) and ch.summary:
            segment = f'【{ch.title}】{ch.summary[:500]}\n'
        elif len(ch_content) > 1000:
            segment = f'【{ch.title}】{ch_content[:800]}…{ch_content[-200:]}\n'
        else:
            segment = f'【{ch.title}】{ch_content}\n'
        if len(chapter_text) + len(segment) > max_chars:
            remaining = max_chars - len(chapter_text)
            if remaining > 200:
                chapter_text += segment[:remaining]
            break
        chapter_text += segment
    return chapter_text, len(volume_chapters)

def _collect_dimension_source(bb, volume_title=''):
    """当无章节时，从设定/大纲/剧情等主要维度收集基础数据作为识别来源。
    返回 (source_text, source_label)。完整注入各维度，不截断。"""
    if not bb:
        return '', ''
    parts = []
    if bb.concept and bb.concept.strip():
        parts.append(f'【构思】\n{bb.concept.strip()}')
    if bb.key_rules and bb.key_rules.strip():
        parts.append(f'【设定/核心规则】\n{bb.key_rules.strip()}')
    if bb.worldbuilding and bb.worldbuilding.strip():
        parts.append(f'【世界观】\n{bb.worldbuilding.strip()}')
    if bb.plot_design and bb.plot_design.strip():
        parts.append(f'【大纲/总纲】\n{bb.plot_design.strip()}')
    if bb.timeline and bb.timeline.strip():
        # timeline 可能是 JSON（卷列表）或纯文本
        tl_text = bb.timeline.strip()
        try:
            tl_parsed = json.loads(tl_text)
            if isinstance(tl_parsed, list):
                tl_lines = []
                for v in tl_parsed:
                    if isinstance(v, dict):
                        vol_name = v.get('volume', '')
                        vol_content = v.get('content', '') or v.get('outline', '') or v.get('plot', '')
                        if vol_content:
                            tl_lines.append(f'卷“{vol_name}”: {vol_content}')
                if tl_lines:
                    tl_text = '\n'.join(tl_lines)
        except (json.JSONDecodeError, ValueError):
            pass
        parts.append(f'【剧情/时间线】\n{tl_text}')
    if bb.character_profiles and bb.character_profiles.strip():
        # 人物维度转自然语言
        cp = bb.character_profiles.strip()
        if cp.startswith('['):
            try:
                from chat_collab_bp import _character_profiles_to_text
                cp = _character_profiles_to_text(cp)
            except Exception:
                pass
        parts.append(f'【人物档案】\n{cp}')
    if bb.foreshadowing and bb.foreshadowing.strip():
        parts.append(f'【伏笔】\n{bb.foreshadowing.strip()}')
    if bb.locations and bb.locations.strip():
        parts.append(f'【地点】\n{bb.locations.strip()}')
    if bb.style_guide and bb.style_guide.strip():
        parts.append(f'【文风】\n{bb.style_guide.strip()}')

    source_text = '\n\n'.join(parts)
    label = '设定/大纲/剧情维度' if source_text else ''
    if volume_title and source_text:
        label = f'设定/大纲/剧情维度（针对“{volume_title}”）'
    return source_text, label

def _get_volume_list(bb):
    """从 bible.timeline 解析卷列表，返回 [{volume_id, volume, volume_index}]"""
    if not bb or not bb.timeline:
        return []
    try:
        parsed = json.loads(bb.timeline)
        if isinstance(parsed, list):
            return [v for v in parsed if isinstance(v, dict)]
    except (json.JSONDecodeError, ValueError):
        pass
    return []

def _upsert_volume_entry(bb, field_name, entry):
    """在 bible.<field_name> 的 JSON 数组中按 volume_id/volume upsert 一条记录。
    P1-5: 对 inventory 字段，若旧值是纯文本（非JSON），先迁移为 [{volume:'', data:旧文本}] 再 upsert。"""
    data_list = []
    raw = getattr(bb, field_name, '') or ''
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, list):
            data_list = parsed
        elif isinstance(parsed, str) and parsed.strip():
            # 旧纯文本：迁移为单元素数组
            data_list = [{'volume': '历史数据', 'volume_id': '', 'data': parsed}]
        else:
            data_list = []
    except (json.JSONDecodeError, ValueError):
        # P1-5: 纯文本兜底——把旧文本包成单元素数组，避免静默丢弃
        if raw.strip():
            data_list = [{'volume': '历史数据', 'volume_id': '', 'data': raw}]
        else:
            data_list = []

    vid = str(entry.get('volume_id', ''))
    vname = str(entry.get('volume', ''))
    found_idx = -1
    for i, v in enumerate(data_list):
        if not isinstance(v, dict):
            continue
        ev_vid = str(v.get('volume_id', ''))
        ev_vol = str(v.get('volume', ''))
        if (vid and ev_vid == vid) or (vname and ev_vol == vname):
            found_idx = i
            break
    if found_idx >= 0:
        data_list[found_idx] = {**data_list[found_idx], **entry}
    else:
        data_list.append(entry)
    data_list.sort(key=lambda v: int(v.get('volume_index', 0) or _extract_volume_index(v.get('volume', v.get('volume_id', '0'))) or 0))
    setattr(bb, field_name, json.dumps(data_list, ensure_ascii=False, indent=2))
    return data_list

@app.route('/api/books/<book_id>/ai-analyze-character-volume', methods=['POST'])
@login_required
def ai_analyze_character_volume(book_id):
    """AI识别指定卷的人物档案。按卷分析章节内容，识别人物并写入 character_volumes。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    volume_id = data.get('volume_id', '')
    volume_title = data.get('volume_title', '')
    skill_pack_ids = data.get('skill_pack_ids', [])

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('recognition') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    chapter_text, ch_count = _collect_volume_chapters(book_id, volume_id)
    source_label = ''
    if not chapter_text or ch_count == 0:
        # 无章节时，从设定/大纲/剧情维度提取基础数据
        chapter_text, source_label = _collect_dimension_source(bb, volume_title or '全部章节')
        if not chapter_text:
            return jsonify({'error': '该卷暂无章节，且设定/大纲/剧情维度也为空，无法识别人物。请先填写设定、大纲或剧情维度。'}), 400

    # 上下文：全局人物档案 + 设定 + 该卷剧情
    ctx_parts = []
    if bb.character_profiles:
        ctx_parts.append(f'【全局人物档案（参考，避免重复识别）】\n{bb.character_profiles[:1000]}')
    if bb.key_rules:
        ctx_parts.append(f'【核心规则】\n{bb.key_rules[:600]}')
    if bb.worldbuilding:
        ctx_parts.append(f'【世界观设定】\n{bb.worldbuilding[:600]}')
    # 该卷剧情
    if bb.timeline:
        try:
            vols = json.loads(bb.timeline)
            if isinstance(vols, list):
                for v in vols:
                    if isinstance(v, dict) and (str(v.get('volume_id', '')) == str(volume_id) or v.get('volume') == volume_title):
                        ctx_parts.append(f'【该卷剧情（参考）】\n{(v.get("main_plot") or "")[:500]}')
                        break
        except (json.JSONDecodeError, ValueError):
            pass
    extra_ctx = '\n\n'.join(ctx_parts)

    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['character_cognition', 'tomato_character'], mode='agent')

    vol_label = volume_title or '全部章节'
    system_prompt = f"""你是专业的小说分析师。请从以下“{vol_label}”的章节内容中，识别本卷出现的所有重要角色（出现2次以上或有台词的角色）。

{extra_ctx and f"【已有参考】{chr(10)}{extra_ctx}" or ""}

严格按JSON对象格式输出（不要任何其他文字）：
{{
  "volume": "{vol_label}",
  "characters": [
    {{
      "name": "角色名",
      "role": "主角/配角/反派/路人",
      "identity": "身份职业",
      "personality": "性格特征（1-2句）",
      "motivation": "本卷中的动机",
      "relationships": "本卷中与其他角色的关系",
      "abilities": "本卷中使用的能力/功法",
      "items": "本卷中持有的重要物品",
      "arc": "本卷中的角色弧线/变化"
    }}
  ]
}}

{skill_note}"""

    user_prompt = f'作品标题：{book.title}\n卷名：{vol_label}\n{source_label and f"（数据来源：{source_label}）" or ""}\n\n以下是该卷内容：\n\n{chapter_text or "（无内容，请根据设定推断）"}'

    content, err = _call_llm(
        [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
        max_tokens=3000, temperature=0.3, task_type='recognition'
    )
    if err:
        return jsonify({'error': err}), 500

    try:
        analysis = json.loads(content)
    except (json.JSONDecodeError, ValueError):
        import re as _re_cv
        m = _re_cv.search(r'\{[\s\S]*\}', content)
        if m:
            analysis = json.loads(m.group())
        else:
            return jsonify({'error': 'AI返回格式无法解析', 'raw': content[:300]}), 500

    chars = analysis.get('characters', []) if isinstance(analysis, dict) else []

    entry = {
        'volume_id': volume_id,
        'volume': vol_label,
        'volume_index': _extract_volume_index(vol_label) or 0,
        'characters': chars,
    }
    data_list = _upsert_volume_entry(bb, 'character_volumes', entry)
    bb.last_synced_at = datetime.now(timezone.utc)
    db.session.commit()

    return jsonify({
        'success': True,
        'volume_data': entry,
        'character_volumes': data_list,
        'bible': bb.to_dict()
    })

@app.route('/api/books/<book_id>/ai-analyze-inventory-volume', methods=['POST'])
@login_required
def ai_analyze_inventory_volume(book_id):
    """AI识别指定卷的物资库。按卷分析章节内容，识别势力/角色拥有的物品、功法、法宝、境界等。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    volume_id = data.get('volume_id', '')
    volume_title = data.get('volume_title', '')
    skill_pack_ids = data.get('skill_pack_ids', [])

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('recognition') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    chapter_text, ch_count = _collect_volume_chapters(book_id, volume_id)
    source_label = ''
    if not chapter_text or ch_count == 0:
        # 无章节时，从设定/大纲/剧情维度提取基础数据
        chapter_text, source_label = _collect_dimension_source(bb, volume_title or '全部章节')
        if not chapter_text:
            return jsonify({'error': '该卷暂无章节，且设定/大纲/剧情维度也为空，无法识别物资。请先填写设定、大纲或剧情维度。'}), 400

    # 上下文
    ctx_parts = []
    if bb.character_profiles:
        ctx_parts.append(f'【人物档案（识别持有者）】\n{bb.character_profiles[:1000]}')
    if bb.key_rules:
        ctx_parts.append(f'【核心规则（能力体系/境界划分）】\n{bb.key_rules[:800]}')
    if bb.worldbuilding:
        ctx_parts.append(f'【世界观设定（势力格局）】\n{bb.worldbuilding[:600]}')
    if bb.timeline:
        try:
            vols = json.loads(bb.timeline)
            if isinstance(vols, list):
                for v in vols:
                    if isinstance(v, dict) and (str(v.get('volume_id', '')) == str(volume_id) or v.get('volume') == volume_title):
                        ctx_parts.append(f'【该卷剧情（参考）】\n{(v.get("main_plot") or "")[:500]}')
                        break
        except (json.JSONDecodeError, ValueError):
            pass
    extra_ctx = '\n\n'.join(ctx_parts)

    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['lock_facts', 'tomato_setting'], mode='agent')

    vol_label = volume_title or '全部章节'
    system_prompt = f"""你是专业的小说世界观分析师。请从以下“{vol_label}”的章节内容中，识别本卷出现的所有势力及角色拥有的物资。
物资类型包括：物品、功法、法宝、境界、灵宠、领地、资源等。

{extra_ctx and f"【已有参考】{chr(10)}{extra_ctx}" or ""}

严格按JSON对象格式输出（不要任何其他文字）：
{{
  "volume": "{vol_label}",
  "items": [
    {{
      "owner": "持有者（角色名/势力名）",
      "owner_type": "角色/势力",
      "name": "物资名称",
      "category": "物品/功法/法宝/境界/灵宠/领地/资源/其他",
      "description": "描述（来源、能力、效果）",
      "status": "获得/持有/失去/消耗",
      "chapter": "首次出现章节"
    }}
  ],
  "realms": [
    {{
      "character": "角色名",
      "realm": "当前境界",
      "progress": "修炼进度/突破节点"
    }}
  ]
}}

{skill_note}"""

    user_prompt = f'作品标题：{book.title}\n卷名：{vol_label}\n{source_label and f"（数据来源：{source_label}）" or ""}\n\n以下是该卷内容：\n\n{chapter_text or "（无内容，请根据设定推断）"}'

    content, err = _call_llm(
        [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
        max_tokens=3000, temperature=0.3, task_type='recognition'
    )
    if err:
        return jsonify({'error': err}), 500

    try:
        analysis = json.loads(content)
    except (json.JSONDecodeError, ValueError):
        import re as _re_iv
        m = _re_iv.search(r'\{[\s\S]*\}', content)
        if m:
            analysis = json.loads(m.group())
        else:
            return jsonify({'error': 'AI返回格式无法解析', 'raw': content[:300]}), 500

    items = analysis.get('items', []) if isinstance(analysis, dict) else []
    realms = analysis.get('realms', []) if isinstance(analysis, dict) else []

    entry = {
        'volume_id': volume_id,
        'volume': vol_label,
        'volume_index': _extract_volume_index(vol_label) or 0,
        'items': items,
        'realms': realms,
    }
    data_list = _upsert_volume_entry(bb, 'inventory', entry)
    bb.last_synced_at = datetime.now(timezone.utc)
    db.session.commit()

    return jsonify({
        'success': True,
        'volume_data': entry,
        'inventory': data_list,
        'bible': bb.to_dict()
    })

@app.route('/api/books/<book_id>/ai-analyze-dynamic-volume', methods=['POST'])
@login_required
def ai_analyze_dynamic_volume(book_id):
    """AI识别指定卷的动态文件分类。按卷汇总章节内容，生成该卷的动态摘要（人物/事件/时间/地点/势力/伏笔/境界/关系）。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    volume_id = data.get('volume_id', '')
    volume_title = data.get('volume_title', '')
    skill_pack_ids = data.get('skill_pack_ids', [])

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('recognition') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    chapter_text, ch_count = _collect_volume_chapters(book_id, volume_id)
    source_label = ''
    if not chapter_text or ch_count == 0:
        # 无章节时，从设定/大纲/剧情维度提取基础数据
        chapter_text, source_label = _collect_dimension_source(bb, volume_title or '全部章节')
        if not chapter_text:
            return jsonify({'error': '该卷暂无章节，且设定/大纲/剧情维度也为空，无法识别动态文件。请先填写设定、大纲或剧情维度。'}), 400

    # 收集该卷区间内的已有动态报告（5章一份）
    dyn_reports = DynamicReport.query.filter_by(book_id=book_id).order_by(DynamicReport.chapter_start).all()
    relevant_reports = []
    # 计算该卷的起止章号（与 _collect_volume_chapters 一致：parent_id 优先，回退顺序遍历）
    all_chs = Chapter.query.filter_by(book_id=book_id).order_by(Chapter.order_index).all()
    if volume_id:
        vol_chs_in_order = _get_volume_chapters_ordered(book_id, volume_id)
        # 计算该卷章节在全章节序列中的序号（1-based）
        non_vol_chs = [c for c in all_chs if not c.is_volume]
        vol_ch_idx = [non_vol_chs.index(c) for c in vol_chs_in_order if c in non_vol_chs]
    else:
        vol_ch_idx = list(range(len([c for c in all_chs if not c.is_volume])))

    if vol_ch_idx:
        # 章节序号从1开始
        ch_start = vol_ch_idx[0] + 1
        ch_end = vol_ch_idx[-1] + 1
        for r in dyn_reports:
            if r.chapter_end >= ch_start and r.chapter_start <= ch_end:
                relevant_reports.append(r)

    reports_text = '\n\n'.join([f'【{r.title}】\n{(r.content or "")[:500]}' for r in relevant_reports]) if relevant_reports else '（无已生成报告）'

    ctx_parts = []
    if bb.character_profiles:
        ctx_parts.append(f'【人物档案】\n{bb.character_profiles[:600]}')
    if bb.key_rules:
        ctx_parts.append(f'【核心规则】\n{bb.key_rules[:600]}')
    extra_ctx = '\n\n'.join(ctx_parts)

    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['lock_facts', 'narrative_debt', 'foreshadow_register'], mode='agent')

    vol_label = volume_title or '全部章节'
    system_prompt = f"""你是专业的小说防遗忘系统分析师。请从以下“{vol_label}”的章节内容及已有动态报告中，生成本卷的动态分类摘要。

【已有动态报告（5章一份）】
{reports_text}

{extra_ctx and f"【已有参考】{chr(10)}{extra_ctx}" or ""}

严格按JSON对象格式输出（不要任何其他文字）：
{{
  "volume": "{vol_label}",
  "characters": "本卷登场人物及状态变化（200字内）",
  "events": "本卷关键事件脉络（200字内）",
  "timeline": "本卷时间线要点（150字内）",
  "locations": "本卷涉及地点（100字内）",
  "factions": "本卷势力动态（100字内）",
  "foreshadowing": "本卷埋设/回收的伏笔（150字内）",
  "realms": "本卷境界/能力变化（100字内）",
  "relationships": "本卷人物关系变化（100字内）",
  "summary": "本卷综合动态摘要（300字内）"
}}

{skill_note}"""

    user_prompt = f'作品标题：{book.title}\n卷名：{vol_label}\n{source_label and f"（数据来源：{source_label}）" or ""}\n\n以下是该卷内容：\n\n{chapter_text or "（无章节内容，依据报告生成）"}'

    content, err = _call_llm(
        [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
        max_tokens=2500, temperature=0.3, task_type='recognition'
    )
    if err:
        return jsonify({'error': err}), 500

    analysis, parse_err = _extract_json_from_llm(content, expect='object')
    if analysis is None:
        return jsonify({'error': 'AI返回格式无法解析', 'raw': content[:300], 'parse_error': parse_err}), 500

    entry = {
        'volume_id': volume_id,
        'volume': vol_label,
        'volume_index': _extract_volume_index(vol_label) or 0,
        'data': analysis if isinstance(analysis, dict) else {},
    }
    data_list = _upsert_volume_entry(bb, 'dynamic_volumes', entry)
    bb.last_synced_at = datetime.now(timezone.utc)
    db.session.commit()

    return jsonify({
        'success': True,
        'volume_data': entry,
        'dynamic_volumes': data_list,
        'bible': bb.to_dict()
    })

@app.route('/api/books/<book_id>/ai-analyze-foreshadowing-volume', methods=['POST'])
@login_required
def ai_analyze_foreshadowing_volume(book_id):
    """AI识别指定卷的伏笔。按卷分析章节内容，识别本卷埋设/回收的伏笔并写入 foreshadowing_volumes。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    volume_id = data.get('volume_id', '')
    volume_title = data.get('volume_title', '')
    skill_pack_ids = data.get('skill_pack_ids', [])

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('recognition') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    chapter_text, ch_count = _collect_volume_chapters(book_id, volume_id)
    source_label = ''
    if not chapter_text or ch_count == 0:
        # 无章节时，从设定/大纲/剧情维度提取基础数据
        chapter_text, source_label = _collect_dimension_source(bb, volume_title or '全部章节')
        if not chapter_text:
            return jsonify({'error': '该卷暂无章节，且设定/大纲/剧情维度也为空，无法识别伏笔。请先填写设定、大纲或剧情维度。'}), 400

    ctx_parts = []
    if bb.foreshadowing:
        ctx_parts.append(f'【全局伏笔档案（参考，避免重复）】\n{bb.foreshadowing[:800]}')
    if bb.key_rules:
        ctx_parts.append(f'【核心规则】\n{bb.key_rules[:400]}')
    extra_ctx = '\n\n'.join(ctx_parts)

    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['foreshadow_register', 'narrative_debt'], mode='agent')

    vol_label = volume_title or '全部章节'
    system_prompt = f"""你是专业的小说伏笔分析师。请从以下“{vol_label}”的章节内容中，识别本卷埋设的伏笔、回收的伏笔、以及尚未回收的悬念。

{extra_ctx and f"【已有参考】{chr(10)}{extra_ctx}" or ""}

严格按JSON对象格式输出（不要任何其他文字）：
{{
  "volume": "{vol_label}",
  "planted": [
    {{"content": "伏笔内容", "chapter": "埋设章节/位置", "purpose": "埋设目的", "status": "待回收"}}
  ],
  "resolved": [
    {{"content": "伏笔内容", "planted_at": "埋设位置", "resolved_at": "回收位置", "effect": "回收效果"}}
  ],
  "pending": [
    {{"content": "未回收悬念", "planted_at": "埋设位置", "importance": "高/中/低"}}
  ],
  "summary": "本卷伏笔综述（150字内）"
}}

{skill_note}"""

    user_prompt = f'作品标题：{book.title}\n卷名：{vol_label}\n{source_label and f"（数据来源：{source_label}）" or ""}\n\n以下是该卷内容：\n\n{chapter_text or "（无内容，请根据设定推断）"}'

    content, err = _call_llm(
        [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
        max_tokens=2500, temperature=0.3, task_type='recognition'
    )
    if err:
        return jsonify({'error': err}), 500

    try:
        analysis = json.loads(content)
    except (json.JSONDecodeError, ValueError):
        import re as _re_fv
        m = _re_fv.search(r'\{[\s\S]*\}', content)
        if m:
            analysis = json.loads(m.group())
        else:
            return jsonify({'error': 'AI返回格式无法解析', 'raw': content[:300]}), 500

    entry = {
        'volume_id': volume_id,
        'volume': vol_label,
        'volume_index': _extract_volume_index(vol_label) or 0,
        'data': analysis if isinstance(analysis, dict) else {},
    }
    data_list = _upsert_volume_entry(bb, 'foreshadowing_volumes', entry)
    bb.last_synced_at = datetime.now(timezone.utc)
    db.session.commit()

    return jsonify({
        'success': True,
        'volume_data': entry,
        'foreshadowing_volumes': data_list,
        'bible': bb.to_dict()
    })

@app.route('/api/books/<book_id>/ai-analyze-locations-volume', methods=['POST'])
@login_required
def ai_analyze_locations_volume(book_id):
    """AI识别指定卷的地点/地图。按卷分析章节内容，识别本卷涉及的地点并写入 locations_volumes。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    volume_id = data.get('volume_id', '')
    volume_title = data.get('volume_title', '')
    skill_pack_ids = data.get('skill_pack_ids', [])

    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
        db.session.commit()

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('recognition') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    chapter_text, ch_count = _collect_volume_chapters(book_id, volume_id)
    source_label = ''
    if not chapter_text or ch_count == 0:
        # 无章节时，从设定/大纲/剧情维度提取基础数据
        chapter_text, source_label = _collect_dimension_source(bb, volume_title or '全部章节')
        if not chapter_text:
            return jsonify({'error': '该卷暂无章节，且设定/大纲/剧情维度也为空，无法识别地点。请先填写设定、大纲或剧情维度。'}), 400

    ctx_parts = []
    if bb.locations:
        ctx_parts.append(f'【全局地点档案（参考）】\n{bb.locations[:800]}')
    if bb.worldbuilding:
        ctx_parts.append(f'【世界观设定】\n{bb.worldbuilding[:500]}')
    extra_ctx = '\n\n'.join(ctx_parts)

    # P2-10: 'world_setting' 是幽灵key，替换为 'tomato_setting'
    skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['lock_facts', 'tomato_setting'], mode='agent')

    vol_label = volume_title or '全部章节'
    system_prompt = f"""你是专业的小说地图分析师。请从以下“{vol_label}”的章节内容中，识别本卷涉及的所有地点、场景、地理信息。

{extra_ctx and f"【已有参考】{chr(10)}{extra_ctx}" or ""}

严格按JSON对象格式输出（不要任何其他文字）：
{{
  "volume": "{vol_label}",
  "locations": [
    {{"name": "地点名", "type": "城市/山脉/秘境/建筑/其它", "description": "地点描述", "events": "该地点发生的重要事件", "importance": "高/中/低"}}
  ],
  "regions": [
    {{"name": "区域名", "scope": "范围描述", "feature": "区域特征"}}
  ],
  "summary": "本卷地理概况（150字内）"
}}

{skill_note}"""

    user_prompt = f'作品标题：{book.title}\n卷名：{vol_label}\n{source_label and f"（数据来源：{source_label}）" or ""}\n\n以下是该卷内容：\n\n{chapter_text or "（无内容，请根据设定推断）"}'

    content, err = _call_llm(
        [{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
        max_tokens=2500, temperature=0.3, task_type='recognition'
    )
    if err:
        return jsonify({'error': err}), 500

    try:
        analysis = json.loads(content)
    except (json.JSONDecodeError, ValueError):
        import re as _re_lv
        m = _re_lv.search(r'\{[\s\S]*\}', content)
        if m:
            analysis = json.loads(m.group())
        else:
            return jsonify({'error': 'AI返回格式无法解析', 'raw': content[:300]}), 500

    entry = {
        'volume_id': volume_id,
        'volume': vol_label,
        'volume_index': _extract_volume_index(vol_label) or 0,
        'data': analysis if isinstance(analysis, dict) else {},
    }
    data_list = _upsert_volume_entry(bb, 'locations_volumes', entry)
    bb.last_synced_at = datetime.now(timezone.utc)
    db.session.commit()

    return jsonify({
        'success': True,
        'volume_data': entry,
        'locations_volumes': data_list,
        'bible': bb.to_dict()
    })

@app.route('/api/books/<book_id>/clear-timeline', methods=['POST'])
@login_required
def clear_timeline(book_id):
    """一键清空剧情分卷大纲（timeline 字段），不影响章节表。"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)
    bb.timeline = ''
    bb.last_synced_at = datetime.now(timezone.utc)
    db.session.commit()
    return jsonify({'success': True, 'bible': bb.to_dict()})

@app.route('/api/books/<book_id>/dynamic-memory', methods=['GET'])
@login_required
def get_dynamic_memory(book_id):
    """获取动态文件库（5个JSON文件）"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    dm = DynamicMemory.query.filter_by(book_id=book_id).first()
    if not dm:
        # 自动初始化空模板
        dm = DynamicMemory(book_id=book_id)
        for key in DynamicMemory.FILE_KEYS:
            setattr(dm, key, DynamicMemory.get_empty_template(key))
        db.session.add(dm)
        db.session.commit()

    return jsonify(dm.to_dict())

@app.route('/api/books/<book_id>/dynamic-memory/<file_key>', methods=['PUT'])
@login_required
def update_dynamic_memory_file(book_id, file_key):
    """更新单个动态文件"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    if file_key not in DynamicMemory.FILE_KEYS:
        return jsonify({'error': f'未知文件: {file_key}'}), 400

    data = request.get_json() or {}
    content = data.get('content', '')

    # 验证JSON格式
    if content.strip():
        try:
            json.loads(content)
        except json.JSONDecodeError as e:
            return jsonify({'error': f'JSON格式错误: {str(e)}'}), 400

    dm = DynamicMemory.query.filter_by(book_id=book_id).first()
    if not dm:
        dm = DynamicMemory(book_id=book_id)
        for key in DynamicMemory.FILE_KEYS:
            if key != file_key:
                setattr(dm, key, DynamicMemory.get_empty_template(key))
        db.session.add(dm)

    setattr(dm, file_key, content)
    db.session.commit()

    return jsonify({'success': True, 'file_key': file_key, 'content': content})

@app.route('/api/books/<book_id>/dynamic-memory/init', methods=['POST'])
@login_required
def init_dynamic_memory(book_id):
    """初始化动态文件库（用空模板填充所有文件）"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    dm = DynamicMemory.query.filter_by(book_id=book_id).first()
    if not dm:
        dm = DynamicMemory(book_id=book_id)
        db.session.add(dm)

    for key in DynamicMemory.FILE_KEYS:
        setattr(dm, key, DynamicMemory.get_empty_template(key))
    db.session.commit()

    return jsonify({'success': True, 'data': dm.to_dict()})

@app.route('/api/books/<book_id>/dynamic-memory/ai-generate', methods=['POST'])
@login_required
def ai_generate_dynamic_memory(book_id):
    """AI生成/更新动态文件库（基于已有章节内容）"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.model if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    data = request.get_json() or {}
    file_key = data.get('file_key', '')
    if file_key not in DynamicMemory.FILE_KEYS:
        return jsonify({'error': f'未知文件: {file_key}'}), 400

    # 获取作品信息
    bible = BookBible.query.filter_by(book_id=book_id).first()
    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()

    # 拼接章节内容摘要
    chapter_summaries = []
    total_chars = 0
    max_chars = 8000
    for ch in chapters:
        summary = f'【{ch.title}】{(ch.content or "")[:500]}'
        if total_chars + len(summary) > max_chars:
            break
        chapter_summaries.append(summary)
        total_chars += len(summary)

    chapters_text = '\n\n'.join(chapter_summaries) if chapter_summaries else '暂无章节内容'

    file_descriptions = {
        'narrative_engine': '叙事引擎：包含state（书名/当前章/卷/节点/mc状态/最近事件/活跃钩子/待处理伏笔/代价点数）+ timeline[]（章号/故事天/事件/地点/角色/重要性）+ chapters[]（章号/标题/字数/摘要/关键事件/伏笔操作/钩子/爽点类型）',
        'foreshadowing_tracker': '伏笔追踪器：包含foreshadowing[]（id/类型/状态/埋设章/当前年龄/最大年龄/紧急度/描述/激活条件/计划回收章）+ scan_rules（短/中/长周期阈值+自动告警级别）',
        'character_ecosystem': '角色生态系统：包含characters[]（id/名字/角色/状态{能力/身体/精神}/CDL{驱动/恐惧/成长}/最近出场/台词量/关系温度）+ relationships[]（from/to/类型/等级/最近互动/趋势）',
        'ability_world': '能力与世界观：包含ability_log[]（章号/角色/变化/新能力/触发条件/代价变动）+ world_facts[]（id/事实/建立章/矛盾记录）',
        'health_dashboard': '健康度仪表盘：包含thread_health（四线各含最近推进章+状态）+ chapter_type_distribution + foreshadowing_aging + ai_flavor_trend[] + dialogue_ratio_trend[] + character_appearances + alerts[]',
    }

    # 获取当前文件内容
    dm = DynamicMemory.query.filter_by(book_id=book_id).first()
    current_content = ''
    if dm:
        current_content = getattr(dm, file_key, '') or ''

    system_prompt = f"""你是专业的小说分析师和创作助手。请根据小说的章节内容、设定信息，生成或更新“{file_descriptions.get(file_key, file_key)}”的JSON数据。

要求：
1. 严格输出有效的JSON格式，不要有任何其他文字
2. 中文字段名，UTF-8编码
3. 基于已有章节内容进行分析和提取
4. 如果已有当前文件内容，在原有基础上更新而非完全重写
5. 未涉及的字段保持原值或空值"""

    user_content = f"""作品标题：{book.title}
题材：{_get_genre_label(book)}
类型：{book.book_type or '小说'}

构思：{bible.concept if bible else '无'}
世界观：{(bible.worldbuilding[:300] if bible and bible.worldbuilding else '无')}
人物：{(bible.character_profiles[:300] if bible and bible.character_profiles else '无')}
大纲：{(bible.plot_design[:300] if bible and bible.plot_design else '无')}

已有章节内容摘要：
{chapters_text}

当前{file_key}文件内容：
{current_content if current_content else '（空，请初始化）'}

请生成完整的JSON数据："""

    try:
        base = base_url.rstrip('/')
        if not base.endswith('/v1'):
            base += '/v1'
        resp = requests.post(f'{base}/chat/completions',
            headers=build_auth_headers(api_key),
            json={
                'model': model,
                'messages': [
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': user_content}
                ],
                'temperature': 0.4,
                'max_tokens': 4000,
                'response_format': {'type': 'json_object'}
            },
            timeout=180)

        result = resp.json()
        content = result['choices'][0]['message']['content']

        # 验证返回的JSON
        try:
            parsed = json.loads(content)
            content = json.dumps(parsed, ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            pass  # 保留原始内容

        # 保存到数据库
        if not dm:
            dm = DynamicMemory(book_id=book_id)
            db.session.add(dm)
        setattr(dm, file_key, content)
        db.session.commit()

        return jsonify({
            'success': True,
            'file_key': file_key,
            'content': content,
            'data': dm.to_dict()
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ==== 动态报告 API（防遗忘摘要系统） ====

DYNAMIC_REPORT_INTERVAL = 5  # 每5章生成一份报告

def _generate_dynamic_report_content(book_id, chapter_start, chapter_end, skill_pack_ids=None):
    """内部函数：调用AI生成动态报告内容。
    skill_pack_ids: 可选，技能包ID列表，用于注入提示词（P0-2修复）"""
    book = Book.query.get(book_id)
    if not book:
        return None, 'Book not found'

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.model if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return None, '请先配置 AI 模型 API Key'

    bible = BookBible.query.filter_by(book_id=book_id).first()
    # 获取指定范围的章节（按order_index排序，非卷标）
    all_chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    # chapter_start/chapter_end 是1-based的章号
    target_chapters = all_chapters[chapter_start - 1:chapter_end] if chapter_start <= len(all_chapters) else []

    if not target_chapters:
        return None, '指定范围内无章节'

    # P0-2: 注入技能包提示词（narrative_debt/foreshadow_register 用于动态摘要生成）
    skill_note = ''
    if skill_pack_ids:
        skill_note = _get_skill_prompts_by_category(skill_pack_ids, 'master', ['narrative_debt', 'foreshadow_register', 'lock_facts'], mode='agent')

    # 拼接章节内容（每章截取前800字，控制总量）
    chapters_text = []
    total_chars = 0
    max_chars = 6000
    for ch in target_chapters:
        snippet = f'【{ch.title}】\n{(ch.content or "")[:800]}'
        if total_chars + len(snippet) > max_chars:
            remaining = max_chars - total_chars
            if remaining > 100:
                snippet = snippet[:remaining]
                chapters_text.append(snippet)
            break
        chapters_text.append(snippet)
        total_chars += len(snippet)

    chapters_content = '\n\n'.join(chapters_text)

    system_prompt = f"""你是专业的小说分析师。请仔细阅读以下第{chapter_start}章到第{chapter_end}章的内容，汇总成一份简洁的报告。

报告必须包含以下要素（如果出现的话），总字数不超过500字：

1. 【人物】本章新出场或有关键表现的角色及其状态变化
2. 【事件】关键剧情事件和转折
3. 【时间】故事内时间线推进
4. 【地点】涉及的重要地点
5. 【势力】出现或变动的势力/组织
6. 【伏笔】埋设或回收的伏笔
7. 【境界】角色境界/实力变化
8. 【关系】人物关系变化
9. 【物资】主要角色当前拥有的重要物品、装备、功法、丹药等（按角色列出）

格式要求：
- 用简洁的条目式写法，每类1-3条
- 只记录关键信息，不展开描述
- 不要写废话和过渡句
- 直接输出报告内容，不要加标题和前后缀

{skill_note}"""

    user_content = f"""作品：{book.title}
题材：{_get_genre_label(book)}

第{chapter_start}章到第{chapter_end}章内容：
{chapters_content}

请生成动态报告（≤500字）："""

    # P2-8：动态报告失败重试（指数退避，最多2次重试），全失败则降级为章节摘要拼接
    import time as _time
    max_retries = 2
    last_error = None
    try:
        base = base_url.rstrip('/')
        if not base.endswith('/v1'):
            base += '/v1'
        for attempt in range(max_retries + 1):
            try:
                resp = requests.post(f'{base}/chat/completions',
                    headers=build_auth_headers(api_key),
                    json={
                        'model': model,
                        'messages': [
                            {'role': 'system', 'content': system_prompt},
                            {'role': 'user', 'content': user_content}
                        ],
                        'temperature': 0.3,
                        'max_tokens': 1200,
                    },
                    timeout=120)
                result = resp.json()
                content = result['choices'][0]['message']['content'].strip()
                if content and len(content) > 50:
                    return content, None
                last_error = f'LLM 返回内容过短（{len(content)}字）'
            except Exception as e:
                last_error = str(e)
            # 指数退避：1s, 2s
            if attempt < max_retries:
                _time.sleep(2 ** attempt)
        # 全部重试失败，降级为章节摘要拼接
        fallback_parts = [f'【降级摘要·第{chapter_start}-{chapter_end}章】（LLM生成失败，自动拼接）']
        for ch in target_chapters:
            snippet = (ch.content or '')[:120].replace('\n', ' ').strip()
            if snippet:
                fallback_parts.append(f'- {ch.title}：{snippet}')
        fallback_content = '\n'.join(fallback_parts)[:800]
        return fallback_content, None
    except Exception as e:
        # 连降级都失败的极端情况
        return None, f'{str(e)} | last_error={last_error}'

def _check_and_auto_generate_report(book_id):
    """检查是否需要自动生成动态报告（每5章触发）"""
    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    chapter_count = len(chapters)
    if chapter_count == 0:
        return None

    # 计算当前应该有报告的最后一个区间
    # 例如：5章 -> 区间1-5，10章 -> 区间1-5和6-10
    current_end = (chapter_count // DYNAMIC_REPORT_INTERVAL) * DYNAMIC_REPORT_INTERVAL
    if current_end == 0:
        return None

    # 检查该区间是否已有报告
    existing = DynamicReport.query.filter_by(
        book_id=book_id, chapter_start=current_end - DYNAMIC_REPORT_INTERVAL + 1,
        chapter_end=current_end
    ).first()

    if existing:
        return None  # 已有报告，不重复生成

    # 需要生成新报告
    chapter_start = current_end - DYNAMIC_REPORT_INTERVAL + 1
    content, error = _generate_dynamic_report_content(book_id, chapter_start, current_end)
    if error:
        return {'error': error}

    title = f'动态-({chapter_start}-{current_end}章)'
    report = DynamicReport(
        book_id=book_id, title=title, content=content or '',
        chapter_start=chapter_start, chapter_end=current_end,
        auto_generated=True
    )
    db.session.add(report)
    db.session.commit()

    # 借鉴 PlotPilot 检查点快照：每5章自动备份 BookBible + DynamicMemory 关键状态
    # 用户可回滚到某章节点重新创作，避免"写崩了无法恢复"
    try:
        _create_state_snapshot(book_id, current_end)
    except Exception as snap_err:
        try:
            app.logger.warning(f'快照创建失败（不影响主流程）: {snap_err}')
        except Exception:
            pass

    return {'report': report.to_dict()}

def _create_state_snapshot(book_id, chapter_end):
    """创建叙事状态检查点快照（借鉴 PlotPilot checkpoint）。
    备份 BookBible 关键字段 + DynamicMemory 5文件，存入 bb.state_snapshots。
    每个快照含：snapshot_id / chapter_end / created_at / bible_fields / dynamic_memory。
    最多保留 20 个快照（超出按时间淘汰最旧的）。"""
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        return
    # 备份 BookBible 关键叙事字段（不含快照自身，避免递归膨胀）
    bible_fields = {}
    for field in ['worldbuilding', 'character_profiles', 'timeline', 'foreshadowing',
                  'style_guide', 'key_rules', 'locations', 'concept', 'plot_design',
                  'relation_graph', 'inventory', 'character_volumes', 'dynamic_volumes',
                  'foreshadowing_volumes', 'locations_volumes', 'foreshadowing_graph',
                  'outline_hierarchy', 'chapter_changes_log']:
        bible_fields[field] = getattr(bb, field, '') or ''
    # 备份 DynamicMemory 5文件
    dm_data = {}
    try:
        dm = DynamicMemory.query.filter_by(book_id=book_id).first()
        if dm:
            for key in ['narrative_engine', 'foreshadowing_tracker', 'character_ecosystem',
                         'ability_world', 'health_dashboard']:
                dm_data[key] = getattr(dm, key, '') or ''
    except Exception:
        pass
    snapshot = {
        'snapshot_id': str(uuid.uuid4())[:8],
        'chapter_end': chapter_end,
        'created_at': datetime.now(timezone.utc).isoformat(),
        'bible_fields': bible_fields,
        'dynamic_memory': dm_data,
    }
    existing = []
    try:
        existing = json.loads(bb.state_snapshots or '[]') if bb.state_snapshots else []
    except Exception:
        existing = []
    existing.append(snapshot)
    # 最多保留 20 个，淘汰最旧
    if len(existing) > 20:
        existing = existing[-20:]
    bb.state_snapshots = json.dumps(existing, ensure_ascii=False)
    db.session.commit()

@app.route('/api/books/<book_id>/dynamic-reports', methods=['GET'])
@login_required
def list_dynamic_reports(book_id):
    """获取所有动态报告"""
    reports = DynamicReport.query.filter_by(book_id=book_id).order_by(DynamicReport.chapter_start).all()
    return jsonify([r.to_dict() for r in reports])

@app.route('/api/books/<book_id>/dynamic-reports', methods=['POST'])
@login_required
def create_dynamic_report(book_id):
    """手动创建动态报告"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    chapter_start = data.get('chapter_start', 0)
    chapter_end = data.get('chapter_end', 0)

    if not chapter_start or not chapter_end or chapter_end < chapter_start:
        return jsonify({'error': '请指定有效的章节范围'}), 400

    # 如果没有提供content，调用AI生成
    content = data.get('content', '')
    if not content:
        content, error = _generate_dynamic_report_content(book_id, chapter_start, chapter_end)
        if error:
            return jsonify({'error': error}), 500

    title = data.get('title', f'动态-({chapter_start}-{chapter_end}章)')
    report = DynamicReport(
        book_id=book_id, title=title, content=content,
        chapter_start=chapter_start, chapter_end=chapter_end,
        auto_generated=False
    )
    db.session.add(report)
    db.session.commit()
    return jsonify(report.to_dict()), 201

@app.route('/api/books/<book_id>/dynamic-reports/batch-generate', methods=['POST'])
@login_required
def batch_generate_dynamic_reports(book_id):
    """按卷批量生成动态报告（每5章一份）。AI识别按钮选择某卷后，自动生成该卷内所有
    尚未生成的5章区间动态报告，减少一个个手动添加的麻烦。
    参数：volume_id（可选，空则全卷）、volume_title、skill_pack_ids、overwrite（是否覆盖已存在，默认false）"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    volume_id = data.get('volume_id', '')
    volume_title = data.get('volume_title', '')
    skill_pack_ids = data.get('skill_pack_ids', [])
    overwrite = data.get('overwrite', False)

    config = AIConfig.get_active()
    if not config or not config.api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    # 收集全作品非卷章节（按 order_index，用于计算全局 1-based 章号）
    all_chs = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    if not all_chs:
        return jsonify({'error': '该作品暂无章节，无法批量生成动态报告。'}), 400

    # 确定该卷章节的全局1-based起止章号
    if volume_id:
        # 用共享助手取该卷章节：parent_id 优先，回退顺序遍历（兼容未设置 parent_id 的旧数据）
        vol_chs = _get_volume_chapters_ordered(book_id, volume_id)
        if not vol_chs:
            return jsonify({'error': f'卷“{volume_title or volume_id}”内暂无章节（请确认章节已归入该卷）'}), 400
        # 计算这些章节在 all_chs 中的全局序号（1-based）
        ch_id_to_idx = {c.id: i for i, c in enumerate(all_chs)}
        vol_ch_idx = [ch_id_to_idx[c.id] for c in vol_chs if c.id in ch_id_to_idx]
        if not vol_ch_idx:
            return jsonify({'error': f'卷“{volume_title or volume_id}”内暂无章节'}), 400
        global_start = min(vol_ch_idx) + 1
        global_end = max(vol_ch_idx) + 1
    else:
        global_start = 1
        global_end = len(all_chs)

    # 按5章一份切分区间
    intervals = []
    s = global_start
    while s <= global_end:
        e = min(s + DYNAMIC_REPORT_INTERVAL - 1, global_end)
        intervals.append((s, e))
        s = e + 1

    if not intervals:
        return jsonify({'error': '该卷章节范围无效'}), 400

    # 查询已存在的报告（按chapter_start/chapter_end匹配），决定是否跳过/覆盖
    existing = DynamicReport.query.filter_by(book_id=book_id).all()
    existing_map = {(r.chapter_start, r.chapter_end): r for r in existing}

    generated = []
    skipped = []
    errors = []
    for (cs, ce) in intervals:
        key = (cs, ce)
        if key in existing_map and not overwrite:
            skipped.append({'chapter_start': cs, 'chapter_end': ce, 'reason': '已存在'})
            continue
        # 调用AI生成内容（P0-2: 传入 skill_pack_ids）
        content, err = _generate_dynamic_report_content(book_id, cs, ce, skill_pack_ids=skill_pack_ids)
        if err:
            errors.append({'chapter_start': cs, 'chapter_end': ce, 'error': err})
            continue
        if key in existing_map and overwrite:
            # 覆盖已存在报告
            r = existing_map[key]
            r.content = content
            r.title = f'动态-({cs}-{ce}章)'
            r.auto_generated = False
        else:
            # 新建报告
            r = DynamicReport(
                book_id=book_id, title=f'动态-({cs}-{ce}章)', content=content,
                chapter_start=cs, chapter_end=ce, auto_generated=False
            )
            db.session.add(r)
        db.session.commit()
        generated.append(r.to_dict())

    return jsonify({
        'success': True,
        'volume_title': volume_title or '全部章节',
        'chapter_range': [global_start, global_end],
        'total_intervals': len(intervals),
        'generated_count': len(generated),
        'skipped_count': len(skipped),
        'error_count': len(errors),
        'generated': generated,
        'skipped': skipped,
        'errors': errors,
    })

@app.route('/api/books/<book_id>/dynamic-reports/<report_id>', methods=['PUT'])
@login_required
def update_dynamic_report(book_id, report_id):
    """更新动态报告"""
    report = DynamicReport.query.filter_by(id=report_id, book_id=book_id).first()
    if not report:
        return jsonify({'error': 'Report not found'}), 404

    data = request.get_json() or {}
    for field in ['title', 'content', 'chapter_start', 'chapter_end']:
        if field in data:
            setattr(report, field, data[field])
    db.session.commit()
    return jsonify(report.to_dict())

@app.route('/api/books/<book_id>/dynamic-reports/<report_id>', methods=['DELETE'])
@login_required
def delete_dynamic_report(book_id, report_id):
    """删除动态报告"""
    report = DynamicReport.query.filter_by(id=report_id, book_id=book_id).first()
    if not report:
        return jsonify({'error': 'Report not found'}), 404
    db.session.delete(report)
    db.session.commit()
    return jsonify({'success': True})

@app.route('/api/books/<book_id>/dynamic-reports/batch-delete', methods=['POST'])
@login_required
def batch_delete_dynamic_reports(book_id):
    """批量删除动态报告"""
    data = request.get_json() or {}
    report_ids = data.get('report_ids') or []
    if not isinstance(report_ids, list) or not report_ids:
        return jsonify({'error': '请提供要删除的报告ID列表'}), 400

    deleted = DynamicReport.query.filter(
        DynamicReport.book_id == book_id,
        DynamicReport.id.in_(report_ids)
    ).all(synchronize_session=False)
    deleted_ids = [r.id for r in deleted]
    for r in deleted:
        db.session.delete(r)
    db.session.commit()
    return jsonify({'success': True, 'deleted_count': len(deleted_ids), 'deleted_ids': deleted_ids})

@app.route('/api/books/<book_id>/dynamic-reports/<report_id>/regenerate', methods=['POST'])
@login_required
def regenerate_dynamic_report(book_id, report_id):
    """重新生成动态报告内容（AI）。支持可选 skill_pack_ids 注入提示词。"""
    report = DynamicReport.query.filter_by(id=report_id, book_id=book_id).first()
    if not report:
        return jsonify({'error': 'Report not found'}), 404

    data = request.get_json(silent=True) or {}
    skill_pack_ids = data.get('skill_pack_ids', [])
    content, error = _generate_dynamic_report_content(book_id, report.chapter_start, report.chapter_end, skill_pack_ids=skill_pack_ids)
    if error:
        return jsonify({'error': error}), 500

    report.content = content or ''
    db.session.commit()
    return jsonify(report.to_dict())

@app.route('/api/books/<book_id>/dynamic-reports/auto-check', methods=['POST'])
@login_required
def auto_check_dynamic_report(book_id):
    """检查并自动生成动态报告（章节保存后触发）"""
    result = _check_and_auto_generate_report(book_id)
    if result is None:
        return jsonify({'success': True, 'message': '无需生成新报告', 'report': None})
    if 'error' in result:
        return jsonify({'success': False, 'error': result['error']}), 500
    return jsonify({'success': True, 'message': '已自动生成动态报告', 'report': result['report']})

@app.route('/api/books/<book_id>/dynamic-reports/context', methods=['GET'])
@login_required
def get_dynamic_report_context(book_id):
    """获取最近的动态报告内容（用于AI创作时注入上下文，减少token）"""
    # 返回最近10份报告（覆盖更长前文记忆，保证剧情连贯）
    reports = DynamicReport.query.filter_by(book_id=book_id).order_by(
        DynamicReport.chapter_start.desc()
    ).limit(10).all()
    # 按正序返回
    reports.reverse()
    return jsonify({
        'reports': [r.to_dict() for r in reports],
        'context_text': '\n\n'.join([f'【{r.title}】\n{r.content}' for r in reports if r.content])
    })

@app.route('/api/books/<book_id>/ai-analyze-from-reports', methods=['POST'])
@login_required
def ai_analyze_from_reports(book_id):
    """从动态文件报告提取维度信息（地图/关系图谱/地点图谱/境界图谱等），节省token"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json() or {}
    dimension = data.get('dimension', '')
    if not dimension:
        return jsonify({'error': '缺少 dimension 参数'}), 400

    # 维度 → bible字段 映射
    dim_field_map = {
        'locations': 'locations',
        'relationGraph': 'relation_graph',
        'locationGraph': 'locations',
        'realmGraph': 'worldbuilding',
    }
    field = dim_field_map.get(dimension)
    if not field:
        return jsonify({'error': f'不支持的维度: {dimension}'}), 400

    dim_labels = {
        'locations': '地点/地图',
        'relationGraph': '人物关系',
        'locationGraph': '地点关系',
        'realmGraph': '境界体系',
    }
    dim_label = dim_labels.get(dimension, dimension)

    config = AIConfig.get_active()
    api_key = config.api_key if config and config.api_key else os.environ.get('USER_LLM_API_KEY', '')
    base_url = config.base_url if config else os.environ.get('USER_LLM_BASE_URL', 'https://api.deepseek.com/v1')
    model = config.get_model_for_task('recognition') if config else os.environ.get('USER_LLM_MODEL', 'deepseek-chat')

    if not api_key:
        return jsonify({'error': '请先配置 AI 模型 API Key'}), 400

    # 获取已有的bible字段内容
    bible = BookBible.query.filter_by(book_id=book_id).first()
    existing_value = getattr(bible, field, '') if bible else ''

    # 按维度组装数据源（用户要求：不同图谱从不同维度读取数据供AI识别）
    # 1. 关系图谱：从“人物及关系”+“剧情”维度读取，再补充动态文件
    # 2. 地点图谱：首先从“设定”+“大纲”维度读取，再从动态文件补充
    # 3. 境界图谱：首先从“设定”+“大纲”维度读取，再从动态文件补充
    # 4. 地图(locations)：保持动态文件优先，回退章节内容

    def _bible_val(attr):
        return getattr(bible, attr, '') if bible else ''

    # 动态文件报告（补充数据源）
    reports = DynamicReport.query.filter_by(book_id=book_id).order_by(
        DynamicReport.chapter_start
    ).all()
    dynamic_text = '\n\n'.join([f'【{r.title}】\n{r.content}' for r in reports if r.content]) if reports else ''

    # 章节内容（最终回退）
    chapter_text = ''
    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()
    if chapters:
        max_chars = 8000
        for ch in chapters:
            segment = f'【{ch.title}】\n{(ch.content or "")[:800]}\n\n'
            if len(chapter_text) + len(segment) > max_chars:
                break
            chapter_text += segment

    source_parts = []  # [(标签, 内容)]
    if dimension == 'relationGraph':
        # 关系图谱：人物及关系 + 剧情 + 动态文件
        cp = _bible_val('character_profiles')
        tl = _bible_val('timeline')
        if cp.strip():
            source_parts.append(('人物及关系维度', cp[:3000]))
        if tl.strip():
            source_parts.append(('剧情维度', tl[:3000]))
        if dynamic_text.strip():
            source_parts.append(('动态文件补充', dynamic_text[:3000]))
        elif chapter_text.strip():
            source_parts.append(('章节内容补充', chapter_text[:2000]))
    elif dimension in ('locationGraph', 'realmGraph'):
        # 地点图谱/境界图谱：设定 + 大纲 + 动态文件
        wb = _bible_val('worldbuilding')
        kr = _bible_val('key_rules')
        pd = _bible_val('plot_design')
        if wb.strip():
            source_parts.append(('设定维度(世界观)', wb[:3000]))
        if kr.strip():
            source_parts.append(('设定维度(核心规则)', kr[:2000]))
        if pd.strip():
            source_parts.append(('大纲维度', pd[:3000]))
        if dynamic_text.strip():
            source_parts.append(('动态文件补充', dynamic_text[:3000]))
        elif chapter_text.strip():
            source_parts.append(('章节内容补充', chapter_text[:2000]))
    else:
        # locations 及其他：动态文件优先，回退章节内容
        if dynamic_text.strip():
            source_parts.append(('动态文件报告', dynamic_text[:8000]))
        elif chapter_text.strip():
            source_parts.append(('章节内容', chapter_text[:8000]))
        else:
            # 没有任何数据源时，尝试从设定/大纲补充
            for attr, label in [('worldbuilding', '世界观设定'), ('key_rules', '核心规则'), ('plot_design', '大纲'), ('character_profiles', '人物及关系'), ('timeline', '剧情')]:
                v = _bible_val(attr)
                if v.strip():
                    source_parts.append((label, v[:3000]))

    if not source_parts:
        return jsonify({'error': '没有可用的数据源，请先在相关维度填写内容或生成动态文件'}), 400

    source_type = '、'.join([p[0] for p in source_parts])
    source_text = '\n\n'.join([f'--- {label} ---\n{content}' for label, content in source_parts])

    # 不同维度的提取提示
    dim_prompts = {
        'locations': """请从以下内容中提取所有地点信息，按三级分类整理（大区域/城市/场景）。
输出JSON格式：
{"locations": [{"name":"大区域名","desc":"描述","children":[{"name":"城市名","desc":"描述","children":[{"name":"场景名","desc":"描述"}]}]}]}
如果没有明确地点信息，输出空数组。""",

        'relationGraph': """请从以下内容中提取所有人物及其关系，整理为关系图谱数据。
重要规则：
- 只提取真实的人物姓名作为节点，绝对不要把"关系"、"好友"、"敌人"、"师徒"等关系类型词当作人物节点。
- 每个人物用姓名开头，关系单独列为"关系: A与B-关系类型"。

输出JSON格式：
{"relation_graph": "人物1: 姓名|身份|性格|动机\\n人物2: 姓名|身份|性格|动机\\n关系: A与B-关系类型"}
如果没有人物信息，输出空字符串。""",

        'locationGraph': """请从以下内容中提取地点之间的关联关系。
输出JSON格式：
{"locations": [{"name":"大区域名","desc":"描述","children":[{"name":"城市名","desc":"描述","children":[{"name":"场景名","desc":"描述"}]}]}]}
按层级整理地点体系。""",

        'realmGraph': """请从以下内容中提取境界/等级/实力体系信息。
输出JSON格式：
{"worldbuilding": "境界体系:\\n第一级: xxx\\n第二级: xxx\\n...\\n能力规则: xxx"}
如果没有境界信息，输出空字符串。""",

        'worldbuilding': """请从以下内容中提取世界观设定，包括境界体系、力量规则、社会结构等。
输出纯文本格式，200-400字。""",
        'character_profiles': """请从以下内容中提取人物档案和关系。
输出纯文本格式，每人一行：姓名|身份|性格|动机|关系。""",
    }

    prompt = dim_prompts.get(dimension, dim_prompts.get(field, f'提取{dim_label}信息'))

    system_prompt = f"""你是专业的小说分析师。请从以下{source_type}中提取“{dim_label}”维度的信息。

已有内容（供参考，在基础上补充而非完全重写）：
{existing_value[:500] if existing_value else '（空）'}

{prompt}"""

    user_content = f"""作品：{book.title}

{source_type}：
{source_text}

请提取{dim_label}信息："""

    try:
        base = base_url.rstrip('/')
        if not base.endswith('/v1'):
            base += '/v1'

        use_json = dimension in ('locations', 'locationGraph')
        req_body = {
            'model': model,
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_content}
            ],
            'temperature': 0.3,
            'max_tokens': 2000,
        }
        if use_json:
            req_body['response_format'] = {'type': 'json_object'}

        resp = requests.post(f'{base}/chat/completions',
            headers=build_auth_headers(api_key),
            json=req_body,
            timeout=120)

        result = resp.json()
        content = result['choices'][0]['message']['content'].strip()

        # 尝试解析JSON提取字段值
        extracted_value = content
        try:
            parsed = json.loads(content)
            if field in parsed:
                extracted_value = parsed[field] if isinstance(parsed[field], str) else json.dumps(parsed[field], ensure_ascii=False, indent=2)
            elif isinstance(parsed, dict) and len(parsed) == 1:
                # 单字段JSON
                val = list(parsed.values())[0]
                extracted_value = val if isinstance(val, str) else json.dumps(val, ensure_ascii=False, indent=2)
        except (json.JSONDecodeError, KeyError):
            pass  # 保留原始文本

        # 保存到bible
        if not bible:
            bible = BookBible(book_id=book_id)
            db.session.add(bible)
        setattr(bible, field, extracted_value)
        db.session.commit()

        return jsonify({
            'success': True,
            'dimension': dimension,
            'field': field,
            'value': extracted_value,
            'bible': bible.to_dict(),
            'source': source_type,
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/books/<book_id>/sync-analysis', methods=['POST'])
@login_required
def sync_analysis_to_book(book_id):
    """将拆书分析结果同步到作品资料（BookBible），支持仿写/同人文/参考三种模式"""
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    data = request.get_json()
    analysis = data.get('analysis', {})
    mode = data.get('mode', 'reference')  # imitate / fanfic / reference

    if not analysis:
        return jsonify({'error': '缺少分析结果数据'}), 400

    # 更新或创建 BookBible
    bb = BookBible.query.filter_by(book_id=book_id).first()
    if not bb:
        bb = BookBible(book_id=book_id)
        db.session.add(bb)

    updated_fields = []
    mode_label = {'imitate': '仿写参考', 'fanfic': '同人文素材', 'reference': '方法论参考'}.get(mode, '参考')
    header = f'\n\n【拆书分析·{mode_label}】\n'

    # 风格指南：文风特点 + 节奏特点
    style_parts = []
    if analysis.get('style_analysis'):
        style_parts.append(f'文风特点：{analysis["style_analysis"]}')
    if analysis.get('rhythm_analysis'):
        style_parts.append(f'节奏特点：{analysis["rhythm_analysis"]}')
    if analysis.get('target_platform'):
        style_parts.append(f'目标平台：{analysis["target_platform"]}')
    if analysis.get('genre_tags'):
        style_parts.append(f'题材标签：{"、".join(analysis["genre_tags"])}')
    if style_parts:
        existing = bb.style_guide or ''
        bb.style_guide = (existing + header if existing else '') + '\n'.join(style_parts)
        updated_fields.append('style_guide')

    # 大纲设计：结构特点
    if analysis.get('structure_analysis'):
        existing = bb.plot_design or ''
        content = f'结构参考：{analysis["structure_analysis"]}'
        bb.plot_design = (existing + header if existing else '') + content
        updated_fields.append('plot_design')

    # 人物档案：人设特点
    if analysis.get('character_design_analysis'):
        existing = bb.character_profiles or ''
        content = f'人设参考：{analysis["character_design_analysis"]}'
        if mode == 'fanfic' and analysis.get('golden_lines'):
            content += f'\n金句参考：{"；".join(analysis["golden_lines"][:5])}'
        bb.character_profiles = (existing + header if existing else '') + content
        updated_fields.append('character_profiles')

    # 伏笔：钩子技巧 + 可学方法
    foreshadow_parts = []
    if analysis.get('hook_techniques'):
        foreshadow_parts.append('钩子技巧：\n' + '\n'.join(f'· {h}' for h in analysis['hook_techniques']))
    if analysis.get('learnable_points'):
        foreshadow_parts.append('可学方法：\n' + '\n'.join(f'· {p}' for p in analysis['learnable_points']))
    if foreshadow_parts:
        existing = bb.foreshadowing or ''
        bb.foreshadowing = (existing + header if existing else '') + '\n'.join(foreshadow_parts)
        updated_fields.append('foreshadowing')

    # 同人文模式：额外同步世界观
    if mode == 'fanfic':
        fanfic_parts = []
        if analysis.get('style_analysis'):
            fanfic_parts.append(f'原作文风：{analysis["style_analysis"]}')
        if analysis.get('structure_analysis'):
            fanfic_parts.append(f'原作结构：{analysis["structure_analysis"]}')
        if analysis.get('character_design_analysis'):
            fanfic_parts.append(f'原作人设：{analysis["character_design_analysis"]}')
        if fanfic_parts:
            existing = bb.worldbuilding or ''
            bb.worldbuilding = (existing + header if existing else '') + '\n'.join(fanfic_parts)
            updated_fields.append('worldbuilding')

    bb.last_synced_at = datetime.now(timezone.utc)
    db.session.commit()

    return jsonify({
        'success': True,
        'updated_fields': updated_fields,
        'bible': bb.to_dict()
    })

# ==== File Upload / 拆书导入导出 ====

UPLOAD_DIR = DATA_DIR / 'uploads'
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_EXTENSIONS = {'txt', 'md', 'docx', 'zip', 'json'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def _detect_and_decode(raw_bytes):
    """自动检测字节流的编码并解码，支持 UTF-8 / GBK / GB18030 等中文编码"""
    # 优先尝试 UTF-8（含 BOM）
    if raw_bytes[:3] == b'\xef\xbb\xbf':
        return raw_bytes[3:].decode('utf-8', errors='replace')
    try:
        text = raw_bytes.decode('utf-8')
        # 检查是否有大量替换字符，如果有说明可能不是 UTF-8
        if text.count('\ufffd') < len(text) * 0.01:
            return text
    except UnicodeDecodeError:
        pass
    # 用 chardet 检测编码
    try:
        import chardet
        result = chardet.detect(raw_bytes)
        enc = result.get('encoding', '')
        if enc:
            try:
                return raw_bytes.decode(enc, errors='replace')
            except (UnicodeDecodeError, LookupError):
                pass
    except ImportError:
        pass
    # 回退到 GB18030（兼容 GBK / GB2312）
    try:
        return raw_bytes.decode('gb18030', errors='replace')
    except Exception:
        return raw_bytes.decode('utf-8', errors='replace')

def extract_text_from_file(filepath, filename):
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    if ext == 'txt':
        with open(filepath, 'rb') as f:
            raw = f.read()
        return _detect_and_decode(raw)
    elif ext in ('md', 'markdown', 'json'):
        with open(filepath, 'rb') as f:
            raw = f.read()
        return _detect_and_decode(raw)
    elif ext == 'docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return '\n'.join([p.text for p in doc.paragraphs])
        except ImportError:
            return '[需要安装python-docx库来解析Word文档]'
    elif ext == 'zip':
        extracted_texts = []
        with zipfile.ZipFile(filepath, 'r') as zf:
            for name in zf.namelist():
                if name.endswith(('/', '\\')):
                    continue
                try:
                    content = zf.read(name)
                    inner_ext = name.rsplit('.', 1)[1].lower() if '.' in name else ''
                    if inner_ext in ('txt', 'md', 'markdown', 'json'):
                        decoded = _detect_and_decode(content)
                        extracted_texts.append(f'=== {name} ===\n{decoded}')
                    elif inner_ext == 'docx':
                        extracted_texts.append(f'=== {name} ===\n[需要python-docx库来解析Word文档]')
                except Exception:
                    pass
        return '\n\n'.join(extracted_texts)
    return ''

@app.route('/api/upload-analyze', methods=['POST'])
def upload_analyze():
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    file = request.files['file']
    if not file.filename or not allowed_file(file.filename):
        return jsonify({'error': '不支持的文件格式，请上传txt/md/docx/zip/json'}), 400

    filename = secure_filename(file.filename)
    filepath = str(UPLOAD_DIR / f'{uuid.uuid4()}_{filename}')
    file.save(filepath)
    try:
        text = extract_text_from_file(filepath, filename)
        return jsonify({'filename': file.filename, 'content': text, 'length': len(text)})
    finally:
        try:
            os.remove(filepath)
        except Exception:
            pass

@app.route('/api/analyze/export', methods=['POST'])
def export_analysis():
    result = request.json
    if not result:
        return jsonify({'error': 'No data'}), 400
    export_text = json.dumps(result, ensure_ascii=False, indent=2)
    bio = BytesIO()
    bio.write(export_text.encode('utf-8'))
    bio.seek(0)
    return send_file(bio, mimetype='application/json', as_attachment=True, download_name='analysis_result.json')

# 【冷启动提速·2026-08-20】schema+seed 版本号：改动数据库结构（新表/新列/迁移）
# 或种子数据（SEED_SKILL_PACKS / 内置模板）时必须递增此版本，老库才会重新走全量初始化。
SCHEMA_SEED_VERSION = '2026-08-30.1'  # 递增：去AI味儿改稿心法新增"悬空动作/被动吃亏"检测第11条，老库需重同步种子

class AppMeta(db.Model):
    """应用元数据 KV 表：记录 schema/seed 版本，支持启动快速路径。"""
    __tablename__ = 'app_meta'
    key = db.Column(db.String(100), primary_key=True)
    value = db.Column(db.Text)

def init_db():
    with app.app_context():
        # 部署冷启动兜底：先确认 PG 连接可用再跑版本门禁，避免"连接失败"被门禁
        # 误判成"新库"而走全量初始化（详见 db_boot_guard.wait_for_db_ready 注释）
        from db_boot_guard import wait_for_db_ready
        wait_for_db_ready()
        # 【冷启动提速】版本门禁：schema 与种子均无变化时，跳过 create_all + 全量 ALTER
        # 迁移 + 种子同步。外部免费 PostgreSQL（冷唤醒/限流）上全量路径 130+ 次往返实测
        # 约 2.5 分钟，期间端口不监听 → Render 报 "No open ports detected"、前端报
        # "无法连接到服务器"。命中本门禁后常规重启降到秒级。
        try:
            row = db.session.execute(db.text(
                "SELECT value FROM app_meta WHERE key = 'schema_seed_version'"
            )).fetchone()
        except Exception:
            db.session.rollback()  # app_meta 尚不存在（首次部署/新库）→ 走全量初始化
            row = None
        if row and row[0] == SCHEMA_SEED_VERSION:
            print(f'[INIT] ✅ schema/seed 版本一致（{SCHEMA_SEED_VERSION}），跳过迁移与种子同步', flush=True)
            _print_db_diagnosis()
            return
        db.create_all()
        # Migration: 逐条独立提交，避免 PostgreSQL 事务污染
        # （PG 中一条 ALTER 失败会使整个事务 aborted，后续语句全失败）
        # 使用 ADD COLUMN IF NOT EXISTS（PostgreSQL 9.6+ / SQLite 3.35+ 均支持，老版本 SQLite 走 except 兜底）
        def _add_column(table, col_def):
            sql = f'ALTER TABLE {table} ADD COLUMN IF NOT EXISTS {col_def}'
            try:
                db.session.execute(db.text(sql))
                db.session.commit()
            except Exception:
                db.session.rollback()

        _add_column('book_bible', 'locations TEXT')
        _add_column('book_bible', 'concept TEXT')
        _add_column('book_bible', 'plot_design TEXT')
        # Migration: 关系图谱专用字段（解耦 character_profiles，避免角色丢失）
        _add_column('book_bible', 'relation_graph TEXT')
        # Migration: 物资库、人物按卷、动态文件按卷
        _add_column('book_bible', 'inventory TEXT')
        _add_column('book_bible', 'character_volumes TEXT')
        _add_column('book_bible', 'dynamic_volumes TEXT')
        # Migration: 伏笔按卷、地图按卷
        _add_column('book_bible', 'foreshadowing_volumes TEXT')
        _add_column('book_bible', 'locations_volumes TEXT')
        _add_column('ai_config', 'recognition_model TEXT')
        # Migration: AI 多配置支持（DEFAULT TRUE 兼容 PG 严格 BOOLEAN 与 SQLite）
        _add_column('ai_config', "name VARCHAR(50) DEFAULT '默认配置'")
        _add_column('ai_config', 'is_active BOOLEAN DEFAULT TRUE')
        # Migration: skill_packs 添加 github_source 和 github_synced_at 字段
        _add_column('skill_packs', "github_source VARCHAR(500) DEFAULT ''")
        _add_column('skill_packs', 'github_synced_at TIMESTAMP')
        # Migration: 防遗忘检查报告历史
        _add_column('book_bible', 'anti_forget_reports TEXT')
        # Migration P0-2: 伏笔 DAG
        _add_column('book_bible', 'foreshadowing_graph TEXT')
        # Migration P1-4: 四级大纲层级
        _add_column('book_bible', 'outline_hierarchy TEXT')
        # Migration P1-6: 章级变更日志
        _add_column('book_bible', 'chapter_changes_log TEXT')
        # Migration: 检查点快照（借鉴 PlotPilot，每5章自动备份叙事状态）
        _add_column('book_bible', 'state_snapshots TEXT')
        # Migration: 技能包三类划分（master/style/review）+ 题材目标 + 优先级
        _add_column('skill_packs', "category VARCHAR(20) DEFAULT 'master'")
        _add_column('skill_packs', "genre_target VARCHAR(50) DEFAULT ''")
        _add_column('skill_packs', 'priority INTEGER DEFAULT 100')
        # Migration: Book 表技能包三类字段（构思/文风/审查 无污染隔离）
        _add_column('books', "master_skill_ids TEXT DEFAULT '[]'")
        _add_column('books', "style_skill_ids TEXT DEFAULT '[]'")
        _add_column('books', "review_skill_ids TEXT DEFAULT '[]'")
        # Migration: 总卷数 + 风格流派（Book 与 BookBible 双写）
        _add_column('books', 'total_volumes INTEGER DEFAULT 10')
        _add_column('books', "novel_styles TEXT DEFAULT '[]'")
        _add_column('book_bible', 'total_volumes INTEGER DEFAULT 10')
        _add_column('book_bible', "novel_styles TEXT DEFAULT '[]'")
        # Migration P1-5: 章节审计-修订闭环快照（修复保存章节 500 错误：老库缺此列）
        _add_column('chapters', 'review_snapshots TEXT')
        # Migration: 章节摘要（修复 'Chapter' object has no attribute 'summary'：老库缺此列）
        _add_column('chapters', 'summary VARCHAR(500)')
        # Migration M1a: 章节伏笔索引 + 事件抽取索引
        _add_column('chapters', 'hooks_set_json TEXT')
        _add_column('chapters', 'events_extracted_json TEXT')
        # Migration M1a: 全书事件日志
        _add_column('book_bible', 'event_log_json TEXT')
        # Migration M1b: 实体注册表
        _add_column('book_bible', 'entity_registry_json TEXT')
        # Migration M4: 失败记录库
        _add_column('book_bible', 'failure_log_json TEXT')
        # Migration M4b: 用户采纳的 prompt 补丁列表 + 忽略的失败 bucket
        _add_column('book_bible', 'prompt_patches_json TEXT')
        _add_column('book_bible', 'ignored_failure_buckets_json TEXT')
        seed_prompt_templates()
        seed_skill_packs()
        # 版本落库：下次启动命中快速路径，跳过全部迁移与种子同步
        try:
            db.session.merge(AppMeta(key='schema_seed_version', value=SCHEMA_SEED_VERSION))
            db.session.commit()
        except Exception:
            db.session.rollback()
        _print_db_diagnosis()

def _print_db_diagnosis():
    """铁律诊断：每次启动打印数据库状态，确认用户数据持久化。"""
    try:
        user_count = User.query.count()
        book_count = Book.query.count()
        uri = app.config.get('SQLALCHEMY_DATABASE_URI', '')
        is_pg = uri.startswith('postgresql')
        db_label = 'PostgreSQL ✅' if is_pg else 'SQLite ⚠️'
        print(f'[铁律] 用户数据持久化检查：{db_label} | users={user_count} | books={book_count}', flush=True)
        if is_pg:
            print(f'[铁律] ✅ 已连接 PostgreSQL，用户数据将持久化，部署/重启不丢失', flush=True)
        else:
            print(f'[铁律] ❌ 检测到 SQLite！本地开发可用，但生产环境会拒绝启动。请配置 DATABASE_URL', flush=True)
        if user_count == 0:
            print(f'[铁律] ℹ️ 用户表为空（新数据库正常；若之前注册过账号说明数据未持久化）', flush=True)
    except Exception as e:
        print(f'[铁律] 诊断失败: {e}', flush=True)

# ==== 前端静态文件托管（生产环境）====
# 当后端直接提供服务时，托管前端构建产物，避免前后端分离部署导致的 /api 请求失败

@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def serve_frontend(path):
    """托管前端 SPA：所有非 /api 请求都返回静态文件或 index.html"""
    # 前端构建产物目录：优先 frontend/dist，其次 backend/static
    dist_dir = FRONTEND_DIST
    if not dist_dir.exists() or not (dist_dir / 'index.html').exists():
        # fallback: backend/static（仓库中提交的预构建产物）
        static_dir = Path(__file__).parent / 'static'
        if (static_dir / 'index.html').exists():
            dist_dir = static_dir
        else:
            # 再 fallback: 项目根目录
            root_dist = Path(__file__).parent.parent.parent
            if (root_dist / 'index.html').exists():
                dist_dir = root_dist
            else:
                return jsonify({'error': '前端构建产物未找到，请先运行 npm run build'}), 404

    # 如果请求的是具体文件且存在，直接返回
    if path:
        file_path = dist_dir / path
        if file_path.is_file():
            return send_from_directory(dist_dir, path)

    # 其他情况返回 index.html（SPA 路由回退）
    index_file = dist_dir / 'index.html'
    if index_file.exists():
        resp = send_from_directory(dist_dir, 'index.html')
        # 强制浏览器清除 HTTP 缓存和 Service Worker 注册（解决 PWA 死锁缓存问题）
        # 不清除 cookies 和 storage，避免用户登出/丢失数据
        resp.headers['Clear-Site-Data'] = '"cache", "serviceworkers"'
        resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
        return resp
    return jsonify({'error': 'index.html not found'}), 404

if __name__ == '__main__':
    init_db()
    # Render 等云平台通过 PORT 环境变量指定端口，本地默认 5000
    port = int(os.environ.get('PORT', 5000))
    # threaded=True 必须：SSE 流式端点（ai_continue_batch_stream / ai_continue_stream）会长时间占用
    # 工作线程。单线程模式下（默认），SSE 连接期间所有其他请求（health/列表/保存等）会被阻塞。
    app.run(host='0.0.0.0', port=port, debug=False, threaded=True)
