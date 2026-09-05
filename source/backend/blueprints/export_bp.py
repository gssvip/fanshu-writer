"""导入/导出蓝图 — 从 app.py 拆出（app.py 巨石外迁第 2 批：export 域，约 560 行）。

路由清单（9 个）：
  导出：/api/books/<id>/export（txt/docx/epub）、/export-zip、/export-full、/api/analyze/export
  导入：/api/books/import-zip、/api/books/import-files、/api/books/<id>/import-chapters
  封面：/api/books/<id>/cover（上传）、/api/covers/<filename>（静态服务）

搬入的纯 import 域 helper（仅本域使用）：
  _natural_sort_key / _strip_leading_title_line / _extract_zip_chapters /
  split_into_chapters / _merge_empty_chapters

依赖方向（无循环）：
  - 顶层仅依赖 flask + auth_utils + 标准库；
  - 模型与跨域 helper（count_words/update_book_stats/extract_text_from_file/
    resort_chapters_by_title/parse_chapter_number/COVERS_DIR）在路由函数体内
    延迟导入，请求期 app 早已加载完毕（与 general_chat.py 同款模式）。
"""
from __future__ import annotations

import json
import os
import re
import shutil
import tempfile
import uuid
import zipfile
from io import BytesIO

from flask import Blueprint, jsonify, request, send_file

from auth_utils import login_required, login_required_download

export_bp = Blueprint('export', __name__)


# ============================================================================
# 章节拆分 helper（纯函数，无模型依赖；parse_chapter_number 延迟导入）
# ============================================================================

def _natural_sort_key(name):
    """文件名自然排序键，支持中文数字章节号（第1章 < 第2章 < ... < 第10章）。
    解析出章节号的按章节号升序排前；无法解析的按文件名字典序排后。"""
    from app import parse_chapter_number
    base = os.path.splitext(name)[0]
    n = parse_chapter_number(base)
    if n is not None:
        return (0, n, base)
    return (1, 0, base)


def _strip_leading_title_line(text, title):
    """若正文第一行与章节标题（文件名）相同或高度相似，去除该行，避免标题重复。
    比较前会标准化：去除空格、常见标点、统一阿拉伯/中文数字。
    仅当第一行较短（≤40字）且与标题高度匹配时才去除，避免误删正文。"""
    if not text or not title:
        return text
    lines = text.split('\n')
    if not lines:
        return text
    first = lines[0].strip()
    if not first or len(first) > 40:
        return text

    def _norm(s):
        # 去除空格、常见标点
        s = re.sub(r'[\s·，。：:、\-_—()（）\[\]【】]+', '', s or '')
        # 中文数字转阿拉伯（仅单位级别，便于 第1章/第一章 匹配）
        cn_map = {'零': '0', '一': '1', '二': '2', '两': '2', '三': '3', '四': '4',
                  '五': '5', '六': '6', '七': '7', '八': '8', '九': '9', '十': '10'}
        for cn, ar in cn_map.items():
            s = s.replace(cn, ar)
        return s

    nf, nt = _norm(first), _norm(title)
    if not nf or not nt:
        return text
    # 完全相同 或 一方包含另一方 → 视为重复标题行
    if nf == nt or nt in nf or nf in nt:
        return '\n'.join(lines[1:]).strip()
    return text


def _extract_zip_chapters(zippath):
    """解压 zip，若含多个文本文件，返回 [{'title','content'}, ...] 按文件名自然排序，
    每个内嵌文件=一章节，文件名（去扩展名）作为章节标题。
    若 zip 内只有 ≤1 个文本文件，返回 None（让调用方走单文件拆分）。"""
    items = []  # [(inner_name, decoded_text)]
    try:
        with zipfile.ZipFile(zippath, 'r') as zf:
            for name in zf.namelist():
                if name.endswith(('/', '\\')):
                    continue
                inner_ext = name.rsplit('.', 1)[1].lower() if '.' in name else ''
                if inner_ext not in ('txt', 'md', 'markdown', 'json'):
                    continue
                try:
                    content = zf.read(name)
                    decoded = _detect_and_decode(content)
                    if decoded.strip():
                        items.append((name, decoded))
                except Exception:
                    pass
    except Exception:
        return None
    if len(items) <= 1:
        return None
    # 按文件名自然排序
    items.sort(key=lambda x: _natural_sort_key(x[0]))
    result = []
    for name, text in items:
        ch_title = os.path.splitext(os.path.basename(name))[0][:100]
        body = _strip_leading_title_line(text, ch_title)
        result.append({'title': ch_title, 'content': body})
    return result


def _detect_and_decode(raw_bytes):
    """字节流编码探测解码（延迟导入 app 同名 helper，保持单一事实源）。"""
    from app import _detect_and_decode as _dad
    return _dad(raw_bytes)


def split_into_chapters(text):
    """将纯文本按章节标记拆分为多个章节，支持多种章节标题格式。
    会自动合并"同章双标题"（如 第20章 / 第二十章 紧挨出现，前者为空）产生的空章。"""

    def _extract(matches, text, strip_prefix=''):
        """从正则匹配列表中提取章节"""
        chapters = []
        for i, m in enumerate(matches):
            title = m.group().strip()
            if strip_prefix:
                title = title.replace(strip_prefix, '').strip()
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            content = text[start:end].strip()
            chapters.append({'title': title[:100] or f'第{i + 1}章', 'content': content})
        return chapters

    # 1. 中文章节标题：第一章、第1章、第001章、第一回 等
    cn_pattern = re.compile(r'^[ \t]*第[零一二三四五六七八九十百千万\d]+[章节回卷][ \t]*.*$', re.MULTILINE)
    cn_matches = list(cn_pattern.finditer(text))
    if len(cn_matches) >= 1:
        return _merge_empty_chapters(_extract(cn_matches, text))

    # 2. Markdown 标题 # xxx 或 ## xxx
    md_pattern = re.compile(r'^#{1,3}[ \t]+.+$', re.MULTILINE)
    md_matches = list(md_pattern.finditer(text))
    if len(md_matches) >= 1:
        return _merge_empty_chapters(_extract(md_matches, text, strip_prefix='#'))

    # 3. 英文 Chapter N / CHAPTER N
    en_pattern = re.compile(r'^[ \t]*[Cc][Hh][Aa][Pp][Tt][Ee][Rr][ \t]*\d+[ \t]*.*$', re.MULTILINE)
    en_matches = list(en_pattern.finditer(text))
    if len(en_matches) >= 1:
        return _merge_empty_chapters(_extract(en_matches, text))

    # 4. 纯数字开头：1. xxx / 1、xxx / 1: xxx（至少2个才拆分，避免误拆）
    num_pattern = re.compile(r'^[ \t]*\d{1,4}[.、:][ \t]*\S.+$', re.MULTILINE)
    num_matches = list(num_pattern.finditer(text))
    if len(num_matches) >= 2:
        return _merge_empty_chapters(_extract(num_matches, text))

    # 5. 【 xxx 】 或 『 xxx 』 格式标题
    bracket_pattern = re.compile(r'^[ \t]*[【『][^】』]+[】』][ \t]*.*$', re.MULTILINE)
    bracket_matches = list(bracket_pattern.finditer(text))
    if len(bracket_matches) >= 2:
        return _merge_empty_chapters(_extract(bracket_matches, text))

    # 无法拆分，作为单个章节
    return None


def _merge_empty_chapters(chapters):
    """合并空章节：处理"同章双标题"导致的空章。
    规则：
    - 若某章 content 为空，且与下一章章节号相同（如 第20章/第二十章），则视为同一章，
      保留下一章（有内容的那个）的标题，删除空章。
    - 若某章 content 为空，且下一章 content 非空，但章节号不同或无法解析，则将下一章
      内容并入当前空章（标题以当前章为主，下一章标题作为副标题保留），避免0字节空章。
    返回合并后的章节列表（至少保留1章）。"""
    from app import parse_chapter_number
    if not chapters:
        return chapters
    result = []
    i = 0
    while i < len(chapters):
        cur = chapters[i]
        cur_num = parse_chapter_number(cur.get('title', ''))
        cur_content = cur.get('content', '').strip()
        # 当前章为空，尝试与下一章合并
        if not cur_content and i + 1 < len(chapters):
            nxt = chapters[i + 1]
            nxt_num = parse_chapter_number(nxt.get('title', ''))
            nxt_content = nxt.get('content', '').strip()
            # 情况A：章节号相同（同章双标题，如 第20章/第二十章）→ 删空章，跳到下一章
            if cur_num is not None and nxt_num is not None and cur_num == nxt_num:
                i += 1
                continue
            # 情况B：下一章有内容 → 内容并入当前章，标题保留当前章（下一章作副标题）
            if nxt_content:
                merged_title = cur.get('title', '')
                # 若两标题不同，拼接副标题
                if nxt.get('title', '') and nxt.get('title', '') != merged_title:
                    merged_title = f'{merged_title} / {nxt.get("title", "")}'
                result.append({'title': merged_title[:100], 'content': nxt_content})
                i += 2  # 跳过下一章
                continue
            # 情况C：下一章也为空 → 删当前空章，处理下一章
            i += 1
            continue
        result.append({'title': cur.get('title', '')[:100], 'content': cur_content})
        i += 1
    # 兜底：若合并后为空（全部空章），保留原第一章避免0章
    if not result:
        result = [{'title': chapters[0].get('title', '')[:100], 'content': chapters[0].get('content', '')}]
    return result


# ============================================================================
# 导出路由
# ============================================================================

@export_bp.route('/api/books/<book_id>/export', methods=['GET'])
@login_required_download
def export_book(book_id):
    from app import Book, Chapter
    fmt = request.args.get('format', 'txt')
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()

    if fmt == 'txt':
        output = f'《{book.title}》\n作者：{book.author}\n\n简介：{book.synopsis}\n\n{"="*50}\n\n'
        for ch in chapters:
            output += f'\n## {ch.title}\n\n{ch.content}\n\n{"="*50}\n'
        buf = BytesIO(output.encode('utf-8'))
        return send_file(buf, mimetype='text/plain', as_attachment=True, download_name=f'{book.title}.txt')

    elif fmt == 'docx':
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'SimSun'
            font.size = Pt(12)

            title_para = doc.add_heading(book.title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f'作者：{book.author}')
            if book.synopsis:
                doc.add_paragraph(f'简介：{book.synopsis}')
            doc.add_page_break()

            for ch in chapters:
                doc.add_heading(ch.title, level=1)
                for para_text in ch.content.split('\n'):
                    if para_text.strip():
                        doc.add_paragraph(para_text.strip())
                doc.add_page_break()

            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                             as_attachment=True, download_name=f'{book.title}.docx')
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    elif fmt == 'epub':
        try:
            from ebooklib import epub

            epub_book = epub.EpubBook()
            epub_book.set_identifier(book.id)
            epub_book.set_title(book.title)
            epub_book.set_language('zh')
            epub_book.add_author(book.author or '佚名')

            epub_chapters = []
            for ch in chapters:
                c = epub.EpubHtml(title=ch.title, file_name=f'ch_{ch.order_index}.xhtml', lang='zh')
                c.content = f'<h1>{ch.title}</h1>' + ''.join(f'<p>{p}</p>' for p in ch.content.split('\n') if p.strip())
                epub_book.add_item(c)
                epub_chapters.append(c)

            epub_book.toc = epub_chapters
            epub_book.add_item(epub.EpubNcx())
            epub_book.add_item(epub.EpubNav())

            style = 'BODY {color: black;}'
            nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
            epub_book.add_item(nav_css)

            epub_book.spine = ['nav'] + epub_chapters

            buf = BytesIO()
            epub.write_epub(buf, epub_book, {})
            buf.seek(0)
            return send_file(buf, mimetype='application/epub+zip', as_attachment=True, download_name=f'{book.title}.epub')
        except Exception as e:
            return jsonify({'error': str(e)}), 500

    return jsonify({'error': 'Unsupported format'}), 400


@export_bp.route('/api/books/<book_id>/export-zip', methods=['GET'])
@login_required_download
def export_book_zip(book_id):
    from app import db, Book, Chapter, Character, Outline, COVERS_DIR
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    tmpdir = tempfile.mkdtemp()
    try:
        metadata = book.to_dict()
        metadata['chapters'] = [c.to_dict(include_content=True) for c in Chapter.query.filter_by(book_id=book_id).order_by(Chapter.order_index).all()]
        metadata['characters'] = [c.to_dict() for c in Character.query.filter_by(book_id=book_id).all()]
        metadata['outlines'] = [o.to_dict() for o in Outline.query.filter_by(book_id=book_id).order_by(Outline.order_index).all()]

        with open(os.path.join(tmpdir, 'book.json'), 'w', encoding='utf-8') as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)

        zippath = os.path.join(tmpdir, f'{book.title}.zip')
        with zipfile.ZipFile(zippath, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(os.path.join(tmpdir, 'book.json'), 'book.json')
            if book.cover_path:
                cover_file = COVERS_DIR / os.path.basename(book.cover_path)
                if cover_file.exists():
                    zf.write(cover_file, f'cover{cover_file.suffix}')

        return send_file(zippath, mimetype='application/zip', as_attachment=True, download_name=f'{book.title}.zip')
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


@export_bp.route('/api/books/<book_id>/export-full', methods=['GET'])
@login_required_download
def export_book_full(book_id):
    """导出小说的全部维度内容（除图谱外）和所有章节为独立文件，打包成zip下载"""
    from app import Book, Chapter, BookBible
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404

    # 以小说标题命名文件夹，简单清理会影响路径的字符
    book_title = (book.title or '未命名').replace('/', '_').replace('\\', '_').replace(':', '_')

    bible = BookBible.query.filter_by(book_id=book_id).first()
    bible_data = bible.to_dict() if bible else {}

    # 维度配置：(文件名, BookBible字段, 中文标题)
    dimensions = [
        ('构思.md', 'concept', '构思'),
        ('设定.md', 'key_rules', '设定'),
        ('大纲.md', 'plot_design', '大纲'),
        ('世界观.md', 'worldbuilding', '世界观'),
        ('人物.md', 'character_profiles', '人物'),
        ('剧情时间线.md', 'timeline', '剧情时间线'),
        ('伏笔.md', 'foreshadowing', '伏笔'),
        ('地点.md', 'locations', '地点'),
        ('风格指南.md', 'style_guide', '风格指南'),
    ]

    # 章节按 order_index 排序，跳过卷标记
    chapters = Chapter.query.filter_by(book_id=book_id, is_volume=False).order_by(Chapter.order_index).all()

    buf = BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        # 写入各维度 md 文件
        for filename, field, title in dimensions:
            content = (bible_data.get(field, '') or '').strip()
            body = content if content else '暂无内容'
            md_text = f'# {title}\n\n---\n\n{body}\n'
            zf.writestr(f'{book_title}/{filename}', md_text.encode('utf-8'))

        # 写入章节子文件夹，每章一个 txt
        for idx, ch in enumerate(chapters, start=1):
            ch_title = ch.title or f'第{idx}章'
            ch_filename = f'第{idx}章-{ch_title}.txt'
            ch_content = ch.content if (ch.content and ch.content.strip()) else '暂无内容'
            zf.writestr(f'{book_title}/章节/{ch_filename}', ch_content.encode('utf-8'))

    buf.seek(0)
    return send_file(buf, mimetype='application/zip', as_attachment=True, download_name=f'{book_title}.zip')


@export_bp.route('/api/analyze/export', methods=['POST'])
def export_analysis():
    result = request.json
    if not result:
        return jsonify({'error': 'No data'}), 400
    export_text = json.dumps(result, ensure_ascii=False, indent=2)
    bio = BytesIO()
    bio.write(export_text.encode('utf-8'))
    bio.seek(0)
    return send_file(bio, mimetype='application/json', as_attachment=True, download_name='analysis_result.json')


# ============================================================================
# 封面（导出 zip 会打包封面资源，归入 export 域统一管理）
# ============================================================================

@export_bp.route('/api/books/<book_id>/cover', methods=['POST'])
def upload_cover(book_id):
    from app import db, Book, COVERS_DIR
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': 'Book not found'}), 404
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No filename'}), 400
    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in ('png', 'jpg', 'jpeg', 'gif', 'webp'):
        return jsonify({'error': 'Invalid image format'}), 400

    filename = f'{book_id}.{ext}'
    filepath = COVERS_DIR / filename
    file.save(filepath)
    book.cover_path = f'/api/covers/{filename}'
    db.session.commit()
    return jsonify({'cover_path': book.cover_path})


@export_bp.route('/api/covers/<filename>')
def serve_cover(filename):
    from app import COVERS_DIR
    # 文件不存在时返回 404 而非让 send_file 抛 FileNotFoundError 变 500
    # （历史遗留：封面/头像路径失效时前端拿到的是 500，无法按 404 降级处理）
    path = COVERS_DIR / filename
    if not path.is_file():
        return jsonify({'error': 'Cover not found'}), 404
    return send_file(path)


# ============================================================================
# 导入路由
# ============================================================================

@export_bp.route('/api/books/import-zip', methods=['POST'])
def import_book_zip():
    from app import db, Book, Chapter, Character, Outline, count_words, update_book_stats
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    file = request.files['file']
    tmpdir = tempfile.mkdtemp()
    try:
        zippath = os.path.join(tmpdir, 'import.zip')
        file.save(zippath)
        with zipfile.ZipFile(zippath, 'r') as zf:
            zf.extractall(tmpdir)

        book_json = os.path.join(tmpdir, 'book.json')
        if not os.path.exists(book_json):
            return jsonify({'error': 'Invalid book package: no book.json'}), 400

        with open(book_json, 'r', encoding='utf-8') as f:
            data = json.load(f)

        book = Book(
            title=data.get('title', '导入书籍'), author=data.get('author', ''),
            genre=data.get('genre', 'other'), book_type=data.get('book_type', 'novel'),
            synopsis=data.get('synopsis', ''), status='draft'
        )
        db.session.add(book)
        db.session.flush()

        for ch_data in data.get('chapters', []):
            ch = Chapter(
                book_id=book.id, title=ch_data.get('title', ''),
                content=ch_data.get('content', ''), order_index=ch_data.get('order_index', 0),
                is_volume=ch_data.get('is_volume', False), parent_id=ch_data.get('parent_id', ''),
                word_count=count_words(ch_data.get('content', ''))
            )
            db.session.add(ch)

        for char_data in data.get('characters', []):
            char = Character(
                book_id=book.id, name=(char_data.get('name', '') or '')[:50],
                role=(char_data.get('role', 'supporting') or 'supporting')[:50], description=char_data.get('description', ''),
                appearance=char_data.get('appearance', ''), personality=char_data.get('personality', ''),
                background=char_data.get('background', ''),
                relationships_json=json.dumps(char_data.get('relationships', []), ensure_ascii=False)
            )
            db.session.add(char)

        for o_data in data.get('outlines', []):
            outline = Outline(
                book_id=book.id, title=o_data.get('title', ''),
                content=o_data.get('content', ''), order_index=o_data.get('order_index', 0),
                level=o_data.get('level', 0), parent_id=o_data.get('parent_id', '')
            )
            db.session.add(outline)

        update_book_stats(book.id)
        return jsonify(book.to_dict()), 201
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


@export_bp.route('/api/books/import-files', methods=['POST'])
@login_required
def import_book_files():
    """从多个文本文件导入创建新作品，支持 txt/md/docx/zip
    导入模式自动识别：
    - 多文件模式（≥2个文件）：每个文件 = 一个章节，文件名作为章节标题，按文件名排序
    - 单文件模式（1个文件）：调用 split_into_chapters 拆分为多个章节
    - zip 模式：解压后若含多个文本文件，按多文件模式处理；否则按单文件拆分
    """
    from app import db, Book, Chapter, count_words, update_book_stats, \
        extract_text_from_file, resort_chapters_by_title
    files = request.files.getlist('files')
    if not files or len(files) == 0:
        return jsonify({'error': '未选择文件'}), 400

    title = request.form.get('title', '').strip()
    book_type = request.form.get('book_type', 'novel')
    genre = request.form.get('genre', 'other')

    tmpdir = tempfile.mkdtemp()
    try:
        all_chapters = []
        # 收集有效文件（保存到临时目录，记录原始文件名）
        collected = []  # [(original_name, filepath, ext)]
        for f in files:
            if not f or not f.filename:
                continue
            original_name = f.filename
            ext = original_name.rsplit('.', 1)[1].lower() if '.' in original_name else ''
            if ext not in ('txt', 'md', 'markdown', 'docx', 'zip', 'json'):
                continue
            safe_name = f'{uuid.uuid4()}.{ext}'
            filepath = os.path.join(tmpdir, safe_name)
            f.save(filepath)
            collected.append((original_name, filepath, ext))

        if not collected:
            return jsonify({'error': '未选择有效文件'}), 400

        # 判断导入模式
        # 情况1：单个 zip 文件 → 解压，若含多个文本文件则按多文件模式（每内嵌文件=一章节）
        if len(collected) == 1 and collected[0][2] == 'zip':
            zip_chapters = _extract_zip_chapters(collected[0][1])
            if zip_chapters:
                all_chapters.extend(zip_chapters)
            else:
                # zip 内无法按文件拆分，回退到单文件拆分
                text = extract_text_from_file(collected[0][1], collected[0][0])
                if text.strip():
                    chapters = split_into_chapters(text)
                    if chapters:
                        all_chapters.extend(chapters)
                    else:
                        all_chapters.append({'title': os.path.splitext(collected[0][0])[0][:100], 'content': text})
        # 情况2：多个文件 → 多文件模式（每个文件=一章节，文件名作为标题，按文件名自然排序）
        elif len(collected) >= 2:
            collected.sort(key=lambda x: _natural_sort_key(x[0]))
            for original_name, filepath, ext in collected:
                text = extract_text_from_file(filepath, original_name)
                if not text.strip():
                    continue
                ch_title = os.path.splitext(original_name)[0][:100]
                # 去除正文开头的重复标题行（若文件第一行与文件名相同，避免标题重复）
                body = _strip_leading_title_line(text, ch_title)
                all_chapters.append({'title': ch_title, 'content': body})
        # 情况3：单个文件 → 单文件拆分模式
        else:
            original_name, filepath, ext = collected[0]
            text = extract_text_from_file(filepath, original_name)
            if text.strip():
                chapters = split_into_chapters(text)
                if chapters:
                    all_chapters.extend(chapters)
                else:
                    all_chapters.append({'title': os.path.splitext(original_name)[0][:100], 'content': text})

        if not all_chapters:
            return jsonify({'error': '未能从文件中提取到有效内容，请检查文件格式或编码'}), 400

        # 从第一个文件名推断标题
        if not title:
            first_file = (collected[0][0] if collected else '导入作品') or '导入作品'
            title = os.path.splitext(first_file)[0][:50]

        book = Book(
            user_id=getattr(request, 'current_user_id', ''),
            title=title, author='', genre=genre, book_type=book_type,
            synopsis=f'从文件导入，共 {len(all_chapters)} 章', status='draft'
        )
        db.session.add(book)
        db.session.flush()

        for idx, ch_data in enumerate(all_chapters):
            ch = Chapter(
                book_id=book.id, title=ch_data['title'][:200] or f'第{idx + 1}章',
                content=ch_data['content'], order_index=idx,
                is_volume=False, parent_id='',
                word_count=count_words(ch_data['content'])
            )
            db.session.add(ch)
        db.session.flush()

        # 按章节标题中的章节号（第N章/第N章等）自动排序
        resort_chapters_by_title(book.id, rebin_volumes=False)

        # 自动按 50 章一组创建卷
        chapter_count = len(all_chapters)
        volume_interval = 50
        vol_count = (chapter_count + volume_interval - 1) // volume_interval
        for vidx in range(vol_count):
            ch_start = vidx * volume_interval + 1
            ch_end = min((vidx + 1) * volume_interval, chapter_count)
            vol = Chapter(
                book_id=book.id,
                title=f'第{vidx + 1}卷（第{ch_start}-{ch_end}章）',
                content='', order_index=chapter_count + vidx,
                is_volume=True, parent_id='', word_count=0,
            )
            db.session.add(vol)
            db.session.flush()
            # 将范围内章节归入该卷
            for cidx in range(vidx * volume_interval, min((vidx + 1) * volume_interval, chapter_count)):
                # 找到刚刚创建的章节（按 order_index 为 cidx）
                ch_obj = Chapter.query.filter_by(book_id=book.id, order_index=cidx, is_volume=False).first()
                if ch_obj:
                    ch_obj.parent_id = vol.id

        update_book_stats(book.id)
        return jsonify(book.to_dict()), 201
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


@export_bp.route('/api/books/<book_id>/import-chapters', methods=['POST'])
@login_required
def append_import_chapters(book_id):
    """追加导入章节到已有作品，支持 txt/md/docx/zip，每个文件可含多章。
    新章节会按当前最大 order_index 顺序追加，不影响已有章节。
    """
    from app import db, Book, Chapter, count_words, update_book_stats, \
        extract_text_from_file, resort_chapters_by_title
    book = Book.query.get(book_id)
    if not book:
        return jsonify({'error': '作品不存在'}), 404
    if book.user_id != request.current_user_id:
        return jsonify({'error': '无权操作该作品'}), 403

    files = request.files.getlist('files')
    if not files or len(files) == 0:
        return jsonify({'error': '未选择文件'}), 400

    tmpdir = tempfile.mkdtemp()
    try:
        new_chapters = []
        for f in files:
            if not f or not f.filename:
                continue
            original_name = f.filename
            ext = original_name.rsplit('.', 1)[1].lower() if '.' in original_name else ''
            if ext not in ('txt', 'md', 'markdown', 'docx', 'zip', 'json'):
                continue
            safe_name = f'{uuid.uuid4()}.{ext}'
            filepath = os.path.join(tmpdir, safe_name)
            f.save(filepath)
            text = extract_text_from_file(filepath, safe_name)
            if not text.strip():
                continue
            chapters = split_into_chapters(text)
            if chapters:
                new_chapters.extend(chapters)
            else:
                ch_title = os.path.splitext(original_name)[0]
                new_chapters.append({'title': ch_title[:100], 'content': text})

        if not new_chapters:
            return jsonify({'error': '未能从文件中提取到有效内容，请检查文件格式或编码'}), 400

        # 计算当前最大 order_index，新章节追加在末尾
        max_order = db.session.query(db.func.max(Chapter.order_index)).filter_by(book_id=book_id).scalar() or 0
        added = 0
        for idx, ch_data in enumerate(new_chapters):
            ch = Chapter(
                book_id=book_id,
                title=ch_data['title'][:200] or f'第{max_order + idx + 1}章',
                content=ch_data['content'],
                order_index=max_order + idx + 1,
                is_volume=False,
                parent_id='',
                word_count=count_words(ch_data['content'])
            )
            db.session.add(ch)
            added += 1
        db.session.flush()

        # 按章节标题中的章节号自动重排（含已有章节），并按 50 章/卷重新归入卷
        resort_chapters_by_title(book_id, rebin_volumes=True)

        update_book_stats(book_id)
        total = Chapter.query.filter_by(book_id=book_id, is_volume=False).count()
        return jsonify({'success': True, 'added': added, 'total': total}), 200
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)
